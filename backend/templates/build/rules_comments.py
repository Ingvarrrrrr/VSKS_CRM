# -*- coding: utf-8 -*-
"""Развешивание Word-комментариев с русскими пояснениями поверх тегов
docxtpl ({{ var }} / {% cond %}) в уже собранном документе.

Вызывается ПОСЛЕДНИМ шагом сборки (после docxedit.save()), когда шаблон
уже прошёл все R*/T* правила и записан на диск как обычный .docx. Дальше
файл открывается python-docx, на runs с тегами вешаются комментарии
(текст — из backend/templates/build/comments_ru.py), и файл пересохраняется.

Идемпотентность: python-docx пишет zip через `ZipFile.writestr()` без
явного ZipInfo → используется текущее время (`time.localtime()`), из-за
чего sha256 «плавает» между сборками. Чтобы --check оставался стабильным,
здесь дата комментария зафиксирована (_FIXED_COMMENT_DATE), а сам zip
пересобирается вручную с фиксированным mtime — по той же схеме, что
docxedit.save().
"""

from __future__ import annotations

import datetime as _dt
import io
import pathlib
import re
import zipfile

from docx import Document

from backend.templates.build.comments_ru import COND_COMMENTS, VAR_COMMENTS

_AUTHOR = "Подсказка"
_INITIALS = "П"

# Зафиксировано, чтобы sha256 не «плавал» между сборками (см. docstring).
_FIXED_COMMENT_DATE = _dt.datetime(2026, 1, 1, tzinfo=_dt.timezone.utc)

_FIXED_MTIME = (2024, 1, 1, 0, 0, 0)

_TAG_RE = re.compile(r"\{\{.*?\}\}|\{%-?.*?-?%\}")
_VAR_NAME_RE = re.compile(r"\{\{\s*([A-Za-z_][A-Za-z0-9_]*)")
_COND_INNER_RE = re.compile(r"\{%-?\s*(.*?)\s*-?%\}")


def _normalize_cond(raw_inner: str) -> str:
    """Убирает {%p ...%}-префикс абзацного тега и схлопывает пробелы."""
    s = re.sub(r"\s+", " ", raw_inner.strip())
    if s == "p":
        return ""
    if s.startswith("p "):
        s = s[2:].strip()
    return s


def _runs_for_span(runs: list, start: int, end: int) -> list | None:
    """Возвращает подсписок runs, покрывающих символьный диапазон [start, end)."""
    idx = 0
    first_i = last_i = None
    for i, r in enumerate(runs):
        r_text = r.text or ""
        r_start, r_end = idx, idx + len(r_text)
        if first_i is None and r_end > start:
            first_i = i
        if r_start < end:
            last_i = i
        idx = r_end
        if r_start >= end:
            break
    if first_i is None or last_i is None:
        return None
    return runs[first_i:last_i + 1]


def _iter_all_paragraphs(doc: Document):
    """Абзацы документа + все ячейки таблиц (рекурсивно, включая вложенные
    таблицы), в порядке документа настолько, насколько это даёт python-docx."""
    for p in doc.paragraphs:
        yield p

    def _walk_table(tbl):
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    yield from _walk_table(nested)

    for tbl in doc.tables:
        yield from _walk_table(tbl)


def apply_comments(doc: Document) -> tuple[int, list[str]]:
    """Проходит все абзацы/ячейки документа, вешает комментарии на теги.

    Возвращает (число_повешенных_комментариев, список_непокрытых_тегов).
    Документ модифицируется на месте (комментарии добавляются через
    doc.add_comment). Сохранение — забота вызывающего кода.
    """
    n_comments = 0
    uncovered: list[str] = []

    for p in _iter_all_paragraphs(doc):
        runs = p.runs
        if not runs:
            continue
        text = "".join(r.text or "" for r in runs)
        if "{{" not in text and "{%" not in text:
            continue

        for m in _TAG_RE.finditer(text):
            raw = m.group(0)
            start, end = m.span()

            if raw.startswith("{{"):
                vm = _VAR_NAME_RE.match(raw)
                name = vm.group(1) if vm else None
                comment_text = VAR_COMMENTS.get(name) if name else None
                if comment_text is None:
                    uncovered.append(raw)
                    continue
            else:
                cm = _COND_INNER_RE.match(raw)
                inner = cm.group(1) if cm else raw
                key = _normalize_cond(inner)
                comment_text = COND_COMMENTS.get(key)
                if comment_text is None:
                    uncovered.append(raw)
                    continue

            span_runs = _runs_for_span(runs, start, end)
            if not span_runs:
                # Тег есть в тексте, но не удалось сопоставить runs —
                # не должно происходить при непустом paragraph.runs, но на
                # всякий случай не падаем, а копим как непокрытый.
                uncovered.append(raw)
                continue

            comment = doc.add_comment(
                span_runs, text=comment_text, author=_AUTHOR, initials=_INITIALS
            )
            comment._comment_elm.date = _FIXED_COMMENT_DATE
            n_comments += 1

    return n_comments, uncovered


def _resave_deterministic(doc: Document, out_path: str | pathlib.Path) -> None:
    """Сохраняет Document в out_path с фиксированным mtime у всех записей
    zip (python-docx пишет через writestr() без ZipInfo → «плавающее»
    время → нестабильный sha256 в build.py --check)."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    with zipfile.ZipFile(buf, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}

    with zipfile.ZipFile(str(out_path), "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name in sorted(entries.keys()):
            zi = zipfile.ZipInfo(name, date_time=_FIXED_MTIME)
            zi.compress_type = zipfile.ZIP_DEFLATED
            zout.writestr(zi, entries[name])


def apply_comments_to_file(path: str | pathlib.Path) -> tuple[int, list[str]]:
    """Открывает готовый .docx по пути `path`, вешает комментарии, сохраняет
    обратно в тот же файл (детерминированно). Возвращает
    (число_комментариев, список_непокрытых_тегов)."""
    doc = Document(str(path))
    n_comments, uncovered = apply_comments(doc)
    _resave_deterministic(doc, path)
    return n_comments, uncovered
