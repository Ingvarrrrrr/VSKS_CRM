"""
Генератор цветных шаблонов Служебных Записок (СЗ) — Phase 23.2.

Запуск:
    py backend/templates/make_service_notes.py

Создаёт 5 файлов в backend/templates/:
    service_note.docx
    service_note_delivery.docx
    service_note_payment.docx
    service_note_procurement.docx
    service_note_advance.docx

Цветовая схема (как в contract.docx):
    {{синие}}   — placeholder'ы (RGBColor #1E40AF, жирный)
    зелёный     — ветка «вариант 1» условного блока
    оранжевый   — ветка «вариант 2» условного блока
    серый курсив — маркеры {% if %} / {% else %} / {% endif %}
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── ЦВЕТА ───────────────────────────────────────────────────────────────────
COLOR_PLACEHOLDER = RGBColor(0x1E, 0x40, 0xAF)   # синий — {{var}}
COLOR_BRANCH_TRUE = RGBColor(0x15, 0x80, 0x3D)   # зелёный — ветка 1
COLOR_BRANCH_FALSE = RGBColor(0xC2, 0x41, 0x0C)  # оранжевый — ветка 2
COLOR_MARKER = RGBColor(0x6B, 0x72, 0x80)         # серый — {% %}

FONT_NAME = "Times New Roman"


# ─── HELPER'Ы ─────────────────────────────────────────────────────────────────

def _base_run(paragraph, text, size=12):
    r = paragraph.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    return r


def add_placeholder(para, text, size=12):
    """{{var}} — синий жирный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_PLACEHOLDER
    r.bold = True
    return r


def add_marker(para, text, size=12):
    """{% if %} / {% else %} / {% endif %} — серый курсив."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_MARKER
    r.italic = True
    return r


def add_branch_true(para, text, size=12):
    """Ветка 1 — зелёный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_BRANCH_TRUE
    return r


def add_branch_false(para, text, size=12):
    """Ветка 2 — оранжевый."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_BRANCH_FALSE
    return r


def add_text(para, text, bold=False, size=12, italic=False):
    r = _base_run(para, text, size)
    r.bold = bold
    r.italic = italic
    return r


# ─── УТИЛИТЫ ДОКУМЕНТА ───────────────────────────────────────────────────────

def page_margins(doc):
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)


def new_para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
             first_line_indent=None, line_spacing=13.8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    return p


def para_simple(doc, text="", bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0, space_after=6, italic=False):
    p = new_para(doc, align=align, space_before=space_before, space_after=space_after)
    if text:
        add_text(p, text, bold=bold, size=size, italic=italic)
    return p


# ─── ЛЕГЕНДА ─────────────────────────────────────────────────────────────────

def add_legend(doc):
    p_leg_title = doc.add_paragraph()
    p_leg_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_leg_title.paragraph_format.space_after = Pt(4)
    p_leg_title.paragraph_format.space_before = Pt(0)
    r_lt = p_leg_title.add_run("🎨 ЛЕГЕНДА — как читать этот шаблон")
    r_lt.bold = True
    r_lt.font.name = FONT_NAME
    r_lt.font.size = Pt(11)
    r_lt.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # phase26-v: legend-примеры в угловых кавычках (без буквальных {{}}/{%%}),
    # потому что docxtpl интерпретирует literal {% %} в тексте шаблона
    # как настоящие jinja-теги и валит «парные теги» если не найдёт пару.
    # Не пиши тут «{{» «}}» «{%» «%}» — даже внутри пояснений!
    LEGEND_LINES = [
        ("placeholder", "«синие переменные»    — автоподстановка из БД"),
        ("true",        "зелёный текст         — вариант 1 в условном блоке"),
        ("false",       "оранжевый текст       — вариант 2"),
        ("marker",      "серые «маркеры»       — технические условия Jinja2 (не трогать)"),
    ]

    for kind, text in LEGEND_LINES:
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pl.paragraph_format.space_before = Pt(0)
        pl.paragraph_format.space_after = Pt(2)
        pl.paragraph_format.first_line_indent = Cm(0.5)
        r = pl.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        if kind == "placeholder":
            r.font.color.rgb = COLOR_PLACEHOLDER
            r.bold = True
        elif kind == "true":
            r.font.color.rgb = COLOR_BRANCH_TRUE
        elif kind == "false":
            r.font.color.rgb = COLOR_BRANCH_FALSE
        elif kind == "marker":
            r.font.color.rgb = COLOR_MARKER
            r.italic = True

    p_warn = doc.add_paragraph()
    p_warn.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_warn.paragraph_format.space_before = Pt(4)
    p_warn.paragraph_format.space_after = Pt(2)
    p_warn.paragraph_format.first_line_indent = Cm(0.5)
    r_warn = p_warn.add_run(
        "⚠️ После открытия в Word удалите эту секцию (от «ЛЕГЕНДА» до строки «—————»).\n"
        "Полный справочник переменных: SubsidiesView → Шаблоны → «Руководство по переменным»"
    )
    r_warn.font.name = FONT_NAME
    r_warn.font.size = Pt(10)
    r_warn.italic = True

    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sep.paragraph_format.space_before = Pt(2)
    p_sep.paragraph_format.space_after = Pt(10)
    r_sep = p_sep.add_run("—" * 40)
    r_sep.font.name = FONT_NAME
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = COLOR_MARKER


# ─── ШАПКА СЗ ────────────────────────────────────────────────────────────────

def add_sz_header(doc, kind):
    """Шапка: кому/от кого + заголовок «СЛУЖЕБНАЯ ЗАПИСКА» с условным подзаголовком."""

    # Правый блок «кому»
    p_to = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_text(p_to, "Президенту ВСКС")

    p_to2 = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    add_placeholder(p_to2, "{{customer_signatory_name_genitive}}")

    p_from = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_text(p_from, "от ")
    add_placeholder(p_from, "{{initiator_role}}")

    p_from2 = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)
    add_placeholder(p_from2, "{{initiator_name}}")

    # Заголовок
    p_title = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_text(p_title, "СЛУЖЕБНАЯ ЗАПИСКА", bold=True, size=14)

    # Условный подзаголовок
    if kind == 'general':
        # Без подзаголовка
        pass
    elif kind == 'delivery':
        p_sub_if = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        add_marker(p_sub_if, "{% if doc_kind == 'delivery' %}", size=10)
        p_sub = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        add_branch_true(p_sub, "о выдаче товаров")
        p_sub_end = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        add_marker(p_sub_end, "{% endif %}", size=10)
    elif kind == 'payment':
        p_sub_if = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        add_marker(p_sub_if, "{% if doc_kind == 'payment' %}", size=10)
        p_sub = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        add_branch_true(p_sub, "об оплате")
        p_sub_end = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        add_marker(p_sub_end, "{% endif %}", size=10)
    elif kind == 'procurement':
        p_sub_if = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        add_marker(p_sub_if, "{% if doc_kind == 'procurement' %}", size=10)
        p_sub = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        add_branch_true(p_sub, "о закупке")
        p_sub_end = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        add_marker(p_sub_end, "{% endif %}", size=10)
    elif kind == 'advance':
        p_sub_if = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        add_marker(p_sub_if, "{% if doc_kind == 'advance' %}", size=10)
        p_sub = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        add_branch_true(p_sub, "об авансовом отчёте")
        p_sub_end = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        add_marker(p_sub_end, "{% endif %}", size=10)

    # Город и дата
    p_city = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
    add_text(p_city, "г. Москва                          «")
    add_placeholder(p_city, "{{today}}")
    add_text(p_city, "»")


# ─── ПОДПИСЬ СЗ ──────────────────────────────────────────────────────────────

def add_sz_signature(doc):
    doc.add_paragraph()
    p_role = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_placeholder(p_role, "{{initiator_role}}")
    p_sign = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_text(p_sign, "___________________ / ")
    add_placeholder(p_sign, "{{initiator_name}}")
    p_date = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
    add_text(p_date, "«")
    add_placeholder(p_date, "{{today}}")
    add_text(p_date, "»")


# ─── ТЕЛА ШАБЛОНОВ ───────────────────────────────────────────────────────────

def add_body_general(doc):
    """Общая СЗ."""
    p = new_para(doc, space_after=6)
    add_text(p, "Прошу Вас рассмотреть возможность ")
    add_placeholder(p, "{{service_subject}}")
    add_text(p, ". Обоснование: [текст пользователя].")

    p2 = new_para(doc, space_after=4)
    add_text(p2, "Объект закупки: ")
    add_placeholder(p2, "{{subject}}")

    p3 = new_para(doc, space_after=4)
    add_text(p3, "Сумма: ")
    add_placeholder(p3, "{{contract_price_words}}")
    add_text(p3, " (")
    add_placeholder(p3, "{{contract_price_num}}")
    add_text(p3, " руб.)")

    p4 = new_para(doc, space_after=4)
    add_text(p4, "Категория ФЭО: ")
    add_placeholder(p4, "{{feo_path}}")

    p5 = new_para(doc, space_after=6)
    add_text(p5, "Субсидия: ")
    add_placeholder(p5, "{{subsidy_name}}")


def add_body_delivery(doc):
    """СЗ на выдачу товара."""
    p = new_para(doc, space_after=6)
    add_text(p, "В соответствии с ")
    add_placeholder(p, "{{contract_number}}")
    add_text(p, " от ")
    add_placeholder(p, "{{contract_date}}")
    add_text(p, " прошу выдать со склада товарно-материальные ценности по ")
    add_placeholder(p, "{{contractor_name}}")
    add_text(p, " согласно прилагаемой спецификации (Приложение №1):")

    # Таблица позиций — цикл
    p_for = new_para(doc, space_after=0)
    add_marker(p_for, "{% tr for item in items %}", size=10)

    tbl = doc.add_table(rows=2, cols=6)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    headers = ["№", "Наименование", "Ед. изм.", "Кол-во", "Цена за ед., руб.", "Сумма, руб."]
    for i, h in enumerate(headers):
        p_h = hdr[i].paragraphs[0]
        r = p_h.add_run(h)
        r.bold = True
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    row_cells = tbl.rows[1].cells
    placeholders = [
        "{{item.num}}", "{{item.name}}", "{{item.unit}}",
        "{{item.quantity}}", "{{item.unit_price}}", "{{item.total_price}}"
    ]
    for i, ph in enumerate(placeholders):
        p_c = row_cells[i].paragraphs[0]
        r = p_c.add_run(ph)
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_PLACEHOLDER
        r.bold = True
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_endfor = new_para(doc, space_before=2, space_after=6)
    add_marker(p_endfor, "{% tr endfor %}", size=10)

    p_total = new_para(doc, space_after=4)
    add_text(p_total, "Итого: ")
    add_placeholder(p_total, "{{contract_price_num}}")
    add_text(p_total, " руб. (")
    add_placeholder(p_total, "{{contract_price_words}}")
    add_text(p_total, ")")

    p_recv = new_para(doc, space_after=6)
    add_text(p_recv, "Получатель: ")
    add_placeholder(p_recv, "{{responsible_person}}")


def add_body_payment(doc):
    """СЗ на оплату."""
    p = new_para(doc, space_after=6)
    add_text(p, "Прошу произвести оплату по договору ")
    add_placeholder(p, "{{contract_number}}")
    add_text(p, " от ")
    add_placeholder(p, "{{contract_date}}")
    add_text(p, " с ")
    add_placeholder(p, "{{contractor_name}}")
    add_text(p, " (ИНН ")
    add_placeholder(p, "{{contractor_inn}}")
    add_text(p, ") на сумму ")
    add_placeholder(p, "{{contract_price_num}}")
    add_text(p, " руб. (")
    add_placeholder(p, "{{contract_price_words}}")
    add_text(p, "), ")

    # НДС — условный блок
    p_vat = new_para(doc, space_after=6)
    add_marker(p_vat, "{% if vat_applicable %}")
    add_branch_true(p_vat, "в т.ч. НДС ")
    add_placeholder(p_vat, "{{vat_rate}}")
    add_branch_true(p_vat, "%")
    add_marker(p_vat, "{% else %}")
    add_branch_false(p_vat, "НДС не облагается на основании ")
    add_placeholder(p_vat, "{{vat_exemption_article}}")
    add_marker(p_vat, "{% endif %}")
    add_text(p_vat, ".")

    p_feo = new_para(doc, space_after=4)
    add_text(p_feo, "Категория ФЭО: ")
    add_placeholder(p_feo, "{{feo_path}}")

    p_sub = new_para(doc, space_after=4)
    add_text(p_sub, "Субсидия: ")
    add_placeholder(p_sub, "{{subsidy_name}}")

    p_req_hdr = new_para(doc, space_before=4, space_after=2)
    add_text(p_req_hdr, "Реквизиты получателя:", bold=True)

    p_rs = new_para(doc, space_after=2)
    add_text(p_rs, "р/с: ")
    add_placeholder(p_rs, "{{contractor_settlement_account}}")

    p_bik = new_para(doc, space_after=2)
    add_text(p_bik, "БИК: ")
    add_placeholder(p_bik, "{{contractor_bik}}")

    p_bank = new_para(doc, space_after=6)
    add_text(p_bank, "Банк: ")
    add_placeholder(p_bank, "{{contractor_bank_name}}")

    p_basis = new_para(doc, space_after=6)
    add_text(p_basis, "Основание: акт сдачи-приёмки от ")
    add_placeholder(p_basis, "{{acceptance_doc_date}}")
    add_text(p_basis, " № ")
    add_placeholder(p_basis, "{{acceptance_doc_number}}")
    add_text(p_basis, ".")


def add_body_procurement(doc):
    """СЗ на закупку."""
    p_hdr = new_para(doc, space_after=4)
    add_text(p_hdr, "Прошу инициировать закупку:", bold=True)

    p_subj = new_para(doc, space_after=4)
    add_text(p_subj, "Предмет: ")
    add_placeholder(p_subj, "{{subject}}")

    # Тип — условный блок
    p_type = new_para(doc, space_after=4)
    add_text(p_type, "Тип: ")
    add_marker(p_type, "{% if subject_kind == 'goods' %}")
    add_branch_false(p_type, "поставка товара")
    add_marker(p_type, "{% else %}")
    add_branch_true(p_type, "оказание услуг")
    add_marker(p_type, "{% endif %}")

    p_method = new_para(doc, space_after=4)
    add_text(p_method, "Способ закупки: ")
    add_placeholder(p_method, "{{purchase_method}}")

    p_term = new_para(doc, space_after=4)
    add_text(p_term, "Срок: ")
    add_placeholder(p_term, "{{service_term}}")

    p_nmcd = new_para(doc, space_after=4)
    add_text(p_nmcd, "НМЦД: ")
    add_placeholder(p_nmcd, "{{total_nmcd}}")

    p_feo = new_para(doc, space_after=4)
    add_text(p_feo, "Категория ФЭО: ")
    add_placeholder(p_feo, "{{feo_path}}")

    p_sub = new_para(doc, space_after=6)
    add_text(p_sub, "Субсидия: ")
    add_placeholder(p_sub, "{{subsidy_name}}")

    p_just = new_para(doc, space_after=6)
    add_text(p_just, "Обоснование необходимости: [пользователь добавит].")

    p_deadline = new_para(doc, space_after=4)
    add_text(p_deadline, "Срок подачи заявок (если конкурс): ")
    add_placeholder(p_deadline, "{{submission_deadline_datetime}}")

    # Место — условный блок
    p_place = new_para(doc, space_after=6)
    add_text(p_place, "Место ")
    add_marker(p_place, "{% if subject_kind == 'goods' %}")
    add_branch_false(p_place, "поставки")
    add_marker(p_place, "{% else %}")
    add_branch_true(p_place, "оказания")
    add_marker(p_place, "{% endif %}")
    add_text(p_place, ": ")
    add_placeholder(p_place, "{{delivery_location}}")


def add_body_advance(doc):
    """СЗ на аванс / авансовый отчёт."""
    p = new_para(doc, space_after=6)
    add_text(p, "Прошу выдать аванс на расходы по ")
    add_placeholder(p, "{{subject}}")
    add_text(p, " в размере ")
    add_placeholder(p, "{{contract_price_num}}")
    add_text(p, " руб. (")
    add_placeholder(p, "{{contract_price_words}}")
    add_text(p, ").")

    p_resp = new_para(doc, space_after=4)
    add_text(p_resp, "Подотчётное лицо: ")
    add_placeholder(p_resp, "{{responsible_person}}")

    p_deadline = new_para(doc, space_after=4)
    add_text(p_deadline, "Срок отчётности: ")
    add_placeholder(p_deadline, "{{service_deadline_date}}")

    p_feo = new_para(doc, space_after=4)
    add_text(p_feo, "Категория ФЭО: ")
    add_placeholder(p_feo, "{{feo_path}}")

    p_sub = new_para(doc, space_after=6)
    add_text(p_sub, "Субсидия: ")
    add_placeholder(p_sub, "{{subsidy_name}}")

    p_comp_hdr = new_para(doc, space_before=4, space_after=2)
    add_text(p_comp_hdr, "Состав расходов (предварительный):", bold=True)

    # Цикл позиций — обычный {% for %}/{% endfor %} (без {% tr %}, т.к. таблицы нет)
    p_for = new_para(doc, space_after=0)
    add_marker(p_for, "{% for item in items %}", size=10)

    p_item = new_para(doc, space_after=2)
    add_placeholder(p_item, "{{item.num}}")
    add_text(p_item, ". ")
    add_placeholder(p_item, "{{item.name}}")
    add_text(p_item, " — ")
    add_placeholder(p_item, "{{item.total_price}}")

    p_endfor = new_para(doc, space_after=6)
    add_marker(p_endfor, "{% endfor %}", size=10)

    # Phase 26-LL: таблица 2 колонки для чеков (по 2 в строке).
    # Используем {% tr for %} / {% tr endfor %} docxtpl — клонирует строку таблицы.
    p_receipts_hdr = new_para(doc, space_before=4, space_after=2)
    add_text(p_receipts_hdr, "Приложенные чеки:", bold=True)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False

    # Ширина каждой ячейки — 6.5 см (итого 13 см, вписывается в A4 с полями)
    left_cell = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]
    left_cell.width = Cm(6.5)
    right_cell.width = Cm(6.5)

    # Левая ячейка: открывающий {% tr for %} + placeholder левого чека
    left_p1 = left_cell.paragraphs[0]
    add_marker(left_p1, "{% tr for pair in receipt_pairs %}", size=10)
    left_p2 = left_cell.add_paragraph()
    left_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_placeholder(left_p2, "{{ pair.left }}", size=10)

    # Правая ячейка: placeholder правого чека + закрывающий {% tr endfor %}
    right_p1 = right_cell.paragraphs[0]
    right_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_placeholder(right_p1, "{{ pair.right }}", size=10)
    right_p2 = right_cell.add_paragraph()
    add_marker(right_p2, "{% tr endfor %}", size=10)


# ─── ГЛАВНАЯ ФУНКЦИЯ ─────────────────────────────────────────────────────────

def make_service_note(kind, output_path):
    """
    Создать шаблон служебной записки.
    kind ∈ {'general', 'delivery', 'payment', 'procurement', 'advance'}
    """
    doc = Document()
    page_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(12)

    add_legend(doc)
    add_sz_header(doc, kind)

    if kind == 'general':
        add_body_general(doc)
    elif kind == 'delivery':
        add_body_delivery(doc)
    elif kind == 'payment':
        add_body_payment(doc)
    elif kind == 'procurement':
        add_body_procurement(doc)
    elif kind == 'advance':
        add_body_advance(doc)
    else:
        raise ValueError(f"Unknown kind: {kind!r}")

    add_sz_signature(doc)

    doc.save(output_path)
    print(f"OK: {output_path}")


# ─── MAIN ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    KINDS = {
        'general':     "service_note.docx",
        'delivery':    "service_note_delivery.docx",
        'payment':     "service_note_payment.docx",
        'procurement': "service_note_procurement.docx",
        'advance':     "service_note_advance.docx",
    }
    for kind, filename in KINDS.items():
        out = os.path.join(TEMPLATES_DIR, filename)
        make_service_note(kind, out)
