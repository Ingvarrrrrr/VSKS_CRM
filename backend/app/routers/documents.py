import os
import re
from io import BytesIO
from datetime import date
from decimal import Decimal
from urllib.parse import quote
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from sqlalchemy.ext.asyncio import AsyncSession
from app.database import get_db
from app.models.purchase import Purchase
from app.models.purchase_item import PurchaseItem
from app.models.contractor import Contractor
from app.models.subsidy import Subsidy
from app.models.subsidy_approver import SubsidyApprover
from app.models.feo_category import FeoCategory
from app.models.event import Event
from app.auth.jwt import get_current_user
from app.services.fio import compose_fio as _compose_fio
from app.services.responsible_role import is_responsible_role, is_blank_person_name, RESPONSIBLE_PLACEHOLDER
from typing import Optional
import logging

logger = logging.getLogger(__name__)


def _sanitize_subject(text: str, max_len: int = 50) -> str:
    """Очищает произвольный текст для использования в имени файла.

    Убирает символы \\ / : * ? " < > | и переносы строк,
    схлопывает пробелы в «_», обрезает до max_len символов.
    """
    if not text:
        return ""
    cleaned = re.sub(r'[\\/:*?"<>|\r\n]+', "", text)
    cleaned = re.sub(r'\s+', "_", cleaned.strip())
    return cleaned[:max_len]


# Phase 26-V: падежные склонения для шаблонов СЗ.
# pymorphy3 — для общих слов (должность), petrovich — для ФИО.
_morph = None
_petro = None

def _get_morph():
    global _morph
    if _morph is None:
        try:
            from pymorphy3 import MorphAnalyzer
            _morph = MorphAnalyzer()
        except Exception:
            _morph = False
    return _morph or None

def _get_petro():
    global _petro
    if _petro is None:
        try:
            from petrovich.main import Petrovich
            from petrovich.enums import Case, Gender
            _petro = (Petrovich(), Case, Gender)
        except Exception:
            _petro = False
    return _petro or None

def _to_gen_word_heuristic(word: str) -> str:
    """Эвристическое склонение слова в родительный падеж без pymorphy3.

    Покрывает 80-90% русских должностей (мужской род, второе склонение).
    Phase 26-SS: fallback после удаления pymorphy3 (OOM на проде, e2dee47).
    """
    if not word or len(word) < 2:
        return word
    lower = word.lower()
    # Сохраняем регистр первой буквы
    cap = word[0].isupper()
    # Окончания:
    # Слова УЖЕ в родительном падеже (атрибуты): «отдела», «управления»,
    # «фирмы» → не трогаем. Эвристика: -а/-я/-ы/-и/-ов/-ей часто = gen.sg/pl.
    if lower.endswith(('ия', 'ой', 'ого', 'его', 'ев', 'ов', 'ей')) and len(lower) > 3:
        result = lower
    elif lower.endswith('ый'):     # «главный» → «главного»
        result = lower[:-2] + 'ого'
    elif lower.endswith('ий'):     # «ведущий» → «ведущего» (после ж/ч/ш/щ → -его),
                                    # «генеральный» → «генерального» (после др. → -ого)
        prev = lower[-3] if len(lower) >= 3 else ''
        result = lower[:-2] + ('его' if prev in ('ж', 'ч', 'ш', 'щ') else 'ого')
    elif lower.endswith('ая'):     # «заведующая» → «заведующей»
        result = lower[:-2] + 'ей'
    elif lower.endswith('я'):      # «дядя» → «дяди»
        result = lower[:-1] + 'и'
    elif lower.endswith('а'):      # «бухгалтера» — уже gen.sg → не трогаем
        result = lower
    elif lower.endswith('ь'):      # «руководитель» → «руководителя»
        result = lower[:-1] + 'я'
    elif lower.endswith('й'):      # «специалистей» — редко
        result = lower[:-1] + 'я'
    elif lower.endswith(('о', 'е', 'у', 'ы', 'э', 'ю')):
        # Несклоняемые: фамилии типа «Лопатко», иностранные «такси»
        result = lower
    else:                          # consonant
        # «специалист» → «специалиста», «директор» → «директора»
        result = lower + 'а'

    if cap:
        result = result[0].upper() + result[1:]
    return result


def _to_gen_word(word: str) -> str:
    """Склонить одно слово в родительный падеж: pymorphy3 → эвристика."""
    if not word:
        return word
    morph = _get_morph()
    if morph:
        try:
            parsed = morph.parse(word)[0]
            inflected = parsed.inflect({'gent'})
            if inflected:
                out = inflected.word
                if word[0].isupper():
                    out = out[0].upper() + out[1:]
                return out
        except Exception:
            pass
    # Phase 26-SS: fallback на эвристику если pymorphy3 нет (e2dee47 OOM revert)
    return _to_gen_word_heuristic(word)

def _to_gen_phrase(phrase: str) -> str:
    """Склонить фразу (должность) в родительный падеж — каждое слово отдельно."""
    if not phrase:
        return phrase
    import re as _re
    parts = _re.split(r'(\s+|-)', phrase)
    return ''.join(_to_gen_word(p) if p.strip() and p != '-' else p for p in parts)

def _inflect_phrase_genitive(phrase: str) -> str:
    """
    Склоняет фразу в родительный падеж пословно (для subject/service_name).
    Примеры:
      «Канцелярские принадлежности» → «канцелярских принадлежностей»
      «Оказание полиграфических услуг» → «оказания полиграфических услуг»
    Результат ВСЕГДА в нижнем регистре (для вставки внутри фразы:
    «Прошу осуществить закупку канцелярских принадлежностей.»).
    Аббревиатуры (ООО), числа, латиница — пропускаются как есть.
    Fallback на per-word эвристику если pymorphy3 недоступен.
    """
    if not phrase or not phrase.strip():
        return ""
    morph = _get_morph()
    import re as _re2
    # Токенизация: русские слова (с дефисом) / прочие токены / пробелы
    tokens = _re2.findall(r"[А-Яа-яЁё]+(?:-[А-Яа-яЁё]+)*|\S+|\s+", phrase)
    out = []
    for tok in tokens:
        if _re2.fullmatch(r"[А-Яа-яЁё]+(?:-[А-Яа-яЁё]+)*", tok):
            # Аббревиатуры (все буквы заглавные, ≥2 символа): не склонять
            if len(tok) >= 2 and tok.isupper():
                out.append(tok)
                continue
            inflected = None
            if morph:
                try:
                    parses = morph.parse(tok.lower())
                    # Pymorphy3 возвращает все возможные разборы по убыванию вероятности.
                    # Для омонимов («принадлежности» = nomn.plur ИЛИ gent.sg) выбираем
                    # разбор в именительном падеже — иначе inflect({'gent'}) у уже-gent
                    # вернёт то же слово (баг «принадлежности» → «принадлежности»).
                    parsed = None
                    for p in parses:
                        if 'nomn' in p.tag:
                            parsed = p
                            break
                    if not parsed and parses:
                        parsed = parses[0]
                    if parsed:
                        infl = parsed.inflect({'gent'})
                        if infl:
                            inflected = infl.word
                except Exception:
                    pass
            # Fallback: per-word эвристика (если pymorphy3 не сработал)
            if not inflected:
                inflected = _to_gen_word_heuristic(tok.lower())
            out.append(inflected.lower())
        else:
            out.append(tok)
    return "".join(out)


def _to_gen_fio(full_name: str) -> str:
    """Склонить ФИО (Фамилия Имя Отчество) в родительный падеж через petrovich."""
    petro_pack = _get_petro()
    if not petro_pack or not full_name:
        return full_name
    petro, Case, Gender = petro_pack
    parts = full_name.strip().split()
    if not parts:
        return full_name
    try:
        gender = Gender.MALE
        if len(parts) >= 3 and parts[2].endswith(('вна', 'чна')):
            gender = Gender.FEMALE
        elif len(parts) >= 3 and parts[2].endswith('вич'):
            gender = Gender.MALE
        result = []
        if len(parts) >= 1:
            result.append(petro.lastname(parts[0], Case.GENITIVE, gender))
        if len(parts) >= 2:
            result.append(petro.firstname(parts[1], Case.GENITIVE, gender))
        if len(parts) >= 3:
            result.append(petro.middlename(parts[2], Case.GENITIVE, gender))
        return ' '.join(result)
    except Exception:
        return _to_gen_phrase(full_name)

router = APIRouter(prefix="/api/purchases", tags=["documents"])
guide_router = APIRouter(prefix="/api/documents", tags=["documents"])

TEMPLATES_DIR = "/app/templates"
SUBSIDY_TEMPLATES_DIR = "/app/uploads/templates"

# Phase 26-ggg: sentinel-маркер для post-render таблицы чеков.
# Пользователь ставит {{ receipts_table }} в шаблон одним параграфом;
# context подставляет эту строку; postprocess находит её и заменяет
# параграф на настоящую docx-таблицу с PNG чеков в каждой ячейке.
RECEIPTS_TABLE_MARKER = "[[RECEIPTS_TABLE_2COL]]"


def _insert_receipts_table_if_marker(
    doc,
    png_paths: list,
    *,
    cols: int = 2,
    col_width_cm: float = 8.0,
    img_width_cm: float = 7.5,
) -> bool:
    """Find paragraph containing RECEIPTS_TABLE_MARKER and replace it with
    a real 2-col docx table containing the receipt PNGs.

    Idempotent: if marker not present — no-op. If png_paths empty — marker
    is cleared but no table is inserted. Returns True if table was inserted.

    Why this exists: docxtpl InlineImage in a paragraph creates inline
    images that clip in narrow cells / overlap with surrounding text
    (Word renders very tall pictures behind subsequent paragraphs). A
    proper <w:tbl> with fixed col widths and an InlineShape inside each
    <w:tc> guarantees layout integrity.
    """
    from docx.shared import Cm as _Cm

    target_para = None
    for p in doc.paragraphs:
        if RECEIPTS_TABLE_MARKER in p.text:
            target_para = p
            break
    if target_para is None:
        return False

    if not png_paths:
        # marker present but no receipts — clear marker text and leave para
        for r in target_para.runs:
            if RECEIPTS_TABLE_MARKER in r.text:
                r.text = r.text.replace(RECEIPTS_TABLE_MARKER, "")
        return False

    rows = (len(png_paths) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    for row in table.rows:
        for c_idx in range(cols):
            row.cells[c_idx].width = _Cm(col_width_cm)

    for idx, png_path in enumerate(png_paths):
        r, c = divmod(idx, cols)
        cell = table.rows[r].cells[c]
        # Cell creation always gives one empty paragraph — clear its runs
        # then add the picture inside a fresh run.
        para = cell.paragraphs[0]
        for run in list(para.runs):
            run.text = ""
        para.add_run().add_picture(png_path, width=_Cm(img_width_cm))

    # Move the freshly appended table from end-of-doc to the marker location,
    # then delete the marker paragraph.
    target_xml = target_para._p
    table_xml = table._tbl
    target_xml.addnext(table_xml)
    target_xml.getparent().remove(target_xml)
    return True

DOC_TYPES = {
    "service_note_delivery": ("service_note_delivery.docx", "Служебная_записка_выдача"),
    "service_note_payment":  ("service_note_payment.docx",  "Служебная_записка_оплата"),
    # Phase 19.05: dedicated SZ for procurement (distinct from generic service_note)
    "service_note_procurement": ("service_note_procurement.docx", "Служебная_записка_закупка"),
    # Phase 19.07: СЗ на аванс
    "service_note_advance":  ("service_note_advance.docx",  "Служебная_записка_аванс"),
    # Legacy — kept for backwards compat with existing uploaded subsidy overrides.
    "contract_tz":           ("contract_tz.docx",           "Договор_с_ТЗ"),
    # tech_spec falls back to contract_tz.docx — separation kept for future
    # when a dedicated tech_spec template is uploaded, but both resolve to
    # the same file today so there is no confusing "empty ТЗ slot" in UI.
    "tech_spec":             ("contract_tz.docx",           "Техническое_задание"),
    # Phase 19.05: split ТЗ into request-of-prices and contract-appendix variants.
    # Default template file is a copy of contract_tz.docx; admins upload
    # per-subsidy overrides via SubsidiesView.
    "tech_spec_request":     ("tech_spec_request.docx",     "ТЗ_запрос_цен"),
    "tech_spec_contract":    ("tech_spec_contract.docx",    "ТЗ_к_договору"),
    "contract":              ("contract.docx",              "Договор"),
    # Phase 23.1: contract_services / contract_goods merged into universal contract.docx
    # (removed — subject_kind auto-detected from purchase_items.product.item_kind)
    "approval_sheet":        ("approval_sheet.docx",        "Лист_согласования"),
    "order_purchase":        ("order_purchase.docx",        "Приказ_о_закупке"),
    # Phase 28: typed contract forms per-subsidy
    # Форма «услуги» объединена в один файл (contract_services.docx) — большая/малая
    # отчётность теперь не отдельные шаблоны договора, а отдельная методичка
    # (methodology_large / methodology_small), приклеиваемая к готовому договору.
    "contract_services":             ("contract_services.docx",            "Договор_услуг"),
    # Алиасы на новый объединённый файл — НЕ удалять: у уже существующих закупок
    # doc_type мог быть сохранён/запрошен под старым именем, без алиаса — 404.
    "contract_services_large":      ("contract_services.docx",      "Договор_услуги_крупный"),
    "contract_services_small":      ("contract_services.docx",      "Договор_услуги_малый"),
    "contract_services_food":       ("contract_services_food.docx",       "Договор_услуги_питание"),
    "methodology_large":             ("methodology_large.docx",            "Методические_рекомендации_большие"),
    "methodology_small":             ("methodology_small.docx",            "Методические_рекомендации_малые"),
    "contract_goods_single":        ("contract_goods_single.docx",        "Договор_поставка_единственный"),
    "contract_gph_individual":      ("contract_gph_individual.docx",      "Договор_ГПХ_физлицо"),
    "contract_gph_individual_rid":  ("contract_gph_individual_rid.docx",  "Договор_ГПХ_физлицо_РИД"),
    "contract_repair_vehicle":      ("contract_repair_vehicle.docx",      "Договор_ремонт_ТС"),
    "contract_repair_framework":    ("contract_repair_framework.docx",    "Договор_ремонт_рамочный"),
    # Fabrikant ЭТП package — four template types for запрос цен на Фабрикант
    "fabrikant_instruction":        ("fabrikant_instruction.docx",        "Фабрикант_инструкция"),
    "fabrikant_application_form":   ("fabrikant_application_form.docx",   "Фабрикант_форма_заявки"),
    "fabrikant_documentation":      ("fabrikant_documentation.docx",      "Фабрикант_документация"),
    "fabrikant_contract_project":   ("fabrikant_contract_project.docx",   "Фабрикант_проект_договора"),
}

# Требование владельца («Плановые не равно Договор»): для этих типов документ
# печатает ТОЛЬКО позиции и суммы "Как в договоре" (ContractItem) — без
# молчаливого отката на плановые purchase_items / НМЦК. Лист согласования
# входит сюда же: он визирует именно то, что уйдёт в договор.
#
# НЕ включены намеренно:
#   - tech_spec_request — ТЗ для ЗАПРОСА цен, по смыслу плановый документ;
#   - service_note_* — служебные записки оформляются ДО заключения договора;
#   - fabrikant_* — пакет документов для тендерной процедуры на Фабрикант,
#     публикуется ДО выбора поставщика, договора ещё не существует;
#   - order_purchase — приказ о закупке, тоже предшествует договору.
# tech_spec_request и tech_spec_contract сейчас оба резолвятся в один файл
# contract_tz.docx (см. DOC_TYPE_FALLBACK_FILES), но теперь ведут себя
# по-разному — это ожидаемо и намеренно.
CONTRACT_FAMILY_DOC_TYPES = {
    "contract_services",
    # Алиасы — держим в семье тоже, чтобы гейт «Плановые не равно Договор»
    # срабатывал одинаково, если старый doc_type всё же запрошен.
    "contract_services_large",
    "contract_services_small",
    "contract_services_food",
    "contract_goods_single",
    "contract_gph_individual",
    "contract_gph_individual_rid",
    "contract_repair_vehicle",
    "contract_repair_framework",
    "contract",
    "contract_tz",
    "tech_spec_contract",
    "approval_sheet",
}

# Семь типовых форм договора (+ 2 legacy-алиаса на contract_services) — только
# к НИМ приклеивается методичка (Purchase.methodology), см. generate_document
# ниже. 'contract' (старый универсальный шаблон), 'approval_sheet',
# 'contract_tz', 'tech_spec_contract' — не подписываемый контрагентом текст
# договора, методичка к ним не приклеивается.
CONTRACT_TYPED_FORM_DOC_TYPES = {
    "contract_services",
    "contract_services_large",
    "contract_services_small",
    "contract_services_food",
    "contract_goods_single",
    "contract_gph_individual",
    "contract_gph_individual_rid",
    "contract_repair_vehicle",
    "contract_repair_framework",
}

# Phase 19.05: fallback map — if a dedicated template file is missing,
# fall back to the legacy file so the endpoint still works before admins
# upload per-subsidy overrides.
DOC_TYPE_FALLBACK_FILES = {
    "service_note_procurement": "service_note.docx",
    # Phase 19.07: fall back to generic service_note until a dedicated
    # advance template is uploaded (per-subsidy or globally).
    "service_note_advance":     "service_note.docx",
    "tech_spec_request":        "contract_tz.docx",
    "tech_spec_contract":       "contract_tz.docx",
    # Phase 28: new typed contract forms fall back to universal contract.docx
    # until per-subsidy templates are uploaded. Prevents 404 for old purchases.
    "contract_services_large":     "contract.docx",
    "contract_services_small":     "contract.docx",
    "contract_services_food":      "contract.docx",
    "contract_goods_single":       "contract.docx",
    "contract_gph_individual":     "contract.docx",
    "contract_gph_individual_rid": "contract.docx",
    "contract_repair_vehicle":     "contract.docx",
    "contract_repair_framework":   "contract.docx",
}


def _resolve_doc_template_path(doc_type: str, subsidy_id: Optional[int]) -> tuple[str, str, str]:
    """Резолвит путь к файлу шаблона doc_type с единым приоритетом:

      1) субсидийный override (uploads/templates/subsidies/<id>/<doc_type>.docx),
      2) глобальный файл backend/templates/<doc_type>.docx,
      3) DOC_TYPE_FALLBACK_FILES[doc_type], если глобальный файл отсутствует.

    Используется и для основного документа, и для приклеиваемой методички —
    единственное место с этой логикой, чтобы приоритет «сначала субсидия,
    иначе глобальный» не разъезжался между местами вызова.

    Возвращает (template_path, template_file, filename_base). template_path
    может не существовать — вызывающий код обязан проверить os.path.exists.
    """
    template_file, filename_base = DOC_TYPES[doc_type]
    template_path = os.path.join(TEMPLATES_DIR, template_file)

    if not os.path.exists(template_path):
        fallback = DOC_TYPE_FALLBACK_FILES.get(doc_type)
        if fallback:
            fallback_path = os.path.join(TEMPLATES_DIR, fallback)
            if os.path.exists(fallback_path):
                template_path = fallback_path
                template_file = fallback

    if subsidy_id:
        subsidy_template = os.path.join(
            SUBSIDY_TEMPLATES_DIR, "subsidies", str(subsidy_id), f"{doc_type}.docx"
        )
        if os.path.exists(subsidy_template):
            template_path = subsidy_template

    return template_path, template_file, filename_base


# Fields required to generate a FADM contract; maps field_path → label
_CONTRACT_REQUIRED_FIELDS = {
    "contractor": "Контрагент",
    "contractor.org_type": "Тип контрагента (Юр.лицо / ИП / Самозанятый)",
    "contractor.inn": "ИНН контрагента",
    "contractor.address": "Юридический адрес контрагента",
    "contractor.settlement_account": "Расчётный счёт контрагента",
    "contractor.bank_name": "Наименование банка",
    "contractor.bik": "БИК банка",
    "contract_number": "Номер договора",
    "contract_date": "Дата договора",
    "subject": "Предмет договора (услуги)",
    "contract_price": "Цена договора",
    "service_period_type": "Тип срока (период / дата)",
    "execution_term": "Срок оказания услуг",
}

_BASIS_LABELS = {
    "plan_schedule": "план закупок",
    "service_note": "служебная записка",
}

# Способ закупки (Purchase.purchase_method) → человекочитаемая подпись.
# Значения см. frontend/src/views/CreateOrderView.vue (v-select purchase_method)
# и ContractsView.vue (purchaseMethodItems): в системе только 3 кода.
# Неизвестный код — не пустая строка, а сам код (см. order_purchase, п.2).
_PURCHASE_METHOD_LABELS = {
    "single": "Единственный поставщик",
    "competitive": "Конкурсная процедура",
    "advance": "Авансовый отчёт",
}

# Уточняющая форма конкурентной процедуры (Purchase.competitive_form).
# Владелец: «"Запрос цен" — это вариант конкурсной процедуры, так же как и
# Аукцион — он же редукцион, а также Конкурс» — НЕ отдельный способ закупки,
# применимо только когда purchase_method == 'competitive'.
_COMPETITIVE_FORM_LABELS = {
    "price_request": "Запрос цен",
    "auction": "Аукцион (редукцион)",
    "tender": "Конкурс",
}


def _purchase_method_label(p) -> str:
    """Подпись способа закупки для п.2 приказа/листа согласования.

    Если способ — конкурентная процедура и заполнена уточняющая форма
    (competitive_form), подписью становится конкретная форма («Запрос цен» /
    «Аукцион (редукцион)» / «Конкурс»), а не общее «Конкурсная процедура».
    Форма не выбрана — старое поведение (общая подпись способа).
    """
    method = p.purchase_method or ""
    if method == "competitive":
        form = getattr(p, "competitive_form", None)
        form_label = _COMPETITIVE_FORM_LABELS.get(form or "")
        if form_label:
            return form_label
    return _PURCHASE_METHOD_LABELS.get(method, method)


# Приказ о закупке и лист согласования визируют способ закупки как отдельный
# распорядительный пункт — без него документ юридически бессмысленный
# («2. Определить способ закупки: .»). Владелец подтвердил для приказа
# уверенно, для листа согласования — «наверное тоже надо» (менее строго).
PURCHASE_METHOD_REQUIRED_DOC_TYPES = {"order_purchase", "approval_sheet"}


def _require_purchase_method_for_doc(p, doc_type: str) -> None:
    """422, если запрошен приказ о закупке/лист согласования, а способ закупки не выбран."""
    if doc_type not in PURCHASE_METHOD_REQUIRED_DOC_TYPES:
        return
    if p.purchase_method:
        return
    raise HTTPException(
        422,
        detail={
            "code": "PURCHASE_METHOD_REQUIRED",
            "message": "Не выбран способ закупки",
            "hint": (
                "Способ закупки — обязательный пункт приказа о закупке и листа "
                "согласования: без него документ считается недействительным. "
                "Откройте карточку закупки и заполните поле «Способ закупки». "
                "Если это конкурентная процедура, укажите ещё и её форму — "
                "запрос цен, аукцион или конкурс."
            ),
            "missing_fields": ["purchase_method"],
            "doc_type": doc_type,
        },
    )


def _fmt_date(d) -> str:
    if not d:
        return ""
    if isinstance(d, str):
        try:
            d = date.fromisoformat(d)
        except ValueError:
            return d
    return d.strftime("%d.%m.%Y")


def _clean_id(v) -> str:
    """Strip trailing .0 from INN/KPP imported as float strings."""
    if not v:
        return ""
    s = str(v).strip()
    return s[:-2] if s.endswith(".0") else s


def _fmt_money(v) -> str:
    if v is None:
        return ""
    return f"{float(v):,.2f}".replace(",", " ").replace(".", ",")


def _fmt_money_plain(v) -> str:
    """Money without currency symbol, with space thousand separator, comma decimal."""
    if v is None:
        return ""
    return f"{float(v):,.2f}".replace(",", " ").replace(".", ",")


def _merge_identical_items(items):
    """Склеивает позиции, у которых совпадает всё, кроме категории ФЭО.

    Возвращает список кортежей (representative_item, merged_quantity, merged_total_price)
    в порядке первого появления. Цена за единицу в группе одна и та же, поэтому
    суммы документа склейка не меняет.
    """
    groups: dict = {}
    order: list = []
    for item in items or []:
        # feo_category_id/feo_planned_item_id намеренно не входят в ключ — по ним и склеиваем
        key = (
            item.product_id,
            (item.item_name or "").strip().lower(),
            (item.item_type or ""),
            (item.unit or ""),
            str(item.unit_price) if item.unit_price is not None else "",
            (item.country_origin or ""),
            (getattr(item, "vat_rate", None) or ""),
        )
        if key not in groups:
            groups[key] = [item, None, None]  # [representative, quantity, total_price]
            order.append(key)
        g = groups[key]
        if item.quantity is not None:
            g[1] = (g[1] if g[1] is not None else Decimal("0")) + Decimal(str(item.quantity))
        if item.total_price is not None:
            g[2] = (g[2] if g[2] is not None else Decimal("0")) + Decimal(str(item.total_price))
    return [tuple(groups[key]) for key in order]


# Signatory position: extract from signatory field if "Директор ФИО" format
def _signatory_position(signatory: str) -> str:
    if not signatory:
        return ""
    parts = signatory.strip().split()
    if len(parts) > 1 and not parts[0][0].isupper():
        return parts[0]
    return "Директор"


# Phase 23: helpers for signatory formatting — used by both paths of _signatory_split

def _fio_to_genitive(name_full: str) -> str:
    """Rough genitive form of a full name: Козеев Евгений Викторович → Козеева Евгения Викторовича."""
    def _word(w: str) -> str:
        if w.endswith("ич"):  # Иванович → Ивановича
            return w + "а"
        if w.endswith("ий"):  # Евгений → Евгения
            return w[:-2] + "ия"
        if w.endswith("ья"):  # Илья → Ильи
            return w[:-2] + "ьи"
        if w.endswith("а") and len(w) > 2:
            return w[:-1] + "ы"
        # Last consonant cluster: add -а
        vowels = set("аеёиоуыьъэюяАЕЁИОУЫЭЮЯ")
        if w and w[-1] not in vowels and w[-1] not in "ьъ":
            return w + "а"
        return w  # fallback: as-is
    return " ".join(_word(w) for w in name_full.split()) if name_full else name_full


def _fio_to_initials(name_full: str) -> str:
    """Return "Фамилия И.О." form; falls back to name_full when < 2 words."""
    words = name_full.split() if name_full else []
    if len(words) >= 3:
        return f"{words[0]} {words[1][0]}.{words[2][0]}."
    if len(words) == 2:
        return f"{words[0]} {words[1][0]}."
    return name_full


def _fio_to_initials_prefix(name_full: str) -> str:
    """Return "И.О. Фамилия" form (инициалы впереди — для строки подписи).
    Falls back to name_full when < 2 words."""
    words = name_full.split() if name_full else []
    if len(words) >= 3:
        return f"{words[1][0]}.{words[2][0]}. {words[0]}"
    if len(words) == 2:
        return f"{words[1][0]}. {words[0]}"
    return name_full


# Phase 23: split "Президент Козеев Евгений Викторович" into structured parts
def _signatory_split(
    signatory: str,
    last: str | None = None,
    first: str | None = None,
    middle: str | None = None,
    position: str | None = None,
) -> dict:
    """Split signatory string into structured dict.

    If structured parts (last/first/middle) are provided — use them directly
    instead of applying heuristics to *signatory*.  *position* is taken from
    the argument when provided.

    Returns:
        position      — должность подписанта
        name_full     — ФИО (Фамилия Имя Отчество)
        name_genitive — rough genitive form
        name_initials — "Козеев Е.В."
    """
    if last:
        # Structured path — no heuristics needed
        name_full = _compose_fio(last, first, middle) or signatory or ""
        return {
            "position": position or "",
            "name_full": name_full,
            "name_genitive": _fio_to_genitive(name_full),
            "name_initials": _fio_to_initials(name_full),
        }

    # Legacy heuristic path (no structured parts available)
    if not signatory:
        return {"position": "", "name_full": "", "name_genitive": "", "name_initials": ""}
    parts = signatory.strip().split()
    if len(parts) <= 1:
        return {"position": "", "name_full": signatory, "name_genitive": signatory, "name_initials": signatory}
    # Heuristic: ФИО = last 3 words if ≥4 words total, else last 2
    if len(parts) >= 4:
        pos_words = parts[:-3]
        fio_words = parts[-3:]  # Фамилия Имя Отчество
    elif len(parts) == 3:
        # Could be "Иванов Иван Иванович" (no position) or "Директор Иванов Иван"
        # Heuristic: first word starts with uppercase → probably all ФИО or pos+2
        pos_words = []
        fio_words = parts
    else:
        pos_words = parts[:1]
        fio_words = parts[1:]

    resolved_position = position or " ".join(pos_words)
    name_full = " ".join(fio_words)

    return {
        "position": resolved_position,
        "name_full": name_full,
        "name_genitive": _fio_to_genitive(name_full),
        "name_initials": _fio_to_initials(name_full),
    }


def _sum_items_price(p) -> float:
    """Fallback: manually sum item.total_price when p.total_nmck is NULL."""
    items = getattr(p, "items", None) or []
    total = 0.0
    for it in items:
        try:
            total += float(getattr(it, "total_price", 0) or 0)
        except Exception:
            pass
    return total


def _build_acceptance_doc_context(p, doc_type: str, doc_indices_csv: str | None) -> dict:
    """Phase 27.2-01: build acceptance_doc_* context entries.

    For service_note_payment / service_note_advance:
      - читаем acceptance_docs JSONB (source of truth после Phase 26-H);
      - если doc_indices задан — фильтруем по индексам;
      - first doc → name/number/date; сумма всех selected → amount.
    Для остальных doc_type и при отсутствии JSONB-данных — legacy fallback.
    """
    SZ_TYPES = ("service_note_payment", "service_note_advance")

    def _legacy_amount():
        return _fmt_money(
            p.acceptance_doc_amount
            or p.contract_price
            or p.planned_total_price
            or p.total_nmck
            or _sum_items_price(p)
            or 0
        )

    if doc_type in SZ_TYPES:
        raw_docs: list = p.acceptance_docs or []
        if raw_docs:
            # Filter by doc_indices if provided
            if doc_indices_csv:
                try:
                    indices = [int(i.strip()) for i in doc_indices_csv.split(",") if i.strip().isdigit()]
                    selected = [raw_docs[i] for i in indices if 0 <= i < len(raw_docs)]
                except Exception:
                    selected = raw_docs
            else:
                selected = raw_docs

            if selected:
                def _fmt_doc_date(raw_date) -> str:
                    if isinstance(raw_date, str) and raw_date:
                        try:
                            from datetime import date as _date
                            return _date.fromisoformat(raw_date).strftime("%d.%m.%Y")
                        except ValueError:
                            return raw_date
                    return _fmt_date(raw_date)

                total_amount = sum(float(d.get("amount") or 0) for d in selected)
                return {
                    "acceptance_doc_name":   "; ".join(d.get("name") or "" for d in selected),
                    "acceptance_doc_number": "; ".join(d.get("number") or "" for d in selected),
                    "acceptance_doc_date":   "; ".join(_fmt_doc_date(d.get("date") or "") for d in selected),
                    "acceptance_doc_amount": _fmt_money(total_amount) if total_amount else _legacy_amount(),
                }

    # Legacy fallback (all other doc_types OR no JSONB data)
    return {
        "acceptance_doc_name":   p.acceptance_doc_name or "",
        "acceptance_doc_number": p.acceptance_doc_number or "",
        "acceptance_doc_date":   _fmt_date(p.acceptance_doc_date) or "",
        "acceptance_doc_amount": (
            _fmt_money(p.acceptance_doc_amount)
            if p.acceptance_doc_amount
            else _legacy_amount()
        ),
    }


async def _build_contract_items_context(p, db) -> dict:
    """Phase 27.1 CD-5: build docxtpl context entries for {{contract_items}} loop.

    Returns dict with keys:
      - contract_items: list[dict] for loop in template (fields num/name/quantity/unit/unit_price/total)
      - contract_items_total: formatted string like "150 000,00 ₽"
      - contract_items_total_numeric: float for arithmetic in template
      - contract_item_count: int, len of contract_items list

    Fallback (D-08 deprecated alias): if purchase has no contract_items —
    populate from purchase_items so legacy templates keep working.
    """
    from app.models.contract_item import ContractItem as _ContractItem
    ci_query = await db.execute(
        select(_ContractItem)
        .where(_ContractItem.purchase_id == p.id)
        .order_by(_ContractItem.id)
    )
    contract_items_db = ci_query.scalars().all()

    result_list = []
    total_numeric = 0.0
    if contract_items_db:
        # Primary path: contract_items exist — use them
        for idx, ci in enumerate(contract_items_db, start=1):
            tot = float(ci.total) if ci.total else 0.0
            result_list.append({
                "num": idx,
                "name": ci.name or "",
                "quantity": float(ci.quantity) if ci.quantity else "",
                "unit": ci.unit or "",
                "unit_price": _fmt_money(ci.unit_price),
                "total": _fmt_money(ci.total),
                "total_numeric": tot,
            })
            total_numeric += tot
    else:
        # D-08 fallback: contract_items empty — use purchase_items as deprecated alias
        for idx, (item, qty, total) in enumerate(_merge_identical_items(getattr(p, "items", None) or []), start=1):
            tot = float(total) if total else 0.0
            result_list.append({
                "num": idx,
                "name": item.item_name or "",
                "quantity": float(qty) if qty else "",
                "unit": item.unit or "",
                "unit_price": _fmt_money(item.unit_price),
                "total": _fmt_money(total),
                "total_numeric": tot,
            })
            total_numeric += tot

    return {
        "contract_items": result_list,
        "contract_items_total": _fmt_money(total_numeric),
        "contract_items_total_numeric": total_numeric,
        "contract_item_count": len(result_list),
    }


def _contract_items_total(p) -> float:
    """Сумма ContractItem.total закупки (relationship должен быть eager-loaded)."""
    total = 0.0
    for ci in (getattr(p, "contract_items", None) or []):
        try:
            total += float(ci.total or 0)
        except Exception:
            pass
    return total


FEO_PATH_UNRESOLVED_LABEL = "Категория ФЭО не определена (позиция договора без привязки к плановой)"


def _build_contract_item_feo_paths(contract_items, plan_items, feo_path_nodes) -> list:
    """Путь ФЭО по ДОГОВОРНЫМ позициям для листа согласования.

    Требование владельца (2026-08-30): «В "Путь ФЭО" должны прописываться
    договорные категории, они должны совпадать с Плановыми». У ContractItem
    своего поля ФЭО нет — категория берётся у плановой PurchaseItem, из
    которой договорная позиция скопирована (``source_item_id``). Этим
    совпадение договорных категорий с плановыми гарантировано по построению:
    это буквально та же запись категории, а не пересчёт/эвристика.

    Договорная позиция без ``source_item_id`` (заведена вручную, а не
    копированием из плана) не может унаследовать чужую категорию — попадает
    в отдельную группу с меткой :data:`FEO_PATH_UNRESOLVED_LABEL`, чтобы в
    документе было видно, что категория не определена, вместо того чтобы
    либо промолчать, либо подставить произвольную.

    Args:
        contract_items: iterable ContractItem (source_item_id, total).
        plan_items: iterable PurchaseItem (id, feo_category_id) — используется
            только как справочник id → feo_category_id, суммы отсюда не берутся.
        feo_path_nodes: callable(category_id) -> list[FeoCategory] (root → leaf),
            та же функция, что строит feo_path/feo_level_* выше по контексту.

    Returns:
        [(path_str, Decimal total), ...] — по одной строке на категорию, в
        порядке первого появления среди contract_items; группа «не
        определена» (если есть) — последней.
    """
    plan_feo_by_id = {
        it.id: it.feo_category_id for it in (plan_items or []) if getattr(it, "feo_category_id", None)
    }
    cat_sums: dict = {}
    order: list = []
    unresolved_sum = Decimal("0")
    has_unresolved = False
    for ci in (contract_items or []):
        amt = Decimal(str(ci.total)) if ci.total is not None else Decimal("0")
        source_id = getattr(ci, "source_item_id", None)
        cid = plan_feo_by_id.get(source_id) if source_id else None
        if cid:
            if cid not in cat_sums:
                cat_sums[cid] = Decimal("0")
                order.append(cid)
            cat_sums[cid] += amt
        else:
            has_unresolved = True
            unresolved_sum += amt

    item_feo_paths: list = []
    for cid in order:
        cat_path = " → ".join(n.name.strip() for n in feo_path_nodes(cid))
        if cat_path:
            item_feo_paths.append((cat_path, cat_sums[cid]))
    if has_unresolved:
        item_feo_paths.append((FEO_PATH_UNRESOLVED_LABEL, unresolved_sum))
    return item_feo_paths


def _build_items_list_from_contract_items(p, resolve_photo=None) -> list[dict]:
    """items_list для договорных документов — ТОЛЬКО из ContractItem.

    Поля те же, что у items_list из purchase_items (num/name/description/
    type/item_kind/quantity/unit/unit_price/total_price/total/photo/code/
    norm_hours), чтобы существующие шаблоны (contract_tz.docx,
    contract_repair_framework.docx и др.), перебирающие {{items}}, работали
    без правки .docx. Никакого отката на purchase_items здесь нет —
    «Плановые не равно Договор» (требование владельца).
    """
    items_list: list[dict] = []
    for idx, ci in enumerate(getattr(p, "contract_items", None) or [], start=1):
        product = getattr(ci, "product", None)
        photo_url = getattr(product, "photo_url", None) if product else None
        items_list.append({
            "num": idx,
            "name": ci.name or "",
            "description": ((getattr(product, "description", None) if product else "") or ""),
            "type": "",
            "item_kind": (getattr(product, "item_kind", None) if product else None) or "товар",
            "quantity": float(ci.quantity) if ci.quantity else "",
            "unit": ci.unit or "",
            "unit_price": _fmt_money(ci.unit_price),
            "total_price": _fmt_money(ci.total),
            "total": _fmt_money(ci.total),
            "photo": resolve_photo(photo_url) if resolve_photo else "",
            # Поля для repair_framework (появятся позже, пока заглушки)
            "code": "",
            "norm_hours": "",
        })
    return items_list


def _build_items_list_from_purchase_items(p, tz_override_mode=None, resolve_photo=None) -> list[dict]:
    """items_list для плановых документов — из purchase_items (как раньше).

    Извлечено без изменения поведения из тела generate_document(), чтобы
    логика выбора источника (план vs договор) была тестируемой чистой
    функцией и переиспользовалась в обоих местах построения контекста.
    """
    description_mode = tz_override_mode or getattr(p, "description_mode", None) or "exact"
    items_list: list[dict] = []
    for idx, (item, qty, total) in enumerate(_merge_identical_items(getattr(p, "items", None) or []), start=1):
        photo_url = item.product.photo_url if item.product else None
        items_list.append({
            "num": idx,
            "name": item.item_name or "",
            "description": (
                (item.product.description_44fz if description_mode == "44fz" else item.product.description)
                if item.product else ""
            ) or "",
            "type": item.item_type or "",
            "item_kind": (item.product.item_kind if item.product else None) or "товар",
            "quantity": float(qty) if qty else "",
            "unit": item.unit or "",
            "unit_price": _fmt_money(item.unit_price),
            "total_price": _fmt_money(total),
            "total": _fmt_money(total),
            "photo": resolve_photo(photo_url) if resolve_photo else "",
            # Поля для repair_framework (появятся позже, пока заглушки)
            "code": "",
            "norm_hours": "",
        })
    return items_list


def _resolve_doc_amount(p, doc_type: str) -> tuple[float, bool]:
    """Вернуть (doc_amount_val, amount_is_planned) — сумму документа.

    Для CONTRACT_FAMILY_DOC_TYPES сумма берётся ТОЛЬКО из p.contract_price
    или суммы ContractItem, БЕЗ отката на НМЦК/план/сумму плановых позиций
    (требование владельца: «Плановые не равно Договор»). total_nmck/nmck
    сюда не входят — они остаются плановыми полями в контексте документа.

    Для остальных типов — старое поведение: контракт → НМЦК/план → сумма
    плановых позиций.
    """
    if doc_type in CONTRACT_FAMILY_DOC_TYPES:
        amount = float(p.contract_price or 0) or _contract_items_total(p)
        return amount, False
    items_sum_val = float(sum(Decimal(str(it.total_price or 0)) for it in (getattr(p, "items", None) or [])))
    doc_amount_val = (float(p.contract_price or 0) or float(getattr(p, "total_nmck", None) or 0)
                      or float(getattr(p, "nmck", None) or 0) or float(getattr(p, "planned_total_price", None) or 0)
                      or items_sum_val)
    amount_is_planned = not bool(p.contract_price)
    return doc_amount_val, amount_is_planned


def _require_contract_items_for_doc(p, doc_type: str) -> None:
    """422, если запрошен договорной документ, а позиции договора не заполнены.

    Молчаливый откат на purchase_items здесь недопустим — «Плановые не равно
    Договор» (требование владельца). Код ошибки CONTRACT_ITEMS_REQUIRED тот
    же, что и в purchase_transitions.py (переход в «Заключён договор»), для
    единообразия обработки на фронте.
    """
    if doc_type not in CONTRACT_FAMILY_DOC_TYPES:
        return
    if getattr(p, "contract_items", None):
        return
    raise HTTPException(
        422,
        detail={
            "code": "CONTRACT_ITEMS_REQUIRED",
            "message": "В закупке пока не заполнены позиции договора",
            "hint": (
                "Договор печатается по списку «Как в договоре». Это отдельный список: "
                "в нём стоят наименования, количество и цены, о которых вы договорились "
                "с поставщиком, — они могут отличаться от того, что заявляли изначально. "
                "Сейчас этот список пуст, поэтому подставить в договор нечего.\n\n"
                "Что сделать: в блоке позиций нажмите «Скопировать из заявки» — плановые "
                "строки перенесутся в договорные. Затем поправьте наименования и цены так, "
                "как в подписываемом договоре, и сформируйте документ заново."
            ),
            "missing_fields": ["contract_items"],
            "doc_type": doc_type,
        },
    )


async def _resolve_user_dept(user, db, org_id: Optional[int] = None) -> str:
    """Возвращает название отдела пользователя для шаблона СЗ.

    Бизнес-правило: должность/отдел берутся per-org, по той организации,
    к которой привязана субсидия закупки. Если юзер в этой org в нескольких
    отделах — берём первый (по user_organizations.id).

    Приоритет:
      1) Department.name через user_organizations где org_id == org_id (если задан)
      2) Department.name через user_organizations первая запись (любая org)
      3) User.department (legacy строковое поле)
      4) "" если ничего нет
    """
    if user is None:
        return ""
    try:
        from app.models.user_organization import UserOrganization
        from app.models.department import Department
        from sqlalchemy import select as _sel
        # 1. В пределах org закупки
        if org_id is not None:
            res = await db.execute(
                _sel(Department.name)
                .join(UserOrganization, UserOrganization.dept_id == Department.id)
                .where(
                    UserOrganization.user_id == user.id,
                    UserOrganization.org_id == org_id,
                )
                .order_by(UserOrganization.id)
                .limit(1)
            )
            name = res.scalar_one_or_none()
            if name:
                return name
        # 2. Любая первая
        res = await db.execute(
            _sel(Department.name)
            .join(UserOrganization, UserOrganization.dept_id == Department.id)
            .where(UserOrganization.user_id == user.id)
            .order_by(UserOrganization.id)
            .limit(1)
        )
        name = res.scalar_one_or_none()
        if name:
            return name
    except Exception:
        pass
    # 3. Legacy строковое поле
    return getattr(user, "department", None) or ""


async def _resolve_user_position(user, db, org_id: Optional[int] = None) -> str:
    """Возвращает должность пользователя для шаблона СЗ.

    Приоритет:
      1) user_organizations.position где org_id == org_id (если задан и не пустой)
      2) user_organizations.position первая (любая org)
      3) User.position (legacy)
      4) "" если нет
    """
    if user is None:
        return ""
    try:
        from app.models.user_organization import UserOrganization
        from sqlalchemy import select as _sel
        if org_id is not None:
            res = await db.execute(
                _sel(UserOrganization.position)
                .where(
                    UserOrganization.user_id == user.id,
                    UserOrganization.org_id == org_id,
                )
                .order_by(UserOrganization.id)
                .limit(1)
            )
            pos = res.scalar_one_or_none()
            if pos:
                return pos
        res = await db.execute(
            _sel(UserOrganization.position)
            .where(UserOrganization.user_id == user.id)
            .order_by(UserOrganization.id)
            .limit(1)
        )
        pos = res.scalar_one_or_none()
        if pos:
            return pos
    except Exception:
        pass
    return getattr(user, "position", None) or ""


def _format_service_term(p) -> str:
    """Build the human-readable service-term string for docx templates.

    Phase 19 — three modes:
      - range:    "с 01.05.2026 по 31.05.2026"
      - duration: "в течение 30 календарных дней после заключения договора"
      - deadline: "до 30.06.2026 включительно"
    """
    mode = getattr(p, "service_term_mode", None)
    if mode == "range" and p.service_start_date and p.service_end_date:
        return (
            f"с {p.service_start_date.strftime('%d.%m.%Y')} "
            f"по {p.service_end_date.strftime('%d.%m.%Y')}"
        )
    if mode == "duration" and getattr(p, "service_term_days", None):
        type_name = {
            "calendar": "календарных",
            "working":  "рабочих",
        }.get(getattr(p, "service_term_type", None) or "calendar", "календарных")
        return f"в течение {p.service_term_days} {type_name} дней после заключения договора"
    if mode == "deadline" and getattr(p, "service_deadline_date", None):
        return f"до {p.service_deadline_date.strftime('%d.%m.%Y')} включительно"
    return ""


_ONES = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять",
         "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
         "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
_TENS = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят",
         "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
_HUNDREDS = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
             "шестьсот", "семьсот", "восемьсот", "девятьсот"]
_ONES_F = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять",
           "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
           "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]


def _chunk_to_words(n: int, feminine: bool = False) -> str:
    parts = []
    h = n // 100
    r = n % 100
    t = r // 10
    o = r % 10
    if h:
        parts.append(_HUNDREDS[h])
    if r < 20:
        w = (_ONES_F[r] if feminine else _ONES[r])
        if w:
            parts.append(w)
    else:
        if t:
            parts.append(_TENS[t])
        w = (_ONES_F[o] if feminine else _ONES[o])
        if w:
            parts.append(w)
    return " ".join(parts)


def _rubles_to_words(amount) -> str:
    """Convert numeric amount to Russian words (рублей XX копеек)."""
    if amount is None:
        return ""
    try:
        val = round(float(amount), 2)
    except (TypeError, ValueError):
        return ""
    rubles = int(val)
    kopecks = round((val - rubles) * 100)
    billions = rubles // 1_000_000_000
    millions = (rubles % 1_000_000_000) // 1_000_000
    thousands = (rubles % 1_000_000) // 1_000
    rest = rubles % 1_000

    parts = []
    if billions:
        w = _chunk_to_words(billions)
        if w:
            parts.append(w)
        if billions % 10 == 1 and billions % 100 != 11:
            parts.append("миллиард")
        elif 2 <= billions % 10 <= 4 and not (11 <= billions % 100 <= 14):
            parts.append("миллиарда")
        else:
            parts.append("миллиардов")
    if millions:
        w = _chunk_to_words(millions)
        if w:
            parts.append(w)
        if millions % 10 == 1 and millions % 100 != 11:
            parts.append("миллион")
        elif 2 <= millions % 10 <= 4 and not (11 <= millions % 100 <= 14):
            parts.append("миллиона")
        else:
            parts.append("миллионов")
    if thousands:
        w = _chunk_to_words(thousands, feminine=True)
        if w:
            parts.append(w)
        if thousands % 10 == 1 and thousands % 100 != 11:
            parts.append("тысяча")
        elif 2 <= thousands % 10 <= 4 and not (11 <= thousands % 100 <= 14):
            parts.append("тысячи")
        else:
            parts.append("тысяч")
    if rest or not parts:
        w = _chunk_to_words(rest, feminine=True)
        if w:
            parts.append(w)
    if rubles == 0:
        parts = ["ноль"]

    # Ruble ending
    r10 = rubles % 10
    r100 = rubles % 100
    if r10 == 1 and r100 != 11:
        rub_word = "рубль"
    elif 2 <= r10 <= 4 and not (11 <= r100 <= 14):
        rub_word = "рубля"
    else:
        rub_word = "рублей"

    parts.append(rub_word)
    parts.append(f"{kopecks:02d}")
    k10 = kopecks % 10
    k100 = kopecks % 100
    if k10 == 1 and k100 != 11:
        parts.append("копейка")
    elif 2 <= k10 <= 4 and not (11 <= k100 <= 14):
        parts.append("копейки")
    else:
        parts.append("копеек")

    return " ".join(parts)


def _validate_contract_fields(p, c) -> list[str]:
    """Return list of missing required fields for contract generation."""
    missing = []
    if not c:
        return ["Контрагент не указан в закупке"]
    checks = [
        (c.org_type, "Тип контрагента (org_type)"),
        (c.inn, "ИНН контрагента"),
        (c.address, "Юридический адрес контрагента"),
        (c.settlement_account, "Расчётный счёт контрагента"),
        (c.bank_name, "Наименование банка"),
        (c.bik, "БИК банка"),
        (p.contract_number, "Номер договора"),
        (p.contract_date, "Дата договора"),
        (p.subject, "Предмет договора"),
        (p.contract_price, "Цена договора"),
        (p.service_period_type, "Тип срока (period/date)"),
        (p.execution_term, "Срок оказания услуг / дата"),
    ]
    if c.org_type == "Юр.лицо":
        checks += [
            (c.kpp, "КПП контрагента"),
            (c.ogrn, "ОГРН контрагента"),
            (c.signatory, "ФИО подписанта"),
            (c.signatory_basis, "Основание полномочий подписанта"),
        ]
    elif c.org_type == "ИП":
        checks += [
            (c.ogrn, "ОГРНИП контрагента"),
        ]
    for val, label in checks:
        if not val and val != 0:
            missing.append(label)
    return missing


@router.get("/{pid}/documents/{doc_type}")
async def generate_document(
    pid: int,
    doc_type: str,
    approver_ids: Optional[str] = Query(default=None, description="ID согласующих через запятую"),
    initiator_id: Optional[int] = Query(default=None, description="ID инициатора служебной записки"),
    responsible_name: Optional[str] = Query(default=None, description="ФИО ответственного исполнителя (переопределяет поле закупки)"),
    tz_override_mode: Optional[str] = Query(default=None, description="Переопределить режим ТЗ: 'exact' или '44fz'"),
    merge: Optional[str] = Query(default=None, description="Phase 19.06: merge with another doc_type (e.g. 'tech_spec_contract') — appends its paragraphs/tables as a new section after the primary doc"),
    doc_indices: Optional[str] = Query(default=None, description="Phase 27.2: CSV индексов acceptance_docs для СЗ на оплату/аванс (напр. '0,2,3')"),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    if doc_type not in DOC_TYPES:
        raise HTTPException(400, f"Неизвестный тип документа: {doc_type}. Доступны: {', '.join(DOC_TYPES)}")

    # Резолв шаблона: субсидийный override нужен по subsidy_id закупки, которая
    # ещё не загружена — на этом шаге резолвим без него (subsidy_id=None), только
    # чтобы рано отбить полностью отсутствующий шаблон (глобальный + fallback).
    # Полный резолв (с субсидийным override) повторяется ниже, после загрузки p.
    template_path, template_file, filename_base = _resolve_doc_template_path(doc_type, None)
    if not os.path.exists(template_path):
        raise HTTPException(
            404,
            f"Шаблон '{template_file}' не найден. "
            f"Поместите файл в backend/templates/{template_file}"
        )

    # Load purchase with related data
    from app.models.contract_item import ContractItem as _ContractItem  # Phase 27.1 CD-5
    result = await db.execute(
        select(Purchase)
        .options(
            selectinload(Purchase.items).selectinload(PurchaseItem.product),
            selectinload(Purchase.contractor),
            selectinload(Purchase.feo_category),
            selectinload(Purchase.contract_items),  # Phase 27.1 CD-5: eager-load for docxtpl context
            selectinload(Purchase.assigned_user),  # B-dedup: для авто-инициалов responsible_person
            selectinload(Purchase.service_note_author),  # fallback «Ответственный исполнитель» = автор СЗ
        )
        .where(Purchase.id == pid)
    )
    p = result.scalar_one_or_none()
    if not p:
        raise HTTPException(404, "Закупка не найдена")

    # Требование владельца («Плановые не равно Договор»): для договорных
    # типов документа позиции/суммы обязаны быть заполнены в ContractItem —
    # никакого молчаливого отката на plan/НМЦК. Гейт стоит максимально рано,
    # до тяжёлых запросов ниже.
    _require_contract_items_for_doc(p, doc_type)
    _require_purchase_method_for_doc(p, doc_type)

    # Override template path with subsidy-specific template if available
    template_path, template_file, filename_base = _resolve_doc_template_path(doc_type, p.subsidy_id)

    subsidy_r = await db.execute(select(Subsidy).where(Subsidy.id == p.subsidy_id))
    subsidy = subsidy_r.scalar_one_or_none()

    # Phase 23: Customer = Organization owning the subsidy. Loads linked Contractor (FK)
    # for banking details (r/s, bank_name, BIK, k/s).
    customer_org = None
    customer_ctr = None  # Contractor linked to org via FK
    if subsidy and subsidy.org_id:
        from app.models.organization import Organization
        org_r = await db.execute(select(Organization).where(Organization.id == subsidy.org_id))
        customer_org = org_r.scalar_one_or_none()
        if customer_org and customer_org.contractor_id:
            from app.models.contractor import Contractor as _CtrCust
            ctr_r = await db.execute(select(_CtrCust).where(_CtrCust.id == customer_org.contractor_id))
            customer_ctr = ctr_r.scalar_one_or_none()

    # Load event if linked
    event = None
    if p.event_id:
        ev_r = await db.execute(select(Event).where(Event.id == p.event_id))
        event = ev_r.scalar_one_or_none()

    # Load approvers if requested
    selected_approvers = []
    if approver_ids:
        ids = [int(x.strip()) for x in approver_ids.split(",") if x.strip().isdigit()]
        if ids:
            res = await db.execute(
                select(SubsidyApprover)
                .where(SubsidyApprover.id.in_(ids))
                .order_by(SubsidyApprover.order_num, SubsidyApprover.id)
            )
            selected_approvers = res.scalars().all()

    # Load initiator if requested (frontend always sends User.id, not SubsidyApprover.id).
    # Бизнес-правило: за другого человека делать СЗ может только тот, кому
    # подчинён этот человек (видим через _get_visible_user_ids — тот же scope,
    # что для постановки задач). Самого себя — всегда можно.
    initiator = None
    if initiator_id and initiator_id != getattr(current_user, "id", None):
        from app.routers.task_visibility import _get_visible_user_ids
        visible = await _get_visible_user_ids(current_user, db)
        if visible is not None and initiator_id not in visible:
            raise HTTPException(
                status_code=403,
                detail={
                    "code": "INITIATOR_FORBIDDEN",
                    "message": (
                        "Указанный инициатор вам не подчинён. Сделать служебную "
                        "записку за другого человека может только его руководитель."
                    ),
                },
            )
    # Стратегия резолва:
    #   1) SubsidyApprover.user_id == initiator_id (привязанный approver субсидии) →
    #      берём full_name/role_name из карточки approver'а
    #   2) Иначе SimpleNamespace из User (full_name/position/department)
    if initiator_id:
        # 1. Привязанный approver субсидии (если есть)
        res = await db.execute(
            select(SubsidyApprover)
            .where(SubsidyApprover.user_id == initiator_id)
            .options(selectinload(SubsidyApprover.user))
            .order_by(SubsidyApprover.id)
            .limit(1)
        )
        initiator = res.scalar_one_or_none()
        # 2. Виртуальный approver из User
        if initiator is None:
            from app.models.user import User as UserModel
            from types import SimpleNamespace
            u = (
                await db.execute(select(UserModel).where(UserModel.id == initiator_id))
            ).scalar_one_or_none()
            if u:
                initiator = SimpleNamespace(
                    full_name=u.full_name or u.username,
                    role_name=u.position or "",
                    user=u,
                )

        # Membership-check: инициатор должен состоять в организации, к которой
        # привязана субсидия закупки. Без этого должность/отдел не определены,
        # СЗ от имени постороннего сотрудника недопустима.
        subsidy_org_id = getattr(subsidy, "org_id", None) if subsidy else None
        init_user = initiator.user if (initiator and getattr(initiator, "user", None)) else None
        if subsidy_org_id and init_user:
            from app.models.user_organization import UserOrganization
            mem = (await db.execute(
                select(UserOrganization.id).where(
                    UserOrganization.user_id == init_user.id,
                    UserOrganization.org_id == subsidy_org_id,
                ).limit(1)
            )).first()
            # Legacy fallback — User.org_id (primary)
            if not mem and getattr(init_user, "org_id", None) != subsidy_org_id:
                org_name = getattr(customer_org, "name", "") if customer_org else ""
                raise HTTPException(
                    status_code=403,
                    detail={
                        "code": "INITIATOR_NOT_IN_ORG",
                        "message": (
                            f"Инициатор не состоит в организации"
                            + (f" «{org_name}»" if org_name else "")
                            + ", к которой привязана субсидия закупки. Выберите "
                            "сотрудника этой организации."
                        ),
                    },
                )

    # If no approvers specified — load defaults for this subsidy
    if not selected_approvers and p.subsidy_id:
        res = await db.execute(
            select(SubsidyApprover)
            .where(SubsidyApprover.subsidy_id == p.subsidy_id, SubsidyApprover.is_default == True)
            .order_by(SubsidyApprover.order_num, SubsidyApprover.id)
        )
        selected_approvers = res.scalars().all()

    # Build FEO category path (root → ... → selected) + individual levels
    feo_path = ""
    feo_level_1 = ""
    feo_level_2 = ""
    feo_level_3 = ""
    item_feo_paths: list = []  # [(path_str, Decimal total)] по позициям со своей категорией ФЭО

    # Требование владельца (2026-08-30): «В "Путь ФЭО" должны прописываться
    # договорные категории, они должны совпадать с Плановыми». Лист
    # согласования и остальные CONTRACT_FAMILY_DOC_TYPES визируют ДОГОВОРНЫЕ
    # позиции (ContractItem) — цены там уже договорные, поэтому путь ФЭО
    # тоже обязан идти от договорных позиций, иначе строки «цена / категория»
    # расходятся. См. _build_contract_item_feo_paths().
    _is_contract_family = doc_type in CONTRACT_FAMILY_DOC_TYPES
    if _is_contract_family:
        _plan_feo_by_id = {it.id: it.feo_category_id for it in (p.items or []) if it.feo_category_id}
        item_feo_ids = list(dict.fromkeys(
            _plan_feo_by_id[ci.source_item_id]
            for ci in (p.contract_items or [])
            if ci.source_item_id and ci.source_item_id in _plan_feo_by_id
        ))
    else:
        item_feo_ids = list(dict.fromkeys(it.feo_category_id for it in (p.items or []) if it.feo_category_id))

    if p.feo_category_id or item_feo_ids or (_is_contract_family and (p.contract_items or [])):
        feo_res = await db.execute(select(FeoCategory))
        all_feo = {f.id: f for f in feo_res.scalars().all()}

        def _feo_path_nodes(node_id):
            path_nodes: list = []
            visited: set = set()
            while node_id and node_id not in visited:
                visited.add(node_id)
                cat = all_feo.get(node_id)
                if not cat:
                    break
                path_nodes.append(cat)
                node_id = cat.parent_id
            path_nodes.reverse()  # root → leaf
            return path_nodes

        if p.feo_category_id:
            path_nodes = _feo_path_nodes(p.feo_category_id)
            feo_path = " → ".join(n.name.strip() for n in path_nodes)
            if len(path_nodes) >= 1: feo_level_1 = path_nodes[0].name.strip()
            if len(path_nodes) >= 2: feo_level_2 = path_nodes[1].name.strip()
            if len(path_nodes) >= 3: feo_level_3 = path_nodes[2].name.strip()

        if _is_contract_family:
            item_feo_paths = _build_contract_item_feo_paths(
                p.contract_items or [], p.items or [], _feo_path_nodes,
            )
        elif item_feo_ids:
            cat_sums: dict = {cid: Decimal("0") for cid in item_feo_ids}
            for it in (p.items or []):
                if it.feo_category_id and it.total_price is not None:
                    cat_sums[it.feo_category_id] += Decimal(str(it.total_price))
            for cid in item_feo_ids:
                cat_path = " → ".join(n.name.strip() for n in _feo_path_nodes(cid))
                if cat_path:
                    item_feo_paths.append((cat_path, cat_sums[cid]))

    # B-dedup: формат ФИО → инициалы.
    # "Иванова Ирина Владиславовна"   → "Иванова И.В."
    # "Кулиев Гасан Валех оглы"        → "Кулиев Г.В." (4-я часть отбрасывается)
    # "Иванова-Петрова Анна Сергеевна" → "Иванова-Петрова А.С." (дефис = одна фамилия)
    def _format_initials(full: str) -> str:
        if not full:
            return ""
        parts = (full or "").strip().split()
        if not parts:
            return ""
        surname = parts[0]
        initials = []
        for p_word in parts[1:3]:  # только имя + отчество (3-я часть «оглы»/«кызы» отбрасывается)
            if p_word and p_word[0].isalpha():
                initials.append(p_word[0].upper() + ".")
        return f"{surname} {''.join(initials)}".strip()

    # Resolved responsible person: priority = ?responsible_name → assigned_user.full_name
    # → p.responsible_person → автор служебной записки (последний фолбэк, чтобы клетка
    # «Ответственный исполнитель» никогда не оставалась привязанной к фиксированному
    # человеку из настроек субсидии — см. app/services/responsible_role.py)
    assigned_full = (getattr(p.assigned_user, "full_name", None) or "") if getattr(p, "assigned_user", None) else ""
    service_note_author_full = (
        (getattr(p.service_note_author, "full_name", None) or "")
        if getattr(p, "service_note_author", None) else ""
    )
    raw_responsible = responsible_name or assigned_full or p.responsible_person or service_note_author_full or ""
    resolved_responsible = _format_initials(raw_responsible) if raw_responsible else ""
    resolved_responsible_full = raw_responsible  # для шаблонов которым нужно полное ФИО

    # Build docxtpl template object early (needed for InlineImage)
    try:
        from docxtpl import DocxTemplate, InlineImage
        from docx.shared import Cm as _Cm

        # Templates are normalized at upload time (subsidies.py
        # _normalize_docx_template). For files uploaded BEFORE that fix
        # landed we lazily rewrite them on first render: scan + replace +
        # atomic rename, so subsequent renders hit the fast path with no
        # extra IO and InlineImage placeholders sit in a single contiguous
        # run (required for the drawing element to be inserted).
        try:
            import zipfile as _zf
            needs_norm = False
            with _zf.ZipFile(template_path, 'r') as _zin:
                for _name in _zin.namelist():
                    if _name == 'word/document.xml' or _name.startswith('word/header') or _name.startswith('word/footer'):
                        _bytes = _zin.read(_name)
                        if b'<w:proofErr' in _bytes or b'<w:bookmark' in _bytes or b'<w:commentRange' in _bytes or b'<w:lastRenderedPageBreak' in _bytes:
                            needs_norm = True
                            break
            if needs_norm:
                from app.routers.subsidies import _normalize_docx_template as _norm
                _stats = _norm(template_path)
                logger.info("lazy normalize on render %s: %s", template_path, _stats)
        except Exception as _scan_e:
            logger.warning("lazy normalize scan failed for %s: %s", template_path, _scan_e)

        tpl = DocxTemplate(template_path)
    except HTTPException:
        raise
    except Exception as e:
        import traceback as _tb
        logger.exception("Document template load failed for purchase %s, doc_type=%s", pid, doc_type)
        raise HTTPException(500, detail={
            "code": "DOCUMENT_GENERATION_FAILED",
            "message": f"Не удалось загрузить шаблон «{doc_type}»: {type(e).__name__}",
            "doc_type": doc_type,
            "purchase_id": pid,
            "error_class": type(e).__name__,
            "error_raw": str(e),
            "traceback": "".join(_tb.format_exception(type(e), e, e.__traceback__))[:4000],
            "hint": (
                "Ошибка при загрузке файла шаблона docxtpl. "
                "Возможные причины: повреждённый .docx файл шаблона, "
                "отсутствие библиотеки docxtpl/python-docx. "
                "Передайте администратору error_class + traceback."
            ),
        })

    UPLOADS_DIR = "/app/uploads/products"

    def _resolve_photo(photo_url):
        """Return InlineImage or empty string."""
        import tempfile, urllib.request as _ur
        if not photo_url:
            return ""
        url = str(photo_url).strip()
        local_path = None

        if url.startswith("/api/products/photos/"):
            fname = url.split("/")[-1]
            local_path = f"{UPLOADS_DIR}/{fname}"
        elif url.isdigit():
            for ext in ("jpg", "jpeg", "png"):
                p = f"{UPLOADS_DIR}/product_{url}.{ext}"
                if os.path.exists(p):
                    local_path = p
                    break
        elif url.startswith("http://") or url.startswith("https://"):
            # Download external image; convert webp via Pillow if available
            try:
                with _ur.urlopen(url, timeout=5) as r:
                    ct = r.headers.get("Content-Type", "").split(";")[0].strip().lower()
                    raw = r.read()
                is_webp = "webp" in ct or url.lower().endswith(".webp")
                if is_webp:
                    try:
                        from PIL import Image as _Img
                        import io as _io
                        img = _Img.open(_io.BytesIO(raw)).convert("RGB")
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
                        img.save(tmp, format="JPEG", quality=85)
                        tmp.close()
                        local_path = tmp.name
                    except Exception:
                        return ""  # Pillow not available or conversion failed
                else:
                    ext_map = {"image/jpeg": ".jpg", "image/jpg": ".jpg", "image/png": ".png",
                               "image/gif": ".gif", "image/bmp": ".bmp", "image/tiff": ".tiff"}
                    suffix = ext_map.get(ct, ".jpg")
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                    tmp.write(raw)
                    tmp.close()
                    local_path = tmp.name
            except Exception:
                return ""

        if not local_path or not os.path.exists(local_path):
            return ""
        try:
            return InlineImage(tpl, local_path, width=_Cm(2.5))
        except Exception:
            return ""

    # Build template context
    # «Плановые не равно Договор»: для CONTRACT_FAMILY_DOC_TYPES позиции
    # берутся ТОЛЬКО из ContractItem (гейт _require_contract_items_for_doc
    # выше уже гарантировал, что contract_items не пуст). Для остальных —
    # старое поведение (purchase_items).
    if doc_type in CONTRACT_FAMILY_DOC_TYPES:
        items_list = _build_items_list_from_contract_items(p, resolve_photo=_resolve_photo)
    else:
        items_list = _build_items_list_from_purchase_items(
            p, tz_override_mode=tz_override_mode, resolve_photo=_resolve_photo,
        )

    # Phase 23.1: auto-detect subject_kind for universal contract.docx
    # 'services' if ALL items have item_kind='услуга', otherwise 'goods' (default)
    subject_kind = "goods"
    if items_list:
        kinds = {it.get("item_kind", "товар").lower() for it in items_list}
        if kinds == {"услуга"}:
            subject_kind = "services"

    # Load existing PurchaseApproval records (electronic signatures)
    from app.models.purchase_approval import PurchaseApproval
    pa_res = await db.execute(
        select(PurchaseApproval)
        .where(PurchaseApproval.purchase_id == pid, PurchaseApproval.status == "approved")
        .order_by(PurchaseApproval.order_num)
    )
    # Map subsidy_approver_id → PurchaseApproval (for signature lookup)
    approval_map: dict[int, PurchaseApproval] = {}
    for pa in pa_res.scalars().all():
        if pa.subsidy_approver_id:
            approval_map[pa.subsidy_approver_id] = pa

    def _base64_to_inline(tpl_obj, b64_data: str):
        """Convert base64 PNG data URL to docxtpl InlineImage."""
        import tempfile, base64, re as _re
        try:
            from docxtpl import InlineImage
            from docx.shared import Cm as _Cm
            m = _re.match(r"data:image/\w+;base64,(.+)", b64_data, _re.DOTALL)
            if not m:
                return ""
            raw = base64.b64decode(m.group(1))
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            tmp.write(raw)
            tmp.close()
            return InlineImage(tpl_obj, tmp.name, width=_Cm(3.0))
        except Exception:
            return ""

    # Collect unique product categories from items
    item_categories = list(dict.fromkeys(
        item.product.category
        for item in (p.items or [])
        if item.product and item.product.category
    ))
    item_categories_str = ", ".join(item_categories)

    approvers_list = []
    for i, a in enumerate(selected_approvers):
        full_name = a.full_name or ""
        # «Ответственный исполнитель» — роль-слот (app/services/responsible_role.py):
        # хранимое в subsidy_approvers ФИО для неё ИГНОРИРУЕТСЯ ВСЕГДА, даже если там
        # почему-то оказалось живое имя — источник истины только резолв по закупке.
        # Для остальных ролей подставляем резолв только если сохранённое ФИО пустое/плейсхолдер.
        if is_responsible_role(a.role_name) or is_blank_person_name(full_name):
            full_name = resolved_responsible or RESPONSIBLE_PLACEHOLDER
        if getattr(a, "show_feo_path", False) and item_feo_paths:
            note = "; ".join(f"{path} — {_fmt_money(total)} ₽" for path, total in item_feo_paths)
        elif getattr(a, "show_feo_path", False) and feo_path:
            item_type_label = {"товар": "Товары", "услуга": "Услуги"}.get(p.item_type or "", "")
            note = feo_path + (f" ({item_type_label})" if item_type_label else "")
        else:
            note = ""

        # Electronic signature
        pa = approval_map.get(a.id)
        signature_img = ""
        decided_date = ""
        if pa and pa.signature_data and pa.signature_algorithm == "visual":
            signature_img = _base64_to_inline(tpl, pa.signature_data)
            if pa.decided_at:
                decided_date = pa.decided_at.strftime("%d.%m.%Y")

        approvers_list.append({
            "num": i + 1,
            "role_name": a.role_name,
            "full_name": full_name,
            "signature_img": signature_img,
            "decided_date": decided_date,
            "note": note,
            # Phase 26-V: родительный падеж
            "full_name_gen": _to_gen_fio(full_name),
            "role_name_gen": _to_gen_phrase(a.role_name or ""),
        })

    c = p.contractor

    # ── Contract-specific context helpers ────────────────────────────────────
    def _contract_date_parts():
        d = p.contract_date
        if not d:
            return "", "", ""
        if isinstance(d, str):
            try:
                d = date.fromisoformat(d)
            except ValueError:
                return "", "", ""
        months_ru = ["января", "февраля", "марта", "апреля", "мая", "июня",
                     "июля", "августа", "сентября", "октября", "ноября", "декабря"]
        return str(d.day).zfill(2), months_ru[d.month - 1], str(d.year)

    cd_day, cd_month, cd_year = _contract_date_parts()

    # Short name: extract parenthetical or use first word
    def _short_name(full_name: str) -> str:
        import re
        m = re.search(r'[«""]([^»""]+)[»""]', full_name)
        if m:
            return m.group(1)
        parts = full_name.split()
        return parts[-1] if parts else full_name

    # Сумма закупки для документов.
    # items_sum_val — сумма ПЛАНОВЫХ позиций, нужна только как последний
    # фолбэк для total_nmcd/total_nmck/nmck ниже (они остаются плановыми
    # полями всегда, независимо от doc_type — НМЦК по определению начальная
    # плановая цена).
    # doc_amount_val/amount_is_planned — сумма ДОКУМЕНТА: для
    # CONTRACT_FAMILY_DOC_TYPES («Плановые не равно Договор») — ТОЛЬКО
    # p.contract_price / сумма ContractItem, без отката на НМЦК/план. Для
    # остальных типов — старое поведение (план/НМЦК до заключения договора).
    items_sum_val = float(sum(Decimal(str(it.total_price or 0)) for it in (p.items or [])))
    doc_amount_val, amount_is_planned = _resolve_doc_amount(p, doc_type)

    # VAT calculations
    vat_app = bool(p.vat_applicable)
    vat_rate_val = p.vat_rate or 20
    price_val = doc_amount_val
    if vat_app and price_val:
        vat_amount_val = price_val * vat_rate_val / (100 + vat_rate_val)
    else:
        vat_amount_val = 0.0

    # НДС info for approval sheet
    # Phase 26-TT: для авансового отчёта не придумывать «НДС не облагается» если
    # в чеке/items нет данных VAT — пишем только то, что реально есть.
    is_advance = (p.purchase_method == 'advance')
    items_with_vat = [it for it in (p.items or []) if getattr(it, 'vat_rate', None)]

    if vat_app:
        vat_info_line = f"В том числе НДС {vat_rate_val}%: {_fmt_money_plain(vat_amount_val)} руб."
    elif is_advance:
        # Авансовый: данные из чеков ФНС; если в чеке нет НДС — не пишем ничего лишнего.
        if items_with_vat:
            # Per-item VAT — показать сводку по факту
            unique_rates = sorted({str(it.vat_rate) for it in items_with_vat if it.vat_rate})
            vat_info_line = f"НДС по позициям: {', '.join(unique_rates)}"
        else:
            vat_info_line = ""  # пусто — не придумываем
    else:
        art = (p.vat_exemption_article or "").strip()
        vat_info_line = f"НДС не облагается" + (f" ({art})" if art else "")

    context = {
        # Закупка
        "purchase_number": p.purchase_number or "",
        "registry_number": p.registry_number or "",
        "purchase_method": _PURCHASE_METHOD_LABELS.get(p.purchase_method or "", p.purchase_method or ""),
        # order_purchase (приказ на закупку), п.2: способ закупки прописью.
        # Неизвестный код способа закупки — подставляем сам код, не пустую строку.
        # Конкурентная процедура с заполненной формой (competitive_form) —
        # подпись конкретной формы («Запрос цен» / «Аукцион (редукцион)» /
        # «Конкурс»), не общее «Конкурсная процедура» (см. _purchase_method_label).
        "purchase_method_label": _purchase_method_label(p),
        "subject": p.subject or "",
        "status": p.status or "",
        "purchase_basis": _BASIS_LABELS.get(p.purchase_basis or "", ""),
        "responsible_person": resolved_responsible,  # инициалы (по умолчанию)
        "responsible_person_full": resolved_responsible_full,  # полное ФИО (для шаблонов где надо)
        # Субсидия
        "subsidy_name": subsidy.name if subsidy else "",
        "subsidy_year": subsidy.year if subsidy else "",
        "subsidy_budget": _fmt_money(subsidy.budget) if subsidy else "",
        # Контрагент — основные
        "contractor_name": (c.name or "") if c else "",
        "contractor_inn": _clean_id(c.inn) if c else "",
        "contractor_kpp": _clean_id(c.kpp) if c else "",
        "contractor_address": (c.address or "") if c else "",
        "contractor_postal_address": (c.postal_address or "") if c else "",
        "contractor_ogrn": (c.ogrn or "") if c else "",
        "contractor_phone": (c.phone or "") if c else "",
        "contractor_email": (c.email or "") if c else "",
        # Контрагент — подписант
        "contractor_signatory": (c.signatory or "") if c else "",
        "contractor_signatory_basis": (c.signatory_basis or "") if c else "",
        # Контрагент — банк
        "contractor_settlement_account": (c.settlement_account or "") if c else "",
        "contractor_bank_name": (c.bank_name or "") if c else "",
        "contractor_bik": (c.bik or "") if c else "",
        "contractor_correspondent_account": (c.correspondent_account or "") if c else "",
        # Комбинированные поля контрагента для шаблонов
        "contractor_bank_details": (c.bank_name or "") if c else "",
        "contractor_signatory_line": (
            f"{c.signatory}, действует на основании {c.signatory_basis}"
            if c and c.signatory and c.signatory_basis
            else ((c.signatory or "") if c else "")
        ),
        # FEO
        "feo_category_name": p.feo_category.name if p.feo_category else "",
        "feo_path":    feo_path,
        "feo_level_1": feo_level_1,
        "feo_level_2": feo_level_2,
        "feo_level_3": feo_level_3,
        # Финансы
        # Phase 19: total_nmcd is the new canonical name (НМЦД — начальная
        # максимальная цена договора). total_nmck is kept as a deprecated
        # alias so existing templates keep rendering. Use total_nmcd in new
        # templates.
        # Phase: суммы в шапке документов не должны зависеть от стадии закупки —
        # фолбэк на doc_amount_val/items_sum_val, см. расчёт выше (перед НДС).
        "total_nmcd": _fmt_money(p.total_nmck or p.nmck or p.planned_total_price or items_sum_val),
        "total_nmck": _fmt_money(p.total_nmck or p.nmck or p.planned_total_price or items_sum_val),
        "nmck": _fmt_money(p.nmck or p.total_nmck or items_sum_val),
        "contract_price": _fmt_money(p.contract_price or doc_amount_val),
        "economy": _fmt_money(p.economy),
        "price_increase": _fmt_money(p.price_increase),
        # Договор
        "contract_number": p.contract_number or "",
        "contract_date": _fmt_date(p.contract_date) or "__.__._____ г.",
        "execution_term": _fmt_date(p.execution_term),
        "execution_term_changed": _fmt_date(p.execution_term_changed),
        "delivery_date": _fmt_date(p.delivery_date),
        "country_origin": p.country_origin or "",
        # Акт приёмки. Phase 26-lll: УБРАН fallback acceptance_doc_* → contract_*.
        # Phase 27.2-01: для СЗ на оплату/аванс читаем acceptance_docs JSONB
        # (source of truth после Phase 26-H). Legacy plain-поля — только fallback.
        **_build_acceptance_doc_context(p, doc_type, doc_indices),
        # Платёж
        "payment_doc_number": p.payment_doc_number or "",
        "payment_doc_date": _fmt_date(p.payment_doc_date),
        "payment_amount": _fmt_money(p.payment_amount),
        "payment_federal": _fmt_money(p.payment_federal),
        # Позиции
        "items": items_list,
        "items_count": len(items_list),
        "item_names": p.subject or ", ".join(i["name"] for i in items_list if i["name"]),
        "item_categories": item_categories_str,
        # Согласующие
        "approvers": approvers_list,
        # Инициатор: ФИО берётся как есть; должность и отдел резолвятся per-org
        # — по организации, к которой привязана субсидия закупки. Если у юзера
        # несколько отделов в этой org — первый.
        "initiator_name": initiator.full_name if initiator else "",
        "initiator_role": (
            await _resolve_user_position(
                initiator.user if (initiator and getattr(initiator, "user", None)) else None,
                db,
                getattr(subsidy, "org_id", None) if subsidy else None,
            )
            or (initiator.role_name if initiator else "")
        ),
        "initiator_dept": await _resolve_user_dept(
            initiator.user if (initiator and getattr(initiator, "user", None)) else None,
            db,
            getattr(subsidy, "org_id", None) if subsidy else None,
        ),
        # Мероприятие
        "event_name": event.name if event else "",
        # Тип договора
        "contract_type": {"single": "Единственный поставщик", "framework_cumulative": "Рамочный (накопительный)", "framework_with_amount": "Рамочный (с суммой)"}.get(p.purchase_contract_type or "", p.purchase_contract_type or ""),
        # Служебные
        "today": _fmt_date(date.today()),
        "today_iso": date.today().isoformat(),
        # ── Расширенные поля для договорных шаблонов ─────────────────────────
        # Дата по частям
        "contract_date_day":   cd_day,
        "contract_date_month": cd_month,
        "contract_date_year":  cd_year,
        # Тип контрагента
        "contractor_org_type": (c.org_type or "") if c else "",
        # Phase 27.2-08: краткое название = поле "Краткое наименование *" из карточки контрагента
        # напрямую (Contractor.name), без вытаскивания из кавычек.
        "contractor_short_name": (c.name or "") if c else "",
        "contractor_signatory_position": _signatory_position(c.signatory) if c else "",
        # Предмет (сервисное имя)
        "service_name": p.subject or "",
        "service_name_gen": _inflect_phrase_genitive(p.subject or ""),
        # Срок оказания услуг
        # Phase 19: service_start_date / service_end_date now prefer the real
        # Purchase columns (used by 'range' mode). If those are empty we fall
        # back to the legacy mapping (contract_date / execution_term).
        "period_type": p.service_period_type or "period",
        "service_start_date": _fmt_date(p.service_start_date) or _fmt_date(p.contract_date),
        "service_end_date":   _fmt_date(p.service_end_date)   or _fmt_date(p.execution_term),
        "service_date":       _fmt_date(p.execution_term),
        # Phase 19: extended service-term context
        "service_term":            _format_service_term(p),
        "service_term_mode":       p.service_term_mode or "",
        "service_term_days":       p.service_term_days or "",
        "service_term_type":       p.service_term_type or "",
        "service_term_type_name":  {"calendar": "календарных", "working": "рабочих"}.get(p.service_term_type or "", ""),
        "service_deadline_date":   _fmt_date(p.service_deadline_date),
        # Дата окончания срока действия договора (п. 8.1) — поле Purchase.contract_end_date
        "contract_end_date":       _fmt_date(p.contract_end_date),
        # Phase 19: submission deadline (дата+время завершения приёма заявок)
        "submission_deadline_date":     p.submission_deadline.date().isoformat() if p.submission_deadline else "",
        "submission_deadline_time":     p.submission_deadline.strftime("%H:%M") if p.submission_deadline else "",
        "submission_deadline_datetime": p.submission_deadline.strftime("%d.%m.%Y %H:%M") if p.submission_deadline else "",
        # Phase 19: delivery location
        "delivery_location": p.delivery_location or "",
        # Phase 19: agreement text from subsidy
        "subsidy_agreement_text": (subsidy.agreement_text if (subsidy and subsidy.agreement_text) else ""),
        # Третьи лица
        "third_party_involved": bool(p.third_party_involved),
        # НДС
        "vat_applicable":        vat_app,
        "vat_rate":              vat_rate_val,
        "vat_amount_num":        _fmt_money_plain(vat_amount_val),
        "vat_amount_words":      _rubles_to_words(vat_amount_val),
        "vat_exemption_article": p.vat_exemption_article or "",
        "vat_info_line":         vat_info_line,
        # Цена прописью (фолбэк на doc_amount_val, если договор ещё не заключён)
        "contract_price_num":   _fmt_money_plain(p.contract_price or doc_amount_val),
        "contract_price_words": _rubles_to_words(p.contract_price or doc_amount_val),
        # Phase: сумма закупки для документов, не зависящая от стадии (план/договор).
        "doc_amount":         _fmt_money(doc_amount_val),
        "doc_amount_num":     _fmt_money_plain(doc_amount_val),
        "doc_amount_words":   _rubles_to_words(doc_amount_val),
        "amount_is_planned":  amount_is_planned,
        "amount_source_label": "НМЦК (план)" if amount_is_planned else "Цена договора",
        # Phase 23: service_subject alias (same as subject but clearer name in services template)
        "service_subject": p.subject or "",
        # Phase 23.1: subject_kind for universal contract.docx auto-switch
        "subject_kind": subject_kind,
    }

    # Phase 28: расширенные ключи для типовых договоров ─────────────────────────
    # Все ключи NULL-safe (пустая строка / ноль / пустой список при отсутствии данных).
    try:
        # Реквизиты субсидии-грантодателя (для рамочных и региональных договоров)
        # Phase 28: grantor_name / ministry_name берём из новых полей модели Subsidy
        context["subsidy_grantor_name"] = (subsidy.grantor_name or "").strip() if subsidy else ""
        context["subsidy_ministry_name"] = (subsidy.ministry_name or "").strip() if subsidy else ""
        # agreement_date: новое поле отсутствует в модели — используем basis_doc_date
        _subsidy_agreement_date_obj = (
            subsidy.basis_doc_date if subsidy and subsidy.basis_doc_date else None
        )
        context["subsidy_agreement_date"] = _fmt_date(_subsidy_agreement_date_obj) if _subsidy_agreement_date_obj else ""

        # Реквизиты физ.лица (ГПХ) — читаем из контрагента
        context["contractor_passport_series"]       = (c.passport_series or '') if c else ''
        context["contractor_passport_number"]       = (c.passport_number or '') if c else ''
        context["contractor_passport_issuer"]       = (c.passport_issuer or '') if c else ''
        context["contractor_passport_issued_date"]  = _fmt_date(c.passport_issued_date) if c and c.passport_issued_date else ''
        context["contractor_snils"]                 = (c.snils or '') if c else ''
        context["contractor_registration_address"]  = (c.registration_address or c.address or c.postal_address or '') if c else ''
        context["contractor_birth_date"]            = _fmt_date(c.birth_date) if c and c.birth_date else ''

        # Комиссия закупки (для протокола) — из полей закупки Phase 28
        context["commission_member_1_name"] = (p.commission_member_1_name or "").strip()
        context["commission_member_2_name"] = (p.commission_member_2_name or "").strip()
        context["commission_member_3_name"] = (p.commission_member_3_name or "").strip()
        _commission_members = [
            {"name": p.commission_member_1_name, "role": "Член комиссии"} if p.commission_member_1_name else None,
            {"name": p.commission_member_2_name, "role": "Член комиссии"} if p.commission_member_2_name else None,
            {"name": p.commission_member_3_name, "role": "Член комиссии"} if p.commission_member_3_name else None,
        ]
        context["commission_members"] = [x for x in _commission_members if x]

        # Прочие условия договора — из полей закупки Phase 28
        context["advance_amount"]              = _fmt_money(p.advance_amount) if p.advance_amount else ""
        context["acceptance_term_days"]        = p.acceptance_term_days if p.acceptance_term_days is not None else 5
        context["penalty_rate"]                = str(p.penalty_rate) if p.penalty_rate is not None else "0.1"
        context["procurement_protocol_number"] = (p.procurement_protocol_number or "").strip()
        context["procurement_order_number"]    = (p.procurement_order_number or "").strip()
        context["repair_request_number"]       = (p.repair_request_number or "").strip()
        context["contractor_ogrnip_date"]      = _fmt_date(p.contractor_ogrnip_date) if p.contractor_ogrnip_date else ""
        # Phase 28: гарантия + ретроактивный договор (комментарии пользователя 2026-05-19)
        context["warranty_period_days"]        = p.warranty_period_days if p.warranty_period_days is not None else 15
        context["is_retroactive"]              = bool(p.is_retroactive)
        # delivery_by_supplier=True — поставщик доставляет; False — самовывоз.
        # has_stages=True — в Приложении №1 есть этапы оказания услуг.
        context["delivery_by_supplier"] = bool(getattr(p, "delivery_by_supplier", True))
        context["has_stages"]           = bool(getattr(p, "has_stages", False))
        # Phase 28: subsidy-specific clauses (пункты из субсидии — раздельный учёт и т.п.)
        context["subsidy_extra_clause_1"]      = (subsidy.extra_contract_clause_1 or '').strip() if subsidy else ''
        context["subsidy_extra_clause_2"]      = (subsidy.extra_contract_clause_2 or '').strip() if subsidy else ''
    except Exception as _e28:
        import logging as _log28
        _log28.getLogger(__name__).warning("Phase 28 context keys failed: %s", _e28)

    # Phase 26-V: родительный падеж для инициатора и ответственного
    context["initiator_name_gen"] = _to_gen_fio(context.get("initiator_name", ""))
    context["initiator_position_gen"] = _to_gen_phrase(context.get("initiator_role", ""))
    context["responsible_name_gen"] = _to_gen_fio(resolved_responsible)
    context["responsible_position_gen"] = _to_gen_phrase(p.responsible_position or "" if hasattr(p, "responsible_position") else "")

    # Phase 27.2-10: дедуп тавтологии "отдел Отдела ..." + gent для отдела
    # Если в должности уже есть корень "отдел*" — убираем leading "Отдел*" из имени отдела.
    # Пример: «Заместитель начальника отдела» + «Отдела МТО» → «МТО».
    import re as _re_dept
    def _strip_redundant_dept_word(position: str, dept: str) -> str:
        if not dept or not position:
            return dept or ""
        if not _re_dept.search(r'\bотдел\w*', position.lower()):
            return dept
        return _re_dept.sub(r'^отдел\w*\s+', '', dept, count=1, flags=_re_dept.IGNORECASE)

    _init_pos = context.get("initiator_role", "") or ""
    _init_dept_raw = context.get("initiator_dept", "") or ""
    _init_dept_clean = _strip_redundant_dept_word(_init_pos, _init_dept_raw)
    context["initiator_dept"] = _init_dept_clean
    context["initiator_dept_gen"] = _inflect_phrase_genitive(_init_dept_clean)

    # ── Phase 23: Заказчик (Customer = Organization владелец субсидии + linked Contractor) ──
    def _g(*sources, default=""):
        """Coalesce — return first non-empty value."""
        for s in sources:
            if s:
                return s
        return default

    cust_signatory_full = _g(
        customer_org.signatory if customer_org else None,
        customer_ctr.signatory if customer_ctr else None,
    )
    # Структурированные части: приоритет org, затем ctr
    _cust_last = _g(
        getattr(customer_org, 'signatory_last_name', None) if customer_org else None,
        getattr(customer_ctr, 'signatory_last_name', None) if customer_ctr else None,
    ) or None
    _cust_first = _g(
        getattr(customer_org, 'signatory_first_name', None) if customer_org else None,
        getattr(customer_ctr, 'signatory_first_name', None) if customer_ctr else None,
    ) or None
    _cust_middle = _g(
        getattr(customer_org, 'signatory_middle_name', None) if customer_org else None,
        getattr(customer_ctr, 'signatory_middle_name', None) if customer_ctr else None,
    ) or None
    _cust_position = _g(
        getattr(customer_org, 'signatory_position', None) if customer_org else None,
        getattr(customer_ctr, 'signatory_position', None) if customer_ctr else None,
    ) or None
    cust_sig = _signatory_split(
        cust_signatory_full,
        last=_cust_last, first=_cust_first, middle=_cust_middle, position=_cust_position,
    )
    cust_signatory_basis = _g(
        customer_ctr.signatory_basis if customer_ctr else None,
        "Устава",
    )

    context.update({
        "customer_name":         _g(customer_org.name if customer_org else None,
                                    customer_ctr.name if customer_ctr else None),
        "customer_full_name":    _g(customer_org.full_name if customer_org else None,
                                    customer_ctr.name if customer_ctr else None,
                                    customer_org.name if customer_org else None),
        # Phase 27.2-08: краткое название Заказчика = поле "Краткое наименование" из карточки
        # организации/контрагента напрямую, без вытаскивания из кавычек.
        "customer_short_name":   _g(customer_org.name if customer_org else None,
                                    customer_ctr.name if customer_ctr else None),
        "customer_address":      _g(customer_org.address if customer_org else None,
                                    customer_ctr.address if customer_ctr else None),
        "customer_postal_address": _g(customer_ctr.postal_address if customer_ctr else None,
                                      customer_org.address if customer_org else None),
        "customer_inn":          _clean_id(_g(customer_org.inn if customer_org else None,
                                              customer_ctr.inn if customer_ctr else None)),
        "customer_kpp":          _clean_id(_g(customer_org.kpp if customer_org else None,
                                              customer_ctr.kpp if customer_ctr else None)),
        "customer_ogrn":         _g(customer_org.ogrn if customer_org else None,
                                    customer_ctr.ogrn if customer_ctr else None),
        "customer_bank_name":    _g(customer_ctr.bank_name if customer_ctr else None),
        "customer_settlement_account":    _g(customer_ctr.settlement_account if customer_ctr else None),
        "customer_correspondent_account": _g(customer_ctr.correspondent_account if customer_ctr else None),
        "customer_bik":          _g(customer_ctr.bik if customer_ctr else None),
        # Лицевой счёт Заказчика
        "customer_personal_account": _g(customer_ctr.personal_account if customer_ctr else None),
        "customer_phone":        _g(customer_ctr.phone if customer_ctr else None),
        "customer_email":        _g(customer_ctr.email if customer_ctr else None),
        # Подписант Заказчика
        "customer_signatory":                cust_signatory_full,
        "customer_signatory_position":       cust_sig["position"],
        "customer_signatory_name":           cust_sig["name_full"],
        "customer_signatory_name_genitive":  cust_sig["name_genitive"],
        "customer_signatory_initials":       cust_sig["name_initials"],
        # Инициалы впереди: «Е.В. Козеев» (для строки подписи _________________ / И.О. Фамилия)
        "customer_signatory_name_initials":  _fio_to_initials_prefix(cust_sig["name_full"]),
        "customer_signatory_basis":          cust_signatory_basis,
        # Город заключения — из поля Organization.contract_city; fallback «Москва»
        "contract_city": (customer_org.contract_city or "Москва") if customer_org else "Москва",
    })

    # Fabrikant context keys ────────────────────────────────────────────────────
    # notice_number — номер процедуры на ЭТП (platform_number), а не наш internal
    # requestId (external_id); откат на external_id, если platform_number ещё
    # не проставлен площадкой на момент рендера
    _notice_number = ""
    try:
        from app.models.platform_publication import PlatformPublication as _PlatPub
        _pub_q = await db.execute(
            select(_PlatPub.platform_number, _PlatPub.external_id)
            .where(
                _PlatPub.purchase_id == p.id,
                _PlatPub.platform == "fabrikant",
            )
            .order_by(_PlatPub.id.desc())
            .limit(1)
        )
        _pub_row = _pub_q.first()
        if _pub_row:
            _notice_number = _pub_row[0] or _pub_row[1] or ""
    except Exception:
        pass
    context["notice_number"] = _notice_number

    # subsidy_agreement_number — из Subsidy.agreement_number
    context["subsidy_agreement_number"] = (subsidy.agreement_number or "") if subsidy else ""

    # payment_term_days — из Purchase.payment_term_days; fallback 10
    context["payment_term_days"] = p.payment_term_days if (p.payment_term_days is not None) else 10

    # applications_review_date — из Purchase.applications_review_date;
    # fallback: submission_deadline + 1 рабочий день (пн-пт, без учёта праздников)
    if p.applications_review_date:
        context["applications_review_date"] = _fmt_date(p.applications_review_date)
    elif p.submission_deadline:
        from datetime import timedelta as _td
        _next = p.submission_deadline.date() + _td(days=1)
        while _next.weekday() >= 5:  # 5=сб, 6=вс
            _next += _td(days=1)
        context["applications_review_date"] = _next.strftime("%d.%m.%Y")
    else:
        context["applications_review_date"] = ""

    # Phase 23: расширенные поля подписанта Исполнителя (name_genitive, initials, ogrnip)
    ctr_sig = _signatory_split(
        c.signatory if c else "",
        last=getattr(c, 'signatory_last_name', None) if c else None,
        first=getattr(c, 'signatory_first_name', None) if c else None,
        middle=getattr(c, 'signatory_middle_name', None) if c else None,
        position=getattr(c, 'signatory_position', None) if c else None,
    )
    context.update({
        "contractor_signatory_name":          ctr_sig["name_full"],
        "contractor_signatory_name_genitive": ctr_sig["name_genitive"],
        "contractor_signatory_initials":      ctr_sig["name_initials"],
        "contractor_ogrnip":                  (c.ogrn or "") if (c and (c.org_type or "").lower().startswith("ип")) else "",
        # contractor_full_name — полное официальное название (fallback на name)
        "contractor_full_name":               (c.full_name or c.name or "") if c else "",
    })

    # Phase 27.1 CD-5: contract_items loop context (+ D-08 fallback on purchase_items)
    # Phase 26-U: wrap pre-render context building — any exception → structured DOCUMENT_GENERATION_FAILED
    receipt_png_paths: list[str] = []  # Phase 26-ggg: scope outside try для post-render таблицы
    try:
        ci_ctx = await _build_contract_items_context(p, db)
        context.update(ci_ctx)

        # Phase 26-R: receipts images for advance reports' service note
        if p.purchase_method == "advance":
            import tempfile as _tempfile
            from app.models.purchase_receipt import PurchaseReceipt as _PurchaseReceipt
            from app.routers.purchase_receipts import _render_receipt_png as _rrpng
            _receipts_q = await db.execute(
                select(_PurchaseReceipt)
                .where(_PurchaseReceipt.purchase_id == p.id)
                .order_by(_PurchaseReceipt.id.asc())
            )
            _receipts = _receipts_q.scalars().all()
            # Phase 26-fff: одно превью PNG используется для нескольких InlineImage
            # с разной шириной — для разных layouts в шаблоне.
            #   default (6.5cm) → одиночная колонка/полный текст СЗ
            #   small (4.5cm)   → 2-колоночная таблица с узкими ячейками
            #   full  (14cm)    → отдельная страница на чек
            receipt_images = []        # default 6.5 cm — backward compat
            receipt_images_small = []  # 4.5 cm — для 2-col layouts
            receipt_images_full = []   # 14 cm — для full-width layouts
            receipt_png_paths = []     # Phase 26-ggg: пути PNG для post-render таблицы
            for _r in _receipts:
                try:
                    _png_bytes = _rrpng(_r)
                    _tmp = _tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                    _tmp.write(_png_bytes)
                    _tmp.close()
                    receipt_png_paths.append(_tmp.name)
                    receipt_images.append(InlineImage(tpl, _tmp.name, width=_Cm(6.5)))
                    receipt_images_small.append(InlineImage(tpl, _tmp.name, width=_Cm(4.5)))
                    receipt_images_full.append(InlineImage(tpl, _tmp.name, width=_Cm(14)))
                except Exception as _re:
                    import logging as _logging
                    _logging.getLogger(__name__).warning(f"render receipt {_r.id} skipped: {_re}")
            context["receipts"] = receipt_images
            context["receipt_images"] = receipt_images  # alias
            context["receipts_small"] = receipt_images_small
            context["receipts_full"] = receipt_images_full
            # Phase 26-ggg: sentinel-маркер. Пользователь ставит {{ receipts_table }}
            # в шаблон одним параграфом — после render текст становится
            # RECEIPTS_TABLE_MARKER, post-process находит и заменяет на
            # настоящую docx-таблицу 2-col с PNG чеков в каждой ячейке.
            # Решает проблему clip'а inline-images в узких ячейках/параграфах.
            context["receipts_table"] = RECEIPTS_TABLE_MARKER
            # Phase 26-LL: chunked в пары для таблицы 2 колонки в шаблоне СЗ
            receipt_pairs = []
            for _i in range(0, len(receipt_images_small), 2):
                _left = receipt_images_small[_i]
                _right = receipt_images_small[_i + 1] if _i + 1 < len(receipt_images_small) else None
                receipt_pairs.append({'left': _left, 'right': _right})
            context["receipt_pairs"] = receipt_pairs
            # Phase 26-RR: split на 2 потока для статичной таблицы 1×2 в шаблоне.
            # left/right берут МАЛЫЕ изображения (4.5 cm) — таблица 2-колоночная,
            # узкие ячейки. Старый шаблон с _Cm(6.5) клипал картинку и пользователь
            # видел пустую узкую полоску.
            context["left_receipts"] = receipt_images_small[::2]
            context["right_receipts"] = receipt_images_small[1::2]
        else:
            context["receipts"] = []
            context["receipt_images"] = []
            context["receipts_small"] = []
            context["receipts_full"] = []
            context["receipt_pairs"] = []
            context["left_receipts"] = []
            context["right_receipts"] = []
            context["receipts_table"] = ""  # marker отсутствует — paragraph будет пустой
            receipt_png_paths = []
    except HTTPException:
        raise
    except Exception as _ctx_exc:
        import traceback as _tb
        logger.exception("Document context build failed (pre-render) for purchase %s, doc_type=%s", pid, doc_type)
        _err_class = type(_ctx_exc).__name__
        _err_msg = str(_ctx_exc)
        _err_tb = "".join(_tb.format_exception(type(_ctx_exc), _ctx_exc, _ctx_exc.__traceback__))[:4000]
        raise HTTPException(500, detail={
            "code": "DOCUMENT_GENERATION_FAILED",
            "message": f"Не удалось сформировать данные для «{doc_type}»: {_err_class}",
            "doc_type": doc_type,
            "purchase_id": pid,
            "error_class": _err_class,
            "error_raw": _err_msg,
            "traceback": _err_tb,
            "hint": (
                "Ошибка при сборке контекста шаблона (до рендеринга). "
                "Возможные причины: ошибка загрузки чеков/изображений, "
                "проблема с данными закупки (FK, пустые обязательные поля), "
                "отсутствующая зависимость (PIL, qrcode). "
                "Передайте администратору error_class + traceback."
            ),
        })

    _fallback_info: dict | None = None  # phase31-02: track silent fallback for response header
    try:
        tpl.render(context)
    except Exception as render_err:
        # phase26-nn: auto-fallback на базовый шаблон если кастомный из БД сломан.
        # Lessons.md (2026-05-15): кастомные шаблоны с {% tr %} вне таблицы или
        # повреждённой структурой валят TemplateSyntaxError → user видит белый экран.
        # Безопаснее упасть на базовый из репо и предупредить.
        is_custom_template = bool(template_path and template_path.startswith(SUBSIDY_TEMPLATES_DIR))
        if is_custom_template:
            base_path = os.path.join(TEMPLATES_DIR, f"{doc_type}.docx")
            if os.path.exists(base_path) and base_path != template_path:
                import logging
                _fallback_reason = f"{type(render_err).__name__}: {str(render_err)[:300]}"
                logging.getLogger(__name__).warning(
                    f"Custom template {template_path} render failed ({_fallback_reason}). "
                    f"Falling back to base template {base_path}."
                )
                _fallback_info = {
                    "original": os.path.basename(template_path),
                    "fallback": os.path.basename(base_path),
                    "reason": _fallback_reason,
                }
                tpl = DocxTemplate(base_path)
                template_path = base_path
                tpl.render(context)
            else:
                raise
        else:
            raise

    try:
        # Post-process: fix approvers table if docxtpl loop didn't render all rows
        if doc_type == "approval_sheet" and len(approvers_list) > 0:
            from docx import Document as _DocxDoc
            from docx.shared import Pt
            from copy import deepcopy
            from lxml import etree

            _buf = BytesIO()
            tpl.save(_buf)
            _buf.seek(0)
            _doc = _DocxDoc(_buf)

            # Find the approvers table (has header row with "Должность" or "ФИО")
            target_table = None
            for _t in _doc.tables:
                hdr = " ".join(c.text.strip() for c in _t.rows[0].cells)
                if "Должность" in hdr or "ФИО" in hdr:
                    target_table = _t
                    break

            if target_table:
                # Count current data rows (skip header)
                current_data_rows = len(target_table.rows) - 1
                needed = len(approvers_list)

                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

                def _set_cell_text(cell_el, text):
                    """Replace all cell content with plain text."""
                    for p_el in list(cell_el.findall(f'{ns}p')):
                        cell_el.remove(p_el)
                    p = etree.SubElement(cell_el, f'{ns}p')
                    r = etree.SubElement(p, f'{ns}r')
                    t = etree.SubElement(r, f'{ns}t')
                    t.text = text or ""

                def _set_cell_two_lines(cell_el, line1, line2):
                    """Replace cell content with two paragraphs."""
                    for p_el in list(cell_el.findall(f'{ns}p')):
                        cell_el.remove(p_el)
                    p1 = etree.SubElement(cell_el, f'{ns}p')
                    r1 = etree.SubElement(p1, f'{ns}r')
                    t1 = etree.SubElement(r1, f'{ns}t')
                    t1.text = line1 or ""
                    p2 = etree.SubElement(cell_el, f'{ns}p')
                    r2 = etree.SubElement(p2, f'{ns}r')
                    t2 = etree.SubElement(r2, f'{ns}t')
                    t2.text = line2 or ""

                if current_data_rows < needed:
                    # Rebuild table: template loop didn't render all rows
                    template_row_el = target_table.rows[1]._tr
                    for row in list(target_table.rows[1:]):
                        target_table._tbl.remove(row._tr)
                    for a in approvers_list:
                        new_tr = deepcopy(template_row_el)
                        cells = new_tr.findall(f'.//{ns}tc')
                        if len(cells) >= 4:
                            _set_cell_text(cells[0], str(a.get("num", "")))
                            _set_cell_two_lines(cells[1], a.get("role_name", ""), a.get("full_name", ""))
                            _set_cell_text(cells[2], "")
                            _set_cell_text(cells[3], a.get("note", ""))
                        target_table._tbl.append(new_tr)
                else:
                    # Rows match — patch note/FEO into last column of each data row
                    for idx, a in enumerate(approvers_list):
                        row_idx = idx + 1  # skip header
                        if row_idx >= len(target_table.rows):
                            break
                        row_el = target_table.rows[row_idx]._tr
                        cells = row_el.findall(f'.//{ns}tc')
                        note_val = a.get("note", "")
                        if note_val and len(cells) >= 4:
                            _set_cell_text(cells[-1], note_val)

                # Phase 26-ggg: receipts table (no-op если маркер отсутствует)
                _insert_receipts_table_if_marker(_doc, receipt_png_paths)
                buf = BytesIO()
                _doc.save(buf)
                buf.seek(0)
            else:
                _insert_receipts_table_if_marker(tpl.docx, receipt_png_paths)
                buf = BytesIO()
                tpl.save(buf)
                buf.seek(0)
        else:
            _insert_receipts_table_if_marker(tpl.docx, receipt_png_paths)
            buf = BytesIO()
            tpl.save(buf)
            buf.seek(0)
    except Exception as e:
        logger.exception("Document generation error for purchase %s, doc_type=%s, template=%s", pid, doc_type, template_path)

        # Phase 23.2: human-friendly error explanations for Jinja/docxtpl errors
        import re as _re
        err_class = type(e).__name__
        err_msg = str(e)
        template_name = os.path.basename(template_path) if template_path else f"{doc_type}.docx"
        is_custom = bool(template_path and template_path.startswith(SUBSIDY_TEMPLATES_DIR))

        detail = {
            "code": "TEMPLATE_RENDER_ERROR",
            "message": f"Не удалось сгенерировать «{template_name}»",
            "template": template_name,
            "template_source": "Шаблон субсидии (загруженный пользователем)" if is_custom else "Глобальный шаблон",
            "error_class": err_class,
            "error_raw": err_msg,
            "hint": None,
        }

        # Pattern: Jinja2 UndefinedError ('X' is undefined)
        m = _re.match(r"'([a-zA-Z_][a-zA-Z0-9_]*)' is undefined", err_msg)
        if m:
            var_name = m.group(1)
            detail["message"] = f"В шаблоне «{template_name}» используется переменная {{{{{var_name}.…}}}} вне цикла"
            loop_hints = {
                "a": "{% tr for a in approvers %}…{{a.full_name}}…{% tr endfor %}  — переменная для согласующих (approval_sheet, лист согласования)",
                "item": "{% tr for item in items %}…{{item.name}}…{% tr endfor %}  — переменная для позиций закупки",
            }
            if var_name in loop_hints:
                detail["hint"] = (
                    f"Переменная «{var_name}» доступна только внутри цикла. "
                    f"Оберните строки/ячейки таблицы в:\n  {loop_hints[var_name]}\n\n"
                    f"Либо удалите кастомный шаблон в UI «Субсидии → Шаблоны → {template_name} → 🗑» и используйте глобальный."
                )
            else:
                detail["hint"] = (
                    f"В словаре переменных нет «{var_name}». Возможно, переменная переименована или удалена. "
                    f"См. справочник «Руководство по переменным» в Subsidies → Шаблоны."
                )

        # Pattern: 'X' has no attribute 'Y'
        m2 = _re.match(r"'(\w+)' (?:object )?has no attribute '(\w+)'", err_msg)
        if not m and m2:
            obj_name = m2.group(1)
            attr = m2.group(2)
            detail["message"] = f"В шаблоне «{template_name}» используется {{{{{obj_name}.{attr}}}}} но поле «{attr}» отсутствует"
            detail["hint"] = (
                f"Возможные причины:\n"
                f"• опечатка в имени переменной — см. «Руководство по переменным»\n"
                f"• данные ещё не заполнены в закупке (например, попытка использовать {{{{{obj_name}.{attr}}}}} когда поле пустое)"
            )

        # Pattern: TemplateSyntaxError
        if err_class == "TemplateSyntaxError":
            detail["message"] = f"Синтаксическая ошибка в шаблоне «{template_name}»: {err_msg}"
            detail["hint"] = (
                "Проверьте парные теги в Word: `{% if ... %}` ↔ `{% endif %}`, "
                "`{% for ... %}` ↔ `{% endfor %}`. Откройте шаблон в Word и убедитесь, "
                "что все условные блоки закрыты."
            )

        # Pattern: file not found / permission
        if isinstance(e, FileNotFoundError) or "no such file" in err_msg.lower():
            detail["message"] = f"Файл шаблона не найден: {template_name}"
            detail["hint"] = "Загрузите шаблон через UI «Субсидии → Шаблоны → Загрузить свой шаблон»."

        raise HTTPException(500, detail=detail)

    # For contract docs: append ТЗ table with items
    if doc_type == "contract" and items_list:
        try:
            from docx import Document as _DocxDoc
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = _DocxDoc(buf)

            # Add page break before ТЗ
            doc.add_page_break()

            # Title
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run("ТЕХНИЧЕСКОЕ ЗАДАНИЕ")
            run.bold = True
            run.font.size = Pt(12)

            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.add_run(f"(Приложение к договору № {p.contract_number or '___'} от {_fmt_date(p.contract_date) or '__.__.____'})")

            doc.add_paragraph()  # spacer

            # Table with columns: №, Фото, Наименование и описание, Кол-во, Ед., Цена ед., Сумма
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'

            # Header row
            hdr_cells = table.rows[0].cells
            headers_tz = ["№", "Фото", "Наименование и описание", "Кол-во", "Ед.", "Цена ед., ₽", "Сумма, ₽"]
            col_widths_cm = [1.0, 2.5, 8.0, 1.8, 1.5, 2.8, 2.8]
            for i, (hdr_cell, hdr_text, w) in enumerate(zip(hdr_cells, headers_tz, col_widths_cm)):
                hdr_cell.width = Cm(w)
                para = hdr_cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(hdr_text)
                run.bold = True
                run.font.size = Pt(9)
                # Blue background
                tc = hdr_cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '1E40AF')
                shd.set(qn('w:color'), 'FFFFFF')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
                # White font
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            # Data rows
            total_sum = 0.0
            for item_data in items_list:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item_data["num"])
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Photo cell: leave empty
                row_cells[1].text = ""
                # Name + description
                name_cell = row_cells[2]
                name_para = name_cell.paragraphs[0]
                name_run = name_para.add_run(str(item_data["name"]))
                name_run.bold = True
                name_run.font.size = Pt(9)
                if item_data.get("description"):
                    desc_para = name_cell.add_paragraph()
                    desc_run = desc_para.add_run(str(item_data["description"]))
                    desc_run.font.size = Pt(8)
                    desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                qty_val = item_data.get("quantity", "")
                row_cells[3].text = str(qty_val) if qty_val != "" else "—"
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[4].text = str(item_data.get("unit", "") or "—")
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[5].text = str(item_data.get("unit_price", "") or "—")
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row_cells[6].text = str(item_data.get("total_price", "") or "—")
                row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # Font size for all data cells
                for cell in row_cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if not run.font.size:
                                run.font.size = Pt(9)

                try:
                    price_str = str(item_data.get("total_price", "")).replace(" ", "").replace(",", ".").replace("₽", "").strip()
                    total_sum += float(price_str) if price_str else 0
                except Exception:
                    pass

            # Total row
            total_row_cells = table.add_row().cells
            # Merge cells 0-5
            total_row_cells[5].merge(total_row_cells[0])
            merged_para = total_row_cells[0].paragraphs[0]
            merged_run = merged_para.add_run("НМЦК итого:")
            merged_run.bold = True
            merged_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total_row_cells[6].text = _fmt_money_plain(total_sum)
            if total_row_cells[6].paragraphs[0].runs:
                total_row_cells[6].paragraphs[0].runs[0].bold = True
            total_row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Save back to buf
            buf2 = BytesIO()
            doc.save(buf2)
            buf2.seek(0)
            buf = buf2
        except Exception as tz_err:
            # Don't fail the whole request if ТЗ append fails — just log
            import traceback
            print(f"ТЗ append error: {tz_err}\n{traceback.format_exc()}")

    # ── Methodology attach: приклеиваем методичку к готовому договору ──────
    # Владелец: методичка (большая/малая) — отдельный документ, физически
    # отсутствующий во всех семи шаблонах договора (см. test_contract_templates.py,
    # _METHODOLOGY_MARKERS). Она приклеивается ПОСЛЕ рендера договора через
    # docxcompose.Composer — итоговый файл содержит и договор, и методичку.
    if doc_type in CONTRACT_TYPED_FORM_DOC_TYPES:
        methodology = getattr(p, "methodology", None)
        if methodology in ("large", "small"):
            methodology_doc_type = f"methodology_{methodology}"
            meth_label = "большие" if methodology == "large" else "малые"
            meth_path, _meth_file, _meth_base = _resolve_doc_template_path(
                methodology_doc_type, p.subsidy_id
            )
            if not os.path.exists(meth_path):
                raise HTTPException(
                    422,
                    detail={
                        "code": "METHODOLOGY_TEMPLATE_MISSING",
                        "message": f"Не найден файл методических рекомендаций ({meth_label})",
                        "hint": (
                            f"В закупке выбрана методичка «{meth_label}», но соответствующий "
                            f"файл шаблона ({DOC_TYPES[methodology_doc_type][0]}) не загружен "
                            "ни на субсидию, ни глобально. Загрузите файл в «Шаблоны документов» "
                            "субсидии или положите его в backend/templates/, либо снимите выбор "
                            "методички в закупке."
                        ),
                        "doc_type": doc_type,
                        "methodology": methodology,
                    },
                )
            try:
                from docxcompose.composer import Composer as _Composer
                from docx import Document as _ComposeDoc

                buf.seek(0)
                _master = _ComposeDoc(buf)
                _composer = _Composer(_master)
                _composer.append(_ComposeDoc(meth_path))
                _composed_buf = BytesIO()
                _composer.save(_composed_buf)
                _composed_buf.seek(0)
                buf = _composed_buf
            except HTTPException:
                raise
            except Exception as meth_err:
                logger.exception(
                    "Methodology attach error for purchase %s, doc_type=%s, methodology=%s",
                    pid, doc_type, methodology,
                )
                raise HTTPException(
                    422,
                    detail={
                        "code": "METHODOLOGY_ATTACH_FAILED",
                        "message": "Не удалось приклеить методические рекомендации к договору",
                        "hint": f"{type(meth_err).__name__}: {meth_err}",
                        "doc_type": doc_type,
                        "methodology": methodology,
                    },
                )

    # ── Phase 19.06: merge with secondary doc ──────────────────────────────
    # When ?merge=<doc_type> is passed, render the secondary template against
    # the same context and append its paragraphs/tables after a page break.
    if merge and merge in DOC_TYPES and merge != doc_type:
        try:
            from docxtpl import DocxTemplate as _Tpl
            from docx import Document as _Docx
            from copy import deepcopy

            secondary_file, secondary_base = DOC_TYPES[merge]
            secondary_path = os.path.join(TEMPLATES_DIR, secondary_file)
            # subsidy override for secondary
            if p.subsidy_id:
                sub_override = os.path.join(
                    SUBSIDY_TEMPLATES_DIR, "subsidies", str(p.subsidy_id), f"{merge}.docx"
                )
                if os.path.exists(sub_override):
                    secondary_path = sub_override
            # fallback if the dedicated secondary file is missing
            if not os.path.exists(secondary_path):
                fb = DOC_TYPE_FALLBACK_FILES.get(merge)
                if fb:
                    fb_path = os.path.join(TEMPLATES_DIR, fb)
                    if os.path.exists(fb_path):
                        secondary_path = fb_path

            if os.path.exists(secondary_path):
                sec_tpl = _Tpl(secondary_path)
                sec_tpl.render(context)
                sec_buf = BytesIO()
                sec_tpl.save(sec_buf)
                sec_buf.seek(0)

                # Append secondary into primary
                buf.seek(0)
                main_doc = _Docx(buf)
                sec_doc = _Docx(sec_buf)

                main_doc.add_page_break()
                # Copy every top-level element (paragraphs, tables, etc.) from body
                for element in sec_doc.element.body:
                    main_doc.element.body.append(deepcopy(element))

                merged_buf = BytesIO()
                main_doc.save(merged_buf)
                merged_buf.seek(0)
                buf = merged_buf
                # Update filename to reflect merge
                filename_base = f"{filename_base}_и_{secondary_base}"
        except Exception as merge_err:
            import traceback
            print(f"Doc merge error ({doc_type} + {merge}): {merge_err}\n{traceback.format_exc()}")

    _subj = _sanitize_subject(getattr(p, "subject", "") or "")
    if _subj:
        safe_name = f"{filename_base}_{_subj}_{p.registry_number or pid}.docx".replace("/", "-").replace(" ", "_")
    else:
        safe_name = f"{filename_base}_{p.registry_number or pid}.docx".replace("/", "-").replace(" ", "_")
    encoded_name = quote(safe_name, safe="-_.~")
    resp_headers: dict = {"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"}
    # phase31-02: surface non-silent fallback via response headers
    if _fallback_info:
        resp_headers["X-Template-Fallback"] = "1"
        resp_headers["X-Template-Fallback-Reason"] = quote(_fallback_info["reason"][:200], safe="")
        resp_headers["X-Template-Fallback-Original"] = quote(_fallback_info["original"], safe="")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=resp_headers,
    )


# ── Fabrikant package ZIP endpoint ───────────────────────────────────────────

# File names for Fabrikant document package (видны пользователю в кабинете
# площадки в разделе документов извещения, поэтому названы по-русски)
FABRIKANT_PKG_FILE_NAMES = [
    ("fabrikant_instruction",      "Инструкция по заполнению заявки.docx", "Инструкция по заполнению заявки"),
    ("fabrikant_application_form", "Заявка (форма).docx",                  "Заявка (форма)"),
    ("fabrikant_documentation",    "Документация к закупке.docx",          "Документация к закупке"),
    ("fabrikant_contract_project", "Проект договора.docx",                 "Проект договора"),
    ("tech_spec_request",          "Техническое задание.docx",             "Техническое задание"),
]


def _strip_tech_spec_legend(docx_bytes: bytes) -> bytes:
    """Remove legend paragraphs from a rendered tech_spec_request docx.

    Deletes paragraphs from the first one containing «ЛЕГЕНДА» up to and
    including the first paragraph consisting solely of em-dashes (—), which
    acts as a visual separator.  If «ЛЕГЕНДА» is not found the original bytes
    are returned unchanged (custom subsidy template without a legend block).

    Wrapped in try/except: on any error returns the original bytes so the
    caller always receives a valid document.
    """
    try:
        from io import BytesIO as _BytesIO
        from docx import Document as _Document

        bio = _BytesIO(docx_bytes)
        doc = _Document(bio)
        paras = doc.paragraphs

        # Find start index: first paragraph containing «ЛЕГЕНДА»
        start_idx = None
        for i, p in enumerate(paras):
            if "ЛЕГЕНДА" in p.text:
                start_idx = i
                break

        if start_idx is None:
            return docx_bytes

        # Find end index: first paragraph after start consisting only of em-dashes
        end_idx = None
        for i in range(start_idx, len(paras)):
            stripped = paras[i].text.strip()
            if stripped and all(ch == "—" for ch in stripped):
                end_idx = i
                break

        if end_idx is None:
            # No separator found — remove from start to end of legend block
            end_idx = start_idx

        # Remove paragraphs in reverse order to keep indices stable
        for i in range(end_idx, start_idx - 1, -1):
            p_el = paras[i]._element
            p_el.getparent().remove(p_el)

        out = _BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception as _exc:
        logger.warning("_strip_tech_spec_legend: failed to strip legend, returning original: %s", _exc)
        return docx_bytes


async def render_fabrikant_package_files(
    db: AsyncSession,
    purchase_id: int,
) -> tuple[list[tuple[str, str, bytes]], list[str]]:
    """Render 5 Fabrikant documents for *purchase_id*.

    Returns:
        rendered: list of (ascii_file_name, ru_title, bytes)
        errors:   list of human-readable error strings for failed docs
    """
    import traceback as _traceback
    from docxtpl import DocxTemplate as _DxTpl
    from app.models.contract_item import ContractItem as _CI

    result = await db.execute(
        select(Purchase)
        .options(
            selectinload(Purchase.items).selectinload(PurchaseItem.product),
            selectinload(Purchase.contractor),
            selectinload(Purchase.feo_category),
            selectinload(Purchase.contract_items),
            selectinload(Purchase.assigned_user),
        )
        .where(Purchase.id == purchase_id)
    )
    p = result.scalar_one_or_none()
    if not p:
        return [], [f"Закупка {purchase_id} не найдена"]

    subsidy_r = await db.execute(select(Subsidy).where(Subsidy.id == p.subsidy_id))
    subsidy = subsidy_r.scalar_one_or_none()

    customer_org = None
    customer_ctr = None
    if subsidy and subsidy.org_id:
        from app.models.organization import Organization as _Org
        org_r = await db.execute(select(_Org).where(_Org.id == subsidy.org_id))
        customer_org = org_r.scalar_one_or_none()
        if customer_org and customer_org.contractor_id:
            from app.models.contractor import Contractor as _CtrC
            ctr_r = await db.execute(select(_CtrC).where(_CtrC.id == customer_org.contractor_id))
            customer_ctr = ctr_r.scalar_one_or_none()

    event = None
    if p.event_id:
        ev_r = await db.execute(select(Event).where(Event.id == p.event_id))
        event = ev_r.scalar_one_or_none()

    feo_path = feo_level_1 = feo_level_2 = feo_level_3 = ""
    if p.feo_category_id:
        feo_res = await db.execute(select(FeoCategory))
        all_feo = {f.id: f for f in feo_res.scalars().all()}
        path_nodes: list = []
        node_id = p.feo_category_id
        visited: set = set()
        while node_id and node_id not in visited:
            visited.add(node_id)
            cat = all_feo.get(node_id)
            if not cat:
                break
            path_nodes.append(cat)
            node_id = cat.parent_id
        path_nodes.reverse()
        feo_path = " → ".join(n.name.strip() for n in path_nodes)
        if len(path_nodes) >= 1: feo_level_1 = path_nodes[0].name.strip()
        if len(path_nodes) >= 2: feo_level_2 = path_nodes[1].name.strip()
        if len(path_nodes) >= 3: feo_level_3 = path_nodes[2].name.strip()

    def _fmt_initials(full: str) -> str:
        if not full:
            return ""
        parts = full.strip().split()
        if not parts:
            return ""
        surname = parts[0]
        initials = [w[0].upper() + "." for w in parts[1:3] if w and w[0].isalpha()]
        return f"{surname} {''.join(initials)}".strip()

    assigned_full = (getattr(p.assigned_user, "full_name", None) or "") if getattr(p, "assigned_user", None) else ""
    raw_responsible = assigned_full or p.responsible_person or ""
    resolved_responsible = _fmt_initials(raw_responsible) if raw_responsible else ""

    # Этот эндпоинт всегда рендерит один и тот же фиксированный набор из 5
    # doc_type — FABRIKANT_PKG_FILE_NAMES (fabrikant_* + tech_spec_request).
    # ВСЕ они намеренно исключены из CONTRACT_FAMILY_DOC_TYPES: это пакет для
    # тендерной процедуры на Фабрикант, публикуется ДО выбора поставщика —
    # ContractItem на этой стадии закупки ещё физически не существует, поэтому
    # items_list здесь по-прежнему строится из purchase_items (план), как и
    # раньше. Sanity-check ниже фейлит явно, если это когда-нибудь изменится
    # (новый doc_type в списке окажется договорным) — молчаливая подмена на
    # плановые данные для настоящего договорного документа недопустима.
    _fabrikant_pkg_doc_keys = {dk for dk, _, _ in FABRIKANT_PKG_FILE_NAMES}
    assert not (_fabrikant_pkg_doc_keys & CONTRACT_FAMILY_DOC_TYPES), (
        "render_fabrikant_package_files теперь рендерит договорной doc_type "
        f"({_fabrikant_pkg_doc_keys & CONTRACT_FAMILY_DOC_TYPES}) — items_list "
        "здесь построен из purchase_items (план), нужно переключить на "
        "_build_items_list_from_contract_items, иначе документ уйдёт с плановыми данными"
    )

    items_list = []
    for idx, (item, qty, total) in enumerate(_merge_identical_items(p.items or []), start=1):
        items_list.append({
            "num": idx,
            "name": item.item_name or "",
            "description": "",
            "type": item.item_type or "",
            "item_kind": (item.product.item_kind if item.product else None) or "товар",
            "quantity": float(qty) if qty else "",
            "unit": item.unit or "",
            "unit_price": _fmt_money(item.unit_price),
            "total_price": _fmt_money(total),
            "total": _fmt_money(total),
            "photo": "",
            # Поля для repair_framework (появятся позже, пока заглушки)
            "code": "",
            "norm_hours": "",
        })

    vat_app = bool(p.vat_applicable)
    vat_rate_val = p.vat_rate or 20
    price_val = float(p.contract_price or 0)
    vat_amount_val = price_val * vat_rate_val / (100 + vat_rate_val) if (vat_app and price_val) else 0.0
    if vat_app:
        vat_info_line = f"В том числе НДС {vat_rate_val}%: {_fmt_money_plain(vat_amount_val)} руб."
    else:
        art = (p.vat_exemption_article or "").strip()
        vat_info_line = "НДС не облагается" + (f" ({art})" if art else "")

    def _cd_parts():
        d = p.contract_date
        if not d:
            return "", "", ""
        if isinstance(d, str):
            try:
                from datetime import date as _d2
                d = _d2.fromisoformat(d)
            except ValueError:
                return "", "", ""
        months_ru = ["января", "февраля", "марта", "апреля", "мая", "июня",
                     "июля", "августа", "сентября", "октября", "ноября", "декабря"]
        return str(d.day).zfill(2), months_ru[d.month - 1], str(d.year)

    cd_day, cd_month, cd_year = _cd_parts()

    def _g2(*sources, default=""):
        for s in sources:
            if s:
                return s
        return default

    cust_signatory_full_z = _g2(
        customer_org.signatory if customer_org else None,
        customer_ctr.signatory if customer_ctr else None,
    )
    _cust_z_last = _g2(
        getattr(customer_org, 'signatory_last_name', None) if customer_org else None,
        getattr(customer_ctr, 'signatory_last_name', None) if customer_ctr else None,
    ) or None
    _cust_z_first = _g2(
        getattr(customer_org, 'signatory_first_name', None) if customer_org else None,
        getattr(customer_ctr, 'signatory_first_name', None) if customer_ctr else None,
    ) or None
    _cust_z_middle = _g2(
        getattr(customer_org, 'signatory_middle_name', None) if customer_org else None,
        getattr(customer_ctr, 'signatory_middle_name', None) if customer_ctr else None,
    ) or None
    _cust_z_position = _g2(
        getattr(customer_org, 'signatory_position', None) if customer_org else None,
        getattr(customer_ctr, 'signatory_position', None) if customer_ctr else None,
    ) or None
    cust_sig_z = _signatory_split(
        cust_signatory_full_z,
        last=_cust_z_last, first=_cust_z_first, middle=_cust_z_middle, position=_cust_z_position,
    )
    cust_signatory_basis_z = _g2(
        customer_ctr.signatory_basis if customer_ctr else None,
        "Устава",
    )
    c = p.contractor
    ctr_sig_z = _signatory_split(
        c.signatory if c else "",
        last=getattr(c, 'signatory_last_name', None) if c else None,
        first=getattr(c, 'signatory_first_name', None) if c else None,
        middle=getattr(c, 'signatory_middle_name', None) if c else None,
        position=getattr(c, 'signatory_position', None) if c else None,
    )

    # Fabrikant notice_number — platform_number (номер процедуры на ЭТП),
    # откат на external_id (наш internal requestId), если ещё не проставлен
    _notice_number = ""
    try:
        from app.models.platform_publication import PlatformPublication as _PlatPubZ
        _pub_q = await db.execute(
            select(_PlatPubZ.platform_number, _PlatPubZ.external_id)
            .where(
                _PlatPubZ.purchase_id == p.id,
                _PlatPubZ.platform == "fabrikant",
            )
            .order_by(_PlatPubZ.id.desc())
            .limit(1)
        )
        _pub_row_z = _pub_q.first()
        if _pub_row_z:
            _notice_number = _pub_row_z[0] or _pub_row_z[1] or ""
    except Exception:
        pass

    # applications_review_date
    if p.applications_review_date:
        _app_review_date = _fmt_date(p.applications_review_date)
    elif p.submission_deadline:
        from datetime import timedelta as _tds
        _next = p.submission_deadline.date() + _tds(days=1)
        while _next.weekday() >= 5:
            _next += _tds(days=1)
        _app_review_date = _next.strftime("%d.%m.%Y")
    else:
        _app_review_date = ""

    ci_ctx_z = await _build_contract_items_context(p, db)

    context = {
        "purchase_number": p.purchase_number or "",
        "registry_number": p.registry_number or "",
        "purchase_method": {"single": "Единственный поставщик", "competitive": "Конкурсная процедура"}.get(p.purchase_method or "", p.purchase_method or ""),
        "subject": p.subject or "",
        "status": p.status or "",
        "responsible_person": resolved_responsible,
        "responsible_person_full": raw_responsible,
        "subsidy_name": subsidy.name if subsidy else "",
        "subsidy_year": subsidy.year if subsidy else "",
        "subsidy_budget": _fmt_money(subsidy.budget) if subsidy else "",
        "subsidy_agreement_number": (subsidy.agreement_number or "") if subsidy else "",
        "contractor_name": (c.name or "") if c else "",
        "contractor_inn": _clean_id(c.inn) if c else "",
        "contractor_kpp": _clean_id(c.kpp) if c else "",
        "contractor_address": (c.address or "") if c else "",
        "contractor_ogrn": (c.ogrn or "") if c else "",
        "contractor_phone": (c.phone or "") if c else "",
        "contractor_email": (c.email or "") if c else "",
        "contractor_signatory": (c.signatory or "") if c else "",
        "contractor_signatory_basis": (c.signatory_basis or "") if c else "",
        "contractor_settlement_account": (c.settlement_account or "") if c else "",
        "contractor_bank_name": (c.bank_name or "") if c else "",
        "contractor_bik": (c.bik or "") if c else "",
        "contractor_correspondent_account": (c.correspondent_account or "") if c else "",
        "contractor_signatory_position": _signatory_position(c.signatory) if c else "",
        "contractor_signatory_name": ctr_sig_z["name_full"],
        "contractor_signatory_name_genitive": ctr_sig_z["name_genitive"],
        "contractor_signatory_initials": ctr_sig_z["name_initials"],
        "contractor_ogrnip": (c.ogrn or "") if (c and (c.org_type or "").lower().startswith("ип")) else "",
        "contractor_full_name": (c.full_name or c.name or "") if c else "",
        "contractor_short_name": (c.name or "") if c else "",
        "feo_category_name": p.feo_category.name if p.feo_category else "",
        "feo_path": feo_path,
        "feo_level_1": feo_level_1,
        "feo_level_2": feo_level_2,
        "feo_level_3": feo_level_3,
        "total_nmcd": _fmt_money(p.total_nmck or p.nmck or p.planned_total_price),
        "total_nmck": _fmt_money(p.total_nmck or p.nmck or p.planned_total_price),
        "nmck": _fmt_money(p.nmck or p.total_nmck),
        "contract_price": _fmt_money(p.contract_price),
        "economy": _fmt_money(p.economy),
        "contract_number": p.contract_number or "",
        "contract_date": _fmt_date(p.contract_date) or "__.__._____ г.",
        "contract_date_day": cd_day,
        "contract_date_month": cd_month,
        "contract_date_year": cd_year,
        "execution_term": _fmt_date(p.execution_term),
        "delivery_date": _fmt_date(p.delivery_date),
        "delivery_location": p.delivery_location or "",
        "country_origin": p.country_origin or "",
        "payment_doc_number": p.payment_doc_number or "",
        "payment_doc_date": _fmt_date(p.payment_doc_date),
        "payment_amount": _fmt_money(p.payment_amount),
        "items": items_list,
        "items_count": len(items_list),
        "item_names": p.subject or ", ".join(i["name"] for i in items_list if i["name"]),
        "approvers": [],
        "initiator_name": "",
        "initiator_role": "",
        "initiator_dept": "",
        "event_name": event.name if event else "",
        "contract_type": {"single": "Единственный поставщик"}.get(p.purchase_contract_type or "", ""),
        "today": _fmt_date(date.today()),
        "today_iso": date.today().isoformat(),
        "service_name": p.subject or "",
        "submission_deadline_date": p.submission_deadline.date().isoformat() if p.submission_deadline else "",
        "submission_deadline_time": p.submission_deadline.strftime("%H:%M") if p.submission_deadline else "",
        "submission_deadline_datetime": p.submission_deadline.strftime("%d.%m.%Y %H:%M") if p.submission_deadline else "",
        "subsidy_agreement_text": (subsidy.agreement_text if (subsidy and subsidy.agreement_text) else ""),
        "vat_applicable": vat_app,
        "vat_rate": vat_rate_val,
        "vat_amount_num": _fmt_money_plain(vat_amount_val),
        "vat_amount_words": _rubles_to_words(vat_amount_val),
        "vat_exemption_article": p.vat_exemption_article or "",
        "vat_info_line": vat_info_line,
        "contract_price_num": _fmt_money_plain(p.contract_price),
        "contract_price_words": _rubles_to_words(p.contract_price),
        "subject_kind": "goods",
        "period_type": p.service_period_type or "period",
        "service_start_date": _fmt_date(p.service_start_date) or _fmt_date(p.contract_date),
        "service_end_date": _fmt_date(p.service_end_date) or _fmt_date(p.execution_term),
        "service_date": _fmt_date(p.execution_term),
        "service_term": _format_service_term(p),
        "service_term_mode": p.service_term_mode or "",
        "service_term_days": p.service_term_days or "",
        "service_term_type": p.service_term_type or "",
        "service_deadline_date": _fmt_date(p.service_deadline_date),
        # Дата окончания срока действия договора (п. 8.1) — поле Purchase.contract_end_date
        "contract_end_date": _fmt_date(p.contract_end_date),
        "third_party_involved": bool(p.third_party_involved),
        "acceptance_doc_name": p.acceptance_doc_name or "",
        "acceptance_doc_number": p.acceptance_doc_number or "",
        "acceptance_doc_date": _fmt_date(p.acceptance_doc_date) or "",
        "acceptance_doc_amount": _fmt_money(p.acceptance_doc_amount) if p.acceptance_doc_amount else "",
        "receipts": [],
        "receipt_images": [],
        "receipts_small": [],
        "receipts_full": [],
        "receipt_pairs": [],
        "left_receipts": [],
        "right_receipts": [],
        "receipts_table": "",
        # Customer (Organisation)
        "customer_name": _g2(customer_org.name if customer_org else None, customer_ctr.name if customer_ctr else None),
        "customer_full_name": _g2(customer_org.full_name if customer_org else None, customer_ctr.name if customer_ctr else None, customer_org.name if customer_org else None),
        "customer_short_name": _g2(customer_org.name if customer_org else None, customer_ctr.name if customer_ctr else None),
        "customer_address": _g2(customer_org.address if customer_org else None, customer_ctr.address if customer_ctr else None),
        "customer_postal_address": _g2(customer_ctr.postal_address if customer_ctr else None, customer_org.address if customer_org else None),
        "customer_inn": _clean_id(_g2(customer_org.inn if customer_org else None, customer_ctr.inn if customer_ctr else None)),
        "customer_kpp": _clean_id(_g2(customer_org.kpp if customer_org else None, customer_ctr.kpp if customer_ctr else None)),
        "customer_ogrn": _g2(customer_org.ogrn if customer_org else None, customer_ctr.ogrn if customer_ctr else None),
        "customer_bank_name": _g2(customer_ctr.bank_name if customer_ctr else None),
        "customer_settlement_account": _g2(customer_ctr.settlement_account if customer_ctr else None),
        "customer_correspondent_account": _g2(customer_ctr.correspondent_account if customer_ctr else None),
        "customer_bik": _g2(customer_ctr.bik if customer_ctr else None),
        # Лицевой счёт Заказчика
        "customer_personal_account": _g2(customer_ctr.personal_account if customer_ctr else None),
        "customer_phone": _g2(customer_ctr.phone if customer_ctr else None),
        "customer_email": _g2(customer_ctr.email if customer_ctr else None),
        "customer_signatory": cust_signatory_full_z,
        "customer_signatory_position": cust_sig_z["position"],
        "customer_signatory_name": cust_sig_z["name_full"],
        "customer_signatory_name_genitive": cust_sig_z["name_genitive"],
        "customer_signatory_initials": cust_sig_z["name_initials"],
        # Инициалы впереди: «Е.В. Козеев» (для строки подписи)
        "customer_signatory_name_initials": _fio_to_initials_prefix(cust_sig_z["name_full"]),
        "customer_signatory_basis": cust_signatory_basis_z,
        "contract_city": (customer_org.contract_city or "Москва") if customer_org else "Москва",
        # Fabrikant-specific
        "notice_number": _notice_number,
        "payment_term_days": p.payment_term_days if (p.payment_term_days is not None) else 10,
        "applications_review_date": _app_review_date,
        # Phase 28 keys (best-effort)
        "subsidy_grantor_name": (subsidy.grantor_name or "").strip() if subsidy else "",
        "subsidy_ministry_name": (subsidy.ministry_name or "").strip() if subsidy else "",
        "subsidy_agreement_date": "",
        "advance_amount": _fmt_money(p.advance_amount) if p.advance_amount else "",
        "acceptance_term_days": p.acceptance_term_days if p.acceptance_term_days is not None else 5,
        "penalty_rate": str(p.penalty_rate) if p.penalty_rate is not None else "0.1",
        "warranty_period_days": p.warranty_period_days if p.warranty_period_days is not None else 15,
        "is_retroactive": bool(p.is_retroactive),
        # Phase 28 T3: условные блоки (поля в модели Purchase пока отсутствуют — T8)
        "delivery_by_supplier": bool(getattr(p, "delivery_by_supplier", True)),
        "has_stages": bool(getattr(p, "has_stages", False)),
        "subsidy_extra_clause_1": (subsidy.extra_contract_clause_1 or '').strip() if subsidy else '',
        "subsidy_extra_clause_2": (subsidy.extra_contract_clause_2 or '').strip() if subsidy else '',
        "commission_members": [],
        "commission_member_1_name": "",
        "commission_member_2_name": "",
        "commission_member_3_name": "",
        "contractor_passport_series": (c.passport_series or '') if c else '',
        "contractor_passport_number": (c.passport_number or '') if c else '',
        "contractor_passport_issuer": (c.passport_issuer or '') if c else '',
        "contractor_passport_issued_date": _fmt_date(c.passport_issued_date) if c and c.passport_issued_date else '',
        "contractor_snils": (c.snils or '') if c else '',
        "contractor_registration_address": (c.registration_address or c.address or '') if c else '',
        "contractor_birth_date": _fmt_date(c.birth_date) if c and c.birth_date else '',
        "contractor_ogrnip_date": _fmt_date(p.contractor_ogrnip_date) if p.contractor_ogrnip_date else "",
        "repair_request_number": (p.repair_request_number or "").strip() if hasattr(p, "repair_request_number") else "",
        "procurement_protocol_number": (p.procurement_protocol_number or "").strip(),
        "procurement_order_number": (p.procurement_order_number or "").strip(),
    }
    context.update(ci_ctx_z)
    context["initiator_name_gen"] = ""
    context["initiator_position_gen"] = ""
    context["responsible_name_gen"] = _to_gen_fio(resolved_responsible)
    context["responsible_position_gen"] = ""
    context["initiator_dept_gen"] = ""
    context["service_name_gen"] = _inflect_phrase_genitive(p.subject or "")
    context["contractor_signatory_line"] = (
        f"{c.signatory}, действует на основании {c.signatory_basis}"
        if c and c.signatory and c.signatory_basis
        else ((c.signatory or "") if c else "")
    )
    context["contractor_bank_details"] = (c.bank_name or "") if c else ""
    context["contractor_org_type"] = (c.org_type or "") if c else ""

    # ── Which templates to render ─────────────────────────────────────────────
    # tech_spec_request with fallback to contract_tz
    _tz_file, _tz_base = DOC_TYPES.get("tech_spec_request", ("tech_spec_request.docx", "ТЗ_запрос_цен"))
    _tz_path = os.path.join(TEMPLATES_DIR, _tz_file)
    if p.subsidy_id:
        _sub_tz = os.path.join(SUBSIDY_TEMPLATES_DIR, "subsidies", str(p.subsidy_id), "tech_spec_request.docx")
        if os.path.exists(_sub_tz):
            _tz_path = _sub_tz
    if not os.path.exists(_tz_path):
        _fb = DOC_TYPE_FALLBACK_FILES.get("tech_spec_request", "contract_tz.docx")
        _fb_path = os.path.join(TEMPLATES_DIR, _fb)
        if os.path.exists(_fb_path):
            _tz_path = _fb_path

    # (doc_key, archive_name used in ZIP) — legacy internal names, kept for ZIP endpoint
    _pkg_docs_legacy = [
        ("fabrikant_instruction",      "Инструкция.docx"),
        ("fabrikant_application_form", "Форма_заявки.docx"),
        ("fabrikant_documentation",    "Документация.docx"),
        ("fabrikant_contract_project", "Проект_договора.docx"),
        ("tech_spec_request",          "ТЗ.docx"),
    ]

    # ── Override file_type → doc_key mapping ─────────────────────────────────
    _override_map = {
        "fabrikant_instruction":      "fabrikant_instruction",
        "fabrikant_application_form": "fabrikant_application_form",
        "fabrikant_documentation":    "fabrikant_documentation",
        "fabrikant_contract_project": "fabrikant_contract_project",
        "tech_spec_request":          "fabrikant_tech_spec",
    }
    from app.models.purchase_file import PurchaseFile as _PF
    _override_keys = list(_override_map.values())
    _pf_res = await db.execute(
        select(_PF)
        .where(
            _PF.purchase_id == purchase_id,
            _PF.file_type.in_(_override_keys),
            _PF.is_active == True,
        )
        .order_by(_PF.id.desc())
    )
    _overrides: dict[str, _PF] = {}
    for _pf in _pf_res.scalars().all():
        if _pf.file_type not in _overrides:
            _overrides[_pf.file_type] = _pf

    # ── Render each doc ───────────────────────────────────────────────────────
    # rendered: (ascii_file_name, ru_title, bytes)
    rendered: list[tuple[str, str, bytes]] = []
    errors: list[str] = []

    for (doc_key, ascii_name, ru_title), (_doc_key2, legacy_arc) in zip(FABRIKANT_PKG_FILE_NAMES, _pkg_docs_legacy):
        override_ft = _override_map.get(doc_key)
        if override_ft and override_ft in _overrides:
            _pf_ov = _overrides[override_ft]
            orig_ext = os.path.splitext(_pf_ov.original_name or _pf_ov.filename or "")[1] or ".docx"
            ov_ascii_name = os.path.splitext(ascii_name)[0] + orig_ext
            if _pf_ov.filepath and os.path.exists(_pf_ov.filepath):
                with open(_pf_ov.filepath, "rb") as _fh:
                    rendered.append((ov_ascii_name, ru_title, _fh.read()))
                continue
            else:
                errors.append(f"{ascii_name}: файл override не найден на диске ({_pf_ov.filepath})")

        if doc_key == "tech_spec_request":
            tpl_path = _tz_path
        else:
            tpl_file, _ = DOC_TYPES.get(doc_key, (f"{doc_key}.docx", doc_key))
            tpl_path = os.path.join(TEMPLATES_DIR, tpl_file)
            if p.subsidy_id:
                _sub_override = os.path.join(SUBSIDY_TEMPLATES_DIR, "subsidies", str(p.subsidy_id), f"{doc_key}.docx")
                if os.path.exists(_sub_override):
                    tpl_path = _sub_override

        if not os.path.exists(tpl_path):
            errors.append(f"{ascii_name}: шаблон не найден ({tpl_path})")
            continue

        try:
            _tpl = _DxTpl(tpl_path)
            _tpl.render(context)
            _buf = BytesIO()
            _tpl.save(_buf)
            _rendered_bytes = _buf.getvalue()
            if doc_key == "tech_spec_request":
                _rendered_bytes = _strip_tech_spec_legend(_rendered_bytes)
            rendered.append((ascii_name, ru_title, _rendered_bytes))
        except Exception as _re:
            _tb = _traceback.format_exc()
            errors.append(f"{ascii_name}: {type(_re).__name__}: {_re}\n{_tb[:500]}")
            logger.warning("fabrikant-package render error for %s (purchase %s): %s", doc_key, purchase_id, _re)

    return rendered, errors


@router.get("/{pid}/fabrikant-package")
async def download_fabrikant_package(
    pid: int,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Render and ZIP 5 documents for ЭТП Fabrikant.

    Individual render failures are tolerated: the failed doc is omitted from
    the ZIP and listed in errors.txt.  If ALL docs fail → HTTP 500.
    """
    import zipfile as _zipfile

    rendered, errors = await render_fabrikant_package_files(db, pid)

    if not rendered:
        raise HTTPException(
            500,
            detail={
                "code": "TEMPLATE_RENDER_ERROR",
                "message": "Не удалось сформировать ни один документ для пакета Фабрикант",
                "error_raw": "\n".join(errors),
            },
        )

    # ZIP archive names: Russian base + actual extension from ascii_name
    _ru_bases = ["Инструкция", "Форма_заявки", "Документация", "Проект_договора", "ТЗ"]
    zip_buf = BytesIO()
    with _zipfile.ZipFile(zip_buf, "w", _zipfile.ZIP_DEFLATED) as zf:
        _ru_iter = iter(_ru_bases)
        for ascii_name, ru_title, data in rendered:
            _ru_base = next(_ru_iter, os.path.splitext(ascii_name)[0])
            _ext = os.path.splitext(ascii_name)[1] or ".docx"
            zf.writestr(_ru_base + _ext, data)
        if errors:
            zf.writestr("errors.txt", "\n\n".join(errors))
    zip_buf.seek(0)

    safe_zip_name = f"Фабрикант_закупка_{pid}.zip"
    encoded_zip_name = quote(safe_zip_name, safe="-_.~")
    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_zip_name}"},
    )


# ── KP xlsx export ───────────────────────────────────────────────────────────

@guide_router.get("/purchases/{pid}/kp-xlsx")
async def download_kp_xlsx(
    pid: int,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Generate xlsx with purchase items (for КП request attachment)."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise HTTPException(500, "openpyxl не установлен")

    result = await db.execute(
        select(Purchase)
        .options(selectinload(Purchase.items).selectinload(PurchaseItem.product))
        .where(Purchase.id == pid)
    )
    p = result.scalar_one_or_none()
    if not p:
        raise HTTPException(404, "Закупка не найдена")

    wb = Workbook()
    ws = wb.active
    ws.title = "Перечень товаров"

    # Header style
    header_fill = PatternFill("solid", fgColor="1E40AF")
    header_font = Font(bold=True, color="FFFFFF", size=10)
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="AAAAAA")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    headers = ["№", "Фото", "Наименование и описание", "Кол-во", "Ед.", "Цена ед., ₽", "Сумма, ₽"]
    col_widths = [5, 12, 50, 10, 8, 16, 16]

    for col, (h, w) in enumerate(zip(headers, col_widths), start=1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = border
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.row_dimensions[1].height = 32

    # Data rows
    data_align = Alignment(vertical="center", wrap_text=True)
    data_align_center = Alignment(horizontal="center", vertical="center")
    data_align_right = Alignment(horizontal="right", vertical="center")

    items = [i for i in (p.items or []) if i.item_name and i.item_name.strip()]
    for idx, item in enumerate(items, start=1):
        row = idx + 1
        description = ""
        if item.product:
            description = item.product.description or ""

        full_name = item.item_name or ""
        if description:
            full_name = full_name + "\n" + description

        ws.cell(row=row, column=1, value=idx).alignment = data_align_center
        ws.cell(row=row, column=2, value="").alignment = data_align_center  # Фото — пусто
        ws.cell(row=row, column=3, value=full_name).alignment = data_align
        ws.cell(row=row, column=4, value=float(item.quantity) if item.quantity else "").alignment = data_align_center
        ws.cell(row=row, column=5, value=item.unit or "").alignment = data_align_center
        price_cell = ws.cell(row=row, column=6, value=float(item.unit_price) if item.unit_price else "")
        price_cell.alignment = data_align_right
        if item.unit_price:
            price_cell.number_format = '# ##0.00'
        total_cell = ws.cell(row=row, column=7, value=float(item.total_price) if item.total_price else "")
        total_cell.alignment = data_align_right
        if item.total_price:
            total_cell.number_format = '# ##0.00'

        ws.row_dimensions[row].height = 40 if description else 20

        for col in range(1, 8):
            ws.cell(row=row, column=col).border = border

    # Total row
    total_row = len(items) + 2
    total_nmck = sum(float(i.total_price or 0) for i in items)
    ws.cell(row=total_row, column=6, value="НМЦК итого:").font = Font(bold=True)
    ws.cell(row=total_row, column=6).alignment = data_align_right
    total_cell = ws.cell(row=total_row, column=7, value=total_nmck)
    total_cell.font = Font(bold=True)
    total_cell.number_format = '# ##0.00'
    total_cell.alignment = data_align_right

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    fname = f"Сравнение_КП_закупка_{p.purchase_number or pid}.xlsx"
    from urllib.parse import quote as _quote
    encoded = _quote(fname, safe="-_.~")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"}
    )


# ── Template markup guide ────────────────────────────────────────────────────

TEMPLATE_VARIABLES = [
    # (template_var, description, example_template, example_result)
    # ── Выбор шаблона ──
    ("", "ВЫБОР ШАБЛОНА", "", ""),
    ("", "  contract.docx — универсальный договор (Phase 23.1): subject_kind определяется автоматически.", "", ""),
    ("", "  Все позиции 'услуга' → договор оказания услуг; иначе → договор поставки.", "", ""),
    ("{{subject_kind}}", "Тип договора (auto): вычисляется из позиций закупки", "{{subject_kind}}", "services"),
    # ── Закупка ──
    ("", "ЗАКУПКА", "", ""),
    ("{{purchase_number}}", "Номер закупки", "{{purchase_number}}", "42"),
    ("{{registry_number}}", "Реестровый номер", "{{registry_number}}", "РЕЕ-2026-00042"),
    ("{{subject}}", "Предмет закупки", "{{subject}}", "Оказание услуг связи"),
    ("{{status}}", "Статус", "{{status}}", "contracted"),
    ("{{purchase_method}}", "Способ закупки", "{{purchase_method}}", "Единственный поставщик"),
    ("{{purchase_method_label}}", "Способ закупки (для order_purchase, п.2)", "{{purchase_method_label}}", "Единственный поставщик"),
    ("{{purchase_basis}}", "Основание", "{{purchase_basis}}", "план закупок"),
    ("{{contract_type}}", "Тип договора", "{{contract_type}}", "Единственный поставщик"),
    ("{{responsible_person}}", "ФИО ответственного исполнителя", "{{responsible_person}}", "Иванов Иван Иванович"),
    ("{{feo_category_name}}", "Категория ФЭО (только выбранный узел)", "{{feo_category_name}}", "Услуги связи"),
    ("{{feo_path}}", "Полный путь ФЭО", "{{feo_path}}", "1. Адм → 1.2 IT → 1.2.3 Услуги связи"),
    ("{{feo_level_1}}", "ФЭО уровень 1 — Направление расходов", "{{feo_level_1}}", "1. Административные расходы"),
    ("{{feo_level_2}}", "ФЭО уровень 2 — Тип расходов", "{{feo_level_2}}", "1.2 IT и коммуникации"),
    ("{{feo_level_3}}", "ФЭО уровень 3 — Конкретизированная категория", "{{feo_level_3}}", "1.2.3 Услуги связи"),
    # ── Субсидия и мероприятие ──
    ("", "СУБСИДИЯ И МЕРОПРИЯТИЕ", "", ""),
    ("{{subsidy_name}}", "Наименование субсидии", "{{subsidy_name}}", "ФАДМ-2026"),
    ("{{subsidy_year}}", "Год субсидии", "{{subsidy_year}}", "2026"),
    ("{{subsidy_budget}}", "Бюджет субсидии", "{{subsidy_budget}}", "15 500 000,00 ₽"),
    ("{{event_name}}", "Название мероприятия", "{{event_name}}", "Сборы спасателей"),
    # ── Заказчик (Phase 23) ──
    ("", "ЗАКАЗЧИК", "", ""),
    ("{{customer_name}}", "Краткое название организации Заказчика", "{{customer_name}}", "АНО «ВСКС»"),
    ("{{customer_full_name}}", "Полное наименование", "{{customer_full_name}}", "Автономная некоммерческая организация «ВСКС»"),
    ("{{customer_short_name}}", "Краткое наименование Заказчика (из карточки организации)", "{{customer_short_name}}", "АНО «ВСКС»"),
    ("{{customer_address}}", "Юридический адрес", "{{customer_address}}", "г. Москва, ул. Ленина, д. 1"),
    ("{{customer_postal_address}}", "Почтовый адрес", "{{customer_postal_address}}", "г. Москва, ул. Ленина, д. 1"),
    ("{{customer_inn}}", "ИНН Заказчика", "{{customer_inn}}", "7700000001"),
    ("{{customer_kpp}}", "КПП Заказчика", "{{customer_kpp}}", "770001001"),
    ("{{customer_ogrn}}", "ОГРН Заказчика", "{{customer_ogrn}}", "1027700000001"),
    ("{{customer_bank_name}}", "Банк Заказчика", "{{customer_bank_name}}", "ПАО «Сбербанк»"),
    ("{{customer_settlement_account}}", "Расчётный счёт Заказчика", "{{customer_settlement_account}}", "40703810400000000001"),
    ("{{customer_correspondent_account}}", "Корр. счёт Заказчика", "{{customer_correspondent_account}}", "30101810400000000225"),
    ("{{customer_bik}}", "БИК банка Заказчика", "{{customer_bik}}", "044525225"),
    ("{{customer_phone}}", "Телефон", "{{customer_phone}}", "+7 (495) 000-00-01"),
    ("{{customer_email}}", "Email", "{{customer_email}}", "info@vsks.ru"),
    ("{{customer_signatory}}", "Полная строка подписанта", "{{customer_signatory}}", "Президент Козеев Е.В."),
    ("{{customer_signatory_position}}", "Должность подписанта", "{{customer_signatory_position}}", "Президент"),
    ("{{customer_signatory_name}}", "ФИО подписанта (им.падеж)", "{{customer_signatory_name}}", "Козеев Евгений Викторович"),
    ("{{customer_signatory_name_genitive}}", "ФИО в род.падеже", "{{customer_signatory_name_genitive}}", "Козеева Евгения Викторовича"),
    ("{{customer_signatory_initials}}", "Фамилия + инициалы", "{{customer_signatory_initials}}", "Козеев Е.В."),
    ("{{customer_signatory_basis}}", "Основание полномочий", "{{customer_signatory_basis}}", "Устава"),
    ("{{contract_city}}", "Город заключения договора", "{{contract_city}}", "Москва"),
    # ── Контрагент ──
    ("", "КОНТРАГЕНТ (ИСПОЛНИТЕЛЬ)", "", ""),
    ("{{contractor_name}}", "Полное наименование контрагента", "{{contractor_name}}", "ООО «Ромашка»"),
    ("{{contractor_short_name}}", "Краткое наименование Подрядчика (из карточки контрагента)", "{{contractor_short_name}}", "ООО «Ромашка»"),
    ("{{contractor_org_type}}", "Тип организации", "{{contractor_org_type}}", "Юр.лицо"),
    ("{{contractor_inn}}", "ИНН контрагента", "{{contractor_inn}}", "7701234567"),
    ("{{contractor_kpp}}", "КПП контрагента", "{{contractor_kpp}}", "770101001"),
    ("{{contractor_ogrn}}", "ОГРН / ОГРНИП", "{{contractor_ogrn}}", "1027701234567"),
    ("{{contractor_address}}", "Юридический адрес", "{{contractor_address}}", "г. Москва, ул. Садовая, д. 5"),
    ("{{contractor_postal_address}}", "Почтовый адрес", "{{contractor_postal_address}}", "г. Москва, ул. Садовая, д. 5"),
    ("{{contractor_phone}}", "Телефон", "{{contractor_phone}}", "+7 (495) 111-22-33"),
    ("{{contractor_email}}", "E-mail", "{{contractor_email}}", "info@romashka.ru"),
    ("{{contractor_signatory}}", "ФИО подписанта", "{{contractor_signatory}}", "Сидоров Пётр Павлович"),
    ("{{contractor_signatory_basis}}", "Основание полномочий", "{{contractor_signatory_basis}}", "Устава"),
    ("{{contractor_signatory_position}}", "Должность подписанта", "{{contractor_signatory_position}}", "Директор"),
    ("{{contractor_signatory_name}}", "ФИО в им.падеже", "{{contractor_signatory_name}}", "Сидоров Пётр Павлович"),
    ("{{contractor_signatory_name_genitive}}", "ФИО в род.падеже", "{{contractor_signatory_name_genitive}}", "Сидорова Петра Павловича"),
    ("{{contractor_signatory_initials}}", "Фамилия + инициалы", "{{contractor_signatory_initials}}", "Сидоров П.П."),
    ("{{contractor_signatory_line}}", "ФИО + основание", "{{contractor_signatory_line}}", "Сидоров П.П., действующий на основании Устава"),
    ("{{contractor_ogrnip}}", "ОГРНИП (только для ИП)", "{{contractor_ogrnip}}", "304770000000001"),
    ("{{service_subject}}", "Предмет услуг (синоним subject)", "{{service_subject}}", "Оказание услуг связи"),
    ("{{contractor_settlement_account}}", "Расчётный счёт", "{{contractor_settlement_account}}", "40702810400000000002"),
    ("{{contractor_bank_name}}", "Наименование банка", "{{contractor_bank_name}}", "ПАО «Сбербанк»"),
    ("{{contractor_bank_details}}", "Реквизиты банка", "{{contractor_bank_details}}", "БИК 044525225, к/с 30101810400000000225"),
    ("{{contractor_bik}}", "БИК банка", "{{contractor_bik}}", "044525225"),
    ("{{contractor_correspondent_account}}", "Корреспондентский счёт", "{{contractor_correspondent_account}}", "30101810400000000225"),
    # ── Финансы ──
    ("", "ФИНАНСЫ", "", ""),
    ("{{total_nmcd}}", "НМЦД — начальная максимальная цена договора", "{{total_nmcd}}", "130 000,00 ₽"),
    ("{{total_nmck}}", "НМЦК (устаревшее)", "{{total_nmck}}", "130 000,00 ₽"),
    ("{{nmck}}", "НМЦК (синоним total_nmck)", "{{nmck}}", "130 000,00 ₽"),
    ("{{contract_price}}", "Цена договора", "{{contract_price}}", "130 000,00 ₽"),
    ("{{contract_price_num}}", "Цена без валюты", "{{contract_price_num}}", "130 000,00"),
    ("{{contract_price_words}}", "Цена прописью", "{{contract_price_words}}", "сто тридцать тысяч рублей 00 копеек"),
    ("{{economy}}", "Экономия", "{{economy}}", "5 000,00 ₽"),
    ("{{price_increase}}", "Увеличение цены", "{{price_increase}}", "0,00 ₽"),
    ("{{doc_amount}}", "Сумма закупки для документов (не зависит от стадии: договор → НМЦК/план → сумма позиций)", "{{doc_amount}}", "130 000,00 ₽"),
    ("{{doc_amount_num}}", "Сумма для документов без валюты", "{{doc_amount_num}}", "130 000,00"),
    ("{{doc_amount_words}}", "Сумма для документов прописью", "{{doc_amount_words}}", "сто тридцать тысяч рублей 00 копеек"),
    ("{{amount_is_planned}}", "true, если сумма плановая (договор ещё не заключён)", "{{amount_is_planned}}", "true"),
    ("{{amount_source_label}}", "Подпись источника суммы", "{{amount_source_label}}", "НМЦК (план)"),
    # ── НДС ──
    ("", "НДС", "", ""),
    ("{{vat_applicable}}", "Облагается НДС", "{{vat_applicable}}", "true"),
    ("{{vat_rate}}", "Ставка НДС", "{{vat_rate}}", "20"),
    ("{{vat_amount_num}}", "Сумма НДС цифрами", "{{vat_amount_num}}", "21 666,67"),
    ("{{vat_amount_words}}", "Сумма НДС прописью", "{{vat_amount_words}}", "двадцать одна тысяча 666 руб. 67 коп."),
    ("{{vat_exemption_article}}", "Статья освобождения от НДС", "{{vat_exemption_article}}", "ст. 149 НК РФ"),
    ("{{vat_info_line}}", "Готовая строка НДС", "{{vat_info_line}}", "В том числе НДС 20%: 21 666,67 руб."),
    # ── Договор ──
    ("", "ДОГОВОР", "", ""),
    ("{{contract_number}}", "Номер договора", "{{contract_number}}", "2026/42"),
    ("{{contract_date}}", "Дата договора", "{{contract_date}}", "15.01.2026"),
    ("{{contract_date_day}}", "День", "{{contract_date_day}}", "15"),
    ("{{contract_date_month}}", "Месяц прописью", "{{contract_date_month}}", "января"),
    ("{{contract_date_year}}", "Год", "{{contract_date_year}}", "2026"),
    ("{{execution_term}}", "Срок исполнения", "{{execution_term}}", "28.02.2026"),
    ("{{execution_term_changed}}", "Изменённый срок", "{{execution_term_changed}}", "31.03.2026"),
    ("{{delivery_date}}", "Дата поставки", "{{delivery_date}}", "20.01.2026"),
    ("{{country_origin}}", "Страна происхождения", "{{country_origin}}", "Российская Федерация"),
    ("{{service_name}}", "Предмет (синоним subject)", "{{service_name}}", "Оказание услуг связи"),
    ("{{service_name_gen}}", "Предмет в родительном падеже («Прошу осуществить закупку ...»)", "{{service_name_gen}}", "оказания услуг связи"),
    ("{{period_type}}", "Тип срока", "{{period_type}}", "date"),
    ("{{service_start_date}}", "Начало оказания услуг", "{{service_start_date}}", "15.01.2026"),
    ("{{service_end_date}}", "Конец оказания услуг", "{{service_end_date}}", "28.02.2026"),
    ("{{service_date}}", "Дата оказания (разовая)", "{{service_date}}", "20.01.2026"),
    ("{{third_party_involved}}", "Привлечение третьих лиц", "{{third_party_involved}}", "false"),
    # ── Phase 19: срок услуг расширенный ──
    ("", "СРОК УСЛУГ (Phase 19)", "", ""),
    ("{{service_term}}", "Готовая строка срока", "{{service_term}}", "с 15.01.2026 по 28.02.2026"),
    ("{{service_term_mode}}", "Режим срока", "{{service_term_mode}}", "range"),
    ("{{service_term_days}}", "Кол-во дней (для mode=duration)", "{{service_term_days}}", "30"),
    ("{{service_term_type}}", "Тип дней", "{{service_term_type}}", "calendar"),
    ("{{service_term_type_name}}", "Тип дней прописью", "{{service_term_type_name}}", "календарных"),
    ("{{service_deadline_date}}", "Срок \"до даты\"", "{{service_deadline_date}}", "28.02.2026"),
    # ── Phase 19: приём заявок ──
    ("", "ПРИЁМ ЗАЯВОК (Phase 19)", "", ""),
    ("{{submission_deadline_date}}", "Дата завершения приёма заявок", "{{submission_deadline_date}}", "10.01.2026"),
    ("{{submission_deadline_time}}", "Время завершения приёма заявок", "{{submission_deadline_time}}", "17:00"),
    ("{{submission_deadline_datetime}}", "Дата+время завершения", "{{submission_deadline_datetime}}", "10.01.2026 17:00"),
    # ── Phase 19: место доставки ──
    ("", "МЕСТО ДОСТАВКИ / ОКАЗАНИЯ УСЛУГ (Phase 19)", "", ""),
    ("{{delivery_location}}", "Место оказания услуг / доставки", "{{delivery_location}}", "г. Москва, ул. Ленина, д. 1"),
    # ── Phase 19: соглашение субсидии ──
    ("", "СОГЛАШЕНИЕ ПО СУБСИДИИ (Phase 19)", "", ""),
    ("{{subsidy_agreement_text}}", "Большой текст соглашения", "{{subsidy_agreement_text}}", "(многострочный текст)"),
    # ── Акт приёмки ──
    ("", "АКТ ПРИЁМКИ", "", ""),
    ("{{acceptance_doc_name}}", "Наименование акта", "{{acceptance_doc_name}}", "АКТ"),
    ("{{acceptance_doc_number}}", "Номер акта", "{{acceptance_doc_number}}", "001"),
    ("{{acceptance_doc_date}}", "Дата акта", "{{acceptance_doc_date}}", "28.02.2026"),
    ("{{acceptance_doc_amount}}", "Сумма акта", "{{acceptance_doc_amount}}", "130 000,00"),
    # ── Платёж ──
    ("", "ПЛАТЁЖ", "", ""),
    ("{{payment_doc_number}}", "Номер платёжного поручения", "{{payment_doc_number}}", "П-00125"),
    ("{{payment_doc_date}}", "Дата ПП", "{{payment_doc_date}}", "05.03.2026"),
    ("{{payment_amount}}", "Сумма платежа", "{{payment_amount}}", "130 000,00"),
    ("{{payment_federal}}", "В т.ч. федеральный бюджет", "{{payment_federal}}", "100 000,00"),
    # ── Инициатор ──
    ("", "ИНИЦИАТОР (для служебных записок)", "", ""),
    ("{{initiator_name}}", "ФИО инициатора", "{{initiator_name}}", "Иванов И.И."),
    ("{{initiator_role}}", "Должность инициатора", "{{initiator_role}}", "начальник отдела"),
    ("{{initiator_dept}}", "Отдел инициатора (если несколько — первый)", "{{initiator_dept}}", "Отдел спасательных операций"),
    # ── Согласующие (цикл) ──
    ("", "СОГЛАСУЮЩИЕ (цикл для таблицы)", "", ""),
    ("{%tr for a in approvers %} ... {%tr endfor %}", "Цикл по согласующим (в строке таблицы)", "{%tr for a in approvers %}<строка>{%tr endfor %}", "повторяется по числу согласующих"),
    ("{{a.num}}", "Порядковый номер (1, 2, 3...)", "{{a.num}}", "1"),
    ("{{a.role_name}}", "Должность согласующего", "{{a.role_name}}", "начальник отдела"),
    ("{{a.full_name}}", "ФИО согласующего", "{{a.full_name}}", "Петров П.П."),
    ("{{a.signature_img}}", "Электронная подпись (картинка)", "{{a.signature_img}}", "(изображение подписи)"),
    ("{{a.decided_date}}", "Дата подписания", "{{a.decided_date}}", "12.05.2026"),
    ("{{a.note}}", "Примечание (путь ФЭО)", "{{a.note}}", "Согласовано"),
    # ── Позиции (цикл) ──
    ("", "ПОЗИЦИИ ЗАКУПКИ (цикл для таблицы)", "", ""),
    ("{%tr for item in items %} ... {%tr endfor %}", "Цикл по позициям (в строке таблицы)", "{%tr for item in items %}<строка>{%tr endfor %}", "повторяется по числу позиций"),
    ("{{item.num}}", "Порядковый номер", "{{item.num}}", "1"),
    ("{{item.name}}", "Наименование товара/услуги", "{{item.name}}", "Услуги мобильной связи"),
    ("{{item.description}}", "Описание (из карточки продукта)", "{{item.description}}", "Корпоративная SIM-карта"),
    ("{{item.type}}", "Тип (товар/услуга)", "{{item.type}}", "услуга"),
    ("{{item.quantity}}", "Количество", "{{item.quantity}}", "10"),
    ("{{item.unit}}", "Единица измерения", "{{item.unit}}", "шт."),
    ("{{item.unit_price}}", "Цена за единицу", "{{item.unit_price}}", "13 000,00"),
    ("{{item.total_price}}", "Сумма строки", "{{item.total_price}}", "130 000,00"),
    ("{{item.photo}}", "Фото товара (картинка)", "{{item.photo}}", "(изображение)"),
    ("{{items_count}}", "Общее количество позиций", "{{items_count}}", "3"),
    ("{{item_names}}", "Перечень названий через запятую", "{{item_names}}", "Услуги связи, Интернет, Хостинг"),
    # ── Чеки (авансовый отчёт) ──
    ("", "ЧЕКИ — АВАНСОВЫЙ ОТЧЁТ (только для закупок с purchase_method=advance)", "", ""),
    ("{{receipts}}", "Список InlineImage чеков (ширина 6.5 см) — стандартная колонка", "{% for r in receipts %}{{ r }}{% endfor %}", "(изображения чеков 6.5 см)"),
    ("{{receipt_images}}", "Алиас receipts — список изображений чеков 6.5 см", "{% for img in receipt_images %}{{ img }}{% endfor %}", "(изображения чеков 6.5 см)"),
    ("{{receipts_small}}", "Чеки 4.5 см — для узких ячеек 2-колоночных таблиц", "{% for r in receipts_small %}{{ r }}{% endfor %}", "(маленькие чеки 4.5 см)"),
    ("{{receipts_full}}", "Чеки 14 см — для full-width страниц (один чек на страницу)", "{% for r in receipts_full %}{{ r }}{% endfor %}", "(полноширинные чеки 14 см)"),
    ("{{left_receipts}}", "Чёткые позиции чеков 4.5 см (1-й, 3-й, ...) — левая колонка таблицы 1×2", "{% for r in left_receipts %}{{ r }}{% endfor %}", "(чеки 1,3,5...)"),
    ("{{right_receipts}}", "Нечёткие позиции чеков 4.5 см (2-й, 4-й, ...) — правая колонка таблицы 1×2", "{% for r in right_receipts %}{{ r }}{% endfor %}", "(чеки 2,4,6...)"),
    # ── Родительный падеж (Phase 26-V) ──
    ("", "РОДИТЕЛЬНЫЙ ПАДЕЖ (Phase 26-V)", "", ""),
    ("{{initiator_name_gen}}", "ФИО инициатора в родительном падеже", "{{initiator_name_gen}}", "Иванова И.И."),
    ("{{initiator_position_gen}}", "Должность инициатора в родительном падеже", "{{initiator_position_gen}}", "начальника отдела"),
    ("{{initiator_dept_gen}}", "Отдел инициатора в родительном падеже (с авто-удалением дубля «Отдел» если он уже есть в должности)", "{{initiator_dept_gen}}", "мто"),
    ("{{responsible_name_gen}}", "ФИО ответственного в родительном падеже", "{{responsible_name_gen}}", "Петрова П.П."),
    ("{{responsible_position_gen}}", "Должность ответственного в родительном падеже", "{{responsible_position_gen}}", "главного специалиста"),
    ("{{a.full_name_gen}}", "ФИО согласующего в родительном падеже (внутри цикла)", "{{a.full_name_gen}}", "Сидорова С.С."),
    ("{{a.role_name_gen}}", "Должность согласующего в родительном падеже (внутри цикла)", "{{a.role_name_gen}}", "директора по правовым вопросам"),
    # ── Служебные ──
    ("", "СЛУЖЕБНЫЕ", "", ""),
    ("{{today}}", "Сегодняшняя дата", "{{today}}", "13.05.2026"),
    ("{{today_iso}}", "Дата ISO", "{{today_iso}}", "2026-05-13"),
]



@guide_router.get("/template-guide")
async def download_template_guide(
    current_user=Depends(get_current_user),
):
    """Download a .docx reference guide with all available template variables."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("Руководство по разметке шаблонов документов", 0)

    p = doc.add_paragraph(
        "Шаблоны документов используют синтаксис Jinja2 (docxtpl). "
        "Переменные заключаются в двойные фигурные скобки: {{variable}}. "
        "Для циклов используется синтаксис: {%tr for item in items %} ... {%tr endfor %}."
    )

    doc.add_heading("Доступные переменные", 1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Переменная"
    hdr[1].text = "Описание"
    hdr[2].text = "Пример записи"
    hdr[3].text = "Результат"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for var, desc, ex_t, ex_r in TEMPLATE_VARIABLES:
        if not var:
            # Section header row
            row = table.add_row().cells
            row[0].merge(row[3])
            p = row[0].paragraphs[0]
            run = p.add_run(desc)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            continue
        row = table.add_row().cells
        row[0].text = var
        row[1].text = desc
        row[2].text = ex_t
        row[3].text = ex_r

    doc.add_heading("Примеры использования в шаблоне Word", 1)

    doc.add_heading("Таблица согласующих", 2)
    doc.add_paragraph(
        '{%tr for a in approvers %}\n'
        '| {{a.num}} | {{a.role_name}} | {{a.full_name}} | {{a.signature_img}} | {{a.decided_date}} |\n'
        '{%tr endfor %}'
    )

    doc.add_heading("Таблица позиций закупки", 2)
    doc.add_paragraph(
        '{%tr for item in items %}\n'
        '| {{item.num}} | {{item.name}} | {{item.quantity}} | {{item.unit}} | {{item.unit_price}} | {{item.total_price}} |\n'
        '{%tr endfor %}'
    )

    doc.add_heading("Уровни субсидии (ФЭО)", 2)
    doc.add_paragraph(
        "Вставьте одну из этих переменных в нужное место шаблона Word:\n\n"
        "  Полный путь одной строкой:\n"
        "    {{feo_path}}\n"
        "  → Пример: «Персонал → Зарплата → Основной ФОТ»\n\n"
        "  Отдельные уровни:\n"
        "    Направление: {{feo_level_1}}\n"
        "    Тип:         {{feo_level_2}}\n"
        "    Категория:   {{feo_level_3}}\n\n"
        "Если закупка без ФЭО-категорий — переменные будут пустыми.\n"
        "Если у субсидии только 1 уровень — feo_level_2 и feo_level_3 пустые."
    )

    doc.add_heading("Условные блоки (НДС)", 2)
    doc.add_paragraph(
        '{% if vat_applicable %}\n'
        'В том числе НДС {{vat_rate}}%: {{vat_amount_num}} руб.\n'
        '{% else %}\n'
        'НДС не облагается {{vat_exemption_article}}\n'
        '{% endif %}'
    )

    doc.add_heading("Частые ошибки", 1)
    doc.add_paragraph(
        "1. Русский текст внутри {{ }} — ошибка: {{ contract_date г. }}\n"
        "   Правильно: {{ contract_date }} г.\n\n"
        "2. Word разбивает переменную на фрагменты — удалите и наберите заново\n\n"
        "3. Пустое значение — поле не заполнено в карточке закупки"
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote('Инструкция_по_шаблонам.docx', safe='-_.~')}"},
    )


@guide_router.get("/template-vars")
async def get_template_vars(current_user=Depends(get_current_user)):
    """Return all template variables with description, example template text, and example result."""
    result = []
    for entry in TEMPLATE_VARIABLES:
        var, desc, ex_t, ex_r = entry
        if not var:
            continue  # skip section headers
        result.append({
            "var": var,
            "description": desc,
            "example_template": ex_t,
            "example_result": ex_r,
        })
    return result
