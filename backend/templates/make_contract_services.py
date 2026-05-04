"""
Генератор шаблона «Договор оказания услуг» (contract_services.docx).

Запуск (один раз):
    py backend/templates/make_contract_services.py

Создаёт backend/templates/contract_services.docx — шаблон docxtpl с Jinja2
placeholder'ами для всех переменных из documents.py (Phase 23).

Цветовая схема:
    {{синие}}   — placeholder'ы (RGBColor #1E40AF, жирный)
    зелёный     — ветка True в {% if %}
    оранжевый   — ветка False / {% else %}
    серый курсив — маркеры {% if %} / {% else %} / {% endif %}
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(TEMPLATES_DIR, "contract_services.docx")

# ─── ЦВЕТА ───────────────────────────────────────────────────────────────────
COLOR_PLACEHOLDER = RGBColor(0x1E, 0x40, 0xAF)   # синий — {{var}}
COLOR_BRANCH_TRUE = RGBColor(0x15, 0x80, 0x3D)   # зелёный — ветка True
COLOR_BRANCH_FALSE = RGBColor(0xC2, 0x41, 0x0C)  # оранжевый — ветка False
COLOR_MARKER = RGBColor(0x6B, 0x72, 0x80)         # серый — {% %}

FONT_NAME = "Times New Roman"


# ─── HELPER'Ы ─────────────────────────────────────────────────────────────────

def _base_run(paragraph, text, size=12):
    """Создаёт run с базовым шрифтом."""
    r = paragraph.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    return r


def add_placeholder(para, text, size=12):
    """{{var}} — синий жирный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_PLACEHOLDER
    r.bold = True
    return r


def add_marker(para, text, size=12):
    """{% if %} / {% else %} / {% endif %} — серый курсив."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_MARKER
    r.italic = True
    return r


def add_branch_true(para, text, size=12):
    """Ветка True — зелёный обычный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_BRANCH_TRUE
    return r


def add_branch_false(para, text, size=12):
    """Ветка False — оранжевый обычный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_BRANCH_FALSE
    return r


def add_text(para, text, bold=False, size=12, italic=False):
    """Обычный чёрный текст."""
    r = _base_run(para, text, size)
    r.bold = bold
    r.italic = italic
    return r


# ─── УТИЛИТЫ ДОКУМЕНТА ───────────────────────────────────────────────────────

def page_margins(doc):
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def new_para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
             first_line_indent=None, line_spacing=13.8):
    """Создаёт пустой параграф с нужным форматированием."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    return p


def para_simple(doc, text="", bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0, space_after=6, italic=False):
    """Параграф с одним обычным run (для простого текста)."""
    p = new_para(doc, align=align, space_before=space_before, space_after=space_after)
    if text:
        add_text(p, text, bold=bold, size=size, italic=italic)
    return p


def section_title(doc, text, num=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    full = f"{num}. {text}" if num else text
    r = p.add_run(full)
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = Pt(12)
    return p


def clause_para(doc, num, space_before=2):
    """Возвращает параграф пункта с номером (жирный) и indent'ом."""
    p = new_para(doc, space_before=space_before, space_after=3,
                 first_line_indent=Cm(1))
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = Pt(12)
    return p


def clause_plain(doc, num, text, space_before=2):
    """Пункт с обычным текстом (без placeholder'ов)."""
    p = clause_para(doc, num, space_before)
    add_text(p, text)
    return p


# ─── СБОРКА ДОКУМЕНТА ────────────────────────────────────────────────────────

doc = Document()
page_margins(doc)

style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = Pt(12)

# ══════════════════════════════════════════════════════════════════════════════
# ЛЕГЕНДА
# ══════════════════════════════════════════════════════════════════════════════

p_leg_title = doc.add_paragraph()
p_leg_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_leg_title.paragraph_format.space_after = Pt(4)
p_leg_title.paragraph_format.space_before = Pt(0)
r_lt = p_leg_title.add_run("🎨 ЛЕГЕНДА — как читать этот шаблон")
r_lt.bold = True
r_lt.font.name = FONT_NAME
r_lt.font.size = Pt(11)
r_lt.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)  # тёмно-синий заголовок

# строки легенды
LEGEND_LINES = [
    ("placeholder", "{{синие переменные}}  — автоподстановка из БД (Заказчик / Исполнитель / Цена / Сроки)"),
    ("true",        "зелёный текст         — вариант 1 (например, «с НДС»)"),
    ("false",       "оранжевый текст       — вариант 2 (например, «НДС не облагается»)"),
    ("marker",      "серые {% маркеры %}   — технические условия Jinja2 (оставить как есть)"),
]

for kind, text in LEGEND_LINES:
    pl = doc.add_paragraph()
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pl.paragraph_format.space_before = Pt(0)
    pl.paragraph_format.space_after = Pt(2)
    pl.paragraph_format.first_line_indent = Cm(0.5)
    r = pl.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = Pt(10)
    if kind == "placeholder":
        r.font.color.rgb = COLOR_PLACEHOLDER
        r.bold = True
    elif kind == "true":
        r.font.color.rgb = COLOR_BRANCH_TRUE
    elif kind == "false":
        r.font.color.rgb = COLOR_BRANCH_FALSE
    elif kind == "marker":
        r.font.color.rgb = COLOR_MARKER
        r.italic = True

p_warn = doc.add_paragraph()
p_warn.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_warn.paragraph_format.space_before = Pt(4)
p_warn.paragraph_format.space_after = Pt(2)
p_warn.paragraph_format.first_line_indent = Cm(0.5)
r_warn = p_warn.add_run(
    "⚠️ После открытия в Word удалите ЭТУ секцию (от «ЛЕГЕНДА» до строки «—————») "
    "если хотите чистый шаблон. Сама разметка должна остаться нетронутой."
)
r_warn.font.name = FONT_NAME
r_warn.font.size = Pt(10)
r_warn.italic = True

p_ref = doc.add_paragraph()
p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_ref.paragraph_format.space_before = Pt(0)
p_ref.paragraph_format.space_after = Pt(2)
p_ref.paragraph_format.first_line_indent = Cm(0.5)
r_ref = p_ref.add_run(
    "Полный справочник переменных: SubsidiesView → Шаблоны → «Руководство по переменным»."
)
r_ref.font.name = FONT_NAME
r_ref.font.size = Pt(10)

p_sep = doc.add_paragraph()
p_sep.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sep.paragraph_format.space_before = Pt(2)
p_sep.paragraph_format.space_after = Pt(10)
r_sep = p_sep.add_run("—" * 40)
r_sep.font.name = FONT_NAME
r_sep.font.size = Pt(10)
r_sep.font.color.rgb = COLOR_MARKER

# ══════════════════════════════════════════════════════════════════════════════
# ЗАГОЛОВОК ДОГОВОРА
# ══════════════════════════════════════════════════════════════════════════════

p_title = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_text(p_title, "ДОГОВОР № ", bold=True, size=14)
add_placeholder(p_title, "{{contract_number}}", size=14)

p_sub = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_text(p_sub, "на оказание услуг", bold=True)

# Город и дата
p_city = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
add_text(p_city, "г. ")
add_placeholder(p_city, "{{contract_city}}")
add_text(p_city, "                                    «")
add_placeholder(p_city, "{{contract_date_day}}")
add_text(p_city, "» ")
add_placeholder(p_city, "{{contract_date_month}}")
add_text(p_city, " ")
add_placeholder(p_city, "{{contract_date_year}}")
add_text(p_city, " г.")

# ══════════════════════════════════════════════════════════════════════════════
# ПРЕАМБУЛА
# ══════════════════════════════════════════════════════════════════════════════

# Заказчик
p_pre = new_para(doc, space_after=6)
add_placeholder(p_pre, "{{customer_full_name}}")
add_text(p_pre, " (")
add_placeholder(p_pre, "{{customer_short_name}}")
add_text(p_pre, "), именуемое в дальнейшем «Заказчик», в лице ")
add_placeholder(p_pre, "{{customer_signatory_position}}")
add_text(p_pre, " ")
add_placeholder(p_pre, "{{customer_signatory_name_genitive}}")
add_text(p_pre, ", действующего на основании ")
add_placeholder(p_pre, "{{customer_signatory_basis}}")
add_text(p_pre, ", с одной стороны, и")

# Исполнитель — условный блок ИП / Юр.лицо
p_cond = new_para(doc, space_after=6)
add_marker(p_cond, "{%- if contractor_org_type == 'ИП' %}")
add_branch_true(p_cond, "Индивидуальный предприниматель ")
add_placeholder(p_cond, "{{contractor_signatory_name}}")
add_branch_true(p_cond, " (ИП ")
add_placeholder(p_cond, "{{contractor_short_name}}")
add_branch_true(p_cond, "), ИНН ")
add_placeholder(p_cond, "{{contractor_inn}}")
add_branch_true(p_cond, ", ОГРНИП ")
add_placeholder(p_cond, "{{contractor_ogrnip}}")
add_marker(p_cond, "{%- else %}")
add_placeholder(p_cond, "{{contractor_full_name}}")
add_branch_false(p_cond, " (")
add_placeholder(p_cond, "{{contractor_short_name}}")
add_branch_false(p_cond, "), в лице ")
add_placeholder(p_cond, "{{contractor_signatory_position}}")
add_branch_false(p_cond, " ")
add_placeholder(p_cond, "{{contractor_signatory_name_genitive}}")
add_branch_false(p_cond, ", действующего на основании ")
add_placeholder(p_cond, "{{contractor_signatory_basis}}")
add_marker(p_cond, "{%- endif %}")
add_text(p_cond, ", именуемое в дальнейшем «Исполнитель», с другой стороны, "
                 "совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем:")

# ══════════════════════════════════════════════════════════════════════════════
# 1. ПРЕДМЕТ ДОГОВОРА
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "ПРЕДМЕТ ДОГОВОРА", "1")

p_11 = clause_para(doc, "1.1")
add_text(p_11, "Исполнитель обязуется оказать услуги по ")
add_placeholder(p_11, "{{service_subject}}")
add_text(p_11, " (далее — Услуги), а Заказчик — принять и оплатить их. "
               "Содержание, объём и требования к Услугам определяются Техническим заданием "
               "(Приложение №1), являющимся неотъемлемой частью настоящего Договора.")

p_12 = clause_para(doc, "1.2")
add_text(p_12, "Срок оказания Услуг: ")
add_placeholder(p_12, "{{service_term}}")
add_text(p_12, ".")

p_13 = clause_para(doc, "1.3")
add_text(p_13, "Договор заключён в рамках реализации ")
add_placeholder(p_13, "{{subsidy_agreement_text}}")
add_text(p_13, ".")

# 1.4 — условный блок: третьи лица
p_14 = clause_para(doc, "1.4")
add_text(p_14, "Услуги оказываются Исполнителем ")
add_marker(p_14, "{% if third_party_involved %}")
add_branch_true(p_14, "с привлечением третьих лиц")
add_marker(p_14, "{% else %}")
add_branch_false(p_14, "своими силами, без привлечения третьих лиц")
add_marker(p_14, "{% endif %}")
add_text(p_14, ".")

# 1.5 — тип предмета договора (TODO: если появится поле is_goods)
# TODO: добавить условный блок {% if is_goods %}товары{% else %}услуги{% endif %}
# когда в модели Purchase появится поле is_goods или аналог.
# Пока используется placeholder {{contract_subject_kind}} без авто-заполнения.
p_15 = clause_para(doc, "1.5")
add_text(p_15, "Вид предмета договора: ")
add_placeholder(p_15, "{{contract_subject_kind}}")
add_text(p_15, ".")

# ══════════════════════════════════════════════════════════════════════════════
# 2. ПРАВА И ОБЯЗАННОСТИ СТОРОН
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "ПРАВА И ОБЯЗАННОСТИ СТОРОН", "2")

clause_plain(doc, "2.1",
             "Исполнитель обязан: оказать Услуги надлежащего качества в установленные сроки; "
             "соблюдать требования действующего законодательства РФ; предоставлять Заказчику "
             "по его требованию информацию о ходе оказания Услуг.")

clause_plain(doc, "2.2",
             "Исполнитель вправе запрашивать у Заказчика документы и сведения, необходимые "
             "для надлежащего оказания Услуг.")

clause_plain(doc, "2.3",
             "Заказчик обязан: своевременно передавать Исполнителю материалы и информацию, "
             "необходимые для оказания Услуг; принять результаты оказанных Услуг и оплатить "
             "их в соответствии с разделом 4 настоящего Договора.")

clause_plain(doc, "2.4",
             "Заказчик вправе осуществлять контроль за ходом и качеством оказания Услуг, "
             "не вмешиваясь в оперативно-хозяйственную деятельность Исполнителя.")

# ══════════════════════════════════════════════════════════════════════════════
# 3. ПОРЯДОК ПРИЁМКИ УСЛУГ
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "ПОРЯДОК ПРИЁМКИ УСЛУГ", "3")

clause_plain(doc, "3.1",
             "По окончании оказания Услуг Исполнитель представляет Заказчику Акт сдачи-приёмки "
             "оказанных услуг (далее — Акт).")

clause_plain(doc, "3.2",
             "Заказчик рассматривает Акт в течение 5 (пяти) рабочих дней с момента его получения. "
             "В случае обнаружения недостатков Заказчик направляет Исполнителю мотивированный "
             "отказ с перечнем замечаний.")

clause_plain(doc, "3.3",
             "Исполнитель устраняет выявленные недостатки в согласованные Сторонами сроки "
             "и представляет Акт повторно.")

clause_plain(doc, "3.4",
             "Датой приёмки Услуг считается дата подписания Акта уполномоченными представителями "
             "обеих Сторон.")

# ══════════════════════════════════════════════════════════════════════════════
# 4. ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЁТОВ
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЁТОВ", "4")

# 4.1 — условный блок: НДС
p_41 = clause_para(doc, "4.1")
add_text(p_41, "Цена Договора составляет ")
add_placeholder(p_41, "{{contract_price_num}}")
add_text(p_41, " (")
add_placeholder(p_41, "{{contract_price_words}}")
add_text(p_41, ") рублей. ")
add_marker(p_41, "{% if vat_applicable %}")
add_branch_true(p_41, "В том числе НДС ")
add_placeholder(p_41, "{{vat_rate}}")
add_branch_true(p_41, "%: ")
add_placeholder(p_41, "{{vat_amount_num}}")
add_branch_true(p_41, " (")
add_placeholder(p_41, "{{vat_amount_words}}")
add_branch_true(p_41, ") рублей.")
add_marker(p_41, "{% else %}")
add_branch_false(p_41, "НДС не облагается на основании ")
add_placeholder(p_41, "{{vat_exemption_article}}")
add_branch_false(p_41, ".")
add_marker(p_41, "{% endif %}")

clause_plain(doc, "4.2",
             "Оплата производится в безналичной форме путём перечисления денежных средств "
             "на расчётный счёт Исполнителя, указанный в разделе 9 настоящего Договора, "
             "в течение 30 (тридцати) календарных дней с даты подписания Акта.")

clause_plain(doc, "4.3",
             "Цена Договора является фиксированной и изменению не подлежит, за исключением "
             "случаев, предусмотренных действующим законодательством Российской Федерации.")

# ══════════════════════════════════════════════════════════════════════════════
# 5. ОТВЕТСТВЕННОСТЬ СТОРОН
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "ОТВЕТСТВЕННОСТЬ СТОРОН", "5")

clause_plain(doc, "5.1",
             "За нарушение сроков оказания Услуг Исполнитель уплачивает Заказчику неустойку "
             "в размере 0,1% от цены Договора за каждый день просрочки.")

clause_plain(doc, "5.2",
             "За нарушение сроков оплаты Заказчик уплачивает Исполнителю пени "
             "в размере 1/300 ставки рефинансирования Банка России за каждый день просрочки.")

clause_plain(doc, "5.3",
             "Стороны освобождаются от ответственности за частичное или полное неисполнение "
             "обязательств, если оно явилось следствием обстоятельств непреодолимой силы "
             "(форс-мажор), возникших после заключения настоящего Договора.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. СРОК ДЕЙСТВИЯ И ПОРЯДОК РАСТОРЖЕНИЯ
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "СРОК ДЕЙСТВИЯ И ПОРЯДОК РАСТОРЖЕНИЯ ДОГОВОРА", "6")

clause_plain(doc, "6.1",
             "Договор вступает в силу с момента подписания Сторонами и действует до полного "
             "исполнения Сторонами принятых на себя обязательств.")

clause_plain(doc, "6.2",
             "Договор может быть расторгнут по соглашению Сторон, а также в одностороннем "
             "порядке по основаниям, предусмотренным действующим законодательством РФ.")

# ══════════════════════════════════════════════════════════════════════════════
# 7. РАЗРЕШЕНИЕ СПОРОВ
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "РАЗРЕШЕНИЕ СПОРОВ", "7")

clause_plain(doc, "7.1",
             "Все споры и разногласия, возникающие в связи с исполнением настоящего Договора, "
             "разрешаются путём переговоров. При недостижении согласия — в Арбитражном суде "
             "г. Москвы в соответствии с действующим законодательством РФ.")

# ══════════════════════════════════════════════════════════════════════════════
# 8. ПРОЧИЕ УСЛОВИЯ
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "ПРОЧИЕ УСЛОВИЯ", "8")

clause_plain(doc, "8.1",
             "Договор составлен в двух экземплярах, имеющих равную юридическую силу, "
             "по одному для каждой из Сторон.")

clause_plain(doc, "8.2",
             "Любые изменения и дополнения к Договору действительны только в случае их "
             "составления в письменной форме и подписания уполномоченными представителями "
             "обеих Сторон.")

clause_plain(doc, "8.3",
             "Приложение №1 (Техническое задание) является неотъемлемой частью Договора.")

# ══════════════════════════════════════════════════════════════════════════════
# 9. АДРЕСА И РЕКВИЗИТЫ СТОРОН
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "АДРЕСА И РЕКВИЗИТЫ СТОРОН", "9")

# Двухколоночная таблица реквизитов
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"

from docx.oxml import OxmlElement as _OE
from docx.oxml.ns import qn as _qn


def set_col_width(cell, cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = _OE("w:tcW")
    tcW.set(_qn("w:w"), str(int(cm * 567)))
    tcW.set(_qn("w:type"), "dxa")
    tcPr.append(tcW)


LEFT = tbl.rows[0].cells[0]
RIGHT = tbl.rows[0].cells[1]
set_col_width(LEFT, 8.5)
set_col_width(RIGHT, 8.5)


def req_line(cell, para_obj, text, bold=False, is_placeholder=False, size=10):
    """Добавляет run в ячейку таблицы реквизитов."""
    r = para_obj.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    r.bold = bold
    if is_placeholder:
        r.font.color.rgb = COLOR_PLACEHOLDER
        r.bold = True
    return r


def req_block_colored(cell, items):
    """
    items — список (text, is_placeholder, bold) или строки.
    Каждый элемент = отдельная строка в ячейке.
    item может быть строкой (обычный текст) или tuple (text, is_placeholder, bold).
    Специальные токены 'IF_IP_START', 'ELSE', 'ENDIF' для маркеров.
    """
    para_idx = 0
    for item in items:
        if para_idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

        if isinstance(item, str):
            r = p.add_run(item)
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            if para_idx == 0:
                r.bold = True
        elif isinstance(item, tuple):
            kind = item[0]
            if kind == "marker":
                r = p.add_run(item[1])
                r.font.name = FONT_NAME
                r.font.size = Pt(9)
                r.font.color.rgb = COLOR_MARKER
                r.italic = True
            elif kind == "branch_true":
                # несколько runs в одной строке
                for seg in item[1]:
                    if seg[0] == "p":
                        r = p.add_run(seg[1])
                        r.font.name = FONT_NAME
                        r.font.size = Pt(10)
                        r.font.color.rgb = COLOR_PLACEHOLDER
                        r.bold = True
                    else:
                        r = p.add_run(seg[1])
                        r.font.name = FONT_NAME
                        r.font.size = Pt(10)
                        r.font.color.rgb = COLOR_BRANCH_TRUE
            elif kind == "branch_false":
                for seg in item[1]:
                    if seg[0] == "p":
                        r = p.add_run(seg[1])
                        r.font.name = FONT_NAME
                        r.font.size = Pt(10)
                        r.font.color.rgb = COLOR_PLACEHOLDER
                        r.bold = True
                    else:
                        r = p.add_run(seg[1])
                        r.font.name = FONT_NAME
                        r.font.size = Pt(10)
                        r.font.color.rgb = COLOR_BRANCH_FALSE
            elif kind == "ph":
                r = p.add_run(item[1])
                r.font.name = FONT_NAME
                r.font.size = Pt(10)
                r.font.color.rgb = COLOR_PLACEHOLDER
                r.bold = True
        para_idx += 1


# ЗАКАЗЧИК — реквизиты
LEFT_ITEMS = [
    "ЗАКАЗЧИК:",
    ("ph", "{{customer_full_name}}"),
    ("ph", "Адрес: {{customer_address}}"),
    ("ph", "Почтовый адрес: {{customer_postal_address}}"),
    ("ph", "ИНН: {{customer_inn}} / КПП: {{customer_kpp}}"),
    ("ph", "ОГРН: {{customer_ogrn}}"),
    ("ph", "Банк: {{customer_bank_name}}"),
    ("ph", "БИК: {{customer_bik}}"),
    ("ph", "р/с: {{customer_settlement_account}}"),
    ("ph", "к/с: {{customer_correspondent_account}}"),
    ("ph", "Тел.: {{customer_phone}}"),
    ("ph", "E-mail: {{customer_email}}"),
    "",
    ("ph", "{{customer_signatory_position}}"),
    ("ph", "________________ / {{customer_signatory_initials}}"),
    "М.П.",
]

req_block_colored(LEFT, LEFT_ITEMS)

# ИСПОЛНИТЕЛЬ — реквизиты с условным блоком ИП / ООО
RIGHT_ITEMS = [
    "ИСПОЛНИТЕЛЬ:",
    ("marker", "{%- if contractor_org_type == 'ИП' %}"),
    ("branch_true", [("t", "ИП "), ("p", "{{contractor_short_name}}")]),
    ("branch_true", [("t", "ИНН: "), ("p", "{{contractor_inn}}")]),
    ("branch_true", [("t", "ОГРНИП: "), ("p", "{{contractor_ogrnip}}")]),
    ("marker", "{%- else %}"),
    ("branch_false", [("p", "{{contractor_full_name}}")]),
    ("branch_false", [("t", "ИНН: "), ("p", "{{contractor_inn}}"), ("t", " / КПП: "), ("p", "{{contractor_kpp}}")]),
    ("branch_false", [("t", "ОГРН: "), ("p", "{{contractor_ogrn}}")]),
    ("marker", "{%- endif %}"),
    ("ph", "Адрес: {{contractor_address}}"),
    ("ph", "Банк: {{contractor_bank_name}}"),
    ("ph", "БИК: {{contractor_bik}}"),
    ("ph", "р/с: {{contractor_settlement_account}}"),
    ("ph", "к/с: {{contractor_correspondent_account}}"),
    ("ph", "Тел.: {{contractor_phone}}"),
    ("ph", "E-mail: {{contractor_email}}"),
    "",
    ("ph", "{{contractor_signatory_position}}"),
    ("ph", "________________ / {{contractor_signatory_initials}}"),
    "М.П.",
]

req_block_colored(RIGHT, RIGHT_ITEMS)

# Приложение
doc.add_paragraph()
p_app = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=6)
add_text(p_app, "Приложение №1: Техническое задание (на __ л.).")

# ── СОХРАНЕНИЕ ───────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"OK: {OUTPUT}")
