from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload
from app.database import get_db
from app.models.contract import Contract
from app.models.contract_subsidy import ContractSubsidy
from app.models.purchase import Purchase
from app.models.contractor import Contractor
from app.schemas.schemas import ContractCreate, ContractOut, ContractSubsidyOut, ContractUpdateResponse, ContractSyncWarnings
from app.auth.jwt import get_current_user, require_role, get_org_filter, ADMIN_ROLES, MANAGER_ROLES, ALL_ROLES
from app.auth.permissions import require_tab, require_action
from app.auth.visibility import build_visibility_clause, get_visible_subsidy_ids
from app.models.subsidy import Subsidy
from typing import List, Optional, Literal
from decimal import Decimal
from io import BytesIO
from datetime import date as date_type
import calendar
from pydantic import BaseModel, Field
from app.routers.purchase_budget import _assign_framework_seq, FRAMEWORK_TYPES

async def _enrich_contract_from_purchases(c: Contract, db: AsyncSession) -> int:
    """Auto-fill ПУСТЫЕ optional fields на Contract из любой связанной Purchase.

    Returns: count of fields filled (для diag). 0 если нечего обогащать.

    Заполняет только NULL поля — не перезаписывает то что user явно ввёл.
    Полей: subject, max_amount, start_date, end_date, purchase_method, item_type, contractor_id.
    """
    from app.models.purchase import Purchase as _P

    # Phase 27.1.12: добавлен contractor_id в проверку all-filled
    # Если все поля заполнены — выйти
    if all([
        c.subject, c.max_amount is not None,
        c.start_date, c.end_date,
        c.purchase_method, c.item_type,
        c.contractor_id is not None,  # NEW
    ]):
        return 0

    p = (await db.execute(
        select(_P).where(_P.contract_id == c.id).limit(1)
    )).scalar_one_or_none()
    if not p:
        return 0

    filled = 0
    # Phase 27.1.12: заполнить contractor_id из Purchase (только если NULL)
    if c.contractor_id is None and p.contractor_id is not None:
        c.contractor_id = p.contractor_id
        filled += 1

    if not c.subject:
        v = p.subject or str(p.purchase_number or '')
        if v:
            c.subject = v
            filled += 1
    # 27.4-19: framework_cumulative (накопительный) НЕ имеет max_amount —
    # пропускаем enrich этого поля чтобы не воскрешать «Превышен лимит» после 27.4-17.
    if c.max_amount is None and c.contract_type != 'framework_cumulative':
        v = p.total_nmck or p.contract_price or p.planned_total_price
        if v is not None:
            c.max_amount = v
            filled += 1
    if not c.start_date and p.contract_date:
        c.start_date = p.contract_date
        filled += 1
    if not c.end_date and p.execution_term:
        c.end_date = p.execution_term
        filled += 1
    if not c.purchase_method:
        c.purchase_method = p.purchase_method if p.purchase_method in ('single', 'competitive') else 'single'
        filled += 1
    if not c.item_type:
        c.item_type = p.item_type or 'товар'
        filled += 1
    return filled


def _framework_approval_state(approval_status: Optional[str]) -> Optional[str]:
    """Purchase.approval_status рамочной ГОЛОВЫ → упрощённое состояние для
    реестра договоров (ContractOut.approval_state), которым фронт решает,
    затемнять ли строку. approved — согласовано; pending — в работе или
    отклонено (ждёт повторного согласования); None — головы ещё нет или
    цепочка не строилась (нет руководителя организации, см.
    purchases.py::_build_framework_chain_approvals) — тогда фронт не затемняет."""
    if approval_status == "approved":
        return "approved"
    if approval_status in ("in_progress", "rejected"):
        return "pending"
    return None


RU_MONTHS_GEN = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]
RU_MONTHS_NOM = [
    "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь",
]


class MonthlyStagesRequest(BaseModel):
    count: int = Field(..., ge=1, le=36)
    amount_per_stage: float
    start_year: int
    start_month: int = Field(..., ge=1, le=12)
    deadline_kind: Literal["end_of_month", "specific_day"]
    deadline_day: Optional[int] = Field(None, ge=1, le=31)
    subject_template: str = "{contract_name} — {month_name} {year}"
    subsidy_id: int
    feo_category_id: Optional[int] = None
    is_likely_needed: bool = True

try:
    from openpyxl import load_workbook as _load_workbook
except ImportError:
    _load_workbook = None

try:
    import pdfplumber as _pdfplumber
except ImportError:
    _pdfplumber = None

router = APIRouter(prefix="/api/contracts", tags=["contracts"])

@router.get("/", response_model=List[ContractOut])
async def list_contracts(
    subsidy_id: Optional[int] = Query(None),
    contract_type: Optional[str] = Query(None),
    purchase_method: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    contractor_id: Optional[int] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    q = select(Contract).options(
        selectinload(Contract.contractor),
        selectinload(Contract.subsidy),
        selectinload(Contract.extra_subsidies).selectinload(ContractSubsidy.subsidy),
    ).order_by(Contract.id.desc())
    if subsidy_id is not None:
        q = q.where(Contract.subsidy_id == subsidy_id)
    if contract_type is not None:
        q = q.where(Contract.contract_type == contract_type)
    if purchase_method is not None:
        q = q.where(Contract.purchase_method == purchase_method)
    if status is not None:
        q = q.where(Contract.status == status)
    if contractor_id is not None:
        q = q.where(Contract.contractor_id == contractor_id)
    vis = await get_visible_subsidy_ids(current_user, db, "contracts")
    if vis is not None:
        q = q.where(Contract.subsidy_id.in_(vis))
    # Phase 28 Bundle 3: user-level visibility filter (was missing — data leak fix).
    clause = await build_visibility_clause(current_user, db, 'contract')
    if clause is not None:
        q = q.where(clause)
    result = await db.execute(q)
    contracts = result.scalars().all()
    out = []
    for c in contracts:
        # phase26-l-2: COALESCE chain contract_price → planned_total_price → total_nmck
        # чтобы legacy-закупки с NULL contract_price тоже учитывались в total_ordered.
        ordered_result = await db.execute(
            select(func.coalesce(
                func.sum(func.coalesce(Purchase.contract_price, Purchase.planned_total_price, Purchase.total_nmck, 0)),
                0
            ))
            .where(Purchase.contract_id == c.id)
        )
        total_ordered = ordered_result.scalar() or Decimal("0")
        # 27.4-18: total_delivered — приоритет JSONB acceptance_docs[].amount
        # (после Phase 26-H legacy delivery_payment_amount/acceptance_doc_amount = NULL
        # для всех новых закупок, агрегат показывал 0 при реально поставленных).
        delivered_rows = (await db.execute(
            select(Purchase.delivery_payment_amount, Purchase.acceptance_docs)
            .where(Purchase.contract_id == c.id)
        )).all()
        total_delivered = Decimal("0")
        for dpa, docs in delivered_rows:
            if dpa is not None:
                total_delivered += Decimal(str(dpa))
                continue
            if docs:
                try:
                    docs_list = docs if isinstance(docs, list) else []
                    for doc_entry in docs_list:
                        if isinstance(doc_entry, dict) and doc_entry.get('amount') is not None:
                            total_delivered += Decimal(str(doc_entry['amount']))
                except Exception:
                    pass
        # SUM(payment_amount) — total paid (only for paid-status purchases)
        paid_result = await db.execute(
            select(func.coalesce(func.sum(Purchase.payment_amount), 0))
            .where(Purchase.contract_id == c.id, Purchase.status == "paid")
        )
        total_paid = paid_result.scalar() or Decimal("0")

        d = ContractOut.model_validate(c)
        d.total_ordered = total_ordered
        d.total_delivered = total_delivered
        d.total_paid = total_paid
        d.total_payment = total_delivered  # legacy compat
        # framework_cumulative — без предельной суммы: остаток не считается (нет лимита)
        if c.contract_type == 'framework_cumulative' and not c.max_amount:
            d.remaining_ordered = None   # нет лимита — нет отрицательного остатка
            d.remaining = None
        else:
            max_amt = c.max_amount or Decimal("0")
            d.remaining_ordered = max_amt - total_ordered              # сколько ещё можно заказать
            d.remaining = d.remaining_ordered  # legacy: now shows ordered remaining
        d.remaining_delivered = total_ordered - total_delivered         # заказано, но не поставлено
        d.remaining_paid = total_delivered - total_paid                 # поставлено, но не оплачено
        if c.contractor:
            d.contractor_name = c.contractor.name
            d.contractor_inn = c.contractor.inn
        elif not d.contractor_name:
            # Phase 27.1.4: fallback — pull contractor from linked purchases
            # (advance imports без явного Contract.contractor_id)
            from app.models.purchase import Purchase as _P
            from app.models.contractor import Contractor as _C
            row = (await db.execute(
                select(_C.name, _C.inn).select_from(_P).join(_C, _C.id == _P.contractor_id)
                .where(_P.contract_id == c.id, _P.contractor_id.isnot(None))
                .limit(1)
            )).first()
            if row:
                d.contractor_name, d.contractor_inn = row
        if c.subsidy:
            d.subsidy_name = c.subsidy.name
        d.extra_subsidies = [
            ContractSubsidyOut(id=es.id, subsidy_id=es.subsidy_id, subsidy_name=es.subsidy.name if es.subsidy else None)
            for es in (c.extra_subsidies or [])
        ]
        # Владелец (2026-09-02): «рамочный договор без закупок внутри должен быть
        # визуально помечен, пока не согласован» — approval_state вычисляется из
        # Purchase.approval_status привязанной рамочной ГОЛОВЫ (is_framework_head,
        # см. purchases.py), отдельной колонки на Contract не заводим.
        if c.contract_type in FRAMEWORK_TYPES:
            head_status = (await db.execute(
                select(Purchase.approval_status).where(
                    Purchase.contract_id == c.id,
                    Purchase.parent_purchase_id.is_(None),
                    Purchase.purchase_contract_type.in_(FRAMEWORK_TYPES),
                ).limit(1)
            )).scalar_one_or_none()
            d.approval_state = _framework_approval_state(head_status)
        out.append(d)
    return out

@router.post("/", response_model=ContractOut)
async def create_contract(
    data: ContractCreate,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_tab('contracts')),
):
    from app.services.subsidy_draft_guard import assert_subsidy_approved_for_binding
    await assert_subsidy_approved_for_binding(db, data.subsidy_id)
    for _sid in (data.extra_subsidy_ids or []):
        await assert_subsidy_approved_for_binding(db, _sid)

    # Duplicate check: same number + contractor + subsidy (org) + date
    if data.contractor_id and data.number:
        dup_q = select(Contract).where(
            Contract.number == data.number,
            Contract.contractor_id == data.contractor_id,
        )
        # Same org (subsidy) — different orgs can have same contract number
        if data.subsidy_id:
            dup_q = dup_q.where(Contract.subsidy_id == data.subsidy_id)
        if data.date:
            dup_q = dup_q.where(Contract.date == data.date)
        dup = (await db.execute(dup_q)).scalar_one_or_none()
        if dup:
            raise HTTPException(
                409,
                f"Договор с таким номером, контрагентом и датой уже существует (ID {dup.id})"
            )
    extra_ids = data.extra_subsidy_ids or []
    dump = data.model_dump(exclude={"extra_subsidy_ids"})
    c = Contract(**dump)
    # 27.4-17: framework_cumulative (накопительный) НЕ имеет предельной суммы —
    # обнуляем max_amount чтобы не получался ложный «Превышен лимит».
    if c.contract_type == 'framework_cumulative':
        c.max_amount = None
    db.add(c)
    await db.flush()

    # Phase 27.1.6: auto-fill empty fields from linked purchases (если уже привязаны)
    await _enrich_contract_from_purchases(c, db)

    for sid in extra_ids:
        db.add(ContractSubsidy(contract_id=c.id, subsidy_id=sid))
    await db.commit()
    await db.refresh(c)
    result = await db.execute(
        select(Contract).options(
            selectinload(Contract.extra_subsidies).selectinload(ContractSubsidy.subsidy)
        ).where(Contract.id == c.id)
    )
    c = result.scalar_one()
    d = ContractOut.model_validate(c)
    d.extra_subsidies = [
        ContractSubsidyOut(id=es.id, subsidy_id=es.subsidy_id, subsidy_name=es.subsidy.name if es.subsidy else None)
        for es in c.extra_subsidies
    ]
    return d

@router.put("/{cid}", response_model=ContractUpdateResponse)
async def update_contract(cid: int, data: ContractCreate, db: AsyncSession = Depends(get_db), current_user=Depends(require_tab('contracts'))):
    result = await db.execute(select(Contract).where(Contract.id == cid))
    c = result.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Not found")

    from app.services.subsidy_draft_guard import assert_subsidy_approved_for_binding
    await assert_subsidy_approved_for_binding(db, data.subsidy_id)
    for _sid in (data.extra_subsidy_ids or []):
        await assert_subsidy_approved_for_binding(db, _sid)

    old_number = c.number
    old_date = c.date
    old_contractor_id = c.contractor_id
    # Duplicate check on update (exclude self, same org)
    if data.contractor_id and data.number:
        dup_q = select(Contract).where(
            Contract.number == data.number,
            Contract.contractor_id == data.contractor_id,
            Contract.id != cid,
        )
        if data.subsidy_id:
            dup_q = dup_q.where(Contract.subsidy_id == data.subsidy_id)
        if data.date:
            dup_q = dup_q.where(Contract.date == data.date)
        dup = (await db.execute(dup_q)).scalar_one_or_none()
        if dup:
            raise HTTPException(409, f"Договор с таким номером и контрагентом уже существует (ID {dup.id})")
    extra_ids = data.extra_subsidy_ids or []
    for k, v in data.model_dump(exclude={"extra_subsidy_ids"}).items():
        setattr(c, k, v)
    # 27.4-17: framework_cumulative (накопительный) НЕ имеет предельной суммы.
    # Обнуляем max_amount даже если пользователь его прислал — иначе ложный «Превышен лимит».
    if c.contract_type == 'framework_cumulative':
        c.max_amount = None

    # Phase 31-04: cascade sync to linked purchases using model_fields_set
    # (distinguishes «not sent» from «explicitly set to None» — pitfall 6)
    set_fields = data.model_fields_set
    cascade_number = 'number' in set_fields
    cascade_date = 'date' in set_fields
    cascade_contractor = 'contractor_id' in set_fields

    n_updated_purchases = 0
    warnings = ContractSyncWarnings()

    if cascade_number or cascade_date or cascade_contractor:
        linked_result = await db.execute(select(Purchase).where(Purchase.contract_id == cid))
        linked_purchases = linked_result.scalars().all()
        n_updated_purchases = len(linked_purchases)

        # Import helper from plan 31-01 (entity change tracking, T-31-04-02)
        from app.routers.entity_changes import record_entity_changes as _rec_changes

        for p in linked_purchases:
            old_vals: dict = {}
            new_vals: dict = {}
            if cascade_number and data.number != old_number:
                old_vals['contract_number'] = p.contract_number
                p.contract_number = data.number
                new_vals['contract_number'] = p.contract_number
            if cascade_date:
                # Always cascade date when in set_fields (includes explicit None = clear)
                old_vals['contract_date'] = p.contract_date
                p.contract_date = data.date
                new_vals['contract_date'] = p.contract_date
            if cascade_contractor and data.contractor_id != old_contractor_id:
                old_vals['contractor_id'] = p.contractor_id
                p.contractor_id = data.contractor_id
                new_vals['contractor_id'] = p.contractor_id

            # Record entity changes for audit trail (batch, committed after main commit)
            if old_vals:
                await _rec_changes(
                    db=db,
                    entity_type='purchase',
                    entity_id=p.id,
                    changed_by_id=current_user.id,
                    changed_by_name=getattr(current_user, 'full_name', None) or getattr(current_user, 'username', None),
                    old_values=old_vals,
                    new_values=new_vals,
                )

        # Warnings: amount_over_max
        if c.max_amount is not None and c.max_amount > 0:
            total_purchases = sum(p.planned_total_price or 0 for p in linked_purchases)
            if total_purchases > c.max_amount:
                warnings.amount_over_max = True

        # Warnings: date_out_of_validity (purchases with execution_term outside contract range)
        out_of_validity: list[int] = []
        for p in linked_purchases:
            if p.execution_term:
                if c.start_date and p.execution_term < c.start_date:
                    out_of_validity.append(p.id)
                elif c.end_date and p.execution_term > c.end_date:
                    out_of_validity.append(p.id)
        warnings.date_out_of_validity = out_of_validity

    # Replace extra subsidies
    old_extras = await db.execute(select(ContractSubsidy).where(ContractSubsidy.contract_id == cid))
    for es in old_extras.scalars().all():
        await db.delete(es)
    for sid in extra_ids:
        db.add(ContractSubsidy(contract_id=cid, subsidy_id=sid))

    # Phase 27.1.6: auto-fill пустые fields из связанных Purchase
    await _enrich_contract_from_purchases(c, db)

    await db.commit()
    result2 = await db.execute(
        select(Contract).options(
            selectinload(Contract.extra_subsidies).selectinload(ContractSubsidy.subsidy)
        ).where(Contract.id == cid)
    )
    c = result2.scalar_one()
    d = ContractOut.model_validate(c)
    d.extra_subsidies = [
        ContractSubsidyOut(id=es.id, subsidy_id=es.subsidy_id, subsidy_name=es.subsidy.name if es.subsidy else None)
        for es in c.extra_subsidies
    ]
    return ContractUpdateResponse(
        contract=d,
        n_updated_purchases=n_updated_purchases,
        warnings=warnings,
    )

def _build_approval_purchase_fields(
    contract: Contract, assigned_user_id: int, responsible_person: Optional[str] = None
) -> dict:
    """Собирает kwargs для Purchase() — «рамочная голова» из уже сохранённого
    Contract. Чистая функция без обращений к БД — вынесена отдельно, чтобы
    маппинг полей был юнит-тестируем без живой сессии
    (см. tests/test_contract_approval_purchase.py).

    Пустой contract.number (falsy) сознательно передаётся как None, а не
    подменяется — при первом формировании документа сработает уже готовая
    логика присвоения временного номера «ВРЕМ-…»
    (is_framework_head + _generate_temp_contract_number в purchases.py).

    responsible_person — ФИО текущего пользователя (если не пусто):
    предзаполняет «Ответственного исполнителя» листа согласования, чтобы у
    свежесозданной рамочной головы клетка не оставалась пустой.
    """
    fields = dict(
        contract_id=contract.id,
        contractor_id=contract.contractor_id,
        subsidy_id=contract.subsidy_id,
        subject=contract.subject,
        contract_number=(contract.number or None),
        contract_date=contract.date,
        contract_price=contract.max_amount,
        purchase_contract_type=contract.contract_type,
        purchase_method=contract.purchase_method,
        parent_purchase_id=None,
        # Начальный статус — как у обычной новой закупки (Purchase.status /
        # PurchaseCreate.status default в purchases.py).
        status="wishes",
        # Без явного владельца закупка невидима для рядового автора в списке
        # (list_purchases фильтрует по visible_user_ids) — как в create_purchase.
        assigned_user_id=assigned_user_id,
    )
    if responsible_person:
        fields["responsible_person"] = responsible_person
    return fields


@router.post("/{contract_id}/approval-purchase")
async def get_or_create_approval_purchase(
    contract_id: int,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_tab('contracts')),
):
    """Найти или создать «рамочную голову» — Purchase, привязанную к этому
    рамочному договору — через которую уже доступны согласование
    необходимости, печать договора и листа согласования.

    Диагноз (2026-09-01): реестр «Договоры» (Contract) — отдельная сущность
    без своего согласования/документооборота; вся эта машинерия живёт на
    Purchase (см. /api/purchases/{pid}/documents/{doc_type},
    purchase_transitions.py). У РАМОЧНЫХ договоров, созданных прямо в реестре
    «Договоры», закупка не заводится автоматически — поэтому кнопок
    согласования/печати у них нет. Этот эндпоинт закрывает разрыв, не
    дублируя генерацию документов здесь.

    Идемпотентно: если у договора уже есть привязанная рамочная голова —
    возвращает её id, ничего не создавая. Проверка «голова ли это» —
    is_framework_head() (purchases.py), единый источник истины, чтобы эта
    проверка не разъехалась с documents.py / purchase_transitions.py.
    """
    # Локальный импорт: purchases.py импортирует ensure_contract_linked ИЗ
    # этого модуля на уровне модуля — импорт is_framework_head сверху дал бы
    # циклический импорт при старте приложения.
    from app.routers.purchases import is_framework_head

    contract = (await db.execute(select(Contract).where(Contract.id == contract_id))).scalar_one_or_none()
    if not contract:
        raise HTTPException(404, "Договор не найден")
    if contract.contract_type not in FRAMEWORK_TYPES:
        raise HTTPException(400, "Согласование и документы через договор доступны только для рамочных договоров")

    linked = (await db.execute(
        select(Purchase).where(Purchase.contract_id == contract_id).order_by(Purchase.id.asc())
    )).scalars().all()
    existing = next((p for p in linked if is_framework_head(p)), None)
    if existing:
        return {"purchase_id": existing.id, "created": False}

    p = Purchase(**_build_approval_purchase_fields(
        contract, current_user.id, responsible_person=(current_user.full_name or None)
    ))
    db.add(p)
    await db.flush()  # get p.id

    if not p.registry_number:
        p.registry_number = f"РЕЕ-{date_type.today().year}-{p.id:05d}"
    if not p.purchase_number:
        max_result = await db.execute(select(func.coalesce(func.max(Purchase.purchase_number), 0)))
        p.purchase_number = max_result.scalar() + 1

    await _assign_framework_seq(p, db)

    # Владелец (2026-09-03), дословно: «просил же сделать по аналогии с
    # Заявкой. То есть выбирают при создании рамочного договора, надо
    # выбрать, кто будет согласовывать то, что этот договор вообще нужен» —
    # раньше здесь (2026-09-02) цепочка согласующих строилась АВТОМАТИЧЕСКИ от
    # руководителя организации, никого не спрашивая. Теперь голова заводится
    # БЕЗ согласующих (approval_status остаётся None) — автор сам добавляет их
    # вручную (POST /purchases/{pid}/approvals/add — доступно всем ролям, как
    # у заявки) или явно жмёт «Построить цепочку»
    # (POST /purchases/{pid}/approvers/cascade, см. purchases.py::
    # _build_framework_chain_approvals — функция осталась, вызывается только
    # по явному действию).
    await db.commit()
    return {"purchase_id": p.id, "created": True}


@router.delete("/{cid}")
async def delete_contract(cid: int, db: AsyncSession = Depends(get_db), _=Depends(require_action('contract.delete'))):
    result = await db.execute(select(Contract).where(Contract.id == cid))
    c = result.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Not found")
    await db.delete(c)
    await db.commit()
    return {"ok": True}


@router.post("/{cid}/merge/{target_id}")
async def merge_contracts(cid: int, target_id: int, db: AsyncSession = Depends(get_db), _=Depends(require_tab('contracts'))):
    """Merge contract cid INTO target_id: move all purchases, then delete cid."""
    if cid == target_id:
        raise HTTPException(400, "Нельзя объединить договор сам с собой")
    source = (await db.execute(select(Contract).where(Contract.id == cid))).scalar_one_or_none()
    target = (await db.execute(select(Contract).where(Contract.id == target_id))).scalar_one_or_none()
    if not source or not target:
        raise HTTPException(404, "Договор не найден")
    # Move purchases from source to target
    linked = (await db.execute(select(Purchase).where(Purchase.contract_id == cid))).scalars().all()
    for p in linked:
        p.contract_id = target_id
        if target.number:
            p.contract_number = target.number
    await db.delete(source)
    await db.commit()
    return {"ok": True, "moved_purchases": len(linked)}


@router.post("/migrate-from-purchases")
async def migrate_contracts_from_purchases(
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('contracts')),
):
    """Create Contract records from purchases.contract_number strings (skip existing)."""
    # Get all purchases with a contract_number set but no contract_id
    q = select(Purchase).where(
        Purchase.contract_number.isnot(None),
        Purchase.contract_number != "",
    )
    result = await db.execute(q)
    purchases = result.scalars().all()

    # Existing contract numbers
    existing_result = await db.execute(select(Contract.number))
    existing_numbers = {row[0] for row in existing_result.all() if row[0]}

    created = 0
    skipped = 0
    # Map: contract_number -> new Contract (for purchases to link to)
    new_contracts: dict[str, Contract] = {}

    for p in purchases:
        num = (p.contract_number or "").strip()
        if not num:
            continue
        if num in existing_numbers or num in new_contracts:
            skipped += 1
            continue
        c = Contract(
            number=num,
            date=p.contract_date if hasattr(p, "contract_date") else None,
            contract_type="single",
            subsidy_id=p.subsidy_id,
            status="active",
        )
        db.add(c)
        new_contracts[num] = c
        created += 1

    if new_contracts:
        await db.flush()
        # Link purchases to their new contracts
        for p in purchases:
            num = (p.contract_number or "").strip()
            if num in new_contracts and p.contract_id is None:
                p.contract_id = new_contracts[num].id

    await db.commit()
    return {"created": created, "skipped": skipped}


# ── OCR helper ─────────────────────────────────────────────────────────────────

def _ocr_pdf_to_rows(content: bytes) -> tuple[list, str | None]:
    """Fallback: convert scanned PDF pages to images, run OCR, parse lines.
    Returns (rows, error_message). error_message is None on success."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError:
        return [], "OCR-библиотеки не установлены (pdf2image, pytesseract). Обратитесь к администратору."

    try:
        images = convert_from_bytes(content, dpi=300)
    except Exception as e:
        return [], f"Не удалось преобразовать PDF в изображения для OCR: {e}"

    if not images:
        return [], "PDF не содержит страниц для распознавания."

    import re
    all_rows = []
    for img in images:
        try:
            text = pytesseract.image_to_string(img, lang='rus+eng')
        except Exception as e:
            return [], f"Ошибка OCR-распознавания: {e}"
        if not text:
            continue
        for line in text.strip().split('\n'):
            line = line.strip()
            if not line:
                continue
            # Split by tabs or multiple spaces
            cells = re.split(r'\t|  {2,}', line)
            cells = [c.strip() for c in cells if c.strip()]
            if cells:
                all_rows.append(cells)
    if not all_rows:
        return [], "OCR не нашёл текст на страницах. Возможно, качество скана слишком низкое."
    return all_rows, None


# ── File parsing helpers ───────────────────────────────────────────────────────

_CONTRACT_HINTS = (
    'номер', 'дата', 'контрагент', 'поставщик', 'исполнитель',
    'предмет', 'сумма', 'субсидия', 'статус', 'тип', 'метод',
    'начал', 'оконч', 'number', 'date', 'contractor', 'amount',
)


def _detect_contract_hdr(rows):
    best_score, best_idx = 0, 0
    for ri, row in enumerate(rows):
        norm = [str(h).strip().lower() if h is not None else "" for h in row]
        score = sum(1 for h in norm if h and any(x in h for x in _CONTRACT_HINTS))
        if score > best_score:
            best_score = score
            best_idx = ri
    return best_idx


def _parse_contract_file(fname: str, content: bytes):
    """Parse xlsx/xls/docx/pdf and return (all_rows, hdr_idx)."""
    fname = fname.lower()

    if fname.endswith('.xls') and not fname.endswith('.xlsx'):
        try:
            import xlrd
        except ImportError:
            raise HTTPException(500, "xlrd не установлен")
        try:
            wb = xlrd.open_workbook(file_contents=content)
        except Exception as e:
            raise HTTPException(400, f"Не удалось прочитать .xls: {e}")
        ws = wb.sheet_by_index(0)
        all_rows = [list(ws.row_values(i)) for i in range(ws.nrows)]

    elif fname.endswith('.xlsx'):
        if not _load_workbook:
            raise HTTPException(500, "openpyxl не установлен")
        try:
            wb = _load_workbook(BytesIO(content), read_only=True, data_only=True)
        except Exception as e:
            raise HTTPException(400, f"Не удалось прочитать .xlsx: {e}")
        ws = wb.active
        all_rows = list(ws.iter_rows(values_only=True))
        wb.close()

    elif fname.endswith(('.docx', '.doc')):
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(500, "python-docx не установлен")
        try:
            doc = Document(BytesIO(content))
        except Exception as e:
            raise HTTPException(400, f"Не удалось прочитать .docx: {e}")
        if not doc.tables:
            raise HTTPException(400, "В документе не найдено таблиц")
        tbl = doc.tables[0]
        all_rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]

    elif fname.endswith('.pdf'):
        if not _pdfplumber:
            raise HTTPException(500, "pdfplumber не установлен")
        try:
            with _pdfplumber.open(BytesIO(content)) as pdf:
                all_rows = []
                has_text = False
                for page in pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        for tbl in tables:
                            all_rows.extend([r for r in tbl if r])
                if not all_rows:
                    for page in _pdfplumber.open(BytesIO(content)).pages:
                        text = page.extract_text()
                        if text:
                            has_text = True
                            for line in text.strip().split('\n'):
                                cells = [c.strip() for c in line.split('\t')]
                                if len(cells) < 2:
                                    cells = [c.strip() for c in line.split('  ') if c.strip()]
                                if cells:
                                    all_rows.append(cells)
        except Exception as e:
            raise HTTPException(400, f"Не удалось прочитать PDF-файл: {e}")
        if not all_rows:
            if not has_text:
                # Scanned PDF — try OCR
                all_rows, ocr_error = _ocr_pdf_to_rows(content)
                if not all_rows:
                    detail = ocr_error or "OCR не смог распознать таблицу."
                    raise HTTPException(
                        400,
                        f"Этот PDF — скан (изображение). {detail} "
                        "Попробуйте сохранить данные в Excel (.xlsx) или Word (.docx)."
                    )
            else:
                raise HTTPException(
                    400,
                    "В PDF найден текст, но не удалось распознать таблицу. "
                    "Попробуйте сохранить данные в Excel (.xlsx) и загрузить его."
                )
    else:
        raise HTTPException(400, "Неподдерживаемый формат. Используйте .xlsx, .xls, .docx, .doc или .pdf")

    if not all_rows:
        raise HTTPException(400, "Файл пустой или не содержит данных")

    hdr_idx = _detect_contract_hdr(all_rows)
    return all_rows, hdr_idx


def _parse_date(val) -> Optional[date_type]:
    if not val:
        return None
    if isinstance(val, date_type):
        return val
    s = str(val).strip()
    for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            from datetime import datetime
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


def _parse_amount(val) -> Optional[Decimal]:
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return Decimal(str(val))
    s = str(val).strip().replace(' ', '').replace('\xa0', '').replace(',', '.')
    s = ''.join(c for c in s if c.isdigit() or c == '.')
    if not s:
        return None
    try:
        return Decimal(s)
    except Exception:
        return None


# ── Import endpoints ──────────────────────────────────────────────────────────

@router.post("/import/preview")
async def contracts_import_preview(
    file: UploadFile = File(...),
    current_user=Depends(require_tab('contracts')),
):
    """Parse uploaded file and return headers + sample rows for column mapping UI."""
    fname = (file.filename or '').lower()
    content = await file.read()
    all_rows, hdr_idx = _parse_contract_file(fname, content)
    hdr_rows = all_rows[hdr_idx:]
    if not hdr_rows:
        raise HTTPException(400, "Файл пустой")

    headers = [
        str(c).strip() if c else f"Столбец {j + 1}"
        for j, c in enumerate(hdr_rows[0])
    ]
    sample = [
        [str(c).strip() if c is not None else "" for c in row]
        for row in hdr_rows[1:min(6, len(hdr_rows))]
    ]
    return {
        "headers": headers,
        "sample": sample,
        "total_rows": max(len(all_rows) - hdr_idx - 1, 0),
        "header_row_offset": hdr_idx,
    }


@router.post("/import/mapped")
async def contracts_import_mapped(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_tab('contracts')),
    col_number: Optional[int] = Query(None),
    col_date: Optional[int] = Query(None),
    col_contractor: Optional[int] = Query(None),
    col_subject: Optional[int] = Query(None),
    col_max_amount: Optional[int] = Query(None),
    col_start_date: Optional[int] = Query(None),
    col_end_date: Optional[int] = Query(None),
    col_contract_type: Optional[int] = Query(None),
    col_purchase_method: Optional[int] = Query(None),
    col_item_type: Optional[int] = Query(None),
    col_status: Optional[int] = Query(None),
    col_notes: Optional[int] = Query(None),
    header_row_offset: int = Query(0),
    subsidy_id: Optional[int] = Query(None),
):
    """Import contracts using user-specified column mapping."""
    if col_number is None:
        raise HTTPException(400, "Не указан столбец «Номер договора»")

    fname = (file.filename or '').lower()
    content = await file.read()
    all_rows, _ = _parse_contract_file(fname, content)
    data_rows = all_rows[header_row_offset + 1:]

    col_map = {
        'number': col_number, 'date': col_date, 'contractor': col_contractor,
        'subject': col_subject, 'max_amount': col_max_amount,
        'start_date': col_start_date, 'end_date': col_end_date,
        'contract_type': col_contract_type, 'purchase_method': col_purchase_method,
        'item_type': col_item_type, 'status': col_status, 'notes': col_notes,
    }

    def _get(row, field):
        idx = col_map.get(field)
        if idx is None or idx >= len(row):
            return None
        v = row[idx]
        return str(v).strip() if v is not None else None

    # Pre-load contractors for matching by name/INN
    ctr_result = await db.execute(select(Contractor))
    all_contractors = ctr_result.scalars().all()
    ctr_by_name = {c.name.strip().lower(): c.id for c in all_contractors if c.name}
    ctr_by_inn = {c.inn.strip(): c.id for c in all_contractors if c.inn}

    created, skipped, errors = 0, 0, []
    for ri, row in enumerate(data_rows, start=1):
        num = _get(row, 'number')
        if not num:
            continue

        # Check duplicate by number
        dup = (await db.execute(
            select(Contract.id).where(Contract.number == num).limit(1)
        )).scalar_one_or_none()
        if dup:
            skipped += 1
            continue

        # Resolve contractor
        contractor_id = None
        ctr_val = _get(row, 'contractor')
        if ctr_val:
            ctr_lower = ctr_val.lower().strip()
            contractor_id = ctr_by_name.get(ctr_lower)
            if not contractor_id:
                contractor_id = ctr_by_inn.get(ctr_val.strip())

        # Resolve contract_type
        ct_raw = (_get(row, 'contract_type') or '').lower()
        if any(x in ct_raw for x in ('рамоч', 'накопит', 'framework')):
            contract_type = 'framework_cumulative'
        elif any(x in ct_raw for x in ('разов', 'единич', 'single')):
            contract_type = 'single'
        else:
            contract_type = 'single'

        c = Contract(
            number=num,
            date=_parse_date(_get(row, 'date')),
            contract_type=contract_type,
            contractor_id=contractor_id,
            subsidy_id=subsidy_id,
            subject=_get(row, 'subject'),
            max_amount=_parse_amount(_get(row, 'max_amount')),
            start_date=_parse_date(_get(row, 'start_date')),
            end_date=_parse_date(_get(row, 'end_date')),
            purchase_method=_get(row, 'purchase_method'),
            item_type=_get(row, 'item_type'),
            status=_get(row, 'status') or 'active',
            notes=_get(row, 'notes'),
        )
        db.add(c)
        created += 1

    if created:
        await db.commit()

    return {"created": created, "skipped": skipped, "errors": errors}


@router.post("/{contract_id}/generate-monthly-stages")
async def generate_monthly_stages(
    contract_id: int,
    body: MonthlyStagesRequest,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_tab('contracts')),
):
    """Сгенерировать N ежемесячных этапов для рамочного договора."""
    contract = (await db.execute(select(Contract).where(Contract.id == contract_id))).scalar_one_or_none()
    if not contract:
        raise HTTPException(404, "Договор не найден")
    if not (contract.contract_type or "").startswith("framework_"):
        raise HTTPException(400, "Генерация этапов доступна только для рамочных договоров (framework_*)")

    contract_name = contract.subject or contract.number or f"Договор #{contract_id}"

    created_items = []
    skipped_items = []

    year = body.start_year
    month = body.start_month

    for i in range(body.count):
        # compute (year, month) for this stage
        cur_year = year + (month - 1 + i) // 12
        cur_month = (month - 1 + i) % 12 + 1

        last_day = calendar.monthrange(cur_year, cur_month)[1]
        if body.deadline_kind == "end_of_month":
            deadline = date_type(cur_year, cur_month, last_day)
        else:
            day = min(body.deadline_day or last_day, last_day)
            deadline = date_type(cur_year, cur_month, day)

        # Idempotency check
        existing = (await db.execute(
            select(Purchase).where(
                Purchase.contract_id == contract_id,
                Purchase.service_deadline_date == deadline,
            )
        )).scalar_one_or_none()
        if existing:
            skipped_items.append({"deadline": deadline.isoformat(), "reason": "already_exists"})
            continue

        month_name_gen = RU_MONTHS_GEN[cur_month - 1]
        month_name_nom = RU_MONTHS_NOM[cur_month - 1]

        subject = (
            body.subject_template
            .replace("{contract_name}", contract_name)
            .replace("{month_name}", month_name_gen)
            .replace("{year}", str(cur_year))
        )
        stage_label = f"{month_name_nom} {cur_year}"

        p = Purchase(
            status="planned",
            contract_id=contract_id,
            subsidy_id=body.subsidy_id,
            feo_category_id=body.feo_category_id,
            contractor_id=contract.contractor_id,
            service_term_mode="deadline",
            service_deadline_date=deadline,
            planned_total_price=Decimal(str(body.amount_per_stage)),
            contract_price=Decimal(str(body.amount_per_stage)),
            subject=subject,
            is_likely_needed=body.is_likely_needed,
            stage_label=stage_label,
            purchase_contract_type=contract.contract_type,
        )
        db.add(p)
        await db.flush()  # get p.id

        await _assign_framework_seq(p, db)

        created_items.append({
            "id": p.id,
            "deadline": deadline.isoformat(),
            "label": stage_label,
            "subject": subject,
        })

    await db.commit()
    return {"created": created_items, "skipped": skipped_items}


# ── Non-router helper ──────────────────────────────────────────────────────────

@router.post("/bulk-enrich-from-purchases", dependencies=[Depends(require_role('admin'))])
async def bulk_enrich_contracts_from_purchases(db: AsyncSession = Depends(get_db)):
    """
    Two-phase batch operation для legacy:

    Phase 1 — create Contract для Purchase'ов с receipts (но без contract_id).
              Это покрывает авансовые импорты из Phase 27.1.4 эпохи.

    Phase 2 — enrich existing Contract'ы где ХОТЯ БЫ ОДНО optional поле NULL.
              Покрывает любые legacy договоры (созданные вручную или через миграцию).
    """
    # Phase 27.1.9: re-link Purchase.contract_id для Purchase'ов где contract_number
    # совпадает с Contract.number но contract_id NULL ИЛИ указывает на другой Contract.
    # Это закрывает gap когда auto-create Contract (Phase 27.1.4) и Purchase
    # разъехались (recompute / manual edit / прочее).
    from app.models.purchase import Purchase as _P2
    from app.models.contract import Contract as _C2
    from sqlalchemy import select as _s2

    purchases_with_num = (await db.execute(
        _s2(_P2).where(_P2.contract_number.isnot(None), _P2.contract_number != '')
    )).scalars().all()

    all_contracts_for_link = (await db.execute(_s2(_C2))).scalars().all()
    contract_by_number: dict = {}
    for c_iter in all_contracts_for_link:
        if c_iter.number:
            contract_by_number.setdefault(c_iter.number.strip(), []).append(c_iter)

    purchases_linked = 0
    for p_iter in purchases_with_num:
        num = (p_iter.contract_number or '').strip()
        if not num:
            continue
        candidates = contract_by_number.get(num, [])
        if not candidates:
            continue
        best = None
        if p_iter.contractor_id:
            best = next((c_c for c_c in candidates if c_c.contractor_id == p_iter.contractor_id), None)
        if not best:
            best = candidates[0]
        if p_iter.contract_id != best.id:
            p_iter.contract_id = best.id
            purchases_linked += 1

    from app.models.purchase_receipt import PurchaseReceipt as _PR
    rows_result = await db.execute(
        select(Purchase).join(_PR, _PR.purchase_id == Purchase.id)
        .where(Purchase.contract_id.is_(None), Purchase.contractor_id.isnot(None))
        .distinct()
    )
    rows = rows_result.scalars().all()

    created = 0
    for p in rows:
        # Idempotency guard: если уже есть Contract с таким же contractor_id + subsidy_id
        # и number = ожидаемому значению — пропустить
        expected_number = str(p.purchase_number) if p.purchase_number else f"AVANS-{p.id}"
        existing_q = await db.execute(
            select(Contract).where(
                Contract.number == expected_number,
                Contract.contractor_id == p.contractor_id,
            ).limit(1)
        )
        if existing_q.scalar_one_or_none():
            continue
        new_contract = Contract(
            contractor_id=p.contractor_id,
            subsidy_id=p.subsidy_id,
            contract_type='single',
            number=expected_number,
            date=p.created_at.date() if getattr(p, 'created_at', None) else None,
            status='active',
            # Phase 27.1.5: заполнить ВСЕ доступные поля из Purchase
            subject=p.subject or str(p.purchase_number or ''),
            max_amount=p.total_nmck or p.contract_price or p.planned_total_price,
            start_date=p.contract_date,
            end_date=p.execution_term,
            purchase_method=p.purchase_method if p.purchase_method in ('single', 'competitive') else 'single',
            item_type=p.item_type or 'товар',
        )
        db.add(new_contract)
        await db.flush()
        p.contract_id = new_contract.id
        created += 1
    # Phase 27.1.6: enrich existing Contracts где ХОТЯ БЫ ОДНО optional поле NULL
    from sqlalchemy import or_
    from app.models.contract import Contract as _C
    contracts_to_enrich = (await db.execute(
        select(_C).where(
            or_(
                _C.subject.is_(None),
                _C.max_amount.is_(None),
                _C.start_date.is_(None),
                _C.end_date.is_(None),
                _C.purchase_method.is_(None),
                _C.item_type.is_(None),
            )
        )
    )).scalars().all()
    enriched = 0
    for c in contracts_to_enrich:
        filled_count = await _enrich_contract_from_purchases(c, db)
        if filled_count > 0:
            enriched += 1
    # Phase 27.1.8: расширенный relink orphan ContractItem.source_item_id с fallback'ами
    from app.models.contract_item import ContractItem as _CI
    from app.models.purchase_item import PurchaseItem as _PI
    from sqlalchemy import select as _select, func as _func

    orphans_q = await db.execute(_select(_CI).where(_CI.source_item_id.isnot(None)))
    orphans_all = orphans_q.scalars().all()
    existing_pi_ids = set((await db.execute(_select(_PI.id))).scalars().all())
    orphans = [ci for ci in orphans_all if ci.source_item_id not in existing_pi_ids]

    relinked = 0
    for ci in orphans:
        if not ci.purchase_id:
            continue

        candidate = None

        # Pass 1: exact name match (был в 27.1.7)
        if ci.name:
            candidate = (await db.execute(
                _select(_PI).where(
                    _PI.purchase_id == ci.purchase_id,
                    _PI.item_name == ci.name,
                ).limit(1)
            )).scalar_one_or_none()

        # Pass 2: case-insensitive trimmed match
        if not candidate and ci.name:
            normalized_ci_name = ci.name.strip().lower()
            candidate = (await db.execute(
                _select(_PI).where(
                    _PI.purchase_id == ci.purchase_id,
                    _func.lower(_func.trim(_PI.item_name)) == normalized_ci_name,
                ).limit(1)
            )).scalar_one_or_none()

        # Pass 3: 1-to-1 fallback — если у Purchase ровно один PI и один orphan CI с тем же purchase_id, безусловно связать
        if not candidate:
            all_pi_in_purchase = (await db.execute(
                _select(_PI).where(_PI.purchase_id == ci.purchase_id)
            )).scalars().all()
            all_ci_in_purchase_orphan = [
                o for o in orphans
                if o.purchase_id == ci.purchase_id
            ]
            if len(all_pi_in_purchase) == 1 and len(all_ci_in_purchase_orphan) == 1:
                candidate = all_pi_in_purchase[0]

        # Pass 4: match по qty + unit_price (если совпадают точно)
        if not candidate and ci.quantity is not None and ci.unit_price is not None:
            candidate = (await db.execute(
                _select(_PI).where(
                    _PI.purchase_id == ci.purchase_id,
                    _PI.quantity == ci.quantity,
                    _PI.unit_price == ci.unit_price,
                ).limit(1)
            )).scalar_one_or_none()

        if candidate:
            ci.source_item_id = candidate.id
            relinked += 1

    # Phase 27.1.17: enrich ContractItem.vat_rate из связанной PurchaseItem (для legacy без vat_rate)
    from app.models.contract_item import ContractItem as _CI_vat
    from app.models.purchase_item import PurchaseItem as _PI_vat
    ci_no_vat = (await db.execute(
        _select(_CI_vat).where(_CI_vat.vat_rate.is_(None), _CI_vat.source_item_id.isnot(None))
    )).scalars().all()
    vat_filled = 0
    for ci in ci_no_vat:
        pi = await db.get(_PI_vat, ci.source_item_id)
        if pi and pi.vat_rate:
            ci.vat_rate = pi.vat_rate
            vat_filled += 1

    await db.commit()
    return {
        "created": created,
        "scanned_purchases": len(rows),
        "enriched_existing": enriched,
        "scanned_contracts": len(contracts_to_enrich),
        "relinked_orphans": relinked,
        "purchases_linked": purchases_linked,
        "vat_filled": vat_filled,
    }


@router.post("/backfill-from-receipts", deprecated=True, dependencies=[Depends(require_role('admin'))])
async def backfill_from_receipts_legacy(db: AsyncSession = Depends(get_db)):
    """Deprecated alias. Use /bulk-enrich-from-purchases."""
    return await bulk_enrich_contracts_from_purchases(db)


async def ensure_contract_linked(p: Purchase, db: AsyncSession) -> None:
    """Find-or-create a Contract for this purchase and link it via purchase.contract_id.

    Uniqueness key: contract_number + contractor_id + contractor INN + contract_date.
    If any of the last 3 are absent they are simply omitted from the lookup
    (making the match less strict when data is incomplete).
    """
    num = (p.contract_number or "").strip()
    if not num:
        return

    # Already linked — sync any changed fields (date, contractor) to contract record
    if p.contract_id:
        existing = await db.get(Contract, p.contract_id)
        if existing and existing.number == num:
            if p.contract_date and existing.date != p.contract_date:
                existing.date = p.contract_date
            if p.contractor_id and existing.contractor_id != p.contractor_id:
                existing.contractor_id = p.contractor_id
            return

    # Resolve contractor INN if we have a contractor_id
    inn: Optional[str] = None
    if p.contractor_id:
        inn_row = await db.execute(
            select(Contractor.inn).where(Contractor.id == p.contractor_id)
        )
        inn = inn_row.scalar_one_or_none()

    # Build lookup query with all 4 uniqueness parameters
    q = (
        select(Contract)
        .outerjoin(Contractor, Contract.contractor_id == Contractor.id)
        .where(Contract.number == num)
    )
    if p.contractor_id:
        q = q.where(Contract.contractor_id == p.contractor_id)
    if p.contract_date:
        q = q.where(Contract.date == p.contract_date)
    if inn:
        q = q.where(Contractor.inn == inn)

    contract = (await db.execute(q)).scalar_one_or_none()

    if not contract:
        contract = Contract(
            number=num,
            date=p.contract_date,
            contract_type=p.purchase_contract_type or "single",
            contractor_id=p.contractor_id,
            subsidy_id=p.subsidy_id,
            status="active",
        )
        db.add(contract)
        await db.flush()

    p.contract_id = contract.id
