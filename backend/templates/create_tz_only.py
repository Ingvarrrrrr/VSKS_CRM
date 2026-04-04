"""Create standalone tz_only.docx template by extracting TZ section from contract_tz.docx."""
import sys
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
from copy import deepcopy

# Start from a copy of contract_tz.docx
src = Document('/app/templates/contract_tz.docx')
paras = src.paragraphs

# Find the TZ section start (Appendix №1)
tz_start = None
for i, p in enumerate(paras):
    if 'Приложение № 1' in p.text and i > 100:
        tz_start = i
        break

if tz_start is None:
    print("ERROR: Could not find TZ section", file=sys.stderr)
    sys.exit(1)

print(f"TZ section starts at paragraph {tz_start}: {repr(paras[tz_start].text[:60])}", file=sys.stderr)

# Create new document
doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21)
section.left_margin = Cm(2.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# Copy TZ paragraphs to new doc
# We need to copy from tz_start to end of document
body = doc.element.body

# Remove the default empty paragraph
for p in doc.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
    doc.element.body.remove(p)

# Copy paragraphs from src starting at tz_start
src_body = src.element.body
src_paras_xml = src_body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')

# Actually copy the whole section of XML (paragraphs and tables after tz_start)
# Need to get the XML elements in order
all_block_elements = [child for child in src_body
                      if child.tag in [
                          '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p',
                          '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl',
                      ]]

# Count paragraphs to find tz_start element index
para_count = 0
tz_start_elem_idx = None
for idx, elem in enumerate(all_block_elements):
    if elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
        if para_count == tz_start:
            tz_start_elem_idx = idx
            break
        para_count += 1

print(f"TZ section element index: {tz_start_elem_idx}", file=sys.stderr)

if tz_start_elem_idx is None:
    print("ERROR: Could not find element index", file=sys.stderr)
    sys.exit(1)

# Copy all elements from tz_start onwards
for elem in all_block_elements[tz_start_elem_idx:]:
    new_elem = deepcopy(elem)
    doc.element.body.append(new_elem)

# Add a final empty paragraph if needed
final_p = OxmlElement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
doc.element.body.append(final_p)

doc.save('/app/templates/tz_only.docx')
print("Created /app/templates/tz_only.docx", file=sys.stderr)

# Verify
doc2 = Document('/app/templates/tz_only.docx')
print(f"tz_only.docx: {len(doc2.paragraphs)} paragraphs, {len(doc2.tables)} tables", file=sys.stderr)
for i, p in enumerate(doc2.paragraphs[:10]):
    if p.text.strip():
        print(f"  P{i}: {repr(p.text[:80])}", file=sys.stderr)
