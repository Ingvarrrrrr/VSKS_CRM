#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для анализа документов и извлечения всех параметров для формы договора/поставки
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Внимание: python-docx не установлен. Установите: pip install python-docx")

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("Внимание: openpyxl не установлен. Установите: pip install openpyxl")

def extract_text_from_docx(file_path):
    """Извлечь весь текст из docx файла"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        
        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
        
        return "\n".join(text)
    except Exception as e:
        print(f"Ошибка при чтении {file_path}: {e}")
        return None

def analyze_excel_sheet(file_path, sheet_name):
    """Анализировать лист Excel"""
    if not EXCEL_AVAILABLE:
        return None
    
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        if sheet_name not in wb.sheetnames:
            return None
        
        ws = wb[sheet_name]
        
        # Получить заголовки (первая строка)
        headers = []
        if ws.max_row > 0:
            for col in range(1, min(ws.max_column + 1, 100)):  # Ограничение до 100 столбцов
                cell_value = ws.cell(row=1, column=col).value
                if cell_value:
                    headers.append(str(cell_value).strip())
                else:
                    headers.append(f"Column_{col}")
        
        return headers
    except Exception as e:
        print(f"Ошибка при чтении Excel {file_path}, лист {sheet_name}: {e}")
        return None

def find_keywords_in_text(text, keywords):
    """Найти ключевые слова в тексте"""
    found = []
    text_lower = text.lower()
    for keyword in keywords:
        if keyword.lower() in text_lower:
            found.append(keyword)
    return found

def main():
    base_dir = Path(".")
    
    # Ключевые слова для поиска полей
    contract_keywords = [
        "номер договора", "дата договора", "субъект", "заказчик", "исполнитель",
        "предмет договора", "сумма договора", "стоимость", "цена",
        "срок исполнения", "срок поставки", "срок оказания услуг",
        "реквизиты", "инн", "кпп", "огрн", "расчетный счет", "корреспондентский счет",
        "бик", "банк", "адрес", "телефон", "email", "контактное лицо",
        "руководитель", "директор", "основание", "устав", "доверенность",
        "нмцк", "начальная цена", "победитель", "участник",
        "гарантийный срок", "условия оплаты", "способ оплаты",
        "аванс", "предоплата", "постоплата", "рассрочка",
        "ндс", "без ндс", "налог", "налогообложение",
        "ответственность сторон", "порядок приемки", "порядок сдачи",
        "форс-мажор", "разрешение споров", "расторжение договора",
        "особые условия", "дополнительные условия", "приложения",
        "техническое задание", "спецификация", "график поставки",
        "график платежей", "график работ", "календарный план",
        "отчетность", "отчет", "акт", "счет", "счет-фактура",
        "упаковка", "транспортировка", "доставка", "погрузка", "разгрузка",
        "качество", "требования", "стандарты", "гост", "ту",
        "количество", "единица измерения", "окей", "штука", "килограмм",
        "закупка", "лот", "позиция", "номенклатура",
        "субсидия", "фео", "категория расходов", "тип расходов",
        "бюджет", "финансирование", "источник финансирования",
        "статус", "стадия", "этап", "выполнено", "в работе", "планируется"
    ]
    
    all_fields = set()
    documents_analyzed = []
    
    print("=" * 80)
    print("АНАЛИЗ ДОКУМЕНТОВ ДЛЯ ИЗВЛЕЧЕНИЯ ПАРАМЕТРОВ")
    print("=" * 80)
    print()
    
    # Анализ Excel файлов
    excel_files = list(base_dir.glob("*.xlsx"))
    for excel_file in excel_files:
        print(f"\n📊 Анализ Excel файла: {excel_file.name}")
        print("-" * 80)
        
        if EXCEL_AVAILABLE:
            # Пробуем прочитать лист GoodsService
            headers = analyze_excel_sheet(excel_file, "GoodsService")
            if headers:
                print(f"Найдено столбцов в GoodsService: {len(headers)}")
                print("Заголовки:")
                for i, header in enumerate(headers[:50], 1):  # Первые 50
                    print(f"  {i}. {header}")
                    all_fields.add(header)
                if len(headers) > 50:
                    print(f"  ... и еще {len(headers) - 50} столбцов")
                documents_analyzed.append(("Excel", excel_file.name, len(headers)))
    
    # Анализ DOCX файлов
    docx_files = list(base_dir.glob("**/*.docx"))
    docx_files.extend(list(base_dir.glob("**/*.doc")))
    
    print(f"\n📄 Анализ DOCX/DOC файлов: найдено {len(docx_files)} файлов")
    print("-" * 80)
    
    for docx_file in docx_files[:10]:  # Анализируем первые 10 файлов
        print(f"\n📄 {docx_file.name}")
        text = extract_text_from_docx(docx_file)
        if text:
            found_keywords = find_keywords_in_text(text, contract_keywords)
            if found_keywords:
                print(f"  Найдены ключевые слова: {', '.join(found_keywords[:10])}")
                for keyword in found_keywords:
                    all_fields.add(keyword)
            documents_analyzed.append(("DOCX", docx_file.name, len(found_keywords) if found_keywords else 0))
    
    # Вывод результатов
    print("\n" + "=" * 80)
    print("ИТОГОВЫЙ СПИСОК ВСЕХ НАЙДЕННЫХ ПОЛЕЙ")
    print("=" * 80)
    
    # Сортируем поля по алфавиту
    sorted_fields = sorted(all_fields)
    
    print(f"\nВсего найдено уникальных полей: {len(sorted_fields)}\n")
    
    # Группируем поля по категориям
    categories = {
        "Основная информация": [],
        "Контрагент": [],
        "Финансы": [],
        "Сроки": [],
        "Состав договора": [],
        "Категоризация ФЭО": [],
        "Статусы": [],
        "Документы": [],
        "Прочее": []
    }
    
    for field in sorted_fields:
        field_lower = field.lower()
        if any(x in field_lower for x in ["номер", "дата", "предмет", "субъект", "заказчик", "исполнитель"]):
            categories["Основная информация"].append(field)
        elif any(x in field_lower for x in ["контрагент", "реквизиты", "инн", "кпп", "огрн", "банк", "счет", "бик", "адрес", "телефон", "email", "руководитель", "директор"]):
            categories["Контрагент"].append(field)
        elif any(x in field_lower for x in ["сумма", "стоимость", "цена", "нмцк", "платеж", "оплата", "аванс", "ндс", "налог", "бюджет", "финансирование"]):
            categories["Финансы"].append(field)
        elif any(x in field_lower for x in ["срок", "дата", "график", "календарный"]):
            categories["Сроки"].append(field)
        elif any(x in field_lower for x in ["товар", "услуга", "номенклатура", "позиция", "количество", "единица", "окей", "спецификация"]):
            categories["Состав договора"].append(field)
        elif any(x in field_lower for x in ["фео", "категория", "тип расходов", "субсидия"]):
            categories["Категоризация ФЭО"].append(field)
        elif any(x in field_lower for x in ["статус", "стадия", "этап", "выполнено", "в работе", "планируется"]):
            categories["Статусы"].append(field)
        elif any(x in field_lower for x in ["документ", "акт", "счет", "отчет", "приложение", "техническое задание"]):
            categories["Документы"].append(field)
        else:
            categories["Прочее"].append(field)
    
    for category, fields in categories.items():
        if fields:
            print(f"\n{category}:")
            for field in fields:
                print(f"  - {field}")
    
    print("\n" + "=" * 80)
    print(f"Проанализировано документов: {len(documents_analyzed)}")
    print("=" * 80)

if __name__ == "__main__":
    main()



