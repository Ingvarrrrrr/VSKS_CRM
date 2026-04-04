"""
Обновляет approval_sheet.docx и service_note.docx:
- approval_sheet: таблица подписантов заменяется на jinja2 цикл по {{ approvers }}
  Структура: header, {%tr for %} (пустая строка-маркер), content row, {%tr endfor %}
- service_note: добавляются переменные {{ purchase_basis }}, {{ initiator_name }}, {{ initiator_role }}
"""
import os
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import lxml.etree as etree

TEMPLATES_DIR = os.path.dirname(os.path.abspath(__file__))


def _clear_cell_set_text(cell_el, text):
    """Очищает ячейку и устанавливает текст в одном параграфе."""
    for p in cell_el.findall(qn('w:p')):
        cell_el.remove(p)
    p_el = etree.SubElement(cell_el, qn('w:p'))
    r_el = etree.SubElement(p_el, qn('w:r'))
    t_el = etree.SubElement(r_el, qn('w:t'))
    t_el.text = text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def update_approval_sheet():
    path = os.path.join(TEMPLATES_DIR, "approval_sheet.docx")
    doc = Document(path)

    # Find table with signatories (header row contains "Должность, ФИО")
    approval_table = None
    for t in doc.tables:
        if t.rows and any("Должность" in cell.text for cell in t.rows[0].cells):
            approval_table = t
            break

    if approval_table is None:
        print("ERROR: Could not find approval table")
        return

    tbl = approval_table._tbl
    header_row = approval_table.rows[0]

    # Remove all rows after header
    rows_to_remove = list(tbl.iterchildren(qn('w:tr')))[1:]
    for row_el in rows_to_remove:
        tbl.remove(row_el)

    # Correct docxtpl structure for row loop:
    # Row A: {%tr for a in approvers %}  (empty marker row — gets consumed by docxtpl)
    # Row B: actual content {{ a.role_name }} etc.  (this row is repeated)
    # Row C: {%tr endfor %}              (empty marker row — gets consumed by docxtpl)

    def make_row(cells_text):
        """Clone header row structure, set cell texts."""
        new_row = deepcopy(header_row._tr)
        row_cells = new_row.findall(qn('w:tc'))
        for i, text in enumerate(cells_text):
            if i < len(row_cells):
                _clear_cell_set_text(row_cells[i], text)
        # Pad missing cells with empty
        for i in range(len(cells_text), len(row_cells)):
            _clear_cell_set_text(row_cells[i], '')
        return new_row

    num_cols = len(header_row.cells)

    # Row A: for-tag marker (empty content)
    for_cells = ['{%tr for a in approvers %}'] + [''] * (num_cols - 1)
    tbl.append(make_row(for_cells))

    # Row B: content template
    content_cells = [
        '{{ a.num }}',
        '{{ a.role_name }}\n{{ a.full_name }}',
        '',         # Подпись (оставляем пустой для подписи)
        '{{ a.note }}',
    ]
    # Pad to num_cols
    while len(content_cells) < num_cols:
        content_cells.append('')
    tbl.append(make_row(content_cells[:num_cols]))

    # Row C: endfor marker
    end_cells = ['{%tr endfor %}'] + [''] * (num_cols - 1)
    tbl.append(make_row(end_cells))

    doc.save(path)
    print(f"approval_sheet.docx updated: dynamic approvers table (4-row structure)")


def update_service_note():
    path = os.path.join(TEMPLATES_DIR, "service_note.docx")
    doc = Document(path)

    # Find table with "ФИО, Контактные данные инициатора"
    info_table = None
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "ФИО" in cell.text and "Контактные" in cell.text:
                    info_table = t
                    break
            if info_table:
                break
        if info_table:
            break

    if info_table:
        for row in info_table.rows:
            if "ФИО" in row.cells[0].text and "Контактные" in row.cells[0].text:
                p = row.cells[0].add_paragraph()
                p.add_run("Инициатор: {{ initiator_role }} {{ initiator_name }}")
                break

    # Add purchase_basis to {{today}} paragraph
    basis_added = False
    for para in doc.paragraphs:
        if "{{today}}" in para.text or ("today" in para.text and "{{" in para.text):
            new_text = para.text + "  |  Основание: {{ purchase_basis }}"
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
            basis_added = True
            break

    if not basis_added:
        doc.paragraphs[0].insert_paragraph_before("Основание: {{ purchase_basis }}")

    doc.save(path)
    print(f"service_note.docx updated: added initiator + purchase_basis")


if __name__ == "__main__":
    update_approval_sheet()
    update_service_note()
    print("Done.")
