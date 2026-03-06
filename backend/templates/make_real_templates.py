"""
Создаёт docxtpl-шаблоны из оригинальных типовых документов.

Перед запуском скопировать файлы в контейнер:
  docker cp "C:/Users/1/VSKS_CRM/типовые документы/Регионы/Служебная записка на закупку (образец).docx" vsks_crm-backend-1:/tmp/sz.docx
  docker cp "C:/Users/1/VSKS_CRM/типовые документы/ЛИСТ_СОГЛАСОВАНИЯ основной 2025г..docx" vsks_crm-backend-1:/tmp/ls.docx
  docker cp "C:/Users/1/VSKS_CRM/типовые документы/Регионы/ПРОЕКТ Договор поставка 2025.docx" vsks_crm-backend-1:/tmp/dp.docx

Затем запустить:
  docker exec vsks_crm-backend-1 sh -c 'cd /app && python templates/make_real_templates.py'
"""
import os, copy
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph as DocxParagraph

TEMPLATES_DIR = '/app/templates'
SZ_SRC = '/tmp/sz.docx'
LS_SRC = '/tmp/ls.docx'
DP_SRC = '/tmp/dp.docx'


# ── Helpers ────────────────────────────────────────────────────────────────────

def para_replace(para, old, new):
    """Replace text in a paragraph, consolidating all runs to first."""
    full = ''.join(r.text for r in para.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(new_full)
    return True


def para_set(para, text):
    """Overwrite paragraph text entirely."""
    full = ''.join(r.text for r in para.runs)
    para_replace(para, full, text) if full else para.add_run(text)


def cell_replace(cell, old, new):
    for p in cell.paragraphs:
        para_replace(p, old, new)


def cell_set(cell, new_text):
    """Set first paragraph of cell to new_text, clear rest."""
    for i, p in enumerate(cell.paragraphs):
        if i == 0:
            full = ''.join(r.text for r in p.runs)
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]:
                    r.text = ''
            else:
                p.add_run(new_text)
        else:
            for r in p.runs:
                r.text = ''


def find_para(doc, text):
    """Return (index, paragraph) of first paragraph containing text."""
    for i, p in enumerate(doc.paragraphs):
        if text in p.text:
            return i, p
    return None, None


def insert_para_after(ref_p_element, parent, text, bold=False, size=None):
    """Insert new paragraph immediately after ref_p_element in parent._element."""
    new_p_xml = copy.deepcopy(ref_p_element)
    # Clear runs
    from docx.oxml.ns import qn
    for r in new_p_xml.findall(qn('w:r')):
        new_p_xml.remove(r)
    ref_p_element.addnext(new_p_xml)
    new_para = DocxParagraph(new_p_xml, parent)
    run = new_para.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return new_para


# ── 1. service_note.docx ───────────────────────────────────────────────────────
def make_service_note():
    doc = Document(SZ_SRC)

    # Date line
    _, p3 = find_para(doc, '«___»')
    if p3:
        para_set(p3, 'от {{today}}')

    # Table 2: main data table
    if len(doc.tables) >= 3:
        t2 = doc.tables[2]
        # [0,1] Предмет закупки
        cell_set(t2.rows[0].cells[1],
            '{{item_names}}\n\nТехническое задание прилагается (Приложение № 1).')
        # [1,1] Сумма
        cell_set(t2.rows[1].cells[1],
            'НМЦК: {{total_nmck}}\nЦена договора: {{contract_price}}\nЭкономия: {{economy}}\nПорядок оплаты: постоплата 100% в течение 10 рабочих дней после подписания акта.')
        # [2,1] Способ закупки
        cell_set(t2.rows[2].cells[1], '{{purchase_method}}')
        # [3,1] Срок
        cell_set(t2.rows[3].cells[1], '{{execution_term}}')
        # [4,1] Обоснование
        cell_set(t2.rows[4].cells[1],
            'Производственная необходимость в рамках реализации проекта субсидии «{{subsidy_name}}».')
        # [5,1] Проект
        cell_set(t2.rows[5].cells[1], '{{subsidy_name}} ({{subsidy_year}})\nРеестровый номер: {{registry_number}}')
        # [6,1] KPI — leave blank or comment
        if t2.rows[6].cells[1].text.strip() == '':
            pass  # leave empty

    # Table 1: инициатор
    if len(doc.tables) >= 2:
        t1 = doc.tables[1]
        cell_set(t1.rows[1].cells[1],
            'Контрагент: {{contractor_name}}\nИНН: {{contractor_inn}}\nТел.: {{contractor_phone}}\nАдрес эл. почты: {{contractor_email}}')

    out = os.path.join(TEMPLATES_DIR, 'service_note.docx')
    doc.save(out)
    print(f'[OK] {out}')


# ── 2. approval_sheet.docx ─────────────────────────────────────────────────────
def make_approval_sheet():
    doc = Document(LS_SRC)

    # Header paragraphs P1-P7
    replacements = [
        ('Предмет договора: Брендированная продукция (16 шт. флаг ШСО)',
         'Предмет договора: {{item_names}}'),
        ('Способ определения поставщика: ',
         'Способ определения поставщика: {{purchase_method}}'),
        ('Контрагент: ООО «Феникс»',
         'Контрагент: {{contractor_name}}'),
        ('ИНН контрагента: 7733872083',
         'ИНН контрагента: {{contractor_inn}}'),
        ('Цена договора (в том числе НДС/ без НДС): ,00 руб., без НДС.',
         'Цена договора: {{contract_price}}, без НДС. НМЦК: {{total_nmck}}'),
        ('Источник финансирования: ',
         'Источник финансирования: {{subsidy_name}}'),
        ('Основание (инициирование закупки): План закупок на год',
         'Основание: реестровый номер {{registry_number}}, договор № {{contract_number}} от {{contract_date}}'),
    ]
    for p in doc.paragraphs:
        for old, new in replacements:
            para_replace(p, old, new)

    # Table 0: header block — replace date "от «01»  2025 г."
    if doc.tables:
        t0 = doc.tables[0]
        for row in t0.rows:
            for cell in row.cells:
                cell_replace(cell, 'от «01»  2025 г.', 'от {{today}}')

    out = os.path.join(TEMPLATES_DIR, 'approval_sheet.docx')
    doc.save(out)
    print(f'[OK] {out}')


# ── 3. contract_tz.docx (Договор поставки + ТЗ) ───────────────────────────────
def make_contract_tz():
    doc = Document(DP_SRC)
    paras = doc.paragraphs

    # ── Paragraph replacements ──────────────────────────────────────────
    # P0: "Договор поставки № "
    para_replace(paras[0], 'Договор поставки № ', 'Договор поставки № {{contract_number}} ')

    # P1: "г. _____   «__» _____ 2025 г."
    para_set(paras[1], '{{contractor_address}}                               {{contract_date}}')

    # P5: contractor block
    para_replace(paras[5], '_____,', '{{contractor_name}},')
    para_replace(paras[5], 'в лице ______,', 'в лице {{contractor_signatory}},')
    para_replace(paras[5], 'на основании _______,', 'на основании {{contractor_signatory_basis}},')

    # P6: procedure description (purchase method)
    para_replace(paras[6],
        'по результатам открытой конкурсной процедуры, размещенной на ____ (Протокол №___ от___) ',
        'в соответствии с порядком проведения закупки ({{purchase_method}}) ')

    # P60: contract price
    _, p60 = find_para(doc, '__________, 00 (______________) рублей 00 копеек')
    if p60:
        para_replace(p60, '__________,00 (______________) рублей 00 копеек, в том числе НДС в размере __.',
                     '{{contract_price}}, без НДС.')
    # fallback with different spacing
    _, p_price = find_para(doc, 'составляет __________')
    if p_price:
        full = ''.join(r.text for r in p_price.runs)
        import re
        new_full = re.sub(r'составляет [_,\s]+\([_,\s]+\) рублей [0-9]+ копеек[^.]*\.',
                          'составляет {{contract_price}}, без НДС.', full)
        if p_price.runs:
            p_price.runs[0].text = new_full
            for r in p_price.runs[1:]:
                r.text = ''

    # P67: subsidy reference
    _, p67 = find_para(doc, 'Соглашения №___')
    if p67:
        para_replace(p67, 'Соглашения №___', 'Соглашения №___')  # keep, no reliable pattern
        para_replace(p67, '4.6. Финансирование договора осуществляется в рамках ', '')

    # P73: delivery term
    _, p73 = find_para(doc, 'поставить товар в срок до')
    if p73:
        para_replace(p73, '«___» _____ 2025 г.', '{{execution_term}}')

    # P78: delivery address
    _, p78 = find_para(doc, 'по адресу')
    if p78:
        para_replace(p78, 'по адресу__________.', 'по адресу {{contractor_postal_address}}.')

    # ── Section 12: Requisites (P166-P169 empty) ───────────────────────
    _, p165 = find_para(doc, '12. Адреса и реквизиты сторон.')
    if p165:
        # Find the empty paragraphs after P165
        idx165, _ = find_para(doc, '12. Адреса и реквизиты сторон.')
        empty_paras = []
        for i in range(idx165 + 1, min(idx165 + 6, len(paras))):
            if paras[i].text.strip() == '':
                empty_paras.append(paras[i])

        reqs = [
            'Покупатель: Всероссийская общественная молодежная организация «ВСКС»\nАдрес: Пр-т Вернадского, д. 78, стр. 8, Москва, 119454\nТел.: (495) 568-00-11  e-mail: info@vsks.ru',
            '',
            'Поставщик: {{contractor_name}}\nАдрес: {{contractor_address}}\nПочтовый адрес: {{contractor_postal_address}}\nИНН/КПП: {{contractor_inn}} / {{contractor_kpp}}  ОГРН: {{contractor_ogrn}}\nр/с {{contractor_settlement_account}}\n{{contractor_bank_name}}\nБИК {{contractor_bik}}  к/с {{contractor_correspondent_account}}\nТел.: {{contractor_phone}}  e-mail: {{contractor_email}}',
            '',
        ]
        for i, p in enumerate(empty_paras[:len(reqs)]):
            if reqs[i]:
                para_set(p, reqs[i])

    # ── TZ Appendix ────────────────────────────────────────────────────
    # P172: "№___________ от_________"
    _, p172 = find_para(doc, '№___________')
    if p172:
        para_set(p172, '№ {{contract_number}} от {{contract_date}}')

    # P174: item subject line "на поставку брендированной продукции"
    _, p174 = find_para(doc, 'на поставку брендированной продукции')
    if p174:
        para_set(p174, 'на поставку: {{item_names}}')

    # After P174 insert items list with {%p for %} loop
    # Find P174 and insert after it
    _, p_subj = find_para(doc, 'на поставку: {{item_names}}')
    if not p_subj:
        _, p_subj = find_para(doc, 'Техническое задание')

    if p_subj:
        body = p_subj._p.getparent()
        ref = p_subj._p

        def ins_after(ref_el, text, bold=False):
            new_el = copy.deepcopy(ref_el)
            from docx.oxml.ns import qn as _qn
            for r in new_el.findall(_qn('w:r')):
                new_el.remove(r)
            ref_el.addnext(new_el)
            p_obj = DocxParagraph(new_el, p_subj._parent)
            run = p_obj.add_run(text)
            run.bold = bold
            return new_el

        # Insert in reverse order (each addnext inserts right after ref)
        # So we add last items first, then prepend

        # Items header
        h = ins_after(ref, 'Перечень поставляемых товаров (работ, услуг):', bold=True)
        # loop end marker (invisible, size 1)
        e = ins_after(h, '{%p endfor %}')
        DocxParagraph(e, p_subj._parent).runs[0].font.size = Pt(1)
        # item details line (this is the repeated paragraph)
        d = ins_after(h, '   Тип: {{item.type}}  |  Кол-во: {{item.quantity}} {{item.unit}}  |  Цена ед.: {{item.unit_price}}  |  Сумма: {{item.total_price}}')
        DocxParagraph(d, p_subj._parent).runs[0].font.size = Pt(9)
        # item description (repeated paragraph)
        desc_el = ins_after(h, '   {{item.description}}')
        DocxParagraph(desc_el, p_subj._parent).runs[0].font.size = Pt(9)
        # item name (repeated paragraph, bold)
        n = ins_after(h, '{{item.num}}. {{item.name}}')
        DocxParagraph(n, p_subj._parent).runs[0].bold = True
        # loop start marker (invisible, size 1)
        s = ins_after(h, '{%p for item in items %}')
        DocxParagraph(s, p_subj._parent).runs[0].font.size = Pt(1)

        # After all items, add total
        ins_after(e, 'Итого НМЦК: {{total_nmck}}')

    # ── Table 0 (main contract signatures) ─────────────────────────────
    if len(doc.tables) >= 1:
        t0 = doc.tables[0]
        # [1,1] Поставщик signatory
        cell_set(t0.rows[1].cells[1],
            '{{contractor_signatory}}\n\n____________________ / М.П.')

    # ── Table 1 (TZ appendix signatures) ───────────────────────────────
    if len(doc.tables) >= 2:
        t1 = doc.tables[1]
        cell_set(t1.rows[0].cells[1],
            'Поставщик\n{{contractor_name}}\n{{contractor_signatory}}\n\n____________________ / М.П.')

    out = os.path.join(TEMPLATES_DIR, 'contract_tz.docx')
    doc.save(out)
    print(f'[OK] {out}')


if __name__ == '__main__':
    missing = []
    for f in [SZ_SRC, LS_SRC, DP_SRC]:
        if not os.path.exists(f):
            missing.append(f)
    if missing:
        print('ОШИБКА: файлы не найдены:', missing)
        print('Скопируйте файлы в контейнер командами из комментария в начале скрипта.')
        raise SystemExit(1)

    make_service_note()
    make_approval_sheet()
    make_contract_tz()
    print('\nВсе шаблоны созданы.')
