"""
Генерирует файл ШАБЛОН_ПЕРЕМЕННЫЕ.docx — справочник всех переменных
для разметки шаблонов docxtpl (Word .docx).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE_DARK  = "1E3A5F"
BLUE_MID   = "2563EB"
BLUE_LIGHT = "DBEAFE"
GREY_LIGHT = "F3F4F6"

SECTIONS = [
    ("СИНТАКСИС docxtpl (Jinja2 для Word)", [
        ("Вставить значение",         "{{ имя_переменной }}",           "Подставляет текст. Пример: {{ contract_number }}"),
        ("Условие IF",                "{% if условие %} ... {% endif %}", "Показывает блок только при выполнении условия"),
        ("Условие IF / ELSE",         "{% if x %} А {% else %} Б {% endif %}", ""),
        ("Цикл по списку",            "{% for item in items %} ... {% endfor %}", "Повторяет блок для каждого элемента"),
        ("Вставить изображение",      "{{ photo }}",                   "Переменная типа InlineImage — вставляет фото"),
        ("Фильтр: первые N символов", "{{ text | truncate(50) }}",      ""),
        ("Фильтр: верхний регистр",   "{{ text | upper }}",             ""),
        ("Пустое если None",          "{{ value or '' }}",              ""),
    ]),

    ("ЗАКУПКА (Purchase)", [
        ("purchase_number",   "{{ purchase_number }}",   "Порядковый номер закупки"),
        ("registry_number",   "{{ registry_number }}",   "Реестровый номер"),
        ("purchase_method",   "{{ purchase_method }}",   "Способ закупки (текст): «Единственный поставщик» / «Конкурсная процедура»"),
        ("subject",           "{{ subject }}",           "Предмет договора / услуги"),
        ("service_name",      "{{ service_name }}",      "То же, что subject — алиас для договора"),
        ("status",            "{{ status }}",            "Статус закупки"),
        ("purchase_basis",    "{{ purchase_basis }}",    "Основание: «план-график» / «служебная записка»"),
        ("responsible_person","{{ responsible_person }}","ФИО ответственного исполнителя"),
        ("country_origin",    "{{ country_origin }}",    "Страна происхождения"),
        ("today",             "{{ today }}",             "Сегодняшняя дата: ДД.ММ.ГГГГ"),
        ("today_iso",         "{{ today_iso }}",         "Сегодняшняя дата: ГГГГ-ММ-ДД"),
    ]),

    ("СУБСИДИЯ (Subsidy)", [
        ("subsidy_name",   "{{ subsidy_name }}",   "Наименование субсидии"),
        ("subsidy_year",   "{{ subsidy_year }}",   "Год субсидии"),
        ("subsidy_budget", "{{ subsidy_budget }}", "Бюджет субсидии (форматированная сумма, ₽)"),
    ]),

    ("КОНТРАГЕНТ — основные реквизиты", [
        ("contractor_name",            "{{ contractor_name }}",            "Полное наименование контрагента"),
        ("contractor_short_name",      "{{ contractor_short_name }}",      "Краткое наименование (из кавычек или последнее слово)"),
        ("contractor_org_type",        "{{ contractor_org_type }}",        "Тип: «Юр.лицо» / «ИП» / «Самозанятый»"),
        ("contractor_inn",             "{{ contractor_inn }}",             "ИНН"),
        ("contractor_kpp",             "{{ contractor_kpp }}",             "КПП (для юр.лиц)"),
        ("contractor_ogrn",            "{{ contractor_ogrn }}",            "ОГРН / ОГРНИП"),
        ("contractor_address",         "{{ contractor_address }}",         "Юридический адрес"),
        ("contractor_postal_address",  "{{ contractor_postal_address }}",  "Почтовый адрес"),
        ("contractor_phone",           "{{ contractor_phone }}",           "Телефон"),
        ("contractor_email",           "{{ contractor_email }}",           "E-mail"),
    ]),

    ("КОНТРАГЕНТ — подписант и банк", [
        ("contractor_signatory",          "{{ contractor_signatory }}",          "ФИО подписанта"),
        ("contractor_signatory_basis",    "{{ contractor_signatory_basis }}",    "Основание полномочий: «Устава», «Доверенности №…»"),
        ("contractor_signatory_position", "{{ contractor_signatory_position }}", "Должность подписанта (из поля signatory)"),
        ("contractor_signatory_line",     "{{ contractor_signatory_line }}",     "Готовая строка: «ФИО, действует на основании …»"),
        ("contractor_settlement_account", "{{ contractor_settlement_account }}", "Расчётный счёт"),
        ("contractor_bank_name",          "{{ contractor_bank_name }}",          "Наименование банка"),
        ("contractor_bank_details",       "{{ contractor_bank_details }}",       "Алиас для bank_name"),
        ("contractor_bik",                "{{ contractor_bik }}",                "БИК банка"),
        ("contractor_correspondent_account", "{{ contractor_correspondent_account }}", "Корреспондентский счёт"),
    ]),

    ("ДОГОВОР — числа и даты", [
        ("contract_number",       "{{ contract_number }}",       "Номер договора"),
        ("contract_date",         "{{ contract_date }}",         "Дата договора: ДД.ММ.ГГГГ"),
        ("contract_date_day",     "{{ contract_date_day }}",     "День договора (двузначный): «05»"),
        ("contract_date_month",   "{{ contract_date_month }}",   "Месяц прописью: «марта»"),
        ("contract_date_year",    "{{ contract_date_year }}",    "Год договора: «2026»"),
        ("execution_term",        "{{ execution_term }}",        "Срок исполнения: ДД.ММ.ГГГГ"),
        ("execution_term_changed","{{ execution_term_changed }}", "Изменённый срок исполнения"),
        ("delivery_date",         "{{ delivery_date }}",         "Дата поставки"),
        ("contract_end_date",     "{{ contract_end_date }}",     "Срок действия договора (из параметров договора)"),
    ]),

    ("ДОГОВОР — суммы", [
        ("contract_price",       "{{ contract_price }}",       "Цена договора, форматированная с ₽"),
        ("contract_price_num",   "{{ contract_price_num }}",   "Цена без ₽, с пробелами: «125 000,00»"),
        ("contract_price_words", "{{ contract_price_words }}", "Цена прописью: «сто двадцать пять тысяч рублей 00 копеек»"),
        ("total_nmck",           "{{ total_nmck }}",           "НМЦК итого с ₽"),
        ("nmck",                 "{{ nmck }}",                 "НМЦК (алиас)"),
        ("economy",              "{{ economy }}",              "Экономия"),
        ("price_increase",       "{{ price_increase }}",       "Удорожание"),
    ]),

    ("ДОГОВОР — НДС", [
        ("vat_applicable",        "{{ vat_applicable }}",        "True/False — применяется ли НДС"),
        ("vat_rate",              "{{ vat_rate }}",              "Ставка НДС: 20 или 10"),
        ("vat_amount_num",        "{{ vat_amount_num }}",        "Сумма НДС без ₽: «20 833,33»"),
        ("vat_amount_words",      "{{ vat_amount_words }}",      "Сумма НДС прописью"),
        ("vat_exemption_article", "{{ vat_exemption_article }}", "Статья НК об освобождении от НДС"),
        ("vat_info_line",         "{{ vat_info_line }}",         "Готовая строка: «В том числе НДС 20%: …» или «НДС не облагается»"),
        ("", "ПРИМЕР использования:", ""),
        ("",
         "{% if vat_applicable %}\n  В т.ч. НДС {{ vat_rate }}%: {{ vat_amount_num }} руб.\n{% else %}\n  {{ vat_info_line }}\n{% endif %}",
         ""),
    ]),

    ("ДОГОВОР — срок оказания услуг", [
        ("period_type",      "{{ period_type }}",      "«period» — период, «date» — разовая дата"),
        ("service_start_date","{{ service_start_date }}","Начало периода: ДД.ММ.ГГГГ"),
        ("service_end_date", "{{ service_end_date }}",  "Конец периода: ДД.ММ.ГГГГ"),
        ("service_date",     "{{ service_date }}",      "Разовая дата оказания услуги"),
        ("third_party_involved", "{{ third_party_involved }}", "True/False — привлекаются ли третьи лица"),
        ("", "ПРИМЕР:", ""),
        ("",
         "{% if period_type == 'period' %}\n  с {{ service_start_date }} по {{ service_end_date }}\n{% else %}\n  {{ service_date }}\n{% endif %}",
         ""),
    ]),

    ("АКТЫ И ПЛАТЁЖ", [
        ("acceptance_doc_name",   "{{ acceptance_doc_name }}",   "Наименование документа приёмки"),
        ("acceptance_doc_number", "{{ acceptance_doc_number }}", "Номер документа приёмки"),
        ("acceptance_doc_date",   "{{ acceptance_doc_date }}",   "Дата документа приёмки"),
        ("acceptance_doc_amount", "{{ acceptance_doc_amount }}", "Сумма по документу приёмки"),
        ("payment_doc_number",    "{{ payment_doc_number }}",    "Номер платёжного поручения"),
        ("payment_doc_date",      "{{ payment_doc_date }}",      "Дата платёжного поручения"),
        ("payment_amount",        "{{ payment_amount }}",        "Сумма оплаты"),
        ("payment_federal",       "{{ payment_federal }}",       "Федеральная часть оплаты"),
    ]),

    ("ПОЗИЦИИ ТЗ (цикл items)", [
        ("items_count",  "{{ items_count }}",  "Количество позиций"),
        ("item_names",   "{{ item_names }}",   "Все наименования через запятую"),
        ("", "--- ЦИКЛ ---", ""),
        ("",
         "{% for item in items %}\n  ...\n{% endfor %}",
         "Повторяет блок для каждой позиции"),
        ("item.num",         "{{ item.num }}",         "Порядковый номер (1, 2, 3…)"),
        ("item.name",        "{{ item.name }}",        "Наименование товара/услуги"),
        ("item.description", "{{ item.description }}", "Описание (точное или по 44-ФЗ)"),
        ("item.type",        "{{ item.type }}",        "Тип позиции"),
        ("item.quantity",    "{{ item.quantity }}",    "Количество"),
        ("item.unit",        "{{ item.unit }}",        "Единица измерения"),
        ("item.unit_price",  "{{ item.unit_price }}",  "Цена за единицу, ₽"),
        ("item.total_price", "{{ item.total_price }}", "Стоимость позиции, ₽"),
        ("item.photo",       "{{ item.photo }}",       "Фото (InlineImage 2.5 см) — вставляется как картинка"),
    ]),

    ("СОГЛАСУЮЩИЕ (цикл approvers)", [
        ("",
         "{% for a in approvers %}\n  ...\n{% endfor %}",
         "Повторяет блок для каждого согласующего"),
        ("a.num",           "{{ a.num }}",           "Порядковый номер"),
        ("a.role_name",     "{{ a.role_name }}",     "Должность/роль"),
        ("a.full_name",     "{{ a.full_name }}",     "ФИО"),
        ("a.signature_img", "{{ a.signature_img }}", "Электронная подпись (InlineImage) — если есть"),
        ("a.decided_date",  "{{ a.decided_date }}",  "Дата подписания"),
        ("a.note",          "{{ a.note }}",          "Примечание (путь ФЭО, если включено)"),
        ("initiator_name",  "{{ initiator_name }}",  "ФИО инициатора служебной записки"),
        ("initiator_role",  "{{ initiator_role }}",  "Должность инициатора"),
    ]),

    ("ФЭО И ПРОЧЕЕ", [
        ("feo_category_name", "{{ feo_category_name }}", "Название раздела ФЭО"),
    ]),

    ("ТИПЫ ДОКУМЕНТОВ (doc_type)", [
        ("service_note",          "service_note.docx",          "Служебная записка на организацию закупки"),
        ("service_note_delivery", "service_note_delivery.docx", "Служебная записка на выдачу"),
        ("service_note_payment",  "service_note_payment.docx",  "Служебная записка на оплату"),
        ("contract_tz",           "contract_tz.docx",           "Техническое задание (отдельный файл)"),
        ("contract",              "contract.docx",              "Договор"),
        ("contract_fadm",         "contract_fadm.docx",         "Договор ФАДМ"),
        ("approval_sheet",        "approval_sheet.docx",        "Лист согласования"),
    ]),

    ("СУБСИДИЯ-СПЕЦИФИЧНЫЕ ШАБЛОНЫ", [
        ("", "Путь:", "backend/templates/subsidies/{subsidy_id}/{doc_type}.docx"),
        ("", "Пример:",
         "backend/templates/subsidies/3/contract.docx — договор для субсидии с id=3"),
        ("", "Приоритет:",
         "Если файл существует — используется вместо базового шаблона. Если нет — берётся стандартный."),
    ]),
]


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def add_section(doc, title, rows):
    # Section heading
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(2)
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Имя переменной", "Синтаксис в шаблоне", "Описание"]):
        set_cell_bg(cell, BLUE_DARK)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    col_widths = [Cm(4.2), Cm(6.8), Cm(7.5)]
    for cell, w in zip(hdr, col_widths):
        cell.width = w

    for i, (var, syntax, desc) in enumerate(rows):
        row = table.add_row()
        bg = GREY_LIGHT if i % 2 == 0 else "FFFFFF"

        # Пустая переменная = строка-разделитель или заголовок
        if var == "" and syntax.startswith("---"):
            # separator header
            for cell in row.cells:
                set_cell_bg(cell, BLUE_LIGHT)
            merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            p = merged.paragraphs[0]
            r = p.add_run(syntax.strip("-").strip())
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if var == "" and syntax.startswith("{%"):
            # code example block spanning all columns
            row.cells[0].text = ""
            cell_code = row.cells[1].merge(row.cells[2])
            for cell in row.cells:
                set_cell_bg(cell, "FEF9C3")  # light yellow
            p = cell_code.paragraphs[0]
            r = p.add_run(syntax)
            r.font.name = "Courier New"
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x78, 0x35, 0x00)
            continue

        if var == "":
            # label row (Путь, Пример и т.п.)
            for cell in row.cells:
                set_cell_bg(cell, BLUE_LIGHT)
            row.cells[0].text = syntax
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            merged = row.cells[1].merge(row.cells[2])
            p = merged.paragraphs[0]
            r = p.add_run(desc)
            r.font.name = "Courier New"
            r.font.size = Pt(9)
            continue

        for cell in row.cells:
            set_cell_bg(cell, bg)

        # Col 0: var name
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(var)
        r0.font.size = Pt(9)
        r0.font.color.rgb = RGBColor(0x16, 0x65, 0x34)  # green

        # Col 1: syntax
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(syntax)
        r1.font.name = "Courier New"
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(0x17, 0x45, 0xAA)  # blue

        # Col 2: description
        p2 = row.cells[2].paragraphs[0]
        r2 = p2.add_run(desc)
        r2.font.size = Pt(9)

        for cell, w in zip(row.cells, col_widths):
            cell.width = w


doc = Document()

# Page margins
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
section = doc.sections[0]
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Title
title = doc.add_heading("Справочник переменных шаблонов VSKS CRM", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

subtitle = doc.add_paragraph("Используйте эти обозначения в Word-шаблонах (.docx) через docxtpl.")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(10)
subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

for section_title, rows in SECTIONS:
    add_section(doc, section_title, rows)

out_path = r"C:\Users\1\VSKS_CRM\backend\templates\ШАБЛОН_ПЕРЕМЕННЫЕ.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
