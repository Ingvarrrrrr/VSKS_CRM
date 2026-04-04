"""
Run inside backend container:
    docker exec vsks_crm-backend-1 sh -c 'cd /app && python templates/generate_templates.py'
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def page_margins(doc):
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)


def add_label_value(doc, label, value, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.size = Pt(size)
    r2 = p.add_run(value)
    r2.font.size = Pt(size)
    return p


def hline(doc):
    """Thin horizontal separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── contract_tz.docx ─────────────────────────────────────────────────────────
# Items rendered as a TABLE with photo column.
# Uses {%tr for item in items %} for row repetition.
# Photo column uses InlineImage — passed from documents.py.
def make_contract_tz():
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE

    def _set_col_width(cell, width_cm):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = _OE('w:tcW')
        tcW.set(_qn('w:w'), str(int(width_cm * 567)))  # 567 twips/cm
        tcW.set(_qn('w:type'), 'dxa')
        tcPr.append(tcW)

    doc = Document()
    page_margins(doc)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ТЕХНИЧЕСКОЕ ЗАДАНИЕ")
    r.bold = True; r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("к договору № {{contract_number}} от {{contract_date}}").font.size = Pt(11)

    doc.add_paragraph()

    add_label_value(doc, "Субсидия:", "{{subsidy_name}} ({{subsidy_year}})")
    add_label_value(doc, "Реестровый номер:", "{{registry_number}}")
    add_label_value(doc, "Способ закупки:", "{{purchase_method}}")
    add_label_value(doc, "Контрагент:", "{{contractor_name}}, ИНН {{contractor_inn}} / КПП {{contractor_kpp}}")
    add_label_value(doc, "Срок исполнения:", "{{execution_term}}")

    doc.add_paragraph()
    hline(doc)

    hdr = doc.add_paragraph()
    r = hdr.add_run("Перечень поставляемых товаров (работ, услуг):")
    r.bold = True; r.font.size = Pt(11)

    # ── Items TABLE with photo ─────────────────────────────────
    # docxtpl requires {%tr for %} and {%tr endfor %} in SEPARATE rows (3-row pattern):
    #   Row 0 (header): Фото | № | ... (blue header)
    #   Row 1 (for):    {%tr for item in items %} | ... (replaced by jinja2 tag)
    #   Row 2 (data):   {{item.photo}} | {{item.num}} | ... (repeated per item)
    #   Row 3 (endfor): {%tr endfor %} | ... (replaced by jinja2 tag)
    # Columns: Фото | № | Наименование / Описание | Тип | Кол-во | Ед. | Цена ед. | Сумма
    COL_WIDTHS = [2.8, 0.8, 7.0, 1.8, 1.4, 1.0, 2.2, 2.2]  # cm
    tbl = doc.add_table(rows=4, cols=8)
    tbl.style = 'Table Grid'

    # Row 0: Header
    hdr_texts = ["Фото", "№", "Наименование / Описание", "Тип", "Кол-во", "Ед.", "Цена ед., руб.", "Сумма, руб."]
    for i, (txt, w) in enumerate(zip(hdr_texts, COL_WIDTHS)):
        cell = tbl.rows[0].cells[i]
        _set_col_width(cell, w)
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True; run.font.size = Pt(9)
        set_cell_bg(cell, "D6E4F7")

    # Row 1: {%tr for item in items %} — this entire row is replaced by jinja2 for tag
    tbl.rows[1].cells[0].paragraphs[0].add_run("{%tr for item in items %}").font.size = Pt(8)

    # Row 2: Data template row (gets repeated per item)
    data_cells = tbl.rows[2].cells
    _set_col_width(data_cells[0], COL_WIDTHS[0])

    p0 = data_cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.add_run("{{item.photo}}").font.size = Pt(9)

    data_cells[1].paragraphs[0].add_run("{{item.num}}").font.size = Pt(9)

    p2 = data_cells[2].paragraphs[0]
    run_name = p2.add_run("{{item.name}}")
    run_name.bold = True; run_name.font.size = Pt(9)
    p2.add_run("\n{{item.description}}").font.size = Pt(8)

    data_cells[3].paragraphs[0].add_run("{{item.type}}").font.size = Pt(9)
    data_cells[4].paragraphs[0].add_run("{{item.quantity}}").font.size = Pt(9)
    data_cells[5].paragraphs[0].add_run("{{item.unit}}").font.size = Pt(9)
    data_cells[6].paragraphs[0].add_run("{{item.unit_price}}").font.size = Pt(9)
    data_cells[7].paragraphs[0].add_run("{{item.total_price}}").font.size = Pt(9)

    # Row 3: {%tr endfor %} — this entire row is replaced by jinja2 endfor tag
    tbl.rows[3].cells[0].paragraphs[0].add_run("{%tr endfor %}").font.size = Pt(8)
    # ── End items TABLE ────────────────────────────────────────

    doc.add_paragraph()

    # Total
    tot = doc.add_paragraph()
    tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = tot.add_run("Итого НМЦК: ")
    r1.bold = True; r1.font.size = Pt(12)
    r2 = tot.add_run("{{total_nmck}}")
    r2.bold = True; r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    doc.add_paragraph()
    add_label_value(doc, "Цена договора:", "{{contract_price}}")
    add_label_value(doc, "Экономия:", "{{economy}}")

    hline(doc)

    # Signatures
    hdr2 = doc.add_paragraph()
    r = hdr2.add_run("Подписи сторон")
    r.bold = True; r.font.size = Pt(11)

    sig = doc.add_table(rows=4, cols=2)
    sig.style = 'Table Grid'

    def sig_cell(row, col, text, bold=False):
        p = sig.rows[row].cells[col].paragraphs[0]
        run = p.add_run(text)
        run.bold = bold; run.font.size = Pt(9)

    sig_cell(0, 0, "Заказчик", bold=True)
    sig_cell(0, 1, "Исполнитель", bold=True)
    sig_cell(1, 0, "Всероссийская общественная молодежная организация «ВСКС»\nИНН: 7704190720 / КПП: 770401001\nОГРН: 1037739183516")
    sig_cell(1, 1, "{{contractor_name}}\nИНН: {{contractor_inn}} / КПП: {{contractor_kpp}}\nОГРН: {{contractor_ogrn}}")
    sig_cell(2, 0,
        "Адрес: Пр-т Вернадского, д. 78, стр. 8, Москва, 119454\n"
        "р/с 40703810138060100002\nПАО Сбербанк г. Москва\n"
        "БИК 044525225 | к/с 30101810400000000225"
    )
    sig_cell(2, 1,
        "Адрес: {{contractor_address}}\n"
        "{{contractor_bank_details}}"
    )
    sig_cell(3, 0, "Председатель Совета _______________  /  М.П.")
    sig_cell(3, 1, "{{contractor_signatory_line}}\nПодпись: _______________ / М.П.")

    path = os.path.join(TEMPLATES_DIR, "contract_tz.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ── service_note.docx ─────────────────────────────────────────────────────────
def make_service_note():
    doc = Document()
    page_margins(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("СЛУЖЕБНАЯ ЗАПИСКА")
    r.bold = True; r.font.size = Pt(14)

    doc.add_paragraph()

    add_label_value(doc, "Дата:", "{{today}}")
    add_label_value(doc, "Субсидия:", "{{subsidy_name}} ({{subsidy_year}})")
    add_label_value(doc, "Реестровый номер:", "{{registry_number}}")
    add_label_value(doc, "Контрагент:", "{{contractor_name}}, ИНН {{contractor_inn}}")
    add_label_value(doc, "Способ закупки:", "{{purchase_method}}")
    add_label_value(doc, "НМЦК:", "{{total_nmck}}")
    add_label_value(doc, "Цена договора:", "{{contract_price}}")
    add_label_value(doc, "Экономия:", "{{economy}}")
    add_label_value(doc, "Страна происхождения:", "{{country_origin}}")
    add_label_value(doc, "Срок исполнения:", "{{execution_term}}")
    add_label_value(doc, "Договор:", "№{{contract_number}} от {{contract_date}}")

    doc.add_paragraph()

    body = doc.add_paragraph()
    body.add_run(
        "Прошу согласовать закупку №{{purchase_number}} (реестр {{registry_number}}) "
        "у контрагента {{contractor_name}} на сумму {{total_nmck}} "
        "в рамках субсидии «{{subsidy_name}}»."
    ).font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph()
    sig.add_run("Исполнитель: _______________  /  ФИО  /  {{today}}").font.size = Pt(10)

    path = os.path.join(TEMPLATES_DIR, "service_note.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ── approval_sheet.docx ───────────────────────────────────────────────────────
def make_approval_sheet():
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE

    def _set_col_width(cell, width_cm):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = _OE('w:tcW')
        tcW.set(_qn('w:w'), str(int(width_cm * 567)))
        tcW.set(_qn('w:type'), 'dxa')
        tcPr.append(tcW)

    doc = Document()
    page_margins(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ЛИСТ СОГЛАСОВАНИЯ")
    r.bold = True; r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Закупка №{{purchase_number}} / Реестр {{registry_number}}").font.size = Pt(11)

    doc.add_paragraph()

    add_label_value(doc, "Субсидия:", "{{subsidy_name}} ({{subsidy_year}})")
    add_label_value(doc, "Контрагент:", "{{contractor_name}}, ИНН {{contractor_inn}}")
    add_label_value(doc, "НМЦК:", "{{total_nmck}}")
    add_label_value(doc, "Цена договора:", "{{contract_price}}")
    add_label_value(doc, "Договор:", "№{{contract_number}} от {{contract_date}}")
    add_label_value(doc, "Срок исполнения:", "{{execution_term}}")

    doc.add_paragraph()

    # Dynamic approvers table — 3-row pattern (for / data / endfor)
    COL_WIDTHS = [1.0, 5.5, 5.5, 5.0]  # cm: №, Должность, ФИО, Подпись/Дата
    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = 'Table Grid'

    # Row 0: Header
    headers = ["№", "Должность", "ФИО", "Подпись / Дата"]
    for i, (h, w) in enumerate(zip(headers, COL_WIDTHS)):
        cell = tbl.rows[0].cells[i]
        _set_col_width(cell, w)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(9)
        set_cell_bg(cell, "D6E4F7")

    # Row 1: {%tr for a in approvers %} — replaced by jinja2 for tag
    tbl.rows[1].cells[0].paragraphs[0].add_run("{%tr for a in approvers %}").font.size = Pt(8)

    # Row 2: Data template row
    data_row = tbl.rows[2]
    _set_col_width(data_row.cells[0], COL_WIDTHS[0])
    data_row.cells[0].paragraphs[0].add_run("{{a.num}}").font.size = Pt(9)
    data_row.cells[1].paragraphs[0].add_run("{{a.role_name}}").font.size = Pt(9)
    data_row.cells[2].paragraphs[0].add_run("{{a.full_name}}").font.size = Pt(9)
    data_row.cells[3].paragraphs[0].add_run("").font.size = Pt(9)

    # Row 3: {%tr endfor %} — replaced by jinja2 endfor tag
    tbl.rows[3].cells[0].paragraphs[0].add_run("{%tr endfor %}").font.size = Pt(8)

    path = os.path.join(TEMPLATES_DIR, "approval_sheet.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ── service_note_delivery.docx ────────────────────────────────────────────────
def make_service_note_delivery():
    doc = Document()
    page_margins(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("СЛУЖЕБНАЯ ЗАПИСКА")
    r.bold = True; r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("на выдачу поставленного товара / услуги").font.size = Pt(11)

    doc.add_paragraph()

    add_label_value(doc, "Дата:", "{{today}}")
    add_label_value(doc, "Субсидия:", "{{subsidy_name}} ({{subsidy_year}})")
    add_label_value(doc, "Реестровый номер:", "{{registry_number}}")
    add_label_value(doc, "Контрагент:", "{{contractor_name}}, ИНН {{contractor_inn}}")
    add_label_value(doc, "Договор:", "№{{contract_number}} от {{contract_date}}")
    add_label_value(doc, "Дата поставки / исполнения:", "{{delivery_date}}")
    add_label_value(doc, "Акт приёмки:", "{{acceptance_doc_name}} №{{acceptance_doc_number}} от {{acceptance_doc_date}}")
    add_label_value(doc, "Сумма по акту:", "{{acceptance_doc_amount}}")

    doc.add_paragraph()

    body = doc.add_paragraph()
    body.add_run(
        "Прошу принять и выдать поставленные товары (работы, услуги) по закупке "
        "№{{purchase_number}} (реестр {{registry_number}}) "
        "у контрагента {{contractor_name}} на основании договора №{{contract_number}} "
        "от {{contract_date}} на сумму {{acceptance_doc_amount}}."
    ).font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph()
    sig.add_run("Исполнитель: _______________  /  ФИО  /  {{today}}").font.size = Pt(10)

    path = os.path.join(TEMPLATES_DIR, "service_note_delivery.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ── service_note_payment.docx ─────────────────────────────────────────────────
def make_service_note_payment():
    doc = Document()
    page_margins(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("СЛУЖЕБНАЯ ЗАПИСКА")
    r.bold = True; r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("на оплату поставленного товара / услуги").font.size = Pt(11)

    doc.add_paragraph()

    add_label_value(doc, "Дата:", "{{today}}")
    add_label_value(doc, "Субсидия:", "{{subsidy_name}} ({{subsidy_year}})")
    add_label_value(doc, "Реестровый номер:", "{{registry_number}}")
    add_label_value(doc, "Контрагент:", "{{contractor_name}}, ИНН {{contractor_inn}}")
    add_label_value(doc, "Договор:", "№{{contract_number}} от {{contract_date}}")
    add_label_value(doc, "Сумма договора:", "{{contract_price}}")
    add_label_value(doc, "Акт приёмки:", "{{acceptance_doc_name}} №{{acceptance_doc_number}} от {{acceptance_doc_date}}")
    add_label_value(doc, "Сумма к оплате:", "{{acceptance_doc_amount}}")

    hline(doc)

    bank = doc.add_paragraph()
    bank.add_run("Реквизиты получателя:").bold = True
    bank.runs[0].font.size = Pt(10)

    add_label_value(doc, "Расчётный счёт:", "{{contractor_settlement_account}}")
    add_label_value(doc, "Банк:", "{{contractor_bank_name}}")
    add_label_value(doc, "БИК:", "{{contractor_bik}}")

    doc.add_paragraph()

    body = doc.add_paragraph()
    body.add_run(
        "Прошу произвести оплату по закупке №{{purchase_number}} (реестр {{registry_number}}) "
        "контрагенту {{contractor_name}} на сумму {{acceptance_doc_amount}} "
        "в соответствии с договором №{{contract_number}} от {{contract_date}} "
        "и актом приёмки {{acceptance_doc_name}} №{{acceptance_doc_number}} от {{acceptance_doc_date}}."
    ).font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph()
    sig.add_run("Исполнитель: _______________  /  ФИО  /  {{today}}").font.size = Pt(10)

    path = os.path.join(TEMPLATES_DIR, "service_note_payment.docx")
    doc.save(path)
    print(f"[OK] {path}")


if __name__ == "__main__":
    make_contract_tz()
    make_service_note()
    make_service_note_delivery()
    make_service_note_payment()
    make_approval_sheet()
    print("All templates generated.")
