"""
Пересоздать шаблон approval_sheet.docx с 5 колонками:
  №  |  Должность / ФИО  |  Подпись  |  Дата  |  Примечание
+ вставить {{a.note}}, {{a.signature_img}}, {{a.decided_date}}
"""
import os
import copy
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SRC  = os.path.join(SCRIPT_DIR, "approval_sheet.docx")
DEST = os.path.join(SCRIPT_DIR, "approval_sheet.docx")  # overwrite in-place

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def set_cell_text(cell, text: str, bold=False, size_pt=9, align="left"):
    """Clear cell and set plain text with given formatting."""
    tc = cell._tc
    # Remove all paragraphs except first
    paras = tc.findall(qn("w:p"))
    for p in paras[1:]:
        tc.remove(p)
    p = paras[0]
    # Clear runs
    for r in p.findall(qn("w:r")):
        p.remove(r)
    for pPr in p.findall(qn("w:pPr")):
        p.remove(pPr)

    # Paragraph properties (alignment)
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), align)
    pPr.append(jc)
    p.insert(0, pPr)

    # Run
    if text:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if bold:
            b = OxmlElement("w:b"); rPr.append(b)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size_pt * 2)); rPr.append(sz)
        szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(size_pt * 2)); rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement("w:t"); t.text = text
        if text.startswith(" ") or text.endswith(" "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)


def clone_tc(source_tc):
    """Deep-copy a table cell."""
    return copy.deepcopy(source_tc)


def make_empty_tc(width_twips: int) -> etree._Element:
    """Create a minimal empty table cell with given width."""
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    tc.append(tcPr)
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)
    tc.append(p)
    return tc


def add_tc_to_row(tr, width_twips: int, text: str = "", bold=False, size_pt=9):
    """Add a new cell at end of row."""
    from docx.table import _Cell
    tc = make_empty_tc(width_twips)
    tr.append(tc)
    cell = _Cell(tc, None)
    if text:
        set_cell_text(cell, text, bold=bold, size_pt=size_pt)
    return cell


def main():
    doc = Document(SRC)

    # Find the approvers table (Table 1 — has {%tr for a in approvers %})
    approvers_table = None
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "{%tr for a in approvers %}" in cell.text or "for a in approvers" in cell.text:
                    approvers_table = t
                    break
            if approvers_table:
                break
        if approvers_table:
            break

    if not approvers_table:
        print("ERROR: Approvers table not found!")
        return

    print(f"Found approvers table with {len(approvers_table.rows)} rows")
    for i, row in enumerate(approvers_table.rows):
        print(f"  Row {i}: {[c.text[:40] for c in row.cells]}")

    tbl = approvers_table._tbl

    # ── Step 1: Update header row (Row 0) ──────────────────────────────────
    header_row = approvers_table.rows[0]
    # Current cols: №, Должность/ФИО, Подпись/Дата  (or №, Должность, ФИО, Подпись/Дата)
    # We need: №, Должность/ФИО, Подпись, Дата, Примечание
    tr0 = header_row._tr
    cells0 = tr0.findall(qn("w:tc"))
    n_cols = len(cells0)
    print(f"Header row has {n_cols} cells")

    if n_cols == 4:
        # Has: №, Должность, ФИО, Подпись/Дата
        # Rename col 3 → Подпись, add col 4 Дата, add col 5 Примечание
        # Rename col 2 → "Должность / ФИО" (merge header text)
        from docx.table import _Cell
        cell_sign = _Cell(cells0[3], None)
        set_cell_text(cell_sign, "Подпись", bold=True, size_pt=9, align="center")

        # Add Дата column
        date_tc = make_empty_tc(1000)
        tr0.append(date_tc)
        date_cell = _Cell(date_tc, None)
        set_cell_text(date_cell, "Дата", bold=True, size_pt=9, align="center")

        # Add Примечание column
        note_tc = make_empty_tc(2200)
        tr0.append(note_tc)
        note_cell = _Cell(note_tc, None)
        set_cell_text(note_cell, "Примечание", bold=True, size_pt=9, align="center")

    elif n_cols == 3:
        # Has: №, Должность/ФИО, Подпись/Дата — split last into Подпись + Дата
        from docx.table import _Cell
        cell_sign = _Cell(cells0[2], None)
        set_cell_text(cell_sign, "Подпись", bold=True, size_pt=9, align="center")

        date_tc = make_empty_tc(1000)
        tr0.append(date_tc)
        date_cell = _Cell(date_tc, None)
        set_cell_text(date_cell, "Дата", bold=True, size_pt=9, align="center")

        note_tc = make_empty_tc(2200)
        tr0.append(note_tc)
        note_cell = _Cell(note_tc, None)
        set_cell_text(note_cell, "Примечание", bold=True, size_pt=9, align="center")

    else:
        print(f"Unexpected column count {n_cols}, adding Дата and Примечание")
        from docx.table import _Cell
        date_tc = make_empty_tc(1000)
        tr0.append(date_tc)
        date_cell = _Cell(date_tc, None)
        set_cell_text(date_cell, "Дата", bold=True, size_pt=9, align="center")

        note_tc = make_empty_tc(2200)
        tr0.append(note_tc)
        note_cell = _Cell(note_tc, None)
        set_cell_text(note_cell, "Примечание", bold=True, size_pt=9, align="center")

    # ── Step 2: Update data row (Row 2 — between for/endfor) ────────────────
    # Row 1 = {%tr for %}, Row 2 = data, Row 3 = {%tr endfor %}
    data_row = approvers_table.rows[2]
    tr2 = data_row._tr
    cells2 = tr2.findall(qn("w:tc"))
    n_data_cols = len(cells2)
    print(f"Data row has {n_data_cols} cells: {[data_row.cells[i].text for i in range(n_data_cols)]}")

    from docx.table import _Cell

    # Update signature cell: add {{a.signature_img}}
    if n_data_cols >= 3:
        sig_cell = _Cell(cells2[n_data_cols - 1], None)
        set_cell_text(sig_cell, "{{a.signature_img}}", size_pt=9, align="center")

    # Add Дата cell
    date_tc2 = make_empty_tc(1000)
    tr2.append(date_tc2)
    date_cell2 = _Cell(date_tc2, None)
    set_cell_text(date_cell2, "{{a.decided_date}}", size_pt=9, align="center")

    # Add Примечание cell
    note_tc2 = make_empty_tc(2200)
    tr2.append(note_tc2)
    note_cell2 = _Cell(note_tc2, None)
    set_cell_text(note_cell2, "{{a.note}}", size_pt=9)

    # ── Step 3: Update for/endfor rows — add empty cells to match column count ─
    for row_idx in [1, 3]:
        tr = approvers_table.rows[row_idx]._tr
        cells = tr.findall(qn("w:tc"))
        # Add 2 empty cells to match new column count
        for _ in range(2):
            empty_tc = make_empty_tc(1000)
            tr.append(empty_tc)

    print("Table updated successfully.")

    # ── Save ────────────────────────────────────────────────────────────────
    doc.save(DEST)
    print(f"Saved: {DEST}")

    # Verify
    doc2 = Document(DEST)
    for t in doc2.tables:
        for row in t.rows:
            for cell in row.cells:
                if "for a in approvers" in cell.text or "a.note" in cell.text or "a.signature" in cell.text:
                    print(f"VERIFY - found template var in: {[c.text[:50] for c in row.cells]}")


if __name__ == "__main__":
    main()
