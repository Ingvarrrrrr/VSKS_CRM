from fastapi import APIRouter, Depends, Query, HTTPException, status, UploadFile, File
from pydantic import BaseModel
from fastapi.responses import StreamingResponse
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession
from app.database import get_db
from app.models.feo_category import FeoCategory
from app.schemas.schemas import FeoCategoryOut, FeoCategoryCreate
from app.auth.jwt import get_current_user, require_role, get_org_filter, ADMIN_ROLES, ALL_ROLES
from app.auth.permissions import require_tab
from app.auth.visibility import get_visible_subsidy_ids
from typing import List, Optional
from decimal import Decimal
from io import BytesIO
try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.worksheet.datavalidation import DataValidation
except ImportError:
    Workbook = None
    load_workbook = None
    DataValidation = None

router = APIRouter(prefix="/api/feo-categories", tags=["feo_categories"])


def _content_disposition(filename: str) -> str:
    """RFC 5987 — кириллица в имени файла недопустима в latin-1 заголовке."""
    from urllib.parse import quote
    ascii_fallback = filename.encode('ascii', 'ignore').decode('ascii').strip() or 'export'
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{quote(filename)}"


def _validate_plan_pair(planned_quantity: Optional[float], planned_amount: Optional[float]) -> None:
    """Правило владельца (2026-08-09): плановое количество и плановая цена за
    единицу — ПАРА. Если задана цена за единицу, обязано быть задано и
    количество (тогда сумма = кол-во × цена считается автоматически). Если
    задано количество без цены — то же самое наоборот. Пусто-пусто — план по
    этому листу просто не задан руками (допустимо, план может считаться из
    детей/факта). Оба заполнены — допустимо.

    «Заполнено» = число > 0 (не просто not None) — совпадает с порогом
    app.services.feo_plan.compute_feo_plan_tree._visit (qty > 0 and amt > 0),
    иначе planned_quantity=0 с ценой прошло бы валидацию, но провалилось бы в
    формуле плана (фолбэк на сумму активных FeoPlannedItem, как и «пусто-пусто» —
    молчаливый и неожиданный для пользователя результат).

    Альтернатива для плана ОБЩЕЙ суммой без разбивки на кол-во/цену (напр.
    «Канцтовары» на 1 000 000 без детализации) — не эта пара полей категории, а
    плановая позиция (FeoPlannedItem, кнопка «Добавить плановую» в панели) с
    суммой amount и БЕЗ quantity.
    """
    qty_filled = planned_quantity is not None and float(planned_quantity) > 0
    amt_filled = planned_amount is not None and float(planned_amount) > 0
    if qty_filled == amt_filled:
        return
    if qty_filled:
        detail = (
            "Задано плановое количество, но не задана плановая стоимость за единицу. "
            "Заполните оба поля («Плановое количество» и «Плановая стоимость за ед.») — "
            "тогда сумма посчитается автоматически, — либо очистите количество и задайте "
            "план общей суммой отдельной плановой позицией (кнопка «Добавить плановую» "
            "в панели) без указания количества."
        )
    else:
        detail = (
            "Задана плановая стоимость за единицу, но не задано плановое количество. "
            "Заполните оба поля («Плановое количество» и «Плановая стоимость за ед.») — "
            "тогда сумма посчитается автоматически, — либо очистите цену и задайте план "
            "общей суммой отдельной плановой позицией (кнопка «Добавить плановую» в "
            "панели) без указания количества."
        )
    raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail=detail)


async def _has_feo_action(current_user, db: AsyncSession, action_key: str) -> bool:
    """Есть ли у пользователя один из feo_budget.* action-ключей (эффективно, с учётом
    ролевой матрицы + персональных/субсидийных оверрайдов). superadmin — всегда True.
    Общий хелпер для мягкого усечения денежных полей в /leaves, /flat, /plan-tree —
    см. задачу владельца 2026-08-06 «остаток средств по каждой категории ФЭО»."""
    if current_user.role == "superadmin":
        return True
    from app.auth.permissions import _get_effective, _active_org
    effective = await _get_effective(current_user, db, _active_org(current_user))
    return action_key in effective


@router.get("/purchase-totals")
async def get_purchase_totals(
    subsidy_id: int = Query(...),
    db: AsyncSession = Depends(get_db),
    _=Depends(get_current_user),
):
    """Фактическая сумма per feo_category_id: только поставленное/оплаченное (по актам)."""
    from app.models.purchase import Purchase
    stmt = (
        select(
            Purchase.feo_category_id,
            func.coalesce(
                func.sum(func.coalesce(Purchase.final_total_amount, Purchase.planned_total_price)), 0
            ).label("total_planned"),
        )
        .where(Purchase.subsidy_id == subsidy_id)
        .where(Purchase.feo_category_id.isnot(None))
        .where(Purchase.status.in_(["delivered", "paid"]))
        .group_by(Purchase.feo_category_id)
    )
    rows = (await db.execute(stmt)).all()
    return {r.feo_category_id: float(r.total_planned) for r in rows}


@router.get("/planned-purchase-totals")
async def get_planned_purchase_totals(
    subsidy_id: int = Query(...),
    db: AsyncSession = Depends(get_db),
    _=Depends(get_current_user),
):
    """Плановая сумма и количество из заявок per feo_category_id: позиции закупок в статусах плана закупок и дальше.

    `total`/`qty` — ВСЕ позиции (для обратной совместимости, режим отображения «из заявок»).
    `total_linked`/`qty_linked` — подмножество позиций, привязанных к плановой позиции
    (`PurchaseItem.feo_planned_item_id IS NOT NULL`): они РАСХОДУЮТ ручной план листа
    (Ур.5), а не складываются с ним поверх — фронт вычитает linked из total, чтобы не
    задваивать план в дереве ФЭО.
    `total_over`/`qty_over` — подмножество НЕпривязанных позиций с `over_plan=true`:
    расход «сверх плана» элемента, прибавляется к плановой сумме безусловно (не входит
    в MAX(план, выбрано) — см. feoPlannedDisplayFor в SubsidiesView.vue и
    app.services.feo_plan.compute_feo_plan_tree). Партиция чистая: total = total_linked +
    total_over + «обычный» расход (total − total_linked − total_over), пересечения
    linked∩over_plan сознательно отнесены к total_linked (привязанные к Ур.5 позиции не
    участвуют в схеме over_plan этого эндпоинта).

    Дополнительно (сессия 2026-08-05, формула v2 — «заказ замещает план, пока не набрано
    количество»): `plan_manual`, `ordered_qty`, `ordered_sum`, `residual`, `forecast`,
    `forecast_over` — из app.services.feo_plan.compute_feo_plan_tree (единый источник,
    та же формула, что и в KPI «Запланировано»/«Свободно»). Присутствуют для КАЖДОЙ
    категории субсидии (и листа, и группы), даже без единой позиции закупки — в отличие
    от total/qty/... выше, которые есть только при наличии хотя бы одной позиции.
    """
    from app.models.purchase import Purchase
    from app.models.purchase_item import PurchaseItem
    from app.routers.purchase_budget import PLANNED_STATUSES
    from app.services.feo_plan import compute_feo_plan_tree
    from sqlalchemy import case, and_, not_

    cat_col = func.coalesce(PurchaseItem.feo_category_id, Purchase.feo_category_id)
    is_linked = PurchaseItem.feo_planned_item_id.isnot(None)
    is_over_unlinked = and_(PurchaseItem.over_plan.is_(True), not_(is_linked))
    stmt = (
        select(
            cat_col.label("cat_id"),
            func.coalesce(func.sum(PurchaseItem.total_price), 0).label("total"),
            func.coalesce(func.sum(PurchaseItem.quantity), 0).label("qty"),
            func.coalesce(func.sum(case((is_linked, PurchaseItem.total_price), else_=0)), 0).label("total_linked"),
            func.coalesce(func.sum(case((is_linked, PurchaseItem.quantity), else_=0)), 0).label("qty_linked"),
            func.coalesce(func.sum(case((is_over_unlinked, PurchaseItem.total_price), else_=0)), 0).label("total_over"),
            func.coalesce(func.sum(case((is_over_unlinked, PurchaseItem.quantity), else_=0)), 0).label("qty_over"),
        )
        .join(Purchase, PurchaseItem.purchase_id == Purchase.id)
        .where(Purchase.subsidy_id == subsidy_id)
        .where(Purchase.status.in_(list(PLANNED_STATUSES)))
        # Выравниваем с app.services.feo_plan.py — остановленные закупки не считаются
        # (решение владельца 2026-08-13).
        .where(Purchase.stopped_at.is_(None))
        .where(cat_col.isnot(None))
        .group_by(cat_col)
    )
    rows = (await db.execute(stmt)).all()
    result = {
        r.cat_id: {
            "total": float(r.total),
            "qty": float(r.qty),
            "total_linked": float(r.total_linked),
            "qty_linked": float(r.qty_linked),
            "total_over": float(r.total_over),
            "qty_over": float(r.qty_over),
        }
        for r in rows
    }

    # Аддитивно: формула v2 плановой суммы (см. docstring) — на КАЖДУЮ категорию
    # субсидии, а не только на те, где уже есть позиции закупок.
    tree = await compute_feo_plan_tree(db, [subsidy_id])
    for cat_id, node in tree.items():
        entry = result.setdefault(cat_id, {
            "total": 0.0, "qty": 0.0, "total_linked": 0.0, "qty_linked": 0.0,
            "total_over": 0.0, "qty_over": 0.0,
        })
        entry["plan_manual"] = node["plan_manual"]
        entry["ordered_qty"] = node["ordered_quantity"]
        entry["ordered_sum"] = node["ordered"]
        entry["residual"] = node["residual"]
        entry["forecast"] = node["forecast"]
        entry["forecast_over"] = node["forecast_over"]

    return result


@router.get("/plan-tree")
async def get_feo_plan_tree(
    subsidy_id: int = Query(...),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Числа по КАЖДОМУ узлу дерева ФЭО (и листу, и группе) — единый источник для
    фронта (сессия 2026-08-05, задача «формула только на бэкенде»: раньше
    SubsidiesView.vue пересчитывал «Плановую сумму»/«Плановое количество» сам
    (feoPlannedDisplayRaw/feoQtyDisplayRaw, MAX(план, выбрано) + сверх_плана —
    СТАРАЯ формула), а KPI «Запланировано» на дашборде/в списке субсидий считал
    _calculate_feo_planned_tree_bulk по НОВОЙ формуле compute_feo_plan_tree —
    два разных числа на одном экране. Теперь фронт только читает готовое отсюда.

    Обёртка над app.services.feo_plan.compute_feo_plan_tree — просто отдаёт её
    per-node словарь наружу, ничего не пересчитывая (см. её подробный docstring
    за формулой «заказ замещает план, когда количество набрано полностью»).

    Response: {cat_id: {plan_manual, ordered_qty, ordered_sum, over, over_quantity,
                         plan, display, residual, forecast, forecast_over,
                         consumed, consumed_quantity, qty_plan, display_quantity}}
    display — то самое число, которое обязано совпасть с KPI «Запланировано»
    (сумма display корневых узлов == _calculate_feo_planned_tree_bulk[subsidy_id]).
    display_quantity — аналог display, но для «Планового количества» узла.

    Плюс АДДИТИВНЫЙ ключ "unassigned" (строкой, не числовой id — чтобы не
    столкнуться с id категории): {amount, purchase_count, purchase_ids} — деньги
    закупок субсидии в PLANNED_STATUSES, у которых НИ сама закупка, НИ одна её
    позиция не привязаны к категории ФЭО (баг «Ведётся работа» видит деньги,
    дерево — нет, см. сессию 2026-08-05). Не участвует в дереве/ИТОГО — чисто
    справочная сумма для отдельной строки на фронте. Отдаётся всегда (нули, если
    таких закупок нет), чтобы фронт не городил доп. ветвление на отсутствие ключа.

    "budget" (задача владельца 2026-08-06, «остаток по каждой категории ФЭО в
    дереве выбора») — финансирование по ФЭО узла (то же значение, что уже
    участвует в формуле display/excess_amount выше), null если не задано.
    free = budget − display считает фронт (см. FeoTreeSelect nodeAmounts).

    "excess_culprit" (задача владельца, план zany-fluttering-mountain.md п.4,
    2026-08-10) — закупка-«виновник», из-за которой узел вышел за ФЭО: null,
    если превышения нет/оно согласовано, иначе {purchase_id, purchase_number,
    item_id, item_name, amount_before, amount_at_crossing, cumulative_after} —
    см. app.services.feo_plan.find_excess_culprit за методом определения.

    Право feo_budget.view_tree_amounts (см. backend/app/__init__.py) гейтит ВСЕ
    денежные поля узла (plan_manual, ordered_sum, over, plan, budget, display,
    residual, forecast, forecast_over, consumed, excess_amount) — без права они
    приходят null (структура ответа не меняется, количественные/булевы поля
    остаются). superadmin — всегда видит.

    Задача владельца, сессия 2026-08-12 («план ≠ факт», продолжение) — три
    дополнительных поля узла (см. app.services.feo_plan.compute_feo_plan_tree
    docstring за формулой):
      manual_plan_entered/excess_plan_over_manual/excess_plan_approved/
      excess_plan_pending/excess_plan_items — ТРЕТИЙ вид превышения: Σ ВСЕХ
      плановых позиций листа/подветки больше «ручного» плана (поля категории
      старого формата + Σ позиций без auto_created). excess_plan_items —
      [{name, amount}] автоматически заведённых позиций, из-за которых вылезли.
      excess_approval_amount/excess_approval_at/excess_approval_by_id/
      excess_approval_by_name — «превышение согласовано»: данные последнего
      approved-запроса PlanExcessApproval по категории (висит, даже если узел
      всё ещё формально в превышении — display включает его целиком).

    Задача владельца, план zany-fluttering-mountain.md (2026-08-13, переключатель
    способа расчёта плана): plan_source ('planned_items' | 'manual_sum') и
    manual_plan_amount — форма категории читает/пишет эту пару (POST/PUT
    /feo-categories). excess_plan_items теперь в новом виде — [{id, name, amount,
    purchases: [{id, registry_number, purchase_number, status, status_label,
    amount, stopped_at}]}], клик по виновнику открывает его закупку(и).
    excess_approval_plan_before/excess_approval_plan_after — «план был X → стал
    Y» для excess_kind='plan_over_manual' (см. app.services.feo_plan и
    app.routers.plan_excess), null для остальных видов превышения.
    """
    from app.services.feo_plan import compute_feo_plan_tree, find_excess_culprit
    from app.models.purchase import Purchase
    from app.models.purchase_item import PurchaseItem
    from app.routers.purchase_budget import PLANNED_STATUSES
    from sqlalchemy import exists, and_

    can_view_amounts = await _has_feo_action(current_user, db, "feo_budget.view_tree_amounts")

    tree = await compute_feo_plan_tree(db, [subsidy_id])
    result: dict = {
        cat_id: {
            "plan_manual": node["plan_manual"],
            "ordered_qty": node["ordered_qty"],
            "ordered_sum": node["ordered_sum"],
            "over": node["over"],
            "over_quantity": node["over_quantity"],
            "plan": node["plan"],
            "budget": node["budget"],
            "display": node["display"],
            "residual": node["residual"],
            "forecast": node["forecast"],
            "forecast_over": node["forecast_over"],
            "consumed": node["consumed"],
            "consumed_quantity": node["consumed_quantity"],
            "qty_plan": node["qty_plan"],
            "display_quantity": node["display_quantity"],
            # Согласование превышения плана над финансированием ФЭО (задача владельца
            # «блокировать пока не согласовано», 2026-08-05) — см. compute_feo_plan_tree
            # и app.routers.plan_excess. excess_amount>0 — есть превышение над node.budget;
            # excess_pending — есть незакрытый запрос на согласование; excess_approved —
            # превышение легализовано (display включает его полностью).
            "excess_amount": node["excess_amount"],
            "excess_pending": node["excess_pending"],
            "excess_approved": node["excess_approved"],
            # Задача владельца «план ≠ факт» (2026-08-06): факт со стадии «Ведётся
            # работа» + независимое превышение «факт дороже плана» — см.
            # compute_feo_plan_tree docstring (шаг C). Отсутствовали здесь —
            # эндпоинт отдавал только явный whitelist полей и терял их молча.
            "fact": node["fact"],
            "fact_quantity": node["fact_quantity"],
            "excess_over_feo": node["excess_over_feo"],
            "excess_fact_over_plan": node["excess_fact_over_plan"],
            "excess_fact_approved": node["excess_fact_approved"],
            "excess_fact_pending": node["excess_fact_pending"],
            # Задача владельца п.2 (2026-08-12, сессия «план ≠ факт», продолжение):
            # ТРЕТИЙ вид превышения — Σ плановых позиций против «ручного» плана,
            # см. compute_feo_plan_tree docstring.
            "manual_plan_entered": node["manual_plan_entered"],
            "excess_plan_over_manual": node["excess_plan_over_manual"],
            "excess_plan_approved": node["excess_plan_approved"],
            "excess_plan_pending": node["excess_plan_pending"],
            "excess_plan_items": node["excess_plan_items"],
            # Владелец, план zany-fluttering-mountain.md (2026-08-13): переключатель
            # способа расчёта плана узла — форма категории читает/пишет эту пару.
            "plan_source": node["plan_source"],
            "manual_plan_amount": node["manual_plan_amount"],
            # Задача владельца п.4 (2026-08-12): «превышение согласовано» — данные
            # последнего approved-запроса по категории (см. compute_feo_plan_tree).
            "excess_approval_amount": node["excess_approval_amount"],
            "excess_approval_at": node["excess_approval_at"],
            "excess_approval_by_id": node["excess_approval_by_id"],
            "excess_approval_by_name": node["excess_approval_by_name"],
            # Владелец, план zany-fluttering-mountain.md (2026-08-13): «план был X →
            # стал Y» на постоянной плашке — см. compute_feo_plan_tree.
            "excess_approval_plan_before": node["excess_approval_plan_before"],
            "excess_approval_plan_after": node["excess_approval_plan_after"],
            # Виновник превышения (задача владельца, план zany-fluttering-mountain.md
            # п.4, 2026-08-10) заполняется НИЖЕ, точечно только для узлов с
            # неснятым excess_amount — find_excess_culprit не бесплатна (доп. запрос),
            # гнать её на каждый узел дерева на каждый вызов (в т.ч. KPI-дашборд) не нужно.
            "excess_culprit": None,
        }
        for cat_id, node in tree.items()
    }

    # Виновник превышения плана над ФЭО — «должна отображаться данная закупка и
    # показать, что из-за неё всё превысило» (владелец). Считаем только для узлов
    # с непогашенным excess_amount (обычно единицы, не весь тысячестрочный тред).
    for cat_id, node in tree.items():
        if (node.get("excess_amount") or 0) > 0.005 and not node.get("excess_approved"):
            result[cat_id]["excess_culprit"] = await find_excess_culprit(db, cat_id, node.get("budget"))

    unassigned_stmt = (
        select(Purchase.id, Purchase.planned_total_price)
        .where(Purchase.subsidy_id == subsidy_id)
        .where(Purchase.status.in_(list(PLANNED_STATUSES)))
        .where(Purchase.feo_category_id.is_(None))
        .where(~exists().where(and_(
            PurchaseItem.purchase_id == Purchase.id,
            PurchaseItem.feo_category_id.isnot(None),
        )))
    )
    unassigned_rows = (await db.execute(unassigned_stmt)).all()
    result["unassigned"] = {
        "amount": sum(float(r.planned_total_price or 0) for r in unassigned_rows),
        "purchase_count": len(unassigned_rows),
        "purchase_ids": [r.id for r in unassigned_rows[:50]],
    }

    if not can_view_amounts:
        MONEY_FIELDS = (
            "plan_manual", "ordered_sum", "over", "plan", "budget", "display",
            "residual", "forecast", "forecast_over", "consumed", "excess_amount",
            "fact", "excess_over_feo", "excess_fact_over_plan", "excess_culprit",
            "manual_plan_entered", "excess_plan_over_manual", "excess_plan_items",
            "excess_approval_amount", "manual_plan_amount",
            "excess_approval_plan_before", "excess_approval_plan_after",
        )
        for cat_id, node in result.items():
            if cat_id == "unassigned":
                node["amount"] = None
                continue
            for f in MONEY_FIELDS:
                if f in node:
                    node[f] = None

    return result


@router.get("/planned-purchase-items")
async def get_planned_purchase_items(
    subsidy_id: int = Query(...),
    db: AsyncSession = Depends(get_db),
    _=Depends(get_current_user),
):
    """Позиции закупок «из заявок» per feo_category_id (статусы плана закупок и дальше).

    Задача владельца «план ≠ факт» (сессия 2026-08-06, шаг 5): панель «Позиции: план
    vs факт» (SubsidiesView.vue, таблица источников группы «из заявок») обязана
    показывать ЗАМОРОЖЕННЫЙ снимок ТЗ (planned_quantity/planned_unit_price/planned_total,
    см. Шаг 1/2 «план ≠ факт») и реальный факт (fact_amount/fact_confirmed) — раньше
    отдавались только «текущие» quantity/unit_price/total_price (которые задним числом
    подменяются ценой по итогам закупки) и фактических данных не было вовсе (фронт
    печатал жёсткую заглушку «ещё не поставлено»). Формула факта — та же единая
    purchase_item_fact_amount, что и в /feo-planned-items/comparison (Table A), чтобы
    два соседних UI-блока не показывали разные числа по одной и той же позиции.
    """
    from app.models.purchase import Purchase
    from app.models.purchase_item import PurchaseItem
    from app.models.product import Product
    from app.models.contract import Contract
    from app.routers.purchase_budget import PLANNED_STATUSES
    from app.services.feo_plan import (
        purchase_item_fact_amount,
        FACT_CONFIRMED_STATUSES,
        _contract_item_totals,
        _purchase_item_totals,
    )

    cat_col = func.coalesce(PurchaseItem.feo_category_id, Purchase.feo_category_id)
    stmt = (
        select(
            cat_col.label("cat_id"),
            PurchaseItem.id,
            PurchaseItem.item_name,
            PurchaseItem.quantity,
            PurchaseItem.unit,
            PurchaseItem.unit_price,
            PurchaseItem.total_price,
            PurchaseItem.planned_quantity,
            PurchaseItem.planned_unit_price,
            PurchaseItem.planned_total,
            PurchaseItem.final_total,
            PurchaseItem.purchase_id,
            PurchaseItem.product_id,
            PurchaseItem.feo_planned_item_id,
            PurchaseItem.wish_item_id,
            Purchase.purchase_number,
            Purchase.registry_number,
            Purchase.status.label("purchase_status"),
            Purchase.wish_id,
            Purchase.contract_id,
            Purchase.purchase_contract_type,
            Purchase.contract_price,
            Contract.contract_type.label("contract_type"),
            Contract.status.label("contract_status"),
            Contract.number.label("contract_number"),
            Product.category.label("product_category"),
            Product.product_type.label("product_type"),
            Product.photo_data.isnot(None).label("product_has_photo"),
            Product.photo_url,
            Product.photo_link,
        )
        .join(Purchase, PurchaseItem.purchase_id == Purchase.id)
        .outerjoin(Product, PurchaseItem.product_id == Product.id)
        .outerjoin(Contract, Purchase.contract_id == Contract.id)
        .where(Purchase.subsidy_id == subsidy_id)
        .where(Purchase.status.in_(list(PLANNED_STATUSES)))
        # Выравниваем с app.services.feo_plan.py — остановленные закупки не считаются
        # (решение владельца 2026-08-13).
        .where(Purchase.stopped_at.is_(None))
        .where(cat_col.isnot(None))
        .order_by(cat_col, PurchaseItem.item_name)
    )
    rows = (await db.execute(stmt)).all()

    # Факт: та же формула, что и в /feo-planned-items/comparison — точное сопоставление
    # через ContractItem.source_item_id, иначе пропорция от purchases.contract_price.
    _item_ids = [r.id for r in rows]
    _purchase_ids = {r.purchase_id for r in rows}
    _ci_totals = await _contract_item_totals(db, _item_ids)
    _purchase_totals = await _purchase_item_totals(db, _purchase_ids)
    # Кол-во/цена факта (не только сумма): берём напрямую из ТОЧНО сопоставленной строки
    # договора (source_item_id), если она есть — та же строка, что дала _ci_totals.
    from app.models.contract_item import ContractItem as _ContractItem
    _ci_qty_price: dict[int, tuple] = {}
    if _item_ids:
        _ci_rows = (await db.execute(
            select(_ContractItem.source_item_id, _ContractItem.quantity, _ContractItem.unit_price)
            .where(_ContractItem.source_item_id.in_(_item_ids))
        )).all()
        for _row in _ci_rows:
            if _row.source_item_id not in _ci_qty_price:
                _ci_qty_price[_row.source_item_id] = (_row.quantity, _row.unit_price)
    # Легковесные объекты-обёртки под purchase_item_fact_amount (принимает ORM-подобные
    # PurchaseItem/Purchase — здесь только Row, поэтому оборачиваем нужные атрибуты).
    class _PiShim:
        __slots__ = ("total_price", "final_total")
    class _PShim:
        __slots__ = ("status", "contract_price", "acceptance_doc_amount")

    result: dict[int, list] = {}
    for r in rows:
        if r.product_id is not None and r.product_has_photo:
            product_photo = f"/api/products/{r.product_id}/photo"
        elif r.product_id is not None:
            product_photo = r.photo_url or r.photo_link or None
        else:
            product_photo = None

        items_count, items_sum = _purchase_totals.get(r.purchase_id, (1, Decimal(str(r.total_price or 0))))
        item_total = Decimal(str(r.total_price or 0))
        if items_count > 1 and items_sum > 0:
            ratio = item_total / items_sum
        elif items_count > 1:
            ratio = Decimal(1) / Decimal(items_count)
        else:
            ratio = Decimal(1)

        pi_shim = _PiShim()
        pi_shim.total_price = r.total_price
        pi_shim.final_total = r.final_total
        p_shim = _PShim()
        p_shim.status = r.purchase_status
        p_shim.contract_price = r.contract_price
        p_shim.acceptance_doc_amount = None  # не выбирается этим запросом — не нужен для FACT_PRICED_STATUSES
        fact_amount, fact_allocated = purchase_item_fact_amount(
            pi_shim, p_shim, ratio, items_count, contract_item_total=_ci_totals.get(r.id)
        )
        fact_confirmed = r.purchase_status in FACT_CONFIRMED_STATUSES

        # Кол-во/цена факта: из точно сопоставленной строки договора, если есть;
        # иначе (пропорциональное распределение или факта ещё нет) — кол-во берём
        # как у ТЗ (обычно не меняется договором), цену — делением суммы на кол-во.
        _ci_qty, _ci_price = _ci_qty_price.get(r.id, (None, None))
        if _ci_qty is not None:
            fact_quantity = float(_ci_qty)
        elif fact_amount is not None:
            fact_quantity = float(r.quantity or 0)
        else:
            fact_quantity = None
        if _ci_price is not None:
            fact_unit_price = float(_ci_price)
        elif fact_amount is not None and fact_quantity:
            fact_unit_price = float(fact_amount) / fact_quantity
        else:
            fact_unit_price = None

        result.setdefault(r.cat_id, []).append({
            "id": r.id,
            "item_name": r.item_name,
            "quantity": float(r.quantity or 0),
            "unit": r.unit,
            "unit_price": float(r.unit_price or 0),
            "total_price": float(r.total_price or 0),
            # Снимок плана (Шаг 1/2 «план ≠ факт») — заморожен с момента объявления
            # закупки; фолбэк на текущие значения для старых позиций без снимка.
            "planned_quantity": float(r.planned_quantity) if r.planned_quantity is not None else float(r.quantity or 0),
            "planned_unit_price": float(r.planned_unit_price) if r.planned_unit_price is not None else float(r.unit_price or 0),
            "planned_total": float(r.planned_total) if r.planned_total is not None else float(r.total_price or 0),
            "fact_amount": float(fact_amount) if fact_amount is not None else None,
            "fact_quantity": fact_quantity,
            "fact_unit_price": fact_unit_price,
            "fact_confirmed": fact_confirmed,
            "fact_allocated": fact_allocated,
            "purchase_id": r.purchase_id,
            "feo_planned_item_id": r.feo_planned_item_id,
            "wish_item_id": r.wish_item_id,
            "purchase_number": r.purchase_number,
            "registry_number": r.registry_number,
            "purchase_status": r.purchase_status,
            "wish_id": r.wish_id,
            "contract_id": r.contract_id,
            "purchase_contract_type": r.purchase_contract_type,
            "contract_type": r.contract_type,
            "contract_status": r.contract_status,
            "contract_number": r.contract_number,
            "category": r.product_category or "Без категории",
            "product_type": r.product_type or "Без вида",
            "product_photo": product_photo,
        })
    return result


@router.get("/leaves")
async def get_feo_leaves(
    subsidy_id: int = Query(...),
    exclude_purchase_id: Optional[int] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Returns leaf FeoCategory nodes (без детей) with aggregated used_amount via feo_category_id.

    Response: [{id, name, parent_id, level, budget, used_amount, residual, path}]
    where path = "Direction › Subcategory › LeafName".

    Право feo_budget.view_leaf (см. backend/app/__init__.py, дефолт — все роли,
    включая employee) гейтит денежные поля (budget, contracted_used, planned_used,
    uncontracted_remaining, spendable_remaining, used_amount, residual) — без права
    приходят 0, структура ответа не меняется (закрытие дыры: раньше эндпоинт вообще
    не проверял права, см. задачу владельца 2026-08-06).
    """
    can_view_leaf = await _has_feo_action(current_user, db, "feo_budget.view_leaf")
    from sqlalchemy import select, func as sqlfunc, case
    from app.models.purchase_item import PurchaseItem
    from app.models.purchase import Purchase as _Purchase
    from app.routers.purchase_budget import CONTRACTED_STATUSES, PLANNED_STATUSES

    # Все FeoCategory для subsidy
    cats_q = select(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id).order_by(FeoCategory.sort_order.nulls_last(), FeoCategory.id)
    all_cats = (await db.execute(cats_q)).scalars().all()
    if not all_cats:
        return []

    cat_by_id = {c.id: c for c in all_cats}
    children_count: dict[int, int] = {}
    for c in all_cats:
        if c.parent_id is not None:
            children_count[c.parent_id] = children_count.get(c.parent_id, 0) + 1

    # Leaves = категории у которых нет детей в feo_categories
    leaves = [c for c in all_cats if children_count.get(c.id, 0) == 0]
    if not leaves:
        return []

    leaf_ids = [c.id for c in leaves]

    # Aggregate contracted_used and planned_used per feo_category_id via conditional sums
    used_q = (
        select(
            PurchaseItem.feo_category_id,
            sqlfunc.coalesce(
                sqlfunc.sum(case((_Purchase.status.in_(list(CONTRACTED_STATUSES)), PurchaseItem.total_price), else_=0)),
                0,
            ).label("contracted_used"),
            sqlfunc.coalesce(
                sqlfunc.sum(case((_Purchase.status.in_(list(PLANNED_STATUSES)), PurchaseItem.total_price), else_=0)),
                0,
            ).label("planned_used"),
        )
        .join(_Purchase, PurchaseItem.purchase_id == _Purchase.id)
        .where(PurchaseItem.feo_category_id.in_(leaf_ids))
    )
    if exclude_purchase_id is not None:
        used_q = used_q.where(PurchaseItem.purchase_id != exclude_purchase_id)
    used_q = used_q.group_by(PurchaseItem.feo_category_id)
    # leaf_used_map: {feo_category_id: (contracted_used, planned_used)}
    leaf_used_map: dict[int, tuple[float, float]] = {}
    for r in (await db.execute(used_q)).all():
        leaf_used_map[r.feo_category_id] = (float(r.contracted_used), float(r.planned_used))

    # Build path "Direction › Subcategory › Leaf"
    def build_path(cat) -> str:
        names = [cat.name]
        cur = cat
        while cur.parent_id is not None and cur.parent_id in cat_by_id:
            cur = cat_by_id[cur.parent_id]
            names.append(cur.name)
        return " \u203a ".join(reversed(names))

    result = []
    for c in leaves:
        budget = float(c.budget or 0)
        contracted_used, planned_used = leaf_used_map.get(c.id, (0.0, 0.0))
        uncontracted_remaining = budget - contracted_used
        spendable_remaining = budget - planned_used
        if not can_view_leaf:
            budget = contracted_used = planned_used = uncontracted_remaining = spendable_remaining = 0.0
        result.append({
            "id": c.id,
            "name": c.name,
            "parent_id": c.parent_id,
            "level": c.level,
            "budget": budget,
            # New metrics
            "contracted_used": contracted_used,
            "planned_used": planned_used,
            "uncontracted_remaining": uncontracted_remaining,
            "spendable_remaining": spendable_remaining,
            # Legacy fields (kept for backward compat): used_amount = planned_used, residual = spendable_remaining
            "used_amount": planned_used,
            "residual": spendable_remaining,
            "path": build_path(c),
        })

    # Сортировка по path для удобства autocomplete
    result.sort(key=lambda x: x["path"])
    return result


@router.get("/plan-positions")
async def get_plan_positions(
    subsidy_id: int = Query(...),
    exclude_purchase_id: Optional[int] = Query(None),
    exclude_wish_id: Optional[int] = Query(None),
    db: AsyncSession = Depends(get_db),
    _=Depends(get_current_user),
):
    """Единый источник «плановых позиций» субсидии.

    Плановая позиция — конечный элемент дерева ФЭО (FeoCategory без детей) с
    заполненными planned_quantity/planned_amount (произведение > 0). Плюс —
    отдельными записями kind='planned_item' — существующие FeoPlannedItem
    (Ур.5, необязательная более глубокая детализация внутри элемента).

    kind:
      'plan_position' — budget IS NULL AND feo_amount IS NULL (план без статьи ФЭО,
                         «мы сами запланировали купить именно это»)
      'feo_article'    — budget и/или feo_amount заполнены (план = статья ФЭО)
      'planned_item'   — FeoPlannedItem внутри элемента

    Response: [{id, name, path, category_id, ancestor_ids, kind, planned_quantity, unit,
                planned_amount, unit_price, consumed, consumed_quantity,
                residual, residual_quantity}]
    planned_amount — ИТОГОВАЯ сумма плана (quantity × unit_price); unit_price —
    цена за единицу (= FeoCategory.planned_amount / FeoPlannedItem.amount per unit).
    ancestor_ids — id всех предков category_id (от корня до непосредственного родителя),
    см. app.services.feo_plan.build_ancestor_ids. Фронт (FeoPlannedItemsSelect.vue)
    матчит категорию, выбранную в дереве заявки/закупки, по `category_id === selected
    || ancestor_ids.includes(selected)` — НАПРЯМУЮ по этому полю, а не обходом
    props.nodes/useFeoLeaves (тот массив обрезан filterFundedNodes: конечные категории
    без собственного budget вырезаются из дерева, даже если у них есть
    planned_quantity/planned_amount — баг «В этой категории нет плановых позиций»,
    сессия 2026-08-05).

    Дополнительно на строках kind='plan_position'/'feo_article' (сессия 2026-08-05,
    формула v2 — «заказ замещает план, пока не набрано количество», см.
    app.services.feo_plan.compute_feo_plan_tree): `ordered_qty`, `ordered_sum`,
    `plan_manual`, `ordered_residual`, `forecast`, `forecast_over`. Существующие
    `consumed`/`consumed_quantity`/`residual`/`residual_quantity` НЕ трогаем (старая
    семантика — используются FeoPlannedItemsSelect.vue/PurchaseItemsEditor.vue для
    предотвращения повторного выбора уже занятого плана); `ordered_residual` — НОВЫЙ
    остаток по формуле v2 (план − ordered_sum), под другим именем, чтобы не столкнуть
    с уже существующим `residual`.

    kind='planned_item' — ВАЖНО (владелец, сессия 2026-08-18): FeoPlannedItem может
    жить не только под конечной категорией (листом), но и на НАПРАВЛЕНИИ/подкатегории
    (FeoCategory с детьми) — так заводят позицию, если решили не спускаться до листа.
    Раньше запрос ниже брал FeoPlannedItem только по leaf_ids — такая позиция участвовала
    в деньгах (compute_feo_plan_tree считает собственные FeoPlannedItem групп, см. его
    докстринг — «Бинт марлевый Навтекс Life» там боевой пример именно такого случая),
    но пропадала из списка выбора: UI показывал «Выбрать плановую позицию», будто
    привязки нет, хотя закупка уже была привязана к существующей записи. Обязаны
    показывать её всегда — иначе пользователь не может понять, куда делась позиция,
    на которую сама система ссылается. path/ancestor_ids для таких строк строятся тем
    же build_category_path/build_ancestor_ids, что и для листьев — они просто идут
    вверх по parent_id, им всё равно, лист cat или группа.

    exclude_purchase_id/exclude_wish_id (сессия 2026-08-17): редактирование
    существующей закупки/заявки — её собственные строки не должны выглядеть
    занявшими план (иначе позиция сравнивается сама с собой: остаток 0, и UI
    повторно вычитает ту же сумму — баг «план 54 318 · выбрано 54 318 · не
    хватает 54 318»). Ровно то же уже делает GET /api/feo-planned-items/residuals.
    Поля `ordered_qty`/`ordered_sum`/`plan_manual`/`ordered_residual`/`forecast`/
    `forecast_over` приходят из compute_feo_plan_tree и исключение НЕ учитывают —
    это осознанно, в UI-подписи «не хватает» они не участвуют.
    """
    from app.models.feo_planned_item import FeoPlannedItem
    from app.services.feo_plan import (
        plan_consumption_by_category, planned_item_consumption, build_category_path,
        build_ancestor_ids, compute_feo_plan_tree,
    )

    all_cats = (await db.execute(
        select(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id)
    )).scalars().all()
    if not all_cats:
        return []

    cat_by_id = {c.id: c for c in all_cats}
    children_count: dict[int, int] = {}
    for c in all_cats:
        if c.parent_id is not None:
            children_count[c.parent_id] = children_count.get(c.parent_id, 0) + 1
    leaves = [c for c in all_cats if children_count.get(c.id, 0) == 0]

    consumption = await plan_consumption_by_category(
        db, [subsidy_id], exclude_purchase_id=exclude_purchase_id, exclude_wish_id=exclude_wish_id
    )
    tree = await compute_feo_plan_tree(db, [subsidy_id])

    result = []
    for c in leaves:
        qty = float(c.planned_quantity) if c.planned_quantity is not None else 0.0
        unit_price = float(c.planned_amount) if c.planned_amount is not None else 0.0
        planned_total = qty * unit_price
        if planned_total <= 0:
            continue
        cons = consumption.get(c.id, {"consumed": 0.0, "consumed_quantity": 0.0})
        consumed = cons["consumed"]
        consumed_qty = cons["consumed_quantity"]
        kind = "plan_position" if (c.budget is None and c.feo_amount is None) else "feo_article"
        node = tree.get(c.id, {})
        result.append({
            "id": c.id,
            "name": c.name,
            "path": build_category_path(c, cat_by_id),
            "category_id": c.id,
            # Предки категории (от корня до непосредственного родителя) — фронт
            # сопоставляет позицию с выбранной в дереве категорией через
            # category_id == selected ИЛИ selected in ancestor_ids, без обхода
            # обрезанного props.nodes (см. build_ancestor_ids).
            "ancestor_ids": build_ancestor_ids(c, cat_by_id),
            "kind": kind,
            "planned_quantity": qty,
            "unit": c.unit,
            "planned_amount": planned_total,
            "unit_price": unit_price,
            "consumed": consumed,
            "consumed_quantity": consumed_qty,
            "residual": planned_total - consumed,
            "residual_quantity": qty - consumed_qty,
            "ordered_qty": node.get("ordered_quantity", 0.0),
            "ordered_sum": node.get("ordered", 0.0),
            "plan_manual": node.get("plan_manual", planned_total),
            "ordered_residual": node.get("residual", planned_total),
            "forecast": node.get("forecast", planned_total),
            "forecast_over": node.get("forecast_over", 0.0),
        })

    # + FeoPlannedItem (Ур.5) — детализация внутри элементов. Владелец 2026-08-18:
    # «"Бинт марлевый Навтекс Life" написано, что находится на уровне "Окружные" —
    # но на уровне "Окружные" его нет. Если пишет, что он там, должен же он как-то
    # отображаться?» — раньше запрос брал FeoPlannedItem ТОЛЬКО под leaves (конечными
    # категориями), но FeoPlannedItem может жить и на НАПРАВЛЕНИИ/группе (категория с
    # детьми) — владельцы заводят такие позиции напрямую на узле верхнего уровня, не
    # спускаясь до листа. compute_feo_plan_tree уже умеет считать деньги по таким
    # группам (см. его докстринг — «Бинт марлевый» там же боевой пример), т.е. в
    # расчётах позиция участвовала всегда; баг был только в том, что её не было видно
    # в списке выбора — интерфейс показывал «Выбрать плановую позицию», будто привязки
    # нет вовсе. Поэтому здесь фильтр по all_cats (вся субсидия), а не по leaf_ids.
    if all_cats:
        fpi_rows = (await db.execute(
            select(FeoPlannedItem)
            .where(FeoPlannedItem.feo_category_id.in_([c.id for c in all_cats]))
            .where(FeoPlannedItem.is_active == True)
            .order_by(FeoPlannedItem.id)
        )).scalars().all()
        if fpi_rows:
            fpi_ids = [it.id for it in fpi_rows]
            fpi_cons = await planned_item_consumption(db, fpi_ids, exclude_purchase_id, exclude_wish_id)
            for it in fpi_rows:
                cat = cat_by_id.get(it.feo_category_id)
                planned_total = float(it.amount or 0)
                qty = float(it.quantity or 0)
                c_cons = fpi_cons.get(it.id, {"used": 0.0, "used_qty": 0.0})
                consumed = c_cons["used"]
                consumed_qty = c_cons["used_qty"]
                result.append({
                    "id": it.id,
                    "name": it.name,
                    "path": build_category_path(cat, cat_by_id) if cat else "",
                    "category_id": it.feo_category_id,
                    "ancestor_ids": build_ancestor_ids(cat, cat_by_id) if cat else [],
                    "kind": "planned_item",
                    "planned_quantity": qty,
                    "unit": it.unit,
                    "planned_amount": planned_total,
                    "unit_price": (planned_total / qty) if qty else None,
                    "consumed": consumed,
                    "consumed_quantity": consumed_qty,
                    "residual": planned_total - consumed,
                    "residual_quantity": qty - consumed_qty,
                })

    result.sort(key=lambda x: x["path"])
    return result


@router.get("/budget-residuals")
async def feo_budget_residuals(
    subsidy_id: int = Query(...),
    category_ids: str = Query("", description="comma-separated leaf FeoCategory ids"),
    exclude_purchase_id: Optional[int] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Остатки бюджета по ФЭО для формы закупки.

    Для пользователей с feo_budget.view_all_levels возвращает направления (ур.1)
    и ancestors предков в каждом листе. Без этого права: directions=[],
    ancestors=[] — только сам листовой остаток.

    budget листа = FeoCategory.budget; used = SUM(PurchaseItem.total_price) по feo_category_id листа.
    Узел выше: budget = собственный .budget если задан, иначе сумма budget детей;
    used = сумма used всех листьев-потомков. residual = budget - used.
    Ответ: {directions:[{id,name,level,path,budget,used,residual}],
            leaves:[{id,name,level,path,budget,used,residual,ancestors:[...]}]}
    (ancestors — сверху вниз: ур.1 первым).
    """
    from app.auth.permissions import _get_effective, _active_org
    effective = await _get_effective(current_user, db, _active_org(current_user))
    can_view_all_levels = (current_user.role == "superadmin") or ("feo_budget.view_all_levels" in effective)
    from sqlalchemy import select as _sel, func as _f, case as _case
    from app.models.purchase_item import PurchaseItem
    from app.models.purchase import Purchase as _Purchase
    from app.routers.purchase_budget import CONTRACTED_STATUSES, PLANNED_STATUSES

    ids = [int(x) for x in category_ids.split(",") if x.strip().isdigit()]
    all_cats = (await db.execute(
        _sel(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id)
    )).scalars().all()
    if not all_cats:
        return {"directions": [], "leaves": []}
    cat_by_id = {c.id: c for c in all_cats}
    children: dict[int, list] = {}
    for c in all_cats:
        if c.parent_id is not None:
            children.setdefault(c.parent_id, []).append(c.id)
    leaf_ids_all = [c.id for c in all_cats if not children.get(c.id)]

    used_q = (
        _sel(
            PurchaseItem.feo_category_id,
            _f.coalesce(
                _f.sum(_case((_Purchase.status.in_(list(CONTRACTED_STATUSES)), PurchaseItem.total_price), else_=0)),
                0,
            ).label("contracted_used"),
            _f.coalesce(
                _f.sum(_case((_Purchase.status.in_(list(PLANNED_STATUSES)), PurchaseItem.total_price), else_=0)),
                0,
            ).label("planned_used"),
        )
        .join(_Purchase, PurchaseItem.purchase_id == _Purchase.id)
        .where(PurchaseItem.feo_category_id.in_(leaf_ids_all))
    )
    if exclude_purchase_id is not None:
        used_q = used_q.where(PurchaseItem.purchase_id != exclude_purchase_id)
    used_q = used_q.group_by(PurchaseItem.feo_category_id)
    # leaf_used: {feo_category_id: (contracted_used, planned_used)}
    leaf_used: dict[int, tuple[float, float]] = {}
    for r in (await db.execute(used_q)).all():
        leaf_used[r.feo_category_id] = (float(r.contracted_used), float(r.planned_used))

    def descendant_leaves(cid):
        ch = children.get(cid)
        if not ch:
            return [cid]
        out = []
        for x in ch:
            out.extend(descendant_leaves(x))
        return out

    def calc_budget(cid):
        c = cat_by_id[cid]
        ch = children.get(cid)
        if not ch:
            return float(c.budget) if c.budget is not None else 0.0
        if c.budget is not None:
            return float(c.budget)
        return sum(calc_budget(x) for x in ch)

    def contracted_of(cid):
        return sum(leaf_used.get(l, (0.0, 0.0))[0] for l in descendant_leaves(cid))

    def planned_of(cid):
        return sum(leaf_used.get(l, (0.0, 0.0))[1] for l in descendant_leaves(cid))

    def node_info(cid):
        c = cat_by_id[cid]
        b = calc_budget(cid)
        cu = contracted_of(cid)
        pu = planned_of(cid)
        return {
            "id": cid, "name": c.name, "level": c.level,
            "budget": b,
            "contracted_used": cu,
            "planned_used": pu,
            "uncontracted_remaining": b - cu,
            "spendable_remaining": b - pu,
            # Legacy fields (backward compat): used = planned_used, residual = spendable_remaining
            "used": pu,
            "residual": b - pu,
        }

    def path_of(cid):
        names = []
        cur = cat_by_id.get(cid)
        while cur is not None:
            names.append(cur.name)
            cur = cat_by_id.get(cur.parent_id) if cur.parent_id else None
        return " \u203a ".join(reversed(names))

    # Направления (ур.1) субсидии — только для пользователей с view_all_levels.
    directions = []
    if can_view_all_levels:
        # Порядок как во вкладке субсидии: sort_order, затем id (не алфавит)
        roots = [c for c in all_cats if c.level == 1 or c.parent_id is None]
        roots.sort(key=lambda c: (c.sort_order is None, c.sort_order or 0, c.id))
        for c in roots:
            d = node_info(c.id)
            d["path"] = path_of(c.id)
            directions.append(d)

    leaves = []
    for lid in ids:
        if lid not in cat_by_id:
            continue
        leaf = node_info(lid)
        leaf["path"] = path_of(lid)
        # ancestors только для пользователей с view_all_levels
        if can_view_all_levels:
            ancestors = []
            cur = cat_by_id[lid]
            while cur.parent_id and cur.parent_id in cat_by_id:
                cur = cat_by_id[cur.parent_id]
                ancestors.append(node_info(cur.id))
            leaf["ancestors"] = list(reversed(ancestors))
        else:
            leaf["ancestors"] = []
        leaves.append(leaf)
    return {"directions": directions, "leaves": leaves}


@router.get("/", response_model=List[FeoCategoryOut])
async def list_categories(
    parent_id: Optional[int] = Query(None),
    level: Optional[int] = Query(None),
    subsidy_id: Optional[int] = Query(None),
    appendix: Optional[str] = Query(None),
    is_active: Optional[bool] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    q = select(FeoCategory)
    vis = await get_visible_subsidy_ids(current_user, db, "feo_categories")
    if vis is not None:
        # ФЭО-категории выбираются внутри форм заявки/закупки, поэтому пикер
        # доступен всем, кто видит субсидию по вкладкам wishes или purchases
        # (напр. роль Менеджер без админской вкладки feo_categories видит субсидию
        # только через «Заявки») — иначе список категорий пуст.
        vis = vis | await get_visible_subsidy_ids(current_user, db, "purchases")
        vis = vis | await get_visible_subsidy_ids(current_user, db, "wishes")
        q = q.where(FeoCategory.subsidy_id.in_(vis))
    if parent_id is not None:
        q = q.where(FeoCategory.parent_id == parent_id)
    if level is not None:
        q = q.where(FeoCategory.level == level)
    if subsidy_id is not None:
        q = q.where(FeoCategory.subsidy_id == subsidy_id)
    if appendix is not None:
        q = q.where(FeoCategory.appendix == appendix)
    if is_active is not None:
        q = q.where(FeoCategory.is_active == is_active)
    result = await db.execute(q.order_by(FeoCategory.sort_order.nulls_last(), FeoCategory.id))
    return result.scalars().all()


@router.get("/flat")
async def get_feo_flat(
    subsidy_id: int = Query(...),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Returns all FeoCategory nodes for a subsidy as a flat list with is_leaf flag.

    Response: [{id, name, parent_id, level, is_leaf, budget, planned_quantity, planned_amount,
                has_budget, has_plan}]
    is_leaf = True if the node has no children within the same subsidy.
    budget = собственная (ручная) сумма финансирования узла, без расчёта по детям.
    planned_quantity/planned_amount — плановые показатели (CRM-план), заполняются НЕЗАВИСИМО
    от budget: конечная категория может иметь план (использована в плане закупок) без
    собственного budget.

    has_budget/has_plan — СТРУКТУРНЫЕ булевы признаки (НЕ денежные величины, НЕ гейтятся
    правом): has_budget истинен, если у узла задан собственный budget > 0; has_plan истинен,
    если planned_quantity × planned_amount > 0, ЛИБО (фолбэк — план переехал в записи
    внутри категории) у узла есть активная FeoPlannedItem с amount > 0. Без этого фолбэка
    мигрированные категории-листья (planned_quantity/planned_amount категории пусты, план
    введён плановыми позициями) пропадают из дерева выбора категории ФЭО в заявке/закупке
    — к ним нельзя привязать закупку, хотя план реально есть. Фронт (useFeoLeaves.filterFundedNodes)
    считает узел значимым по has_budget ИЛИ has_plan (с фолбэком на числовые budget/
    planned_quantity/planned_amount, если бэкенд старый) — иначе такие категории
    вырезались из дерева выбора, хотя реально существуют и используются (баг «категория
    ФЭО была удалена из справочника», сессия 2026-08-05). Важно: has_budget/has_plan
    обязаны приходить ВСЕГДА, независимо от права feo_budget.view_leaf — иначе у роли без
    этого права дерево выбора категорий ФЭО становится полностью пустым (найдено при
    ревью гейта, сессия 2026-08-06).
    Sorted by level, then sort_order, then id.

    Право feo_budget.view_leaf (дефолт — все роли, включая employee) гейтит денежные
    поля budget/planned_amount — без права оба приходят null (уже допустимое по типу
    состояние, см. ниже); planned_quantity — количество, не деньги, не гейтится.
    Закрытие дыры: раньше эндпоинт вообще не проверял права (задача владельца 2026-08-06).
    """
    can_view_leaf = await _has_feo_action(current_user, db, "feo_budget.view_leaf")
    cats_q = (
        select(FeoCategory)
        .where(FeoCategory.subsidy_id == subsidy_id)
        .order_by(FeoCategory.level, FeoCategory.sort_order.nulls_last(), FeoCategory.id)
    )
    all_cats = (await db.execute(cats_q)).scalars().all()
    if not all_cats:
        return []

    # Determine which nodes have children
    has_children: set[int] = set()
    for c in all_cats:
        if c.parent_id is not None:
            has_children.add(c.parent_id)

    # Фолбэк has_plan: план переехал в записи внутри категории (FeoPlannedItem) —
    # у мигрированных категорий-листьев planned_quantity/planned_amount пусты, план
    # введён активными плановыми позициями. Без этого набора has_plan у таких узлов
    # уходит в false, и они пропадают из дерева выбора категории ФЭО на фронте
    # (useFeoLeaves.filterFundedNodes) — к ним нельзя привязать закупку, хотя план
    # реально есть (на боевых данных таких категорий около сотни). Один batch-запрос
    # на все категории ответа — без N+1 в цикле.
    from app.models.feo_planned_item import FeoPlannedItem
    cat_ids = [c.id for c in all_cats]
    cats_with_plan_items: set[int] = set()
    if cat_ids:
        plan_rows = (await db.execute(
            select(FeoPlannedItem.feo_category_id)
            .where(FeoPlannedItem.feo_category_id.in_(cat_ids))
            .where(FeoPlannedItem.is_active.is_(True))
            .where(FeoPlannedItem.amount > 0)
            .group_by(FeoPlannedItem.feo_category_id)
        )).scalars().all()
        cats_with_plan_items = set(plan_rows)

    return [
        {
            "id": c.id,
            "name": c.name,
            "parent_id": c.parent_id,
            "level": c.level,
            "is_leaf": c.id not in has_children,
            "description": c.description,
            "budget": (float(c.budget) if c.budget is not None else None) if can_view_leaf else None,
            "planned_quantity": float(c.planned_quantity) if c.planned_quantity is not None else None,
            "planned_amount": (float(c.planned_amount) if c.planned_amount is not None else None) if can_view_leaf else None,
            # Структурные признаки — НЕ гейтятся правом, см. docstring выше.
            "has_budget": bool(c.budget is not None and float(c.budget) > 0),
            "has_plan": bool(
                (
                    c.planned_quantity is not None
                    and c.planned_amount is not None
                    and float(c.planned_quantity) * float(c.planned_amount) > 0
                )
                or c.id in cats_with_plan_items
            ),
        }
        for c in all_cats
    ]


@router.get("/tree")
async def category_tree(
    subsidy_id: Optional[int] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    from app.models.subsidy import Subsidy
    q = select(FeoCategory)
    if subsidy_id is not None:
        q = q.where(FeoCategory.subsidy_id == subsidy_id)
    org_ids = get_org_filter(current_user)
    if org_ids is not None:
        q = q.join(Subsidy, FeoCategory.subsidy_id == Subsidy.id).where(Subsidy.org_id.in_(org_ids))
    result = await db.execute(q.order_by(FeoCategory.level, FeoCategory.sort_order.nulls_last(), FeoCategory.id))
    all_cats = result.scalars().all()
    by_id = {c.id: {"id": c.id, "parent_id": c.parent_id, "subsidy_id": c.subsidy_id,
                    "level": c.level, "name": c.name, "code": c.code,
                    "appendix": c.appendix, "is_active": c.is_active,
                    "description": c.description,
                    "budget": float(c.budget) if c.budget is not None else None,
                    "feo_quantity": float(c.feo_quantity) if c.feo_quantity is not None else None,
                    "feo_unit": c.feo_unit,
                    "feo_amount": float(c.feo_amount) if c.feo_amount is not None else None,
                    "planned_quantity": float(c.planned_quantity) if c.planned_quantity is not None else None,
                    "planned_amount": float(c.planned_amount) if c.planned_amount is not None else None,
                    "unit": c.unit,
                    "children": []} for c in all_cats}
    roots = []
    for c in all_cats:
        node = by_id[c.id]
        if c.parent_id and c.parent_id in by_id:
            by_id[c.parent_id]["children"].append(node)
        else:
            roots.append(node)
    return roots


class _UnallocatedBody(BaseModel):
    subsidy_id: int
    parent_id: Optional[int] = None


@router.post("/unallocated")
async def get_or_create_unallocated(
    body: _UnallocatedBody,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Найти или создать категорию «Не определена» (или «Нераспределённое») для субсидии.

    Доступна всем авторизованным пользователям с доступом к субсидии
    (wishes / purchases / feo_categories). Не требует require_tab('feo_categories').

    parent_id (опц.) — создать дочернюю «Не определена» под этим родителем.

    Response: {id, name, subsidy_id, parent_id, created: bool}
    """
    from app.models.subsidy import Subsidy

    # Проверить существование субсидии
    sub = (await db.execute(select(Subsidy).where(Subsidy.id == body.subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(status_code=404, detail="Субсидия не найдена")

    # Изоляция по организации (аналогично list_categories)
    vis = await get_visible_subsidy_ids(current_user, db, "feo_categories")
    if vis is not None:
        vis = vis | await get_visible_subsidy_ids(current_user, db, "purchases")
        vis = vis | await get_visible_subsidy_ids(current_user, db, "wishes")
        if body.subsidy_id not in vis:
            raise HTTPException(status_code=403, detail="Нет доступа к этой субсидии")

    # Загрузить родителя, если указан
    parent: Optional[FeoCategory] = None
    if body.parent_id is not None:
        parent = (await db.execute(
            select(FeoCategory).where(FeoCategory.id == body.parent_id)
        )).scalar_one_or_none()
        if not parent:
            raise HTTPException(status_code=404, detail="Родительская категория не найдена")
        if parent.subsidy_id != body.subsidy_id:
            raise HTTPException(
                status_code=422,
                detail="Родительская категория относится к другой субсидии",
            )

    # Найти существующую (все варианты имён для обратной совместимости)
    stmt = (
        select(FeoCategory)
        .where(FeoCategory.subsidy_id == body.subsidy_id)
        .where(FeoCategory.is_active.is_(True))
        .where(func.lower(FeoCategory.name).in_([
            "не определена",
            "нераспределённое",
            "нераспределенное",
        ]))
    )
    if body.parent_id is None:
        stmt = stmt.where(FeoCategory.parent_id.is_(None))
    else:
        stmt = stmt.where(FeoCategory.parent_id == body.parent_id)

    existing = (await db.execute(stmt.order_by(FeoCategory.id).limit(1))).scalars().first()
    if existing:
        return {
            "id": existing.id,
            "name": existing.name,
            "subsidy_id": existing.subsidy_id,
            "parent_id": existing.parent_id,
            "created": False,
        }

    # Создать новую
    new_level = (parent.level + 1) if parent is not None else 1
    new_cat = FeoCategory(
        name="Не определена",
        subsidy_id=body.subsidy_id,
        parent_id=body.parent_id,
        level=new_level,
        sort_order=9999,
        is_active=True,
    )
    db.add(new_cat)
    await db.commit()
    await db.refresh(new_cat)
    return {
        "id": new_cat.id,
        "name": new_cat.name,
        "subsidy_id": new_cat.subsidy_id,
        "parent_id": new_cat.parent_id,
        "created": True,
    }


@router.post("/", response_model=FeoCategoryOut)
async def create_category(
    category_data: FeoCategoryCreate,
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('feo_categories')),
):
    _validate_plan_pair(category_data.planned_quantity, category_data.planned_amount)

    if category_data.parent_id:
        parent_result = await db.execute(
            select(FeoCategory).where(FeoCategory.id == category_data.parent_id)
        )
        parent = parent_result.scalar_one_or_none()
        if not parent:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Родительская категория не найдена")
        level = parent.level + 1
    else:
        level = 1

    new_category = FeoCategory(
        parent_id=category_data.parent_id,
        subsidy_id=category_data.subsidy_id,
        level=level,
        name=category_data.name,
        code=category_data.code,
        appendix=category_data.appendix,
        is_active=category_data.is_active,
        description=category_data.description,
        budget=category_data.budget,
        feo_quantity=category_data.feo_quantity,
        feo_unit=category_data.feo_unit,
        feo_amount=category_data.feo_amount,
        planned_quantity=category_data.planned_quantity,
        planned_amount=category_data.planned_amount,
        unit=category_data.unit,
        plan_source=category_data.plan_source or "planned_items",
        manual_plan_amount=category_data.manual_plan_amount,
    )
    db.add(new_category)
    await db.commit()
    await db.refresh(new_category)
    return new_category


@router.get("/import/template")
async def download_feo_template(
    _=Depends(get_current_user),
):
    """Шаблон Excel для импорта категорий ФЭО (18 колонок, переверстан 2026-08-14).

    Причина переверстки (боевой файл владельца, «Субсидия ДНР 2.xlsx»): в
    37-колоночном шаблоне между заголовком уровня и его числовыми колонками
    стояло по семь числовых колонок — легко промахнуться и набрать название
    следующего уровня в числовой колонке предыдущего (см. level_name_in_number_column
    в _do_feo_import). Новый шаблон разводит уровни (только названия, 4 колонки)
    и числа (2 ОДНА-на-всю-строку пары колонок «по ФЭО»/«плана», а не по 12
    числовых колонок на каждый уровень) — числа сами прикрепляются к самому
    глубокому заполненному уровню строки, либо к плановой позиции, если она
    заполнена (колонка 5).

    Колонки: Субсидия | Уровень 2 | Уровень 3 | Уровень 4 | Плановая позиция
    (папка НЕ создаётся) | Товар/услуга | Количество/Ед.изм./Цена/Сумма по ФЭО
    (одна пара на строку) | Плановое количество/Ед.изм./Цена/Сумма плана (одна
    пара на строку) | Код | Приложение | Активна | Финансирование (устар.).
    Если уровень пропущен, содержимое нижнего поднимается на его место. Если в
    строке нет ни одного уровня, а плановая позиция заполнена — её название
    становится Уровнем 2. Сумма строки приоритетнее кол-во × цена; расхождение —
    предупреждение. Второй лист «Как заполнять» — текстовые правила.
    """
    if Workbook is None:
        raise HTTPException(500, "openpyxl не установлен")
    wb = Workbook()
    ws = wb.active
    ws.title = "Категории ФЭО"
    headers = [
        "Субсидия",                                          # A   1
        "Уровень 2 (Направление расходов по ФЭО)",           # B   2
        "Уровень 3 (Тип расходов по ФЭО)",                    # C   3
        "Уровень 4 (Конкретизированный)",                     # D   4
        "Плановая позиция (папка НЕ создаётся)",              # E   5
        "Товар/услуга",                                       # F   6
        "Количество по ФЭО",                                  # G   7
        "Ед. изм. по ФЭО",                                    # H   8
        "Цена за единицу по ФЭО",                             # I   9
        "Сумма по ФЭО",                                       # J  10
        "Плановое количество",                                # K  11
        "Ед. изм. плана",                                     # L  12
        "Плановая цена за единицу",                           # M  13
        "Сумма плана",                                        # N  14
        "Код",                                                # O  15
        "Приложение",                                         # P  16
        "Активна",                                            # Q  17
        "Финансирование (устар., можно не заполнять)",        # R  18
    ]
    ws.append(headers)

    # Цветовое кодирование заголовков
    fill_cat  = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")   # синий — субсидия/уровни + ФЭО
    fill_plan = PatternFill(start_color="0891B2", end_color="0891B2", fill_type="solid")   # голубой — план
    fill_item = PatternFill(start_color="059669", end_color="059669", fill_type="solid")   # зелёный — плановая позиция
    fill_attr = PatternFill(start_color="7C3AED", end_color="7C3AED", fill_type="solid")   # фиолетовый — атрибуты
    font_w = Font(color="FFFFFF", bold=True, size=10)

    _blue_cols  = {1, 2, 3, 4, 7, 8, 9, 10}    # субсидия, уровни, «по ФЭО» (одна пара на строку)
    _cyan_cols  = {11, 12, 13, 14}             # план (одна пара на строку)
    _green_cols = {5, 6}                       # плановая позиция + тип

    for i, cell in enumerate(ws[1], start=1):
        if i in _blue_cols:
            cell.fill = fill_cat
        elif i in _cyan_cols:
            cell.fill = fill_plan
        elif i in _green_cols:
            cell.fill = fill_item
        else:
            cell.fill = fill_attr
        cell.font = font_w
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 40

    # Примеры (строки 2–6), 18 элементов каждая:
    # col: 1=Субс 2=Ур2 3=Ур3 4=Ур4 5=Плановая_позиция 6=Тип 7=feoQ 8=feoU 9=feoЦена 10=feoСумма 11=planQ 12=planU 13=planЦена 14=planСумма 15=Код 16=Прил 17=Акт 18=Финанс(устар.)

    # Строка 1: только Ур.2+Ур.3 (Ур.4 пропущен), сумма ФЭО оседает на Ур.3 (самом глубоком уровне строки)
    ws.append(["ДНР_2026", "Техническое оснащение", "Оргтехника", "", "", "", "", "", "", "2000000", "", "", "", "", "01.01.01", "Прил. 1", "да", ""])
    # Строка 2: Ур.4 заполнен — числа «по ФЭО» и «плана» садятся на него (кол-во×цена, сумма не задана — считается)
    ws.append(["ДНР_2026", "Техническое оснащение", "Оргтехника", "Закупка компьютеров", "", "", "6", "шт", "150000", "", "6", "шт", "150000", "900000", "01.01.02", "Прил. 1", "да", ""])
    # Строка 3: боевой кейс владельца — Ур.4 пуст, «Плановая позиция» заполнена: числа плана уходят В ПОЗИЦИЮ, а не на Ур.3
    ws.append(["ДНР_2026", "Прочие расходы", "Расходы на содержание и ремонт ТС", "", "УАЗ Патриот У914ВН 180", "Товар", "", "", "", "", "1", "усл", "178779.59", "178779.59", "02.01.01", "Прил. 2", "да", ""])
    # Строка 4: в строке нет ни одного уровня (Ур.2/3/4 пусты) — «Плановая позиция» становится Уровнем 2, позиция НЕ создаётся
    ws.append(["ДНР_2026", "", "", "", "Услуга по заправке техники", "Услуга", "", "", "", "", "500", "л", "60", "30000", "", "", "да", ""])
    # Строка 5: Ур.3 задан, Ур.4 пропущен, позиции нет — сумма плана садится прямо на Ур.3
    ws.append(["ДНР_2026", "Организация мероприятий", "Слёт студентов-спасателей", "", "", "", "", "", "", "", "", "", "", "3500000", "02.02.01", "Прил. 2", "да", ""])

    # Подсказки в строке 7
    hints = [
        "← Точное название как в системе",                                                              # A   1
        "← Направление расходов (создаётся если нет)",                                                  # B   2
        "← Тип расходов (если пусто — атрибуты к Ур.2); пропущенный уровень поднимается вверх",          # C   3
        "← Конкретизированный (если пусто — к Ур.3); пропущенный уровень поднимается вверх",             # D   4
        "← Плановая позиция — папка НЕ создаётся; если уровней в строке вовсе нет, станет Уровнем 2",    # E   5
        "← Товар / Услуга / Работа (выпадающий список)",                                                 # F   6
        "← Кол-во по ФЭО — одна колонка на строку, привязывается к самому глубокому уровню",             # G   7
        "← Ед. изм. по ФЭО",                                                                              # H   8
        "← Цена за единицу по ФЭО (руб.)",                                                                # I   9
        "← Сумма по ФЭО (руб.); если пусто — кол-во × цена; сумма приоритетнее",                          # J  10
        "← Плановое кол-во — одна колонка на строку; если задана «Плановая позиция», уходит в неё",       # K  11
        "← Ед. изм. плана",                                                                               # L  12
        "← Плановая цена за единицу (руб.)",                                                              # M  13
        "← Сумма плана (руб.); если пусто — кол-во × цена; сумма приоритетнее",                           # N  14
        "← Код категории",                                                                                # O  15
        "← Номер приложения",                                                                             # P  16
        "← да/нет",                                                                                       # Q  17
        "← устарело, можно не заполнять — используйте «Сумма по ФЭО»",                                    # R  18
    ]
    for col, hint in enumerate(hints, start=1):
        ws.cell(7, col).value = hint
        ws.cell(7, col).font = Font(italic=True, color="888888", size=8)

    # Ширины колонок (18 штук)
    col_widths = [18, 42, 42, 42, 42, 14, 16, 14, 16, 16, 16, 14, 16, 16, 10, 12, 10, 22]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(1, i).column_letter].width = w
    ws.freeze_panes = "A2"

    # Выпадающий список Товар/Услуга/Работа на колонку F (Товар/услуга), строки данных 2:1000.
    if DataValidation is not None:
        dv_item_type = DataValidation(
            type="list",
            formula1='"Товар,Услуга,Работа"',
            allow_blank=True,
            showErrorMessage=False,
            showInputMessage=True,
        )
        dv_item_type.error = "Значение не из списка — будет принято как есть"
        dv_item_type.errorTitle = "Нестандартное значение"
        dv_item_type.promptTitle = "Товар/услуга"
        dv_item_type.prompt = "Выберите Товар, Услуга или Работа"
        dv_item_type.sqref = "F2:F1000"
        ws.add_data_validation(dv_item_type)

    # Второй лист — текстовые правила заполнения.
    ws2 = wb.create_sheet("Как заполнять")
    ws2.column_dimensions["A"].width = 110
    rules = [
        "Как заполнять шаблон импорта направлений ФЭО",
        "",
        "1. Уровни (Уровень 2 → Уровень 3 → Уровень 4) идут СЛЕВА НАПРАВО. Уровень 2 — обязателен.",
        "2. Пропущенный уровень поджимается вверх: заполнены Уровень 2 и Уровень 4, а Уровень 3 пуст — Уровень 4 становится ребёнком Уровня 2 напрямую.",
        "3. «Плановая позиция» (колонка 5) — конкретный товар/услуга/работа ВНУТРИ категории. Заполненная «Плановая позиция» папку (категорию) НЕ создаёт — только позицию.",
        "4. Если в строке не заполнен НИ ОДИН из уровней (2/3/4), а «Плановая позиция» заполнена — её название становится Уровнем 2 (создаётся направлением), а позиция при этом не создаётся.",
        "5. «Количество/Ед.изм./Цена/Сумма по ФЭО» и «Плановое количество/Ед.изм./Цена/Сумма плана» — ОДНА пара колонок на всю строку, а не по уровням. Если «Плановая позиция» не задана — числа привязываются к самому глубокому заполненному уровню строки. Если позиция задана — числа плана описывают именно её.",
        "6. Сумма (по ФЭО / плана) ПРИОРИТЕТНЕЕ «количество × цена»: если сумма указана явно — используется она; расхождение с расчётом по кол-во × цена попадёт в предупреждения импорта.",
        "7. «Товар/услуга» (колонка 6) — тип плановой позиции: Товар, Услуга или Работа (выпадающий список в ячейке).",
        "8. Не меняйте заголовки колонок — импорт определяет их по названию, а не по порядку.",
        "9. Строка-подсказка (начинается с «←», строка 7 примера) в файл не попадает — служебная, игнорируется при импорте.",
    ]
    for line in rules:
        ws2.append([line])
    ws2["A1"].font = Font(bold=True, size=12)

    wb.active = 0
    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": _content_disposition("Шаблон_импорта_направлений_ФЭО.xlsx")})


@router.post("/import-preview")
async def feo_import_preview(
    file: UploadFile = File(...),
    _=Depends(require_tab('feo_categories')),
):
    """Read Excel/DOCX file and return headers + sample rows for column mapping."""
    fname = (file.filename or "").lower()
    if not fname.endswith((".xlsx", ".xls", ".docx", ".doc", ".pdf")):
        raise HTTPException(400, "Поддерживаются файлы .xlsx, .xls, .docx, .pdf")

    content = await file.read()

    _FEO_HINTS = (
        "субсидия", "наименован", "направлен", "расходов", "уровень",
        "код", "финансирован", "количеств", "ед. изм", "ед.изм",
        "активн", "приложен", "бюджет", "плановый", "тип расх",
    )

    def _detect_hdr(rows):
        best_score, best_idx = 0, 0
        for ri, row in enumerate(rows[:20]):
            norm = [str(h).strip().lower() if h is not None else "" for h in row]
            score = sum(1 for h in norm if h and any(x in h for x in _FEO_HINTS))
            if score > best_score:
                best_score = score
                best_idx = ri
        return best_idx

    try:
        # ── PDF ──
        if fname.endswith(".pdf"):
            try:
                import pdfplumber
            except ImportError:
                raise HTTPException(500, "pdfplumber не установлен")
            pdf = pdfplumber.open(BytesIO(content))
            all_rows = []
            for page in pdf.pages:
                for t in (page.extract_tables() or []):
                    if t:
                        all_rows.extend([[str(c).strip() if c else "" for c in row] for row in t])
            pdf.close()
            if not all_rows:
                raise HTTPException(400, "Не удалось извлечь таблицы из PDF")
            hdr_idx = _detect_hdr(all_rows)
            headers = [str(h).strip() if h else f"Столбец {j+1}" for j, h in enumerate(all_rows[hdr_idx])]
            data = all_rows[hdr_idx + 1:]
            sample = [[str(c) if c else "" for c in r] for r in data[:5]]
            return {"sheets": [{"name": "PDF", "headers": headers, "sample": sample, "total_rows": len(data), "header_row_offset": hdr_idx}]}

        # ── DOCX ──
        if fname.endswith((".docx", ".doc")):
            try:
                from docx import Document as _DDoc
            except ImportError:
                raise HTTPException(500, "python-docx не установлен")
            doc = _DDoc(BytesIO(content))
            all_rows = []
            for table in doc.tables:
                for row in table.rows:
                    all_rows.append([cell.text.strip() for cell in row.cells])
            if not all_rows:
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_rows.append([text])
            if not all_rows:
                raise HTTPException(400, "Не удалось извлечь данные из документа")
            hdr_idx = _detect_hdr(all_rows)
            headers = [str(h).strip() if h else f"Столбец {j+1}" for j, h in enumerate(all_rows[hdr_idx])]
            data = all_rows[hdr_idx + 1:]
            sample = [[str(c) if c else "" for c in r] for r in data[:5]]
            return {"sheets": [{"name": "Document", "headers": headers, "sample": sample, "total_rows": len(data), "header_row_offset": hdr_idx}]}

        # ── XLS ──
        if fname.endswith(".xls"):
            try:
                import xlrd as _xlrd_mod
            except ImportError:
                raise HTTPException(500, "xlrd не установлен")
            wb_xls = _xlrd_mod.open_workbook(file_contents=content)
            sheets = []
            for sheet_name in wb_xls.sheet_names():
                ws_xls = wb_xls.sheet_by_name(sheet_name)
                all_rows = [list(ws_xls.row_values(i)) for i in range(ws_xls.nrows)]
                if not all_rows:
                    continue
                hdr_idx = _detect_hdr(all_rows)
                hdr_rows = all_rows[hdr_idx:]
                if not hdr_rows:
                    continue
                headers = [str(c).strip() if c else f"Столбец {j+1}" for j, c in enumerate(hdr_rows[0])]
                sample = [[str(c).strip() if c is not None else "" for c in row] for row in hdr_rows[1:min(6, len(hdr_rows))]]
                sheets.append({"name": sheet_name, "headers": headers, "sample": sample,
                               "total_rows": ws_xls.nrows - hdr_idx - 1, "header_row_offset": hdr_idx})

        # ── XLSX ──
        else:
            if load_workbook is None:
                raise HTTPException(500, "openpyxl не установлен")
            wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
            sheets = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                all_rows = list(ws.iter_rows(values_only=True))
                if not all_rows:
                    continue
                hdr_idx = _detect_hdr(all_rows)
                hdr_rows = all_rows[hdr_idx:]
                if not hdr_rows:
                    continue
                headers = [str(c).strip() if c else f"Столбец {j+1}" for j, c in enumerate(hdr_rows[0])]
                sample = [[str(c).strip() if c is not None else "" for c in row] for row in hdr_rows[1:min(6, len(hdr_rows))]]
                sheets.append({"name": sheet_name, "headers": headers, "sample": sample,
                               "total_rows": len(all_rows) - hdr_idx - 1, "header_row_offset": hdr_idx})
            wb.close()

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл ({file.filename}): {e}")

    if not sheets:
        raise HTTPException(400, "Файл не содержит листов с данными")

    return {"sheets": sheets}


async def _do_feo_import(
    rows: list,
    c_subsidy: int | None,
    c_lvl2: int | None,
    c_lvl3: int | None,
    c_lvl4: int | None,
    c_lvl5: int | None,
    c_qty: int | None,
    c_unit: int | None,
    c_item_amt: int | None,
    c_code: int | None,
    c_appendix: int | None,
    c_budget: int | None,
    c_active: int | None,
    db: AsyncSession,
    c_qty_lvl2: int | None = None,
    c_qty_lvl3: int | None = None,
    c_qty_lvl4: int | None = None,
    c_unit_lvl2: int | None = None,
    c_unit_lvl3: int | None = None,
    c_unit_lvl4: int | None = None,
    c_amt_lvl2: int | None = None,
    c_amt_lvl3: int | None = None,
    c_amt_lvl4: int | None = None,
    c_feo_qty_lvl2: int | None = None,
    c_feo_qty_lvl3: int | None = None,
    c_feo_qty_lvl4: int | None = None,
    c_feo_unit_lvl2: int | None = None,
    c_feo_unit_lvl3: int | None = None,
    c_feo_unit_lvl4: int | None = None,
    c_feo_amt_lvl2: int | None = None,
    c_feo_amt_lvl3: int | None = None,
    c_feo_amt_lvl4: int | None = None,
    c_feo_sum_lvl2: int | None = None,
    c_feo_sum_lvl3: int | None = None,
    c_feo_sum_lvl4: int | None = None,
    c_plan_sum_lvl2: int | None = None,
    c_plan_sum_lvl3: int | None = None,
    c_plan_sum_lvl4: int | None = None,
    c_item_price: int | None = None,
    # Новый 18-колоночный шаблон (2026-08-14): числа «по ФЭО»/«плана» — ОДНА пара
    # колонок на всю строку (не по уровням), см. блок «плоские числа» ниже.
    c_row_feo_qty: int | None = None,
    c_row_feo_unit: int | None = None,
    c_row_feo_price: int | None = None,
    c_row_feo_sum: int | None = None,
    c_row_plan_qty: int | None = None,
    c_row_plan_unit: int | None = None,
    c_row_plan_price: int | None = None,
    c_row_plan_sum: int | None = None,
    c_item_type: int | None = None,
    default_subsidy_id: int | None = None,
    dry_run: bool = False,
    user=None,
    remap: str = "",
    apply_remap: bool = False,
) -> dict:
    """Core import logic shared by /import и /import-mapped endpoints.

    dry_run=True: вся обработка выполняется, но транзакция откатывается.
    Возвращает {created, updated, skipped, errors, warnings, ...}.

    N2б: `remap` — JSON-список {"old_id": int, "new_path": str} для явного
    переезда несопоставленных узлов на новые (переезд + удаление опустевших
    старых узлов выполняются после основного цикла, см. блок ниже unmatched).
    Фаза переезда/удаления выполняется ТОЛЬКО при apply_remap=True — иначе
    (обычная загрузка без мастера сопоставления) выполняется только анализ.
    """
    import re as _re
    import json as _json

    try:
        # Нормализатор типа плановой позиции (Товар/Услуга/Работа) — общий с
        # app.routers.feo_planned_items, чтобы не разъезжались правила. Локальный
        # fallback ниже — только на случай, если модуль/функция ещё не готовы
        # (параллельная задача добавляет и её, и поле FeoPlannedItem.item_type);
        # предпочтителен импорт, fallback не должен жить долго.
        from app.routers.feo_planned_items import normalize_item_type
    except ImportError:
        def normalize_item_type(v):
            if not v:
                return None
            s = str(v).strip().lower()
            mapping = {
                "товар": "товар", "товары": "товар", "т": "товар",
                "услуга": "услуга", "услуги": "услуга", "у": "услуга",
                "работа": "работа", "работы": "работа", "р": "работа",
            }
            return mapping.get(s)

    if c_lvl2 is None:
        raise HTTPException(400, "Не найден обязательный столбец: 'Уровень 2 (Направление расходов)'")
    if c_subsidy is None and default_subsidy_id is None:
        raise HTTPException(400, "Укажите столбец 'Субсидия' или выберите субсидию назначения")
    if remap and not apply_remap:
        raise HTTPException(400, "Параметр remap передан без apply_remap=true")

    remap_list: list[dict] = []
    if remap:
        try:
            _raw_remap = _json.loads(remap)
            if not isinstance(_raw_remap, list):
                raise ValueError("ожидался список")
            for _item in _raw_remap:
                if not isinstance(_item, dict) or "old_id" not in _item or "new_path" not in _item:
                    raise ValueError("каждый элемент должен содержать old_id и new_path")
                remap_list.append({
                    "old_id": int(_item["old_id"]),
                    "new_path": str(_item["new_path"]).strip(),
                })
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(400, f"Неверный формат параметра remap: {e}")

    def get_cell(row, col: int | None) -> str | None:
        if col is None or col < 0: return None
        if col >= len(row): return None
        v = row[col]
        if v is None: return None
        s = str(v).strip()
        if not s or s.lower() in ('none', 'null'):
            return None
        return s

    def to_bool(v: str | None) -> bool:
        if v is None: return True
        return v.lower() in ("да", "yes", "true", "1", "+")

    def to_dec(v: str | None):
        if not v: return None
        s = str(v).strip()
        if not s or s in ('-', '—', '–', 'None', 'null', 'н/д', 'N/A'):
            return None
        s = s.replace(" ", "").replace("\xa0", "").replace("\u202f", "")
        s = s.replace("₽", "").replace("руб", "").replace("р.", "").replace("р", "")
        s = s.replace(",", ".")
        s = s.rstrip(".")
        if not s:
            return None
        try:
            return Decimal(s)
        except Exception:
            return None

    def _fmt(v) -> str:
        """Число с разделителями разрядов для читаемых предупреждений."""
        try:
            return f"{float(v):,.2f}".replace(",", " ")
        except Exception:
            return str(v)

    def _norm(s: str) -> str:
        """Нормализация имени уровня для сравнения."""
        s2 = _re.sub(r'^\s*\d+([.\)]\d+)*[.\)]?\s*', '', s)
        s2 = s2.lower().strip()
        return _re.sub(r'\s+', ' ', s2)

    ZERO = Decimal("0")
    QUANT = Decimal("0.01")

    from app.models.subsidy import Subsidy
    sub_rows = (await db.execute(select(Subsidy))).scalars().all()
    sub_by_name = {s.name.lower().strip(): s.id for s in sub_rows}

    existing_cats = (await db.execute(select(FeoCategory))).scalars().all()
    cat_cache: dict[tuple, FeoCategory] = {}
    for c in existing_cats:
        cat_cache[(c.subsidy_id, c.parent_id, c.name.lower().strip())] = c

    # --- Снимок дерева ДО импорта (нужен и для отчёта "несопоставленные узлы",
    # и для фазы переезда/удаления N2б). Перенесено сюда с места ниже (после
    # основного цикла) — переезд/удаление должны опираться на состояние дерева
    # ДО того, как основной цикл его изменит.
    existing_by_id: dict[int, FeoCategory] = {c.id: c for c in existing_cats}
    existing_children: dict[int, list[int]] = {}
    for c in existing_cats:
        if c.parent_id is not None:
            existing_children.setdefault(c.parent_id, []).append(c.id)

    def _get_root_id(cat_id: int) -> int:
        cur = existing_by_id.get(cat_id)
        if cur is None:
            return cat_id
        visited: set[int] = set()
        while cur.parent_id is not None and cur.parent_id in existing_by_id:
            if cur.id in visited:
                break
            visited.add(cur.id)
            cur = existing_by_id[cur.parent_id]
        return cur.id

    def _full_path(cat_id: int) -> str:
        chain: list[str] = []
        cur = existing_by_id.get(cat_id)
        visited: set[int] = set()
        while cur is not None and cur.id not in visited:
            chain.append(cur.name)
            visited.add(cur.id)
            if cur.parent_id is None:
                break
            cur = existing_by_id.get(cur.parent_id)
        return " / ".join(reversed(chain))

    def _subtree_ids_local(root_id: int) -> list[int]:
        ids = [root_id]
        stack = [root_id]
        while stack:
            cur_id = stack.pop()
            for ch_id in existing_children.get(cur_id, []):
                ids.append(ch_id)
                stack.append(ch_id)
        return ids

    def _strip_num(s: str) -> str:
        return _re.sub(r'^\s*\d+([.\)]\d+)*[.\)]?\s*', '', s).strip()

    def _canon_path(path: str, *, lower: bool, yo: bool) -> str:
        out = []
        for seg in path.split(" / "):
            s = _strip_num(seg)
            if lower:
                s = _re.sub(r'\s+', ' ', s.lower()).strip()
            if yo:
                s = s.replace('ё', 'е')
            out.append(s)
        return " / ".join(out)

    created = 0; updated = 0; skipped = 0; errors: list[dict] = []
    warnings: list[dict] = []
    created_details: list[dict] = []
    updated_details: list[dict] = []
    skipped_details: list[dict] = []

    def _check_unit_shift(raw: str | None, row_num: int, name: str, col_label: str) -> str | None:
        """Признак сдвига колонок при импорте ФЭО (задача владельца 2026-08-07,
        прод: часть категорий получила ЧИСЛО в feo_categories.unit, а сумму — в
        planned_amount/feo_amount вместо цены за единицу). Если значение, попавшее
        в колонку «Ед. изм.», само выглядит числом (после нормализации пробелов/
        неразрывных пробелов и запятой→точка) — это почти наверняка не единица
        измерения, а число из соседней колонки, уехавшей на одну позицию влево.
        Не молчим: пишем warning нового вида "column_shift" в тот же массив
        `warnings`, что уже уходит на фронт, и НЕ отдаём число дальше как unit —
        вызывающий код оставит прежнее/пустое значение вместо мусора.
        Существующая проверка sum_mismatch (см. ниже) здесь бессильна: она
        сравнивает сумму с кол-во×цена только когда в файле ЕСТЬ отдельная колонка
        «Сумма по ФЭО» — а типичный сдвиг (нет колонки «Ед.изм.» в исходнике)
        одновременно сдвигает и её, оставляя feo_sum пустым, так что сравнивать
        не с чем. Эта проверка — независимый сигнал, не требующий колонки суммы.
        """
        if raw is None:
            return None
        s = raw.strip().replace(" ", "").replace("\xa0", "").replace(" ", "").replace(",", ".")
        if not s:
            return raw
        try:
            float(s)
        except (ValueError, TypeError):
            return raw
        warnings.append({
            "kind": "column_shift",
            "row": row_num,
            "name": name or "",
            "message": f"Строка {row_num}: в колонке «{col_label}» число {raw} — похоже, колонки сдвинуты, проверьте раскладку файла",
        })
        return None

    def _numeric_shift_scan(row_num: int, level_n: int, next_name: str | None, cols_and_labels: list) -> str | None:
        """Обратный случай column_shift (задача владельца, боевой файл «Субсидия
        ДНР 2.xlsx», 82 строки): не единица измерения попала в числовую колонку,
        а НАЗВАНИЕ следующего уровня — стояло в числовой колонке уровня N (между
        заголовком уровня и его собственными числовыми колонками — до семи
        числовых колонок, легко промахнуться). to_dec() на таком тексте молча
        возвращает None и уровень N+1 просто исчезает — соседний уровень N+2
        (например «Уровень 4») поджимается на его место и становится СОСЕДОМ
        уровня N, а не его ребёнком (баг «лишняя папка-двойник»).

        Признак: значение длиннее 15 символов, не приводится к числу, строка не
        служебная (не начинается с «←»). Если колонка «Уровень N+1» пуста —
        текст принимается как её имя (то самое, что «должно было» туда попасть).
        Если колонка «Уровень N+1» уже заполнена — текст игнорируется, тем не
        менее предупреждаем (возможно опечатка/дублирование, стоит проверить строку).
        Возвращает восстановленное имя уровня N+1 (или None, если восстанавливать нечего).
        """
        recovered: str | None = None
        cur_next = next_name
        for col, label in cols_and_labels:
            raw = get_cell(row, col)
            if raw is None or raw.startswith("←"):
                continue
            if len(raw) <= 15 or to_dec(raw) is not None:
                continue
            short = raw if len(raw) <= 40 else raw[:40] + "…"
            if not cur_next:
                warnings.append({
                    "kind": "level_name_in_number_column",
                    "row": row_num,
                    "name": raw,
                    "message": f"Название «{short}» стояло в колонке «{label}» — принято как Уровень {level_n + 1}",
                })
                recovered = raw
                cur_next = raw
            else:
                warnings.append({
                    "kind": "level_name_in_number_column",
                    "row": row_num,
                    "name": raw,
                    "message": f"Название «{short}» стояло в колонке «{label}» — Уровень {level_n + 1} уже заполнен, проверьте строку",
                })
        return recovered

    # Множество id родительских узлов, затронутых импортом — для parent_sum_mismatch
    touched_parents: set[int] = set()

    # --- Для отчёта "несопоставленные узлы" (только анализ, БЕЗ мутаций) ---
    seen_ids: set[int] = set()      # id всех категорий, вернувшихся из find_or_create
    seen_roots: set[int] = set()    # id корневых (ур.2) узлов, затронутых файлом
    new_paths: list[str] = []       # пути узлов, которые описывает файл ("Ур2 / Ур3 / Ур4")
    _new_paths_seen: set[str] = set()
    new_path_cats: dict[str, FeoCategory] = {}  # путь → объект категории (для резолва remap.new_path)

    # --- 2026-08: план больше не пишется прямо в cat.planned_quantity/planned_amount
    # (владелец требует, чтобы план жил ЗАПИСЯМИ внутри категории — FeoPlannedItem,
    # у которой amount — это СУММА, а не цена за единицу). Пока идёт цикл по строкам,
    # копим посчитанный план каждого уровня сюда, а после цикла (когда уже видно,
    # у какой категории есть дети, а у какой — свои позиции Ур.5) превращаем
    # в FeoPlannedItem, см. блок обработки collected_plan ниже.
    collected_plan: dict[int, dict] = {}  # cat.id -> {"qty","unit","amount","row","name"}
    lvl5_leaves: set[int] = set()         # id категорий, которым в ЭТОМ импорте заведены/обновлены позиции Ур.5
    lvl5_sum_by_cat: dict[int, Decimal] = {}  # сумма amount позиций Ур.5 по категории (из этого импорта)

    async def find_or_create(subsidy_id: int, parent_id, name: str, level: int):
        key = (subsidy_id, parent_id, name.lower().strip())
        if key in cat_cache:
            return cat_cache[key], False
        cat = FeoCategory(
            name=name, subsidy_id=subsidy_id, parent_id=parent_id, level=level,
            is_active=True,
        )
        db.add(cat)
        await db.flush()
        cat_cache[key] = cat
        return cat, True

    # --- N2б: пред-проход ТОЛЬКО ради определения затронутых субсидий (для
    # снимка предыдущей редакции ниже) — лёгкая версия резолва subsidy_id,
    # без накопления ошибок (их соберёт основной цикл).
    touched_subsidies: set[int] = set()
    for _row in rows:
        _lvl2 = get_cell(_row, c_lvl2)
        if not _lvl2 or _lvl2.startswith("←"):
            continue
        _sub_name = get_cell(_row, c_subsidy) if c_subsidy is not None else None
        if _sub_name and not _sub_name.startswith("←"):
            _sid = sub_by_name.get(_sub_name.lower().strip()) or default_subsidy_id
        else:
            _sid = default_subsidy_id
        if _sid:
            touched_subsidies.add(_sid)

    # --- N2б: снимок предыдущей редакции — ВСЕГДА и ДО основного цикла.
    # _create_plan_graph_version заново селектит FeoCategory, поэтому вызывать
    # её нужно именно здесь: после основного цикла дерево уже было бы изменено
    # (создание/обновление/удаление), и снимок перестал бы быть "предыдущей"
    # редакцией. Снимок самодостаточен (дерево пишется в JSON инлайном), так
    # что последующее удаление узлов не портит уже сохранённую версию.
    version_created = False
    if user is not None:
        _remap_note_suffix = ""
        if remap_list:
            _pairs = []
            for _rm in remap_list[:5]:
                _pairs.append(f"«{_full_path(_rm['old_id'])}» → «{_rm['new_path']}»")
            _remap_note_suffix = "; перенос узлов: " + "; ".join(_pairs)
            if len(remap_list) > 5:
                _remap_note_suffix += f" и ещё {len(remap_list) - 5}"
        _note = "Загрузка новой редакции разбивки ФЭО" + _remap_note_suffix
        from app.routers.purchases import _create_plan_graph_version
        for _sid in touched_subsidies:
            _v_created = await _create_plan_graph_version(subsidy_id=_sid, db=db, user=user, note=_note)
            if _v_created:
                version_created = True

    for row_num, row in enumerate(rows, start=2):
        lvl2_name = get_cell(row, c_lvl2)

        # Строка-подсказка («← ...») — всегда служебная, ничего в ней (включая
        # числовые колонки) реальными данными не считаем.
        if lvl2_name and lvl2_name.startswith("←"):
            skipped += 1
            skipped_details.append({"row": row_num, "name": lvl2_name, "reason": "служебная строка"})
            continue

        lvl3_name = get_cell(row, c_lvl3)
        lvl4_name = get_cell(row, c_lvl4)
        lvl5_name = get_cell(row, c_lvl5)

        # Защита от сдвига колонок (боевой файл «Субсидия ДНР 2.xlsx»): название
        # уровня N+1, случайно набранное в числовой колонке уровня N, — вернуть
        # на место ДО того, как из lvl2/3/4/5_name соберётся дерево строки.
        lvl3_name = _numeric_shift_scan(row_num, 2, lvl3_name, [
            (c_feo_qty_lvl2, "Кол-во по ФЭО (Ур.2)"),
            (c_feo_amt_lvl2, "Стоимость по ФЭО (Ур.2)"),
            (c_feo_sum_lvl2, "Сумма по ФЭО (Ур.2)"),
            (c_qty_lvl2, "Плановое кол-во (Ур.2)"),
            (c_amt_lvl2, "Плановая стоимость за ед. (Ур.2)"),
            (c_plan_sum_lvl2, "Сумма плана (Ур.2)"),
        ]) or lvl3_name
        lvl4_name = _numeric_shift_scan(row_num, 3, lvl4_name, [
            (c_feo_qty_lvl3, "Кол-во по ФЭО (Ур.3)"),
            (c_feo_amt_lvl3, "Стоимость по ФЭО (Ур.3)"),
            (c_feo_sum_lvl3, "Сумма по ФЭО (Ур.3)"),
            (c_qty_lvl3, "Плановое кол-во (Ур.3)"),
            (c_amt_lvl3, "Плановая стоимость за ед. (Ур.3)"),
            (c_plan_sum_lvl3, "Сумма плана (Ур.3)"),
        ]) or lvl4_name
        lvl5_name = _numeric_shift_scan(row_num, 4, lvl5_name, [
            (c_feo_qty_lvl4, "Кол-во по ФЭО (Ур.4)"),
            (c_feo_amt_lvl4, "Стоимость по ФЭО (Ур.4)"),
            (c_feo_sum_lvl4, "Сумма по ФЭО (Ур.4)"),
            (c_qty_lvl4, "Плановое кол-во (Ур.4)"),
            (c_amt_lvl4, "Плановая стоимость за ед. (Ур.4)"),
            (c_plan_sum_lvl4, "Сумма плана (Ур.4)"),
        ]) or lvl5_name

        # Позиция без уровней → переезжает на Уровень 2 (задача владельца,
        # шаблон 2026-08-14): если Ур.2/3/4 пусты, а «Плановая позиция» заполнена —
        # её название становится направлением (Ур.2), плановая позиция при этом
        # НЕ создаётся — продолжение правила «пропущенный уровень поджимается
        # вверх», доведённое до предела (когда поджиматься уже некуда). Никакой
        # памяти между строками не заводим, к предыдущей строке ничего не цепляем.
        if (not lvl2_name) and (not lvl3_name or lvl3_name.startswith("←")) and (not lvl4_name or lvl4_name.startswith("←")):
            if lvl5_name and not lvl5_name.startswith("←"):
                warnings.append({
                    "kind": "item_promoted_to_level2",
                    "row": row_num,
                    "name": lvl5_name,
                    "message": f"Плановая позиция «{lvl5_name}» — в строке нет ни одного уровня, создана направлением (Ур.2)",
                })
                lvl2_name = lvl5_name
                lvl5_name = None

        if not lvl2_name:
            skipped += 1
            skipped_details.append({
                "row": row_num,
                "name": lvl2_name or "(пустая строка)",
                "reason": "нет наименования (уровень 2 пуст)",
            })
            continue

        sub_name = get_cell(row, c_subsidy) if c_subsidy is not None else None
        if sub_name and not sub_name.startswith("←"):
            subsidy_id = sub_by_name.get(sub_name.lower().strip()) or default_subsidy_id
            if not subsidy_id:
                errors.append({"row": row_num, "name": lvl2_name, "message": f"Субсидия не найдена: '{sub_name}'"})
                continue
        else:
            subsidy_id = default_subsidy_id
            if not subsidy_id:
                skipped += 1
                skipped_details.append({"row": row_num, "name": lvl2_name, "reason": "не указана субсидия назначения"})
                continue

        code      = get_cell(row, c_code)
        appendix  = get_cell(row, c_appendix)
        budget    = to_dec(get_cell(row, c_budget))
        is_active = to_bool(get_cell(row, c_active))

        item_qty    = to_dec(get_cell(row, c_qty))
        item_unit   = get_cell(row, c_unit)
        item_unit   = _check_unit_shift(
            item_unit, row_num, lvl5_name or lvl4_name or lvl3_name or lvl2_name, "Ед. изм. (Ур.5)"
        )
        item_amount = to_dec(get_cell(row, c_item_amt))
        item_price  = to_dec(get_cell(row, c_item_price)) if c_item_price is not None else None

        raw_item_type = get_cell(row, c_item_type) if c_item_type is not None else None
        item_type = normalize_item_type(raw_item_type) if raw_item_type else None
        if raw_item_type and item_type is None:
            warnings.append({
                "kind": "item_type_unknown",
                "row": row_num,
                "name": lvl5_name or lvl2_name,
                "message": f"Тип позиции «{raw_item_type}» не распознан (ожидались: товар/услуга/работа) — не заполнен",
            })

        # Считать все данные по уровням (raw, без приоритизации)
        _lv = [
            {
                "level_src": 2,
                "name": lvl2_name,
                "feo_qty":  to_dec(get_cell(row, c_feo_qty_lvl2))  if c_feo_qty_lvl2  is not None else None,
                "feo_unit": get_cell(row, c_feo_unit_lvl2) if c_feo_unit_lvl2 is not None else get_cell(row, c_unit_lvl2),
                "feo_amt":  to_dec(get_cell(row, c_feo_amt_lvl2))  if c_feo_amt_lvl2  is not None else None,
                "feo_sum":  to_dec(get_cell(row, c_feo_sum_lvl2))  if c_feo_sum_lvl2  is not None else None,
                "plan_qty": to_dec(get_cell(row, c_qty_lvl2))      if c_qty_lvl2      is not None else None,
                "plan_unit":get_cell(row, c_unit_lvl2)             if c_unit_lvl2     is not None else None,
                "plan_amt": to_dec(get_cell(row, c_amt_lvl2))      if c_amt_lvl2      is not None else None,
                "plan_sum": to_dec(get_cell(row, c_plan_sum_lvl2)) if c_plan_sum_lvl2 is not None else None,
            },
            {
                "level_src": 3,
                "name": lvl3_name,
                "feo_qty":  to_dec(get_cell(row, c_feo_qty_lvl3))  if c_feo_qty_lvl3  is not None else None,
                "feo_unit": get_cell(row, c_feo_unit_lvl3) if c_feo_unit_lvl3 is not None else get_cell(row, c_unit_lvl3),
                "feo_amt":  to_dec(get_cell(row, c_feo_amt_lvl3))  if c_feo_amt_lvl3  is not None else None,
                "feo_sum":  to_dec(get_cell(row, c_feo_sum_lvl3))  if c_feo_sum_lvl3  is not None else None,
                "plan_qty": to_dec(get_cell(row, c_qty_lvl3))      if c_qty_lvl3      is not None else None,
                "plan_unit":get_cell(row, c_unit_lvl3)             if c_unit_lvl3     is not None else None,
                "plan_amt": to_dec(get_cell(row, c_amt_lvl3))      if c_amt_lvl3      is not None else None,
                "plan_sum": to_dec(get_cell(row, c_plan_sum_lvl3)) if c_plan_sum_lvl3 is not None else None,
            },
            {
                "level_src": 4,
                "name": lvl4_name,
                "feo_qty":  to_dec(get_cell(row, c_feo_qty_lvl4))  if c_feo_qty_lvl4  is not None else None,
                "feo_unit": get_cell(row, c_feo_unit_lvl4) if c_feo_unit_lvl4 is not None else get_cell(row, c_unit_lvl4),
                "feo_amt":  to_dec(get_cell(row, c_feo_amt_lvl4))  if c_feo_amt_lvl4  is not None else None,
                "feo_sum":  to_dec(get_cell(row, c_feo_sum_lvl4))  if c_feo_sum_lvl4  is not None else None,
                "plan_qty": to_dec(get_cell(row, c_qty_lvl4))      if c_qty_lvl4      is not None else None,
                "plan_unit":get_cell(row, c_unit_lvl4)             if c_unit_lvl4     is not None else None,
                "plan_amt": to_dec(get_cell(row, c_amt_lvl4))      if c_amt_lvl4      is not None else None,
                "plan_sum": to_dec(get_cell(row, c_plan_sum_lvl4)) if c_plan_sum_lvl4 is not None else None,
            },
        ]

        # Обратная совместимость: если нет явного c_feo_qty_lvlN — берём plan_qty за feo_qty
        for lv in _lv:
            if lv["level_src"] == 2 and c_feo_qty_lvl2 is None and lv["feo_qty"] is None and lv["plan_qty"] is not None:
                lv["feo_qty"] = lv["plan_qty"]
            elif lv["level_src"] == 3 and c_feo_qty_lvl3 is None and lv["feo_qty"] is None and lv["plan_qty"] is not None:
                lv["feo_qty"] = lv["plan_qty"]
            elif lv["level_src"] == 4 and c_feo_qty_lvl4 is None and lv["feo_qty"] is None and lv["plan_qty"] is not None:
                lv["feo_qty"] = lv["plan_qty"]

        # --- Плоские числа нового 18-колоночного шаблона (2026-08-14): «Количество/
        # Ед.изм./Цена/Сумма по ФЭО» и «Плановое количество/Ед.изм./Цена/Сумма плана» —
        # ОДНА пара колонок на всю строку, не по уровням. Прикрепляются к САМОМУ
        # ГЛУБОКОМУ заполненному уровню строки, либо — если заполнена «Плановая
        # позиция» — к переменной позиции (item_qty/item_unit/item_price/item_amount).
        # Значения из per-level колонок (уже посчитаны в _lv выше) ИМЕЮТ ПРИОРИТЕТ —
        # присваиваем только там, где ещё None, — так старые 37-колоночные файлы
        # (с явными per-level колонками) ведут себя ровно как раньше.
        _row_feo_qty   = to_dec(get_cell(row, c_row_feo_qty))   if c_row_feo_qty   is not None else None
        _row_feo_unit  = get_cell(row, c_row_feo_unit)          if c_row_feo_unit  is not None else None
        _row_feo_price = to_dec(get_cell(row, c_row_feo_price)) if c_row_feo_price is not None else None
        _row_feo_sum   = to_dec(get_cell(row, c_row_feo_sum))   if c_row_feo_sum   is not None else None
        _row_plan_qty   = to_dec(get_cell(row, c_row_plan_qty))   if c_row_plan_qty   is not None else None
        _row_plan_unit  = get_cell(row, c_row_plan_unit)          if c_row_plan_unit  is not None else None
        _row_plan_price = to_dec(get_cell(row, c_row_plan_price)) if c_row_plan_price is not None else None
        _row_plan_sum   = to_dec(get_cell(row, c_row_plan_sum))   if c_row_plan_sum   is not None else None

        _deepest_lv = next((lv for lv in reversed(_lv) if lv["name"]), None)

        if _deepest_lv is not None and any(v is not None for v in (_row_feo_qty, _row_feo_unit, _row_feo_price, _row_feo_sum)):
            if _deepest_lv["feo_qty"] is None:
                _deepest_lv["feo_qty"] = _row_feo_qty
            if _deepest_lv["feo_unit"] is None:
                _deepest_lv["feo_unit"] = _row_feo_unit
            if _deepest_lv["feo_amt"] is None:
                _deepest_lv["feo_amt"] = _row_feo_price
            if _deepest_lv["feo_sum"] is None:
                _deepest_lv["feo_sum"] = _row_feo_sum

        if any(v is not None for v in (_row_plan_qty, _row_plan_unit, _row_plan_price, _row_plan_sum)):
            if lvl5_name and not lvl5_name.startswith("←"):
                # «Плановая позиция» заполнена — плоский план описывает ЕЁ (переменную
                # позицию), а не категорию; см. item_qty/item_unit/item_price/item_amount ниже.
                if item_qty is None:
                    item_qty = _row_plan_qty
                if item_unit is None:
                    item_unit = _row_plan_unit
                if item_price is None:
                    item_price = _row_plan_price
                if item_amount is None:
                    item_amount = _row_plan_sum
            elif _deepest_lv is not None:
                if _deepest_lv["plan_qty"] is None:
                    _deepest_lv["plan_qty"] = _row_plan_qty
                if _deepest_lv["plan_unit"] is None:
                    _deepest_lv["plan_unit"] = _row_plan_unit
                if _deepest_lv["plan_amt"] is None:
                    _deepest_lv["plan_amt"] = _row_plan_price
                if _deepest_lv["plan_sum"] is None:
                    _deepest_lv["plan_sum"] = _row_plan_sum

        # Оставляем только уровни с непустым именем
        filled = [lv for lv in _lv if lv["name"]]

        # Схлопываем соседние дубли по нормализованному имени
        deduped: list[dict] = []
        for lv in filled:
            if deduped and _norm(deduped[-1]["name"]) == _norm(lv["name"]):
                warnings.append({
                    "kind": "level_duplicate",
                    "row": row_num,
                    "name": lv["name"],
                    "message": f"Ур.{lv['level_src']} и Ур.{deduped[-1]['level_src']} названы одинаково — склеены в один узел",
                })
                # Числа объединяем с приоритетом нижнего непустого
                prev = deduped[-1]
                for k in ("feo_qty", "feo_unit", "feo_amt", "feo_sum", "plan_qty", "plan_unit", "plan_amt", "plan_sum"):
                    if lv[k] is not None:
                        prev[k] = lv[k]
            else:
                # Проверяем пропуск уровня: если предыдущий src=2 а текущий src=4 — Ур.3 был пропущен
                if deduped and lv["level_src"] - deduped[-1]["level_src"] > 1:
                    warnings.append({
                        "kind": "level_gap",
                        "row": row_num,
                        "name": lv["name"],
                        "message": f"Ур.{lv['level_src']} поднят на место Ур.{deduped[-1]['level_src'] + 1} — промежуточный уровень не заполнен",
                    })
                deduped.append(lv)

        # Нет ни одного заполненного уровня — уже пропустили по lvl2_name выше
        try:
            prev_cat = None
            cats_in_row: list[FeoCategory] = []
            leaf_is_new = False  # будет обновлён на последней итерации
            # Базовый level в БД для первого элемента = 1 (совпадает со старым поведением)
            for seq_idx, lv in enumerate(deduped):
                db_level = seq_idx + 1
                parent_id = prev_cat.id if prev_cat else None
                cat, is_new = await find_or_create(subsidy_id, parent_id, lv["name"], db_level)
                leaf_is_new = is_new  # последнее значение = флаг создания листового узла

                if is_new:
                    created += 1
                    label = {1: "направление (ур. 2)", 2: "категория (ур. 3)", 3: "статья (ур. 4)"}.get(db_level, f"уровень {db_level + 1}")
                    created_details.append({"row": row_num, "name": lv["name"], "reason": label})

                # --- ФЭО-поля ---
                feo_qty  = lv["feo_qty"]
                feo_unit = lv["feo_unit"]
                feo_unit = _check_unit_shift(
                    feo_unit, row_num, lv["name"], f"Ед. изм. по ФЭО (Ур.{lv['level_src']})"
                )
                feo_amt  = lv["feo_amt"]
                feo_sum  = lv["feo_sum"]

                # Сумма ФЭО → budget
                if feo_sum is not None:
                    # Проверяем расхождение с кол-во × цена
                    if feo_qty is not None and feo_amt is not None and feo_qty != ZERO:
                        calc = (feo_qty * feo_amt).quantize(QUANT)
                        if abs(calc - feo_sum) > Decimal("0.01"):
                            warnings.append({
                                "kind": "sum_mismatch",
                                "row": row_num,
                                "name": lv["name"],
                                "message": f"Сумма по ФЭО {_fmt(feo_sum)} ≠ кол-во × цена = {_fmt(calc)}; взята сумма из файла",
                            })
                    if cat.budget != feo_sum:
                        cat.budget = feo_sum
                    # Если feo_amt пуст, но есть кол-во — восстановим цену
                    if feo_amt is None and feo_qty is not None and feo_qty != ZERO:
                        feo_amt = (feo_sum / feo_qty).quantize(QUANT)

                if feo_qty is not None and cat.feo_quantity != feo_qty:
                    cat.feo_quantity = feo_qty
                if feo_unit and cat.feo_unit != feo_unit:
                    cat.feo_unit = feo_unit
                if feo_amt is not None and cat.feo_amount != feo_amt:
                    cat.feo_amount = feo_amt

                # --- Плановые поля ---
                # 2026-08: план строки больше НЕ пишется в cat.planned_quantity/
                # cat.planned_amount — только копится в collected_plan. Присвоение
                # полей категории отсюда убрано намеренно: план должен жить
                # ЗАПИСЬЮ FeoPlannedItem внутри категории (там amount — СУММА,
                # не цена за единицу), а не полями категории. Что именно делать
                # с собранным планом (создать позицию листа / обновить существующую /
                # обнулить поля группы) решается ПОСЛЕ цикла по всем строкам —
                # см. блок обработки collected_plan ниже, там уже видно, у какой
                # категории есть дети, а у какой — свои позиции Ур.5.
                plan_qty  = lv["plan_qty"]
                plan_unit = lv["plan_unit"]
                plan_unit = _check_unit_shift(
                    plan_unit, row_num, lv["name"], f"Ед. изм. плана (Ур.{lv['level_src']})"
                )
                plan_amt  = lv["plan_amt"]
                plan_sum  = lv["plan_sum"]

                _pu = plan_unit if plan_unit is not None else feo_unit

                if plan_sum is not None:
                    # Проверяем расхождение
                    if plan_qty is not None and plan_amt is not None and plan_qty != ZERO:
                        calc_ps = (plan_qty * plan_amt).quantize(QUANT)
                        if abs(calc_ps - plan_sum) > Decimal("0.01"):
                            warnings.append({
                                "kind": "sum_mismatch",
                                "row": row_num,
                                "name": lv["name"],
                                "message": f"Сумма плана {_fmt(plan_sum)} ≠ кол-во × цена = {_fmt(calc_ps)}; взята сумма из файла",
                            })
                    if plan_qty is None:
                        warnings.append({
                            "kind": "sum_without_qty",
                            "row": row_num,
                            "name": lv["name"],
                            "message": f"Сумма плана {_fmt(plan_sum)} задана без кол-во; установлено кол-во = 1",
                        })
                    # amount = сумма плана как есть; qty = кол-во из файла, иначе 1
                    eff_plan_qty = plan_qty if (plan_qty is not None and plan_qty != ZERO) else Decimal("1")
                    collected_plan[cat.id] = {
                        "qty": eff_plan_qty,
                        "unit": _pu,
                        "amount": plan_sum,
                        "row": row_num,
                        "name": cat.name,
                        "item_type": item_type,
                        "from_feo_fallback": False,
                    }
                else:
                    # Старое поведение источника данных: план кол-во/ед/цена напрямую.
                    # Если ОБЕ плановые колонки (кол-во и цена) пусты — сумма целиком
                    # взята из чисел «по ФЭО» (feo_qty × feo_amt), а не из плановых
                    # колонок; помечаем происхождение флагом from_feo_fallback, чтобы
                    # предупреждение plan_vs_items_mismatch ниже не выдавало эту сумму
                    # за «план строки» без уточнения, откуда она на самом деле взята.
                    _pq = plan_qty if plan_qty is not None else feo_qty
                    _pa = plan_amt if plan_amt is not None else feo_amt
                    if _pq is not None and _pa is not None:
                        collected_plan[cat.id] = {
                            "qty": _pq,
                            "unit": _pu,
                            "amount": (_pq * _pa).quantize(QUANT),
                            "row": row_num,
                            "name": cat.name,
                            "item_type": item_type,
                            "from_feo_fallback": plan_qty is None and plan_amt is None,
                        }

                if _pu and cat.unit != _pu:
                    cat.unit = _pu

                cats_in_row.append(cat)
                if prev_cat is not None and prev_cat.id is not None:
                    touched_parents.add(prev_cat.id)
                prev_cat = cat

            leaf = cats_in_row[-1]

            # --- Учёт для отчёта "несопоставленные узлы" (только анализ) ---
            root_c = cats_in_row[0]
            if root_c.id is not None:
                seen_roots.add(root_c.id)
            _path_parts: list[str] = []
            for _c in cats_in_row:
                if _c.id is not None:
                    seen_ids.add(_c.id)
                _path_parts.append(_c.name)
                _p = " / ".join(_path_parts)
                if _p not in _new_paths_seen:
                    _new_paths_seen.add(_p)
                    new_paths.append(_p)
                new_path_cats[_p] = _c

            changed = False
            if code is not None and leaf.code != code:
                leaf.code = code; changed = True
            if appendix is not None and leaf.appendix != appendix:
                leaf.appendix = appendix; changed = True
            # budget из легаси-колонки «Финансирование» пишем только если per-level сумма не задана
            if budget is not None:
                # per-level feo_sum уже записан выше; не перезаписываем
                if deduped and deduped[-1].get("feo_sum") is None:
                    if leaf.budget != budget:
                        leaf.budget = budget; changed = True
            if leaf.is_active != is_active:
                leaf.is_active = is_active; changed = True
            if changed and not leaf_is_new:
                updated += 1
                updated_details.append({"row": row_num, "name": leaf.name, "reason": "обновлены поля категории"})

            if lvl5_name and lvl5_name not in ("←", ""):
                from app.models.feo_planned_item import FeoPlannedItem

                # Вычислить итоговую сумму позиции: item_amount приоритетнее
                eff_item_amount = item_amount
                if eff_item_amount is None and item_price is not None:
                    eff_item_qty = item_qty if item_qty is not None else Decimal("1")
                    eff_item_amount = (item_price * eff_item_qty).quantize(QUANT)

                # Категория получила позицию Ур.5 в ЭТОМ импорте — план строки
                # (собранный выше в collected_plan) для неё уже не отдельная
                # позиция, а описание её содержимого; см. блок ниже.
                lvl5_leaves.add(leaf.id)
                lvl5_sum_by_cat[leaf.id] = lvl5_sum_by_cat.get(leaf.id, ZERO) + (eff_item_amount or ZERO)

                existing_item = (await db.execute(
                    select(FeoPlannedItem).where(
                        FeoPlannedItem.feo_category_id == leaf.id,
                        FeoPlannedItem.name == lvl5_name,
                    )
                )).scalar_one_or_none()
                if not existing_item:
                    # item_type (Товар/Услуга/Работа, задача владельца 2026-08-14):
                    # поле добавляется параллельно в модель FeoPlannedItem — hasattr-
                    # проверка, чтобы этот код не падал, пока миграция ещё не применена.
                    _fpi_kwargs = dict(
                        feo_category_id=leaf.id,
                        name=lvl5_name,
                        quantity=item_qty,
                        unit=item_unit,
                        amount=eff_item_amount,
                        is_active=is_active,
                    )
                    if hasattr(FeoPlannedItem, "item_type"):
                        _fpi_kwargs["item_type"] = item_type
                    pi = FeoPlannedItem(**_fpi_kwargs)
                    db.add(pi)
                    await db.flush()
                    created += 1
                    created_details.append({"row": row_num, "name": lvl5_name, "reason": "плановая позиция (ур. 5)"})
                else:
                    ch2 = False
                    if item_qty is not None and existing_item.quantity != item_qty:
                        existing_item.quantity = item_qty; ch2 = True
                    if item_unit is not None and existing_item.unit != item_unit:
                        existing_item.unit = item_unit; ch2 = True
                    if eff_item_amount is not None and existing_item.amount != eff_item_amount:
                        existing_item.amount = eff_item_amount; ch2 = True
                    if item_type is not None and hasattr(existing_item, "item_type") and existing_item.item_type != item_type:
                        existing_item.item_type = item_type; ch2 = True
                    if ch2:
                        updated += 1
                        updated_details.append({"row": row_num, "name": lvl5_name, "reason": "обновлена позиция"})
                    else:
                        skipped += 1
                        skipped_details.append({"row": row_num, "name": lvl5_name, "reason": "без изменений"})

        except Exception as e:
            errors.append({"row": row_num, "name": lvl2_name, "message": str(e)})

    # Собираем актуальные объекты категорий по id через cat_cache (все объекты
    # уже в сессии) — нужно и для обработки collected_plan ниже, и для проверки
    # "родитель vs сумма детей" после неё.
    cat_by_db_id: dict[int, FeoCategory] = {c.id: c for c in cat_cache.values() if c.id is not None}

    # --- Обработка collected_plan: план строки → FeoPlannedItem внутри
    # категории, а не поля категории (см. комментарий у объявления collected_plan
    # выше цикла по строкам). Делается ЗДЕСЬ, после цикла по всем строкам файла,
    # потому что только сейчас окончательно видно: у категории есть подкатегории
    # (группа) или нет (лист), и получила ли она в этом же импорте отдельные
    # позиции Ур.5.
    if collected_plan:
        from app.models.feo_planned_item import FeoPlannedItem

        _plan_ids = list(collected_plan.keys())
        # Кто из собранных категорий — родитель (группа): есть хотя бы один
        # ребёнок — существующий или только что созданный в этом же импорте
        # (find_or_create уже сделал flush, поэтому дети видны через запрос).
        _parent_rows = (await db.execute(
            select(FeoCategory.parent_id).where(FeoCategory.parent_id.in_(_plan_ids))
        )).scalars().all()
        _groups_with_plan = {pid for pid in _parent_rows if pid is not None}

        for _cat_id, _pdata in collected_plan.items():
            _cat_obj = cat_by_db_id.get(_cat_id)
            if _cat_obj is None:
                continue
            _plan_name = _pdata["name"]
            _plan_row = _pdata["row"]

            if _cat_id in _groups_with_plan:
                # Категория-ГРУППА: собственный план группы в compute_feo_plan_tree
                # (app/services/feo_plan.py) вообще не участвует в расчёте — план
                # группы считается только по сумме подкатегорий. Оставлять здесь
                # значения — мёртвые данные, которые незаметно всплывают и ломают
                # числа, если подкатегория потом пропадёт из файла (боевой случай:
                # категория «Микроавтобус (автобус)» после исчезновения подкатегории
                # (DONGFENG) JUNFENG K33 внезапно показала цену 10 130 000 за штуку
                # и превышение на «Транспорт и техника»).
                if _cat_obj.planned_quantity is not None:
                    _cat_obj.planned_quantity = None
                if _cat_obj.planned_amount is not None:
                    _cat_obj.planned_amount = None
                warnings.append({
                    "kind": "group_plan_ignored",
                    "row": _plan_row,
                    "name": _plan_name,
                    "message": f"План строки «{_plan_name}» не записан: у категории есть подкатегории, план группы считается по ним",
                })
                continue

            if _cat_id in lvl5_leaves:
                # Лист уже описан отдельными позициями Ур.5 в этом же импорте —
                # план строки дублировал бы их сумму, отдельную позицию с именем
                # самой категории не создаём.
                if _cat_obj.planned_quantity is not None:
                    _cat_obj.planned_quantity = None
                if _cat_obj.planned_amount is not None:
                    _cat_obj.planned_amount = None
                _items_sum = lvl5_sum_by_cat.get(_cat_id, ZERO)
                if abs(_items_sum - (_pdata["amount"] or ZERO)) > Decimal("0.01"):
                    # Честное указание источника (задача владельца): если этот план
                    # строки на самом деле собран старым фолбэком из чисел «по ФЭО»
                    # (см. from_feo_fallback выше — плановых кол-ва/цены в файле не
                    # было), пользователь не должен думать, что цифра пришла из
                    # плановых колонок — уточняем происхождение прямо в тексте.
                    _fallback_note = (
                        " (взят из чисел по ФЭО, плановые колонки пустые)"
                        if _pdata.get("from_feo_fallback") else ""
                    )
                    warnings.append({
                        "kind": "plan_vs_items_mismatch",
                        "row": _plan_row,
                        "name": _plan_name,
                        "message": (
                            f"План строки «{_plan_name}» = {_fmt(_pdata['amount'])}{_fallback_note}, а сумма позиций Ур.5 "
                            f"= {_fmt(_items_sum)} — расхождение, план строки не записан"
                        ),
                    })
                continue

            # ЛИСТ без своих позиций Ур.5 в этом импорте — план строки описывает
            # саму категорию; реализуем плановой позицией с именем категории.
            # Ищем существующую активную позицию по точному совпадению имени
            # (TRIM+LOWER) — повторная загрузка того же файла обязана найти и
            # обновить именно её, а не плодить дубли.
            _name_norm = (_plan_name or "").strip().lower()
            _existing_items = (await db.execute(
                select(FeoPlannedItem).where(
                    FeoPlannedItem.feo_category_id == _cat_id,
                    FeoPlannedItem.is_active == True,  # noqa: E712
                )
            )).scalars().all()
            _match_item = next(
                (it for it in _existing_items if (it.name or "").strip().lower() == _name_norm),
                None,
            )
            if _match_item is not None:
                _ch = False
                if _pdata["qty"] is not None and _match_item.quantity != _pdata["qty"]:
                    _match_item.quantity = _pdata["qty"]; _ch = True
                if _pdata["unit"] is not None and _match_item.unit != _pdata["unit"]:
                    _match_item.unit = _pdata["unit"]; _ch = True
                if _pdata["amount"] is not None and _match_item.amount != _pdata["amount"]:
                    _match_item.amount = _pdata["amount"]; _ch = True
                if _pdata.get("item_type") is not None and hasattr(_match_item, "item_type") and _match_item.item_type != _pdata["item_type"]:
                    _match_item.item_type = _pdata["item_type"]; _ch = True
                if _ch:
                    updated += 1
                    updated_details.append({"row": _plan_row, "name": _plan_name, "reason": "плановая позиция из плана строки"})
            else:
                _other_active = [it for it in _existing_items if (it.amount or ZERO) != ZERO]
                if _other_active:
                    # У категории уже есть свои плановые позиции — не задваиваем.
                    warnings.append({
                        "kind": "plan_skipped_has_items",
                        "row": _plan_row,
                        "name": _plan_name,
                        "message": f"У категории «{_plan_name}» уже есть плановые позиции — план строки не записан, чтобы не задвоить",
                    })
                else:
                    _pi_kwargs = dict(
                        feo_category_id=_cat_id,
                        name=(_plan_name or "")[:500],
                        quantity=_pdata["qty"],
                        unit=_pdata["unit"],
                        amount=_pdata["amount"],
                        is_active=True,
                        notes="из импорта ФЭО",
                    )
                    if hasattr(FeoPlannedItem, "item_type"):
                        _pi_kwargs["item_type"] = _pdata.get("item_type")
                    _pi = FeoPlannedItem(**_pi_kwargs)
                    db.add(_pi)
                    await db.flush()
                    created += 1
                    created_details.append({"row": _plan_row, "name": _plan_name, "reason": "плановая позиция из плана строки"})

            if _cat_obj.planned_quantity is not None:
                _cat_obj.planned_quantity = None
            if _cat_obj.planned_amount is not None:
                _cat_obj.planned_amount = None

    # Проверка родитель vs сумма ВСЕХ дочерних узлов (до commit)
    for parent_id in touched_parents:
        if parent_id not in cat_by_db_id:
            continue
        parent_cat = cat_by_db_id[parent_id]
        parent_budget = parent_cat.budget or ZERO
        if not parent_budget:
            continue
        # Сумма budget всех прямых детей из cat_cache (весь справочник)
        children_sum = sum(
            (c.budget or ZERO)
            for c in cat_cache.values()
            if c.parent_id == parent_id and c.id is not None
        )
        if abs(parent_budget - children_sum) > Decimal("0.01"):
            warnings.append({
                "kind": "parent_sum_mismatch",
                "row": None,
                "name": parent_cat.name,
                "message": (
                    f"Родитель «{parent_cat.name}»: бюджет {_fmt(parent_budget)} ≠ "
                    f"сумма всех дочерних узлов {_fmt(children_sum)}; победит значение родителя"
                ),
            })

    # --- Отчёт "несопоставленные узлы": ТОЛЬКО анализ, без мутаций/перепривязок ---
    # Дерево родитель→дети (existing_by_id/existing_children) и хелперы
    # (_get_root_id/_full_path/_subtree_ids_local/_strip_num/_canon_path)
    # определены выше, сразу после загрузки existing_cats — они нужны и здесь,
    # и в фазе переезда/удаления ниже.
    unmatched: list[dict] = []
    if seen_roots:
        candidates = [
            c for c in existing_cats
            if c.id not in seen_ids and _get_root_id(c.id) in seen_roots
        ]
        # Предрасчёт канонических форм всех new_paths (для подсказок)
        _np_nonum = [(np, _canon_path(np, lower=False, yo=False)) for np in new_paths]
        _np_norm = [(np, _canon_path(np, lower=True, yo=False)) for np in new_paths]
        _np_yo = [(np, _canon_path(np, lower=True, yo=True)) for np in new_paths]

        for cand in candidates:
            cand_path = _full_path(cand.id)
            subtree_ids = _subtree_ids_local(cand.id)
            load = await _feo_category_load(subtree_ids, db)
            # own_data (свой план/финансирование/поля ФЭО) считается наравне со
            # внешними ссылками — узел с собственными данными не «пустой», даже
            # если на него никто не ссылается (см. own_data в _feo_category_load,
            # причина — боевая пропажа категории (DONGFENG) JUNFENG K33).
            has_refs = any(load[k] for k in (
                "purchases", "purchase_items", "wishes", "wish_items", "products", "feo_planned_items",
                "own_data",
            ))
            kind = "needs_mapping" if has_refs else "empty"

            suggestion = None
            suggestion_reason = None
            if kind == "needs_mapping":
                cand_nonum = _canon_path(cand_path, lower=False, yo=False)
                for np, np_c in _np_nonum:
                    if np != cand_path and np_c == cand_nonum:
                        suggestion, suggestion_reason = np, "отличается нумерацией"
                        break
                if suggestion is None:
                    cand_norm = _canon_path(cand_path, lower=True, yo=False)
                    for np, np_c in _np_norm:
                        if np != cand_path and np_c == cand_norm:
                            suggestion, suggestion_reason = np, "отличается регистром или пробелами"
                            break
                if suggestion is None:
                    cand_yo = _canon_path(cand_path, lower=True, yo=True)
                    for np, np_c in _np_yo:
                        if np != cand_path and np_c == cand_yo:
                            suggestion, suggestion_reason = np, "отличается ё/е"
                            break

            unmatched.append({
                "id": cand.id,
                "path": cand_path,
                "kind": kind,
                "suggestion": suggestion,
                "suggestion_reason": suggestion_reason,
                "load": {
                    "purchases": load["purchases"],
                    "purchase_items": load["purchase_items"],
                    "wishes": load["wishes"],
                    "wish_items": load["wish_items"],
                    "products": load["products"],
                    "feo_planned_items": load["feo_planned_items"],
                },
                "blocking_purchases": load["blocking_purchases"],
            })

    # --- N2б: переезд (remap) + удаление опустевших старых узлов ---
    # Обязательно ДО dry_run rollback/commit — все строки путей ниже нужно
    # материализовать в обычные str/dict, пока ORM-объекты ещё не просрочены.
    # Выполняется ТОЛЬКО при apply_remap=True (см. docstring) — обычная
    # загрузка (без явного согласия из мастера сопоставления) делает только
    # анализ unmatched/new_paths выше, ничего не переносит и не удаляет.
    relinked_count = 0
    deleted_count = 0
    remap_applied: list[dict] = []
    deleted_details: list[dict] = []
    remap_aborted_reason: str | None = None

    if apply_remap:
        _unmatched_ids = {u["id"] for u in unmatched}

        if errors:
            remap_aborted_reason = "переезд отменён: в файле есть ошибки"
        elif any(sd.get("reason") == "не указана субсидия назначения" for sd in skipped_details):
            remap_aborted_reason = "переезд отменён: часть строк без субсидии назначения"
        elif not seen_ids:
            remap_aborted_reason = "переезд отменён: файл не описал ни одного узла"

        _resolved_remap: list[tuple[int, FeoCategory, str]] = []
        if remap_aborted_reason is None:
            for _rm in remap_list:
                _old_id = _rm["old_id"]
                _new_path = _rm["new_path"]
                if _old_id not in _unmatched_ids:
                    remap_aborted_reason = f"переезд отменён: узел не найден среди несопоставленных (id={_old_id})"
                    break
                _new_cat = new_path_cats.get(_new_path)
                if _new_cat is None or _new_cat.id is None:
                    remap_aborted_reason = f"переезд отменён: цель сопоставления не найдена в новой разбивке: «{_new_path}»"
                    break
                _resolved_remap.append((_old_id, _new_cat, _new_path))

        if remap_aborted_reason is None:
            # Шаг A — применить переезды. Пути берём СЕЙЧАС (пока объекты живы).
            for _old_id, _new_cat, _new_path in _resolved_remap:
                _old_path = _full_path(_old_id)
                _counts = await _relink_feo_category(_old_id, _new_cat.id, db)
                relinked_count += sum(_counts.values())
                remap_applied.append({
                    "old_path": _old_path,
                    "new_path": _new_path,
                    "counts": _counts,
                })

            # Шаг B — удалить опустевшие несопоставленные узлы (кроме тех, кого
            # файл всё-таки назвал где-то в поддереве — их каскадом не трогаем).
            already_deleted: set[int] = set()
            # обходим от корня к листьям (по глубине path), иначе ребёнок удалится раньше родителя и родитель останется пустым висяком
            _unmatched_by_depth = sorted(unmatched, key=lambda _c: _c["path"].count(" / "))
            for _cand in _unmatched_by_depth:
                _cand_id = _cand["id"]
                if _cand_id in already_deleted:
                    continue
                subtree = _subtree_ids_local(_cand_id)
                if any(_sid in already_deleted for _sid in subtree):
                    continue
                if any(_sid in seen_ids for _sid in subtree):
                    continue
                load = await _feo_category_load(subtree, db)
                _real_refs = any(load[k] for k in (
                    "purchases", "purchase_items", "wishes", "wish_items", "products", "feo_planned_items",
                ))
                # own_data (свой план/финансирование/поля ФЭО) считается наравне со
                # внешними ссылками — узел с собственными данными НЕ удаляется, даже
                # если в файле его нет и ссылок на него нет. Боевая причина: категория
                # «(DONGFENG) JUNFENG K33» дважды исчезала с прода — в ней был план
                # (planned_quantity/planned_amount) на 10 130 000, но ссылок не было,
                # и старая проверка has_refs их не видела, поэтому узел молча удалялся.
                has_refs = _real_refs or bool(load.get("own_data"))
                if has_refs:
                    if not _real_refs:
                        warnings.append({
                            "kind": "kept_has_own_plan",
                            "row": None,
                            "name": _cand["path"],
                            "message": (
                                f"Категория «{_cand['path']}» не удалена: в ней есть собственный план или "
                                f"финансирование по ФЭО, хотя в файле её нет. Проверьте, не потерялась ли строка в файле"
                            ),
                        })
                    continue
                for _sid in subtree:
                    if _sid == _cand_id:
                        deleted_details.append({"path": _cand["path"], "reason": "нет в новом файле, ссылок нет"})
                    else:
                        deleted_details.append({"path": _full_path(_sid), "reason": f"внутри удаляемого «{_cand['path']}»"})
                await _purge_feo_categories(subtree, db)
                deleted_count += len(subtree)
                already_deleted.update(subtree)

            # Шаг C — вычистить cat_cache от удалённых id, чтобы ниже по коду
            # (если он появится) не переиспользовать протухшие объекты.
            if already_deleted:
                for _k in [k for k, c in cat_cache.items() if c.id in already_deleted]:
                    del cat_cache[_k]

    if dry_run:
        await db.rollback()
    else:
        await db.commit()

    # Схлопываем одинаковые предупреждения по ключу (kind, name, message-без-суффикса)
    # Сохраняем номер строки первого вхождения и порядок появления
    _seen: dict[tuple, int] = {}   # ключ → индекс в _dedup_warnings
    _dedup_counts: list[int] = []
    _dedup_warnings: list[dict] = []
    for w in warnings:
        key = (w.get("kind"), w.get("name"), w.get("message"))
        if key in _seen:
            _dedup_counts[_seen[key]] += 1
        else:
            _seen[key] = len(_dedup_warnings)
            _dedup_warnings.append(dict(w))
            _dedup_counts.append(1)
    for i, cnt in enumerate(_dedup_counts):
        if cnt > 1:
            _dedup_warnings[i]["message"] = _dedup_warnings[i]["message"] + f" (строк: {cnt})"
    warnings = _dedup_warnings

    return {
        "created": created, "updated": updated, "skipped": skipped,
        "errors": errors, "warnings": warnings,
        "created_details": created_details,
        "updated_details": updated_details, "skipped_details": skipped_details,
        "dry_run": dry_run,
        "unmatched": unmatched,
        "new_paths": new_paths,
        "deleted_count": deleted_count,
        "relinked_count": relinked_count,
        "deleted_details": deleted_details,
        "remap_applied": remap_applied,
        "remap_aborted_reason": remap_aborted_reason,
        "version_created": version_created,
        "deletes_applied": bool(apply_remap),
    }


@router.post("/import")
async def import_feo_from_excel(
    file: UploadFile = File(...),
    dry_run: bool = Query(False),
    remap: str = Query(""),
    apply_remap: bool = Query(False),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_tab('feo_categories')),
):
    """Импорт категорий ФЭО из Excel.
    Формат: Субсидия | Уровень 2 | Уровень 3 | Уровень 4 | Код | Приложение | Финансирование | Активна
    Каждая строка задаёт путь в иерархии. Промежуточные узлы создаются автоматически.
    Код/Приложение/Финансирование/Активна применяются к самому глубокому указанному уровню.
    dry_run=true: возвращает {created, updated, skipped, errors, warnings} без записи в БД.
    Переезд (remap) несопоставленных узлов и удаление опустевших старых узлов выполняются
    только при apply_remap=true; иначе выполняется только анализ (unmatched/new_paths).
    Возвращает {created, updated, skipped, errors, warnings}."""
    if load_workbook is None:
        raise HTTPException(500, "openpyxl не установлен")
    if not (file.filename or "").lower().endswith((".xlsx", ".xls")):
        raise HTTPException(400, "Поддерживаются только .xlsx и .xls")

    content = await file.read()
    wb = load_workbook(BytesIO(content), data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < 2:
        raise HTTPException(400, "Файл пустой")

    # Определяем индексы столбцов по заголовку строки 1
    raw_headers = [str(h).strip().lower() if h is not None else "" for h in rows[0]]

    # Каждая колонка достаётся ровно одному полю
    _used_cols: set[int] = set()

    def find_col(keywords: list[str]) -> int | None:
        for kw in keywords:
            for i, h in enumerate(raw_headers):
                if i in _used_cols:
                    continue
                if kw in h:
                    _used_cols.add(i)
                    return i
        return None

    c_subsidy  = find_col(["субсидия"])
    c_lvl2     = find_col(["уровень 2", "направление расходов", "level 2"])
    c_lvl3     = find_col(["уровень 3", "тип расходов", "level 3"])
    c_lvl4     = find_col(["уровень 4", "конкретизир", "level 4"])
    c_lvl5     = find_col(["плановая позиция", "уровень 5", "плановый товар", "level 5"])
    c_qty      = find_col(["количество (ур.5)", "количество ур.5", "кол-во (ур.5)", "кол-во ур.5"])
    # Новый шаблон: «кол-во по фэо» — feo_quantity
    c_feo_qty_lvl2  = find_col(["кол-во по фэо (ур.2)", "кол-во по фэо ур.2"])
    c_feo_qty_lvl3  = find_col(["кол-во по фэо (ур.3)", "кол-во по фэо ур.3"])
    c_feo_qty_lvl4  = find_col(["кол-во по фэо (ур.4)", "кол-во по фэо ур.4"])
    c_feo_unit_lvl2 = find_col(["ед. изм. по фэо (ур.2)", "ед. изм. по фэо ур.2"])
    c_feo_unit_lvl3 = find_col(["ед. изм. по фэо (ур.3)", "ед. изм. по фэо ур.3"])
    c_feo_unit_lvl4 = find_col(["ед. изм. по фэо (ур.4)", "ед. изм. по фэо ур.4"])
    c_feo_amt_lvl2  = find_col(["стоимость по фэо (ур.2)", "стоимость по фэо ур.2"])
    c_feo_amt_lvl3  = find_col(["стоимость по фэо (ур.3)", "стоимость по фэо ур.3"])
    c_feo_amt_lvl4  = find_col(["стоимость по фэо (ур.4)", "стоимость по фэо ур.4"])
    # Сумма по ФЭО (итог строки) — НОВЫЕ колонки; НЕ путать со стоимостью за ед.
    c_feo_sum_lvl2  = find_col(["сумма по фэо (ур.2)", "сумма по фэо ур.2"])
    c_feo_sum_lvl3  = find_col(["сумма по фэо (ур.3)", "сумма по фэо ур.3"])
    c_feo_sum_lvl4  = find_col(["сумма по фэо (ур.4)", "сумма по фэо ур.4"])
    # Плановое кол-во (CRM-план)
    c_qty_lvl2 = find_col(["плановое кол-во (ур.2)", "плановое кол-во ур.2", "кол-во (ур.2)", "кол-во ур.2", "количество (ур.2)"])
    c_qty_lvl3 = find_col(["плановое кол-во (ур.3)", "плановое кол-во ур.3", "кол-во (ур.3)", "кол-во ур.3", "количество (ур.3)"])
    c_qty_lvl4 = find_col(["плановое кол-во (ур.4)", "плановое кол-во ур.4", "кол-во (ур.4)", "кол-во ур.4", "количество (ур.4)"])
    c_unit_lvl2 = find_col(["ед. изм. плана (ур.2)", "ед. изм. плана ур.2", "ед. изм. (ур.2)", "ед.изм. ур.2", "единица ур.2"])
    c_unit_lvl3 = find_col(["ед. изм. плана (ур.3)", "ед. изм. плана ур.3", "ед. изм. (ур.3)", "ед.изм. ур.3", "единица ур.3"])
    c_unit_lvl4 = find_col(["ед. изм. плана (ур.4)", "ед. изм. плана ур.4", "ед. изм. (ур.4)", "ед.изм. ур.4", "единица ур.4"])
    # Плановая стоимость за ед. — НЕ путать с «сумма плана»
    c_amt_lvl2 = find_col(["плановая стоимость за ед. (ур.2)", "плановая стоимость (ур.2)", "стоимость за ед. (ур.2)", "стоимость ур.2"])
    c_amt_lvl3 = find_col(["плановая стоимость за ед. (ур.3)", "плановая стоимость (ур.3)", "стоимость за ед. (ур.3)", "стоимость ур.3"])
    c_amt_lvl4 = find_col(["плановая стоимость за ед. (ур.4)", "плановая стоимость (ур.4)", "стоимость за ед. (ур.4)", "стоимость ур.4"])
    # Сумма плана — отдельные колонки; старые алиасы «плановая сумма / сумма ур.N» сюда, а НЕ в c_amt_lvl*
    c_plan_sum_lvl2 = find_col(["сумма плана (ур.2)", "плановая сумма (ур.2)", "сумма ур.2"])
    c_plan_sum_lvl3 = find_col(["сумма плана (ур.3)", "плановая сумма (ур.3)", "сумма ур.3"])
    c_plan_sum_lvl4 = find_col(["сумма плана (ур.4)", "плановая сумма (ур.4)", "сумма ур.4"])
    # Новый плоский 18-колоночный шаблон (2026-08-14): числа по ФЭО/плану — ОДНА
    # пара колонок на всю строку (не по уровням), см. _do_feo_import, блок «плоские
    # числа». Объявлены ПОСЛЕ всех per-level find_col выше и ДО generic-фолбэков
    # ниже (c_qty/c_unit) — иначе более общие ключи перехватили бы эти колонки
    # раньше специфичных. На старых 37-колоночных файлах колонки этих названий нет
    # (там «Кол-во по ФЭО (Ур.N)» уже разобран per-level выше и помечен used) — эти
    # find_col останутся None, старое поведение не меняется.
    c_row_feo_qty    = find_col(["количество по фэо"])
    c_row_feo_unit   = find_col(["ед. изм. по фэо"])
    c_row_feo_price  = find_col(["цена за единицу по фэо", "цена за ед. по фэо"])
    c_row_feo_sum    = find_col(["сумма по фэо"])
    c_row_plan_qty   = find_col(["плановое количество"])
    c_row_plan_unit  = find_col(["ед. изм. плана"])
    c_row_plan_price = find_col(["плановая цена за единицу", "плановая цена за ед."])
    c_row_plan_sum   = find_col(["сумма плана"])
    c_item_type      = find_col(["товар/услуга", "тип позиции"])
    # Fallback: generic qty column if no specific level columns present
    if c_qty is None and c_qty_lvl2 is None and c_qty_lvl3 is None and c_qty_lvl4 is None and c_feo_qty_lvl2 is None and c_feo_qty_lvl3 is None and c_feo_qty_lvl4 is None:
        c_qty = find_col(["количество", "кол-во", "qty"])
    c_unit      = find_col(["ед. измерения (ур.5)", "ед. изм. (ур.5)", "единица ур.5", "ед. изм", "единица изм", "ед.изм"])
    c_item_price = find_col(["цена за ед. (ур.5)", "цена за ед. ур.5"])
    c_item_amt  = find_col(["сумма по позиции (ур.5)", "сумма (ур.5)", "сумма ур", "плановая стоимость за ед. (ур.5)", "плановая стоимость (ур.5)", "стоимость за ед. (ур.5)", "стоимость ур.5", "сумма плановая"])
    c_code      = find_col(["код"])
    c_appendix  = find_col(["приложение"])
    c_budget    = find_col(["финансирование", "бюджет", "budget"])
    c_active    = find_col(["активна", "активен", "active"])

    return await _do_feo_import(
        rows=rows[1:],
        c_subsidy=c_subsidy, c_lvl2=c_lvl2, c_lvl3=c_lvl3, c_lvl4=c_lvl4,
        c_lvl5=c_lvl5, c_qty=c_qty, c_unit=c_unit, c_item_amt=c_item_amt,
        c_code=c_code, c_appendix=c_appendix, c_budget=c_budget, c_active=c_active,
        c_qty_lvl2=c_qty_lvl2, c_qty_lvl3=c_qty_lvl3, c_qty_lvl4=c_qty_lvl4,
        c_unit_lvl2=c_unit_lvl2, c_unit_lvl3=c_unit_lvl3, c_unit_lvl4=c_unit_lvl4,
        c_amt_lvl2=c_amt_lvl2, c_amt_lvl3=c_amt_lvl3, c_amt_lvl4=c_amt_lvl4,
        c_feo_qty_lvl2=c_feo_qty_lvl2, c_feo_qty_lvl3=c_feo_qty_lvl3, c_feo_qty_lvl4=c_feo_qty_lvl4,
        c_feo_unit_lvl2=c_feo_unit_lvl2, c_feo_unit_lvl3=c_feo_unit_lvl3, c_feo_unit_lvl4=c_feo_unit_lvl4,
        c_feo_amt_lvl2=c_feo_amt_lvl2, c_feo_amt_lvl3=c_feo_amt_lvl3, c_feo_amt_lvl4=c_feo_amt_lvl4,
        c_feo_sum_lvl2=c_feo_sum_lvl2, c_feo_sum_lvl3=c_feo_sum_lvl3, c_feo_sum_lvl4=c_feo_sum_lvl4,
        c_plan_sum_lvl2=c_plan_sum_lvl2, c_plan_sum_lvl3=c_plan_sum_lvl3, c_plan_sum_lvl4=c_plan_sum_lvl4,
        c_item_price=c_item_price,
        c_row_feo_qty=c_row_feo_qty, c_row_feo_unit=c_row_feo_unit,
        c_row_feo_price=c_row_feo_price, c_row_feo_sum=c_row_feo_sum,
        c_row_plan_qty=c_row_plan_qty, c_row_plan_unit=c_row_plan_unit,
        c_row_plan_price=c_row_plan_price, c_row_plan_sum=c_row_plan_sum,
        c_item_type=c_item_type,
        db=db, dry_run=dry_run,
        user=current_user, remap=remap, apply_remap=apply_remap,
    )


@router.post("/import-mapped")
async def import_feo_mapped(
    file: UploadFile = File(...),
    sheet_name: str = Query(""),
    header_row_offset: int = Query(0),
    col_subsidy: int = Query(-1),
    col_lvl2: int = Query(-1),
    col_lvl3: int = Query(-1),
    col_lvl4: int = Query(-1),
    col_lvl5: int = Query(-1),
    col_code: int = Query(-1),
    col_appendix: int = Query(-1),
    col_budget: int = Query(-1),
    col_quantity: int = Query(-1),
    col_unit: int = Query(-1),
    col_item_amt: int = Query(-1),
    col_active: int = Query(-1),
    col_qty_lvl2: int = Query(-1),
    col_qty_lvl3: int = Query(-1),
    col_qty_lvl4: int = Query(-1),
    col_unit_lvl2: int = Query(-1),
    col_unit_lvl3: int = Query(-1),
    col_unit_lvl4: int = Query(-1),
    col_amt_lvl2: int = Query(-1),
    col_amt_lvl3: int = Query(-1),
    col_amt_lvl4: int = Query(-1),
    col_feo_qty_lvl2: int = Query(-1),
    col_feo_qty_lvl3: int = Query(-1),
    col_feo_qty_lvl4: int = Query(-1),
    col_feo_unit_lvl2: int = Query(-1),
    col_feo_unit_lvl3: int = Query(-1),
    col_feo_unit_lvl4: int = Query(-1),
    col_feo_amount_lvl2: int = Query(-1),
    col_feo_amount_lvl3: int = Query(-1),
    col_feo_amount_lvl4: int = Query(-1),
    # Новые параметры — сумма по ФЭО, сумма плана, цена за ед. Ур.5
    col_feo_sum_lvl2: int = Query(-1),
    col_feo_sum_lvl3: int = Query(-1),
    col_feo_sum_lvl4: int = Query(-1),
    col_plan_sum_lvl2: int = Query(-1),
    col_plan_sum_lvl3: int = Query(-1),
    col_plan_sum_lvl4: int = Query(-1),
    col_item_price: int = Query(-1),
    # Новый плоский 18-колоночный шаблон (2026-08-14) — одна пара колонок «по
    # ФЭО»/«плана» на всю строку, плюс тип плановой позиции.
    col_row_feo_qty: int = Query(-1),
    col_row_feo_unit: int = Query(-1),
    col_row_feo_price: int = Query(-1),
    col_row_feo_sum: int = Query(-1),
    col_row_plan_qty: int = Query(-1),
    col_row_plan_unit: int = Query(-1),
    col_row_plan_price: int = Query(-1),
    col_row_plan_sum: int = Query(-1),
    col_item_type: int = Query(-1),
    default_subsidy_id: int = Query(-1),
    dry_run: bool = Query(False),
    remap: str = Query(""),
    apply_remap: bool = Query(False),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_tab('feo_categories')),
):
    """Импорт категорий ФЭО с пользовательским маппингом столбцов.
    dry_run=true: вся обработка выполняется, транзакция откатывается; возвращает предупреждения.
    Переезд (remap) несопоставленных узлов и удаление опустевших старых узлов выполняются
    только при apply_remap=true; иначе выполняется только анализ (unmatched/new_paths).
    """
    if col_lvl2 < 0:
        raise HTTPException(400, "Не указан обязательный столбец: Уровень 2")
    if col_subsidy < 0 and default_subsidy_id <= 0:
        raise HTTPException(400, "Укажите столбец Субсидия или выберите субсидию назначения")

    fname = (file.filename or "").lower()
    content = await file.read()

    try:
        if fname.endswith(".xls"):
            try:
                import xlrd as _xlrd_mod
            except ImportError:
                raise HTTPException(500, "xlrd не установлен")
            wb_xls = _xlrd_mod.open_workbook(file_contents=content)
            ws_names = wb_xls.sheet_names()
            target_sheet = sheet_name if sheet_name in ws_names else ws_names[0]
            ws_xls = wb_xls.sheet_by_name(target_sheet)
            all_rows = [list(ws_xls.row_values(i)) for i in range(ws_xls.nrows)]
        elif fname.endswith(".pdf"):
            try:
                import pdfplumber
            except ImportError:
                raise HTTPException(500, "pdfplumber не установлен")
            pdf = pdfplumber.open(BytesIO(content))
            all_rows = []
            for page in pdf.pages:
                for t in (page.extract_tables() or []):
                    if t:
                        all_rows.extend([[str(c).strip() if c else "" for c in row] for row in t])
            pdf.close()
        elif fname.endswith((".docx", ".doc")):
            try:
                from docx import Document as _DDoc
            except ImportError:
                raise HTTPException(500, "python-docx не установлен")
            doc = _DDoc(BytesIO(content))
            all_rows = []
            for table in doc.tables:
                for row in table.rows:
                    all_rows.append([cell.text.strip() for cell in row.cells])
        else:
            if load_workbook is None:
                raise HTTPException(500, "openpyxl не установлен")
            wb = load_workbook(BytesIO(content), data_only=True)
            ws_names = wb.sheetnames
            target_sheet = sheet_name if sheet_name in ws_names else ws_names[0]
            ws = wb[target_sheet]
            all_rows = list(ws.iter_rows(values_only=True))
            wb.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл: {e}")

    if len(all_rows) <= header_row_offset + 1:
        raise HTTPException(400, "Файл пустой или не содержит данных после строки заголовка")

    data_rows = all_rows[header_row_offset + 1:]

    return await _do_feo_import(
        rows=data_rows,
        c_subsidy=col_subsidy if col_subsidy >= 0 else None,
        c_lvl2=col_lvl2,
        c_lvl3=col_lvl3 if col_lvl3 >= 0 else None,
        c_lvl4=col_lvl4 if col_lvl4 >= 0 else None,
        c_lvl5=col_lvl5 if col_lvl5 >= 0 else None,
        c_qty=col_quantity if col_quantity >= 0 else None,
        c_unit=col_unit if col_unit >= 0 else None,
        c_item_amt=col_item_amt if col_item_amt >= 0 else None,
        c_code=col_code if col_code >= 0 else None,
        c_appendix=col_appendix if col_appendix >= 0 else None,
        c_budget=col_budget if col_budget >= 0 else None,
        c_active=col_active if col_active >= 0 else None,
        c_qty_lvl2=col_qty_lvl2 if col_qty_lvl2 >= 0 else None,
        c_qty_lvl3=col_qty_lvl3 if col_qty_lvl3 >= 0 else None,
        c_qty_lvl4=col_qty_lvl4 if col_qty_lvl4 >= 0 else None,
        c_unit_lvl2=col_unit_lvl2 if col_unit_lvl2 >= 0 else None,
        c_unit_lvl3=col_unit_lvl3 if col_unit_lvl3 >= 0 else None,
        c_unit_lvl4=col_unit_lvl4 if col_unit_lvl4 >= 0 else None,
        c_amt_lvl2=col_amt_lvl2 if col_amt_lvl2 >= 0 else None,
        c_amt_lvl3=col_amt_lvl3 if col_amt_lvl3 >= 0 else None,
        c_amt_lvl4=col_amt_lvl4 if col_amt_lvl4 >= 0 else None,
        c_feo_qty_lvl2=col_feo_qty_lvl2 if col_feo_qty_lvl2 >= 0 else None,
        c_feo_qty_lvl3=col_feo_qty_lvl3 if col_feo_qty_lvl3 >= 0 else None,
        c_feo_qty_lvl4=col_feo_qty_lvl4 if col_feo_qty_lvl4 >= 0 else None,
        c_feo_unit_lvl2=col_feo_unit_lvl2 if col_feo_unit_lvl2 >= 0 else None,
        c_feo_unit_lvl3=col_feo_unit_lvl3 if col_feo_unit_lvl3 >= 0 else None,
        c_feo_unit_lvl4=col_feo_unit_lvl4 if col_feo_unit_lvl4 >= 0 else None,
        c_feo_amt_lvl2=col_feo_amount_lvl2 if col_feo_amount_lvl2 >= 0 else None,
        c_feo_amt_lvl3=col_feo_amount_lvl3 if col_feo_amount_lvl3 >= 0 else None,
        c_feo_amt_lvl4=col_feo_amount_lvl4 if col_feo_amount_lvl4 >= 0 else None,
        c_feo_sum_lvl2=col_feo_sum_lvl2 if col_feo_sum_lvl2 >= 0 else None,
        c_feo_sum_lvl3=col_feo_sum_lvl3 if col_feo_sum_lvl3 >= 0 else None,
        c_feo_sum_lvl4=col_feo_sum_lvl4 if col_feo_sum_lvl4 >= 0 else None,
        c_plan_sum_lvl2=col_plan_sum_lvl2 if col_plan_sum_lvl2 >= 0 else None,
        c_plan_sum_lvl3=col_plan_sum_lvl3 if col_plan_sum_lvl3 >= 0 else None,
        c_plan_sum_lvl4=col_plan_sum_lvl4 if col_plan_sum_lvl4 >= 0 else None,
        c_item_price=col_item_price if col_item_price >= 0 else None,
        c_row_feo_qty=col_row_feo_qty if col_row_feo_qty >= 0 else None,
        c_row_feo_unit=col_row_feo_unit if col_row_feo_unit >= 0 else None,
        c_row_feo_price=col_row_feo_price if col_row_feo_price >= 0 else None,
        c_row_feo_sum=col_row_feo_sum if col_row_feo_sum >= 0 else None,
        c_row_plan_qty=col_row_plan_qty if col_row_plan_qty >= 0 else None,
        c_row_plan_unit=col_row_plan_unit if col_row_plan_unit >= 0 else None,
        c_row_plan_price=col_row_plan_price if col_row_plan_price >= 0 else None,
        c_row_plan_sum=col_row_plan_sum if col_row_plan_sum >= 0 else None,
        c_item_type=col_item_type if col_item_type >= 0 else None,
        default_subsidy_id=default_subsidy_id if default_subsidy_id > 0 else None,
        db=db, dry_run=dry_run,
        user=current_user, remap=remap, apply_remap=apply_remap,
    )


@router.get("/export")
async def export_feo_to_excel(
    subsidy_id: int = Query(...),
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('feo_categories')),
):
    """Экспорт дерева категорий ФЭО в Excel."""
    if Workbook is None:
        raise HTTPException(500, "openpyxl не установлен")
    from app.models.subsidy import Subsidy
    from app.models.purchase import Purchase

    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")

    cats = (await db.execute(
        select(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id).order_by(FeoCategory.sort_order.nulls_last(), FeoCategory.id)
    )).scalars().all()

    # purchase totals
    pt_rows = (await db.execute(
        select(Purchase.feo_category_id, func.coalesce(func.sum(Purchase.planned_total_price), 0).label("total"))
        .where(Purchase.subsidy_id == subsidy_id)
        .where(Purchase.feo_category_id.isnot(None))
        .group_by(Purchase.feo_category_id)
    )).all()
    purchase_totals = {r.feo_category_id: float(r.total) for r in pt_rows}

    # Build tree
    by_id = {c.id: {"cat": c, "children": []} for c in cats}
    roots = []
    for c in cats:
        if c.parent_id and c.parent_id in by_id:
            by_id[c.parent_id]["children"].append(by_id[c.id])
        else:
            roots.append(by_id[c.id])

    def calc_budget(node):
        c = node["cat"]
        if not node["children"]:
            return float(c.budget) if c.budget is not None else None
        child_sum = sum(v for ch in node["children"] if (v := calc_budget(ch)) is not None)
        return child_sum if any(calc_budget(ch) is not None for ch in node["children"]) else (float(c.budget) if c.budget is not None else None)

    def calc_purchased(node):
        c = node["cat"]
        if not node["children"]:
            return purchase_totals.get(c.id, 0.0)
        return sum(calc_purchased(ch) for ch in node["children"])

    wb = Workbook()
    ws = wb.active
    ws.title = sub.name[:31]

    header_fill = PatternFill(start_color="1D4ED8", end_color="1D4ED8", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    l1_fill = PatternFill(start_color="EFF6FF", end_color="EFF6FF", fill_type="solid")
    l2_fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    bold_font = Font(bold=True)
    semi_font = Font(bold=False)

    ws.append(["Наименование", "Код", "Прил.", "Финансирование по ФЭО (₽)", "Фактически запланировано (₽)", "Активна"])
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 32

    def write_node(node, depth=0):
        c = node["cat"]
        indent = "  " * depth
        budget = calc_budget(node)
        purchased = calc_purchased(node)
        row = [
            indent + c.name,
            c.code or "",
            c.appendix or "",
            budget if budget is not None else "",
            purchased if purchased > 0 else "",
            "Да" if c.is_active else "Нет",
        ]
        ws.append(row)
        r = ws.max_row
        fill = l1_fill if c.level == 1 else (l2_fill if c.level == 2 else None)
        font = bold_font if c.level == 1 else (semi_font)
        for col in range(1, 7):
            cell = ws.cell(r, col)
            if fill:
                cell.fill = fill
            cell.font = font
            if col in (4, 5) and isinstance(cell.value, float):
                cell.number_format = '#,##0.00'
        for ch in node["children"]:
            write_node(ch, depth + 1)

    for root in roots:
        write_node(root)

    ws.column_dimensions["A"].width = 50
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 12
    ws.column_dimensions["D"].width = 22
    ws.column_dimensions["E"].width = 22
    ws.column_dimensions["F"].width = 10
    ws.freeze_panes = "A2"

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    safe_name = sub.name.replace(" ", "_").replace("/", "-")[:40]
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": _content_disposition(f"ФЭО_{safe_name}.xlsx")},
    )


@router.put("/{cat_id}", response_model=FeoCategoryOut)
async def update_category(
    cat_id: int,
    category_data: FeoCategoryCreate,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_tab('feo_categories')),
):
    result = await db.execute(select(FeoCategory).where(FeoCategory.id == cat_id))
    cat = result.scalar_one_or_none()
    if not cat:
        raise HTTPException(status_code=404, detail="Категория не найдена")
    _validate_plan_pair(category_data.planned_quantity, category_data.planned_amount)
    _old_plan = (
        cat.budget, cat.feo_quantity, cat.feo_amount, cat.planned_quantity, cat.planned_amount,
        cat.plan_source, cat.manual_plan_amount,
    )
    cat.name = category_data.name
    cat.code = category_data.code
    cat.appendix = category_data.appendix
    cat.is_active = category_data.is_active
    cat.description = category_data.description
    cat.budget = category_data.budget
    cat.feo_quantity = category_data.feo_quantity
    cat.feo_unit = category_data.feo_unit
    cat.feo_amount = category_data.feo_amount
    cat.planned_quantity = category_data.planned_quantity
    cat.planned_amount = category_data.planned_amount
    cat.unit = category_data.unit
    cat.plan_source = category_data.plan_source or "planned_items"
    cat.manual_plan_amount = category_data.manual_plan_amount

    # Задача владельца «план ≠ факт» (шаг D, сессия 2026-08-06): защита от повторения
    # К1 (боевые 16 760 000 — сумма записана в поле «цена за единицу»). НЕ блокируем
    # (владелец мог ввести именно это осознанно), только предупреждаем в ответе, если
    # planned_quantity × planned_amount более чем вдвое больше финансирования ФЭО —
    # похоже на «сумму вместо цены за единицу» при quantity > 1.
    warning: Optional[str] = None
    if (
        cat.planned_quantity is not None and float(cat.planned_quantity) > 1
        and cat.planned_amount is not None
        and cat.budget is not None and float(cat.budget) > 0
    ):
        _qty = float(cat.planned_quantity)
        _amt = float(cat.planned_amount)
        _budget = float(cat.budget)
        _total = _qty * _amt
        if _total > _budget * 2:
            warning = (
                f"Похоже, в поле «Цена за единицу» введена сумма: "
                f"{_qty:g} × {_amt:,.2f} = {_total:,.2f} ₽ при финансировании {_budget:,.2f} ₽."
            )

    # Владелец, план zany-fluttering-mountain.md (2026-08-13): предупреждение о
    # смене режима расчёта плана, если в категории уже есть плановые позиции —
    # переключение на 'manual_sum' не удаляет их (они продолжают участвовать в
    # excess_plan_over_manual/qty_plan), но раньше именно их сумма БЫЛА планом,
    # теперь плановой суммой становится manual_plan_amount, введённая руками.
    if _old_plan[5] != cat.plan_source:
        from app.models.feo_planned_item import FeoPlannedItem as _FeoPlannedItem
        _items_cnt = (await db.execute(
            select(func.count(_FeoPlannedItem.id))
            .where(_FeoPlannedItem.feo_category_id == cat_id)
            .where(_FeoPlannedItem.is_active.is_(True))
        )).scalar_one()
        if _items_cnt:
            _mode_warning = (
                f"Смена способа расчёта плана: у категории уже есть {_items_cnt} "
                f"плановых позиций. "
                + (
                    "Планом теперь становится введённая сумма, а не их Σ — если сумма "
                    "меньше, появится превышение."
                    if cat.plan_source == "manual_sum"
                    else "Планом снова становится Σ плановых позиций."
                )
            )
            warning = f"{warning} {_mode_warning}" if warning else _mode_warning

    _new_plan = (
        cat.budget, cat.feo_quantity, cat.feo_amount, cat.planned_quantity, cat.planned_amount,
        cat.plan_source, cat.manual_plan_amount,
    )
    if _new_plan != _old_plan and cat.subsidy_id:
        from app.routers.purchases import _create_plan_graph_version
        await _create_plan_graph_version(subsidy_id=cat.subsidy_id, db=db, user=current_user, note=f"Авто-версия: изменение плановых показателей ФЭО «{cat.name}»")
    await db.commit()
    await db.refresh(cat)
    if warning:
        out = FeoCategoryOut.model_validate(cat).model_dump()
        out["warning"] = warning
        return out
    return cat


@router.post("/{cat_id}/align-budget-to-plan")
async def align_budget_to_plan(
    cat_id: int,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_role(*ADMIN_ROLES)),
):
    """«Приравнять ФЭО к сумме плана» — задача владельца (2026-08-12): «должна быть
    кнопка, которая приравнивает сумму ФЭО к сумме плана — условно в конце, когда
    подбивают итоги; это может сделать только админ организации, не ниже».

    Права: require_role(*ADMIN_ROLES) — ADMIN_ROLES = (superadmin, account_owner,
    admin, org_admin), см. app.auth.jwt — org_admin и выше, org_admin включён
    ЯВНО (в подобных кортежах ролей в проекте раньше терялся именно он).

    Целевое значение budget = node["plan"] + node["over"] — ПОЛНАЯ плановая сумма
    узла (см. app.services.feo_plan.compute_feo_plan_tree._visit: full_display),
    а НЕ node["display"]. display может быть уже искусственно занижен до
    plan_manual, если по узлу висит несогласованное превышение плана над ФЭО
    (excess_amount>0 и не excess_approved) — использовать урезанный display
    означало бы занизить итоговое ФЭО и молча оставить расхождение вместо того,
    чтобы его устранить. plan+over — это РЕАЛЬНАЯ сумма того, что запланировано/
    заказано по узлу целиком, именно её и нужно «подбить» в ФЭО при закрытии
    итогов.

    Обязана уважать жёсткий потолок субсидии (задача владельца п.3, см.
    app.services.feo_plan.assert_no_unapproved_excess) — если ПОСЛЕ приравнивания
    суммарный план по субсидии окажется больше суммарного финансирования по ФЭО
    (calculate_budget_from_categories), budget откатывается и действие отклоняется
    с цифрами. Это тот же потолок, что не согласуется НИ ПРИ КАКИХ обстоятельствах —
    у align-budget-to-plan для него тоже нет обхода.

    Response: {id, name, subsidy_id, old_budget, new_budget}. В проекте нет
    отдельного механизма истории/audit-лога для feo_categories (проверено —
    BudgetHistory существует, но её entity_type жёстко "subsidy"/"purchase", для
    категорий ФЭО не заводился и заводить его тут не стали, чтобы не путать
    существующих потребителей этой таблицы) — только before/after в этом ответе.
    """
    from app.services.feo_plan import compute_feo_plan_tree
    from app.routers.subsidies import calculate_budget_from_categories

    cat = await db.get(FeoCategory, cat_id)
    if cat is None:
        raise HTTPException(404, "Категория ФЭО не найдена")
    if not cat.subsidy_id:
        raise HTTPException(422, "У категории не задана субсидия")

    tree = await compute_feo_plan_tree(db, [cat.subsidy_id])
    node = tree.get(cat_id)
    if node is None:
        raise HTTPException(404, "Узел ФЭО не найден в дереве плана")

    old_budget = float(cat.budget) if cat.budget is not None else None
    new_budget = Decimal(str(node["plan"] + node["over"])).quantize(Decimal("0.01"))

    cat.budget = new_budget
    await db.flush()

    # Пересчитать дерево/потолок УЖЕ с новым budget — изменение budget само влияет
    # и на display/excess_amount узла, и на общий потолок финансирования субсидии
    # (см. app.routers.subsidies.calculate_budget_from_categories).
    tree_after = await compute_feo_plan_tree(db, [cat.subsidy_id])
    total_plan_after = sum(n["display"] for n in tree_after.values() if n["parent_id"] is None)
    ceiling_after = await calculate_budget_from_categories(db, cat.subsidy_id)

    if ceiling_after and ceiling_after > 0:
        total_plan_after_d = Decimal(str(total_plan_after))
        ceiling_after_d = Decimal(str(ceiling_after))
        if total_plan_after_d - ceiling_after_d > Decimal("0.005"):
            over_d = total_plan_after_d - ceiling_after_d
            await db.rollback()
            raise HTTPException(
                409,
                {
                    "code": "PLAN_OVER_SUBSIDY_CEILING",
                    "message": (
                        f"Приравнять ФЭО к плану нельзя: после этого суммарный план по субсидии "
                        f"составит {total_plan_after_d:,.2f} ₽, а общий потолок финансирования по "
                        f"ФЭО — {ceiling_after_d:,.2f} ₽ (превышение {over_d:,.2f} ₽). Уменьшите "
                        f"финансирование или план по другим категориям субсидии."
                    ),
                    "subsidy_id": cat.subsidy_id,
                    "total_plan": float(total_plan_after_d),
                    "ceiling": float(ceiling_after_d),
                    "over_amount": float(over_d),
                },
            )

    await db.commit()
    await db.refresh(cat)

    return {
        "id": cat.id,
        "name": cat.name,
        "subsidy_id": cat.subsidy_id,
        "old_budget": old_budget,
        "new_budget": float(cat.budget),
    }


async def _collect_subtree_ids(cat_id: int, db: AsyncSession) -> list[int]:
    """Recursively collect all descendant IDs including the given cat_id."""
    ids = [cat_id]
    children = (await db.execute(
        select(FeoCategory.id).where(FeoCategory.parent_id == cat_id)
    )).scalars().all()
    for cid in children:
        ids.extend(await _collect_subtree_ids(cid, db))
    return ids


async def _update_subtree_levels(cat_id: int, level_delta: int, db: AsyncSession):
    """Recursively update levels of all children by delta."""
    children = (await db.execute(
        select(FeoCategory).where(FeoCategory.parent_id == cat_id)
    )).scalars().all()
    for child in children:
        child.level = child.level + level_delta
        await _update_subtree_levels(child.id, level_delta, db)


# Удаление блокируют только закупки, по которым работа реально идёт
# (стадия «Ведётся работа» и далее). Желания/план закупок/подтверждено —
# работа не начата, категория удаляется, привязка обнуляется (FK SET NULL).
BLOCKING_STATUSES = ("work_in_progress", "contracted", "ordered", "delivered", "paid")


async def _collect_blocking_purchases(ids: list[int], db: AsyncSession) -> list:
    """Закупки в блокирующих статусах, привязанные к переданным категориям
    напрямую (Purchase.feo_category_id) либо через позиции (PurchaseItem.feo_category_id).
    Уникальные по id, поля: id, purchase_number, subject, status."""
    from app.models.purchase import Purchase
    from app.models.purchase_item import PurchaseItem

    direct = (await db.execute(
        select(Purchase.id, Purchase.purchase_number, Purchase.subject, Purchase.status).where(
            Purchase.feo_category_id.in_(ids),
            Purchase.status.in_(BLOCKING_STATUSES),
        )
    )).all()
    via_items = (await db.execute(
        select(Purchase.id, Purchase.purchase_number, Purchase.subject, Purchase.status)
        .join(PurchaseItem, PurchaseItem.purchase_id == Purchase.id)
        .where(
            PurchaseItem.feo_category_id.in_(ids),
            Purchase.status.in_(BLOCKING_STATUSES),
        )
    )).all()

    seen: dict[int, object] = {}
    for row in direct:
        seen[row.id] = row
    for row in via_items:
        seen.setdefault(row.id, row)
    return list(seen.values())


async def _purge_feo_categories(ids: list[int], db: AsyncSession):
    """Отвязать все ссылки на переданные категории и удалить их. Не коммитит —
    коммитит вызывающий."""
    from app.models.purchase import Purchase
    from app.models.purchase_item import PurchaseItem
    from app.models.product import Product
    from app.models.feo_planned_item import FeoPlannedItem

    # Отвязать закупки ранних стадий и их позиции от удаляемого поддерева
    await db.execute(
        Purchase.__table__.update().where(Purchase.feo_category_id.in_(ids)).values(feo_category_id=None)
    )
    await db.execute(
        PurchaseItem.__table__.update().where(PurchaseItem.feo_category_id.in_(ids)).values(feo_category_id=None)
    )

    # Nullify FK references in products before deleting
    await db.execute(
        Product.__table__.update().where(Product.feo_category_id.in_(ids)).values(feo_category_id=None)
    )

    # Nullify FK in purchase_items referencing planned_items of these categories
    planned_item_ids = (await db.execute(
        select(FeoPlannedItem.id).where(FeoPlannedItem.feo_category_id.in_(ids))
    )).scalars().all()
    if planned_item_ids:
        await db.execute(
            PurchaseItem.__table__.update().where(
                PurchaseItem.feo_planned_item_id.in_(planned_item_ids)
            ).values(feo_planned_item_id=None)
        )

    # Delete planned items explicitly (in case DB lacks CASCADE)
    await db.execute(
        FeoPlannedItem.__table__.delete().where(FeoPlannedItem.feo_category_id.in_(ids))
    )

    # Delete categories in one statement — feo_categories.parent_id is
    # ON DELETE CASCADE in the DB, so order doesn't matter. Doing it via
    # a table-level DELETE (instead of ORM db.delete per object) avoids
    # lazy-loading the `children` backref (FeoCategory.parent relationship).
    await db.execute(
        FeoCategory.__table__.delete().where(FeoCategory.id.in_(ids))
    )


async def _relink_feo_category(old_id: int, new_id: int, db: AsyncSession) -> dict:
    """Перевести все ссылки со старой категории на новую. Не коммитит.
    Возвращает счётчик переехавших строк по каждой таблице."""
    from app.models.purchase import Purchase
    from app.models.purchase_item import PurchaseItem
    from app.models.product import Product
    from app.models.feo_planned_item import FeoPlannedItem
    from app.models.wish import Wish
    from app.models.wish_item import WishItem

    counts: dict[str, int] = {}

    def _rowcount(result) -> int:
        rc = result.rowcount
        return rc if rc and rc > 0 else 0

    result = await db.execute(
        Purchase.__table__.update().where(Purchase.feo_category_id == old_id).values(feo_category_id=new_id)
    )
    counts["purchases"] = _rowcount(result)

    result = await db.execute(
        PurchaseItem.__table__.update().where(PurchaseItem.feo_category_id == old_id).values(feo_category_id=new_id)
    )
    counts["purchase_items"] = _rowcount(result)

    result = await db.execute(
        Wish.__table__.update().where(Wish.feo_category_id == old_id).values(feo_category_id=new_id)
    )
    counts["wishes"] = _rowcount(result)

    result = await db.execute(
        WishItem.__table__.update().where(WishItem.feo_category_id == old_id).values(feo_category_id=new_id)
    )
    counts["wish_items"] = _rowcount(result)

    result = await db.execute(
        Product.__table__.update().where(Product.feo_category_id == old_id).values(feo_category_id=new_id)
    )
    counts["products"] = _rowcount(result)

    # Плановые позиции — с дедупликацией по имени (регистронезависимо, trim)
    old_items = (await db.execute(
        select(FeoPlannedItem).where(FeoPlannedItem.feo_category_id == old_id)
    )).scalars().all()
    new_items = (await db.execute(
        select(FeoPlannedItem).where(FeoPlannedItem.feo_category_id == new_id)
    )).scalars().all()
    new_by_name = {(i.name or "").strip().lower(): i for i in new_items}

    planned_items_moved = 0
    for item in old_items:
        key = (item.name or "").strip().lower()
        existing = new_by_name.get(key)
        if existing is not None:
            await db.execute(
                PurchaseItem.__table__.update()
                .where(PurchaseItem.feo_planned_item_id == item.id)
                .values(feo_planned_item_id=existing.id)
            )
            await db.execute(
                FeoPlannedItem.__table__.delete().where(FeoPlannedItem.id == item.id)
            )
        else:
            await db.execute(
                FeoPlannedItem.__table__.update()
                .where(FeoPlannedItem.id == item.id)
                .values(feo_category_id=new_id)
            )
            new_by_name[key] = item
        planned_items_moved += 1
    counts["feo_planned_items"] = planned_items_moved

    return counts


async def _feo_category_load(ids: list[int], db: AsyncSession) -> dict:
    """Что висит на переданных категориях: количества по каждой из ссылающихся
    таблиц + список блокирующих закупок с человекочитаемым статусом.
    Позволяет вызывающему решить, пуст ли узел, и показать пользователю причину.

    own_data (2026-08, боевая причина): раньше "пусто" проверялось ТОЛЬКО по
    внешним ссылкам (закупки/позиции/заявки/товары/плановые позиции). Собственные
    данные категории — план (planned_quantity/planned_amount), финансирование
    (budget) и поля ФЭО (feo_quantity/feo_amount) — не считались ничем, поэтому
    категория с планом на 10 130 000 руб., но без единой ссылки, признавалась
    пустой и удалялась молча. Так дважды пропадала с прода категория
    «(DONGFENG) JUNFENG K33». own_data = сколько из переданных id имеют хоть одно
    из этих полей not NULL и не ноль — вызывающий код обязан учитывать его наравне
    с остальными ссылками при решении "пуст ли узел".
    """
    from app.models.purchase import Purchase
    from app.models.purchase_item import PurchaseItem
    from app.models.product import Product
    from app.models.feo_planned_item import FeoPlannedItem
    from app.models.wish import Wish
    from app.models.wish_item import WishItem
    from app.routers.purchase_transitions import STATUS_LABELS
    from sqlalchemy import or_, and_

    purchases_count = (await db.execute(
        select(func.count(Purchase.id)).where(Purchase.feo_category_id.in_(ids))
    )).scalar_one()
    purchase_items_count = (await db.execute(
        select(func.count(PurchaseItem.id)).where(PurchaseItem.feo_category_id.in_(ids))
    )).scalar_one()
    wishes_count = (await db.execute(
        select(func.count(Wish.id)).where(Wish.feo_category_id.in_(ids))
    )).scalar_one()
    wish_items_count = (await db.execute(
        select(func.count(WishItem.id)).where(WishItem.feo_category_id.in_(ids))
    )).scalar_one()
    products_count = (await db.execute(
        select(func.count(Product.id)).where(Product.feo_category_id.in_(ids))
    )).scalar_one()
    planned_items_count = (await db.execute(
        select(func.count(FeoPlannedItem.id)).where(FeoPlannedItem.feo_category_id.in_(ids))
    )).scalar_one()
    own_data_count = (await db.execute(
        select(func.count(FeoCategory.id)).where(
            FeoCategory.id.in_(ids),
            or_(
                and_(FeoCategory.budget.isnot(None), FeoCategory.budget != 0),
                and_(FeoCategory.feo_quantity.isnot(None), FeoCategory.feo_quantity != 0),
                and_(FeoCategory.feo_amount.isnot(None), FeoCategory.feo_amount != 0),
                and_(FeoCategory.planned_quantity.isnot(None), FeoCategory.planned_quantity != 0),
                and_(FeoCategory.planned_amount.isnot(None), FeoCategory.planned_amount != 0),
            ),
        )
    )).scalar_one()

    blocking = await _collect_blocking_purchases(ids, db)
    purchases_list = [
        {
            "id": p.id,
            "purchase_number": p.purchase_number,
            "subject": p.subject,
            "status": p.status,
            "status_label": STATUS_LABELS.get(p.status, p.status),
        }
        for p in blocking
    ]

    return {
        "purchases": purchases_count,
        "purchase_items": purchase_items_count,
        "wishes": wishes_count,
        "wish_items": wish_items_count,
        "products": products_count,
        "feo_planned_items": planned_items_count,
        "own_data": own_data_count,
        "blocking_purchases": purchases_list,
    }


@router.get("/{cat_id}/subtree")
async def get_category_subtree(
    cat_id: int,
    db: AsyncSession = Depends(get_db),
    _=Depends(get_current_user),
):
    """Имя категории + id всего поддерева (для фильтра «закупки этой категории»)."""
    cat = (await db.execute(
        select(FeoCategory).where(FeoCategory.id == cat_id)
    )).scalar_one_or_none()
    if not cat:
        raise HTTPException(status_code=404, detail="Категория не найдена")
    ids = await _collect_subtree_ids(cat_id, db)
    return {"id": cat.id, "name": cat.name, "ids": ids}


@router.delete("/{cat_id}")
async def delete_category(
    cat_id: int,
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('feo_categories')),
):
    cat = (await db.execute(select(FeoCategory).where(FeoCategory.id == cat_id))).scalar_one_or_none()
    if not cat:
        raise HTTPException(status_code=404, detail="Категория не найдена")

    # Collect entire subtree
    all_ids = await _collect_subtree_ids(cat_id, db)

    linked_purchases = await _collect_blocking_purchases(all_ids, db)
    if linked_purchases:
        purchase_ids = [p.id for p in linked_purchases]
        raise HTTPException(
            status_code=409,
            detail={
                "message": (
                    f"Нельзя удалить: {len(linked_purchases)} закупок со стадией "
                    f"«Ведётся работа» и далее привязано к этой категории. "
                    f"Закупки на ранних стадиях (желания, план закупок) удалению не мешают."
                ),
                "purchase_ids": purchase_ids,
                "feo_category_ids": all_ids,
            }
        )

    await _purge_feo_categories(all_ids, db)
    await db.commit()
    deleted_count = len(all_ids)
    return {"ok": True, "deleted_count": deleted_count}


@router.patch("/{cat_id}/move")
async def move_category(
    cat_id: int,
    data: dict,
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('feo_categories')),
):
    """Move category to a new parent. Set parent_id=null to make root."""
    cat = (await db.execute(select(FeoCategory).where(FeoCategory.id == cat_id))).scalar_one_or_none()
    if not cat:
        raise HTTPException(status_code=404, detail="Категория не найдена")

    new_parent_id = data.get("parent_id")
    warning = None

    # Determine new level
    if new_parent_id:
        parent = (await db.execute(select(FeoCategory).where(FeoCategory.id == new_parent_id))).scalar_one_or_none()
        if not parent:
            raise HTTPException(status_code=404, detail="Родительская категория не найдена")
        if parent.subsidy_id != cat.subsidy_id:
            raise HTTPException(status_code=400, detail="Нельзя переместить в другую субсидию")
        new_level = parent.level + 1
    else:
        new_level = 1

    level_delta = new_level - cat.level

    # Check for circular reference
    if new_parent_id:
        subtree_ids = await _collect_subtree_ids(cat_id, db)
        if new_parent_id in subtree_ids:
            raise HTTPException(status_code=400, detail="Нельзя переместить в собственное поддерево")

    # Check if new max depth creates a previously unseen level
    max_child_depth = 0
    all_ids = await _collect_subtree_ids(cat_id, db)
    if len(all_ids) > 1:
        max_existing = (await db.execute(
            select(func.max(FeoCategory.level)).where(FeoCategory.id.in_(all_ids))
        )).scalar() or cat.level
        max_child_depth = max_existing - cat.level
    new_max_level = new_level + max_child_depth
    existing_max = (await db.execute(
        select(func.max(FeoCategory.level)).where(FeoCategory.subsidy_id == cat.subsidy_id)
    )).scalar() or 1
    if new_max_level > existing_max:
        warning = f"Создан новый уровень вложенности: {new_max_level}"

    # Update category
    cat.parent_id = new_parent_id
    cat.level = new_level

    # Recursively update children levels
    if level_delta != 0:
        await _update_subtree_levels(cat_id, level_delta, db)

    await db.commit()
    result = {"ok": True, "new_level": new_level}
    if warning:
        result["warning"] = warning
    return result


@router.patch("/{cat_id}/reorder")
async def reorder_category(
    cat_id: int,
    data: dict,
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('feo_categories')),
):
    """Move a category up/down among its siblings (same parent, same subsidy).
    Body: {"direction": "up"|"down"}. Swaps sort_order with the adjacent sibling."""
    direction = (data.get("direction") or "").lower()
    if direction not in ("up", "down"):
        raise HTTPException(status_code=400, detail="direction должен быть 'up' или 'down'")
    cat = (await db.execute(select(FeoCategory).where(FeoCategory.id == cat_id))).scalar_one_or_none()
    if not cat:
        raise HTTPException(status_code=404, detail="Категория не найдена")
    sib_filter = (
        FeoCategory.parent_id.is_(None) if cat.parent_id is None
        else FeoCategory.parent_id == cat.parent_id
    )
    siblings = (await db.execute(
        select(FeoCategory)
        .where(FeoCategory.subsidy_id == cat.subsidy_id, sib_filter)
        .order_by(FeoCategory.sort_order.nulls_last(), FeoCategory.id)
    )).scalars().all()
    # Normalize any NULL sort_order to a stable increasing sequence
    for i, s in enumerate(siblings):
        if s.sort_order is None:
            s.sort_order = (i + 1) * 10
    idx = next((i for i, s in enumerate(siblings) if s.id == cat_id), None)
    if idx is None:
        raise HTTPException(status_code=404, detail="Категория не найдена среди соседей")
    swap_idx = idx - 1 if direction == "up" else idx + 1
    if swap_idx < 0 or swap_idx >= len(siblings):
        await db.commit()
        return {"ok": True, "moved": False}
    a, b = siblings[idx], siblings[swap_idx]
    a.sort_order, b.sort_order = b.sort_order, a.sort_order
    await db.commit()
    return {"ok": True, "moved": True}
