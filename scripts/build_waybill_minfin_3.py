"""Создание skeleton .docx шаблона Минтранс №3 (путевой лист легкового автомобиля).
Использует docxtpl placeholders {{ name }} для динамической подстановки.
Места для рукописных подписей механика/медика — пустые линии (не canvas)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pathlib

OUT = pathlib.Path(__file__).parent.parent / 'backend' / 'templates' / 'waybill_form_3.docx'
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# Заголовок
title = doc.add_heading('ПУТЕВОЙ ЛИСТ', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('легкового автомобиля')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Шапка: № и дата
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_nr = header.add_run('№ ')
run_nr.bold = True
run_num = header.add_run('{{ waybill_number }}')
run_num.bold = True
header.add_run('  от  ')
header.add_run('{{ waybill_date }}')

doc.add_paragraph()

# Организация
org_p = doc.add_paragraph()
run_org = org_p.add_run('Организация: ')
run_org.bold = True
org_p.add_run('{{ org_name }}')

org_addr = doc.add_paragraph()
run_addr = org_addr.add_run('Адрес: ')
run_addr.bold = True
org_addr.add_run('{{ org_address }}')

doc.add_paragraph()

# Транспортное средство
veh_h = doc.add_paragraph()
veh_h.add_run('ТРАНСПОРТНОЕ СРЕДСТВО').bold = True

veh_t = doc.add_table(rows=4, cols=2)
veh_t.style = 'Light Grid Accent 1'
rows = veh_t.rows
rows[0].cells[0].text = 'Марка, модель'
rows[0].cells[1].text = '{{ vehicle_brand_model }}'
rows[1].cells[0].text = 'Гос. рег. знак'
rows[1].cells[1].text = '{{ vehicle_plate }}'
rows[2].cells[0].text = 'Гаражный №'
rows[2].cells[1].text = '{{ vehicle_garage_number }}'
rows[3].cells[0].text = 'VIN'
rows[3].cells[1].text = '{{ vehicle_vin }}'

doc.add_paragraph()

# Водитель
drv_h = doc.add_paragraph()
drv_h.add_run('ВОДИТЕЛЬ').bold = True

drv_t = doc.add_table(rows=4, cols=2)
drv_t.style = 'Light Grid Accent 1'
rows = drv_t.rows
rows[0].cells[0].text = 'ФИО'
rows[0].cells[1].text = '{{ driver_full_name }}'
rows[1].cells[0].text = 'Категория ВУ'
rows[1].cells[1].text = '{{ driver_license_categories }}'
rows[2].cells[0].text = '№ ВУ'
rows[2].cells[1].text = '{{ driver_license_number }}'
rows[3].cells[0].text = 'Табельный №'
rows[3].cells[1].text = '{{ driver_tab_number }}'

doc.add_paragraph()

# Маршрут
route_h = doc.add_paragraph()
route_h.add_run('МАРШРУТ И ЗАДАНИЕ').bold = True

doc.add_paragraph('Цель поездки: {{ task_description }}')
doc.add_paragraph('{% for stop in route_stops %}')
doc.add_paragraph('  {{ stop.ord }}. {{ stop.name }} — план: {{ stop.planned_time }} {% if stop.actual_time %}факт: {{ stop.actual_time }}{% endif %}')
doc.add_paragraph('{% endfor %}')

doc.add_paragraph()
doc.add_paragraph('Плановый пробег: {{ planned_mileage_km }} км')
doc.add_paragraph('Фактический пробег: {{ actual_mileage_km }} км')

doc.add_paragraph()

# Спидометр и топливо
fuel_h = doc.add_paragraph()
fuel_h.add_run('СПИДОМЕТР И ТОПЛИВО').bold = True

fuel_t = doc.add_table(rows=4, cols=3)
fuel_t.style = 'Light Grid Accent 1'
rows = fuel_t.rows
rows[0].cells[0].text = ''
rows[0].cells[1].text = 'При выезде'
rows[0].cells[2].text = 'При возврате'
rows[1].cells[0].text = 'Спидометр, км'
rows[1].cells[1].text = '{{ odometer_start }}'
rows[1].cells[2].text = '{{ odometer_finish }}'
rows[2].cells[0].text = 'Топливо, л'
rows[2].cells[1].text = '{{ fuel_start }}'
rows[2].cells[2].text = '{{ fuel_end }}'
rows[3].cells[0].text = 'Заправлено за смену, л'
rows[3].cells[1].text = '{{ fuel_issued }}'
rows[3].cells[2].text = ''

doc.add_paragraph()

# Груз
cargo_h = doc.add_paragraph()
cargo_h.add_run('ГРУЗ').bold = True
doc.add_paragraph('Описание: {{ cargo_description }}')
doc.add_paragraph('Пассажиры: {{ passengers_count }}')

doc.add_paragraph()

# ПРЕДРЕЙСОВЫЙ ОСМОТР
pre_h = doc.add_paragraph()
pre_h.add_run('ПРЕДРЕЙСОВЫЙ ОСМОТР').bold = True

# Механик
doc.add_paragraph('Технический осмотр:')
doc.add_paragraph('Механик: {{ pre_trip_mechanic_name }}')
doc.add_paragraph('Время: {{ pre_trip_mechanic_inspected_at }}')
doc.add_paragraph('Результат: {{ pre_trip_mechanic_result }}')
mech_sig = doc.add_paragraph()
mech_sig.add_run('Подпись механика: ').bold = True
mech_sig.add_run('_' * 30)

# Медик
doc.add_paragraph()
doc.add_paragraph('Медицинский осмотр:')
doc.add_paragraph('Медик: {{ pre_trip_doctor_name }}')
doc.add_paragraph('Время: {{ pre_trip_doctor_inspected_at }}')
doc.add_paragraph('Результат: {{ pre_trip_doctor_result }}')
doc_sig = doc.add_paragraph()
doc_sig.add_run('Подпись медика: ').bold = True
doc_sig.add_run('_' * 30)

doc.add_paragraph()
doc.add_paragraph()

# ПОСЛЕРЕЙСОВЫЙ ОСМОТР
post_h = doc.add_paragraph()
post_h.add_run('ПОСЛЕРЕЙСОВЫЙ ОСМОТР').bold = True

doc.add_paragraph('Технический осмотр:')
doc.add_paragraph('Механик: {{ post_trip_mechanic_name }}')
doc.add_paragraph('Время: {{ post_trip_mechanic_inspected_at }}')
doc.add_paragraph('Результат: {{ post_trip_mechanic_result }}')
mech_sig2 = doc.add_paragraph()
mech_sig2.add_run('Подпись механика: ').bold = True
mech_sig2.add_run('_' * 30)

doc.add_paragraph()
doc.add_paragraph('Медицинский осмотр:')
doc.add_paragraph('Медик: {{ post_trip_doctor_name }}')
doc.add_paragraph('Время: {{ post_trip_doctor_inspected_at }}')
doc.add_paragraph('Результат: {{ post_trip_doctor_result }}')
doc_sig2 = doc.add_paragraph()
doc_sig2.add_run('Подпись медика: ').bold = True
doc_sig2.add_run('_' * 30)

doc.add_paragraph()
doc.add_paragraph()

# Подпись водителя
drv_sig_h = doc.add_paragraph()
drv_sig_h.add_run('ПОДПИСЬ ВОДИТЕЛЯ').bold = True
doc.add_paragraph('Водитель сдал путевой лист: {{ driver_full_name }}')
doc.add_paragraph('Дата и время сдачи: {{ driver_signed_at }}')
drv_sig = doc.add_paragraph()
drv_sig.add_run('Подпись водителя: ').bold = True
drv_sig.add_run('{{ driver_signature_image }}')  # Здесь подставится InlineImage в docxtpl

doc.save(OUT)
print(f'Шаблон Минтранс №3 создан: {OUT}')
print(f'   Размер: {OUT.stat().st_size} bytes')
