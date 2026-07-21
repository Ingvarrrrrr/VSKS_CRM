"""
Generates four Fabrikant docxtpl templates for the GALA procurement system.

Run from any working directory:
    python backend/templates/make_fabrikant_templates.py

Creates:
    backend/templates/fabrikant_instruction.docx
    backend/templates/fabrikant_application_form.docx
    backend/templates/fabrikant_documentation.docx
    backend/templates/fabrikant_contract_project.docx

Template variable reference (context keys added in documents.py):
    notice_number, subject, customer_full_name, customer_name,
    customer_address, customer_inn, customer_kpp, customer_ogrn,
    customer_bank_name, customer_settlement_account, customer_bik,
    customer_correspondent_account, customer_signatory,
    customer_signatory_position, customer_signatory_initials,
    customer_signatory_basis, contract_city, total_nmcd, contract_price,
    contract_price_words, contract_number, contract_date,
    contract_date_day, contract_date_month, contract_date_year,
    delivery_location, execution_term, submission_deadline_datetime,
    applications_review_date, payment_term_days, subsidy_agreement_number,
    subsidy_name, contractor_full_name, contractor_name, contractor_inn,
    contractor_kpp, contractor_ogrn, contractor_address,
    contractor_settlement_account, contractor_bank_name, contractor_bik,
    contractor_correspondent_account, contractor_signatory,
    contractor_signatory_position, contractor_signatory_initials,
    contractor_signatory_basis, items (list of dicts), today, vat_info_line
"""

import os
import shutil

BASE = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__))))
SRC_DIR = os.path.normpath(os.path.join(BASE, '..', '..', 'Доработки', 'Настраиваем ФАБРИКАНТ'))


def _add(doc, text, bold=False, size=None, align=None, space_before=0, space_after=2, indent=0):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p


def make_instruction():
    """Copy instruction as-is — it is a static document without variables."""
    src = os.path.join(SRC_DIR, 'Инструкция (краткая) (1) (2).docx')
    dst = os.path.join(BASE, 'fabrikant_instruction.docx')
    shutil.copy2(src, dst)
    print(f'  instruction: copied from source -> {dst}')


def make_application_form():
    """
    fabrikant_application_form.docx — Приложение 2: Форма заявки участника.
    Variables: notice_number, subject
    All other fields (организация, ИНН, банк и т.д.) заполняются участником вручную
    и остаются пустыми строками в шаблоне.
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Приложение №2').bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('к запросу цен на право заключения').italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('договора поставки').italic = True

    _add(doc, '', space_before=6)

    _add(doc, 'ФОРМА ЗАЯВКИ УЧАСТНИКА', bold=True, align='center', size=13, space_before=6, space_after=4)
    _add(doc, 'на право заключения договора поставки', align='center', space_after=6)

    p = doc.add_paragraph()
    p.add_run('№ извещения: ')
    p.add_run('{{ notice_number }}').bold = True

    p = doc.add_paragraph()
    p.add_run('Предмет закупки: ')
    p.add_run('{{ subject }}').bold = True

    _add(doc, '', space_before=4)

    p = doc.add_paragraph()
    p.add_run(
        'Ознакомившись с запросом цен, приглашением принять участие в запросе цен '
        'и принимая его условия, нижеподписавшийся: '
        '_______________________________________________________________________________________ '
        '(наименование и организационно-правовая форма участника закупки), '
        'действующий при наличии соответствующих полномочий (доверенность, устав, иное), '
        'предлагает поставить товар в полном соответствии с условиями, '
        'указанными в запросе цен и согласно следующим условиям.'
    )

    _add(doc, '1. Сведения об участнике', bold=True, space_before=6, space_after=2)

    # Сведения table
    tbl = doc.add_table(rows=9, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ('Наименование (для юр. лиц)', ''),
        ('Сведения об учредительных документах (наименование и реквизиты)', ''),
        ('Фирменное наименование (для юр. лиц)', ''),
        ('ИНН (для юр. лиц)', ''),
        ('Почтовый адрес (для юр. лиц)', ''),
        ('Фактический адрес (для юр. лиц)', ''),
        ('Телефон (для справок)', ''),
        ('Контактное лицо (для справок)', ''),
        ('Расчётный счёт', ''),
    ]
    for ri, (label, val) in enumerate(rows_data):
        tbl.rows[ri].cells[0].paragraphs[0].add_run(label)
        tbl.rows[ri].cells[1].paragraphs[0].add_run(val)

    _add(doc, '2. Спецификация предложения:', bold=True, space_before=6, space_after=2)

    spec = doc.add_table(rows=2, cols=7)
    spec.style = 'Table Grid'
    hdr = spec.rows[0].cells
    hdr_labels = ['№', 'Наименование', 'Ед. изм.', 'Кол-во', 'Цена за ед., руб.', 'Сумма, руб.', 'Страна']
    for ci, lbl in enumerate(hdr_labels):
        r = hdr[ci].paragraphs[0].add_run(lbl)
        r.bold = True

    # Items loop — rendered as paragraphs, not table rows, per project rules
    _add(doc, '{% for item in items %}', space_before=2, space_after=0)
    _add(doc, '{{ item.num }}. {{ item.name }}, {{ item.quantity }} {{ item.unit }}, '
              '{{ item.unit_price }} руб./ед., итого {{ item.total_price }} руб., страна происхождения: _____',
         space_before=0, space_after=0, indent=0.5)
    _add(doc, '{% endfor %}', space_before=0, space_after=4)

    _add(doc, '2.2. Предложение по ценовым условиям. В соответствии с техническим заданием предлагаем:', space_before=2)

    price_tbl = doc.add_table(rows=3, cols=3)
    price_tbl.style = 'Table Grid'
    ph = price_tbl.rows[0].cells
    ph[0].paragraphs[0].add_run('№').bold = True
    ph[1].paragraphs[0].add_run('Наименование показателя').bold = True
    ph[2].paragraphs[0].add_run('Значение показателя').bold = True
    price_tbl.rows[1].cells[0].paragraphs[0].add_run('1.')
    price_tbl.rows[1].cells[1].paragraphs[0].add_run('Цена за единицу')
    price_tbl.rows[1].cells[2].paragraphs[0].add_run('')
    price_tbl.rows[2].cells[0].paragraphs[0].add_run('2.')
    price_tbl.rows[2].cells[1].paragraphs[0].add_run('Итоговая цена')
    price_tbl.rows[2].cells[2].paragraphs[0].add_run('{{ contract_price }} руб.')

    _add(doc, '3. Подтверждение условий участия (заявитель):', bold=True, space_before=6, space_after=2)
    _add(doc, '_______________________________________________________________________________', space_before=2)
    _add(doc, 'Подтверждаю, что к настоящему моменту установленные заявителем критерии:', space_before=2)
    _add(doc, '3.1. В соответствии с п.7.1. Главы 7 Положения о закупках, в части:')
    _add(doc, '- отсутствие сведений о поставщике, перечень заявителей участника (установленных критериев участника);', indent=0.5)
    _add(doc, '- наличие лицензий на осуществление определённой деятельности;', indent=0.5)
    _add(doc, '3.2. В соответствии с п.7.2. Главы 7 Положения о закупках;', space_before=2)

    _add(doc, '')
    _add(doc, 'Срок поставки составляет ___ (не более чем 30) рабочих дней со дня подписания настоящего договора поставки.', space_before=4)

    _add(doc, '', space_before=8)
    _add(doc, '_________________________________ (ФИО, должность, ЭЦП (при наличии))', space_before=6)

    dst = os.path.join(BASE, 'fabrikant_application_form.docx')
    doc.save(dst)
    print(f'  application_form: created -> {dst}')


def make_documentation():
    """
    fabrikant_documentation.docx — Приложение 1: Документация к запросу цен.
    Информационная карта с переменными полями.
    Variables: customer_full_name, customer_address, subject, total_nmcd,
               subsidy_agreement_number, submission_deadline_datetime,
               applications_review_date, payment_term_days, notice_number,
               vat_info_line, contract_city
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Приложение №1').bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('к запросу цен на право заключения').italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('договора поставки').italic = True

    _add(doc, '')
    _add(doc, 'ДОКУМЕНТАЦИЯ', bold=True, align='center', size=13, space_before=6, space_after=2)
    _add(doc, 'запроса цен на право заключения договора поставки', align='center', space_after=2)
    _add(doc, 'Идентификационный номер закупки «{{ notice_number }}»', align='center', space_after=6)

    _add(doc, '{% if subsidy_agreement_number %}', space_before=0, space_after=0)
    _add(doc, 'Соглашение о предоставлении субсидии: {{ subsidy_agreement_number }}', space_before=0, space_after=0)
    _add(doc, '{% endif %}', space_before=0, space_after=4)

    _add(doc, 'Настоящая документация разработана в целях установления требований к Участникам закупки. '
              'Документация носит характер приглашения, на основании которого лица вправе направить '
              'предложения, и не является офертой в смысле ст. 437 ГК РФ.', space_before=4)

    _add(doc, 'РАЗДЕЛ 1. ОБЩАЯ ИНФОРМАЦИЯ О ЗАКУПКЕ', bold=True, space_before=8, space_after=4)

    # Information card table
    info_rows = [
        ('1.', 'Наименование, организационно-правовая форма, место нахождения, ИНН, КПП\nзаказчика',
         '{{ customer_full_name }}\nИНН: {{ customer_inn }}, КПП: {{ customer_kpp }}'),
        ('2.', 'Способ закупки',
         'Запрос цен в соответствии с Положением о закупках'),
        ('3.', 'Предмет закупки, включая технические характеристики и характеристики',
         '{{ subject }}'),
        ('4.', 'Место доставки (выполнения) работ, оказания услуг, место поставки товаров, '
               'почтовый адрес',
         '{{ customer_address }}'),
        ('5.', 'Дата начала, порядок и сроки исполнения контракта/договора; порядок оказания '
               'услуг, поставки товара',
         'В течение 3-х месяцев с момента подписания договора поставки. '
         '(см. условия договора поставки)'),
        ('6.', 'Начальная (максимальная) цена договора; или цена единицы товара, работы, услуги; '
               'обоснование цены',
         '{{ total_nmcd }} руб. {{ vat_info_line }}'),
        ('7.', 'Форма и место подачи заявок на участие в закупочной процедуре, '
               'номер лота и идентификационный номер лотов',
         'Сайт: www.fabrikant.ru. Лот №{{ notice_number }}'),
        ('8.', 'Обеспечение (обеспечительный) платёж',
         'Не устанавливается'),
        ('9.', 'Срок подачи заявок на участие в закупочной процедуре: дата и время начала '
               'подачи заявок и дата и время окончания подачи заявок',
         'Адрес, указанный Организатором: площадка Fabrikant'),
        ('9.',
         'Срок подачи заявок на участие в закупочной процедуре: дата и время начала '
         'подачи заявок и дата и время окончания подачи заявок',
         '{{ submission_deadline_datetime }}'),
        ('10.', 'Дата рассмотрения заявок',
         '{{ applications_review_date }}'),
        ('11.', 'Дата подведения итогов закупки',
         '{{ applications_review_date }}'),
        ('12.', 'Срок и порядок оплаты',
         'В течение {{ payment_term_days }} ({% if payment_term_days == 10 %}десяти{% else %}{{ payment_term_days }}{% endif %}) '
         'рабочих дней после подписания Акта приёмки-передачи товара'),
        ('13.', 'Место подведения итогов закупки',
         'г. {{ contract_city }}'),
    ]

    tbl = doc.add_table(rows=len(info_rows) + 1, cols=3)
    tbl.style = 'Table Grid'

    # Header row
    hdr = tbl.rows[0].cells
    for ci, lbl in enumerate(['№ п/п', 'Наименование п/п', 'Значение п/п']):
        hdr[ci].paragraphs[0].add_run(lbl).bold = True

    for ri, (num, param, val) in enumerate(info_rows):
        row = tbl.rows[ri + 1]
        row.cells[0].paragraphs[0].add_run(num)
        row.cells[1].paragraphs[0].add_run(param)
        row.cells[2].paragraphs[0].add_run(val)

    _add(doc, '')
    _add(doc, 'РАЗДЕЛ 2. ТРЕБОВАНИЯ К УЧАСТНИКАМ', bold=True, space_before=8, space_after=4)
    _add(doc,
         '1. Заявки принимаются в соответствии с настоящим приглашением и условиями '
         'закупочной процедуры (форма, место) для всех Участников (без установки в '
         'приглашении и форме).')
    _add(doc,
         '2. Заявка, предусмотренная договором, не применяется, согласно нормам 447-449.1, '
         '1057-1061 ГК РФ.')
    _add(doc, '3. Заказчик: {{ customer_full_name }}')
    _add(doc, '4. Организатор закупки: {{ customer_full_name }}')
    _add(doc, '5. Заявки принимаются в форме подачи цен для всех участников '
              '(форма и места) в среде Электронной Торговой Площадки Fabrikant.ru.')
    _add(doc, '6. Участники вправе направить в сроки, определённые организатором '
              'запроса цен. Участник вправе направить только одну заявку. '
              'Участник, направивший заявку на участие в запросе цен до истечения '
              'срока его проведения и истечения срока приёма заявок, вправе отозвать заявку. '
              'Участник, отозвавший заявку на участие в запросе цен до истечения '
              'срока его проведения, вправе подать другую заявку.')
    _add(doc, '7. Заявки принимаются до-время подписания электронного документа (при подаче '
              'заявки до дня рассмотрения, представляется форма настоящей форме формы '
              'Приложение №2 к запросу цен «Форма Заявки». '
              'Необходимо в настоящей форме установить наименование и место подписания в срок. '
              'Размещение в форм на ЭТП Fabrikant.ru')
    _add(doc, '8. Заявки на ЭТП должны быть поданы участниками (в том числе заявки должны '
              'быть направлены на доступных ресурсах, привлечённые на иные ресурсы товаров '
              '(работ, услуг), полученные на официальном платформе), после официального '
              'опубликования официального запроса цен. Форм')
    _add(doc, '9. Подведение до рассмотрения заявок в срок, не позднее 10 (десяти) рабочих '
              'дней, начиная со дня подачи заявок формы Участниками, '
              'организатор направит в направленный адрес участника, '
              'уведомление о проведении, участниках в нём, на настоящую форму — отказе по '
              'настоящей Форме')
    _add(doc, '')
    _add(doc, '______________________', space_before=8)
    _add(doc, '* Требования, указанные в Разделе 2 настоящей Документации являются формой условий раздела 1 Положения.')

    dst = os.path.join(BASE, 'fabrikant_documentation.docx')
    doc.save(dst)
    print(f'  documentation: created -> {dst}')


def make_contract_project():
    """
    fabrikant_contract_project.docx — Проект договора поставки.
    Variables: contract_city, contract_date_day, contract_date_month, contract_date_year,
               customer_full_name, customer_signatory_position, customer_signatory,
               customer_signatory_basis, contractor_full_name, contractor_signatory_position,
               contractor_signatory, contractor_signatory_basis,
               total_nmcd (НМЦД), contract_price, contract_price_words,
               payment_term_days, execution_term, delivery_location,
               customer_inn, customer_kpp, customer_ogrn, customer_address,
               customer_bank_name, customer_settlement_account, customer_bik,
               customer_correspondent_account, customer_signatory_initials,
               contractor_inn, contractor_kpp, contractor_ogrn, contractor_address,
               contractor_bank_name, contractor_settlement_account, contractor_bik,
               contractor_correspondent_account, contractor_signatory_initials,
               items (list), vat_info_line
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Title
    _add(doc, 'ДОГОВОР ПОСТАВКИ в', bold=True, align='center', size=13, space_before=0, space_after=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.add_run('г. {{ contract_city }}')
    p.add_run('\t\t\t\t')
    p.add_run('«{{ contract_date_day }}» {{ contract_date_month }} {{ contract_date_year }} г.')

    _add(doc, '')

    # Parties
    p = doc.add_paragraph()
    p.add_run(
        '{{ customer_full_name }} ({{ customer_inn }} лет), именуемое в дальнейшем «Заказчик», '
        'действующее в соответствии с Законодательством, в лице '
        '{{ customer_signatory_position }} {{ customer_signatory }}, '
        'действующего на основании {{ customer_signatory_basis }}, с одной стороны, и'
    )

    p = doc.add_paragraph()
    p.add_run(
        '{{ contractor_full_name }}, именуемое в дальнейшем «Поставщик», '
        'в лице {{ contractor_signatory_position }} {{ contractor_signatory }}, '
        'действующего на основании {{ contractor_signatory_basis }}, с другой стороны,'
    )

    _add(doc,
         'на основании итогов запроса цен, проведённых на ЭТП Fabrikant (далее по тексту Площадка) '
         'и настоящее соглашение считается заключённым (далее по тексту) в следующем:')

    # Section 1
    _add(doc, '1. Предмет договора', bold=True, space_before=6, space_after=2)
    _add(doc,
         '1.1. Поставщик обязуется поставить и передать в собственность заказчику '
         '(в т.ч. соответствующих дополнительных приложений и условий) Товар в ассортименте, '
         'количестве и по ценам согласно настоящего договора и Спецификации '
         '(Приложение № 1 к договору).', indent=0.5)
    _add(doc,
         '1.2. Маркировка, упаковка, конструкция, разгрузка, транспортировка, установка, '
         'передача на сборку (при необходимости), документация Товаров, '
         'а также иные расходы и условия поставки таких Товаров (Приложение № 1 к договору).', indent=0.5)
    _add(doc, '1.3. Товар, поставляемый в рамках Договора, должен быть новым, '
              'ранее не бывшим в эксплуатации.', indent=0.5)
    _add(doc,
         '1.4. Соответствие (стандартизация) в рамках Договора поставки '
         'в соответствии с аналогом ни с кем из сторонних нарушителей по своему усмотрению, '
         'а также иных условий одной (или несколькими) части возможностей рекламации '
         '(нарушение нормы монопольного статуса). Требования проверки, нарушения '
         'стандартизации, что ранее это коммерческой деятельностью рекламация не '
         'может и не подлежащих по рекламациям: изменение в пользу', indent=0.5)

    # Section 2
    _add(doc, '2. Поставка, транспортировка и передача по факту', bold=True, space_before=6, space_after=2)
    _add(doc,
         '2.1. Весь Товар, поставляемый по Договору (если он не соответствует '
         'и/или не соответствующей таможенной документальной документации по '
         'Спецификации (Приложение № 1 к договору)), должен в Договоре (ст. 479 ГК РФ).',
         indent=0.5)
    _add(doc,
         'Поставщик обязуется до поставки Товара и передачи в собственность продавца '
         'Покупателю в срок поставки «Дата Поставки», указанной в Договоре и Спецификации '
         'поставляемой продукции, установленному разделом Порядка (Приложения).', indent=0.5)
    _add(doc,
         '2.2. Условия и транспортировке (доставку) условия по Договору Поставки '
         'в Спецификации (Приложение № 1 к договору).', indent=0.5)
    _add(doc,
         '2.3. Приёмка поставленного товара.\n'
         '2.3.1. Покупатель, получивший уведомление о поставке в соответствии условий '
         'Поставщика, обязан принять, если получив товар ранее было зафиксировано '
         'предоставление Поставщиком с документальностью, поставленного количества '
         'Товаров, предусмотренного ст. 480 ГК РФ, либо товара товаром, '
         'установленных условиями количественными и ассортиментной документации по '
         'Договору, для проведения выявленных по необходимым условиям '
         'Документации по необходимому факту и необходимой в Договоре.', indent=0.5)

    # Section 3
    _add(doc, '3. Качество товара и гарантийный срок', bold=True, space_before=6, space_after=2)
    _add(doc,
         '3.1. По Договору поставляемый (передаваемый) в соответствии Товар должен '
         'быть надлежащего качества и ранее не бывавшим в эксплуатации.', indent=0.5)
    _add(doc,
         '3.2. Качество товара должно соответствовать технической документации '
         '(в том числе ГОСТам, стандартам), установленной в типовых условиях, а также '
         'соответствующими типами условий и документации, требованиями определённой '
         'условием нашим типом условий по техническим, т.е. не являющимися действующими.', indent=0.5)

    # Section 4
    _add(doc, '4. Цена и порядок оплаты.', bold=True, space_before=6, space_after=2)
    _add(doc,
         '4.1. Общая стоимость Товаров, поставляемых по Договору (цена и иная Стоимость), '
         'составляет {{ contract_price }},00 ({{ contract_price_words }}) рублей 00 копеек, '
         'в том числе НДС в размере {{ vat_info_line }}', indent=0.5)
    _add(doc,
         'Общая стоимость товара является в рамках НДС установленных, включённой в '
         'сумму общего количественного учёта, определённой на Плановый количественный '
         'показатель цен (для таможенного бизнеса, несоответствия (категории), '
         'транспортный ценовой порядка, а также иные нормы и правом, возникающем '
         'применительно к настоящему Договору.', indent=0.5)
    _add(doc, '4.2. Стоимость на 1 (одну) единицу Товара представлена в Спецификации '
              '(Приложение № 1 к договору).', indent=0.5)
    _add(doc,
         '4.3. Оплата производится в размере 100% стоимости в безналичной форме после '
         'принятия Заказчиком Товара на основании подписанного Акта приёма-передачи '
         'в срок {{ payment_term_days }} ({% if payment_term_days == 10 %}десяти{% else %}{{ payment_term_days }}{% endif %}) '
         'рабочих дней с момента подписания Акта приёмки по п. 2.4.1. Договора. '
         'При зачёте суммы оплачиваются пени и штрафные санкции применительно к '
         'настоящей сумме по необходимому платежу.', indent=0.5)
    _add(doc,
         '4.4. Оплата должна осуществляться по основе (формата) договора, '
         'установленного (исполненного) описанием Договора и суммы стоимости, '
         'изложенной в п. 4.1. Договора и выставленным должным платёжным по платежам.', indent=0.5)
    _add(doc,
         '{% if subsidy_agreement_number %}'
         '4.6. Финансирование договора осуществляется в рамках субсидии '
         '№{{ subsidy_agreement_number }}.'
         '{% endif %}', indent=0.5)

    # Section 5
    _add(doc, '5. Место и сроки поставки.', bold=True, space_before=6, space_after=2)
    _add(doc, '5.1. Срок поставки.', bold=False, space_before=2, space_after=1, indent=0.5)
    _add(doc,
         '5.1.1. Поставщик должен поставить товар в полном наборе и форме, '
         'установленной Поставщиком и/или Условиями Поставки (Приложение № 1 к договору).',
         indent=0.5)
    _add(doc,
         '5.1.2. Конечный срок поставки Товара установлен до {{ execution_term }}.',
         indent=0.5)
    _add(doc,
         '5.2. Адрес доставки: {{ delivery_location }}.',
         indent=0.5)

    # Section 6 — Гарантии
    _add(doc, '6. Гарантия товара.', bold=True, space_before=6, space_after=2)
    _add(doc,
         '6.1. Поставщик гарантирует права в части и качестве в форме, '
         'установленной Поставщиком (в т.ч. фирменными условиями).',
         indent=0.5)

    # Section 7 — Ответственность
    _add(doc, '7. Ответственность сторон.', bold=True, space_before=6, space_after=2)
    _add(doc,
         '7.2.1. В случае нарушения Поставщиком сроков поставки обязан уплатить '
         'Заказчику штрафную санкцию (пени) в размере 0,1 % от суммы неисполненного '
         'обязательства в день просрочки.', indent=0.5)
    _add(doc,
         '7.2.2. В случае нарушения Поставщиком условий в части возмещения Покупатель '
         'обязан уплатить штрафную санкцию (пени) в размере 0,1 % от стоимости товара '
         'в неисполненном дне просрочки.', indent=0.5)

    # Section 11 — Реквизиты
    _add(doc, '11. Заключительные положения.', bold=True, space_before=6, space_after=2)
    _add(doc, '11.1. Договор вступает в силу с момента подписания сторонами и действует '
              'до {{ execution_term }} (включительно).', indent=0.5)
    _add(doc, '11.7. Договор составлен в 2 (двух) экземплярах, которые заключают '
              'юридическую силу, по одному — каждой стороне, '
              'при наличии одним заключения юридической силы.', indent=0.5)
    _add(doc, '11.9. Приложения:', space_before=4, space_after=2, indent=0.5)
    _add(doc, '- Приложение № 1 к Спецификации Товара.', indent=1)

    # Section 12 — Реквизиты и подписи
    _add(doc, '12. Адреса и банковские реквизиты сторон.', bold=True, space_before=8, space_after=4)

    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.style = 'Table Grid'
    left = sig_tbl.rows[0].cells[0]
    right = sig_tbl.rows[0].cells[1]

    def _fill_reqs(cell, party, prefix):
        lines = [
            (f'{party}', True),
            (f'{{ {prefix}_full_name }}', False),
            (f'ИНН: {{ {prefix}_inn }}  КПП: {{ {prefix}_kpp }}', False),
            (f'ОГРН: {{ {prefix}_ogrn }}', False),
            (f'Адрес: {{ {prefix}_address }}', False),
            (f'Р/с: {{ {prefix}_settlement_account }}', False),
            (f'Банк: {{ {prefix}_bank_name }}', False),
            (f'БИК: {{ {prefix}_bik }}', False),
            (f'К/с: {{ {prefix}_correspondent_account }}', False),
            ('', False),
            ('_________________ / ', False),
            (f'{{ {prefix}_signatory_initials }}', False),
        ]
        for i, (txt, bold) in enumerate(lines):
            if i == 0:
                p = cell.paragraphs[0]
                r = p.add_run(txt)
                r.bold = bold
            else:
                p = cell.add_paragraph()
                r = p.add_run(txt)
                r.bold = bold

    _fill_reqs(left, 'ЗАКАЗЧИК', 'customer')
    _fill_reqs(right, 'ПОСТАВЩИК', 'contractor')

    # Appendix 1 — Specification
    doc.add_page_break()
    _add(doc, 'Приложение № 1', bold=True, align='center', space_before=4)
    _add(doc, 'к Договору поставки', align='center')
    _add(doc, '№{{ contract_number }} от {{ contract_date }}', align='center', space_after=6)
    _add(doc, 'СПЕЦИФИКАЦИЯ ТОВАРА', bold=True, align='center', size=12)
    _add(doc, 'по Договору поставки', align='center', space_after=4)

    spec_tbl = doc.add_table(rows=2, cols=7)
    spec_tbl.style = 'Table Grid'
    hdr_cells = spec_tbl.rows[0].cells
    for ci, lbl in enumerate(['№', 'Наименование товара', 'Ед. изм.', 'Кол-во', 'Цена за ед., руб.', 'Сумма, руб.', 'Страна происхождения']):
        r = hdr_cells[ci].paragraphs[0].add_run(lbl)
        r.bold = True
    # Number row
    for ci, lbl in enumerate(['1', '2', '3', '4', '5', '6', '7']):
        spec_tbl.rows[1].cells[ci].paragraphs[0].add_run(lbl)

    _add(doc, '{% for item in items %}', space_before=2, space_after=0)
    _add(doc,
         '{{ item.num }}. {{ item.name }}, кол-во: {{ item.quantity }} {{ item.unit }}, '
         '{{ item.unit_price }} руб./ед., итого: {{ item.total_price }} руб.',
         indent=0.5, space_before=0, space_after=0)
    _add(doc, '{% endfor %}', space_before=0, space_after=4)

    _add(doc, 'Итого: {{ contract_price }} руб. ({{ contract_price_words }})', bold=True, space_before=4)
    _add(doc, '{{ vat_info_line }}', space_before=2)
    _add(doc, 'Адрес поставки: {{ delivery_location }}', space_before=2)
    _add(doc, 'Срок поставки: {{ execution_term }}', space_before=2)
    _add(doc, '')

    sig_tbl2 = doc.add_table(rows=1, cols=2)
    sig_tbl2.style = 'Table Grid'
    left2 = sig_tbl2.rows[0].cells[0]
    right2 = sig_tbl2.rows[0].cells[1]
    left2.paragraphs[0].add_run('ЗАКАЗЧИК').bold = True
    left2.add_paragraph('{{ customer_signatory_position }}')
    left2.add_paragraph('')
    left2.add_paragraph('_________________ / {{ customer_signatory_initials }}')
    right2.paragraphs[0].add_run('ПОСТАВЩИК').bold = True
    right2.add_paragraph('{{ contractor_signatory_position }}')
    right2.add_paragraph('')
    right2.add_paragraph('_________________ / {{ contractor_signatory_initials }}')

    dst = os.path.join(BASE, 'fabrikant_contract_project.docx')
    doc.save(dst)
    print(f'  contract_project: created -> {dst}')


if __name__ == '__main__':
    print('Generating Fabrikant docxtpl templates...')
    print(f'  Output dir: {BASE}')
    print(f'  Source dir: {SRC_DIR}')
    make_instruction()
    make_application_form()
    make_documentation()
    make_contract_project()
    print('Done. All 4 templates created.')
