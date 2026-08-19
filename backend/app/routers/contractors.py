from urllib.parse import quote as _url_quote
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Body, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select, func, distinct, or_
from sqlalchemy.ext.asyncio import AsyncSession
from app.database import get_db
from app.models.contractor import Contractor
from app.schemas.schemas import ContractorCreate, ContractorOut
from app.auth.jwt import require_role, get_current_user, get_org_filter, get_single_org_id, ADMIN_ROLES, MANAGER_ROLES, ALL_ROLES
from app.auth.permissions import require_tab
from app.models.user import User
from app.services.fio import compose_fio, split_position_and_fio
from typing import List, Optional
from io import BytesIO

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    import pdfplumber as _pdfplumber
except ImportError:
    _pdfplumber = None


def _ocr_pdf_to_rows(content: bytes) -> tuple[list, str | None]:
    """Fallback: convert scanned PDF pages to images, run OCR, parse lines.
    Returns (rows, error_message). error_message is None on success."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError:
        return [], "OCR-библиотеки не установлены (pdf2image, pytesseract). Обратитесь к администратору."
    try:
        images = convert_from_bytes(content, dpi=300)
    except Exception as e:
        return [], f"Не удалось преобразовать PDF в изображения для OCR: {e}"
    if not images:
        return [], "PDF не содержит страниц для распознавания."
    import re
    all_rows = []
    for img in images:
        try:
            text = pytesseract.image_to_string(img, lang='rus+eng')
        except Exception as e:
            return [], f"Ошибка OCR-распознавания: {e}"
        if not text:
            continue
        for line in text.strip().split('\n'):
            line = line.strip()
            if not line:
                continue
            cells = re.split(r'\t|  {2,}', line)
            cells = [c.strip() for c in cells if c.strip()]
            if cells:
                all_rows.append(cells)
    if not all_rows:
        return [], "OCR не нашёл текст на страницах. Возможно, качество скана слишком низкое."
    return all_rows, None


router = APIRouter(prefix="/api/contractors", tags=["contractors"])


# ---------------------------------------------------------------------------
# Shared file parsing helpers
# ---------------------------------------------------------------------------

_CONTRACTOR_HINTS = (
    'назван', 'наимен', 'инн', 'inn', 'кпп', 'огрн', 'адрес',
    'email', 'телефон', 'банк', 'бик', 'контакт', 'подписант',
)


def _detect_hdr(rows):
    """Return index of the row most likely to be a header row."""
    best_score, best_idx = 0, 0
    for ri, row in enumerate(rows):
        norm = [str(h).strip().lower() if h is not None else "" for h in row]
        score = sum(1 for h in norm if h and any(x in h for x in _CONTRACTOR_HINTS))
        if score > best_score:
            best_score = score
            best_idx = ri
    return best_idx


def _map_kv_key(key: str) -> str | None:
    """Map a key-value card label to a contractor field name.
    Returns the field name string or None if the key is not recognised."""
    k = key.strip().lower()
    if not k:
        return None
    # Order matters — more specific checks first
    if 'инн' in k or 'inn' in k:
        return 'инн'
    if 'кпп' in k or 'kpp' in k:
        return 'кпп'
    if 'огрн' in k:
        return 'огрн'
    if 'бик' in k or 'bik' in k:
        return 'бик'
    if 'расч' in k or 'р/с' in k:
        return 'расчётный счёт'
    if 'корр' in k or 'к/с' in k:
        return 'корр. счёт'
    if 'банк' in k and 'реквиз' not in k:
        return 'банк'
    if 'назван' in k or 'наимен' in k or 'учредит' in k:
        return 'наименование'
    if 'юридич' in k and 'адрес' in k:
        return 'адрес'
    if 'почтов' in k:
        return 'почтовый адрес'
    if 'телефон' in k or (k.startswith('тел') and len(k) < 10):
        return 'телефон'
    if 'email' in k or 'e-mail' in k:
        return 'email'
    if 'подписант' in k or 'уполномоч' in k or 'представит' in k:
        return 'подписант'
    if 'должност' in k:
        return 'должность'
    if 'сайт' in k or 'website' in k:
        return 'сайт'
    if 'окпо' in k:
        return 'окпо'
    if 'оквэд' in k:
        return 'оквэд'
    if 'единый' in k and 'казначейск' in k:
        return 'единый казначейский счёт'
    if 'казначейск' in k:
        return 'казначейский счёт'
    return None


# Mapping from Russian card labels to ContractorCreate field names
_KV_LABEL_TO_FIELD: dict[str, str] = {
    'инн': 'inn',
    'кпп': 'kpp',
    'огрн': 'ogrn',
    'бик': 'bik',
    'расчётный счёт': 'settlement_account',
    'корр. счёт': 'correspondent_account',
    'банк': 'bank_name',
    'наименование': 'name',
    'адрес': 'address',
    'почтовый адрес': 'postal_address',
    'телефон': 'phone',
    'email': 'email',
    'подписант': 'signatory',
    'должность': 'signatory_position',
    'сайт': 'website',
    'окпо': 'okpo',
    'оквэд': 'okved',
    'казначейский счёт': 'treasury_account',
    'единый казначейский счёт': 'single_treasury_account',
}

# Human-readable column header for each field used in preview response
_FIELD_TO_HEADER: dict[str, str] = {
    'inn': 'ИНН',
    'kpp': 'КПП',
    'ogrn': 'ОГРН',
    'bik': 'БИК',
    'settlement_account': 'Расчётный счёт',
    'correspondent_account': 'Корр. счёт',
    'bank_name': 'Банк',
    'name': 'Наименование',
    'address': 'Адрес',
    'postal_address': 'Почтовый адрес',
    'phone': 'Телефон',
    'email': 'Email',
    'signatory': 'Подписант',
    'website': 'Сайт',
    'registration_date': 'Дата регистрации',
    'okpo': 'ОКПО',
    'okved': 'ОКВЭД',
    'treasury_account': 'Казначейский счёт',
    'single_treasury_account': 'Единый казначейский счёт',
    'signatory_position': 'Должность',
}


def _try_parse_kv_docx(doc) -> tuple[list, int] | None:
    """Try to parse a DOCX document as a key-value requisites card.

    Returns (all_rows, hdr_idx) where all_rows = [header_row, value_row]
    with hdr_idx = 0, ready to be returned from _parse_file_to_rows.

    Returns None if the document does not look like a key-value card
    (e.g. it is a proper table with data in multiple data rows).
    """
    # Collect all key→value pairs from every table in the document
    kv_pairs: list[tuple[str, str]] = []
    two_col_count = 0
    total_rows = 0

    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            # Remove duplicate adjacent cells (merged cells repeat in python-docx)
            deduped = []
            for c in cells:
                if not deduped or c != deduped[-1]:
                    deduped.append(c)
            non_empty = [c for c in deduped if c]
            total_rows += 1
            if len(non_empty) == 2:
                two_col_count += 1
                kv_pairs.append((non_empty[0], non_empty[1]))
            elif len(non_empty) == 1:
                # Single-cell rows are OK (section headers) — do not disqualify
                pass
            # Rows with 3+ distinct non-empty cells suggest a real data table
            elif len(non_empty) >= 3:
                return None  # Looks like a multi-column table — not a card

    if total_rows == 0:
        return None

    # If fewer than half the rows are 2-column key-value rows, it's not a card
    if two_col_count < max(2, total_rows * 0.4):
        return None

    # Also check that the keys look like field labels (not values)
    recognised = sum(1 for k, _ in kv_pairs if _map_kv_key(k) is not None)
    if recognised < 2:
        return None  # Too few recognisable field labels — not a requisites card

    # Build a synthetic header row + single data row
    # Use only the first match for each field to avoid duplicates
    headers: list[str] = []
    values: list[str] = []
    seen_fields: set[str] = set()

    for raw_key, raw_val in kv_pairs:
        label = _map_kv_key(raw_key)
        if label is None:
            continue
        field = _KV_LABEL_TO_FIELD.get(label)
        if field is None:
            continue
        if field in seen_fields:
            continue  # Take first match only
        seen_fields.add(field)
        headers.append(_FIELD_TO_HEADER.get(field, field))
        values.append(raw_val)

    # Post-process: if ОГРН value contains a dd.mm.yyyy date (e.g. "1239500010639, 11.12.2023"),
    # split it — keep digits as ОГРН and extract registration_date if not already captured.
    try:
        import re as _re_kv
        ogrn_header = _FIELD_TO_HEADER.get('ogrn', 'ОГРН')
        if ogrn_header in headers:
            idx = headers.index(ogrn_header)
            ogrn_val = values[idx]
            dm = _re_kv.search(r'(\d{2}\.\d{2}\.\d{4})', ogrn_val)
            if dm:
                # Keep only digits for ОГРН
                values[idx] = _re_kv.sub(r'\D', '', ogrn_val.split(',')[0])
                if 'registration_date' not in seen_fields:
                    headers.append(_FIELD_TO_HEADER.get('registration_date', 'Дата регистрации'))
                    values.append(dm.group(1))
                    seen_fields.add('registration_date')
    except Exception:
        pass  # defensive — don't break existing parsing

    # Fallback: if name not found, scan paragraphs
    if 'name' not in seen_fields:
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) > 5:
                headers.insert(0, _FIELD_TO_HEADER['name'])
                values.insert(0, text)
                break

    if not headers:
        return None

    all_rows = [headers, values]
    return all_rows, 0  # hdr_idx = 0


def _try_parse_paragraphs_docx(doc) -> tuple[list, int] | None:
    """Parse plain-paragraph requisites card (no tables).

    Scans paragraph text via regex patterns for ИНН, КПП, ОГРН, БИК,
    р/с, к/с, name, address, email, phone, signatory.
    Returns (all_rows, 0) — header row + value row — or None if too few fields found.
    """
    import re as _re
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return None
    # Also scan paragraphs inside tables (some docs use 1-cell layout-tables)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        paragraphs.append(t)
    full_text = "\n".join(paragraphs)

    fields: dict[str, str] = {}

    def _digits(s: str) -> str:
        return _re.sub(r'\D', '', s)

    # Compact "ИНН/КПП:" combo (e.g. "ИНН/КПП: 7731178803/772901001")
    m = _re.search(r'ИНН\s*/?\s*КПП\s*[:№]?\s*(\d{10,12})\s*/\s*(\d{9})', full_text, _re.IGNORECASE)
    if m:
        fields['inn'] = m.group(1)
        fields['kpp'] = m.group(2)

    # Individual labelled fields
    PATTERNS = [
        ('inn',    r'ИНН[\s:№]+(\d{10,12})\b'),
        ('kpp',    r'КПП[\s:№]+(\d{9})\b'),
        ('ogrn',   r'ОГРН(?:ИП)?[\s:№]+(\d{13,15})\b'),
        ('bik',    r'БИК[\s:№]+(\d{9})\b'),
    ]
    for field, pat in PATTERNS:
        if field in fields:
            continue
        m = _re.search(pat, full_text, _re.IGNORECASE)
        if m:
            fields[field] = m.group(1)

    # Bank accounts: collect ALL 20-digit numbers, classify by prefix.
    # 30101... → correspondent (к/с), 4xxxxx → settlement (р/с).
    # This is more robust than label-matching ("р/счёт", "р/с", "расч/сч." etc).
    accounts_seen: set[str] = set()
    for m in _re.finditer(r'\b(\d{20})\b', full_text):
        acc = m.group(1)
        if acc in accounts_seen:
            continue
        accounts_seen.add(acc)
        if acc.startswith('30101') and 'correspondent_account' not in fields:
            fields['correspondent_account'] = acc
        elif acc.startswith('4') and 'settlement_account' not in fields:
            fields['settlement_account'] = acc

    # Name — first quoted phrase or first paragraph if it looks like an org name
    m = _re.search(r'((?:ООО|ОАО|ЗАО|ПАО|АО|ИП|НКО|ОООО|АНО)\s*["«»]?[^"\n«»]+["»]?)', full_text)
    if m:
        fields['name'] = m.group(1).strip().strip(',').strip()
    elif paragraphs:
        # First non-trivial paragraph as name
        for p in paragraphs:
            if len(p) >= 5 and not _re.fullmatch(r'[\d\s\-+\(\)]+', p):
                fields['name'] = p.strip()
                break

    # Email
    m = _re.search(r'\b([\w.+-]+@[\w-]+\.[\w.-]+)\b', full_text)
    if m:
        fields['email'] = m.group(1)

    # Phone
    m = _re.search(r'(\+?7?\s*\(?\d{3,4}\)?[\s-]?\d{3}[\s-]?\d{2}[\s-]?\d{2})', full_text)
    if m:
        fields['phone'] = m.group(1).strip()

    # Address — line containing "г. " or 6-digit postal code + comma
    for p in paragraphs:
        clean = _re.sub(r'^(?:Юр\.?\s*адрес|Адрес|Юридический\s+адрес|Фактический\s+адрес|Местонахождение)[\s:]+', '', p, flags=_re.IGNORECASE)
        if _re.search(r'\b\d{6}\b', clean) or _re.search(r'\bг\.\s', clean):
            if 10 < len(clean) < 250 and 'address' not in fields:
                fields['address'] = clean.strip()
                break

    # Bank name — line starting with "Банк:" / "в банке" etc, or following BIK
    m = _re.search(r'(?:Банк|в\s+банке)[\s:]+([^\n]+?)(?=\s*БИК|\s*к/?с|\s*$|\n)', full_text, _re.IGNORECASE)
    if m:
        fields['bank_name'] = m.group(1).strip().strip(',').strip()

    # Signatory — "Подписант: ..." or "Генеральный директор ФИО"
    m = _re.search(r'(?:Подписант|Уполномоченное\s+лицо|Представитель)[\s:]+([^\n]+)', full_text, _re.IGNORECASE)
    if m:
        fields['signatory'] = m.group(1).strip()
    else:
        m = _re.search(r'((?:Генеральный\s+директор|Директор|Председатель|Президент|Главный\s+бухгалтер|Управляющий)\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.?\s*[А-ЯЁ]\.?)', full_text)
        if m:
            fields['signatory'] = m.group(1).strip()

    if len(fields) < 2:
        return None  # too little extracted — not a requisites card

    # Build header/value rows in stable order
    order = ['name', 'inn', 'kpp', 'ogrn', 'address', 'phone', 'email',
             'settlement_account', 'correspondent_account', 'bik', 'bank_name', 'signatory',
             'okpo', 'okved', 'treasury_account', 'single_treasury_account',
             'signatory_position', 'website', 'registration_date']
    headers = [_FIELD_TO_HEADER.get(f, f) for f in order if f in fields]
    values = [fields[f] for f in order if f in fields]
    return [headers, values], 0


def _parse_file_to_rows(fname: str, content: bytes):
    """Parse xlsx/xls/docx/doc/pdf and return (all_rows, hdr_idx)."""
    fname = fname.lower()

    if fname.endswith('.xls') and not fname.endswith('.xlsx'):
        try:
            import xlrd as _xlrd_mod
        except ImportError:
            raise HTTPException(500, "xlrd не установлен")
        try:
            wb_xls = _xlrd_mod.open_workbook(file_contents=content)
        except Exception as e:
            raise HTTPException(400, f"Не удалось прочитать .xls файл: {e}")
        ws_xls = wb_xls.sheet_by_index(0)
        all_rows = [list(ws_xls.row_values(i)) for i in range(ws_xls.nrows)]

    elif fname.endswith('.xlsx'):
        if not load_workbook:
            raise HTTPException(500, "openpyxl не установлен")
        try:
            wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
        except Exception as e:
            raise HTTPException(400, f"Не удалось прочитать .xlsx файл: {e}")
        ws = wb.active
        all_rows = list(ws.iter_rows(values_only=True))
        wb.close()

    elif fname.endswith(('.docx', '.doc')):
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(500, "python-docx не установлен")
        try:
            doc = Document(BytesIO(content))
        except Exception as e:
            raise HTTPException(400, f"Не удалось прочитать .docx файл: {e}")

        # Try key-value card format from tables (карточка реквизитов в таблице)
        if doc.tables:
            kv_result = _try_parse_kv_docx(doc)
            if kv_result is not None:
                return kv_result

        # Fallback: parse paragraphs via regex (карточка-абзацы без таблиц)
        para_result = _try_parse_paragraphs_docx(doc)
        if para_result is not None:
            return para_result

        if not doc.tables:
            raise HTTPException(
                400,
                "В документе .docx не найдено ни таблиц, ни узнаваемых полей "
                "(ИНН/КПП/ОГРН/название). Проверьте что документ содержит реквизиты."
            )

        tbl = doc.tables[0]
        all_rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]

    elif fname.endswith('.pdf'):
        if not _pdfplumber:
            raise HTTPException(500, "pdfplumber не установлен")
        has_text = False
        try:
            with _pdfplumber.open(BytesIO(content)) as pdf:
                all_rows = []
                for page in pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        for tbl in tables:
                            all_rows.extend([r for r in tbl if r])
                if not all_rows:
                    for page in _pdfplumber.open(BytesIO(content)).pages:
                        text = page.extract_text()
                        if text:
                            has_text = True
                            for line in text.strip().split('\n'):
                                cells = [c.strip() for c in line.split('\t')]
                                if len(cells) < 2:
                                    cells = [c.strip() for c in line.split('  ') if c.strip()]
                                if cells:
                                    all_rows.append(cells)
        except Exception as e:
            raise HTTPException(400, f"Не удалось прочитать PDF-файл: {e}")
        if not all_rows:
            if not has_text:
                # Scanned PDF — try OCR
                all_rows, ocr_error = _ocr_pdf_to_rows(content)
                if not all_rows:
                    detail = ocr_error or "OCR не смог распознать таблицу."
                    raise HTTPException(
                        400,
                        f"Этот PDF — скан (изображение). {detail} "
                        "Попробуйте сохранить данные в Excel (.xlsx) или Word (.docx)."
                    )
            else:
                raise HTTPException(
                    400,
                    "В PDF найден текст, но не удалось распознать таблицу. "
                    "Попробуйте сохранить данные в Excel (.xlsx) и загрузить его."
                )

    else:
        raise HTTPException(400, "Неподдерживаемый формат файла. Используйте .xlsx, .xls, .docx, .doc или .pdf")

    if not all_rows:
        raise HTTPException(400, "Файл пустой или не содержит данных")

    hdr_idx = _detect_hdr(all_rows)
    return all_rows, hdr_idx


@router.get("/product-categories")
async def list_all_product_categories(
    db: AsyncSession = Depends(get_db),
    _: User = Depends(get_current_user),
):
    """All unique product categories from products + manual contractor categories."""
    from app.models.product import Product
    # From products table
    prod_res = await db.execute(
        select(distinct(Product.category))
        .where(Product.category.isnot(None), Product.category != '')
    )
    cats = {r[0] for r in prod_res}
    # From manual contractor categories
    ctr_res = await db.execute(
        select(Contractor.manual_product_categories)
        .where(Contractor.manual_product_categories.isnot(None))
    )
    for row in ctr_res:
        for c in (row[0] or []):
            if c and c != 'Все':
                cats.add(c)
    return sorted(cats)


@router.get("/with-stats")
async def list_contractors_with_stats(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
    search: str = Query(None),
    offset: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
    category: str = Query(None),
):
    """Contractors with product categories, server-side pagination + search."""
    from app.models.purchase import Purchase
    from app.models.purchase_item import PurchaseItem
    from app.models.product import Product

    q = select(Contractor).order_by(Contractor.name)
    # Контрагенты — общий справочник юрлиц (см. фикс в list_contractors:1071).
    # Org-фильтр снят: ContractorsView и autocomplete показывают всех.
    if search:
        term = f"%{search}%"
        q = q.where(or_(Contractor.name.ilike(term), Contractor.inn.ilike(term)))
    if category:
        # Filter by manual_product_categories JSONB contains, or "Все"
        from sqlalchemy.dialects.postgresql import JSONB
        from sqlalchemy import cast, literal
        q = q.where(or_(
            Contractor.manual_product_categories.op('?')(category),
            Contractor.manual_product_categories.op('?')('Все'),
        ))

    # Total count for pagination
    from sqlalchemy import func as safunc
    count_q = select(safunc.count()).select_from(q.subquery())
    total = (await db.execute(count_q)).scalar() or 0

    contractors = (await db.execute(q.offset(offset).limit(limit))).scalars().all()
    c_ids = [c.id for c in contractors]

    # Product categories per contractor (only for current page)
    prod_cat_map = {}
    if c_ids:
        prod_stmt = (
            select(
                Purchase.contractor_id,
                func.array_agg(distinct(Product.category)).label("product_categories"),
            )
            .join(PurchaseItem, PurchaseItem.purchase_id == Purchase.id)
            .join(Product, Product.id == PurchaseItem.product_id)
            .where(Purchase.contractor_id.in_(c_ids))
            .where(Product.category.isnot(None))
            .where(Product.category != '')
            .group_by(Purchase.contractor_id)
        )
        prod_rows = (await db.execute(prod_stmt)).all()
        prod_cat_map = {
            row.contractor_id: [x for x in (row.product_categories or []) if x]
            for row in prod_rows
        }

    result = []
    for c in contractors:
        c_dict = ContractorOut.model_validate(c).model_dump()
        manual = c.manual_product_categories or []
        auto = prod_cat_map.get(c.id, [])
        if "Все" in manual:
            c_dict["product_categories"] = ["Все"]
        else:
            merged = list(dict.fromkeys(manual + auto))
            c_dict["product_categories"] = merged
        result.append(c_dict)
    return {"items": result, "total": total}


def _split_signatory(raw, position=None):
    """Тонкая обёртка над split_position_and_fio из fio.py.

    Сохраняет прежнюю сигнатуру и возврат (signatory_fio, signatory_position)
    для обратной совместимости со всеми вызывающими местами.
    """
    from app.services.fio import split_position_and_fio, compose_fio
    last, first, middle, pos = split_position_and_fio(raw, position)
    fio = compose_fio(last, first, middle) or raw
    return (fio, pos)


async def _check_npd_status(inn: str) -> dict:
    """Статус плательщика НПД (самозанятого) в реестре ФНС.

    ЕГРЮЛ/ЕГРИП самозанятых НЕ содержит — по такому ИНН egrul.nalog.ru отдаёт
    пустой rows, хотя человек реально работает и его находит проверка на
    npd.nalog.ru. Это единственный публичный источник по НПД, и он отдаёт
    только факт статуса — ни ФИО, ни адреса там нет.

    Возвращает {'state': 'yes' | 'no' | 'invalid' | 'unknown', 'message': str}.
    'invalid' — ИНН не проходит проверку контрольной цифры (ФНС: validation.failed).
    'unknown' — сервис недоступен либо упёрлись в его лимит запросов с одного IP
    (ФНС отдаёт taxpayer.status.service.limited.error); отличать от 'no' важно,
    иначе пользователю соврём, что человек не самозанятый.
    """
    import httpx
    import logging
    from datetime import date as _date

    logger = logging.getLogger(__name__)
    try:
        async with httpx.AsyncClient(timeout=10, verify=False) as client:
            resp = await client.post(
                "https://statusnpd.nalog.ru/api/v1/tracker/taxpayer_status",
                json={"inn": inn, "requestDate": _date.today().isoformat()},
            )
            data = resp.json()
    except Exception as e:
        logger.warning("NPD status check failed for INN %s: %s", inn, e)
        return {"state": "unknown", "message": "сервис проверки самозанятых ФНС не ответил"}

    if data.get("code"):
        logger.warning("NPD status refused for INN %s: %s", inn, data)
        _msg = data.get("message") or str(data.get("code"))
        if data.get("code") == "validation.failed":
            return {"state": "invalid", "message": _msg}
        return {"state": "unknown", "message": _msg}
    if data.get("status") is True:
        return {"state": "yes", "message": data.get("message") or ""}
    return {"state": "no", "message": data.get("message") or ""}


@router.get("/lookup-inn/{inn}")
async def lookup_inn(
    inn: str,
    force_egrul: bool = Query(False),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Lookup company data by INN: first local DB (unless force_egrul=1), then FNS EGRUL/EGRIP API."""
    import httpx
    import logging
    import re as _re_lookup
    logger = logging.getLogger(__name__)

    inn = inn.strip()
    if not inn or len(inn) not in (10, 12):
        raise HTTPException(400, "ИНН должен быть 10 (юр.лицо) или 12 (ИП) цифр")

    # Step 0: check local DB first (skip if force_egrul requested)
    if not force_egrul:
        local = await db.execute(
            select(Contractor).where(Contractor.inn == inn).limit(1)
        )
        local_contractor = local.scalar_one_or_none()
    else:
        local_contractor = None
    if local_contractor:
        # Defensive filter: if phone field contains the INN (data entry error) — suppress it
        raw_phone = local_contractor.phone
        _lc_phone = (raw_phone or '').lower().strip()
        if raw_phone and (
            _lc_phone == inn
            or _lc_phone.startswith('инн')
            or _re_lookup.sub(r'\D', '', raw_phone) == inn
        ):
            raw_phone = None
            logger.warning("lookup_inn: phone field contained INN for contractor %s, suppressed", inn)
        _lc_sig, _lc_pos = _split_signatory(
            local_contractor.signatory, getattr(local_contractor, "signatory_position", None)
        )
        return {
            "id": local_contractor.id,
            "name": local_contractor.name,
            "inn": local_contractor.inn,
            "kpp": local_contractor.kpp,
            "ogrn": local_contractor.ogrn,
            "address": local_contractor.address,
            "postal_address": local_contractor.postal_address,
            "org_type": local_contractor.org_type,
            "signatory": _lc_sig,
            "signatory_position": _lc_pos,
            "signatory_last_name": local_contractor.signatory_last_name,
            "signatory_first_name": local_contractor.signatory_first_name,
            "signatory_middle_name": local_contractor.signatory_middle_name,
            "signatory_basis": local_contractor.signatory_basis,
            "contact_person": local_contractor.contact_person,
            "phone": raw_phone,
            "email": local_contractor.email,
            "org_phone": local_contractor.org_phone,
            "org_email": local_contractor.org_email,
            "settlement_account": local_contractor.settlement_account,
            "bank_name": local_contractor.bank_name,
            "bik": local_contractor.bik,
            "correspondent_account": local_contractor.correspondent_account,
            "bank_details": local_contractor.bank_details,
            "_source": "local",
        }

    # FNS public API (no auth required)
    url = f"https://egrul.nalog.ru/search-result/{inn}"
    search_url = "https://egrul.nalog.ru/"

    try:
        async with httpx.AsyncClient(timeout=15, verify=False) as client:
            # Step 1: initiate search
            resp1 = await client.post(search_url, json={"query": inn, "region": "", "page": ""})
            token = resp1.json().get("t")
            if not token:
                raise HTTPException(502, "ФНС не вернула токен поиска")

            # Step 2: get results (may need retry)
            import asyncio
            for attempt in range(5):
                await asyncio.sleep(1)
                resp2 = await client.get(f"https://egrul.nalog.ru/search-result/{token}")
                data = resp2.json()
                rows = data.get("rows", [])
                if rows:
                    break

            if not rows:
                # Самозанятых (плательщиков НПД) в ЕГРЮЛ/ЕГРИП нет вообще —
                # прежде чем сказать «не найден», спрашиваем реестр НПД.
                npd = await _check_npd_status(inn) if len(inn) == 12 else {"state": "no", "message": ""}
                if npd["state"] == "yes":
                    return {
                        "inn": inn,
                        "org_type": "Самозанятый",
                        "status": npd["message"],
                        "_source": "npd",
                        "_notice": (
                            f"ИНН {inn} — самозанятый (плательщик налога на профессиональный доход). "
                            "В ЕГРЮЛ/ЕГРИП таких записей нет, поэтому ФИО, адрес и банковские "
                            "реквизиты придётся заполнить вручную."
                        ),
                    }
                if npd["state"] == "invalid":
                    raise HTTPException(
                        status_code=404,
                        detail={
                            "code": "INN_NOT_FOUND",
                            "message": (
                                f"ИНН {inn} некорректен — не проходит проверку контрольной цифры ФНС. "
                                "Скорее всего, в номере опечатка."
                            ),
                            "hint": "Сверьте ИНН с документом контрагента.",
                        },
                    )
                if npd["state"] == "unknown":
                    raise HTTPException(
                        status_code=404,
                        detail={
                            "code": "INN_NOT_FOUND",
                            "message": (
                                f"ИНН {inn} не найден в ЕГРЮЛ/ЕГРИП. Проверить, не самозанятый ли это, "
                                f"сейчас не получилось: {npd['message']}."
                            ),
                            "hint": "Повторите попытку через минуту либо заполните карточку вручную.",
                        },
                    )
                raise HTTPException(
                    status_code=404,
                    detail={
                        "code": "INN_NOT_FOUND",
                        "message": (
                            f"ИНН {inn} не найден: его нет ни в ЕГРЮЛ/ЕГРИП, ни в реестре самозанятых. "
                            "Проверьте правильность ввода."
                        ),
                        "hint": "Если организация существует, попробуйте найти её по названию или КПП.",
                    },
                )

            row = rows[0]  # first match
            # ЕГРЮЛ возвращает руководителя как "ДОЛЖНОСТЬ: ФИО"
            # (например "ПРЕДСЕДАТЕЛЬ: Девлишева Максим Махмович").
            _signatory, _signatory_position = _split_signatory(row.get("g"))
            _eg_last, _eg_first, _eg_middle, _ = split_position_and_fio(
                row.get("g"), _signatory_position
            )
            result = {
                "name": row.get("c") or row.get("n"),  # c=short name, n=full name
                "full_name": row.get("n"),              # n=full legal name
                "inn": row.get("i"),
                "ogrn": row.get("o"),
                "kpp": row.get("p"),
                "address": row.get("a"),
                # Доработка 5 мая: НЕ выставляем "Юр.лицо" по умолчанию.
                # ФНС возвращает запись и для ИП (12 цифр ИНН) и для ЮЛ (10 цифр).
                # Для 10-значного ИНН однозначно подразумевается ЮЛ → "Юр.лицо";
                # для 12-значного определяем по статусу/типу выгрузки. Если в выгрузке
                # FNS поля o/p отсутствуют (для физлиц) → не выставляем тип, пусть
                # пользователь сам выберет «Самозанятый/Физ.лицо/ИП».
                "org_type": (
                    "ИП" if len(inn) == 12 else (
                        "Юр.лицо" if row.get("o") else None
                    )
                ),
                "signatory": _signatory,  # director/head ФИО (без должности)
                "signatory_position": _signatory_position,  # должность подписанта
                "signatory_last_name": _eg_last,
                "signatory_first_name": _eg_first,
                "signatory_middle_name": _eg_middle,
                "status": row.get("s"),  # status text
                "registration_date": row.get("r"),
            }
            return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error("FNS lookup error for INN %s: %s", inn, e)
        raise HTTPException(502, f"Ошибка запроса к ФНС: {str(e)[:200]}")


@router.post("/enrich-all-from-egrul")
async def enrich_all_contractors_from_egrul(
    limit: int = Query(50, ge=1, le=500, description="Сколько контрагентов обработать за один запуск"),
    db: AsyncSession = Depends(get_db),
    _user=Depends(get_current_user),
):
    """Phase 26-CCC: bulk backfill.

    Для всех Contractor где name длиннее 80 символов И full_name пустое
    И ИНН задан → попытка ЕГРЮЛ enrich (короткое name в Contractor.name,
    длинное в Contractor.full_name + ОГРН/КПП/адрес/форма/подписант).

    Идемпотентно: не перезаписывает уже заполненные поля.
    Запускать batch'ами через ?limit=N (max 500). Возвращает stats.
    """
    from app.routers.purchase_receipts import _create_or_enrich_contractor_from_receipt

    rows = (await db.execute(
        select(Contractor)
        .where(func.length(Contractor.name) > 80)
        .where(Contractor.full_name.is_(None))
        .where(Contractor.inn.is_not(None))
        .limit(limit)
    )).scalars().all()

    enriched = 0
    skipped = 0
    failed = 0
    for c in rows:
        try:
            new_c = await _create_or_enrich_contractor_from_receipt(c.inn, c.name, db)
            # Применяем только если ЕГРЮЛ реально дал более короткое имя
            if new_c.name and new_c.name != c.name and len(new_c.name) < len(c.name):
                # Сохраняем старое длинное в full_name если оно ещё пустое
                if not c.full_name:
                    c.full_name = c.name
                c.name = new_c.name
                # Если ЕГРЮЛ вернул более точное full_name — используем его
                if new_c.full_name:
                    c.full_name = new_c.full_name
                if not c.ogrn and new_c.ogrn:
                    c.ogrn = new_c.ogrn
                if not c.kpp and new_c.kpp:
                    c.kpp = new_c.kpp
                if not c.address and new_c.address:
                    c.address = new_c.address
                if not c.org_type and new_c.org_type:
                    c.org_type = new_c.org_type
                if not c.signatory and new_c.signatory:
                    c.signatory = new_c.signatory
                enriched += 1
            else:
                skipped += 1
        except Exception:
            failed += 1

    await db.commit()
    return {
        "ok": True,
        "enriched": enriched,
        "skipped_no_change": skipped,
        "failed": failed,
        "total_processed": len(rows),
    }


@router.post("/enrich-from-receipts")
async def enrich_contractors_from_receipts(
    limit: int = Query(50, ge=1, le=500, description="Сколько ИНН из чеков обработать за один запуск"),
    db: AsyncSession = Depends(get_db),
    _user=Depends(get_current_user),
):
    """Phase 26-CCC-2: обогащение через ЕГРЮЛ для ВСЕХ контрагентов,
    которые встречаются в PurchaseReceipt.seller_inn.

    Отличия от /enrich-all-from-egrul:
    - Берёт ИНН не из Contractor.inn, а из уникальных PurchaseReceipt.seller_inn.
    - ПЕРЕЗАПИСЫВАЕТ Contractor.name на короткое из ЕГРЮЛ (даже если текущее
      короткое — синхронизируем с актуальным ЕГРЮЛ-значением).
    - Заполняет full_name и остальные пустые поля.
    """
    from app.routers.purchase_receipts import _create_or_enrich_contractor_from_receipt
    from app.models.purchase_receipt import PurchaseReceipt

    # Уникальные ИНН продавцов из чеков
    inns_q = await db.execute(
        select(PurchaseReceipt.seller_inn)
        .where(PurchaseReceipt.seller_inn.is_not(None))
        .distinct()
        .limit(limit)
    )
    inns = [r[0] for r in inns_q.all() if r[0]]

    enriched = 0
    name_shortened = 0
    not_found = 0
    failed = 0
    for inn in inns:
        try:
            c_row = (await db.execute(
                select(Contractor).where(Contractor.inn == inn).limit(1)
            )).scalar_one_or_none()
            if not c_row:
                not_found += 1
                continue
            new_c = await _create_or_enrich_contractor_from_receipt(inn, c_row.name, db)
            if not new_c.name:
                failed += 1
                continue
            # ПЕРЕЗАПИСЫВАЕМ name на короткое из ЕГРЮЛ (Phase 26-CCC-2)
            if new_c.name != c_row.name:
                # Сохраняем длинное в full_name если оно ещё пустое
                if not c_row.full_name and len(c_row.name) > len(new_c.name):
                    c_row.full_name = c_row.name
                c_row.name = new_c.name
                name_shortened += 1
            # full_name обновляем если ЕГРЮЛ дал более полное
            if new_c.full_name and (not c_row.full_name or len(new_c.full_name) > len(c_row.full_name or '')):
                c_row.full_name = new_c.full_name
            # Прочие поля — только если пусты
            if not c_row.ogrn and new_c.ogrn:
                c_row.ogrn = new_c.ogrn
            if not c_row.kpp and new_c.kpp:
                c_row.kpp = new_c.kpp
            if not c_row.address and new_c.address:
                c_row.address = new_c.address
            if not c_row.org_type and new_c.org_type:
                c_row.org_type = new_c.org_type
            if not c_row.signatory and new_c.signatory:
                c_row.signatory = new_c.signatory
            enriched += 1
        except Exception:
            failed += 1

    await db.commit()
    return {
        "ok": True,
        "total_inns_from_receipts": len(inns),
        "enriched": enriched,
        "name_shortened": name_shortened,
        "contractor_not_found": not_found,
        "failed": failed,
    }


@router.post("/sync-denormalized-names")
async def sync_denormalized_contractor_names(
    db: AsyncSession = Depends(get_db),
    _user=Depends(get_current_user),
):
    """Phase 26-CCC-3: синхронизировать PurchaseItem.contractor_name
    (denormalized snapshot из чека) с актуальным Contractor.name по contractor_id.

    Нужно после обогащения Contractor.name через ЕГРЮЛ — иначе в UI реестров
    остаются длинные старые имена из чеков, хотя в Contractor уже короткое.
    """
    from sqlalchemy import text as _text
    result = await db.execute(_text("""
        UPDATE purchase_items
        SET contractor_name = c.name
        FROM contractors c
        WHERE purchase_items.contractor_id = c.id
          AND (purchase_items.contractor_name IS DISTINCT FROM c.name)
    """))
    await db.commit()
    return {"ok": True, "updated_rows": result.rowcount}


@router.post("/enrich-from-fns/{contractor_id}")
async def enrich_contractor_from_fns(
    contractor_id: int,
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('contractors')),
):
    """Fetch data from FNS by contractor's INN and fill empty fields."""
    res = await db.execute(select(Contractor).where(Contractor.id == contractor_id))
    c = res.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Контрагент не найден")
    if not c.inn:
        raise HTTPException(400, "У контрагента не заполнен ИНН")

    # Reuse lookup
    fns_data = await lookup_inn(c.inn, _)

    updated_fields = []
    field_map = {
        'name': 'name', 'kpp': 'kpp', 'ogrn': 'ogrn',
        'address': 'address', 'org_type': 'org_type', 'signatory': 'signatory',
    }
    for fns_field, model_field in field_map.items():
        fns_val = fns_data.get(fns_field)
        if fns_val and not getattr(c, model_field, None):
            setattr(c, model_field, fns_val)
            updated_fields.append(model_field)

    await db.commit()
    await db.refresh(c)
    return {"updated_fields": updated_fields, "contractor": ContractorOut.model_validate(c)}


@router.post("/enrich-all-fns")
async def enrich_all_contractors_from_fns(
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_tab('contractors')),
):
    """Bulk-enrich all contractors that have a Russian INN (10 or 12 digits) from FNS."""
    import asyncio
    import re
    import logging
    logger = logging.getLogger(__name__)

    result = await db.execute(select(Contractor).where(Contractor.inn.isnot(None), Contractor.inn != ""))
    contractors = result.scalars().all()

    russian_inn_re = re.compile(r"^\d{10}(\d{2})?$")
    candidates = [c for c in contractors if russian_inn_re.match((c.inn or "").strip())]

    updated_count = 0
    skipped_count = 0
    errors = []

    field_map = {
        'name': 'name', 'kpp': 'kpp', 'ogrn': 'ogrn',
        'address': 'address', 'org_type': 'org_type', 'signatory': 'signatory',
    }

    for c in candidates:
        try:
            fns_data = await lookup_inn(c.inn.strip(), current_user)
            changed = False
            for fns_field, model_field in field_map.items():
                fns_val = fns_data.get(fns_field)
                if fns_val and not getattr(c, model_field, None):
                    setattr(c, model_field, fns_val)
                    changed = True
            if changed:
                updated_count += 1
            else:
                skipped_count += 1
        except HTTPException:
            skipped_count += 1
        except Exception as e:
            errors.append({"inn": c.inn, "error": str(e)[:100]})
            logger.warning("FNS enrich error for INN %s: %s", c.inn, e)

    await db.commit()
    return {
        "total": len(candidates),
        "updated": updated_count,
        "skipped": skipped_count,
        "errors": errors,
    }


@router.post("/parse-file")
async def parse_contractor_file(
    file: UploadFile = File(...),
    _=Depends(get_current_user),
):
    """Parse a contractor card file (xlsx, docx, pdf) and return extracted fields."""
    fname = (file.filename or '').lower()
    content = await file.read()
    try:
        all_rows, hdr_idx = _parse_file_to_rows(fname, content)
    except HTTPException as e:
        raise
    except Exception as e:
        raise HTTPException(400, f"Ошибка чтения файла: {e}")

    if not all_rows or len(all_rows) <= hdr_idx:
        raise HTTPException(400, "Файл не содержит данных")

    # If it's a kv-card (2 rows: header + values), extract directly
    headers = all_rows[hdr_idx] if len(all_rows) > hdr_idx else []
    values = all_rows[hdr_idx + 1] if len(all_rows) > hdr_idx + 1 else []

    result = {}
    field_hints = {
        'name': ('назван', 'наимен', 'name', 'органи', 'учредит'),
        'inn': ('инн', 'inn'),
        'kpp': ('кпп', 'kpp'),
        'ogrn': ('огрн', 'ogrn'),
        'address': ('юридич', 'адрес', 'address'),
        'postal_address': ('почтов', 'postal'),
        'phone': ('телефон', 'phone', 'тел'),
        'email': ('email', 'e-mail'),
        'signatory': ('подписант', 'signatory', 'директор', 'руководит', 'уполномоч'),
        'bik': ('бик', 'bik'),
        'settlement_account': ('расч', 'р/с', 'settlement'),
        'correspondent_account': ('корр', 'к/с', 'correspondent'),
        'bank_name': ('банк', 'bank'),
        'org_type': ('тип', 'форма', 'org_type'),
    }

    for j, h in enumerate(headers):
        h_str = str(h).strip().lower() if h else ''
        if not h_str or j >= len(values):
            continue
        val = str(values[j]).strip() if values[j] is not None else ''
        if not val or val.lower() in ('none', 'null', '-', '—', ''):
            continue
        for field, hints in field_hints.items():
            if field not in result and any(x in h_str for x in hints):
                # INN/KPP length limits
                if field == 'inn' and len(val.replace(' ', '')) > 12:
                    continue
                if field == 'kpp' and len(val.replace(' ', '')) > 9:
                    continue
                result[field] = val
                break

    return result


@router.get("/duplicates-by-inn")
async def list_duplicates_by_inn(
    db: AsyncSession = Depends(get_db),
    current_user=Depends(require_tab('contractors')),
):
    """
    Группирует контрагентов с непустым ИНН, возвращает только группы где count>1.
    Сортировка: max count first, потом по ИНН.
    """
    groups_rows = (await db.execute(
        select(Contractor.inn, func.count(Contractor.id).label('cnt'))
        .where(Contractor.inn.isnot(None))
        .where(func.trim(Contractor.inn) != '')
        .group_by(Contractor.inn)
        .having(func.count(Contractor.id) > 1)
        .order_by(func.count(Contractor.id).desc(), Contractor.inn)
    )).all()

    if not groups_rows:
        return {"groups": [], "total_groups": 0, "total_extra": 0}

    inns = [r[0] for r in groups_rows]
    details = (await db.execute(
        select(Contractor).where(Contractor.inn.in_(inns)).order_by(Contractor.inn, Contractor.id)
    )).scalars().all()

    by_inn: dict = {}
    for c in details:
        by_inn.setdefault(c.inn, []).append({
            "id": c.id,
            "name": c.name or "",
            "full_name": c.full_name or "",
            "kpp": c.kpp or "",
            "address": c.address or "",
            "ogrn": c.ogrn or "",
            "org_type": c.org_type or "",
        })

    out = []
    total_extra = 0
    for inn, cnt in groups_rows:
        contractors = by_inn.get(inn, [])
        out.append({"inn": inn, "count": cnt, "contractors": contractors})
        total_extra += cnt - 1

    return {"groups": out, "total_groups": len(out), "total_extra": total_extra}


@router.get("/", response_model=List[ContractorOut])
async def list_contractors(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
    search: str = Query(None),
    limit: int = Query(200, ge=1, le=5000),
):
    q = select(Contractor)
    # Контрагенты — общий справочник юрлиц (не sensitive data). Org-фильтр
    # убран: при создании закупки/договора любой пользователь должен видеть
    # любого контрагента в autocomplete (фидбек Филиппов 01.06).
    # Раньше: org_id IN user_orgs OR org_id IS NULL → не-админы не видели
    # контрагентов привязанных к другим орг (КРАФТВЭЙ, ЛИНИЯ ГРАФИК и т.п.).
    from sqlalchemy import case
    # Приоритет: фирмы с российским ИНН (ровно 10 или 12 цифр) — выше остальных.
    # Фидбек пользователя: искать приоритетно по российским ИНН, а не все подряд.
    ru_inn_priority = case(
        (Contractor.inn.op('~')(r'^[0-9]{10}$'), 0),
        (Contractor.inn.op('~')(r'^[0-9]{12}$'), 0),
        else_=1,
    )
    if search:
        term = f"%{search}%"
        prefix = f"{search}%"
        relevance = case(
            (Contractor.name.ilike(prefix), 0),
            (Contractor.inn.ilike(prefix), 1),
            (Contractor.name.ilike(term), 2),
            else_=3,
        )
        q = q.where(or_(Contractor.name.ilike(term), Contractor.inn.ilike(term)))
        q = q.order_by(ru_inn_priority, relevance, Contractor.name)
    else:
        q = q.order_by(ru_inn_priority, Contractor.name)
    result = await db.execute(q.limit(limit))
    return result.scalars().all()


@router.get("/{cid}", response_model=ContractorOut)
async def get_contractor(
    cid: int,
    db: AsyncSession = Depends(get_db),
    _=Depends(get_current_user),
):
    c = await db.get(Contractor, cid)
    if not c:
        raise HTTPException(404, "Контрагент не найден")
    return c


@router.post("/", response_model=ContractorOut)
async def create_contractor(
    data: ContractorCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_tab('contractors')),
):
    d = data.model_dump()
    d['org_id'] = get_single_org_id(current_user) or current_user.org_id

    # phase26-oo: дедуп по ИНН — фронт-quick-add часто вызывает POST даже когда
    # контрагент уже есть в БД (race condition или забытый lookup). Возвращаем
    # существующего вместо создания дубля.
    inn = (d.get('inn') or '').strip()
    if inn:
        existing_q = await db.execute(
            select(Contractor).where(Contractor.inn == inn)
        )
        existing = existing_q.scalars().first()
        if existing:
            # Обновим пустые поля если в новом запросе они заполнены (best-effort merge)
            updated = False
            for field in ('name', 'kpp', 'ogrn', 'address', 'phone', 'email', 'signatory'):
                new_val = d.get(field)
                old_val = getattr(existing, field, None)
                if new_val and not old_val:
                    setattr(existing, field, new_val)
                    updated = True
            if updated:
                await db.commit()
                await db.refresh(existing)
            return existing

    c = Contractor(**d)
    # Пересобираем signatory из структурированных частей (только ФИО, без должности)
    if any(d.get(f) for f in ('signatory_last_name', 'signatory_first_name', 'signatory_middle_name')):
        c.signatory = compose_fio(
            d.get('signatory_last_name'),
            d.get('signatory_first_name'),
            d.get('signatory_middle_name'),
        )
    db.add(c)
    await db.commit()
    await db.refresh(c)
    return c


@router.patch("/{cid}/email", response_model=ContractorOut)
async def patch_contractor_email(
    cid: int,
    email: str = Body(..., embed=True),
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('contractors')),
):
    c = (await db.execute(select(Contractor).where(Contractor.id == cid))).scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Not found")
    c.email = email
    await db.commit()
    await db.refresh(c)
    return c


@router.put("/{cid}", response_model=ContractorOut)
async def update_contractor(
    cid: int,
    data: ContractorCreate,
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('contractors'))
):
    result = await db.execute(select(Contractor).where(Contractor.id == cid))
    c = result.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Not found")
    d = data.model_dump()
    for k, v in d.items():
        setattr(c, k, v)
    # Пересобираем signatory из структурированных частей (только ФИО, без должности)
    if any(d.get(f) for f in ('signatory_last_name', 'signatory_first_name', 'signatory_middle_name')):
        c.signatory = compose_fio(
            d.get('signatory_last_name'),
            d.get('signatory_first_name'),
            d.get('signatory_middle_name'),
        )
    await db.commit()
    await db.refresh(c)
    return c


@router.delete("/bulk")
async def bulk_delete_contractors(
    payload: dict,
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('contractors'))
):
    from app.models.purchase import Purchase
    ids = payload.get("ids", [])
    if not ids:
        return {"deleted": 0, "skipped_linked": 0, "skipped_not_found": 0}

    # IDs referenced in purchases — cannot delete
    linked_result = await db.execute(
        select(Purchase.contractor_id).where(Purchase.contractor_id.in_(ids)).distinct()
    )
    linked_ids = {r[0] for r in linked_result}

    safe_ids = [i for i in ids if i not in linked_ids]

    result = await db.execute(select(Contractor).where(Contractor.id.in_(safe_ids)))
    contractors_found = result.scalars().all()
    for c in contractors_found:
        await db.delete(c)
    await db.commit()

    return {
        "deleted": len(contractors_found),
        "skipped_linked": len(linked_ids),
        "skipped_not_found": len(safe_ids) - len(contractors_found),
    }


@router.delete("/{cid}")
async def delete_contractor(
    cid: int,
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('contractors'))
):
    result = await db.execute(select(Contractor).where(Contractor.id == cid))
    c = result.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Not found")
    await db.delete(c)
    await db.commit()
    return {"ok": True}


@router.get("/import/template")
async def contractors_import_template(_=Depends(require_tab('contractors'))):
    """Download xlsx template for bulk contractor import."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise HTTPException(500, "openpyxl не установлен")

    wb = Workbook()
    ws = wb.active
    ws.title = "Контрагенты"

    headers = [
        "Наименование", "ИНН", "КПП", "ОГРН",
        "Адрес местонахождения", "Почтовый адрес",
        "Подписант", "Основание",
        "Контактное лицо", "Телефон", "Email",
        "Расчётный счёт", "Банк", "БИК", "Корр. счёт",
        "Банковские реквизиты",
    ]
    required = {"Наименование", "ИНН"}

    header_fill = PatternFill("solid", fgColor="1E40AF")
    req_fill    = PatternFill("solid", fgColor="1D4ED8")
    header_font = Font(bold=True, color="FFFFFF")

    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.font = header_font
        cell.fill = req_fill if h in required else header_fill
        cell.alignment = Alignment(horizontal="center")

    # Example row
    example = [
        "ООО Пример", "1234567890", "123456789", "1234567890123",
        "г. Москва, ул. Примерная, д. 1", "г. Москва, ул. Почтовая, д. 2",
        "Иванов Иван Иванович, Генеральный директор", "Устава",
        "Петров Пётр Петрович", "+7 (999) 000-00-00", "example@mail.ru",
        "40702810000000000000", "ПАО Сбербанк", "044525225", "30101810400000000225",
        "",
    ]
    for ci, val in enumerate(example, 1):
        ws.cell(row=2, column=ci, value=val)

    col_widths = [35, 14, 12, 16, 40, 40, 45, 20, 30, 20, 25, 24, 30, 12, 24, 35]
    for ci, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=1, column=ci).column_letter].width = w

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{_url_quote('Шаблон_импорта_контрагентов.xlsx', safe='-_.~')}"},
    )


@router.post("/import/excel")
async def import_contractors_excel(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    _=Depends(require_tab('contractors'))
):
    """Bulk import contractors from Excel. First row must be headers."""
    if not (file.filename or '').lower().endswith(('.xlsx', '.xls')):
        raise HTTPException(400, "Поддерживаются только файлы .xlsx / .xls")

    if not load_workbook:
        raise HTTPException(500, "openpyxl не установлен")

    content = await file.read()
    try:
        wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл: {e}")

    ws = wb.active

    # Detect header row
    header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
    if not header_row:
        raise HTTPException(400, "Файл пустой")

    def _norm(v) -> str:
        return str(v).strip().lower() if v else ''

    col: dict[str, int] = {}
    for i, h in enumerate(header_row):
        h_str = _norm(h)
        if any(x in h_str for x in ('назван', 'наимен', 'name', 'органи')):
            col.setdefault('name', i)
        elif any(x in h_str for x in ('инн', 'inn', 'идентиф')):
            col.setdefault('inn', i)
        elif any(x in h_str for x in ('кпп', 'kpp')):
            col.setdefault('kpp', i)
        elif any(x in h_str for x in ('огрн', 'ogrn')):
            col.setdefault('ogrn', i)
        elif any(x in h_str for x in ('почтов', 'postal')):
            col.setdefault('postal_address', i)
        elif any(x in h_str for x in ('адрес', 'address')):
            col.setdefault('address', i)
        elif any(x in h_str for x in ('основан', 'basis', 'устав', 'действует')):
            col.setdefault('signatory_basis', i)
        elif any(x in h_str for x in ('подписант', 'signatory', 'директор', 'руководит')):
            col.setdefault('signatory', i)
        elif any(x in h_str for x in ('контакт', 'contact', 'лицо')):
            col.setdefault('contact_person', i)
        elif any(x in h_str for x in ('телефон', 'phone', 'тел.')):
            col.setdefault('phone', i)
        elif 'email' in h_str or 'e-mail' in h_str or 'mail' in h_str:
            col.setdefault('email', i)
        elif any(x in h_str for x in ('расч', 'р/с', 'р/с', 'settlement')):
            col.setdefault('settlement_account', i)
        elif any(x in h_str for x in ('корр', 'к/с', 'correspondent')):
            col.setdefault('correspondent_account', i)
        elif any(x in h_str for x in ('бик', 'bik')):
            col.setdefault('bik', i)
        elif any(x in h_str for x in ('банк', 'bank')):
            col.setdefault('bank_name', i)
        elif any(x in h_str for x in ('реквизит',)):
            col.setdefault('bank_details', i)

    if 'name' not in col:
        raise HTTPException(
            400,
            "Не найдена колонка с наименованием. "
            "Убедитесь что первая строка — заголовки с колонкой «Наименование» или «name»."
        )

    _limits = {
        'inn': 12, 'kpp': 9, 'ogrn': 20, 'bik': 20,
        'settlement_account': 100, 'correspondent_account': 100,
        'phone': 50, 'email': 255, 'contact_person': 255,
        'signatory': 255, 'signatory_basis': 500, 'bank_name': 500,
    }

    def _cell(row, field):
        idx = col.get(field)
        if idx is None or idx >= len(row):
            return None
        v = row[idx]
        if v is None:
            return None
        s = str(v).strip()
        if not s or s.lower() in ('none', 'null', '-', '—'):
            return None
        # Clean ".0" suffix from numeric fields (Excel float issue)
        if field in ('inn', 'kpp', 'ogrn', 'bik') and s.endswith('.0'):
            s = s[:-2]
        limit = _limits.get(field)
        return s[:limit] if limit else s

    created = 0
    skipped = 0
    errors_list = []

    # Collect existing INNs for dedup
    inn_result = await db.execute(select(Contractor.inn).where(Contractor.inn.isnot(None)))
    existing_inns = {r[0] for r in inn_result}

    for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        name = _cell(row, 'name')
        if not name:
            skipped += 1
            continue

        inn = _cell(row, 'inn')
        if inn and inn in existing_inns:
            skipped += 1
            continue

        c = Contractor(
            name=name,
            inn=inn,
            kpp=_cell(row, 'kpp'),
            ogrn=_cell(row, 'ogrn'),
            address=_cell(row, 'address'),
            postal_address=_cell(row, 'postal_address'),
            signatory=_cell(row, 'signatory'),
            signatory_basis=_cell(row, 'signatory_basis'),
            contact_person=_cell(row, 'contact_person'),
            phone=_cell(row, 'phone'),
            email=_cell(row, 'email'),
            settlement_account=_cell(row, 'settlement_account'),
            bank_name=_cell(row, 'bank_name'),
            bik=_cell(row, 'bik'),
            correspondent_account=_cell(row, 'correspondent_account'),
            bank_details=_cell(row, 'bank_details'),
        )
        db.add(c)
        if inn:
            existing_inns.add(inn)
        created += 1

    await db.commit()
    return {"created": created, "skipped": skipped}


# ---------------------------------------------------------------------------
# New multi-format import endpoints
# ---------------------------------------------------------------------------

@router.post("/import/preview")
async def contractors_import_preview(
    file: UploadFile = File(...),
    _=Depends(require_tab('contractors')),
):
    """Parse uploaded file and return headers + sample rows for column mapping UI."""
    fname = (file.filename or '').lower()
    content = await file.read()

    try:
        all_rows, hdr_idx = _parse_file_to_rows(fname, content)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл: {e}")

    hdr_rows = all_rows[hdr_idx:]
    if not hdr_rows:
        raise HTTPException(400, "Файл пустой или не содержит данных после заголовка")

    headers = [
        str(c).strip() if c else f"Столбец {j + 1}"
        for j, c in enumerate(hdr_rows[0])
    ]
    sample = [
        [str(c).strip() if c is not None else "" for c in row]
        for row in hdr_rows[1:min(4, len(hdr_rows))]
    ]
    total_rows = len(all_rows) - hdr_idx - 1

    return {
        "headers": headers,
        "sample": sample,
        "total_rows": max(total_rows, 0),
        "header_row_offset": hdr_idx,
    }


@router.post("/import/mapped")
async def contractors_import_mapped(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_tab('contractors')),
    col_name: Optional[int] = Query(None),
    col_inn: Optional[int] = Query(None),
    col_kpp: Optional[int] = Query(None),
    col_ogrn: Optional[int] = Query(None),
    col_address: Optional[int] = Query(None),
    col_postal_address: Optional[int] = Query(None),
    col_signatory: Optional[int] = Query(None),
    col_signatory_basis: Optional[int] = Query(None),
    col_contact_person: Optional[int] = Query(None),
    col_phone: Optional[int] = Query(None),
    col_email: Optional[int] = Query(None),
    col_org_phone: Optional[int] = Query(None),
    col_org_email: Optional[int] = Query(None),
    col_settlement_account: Optional[int] = Query(None),
    col_bank_name: Optional[int] = Query(None),
    col_bik: Optional[int] = Query(None),
    col_correspondent_account: Optional[int] = Query(None),
    col_bank_details: Optional[int] = Query(None),
    col_org_type: Optional[int] = Query(None),
    col_manual_product_categories: Optional[int] = Query(None),
    header_row_offset: int = Query(0),
):
    """Import contractors using user-specified column mapping."""
    if col_name is None:
        raise HTTPException(400, "Не указан столбец «Наименование»")

    fname = (file.filename or '').lower()
    content = await file.read()

    try:
        all_rows, _ = _parse_file_to_rows(fname, content)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл: {e}")

    # Skip header row
    skip = header_row_offset + 1
    data_rows = all_rows[skip:]

    _limits = {
        'inn': 12, 'kpp': 9, 'ogrn': 20, 'bik': 20,
        'settlement_account': 100, 'correspondent_account': 100,
        'phone': 50, 'email': 255, 'contact_person': 255,
        'signatory': 255, 'signatory_basis': 500, 'bank_name': 500,
    }

    col_map = {
        'name': col_name,
        'inn': col_inn,
        'kpp': col_kpp,
        'ogrn': col_ogrn,
        'address': col_address,
        'postal_address': col_postal_address,
        'signatory': col_signatory,
        'signatory_basis': col_signatory_basis,
        'contact_person': col_contact_person,
        'phone': col_phone,
        'email': col_email,
        'org_phone': col_org_phone,
        'org_email': col_org_email,
        'settlement_account': col_settlement_account,
        'bank_name': col_bank_name,
        'bik': col_bik,
        'correspondent_account': col_correspondent_account,
        'bank_details': col_bank_details,
        'org_type': col_org_type,
        'manual_product_categories': col_manual_product_categories,
    }

    def _get(row, field):
        idx = col_map.get(field)
        if idx is None or idx >= len(row):
            return None
        v = row[idx]
        if v is None:
            return None
        s = str(v).strip()
        if not s or s.lower() in ('none', 'null', '-', '—'):
            return None
        # Clean ".0" suffix from numeric fields (Excel float issue)
        if field in ('inn', 'kpp', 'ogrn', 'bik') and s.endswith('.0'):
            s = s[:-2]
        limit = _limits.get(field)
        return s[:limit] if limit else s

    # Collect existing contractors by INN for merge
    inn_result = await db.execute(
        select(Contractor).where(Contractor.inn.isnot(None))
    )
    existing_by_inn: dict[str, Contractor] = {c.inn: c for c in inn_result.scalars().all() if c.inn}

    _updatable_fields = [
        'kpp', 'ogrn', 'address', 'postal_address', 'signatory', 'signatory_basis',
        'contact_person', 'phone', 'email', 'org_phone', 'org_email',
        'settlement_account', 'bank_name', 'bik', 'correspondent_account',
        'bank_details', 'org_type',
    ]

    created = 0
    updated = 0
    skipped_empty = 0
    errors_list = []
    update_details = []

    for row_num, row in enumerate(data_rows, start=2):
        try:
            name = _get(row, 'name')
            if not name:
                skipped_empty += 1
                continue

            inn = _get(row, 'inn')

            # Parse categories: comma/semicolon separated string → JSON array
            cats_raw = _get(row, 'manual_product_categories')
            cats = None
            if cats_raw:
                cats = [c.strip() for c in cats_raw.replace(';', ',').split(',') if c.strip()]

            # If contractor with this INN exists — merge new data into it
            if inn and inn in existing_by_inn:
                existing = existing_by_inn[inn]
                changed_fields = []
                for field in _updatable_fields:
                    new_val = _get(row, field)
                    if new_val and not getattr(existing, field, None):
                        setattr(existing, field, new_val)
                        changed_fields.append(field)
                # Merge categories
                if cats:
                    old_cats = existing.manual_product_categories or []
                    merged = list(set(old_cats + cats))
                    if merged != old_cats:
                        existing.manual_product_categories = merged
                        changed_fields.append('categories')
                # Update name if existing is shorter/empty
                if name and (not existing.name or len(name) > len(existing.name)):
                    existing.name = name
                    changed_fields.append('name')
                if changed_fields:
                    updated += 1
                    update_details.append(f"ИНН {inn}: дополнены {', '.join(changed_fields)}")
                continue

            c = Contractor(
                name=name,
                inn=inn,
                kpp=_get(row, 'kpp'),
                ogrn=_get(row, 'ogrn'),
                address=_get(row, 'address'),
                postal_address=_get(row, 'postal_address'),
                signatory=_get(row, 'signatory'),
                signatory_basis=_get(row, 'signatory_basis'),
                contact_person=_get(row, 'contact_person'),
                phone=_get(row, 'phone'),
                email=_get(row, 'email'),
                org_phone=_get(row, 'org_phone'),
                org_email=_get(row, 'org_email'),
                settlement_account=_get(row, 'settlement_account'),
                bank_name=_get(row, 'bank_name'),
                bik=_get(row, 'bik'),
                correspondent_account=_get(row, 'correspondent_account'),
                bank_details=_get(row, 'bank_details'),
                org_type=_get(row, 'org_type'),
                org_id=get_single_org_id(current_user) or current_user.org_id,
                manual_product_categories=cats,
            )
            db.add(c)
            if inn:
                existing_by_inn[inn] = c
            created += 1
        except Exception as e:
            errors_list.append(f"Строка {row_num}: {str(e)}")

    await db.commit()
    return {
        "created": created,
        "updated": updated,
        "skipped": skipped_empty,
        "skipped_empty": skipped_empty,
        "update_details": update_details[:50],
        "errors": errors_list,
    }
