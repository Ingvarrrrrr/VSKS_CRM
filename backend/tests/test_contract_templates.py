# -*- coding: utf-8 -*-
"""Phase 28 T5: приёмочные тесты шаблонов договоров.

42 кейса (7 шаблонов × 6 комбо флагов) + отдельные целевые тесты.
Offline, синхронные, без БД и сети.

Владелец отменил слияние форм под флагами (Этап 1: large_reporting,
food_service, rid_transfer) — итоговый состав форм СЕМЬ отдельных файлов:
contract_goods_single, contract_services, contract_services_food,
contract_gph_individual, contract_gph_individual_rid, contract_repair_vehicle,
contract_repair_framework. Флагами внутри файлов остаются только простые
реквизиты сделки (НДС/без НДС, ИП/ЮЛ, третьи лица, ретро, доставка/самовывоз,
этапность — правила R1–R6).

Методические рекомендации вынесены в отдельные документы (methodology_large,
methodology_small) и физически отсутствуют во всех семи договорных шаблонах.
Абзацы про питание (contract_services_food) и про передачу прав на РИД
(contract_gph_individual_rid) присутствуют в соответствующих файлах
БЕЗУСЛОВНО — никакой Jinja-условности по отменённым флагам.

Логика перенесена из .tmp_p28_strict.py.
"""
import os
import re
import tempfile

import docx
import pytest
from docxtpl import DocxTemplate

# ---------------------------------------------------------------------------
# Пути
# ---------------------------------------------------------------------------

_THIS_DIR = os.path.dirname(__file__)
_TEMPLATES_DIR = os.path.normpath(os.path.join(_THIS_DIR, "..", "templates"))

DOC_TYPES = [
    "contract_goods_single",
    "contract_services",
    "contract_services_food",
    "contract_gph_individual",
    "contract_gph_individual_rid",
    "contract_repair_vehicle",
    "contract_repair_framework",
]

COMBOS = [
    ("НДС20/ЮЛ",  {"vat_rate": 20, "vat_applicable": True,  "contractor_org_type": "ЮЛ"}),
    ("НДС5/ИП",   {"vat_rate": 5,  "vat_applicable": True,  "contractor_org_type": "ИП"}),
    ("безНДС/ИП", {"vat_rate": 0,  "vat_applicable": False, "contractor_org_type": "ИП"}),
    ("без3лиц",   {"third_party_involved": False}),
    ("ретро",     {"is_retroactive": True}),
    ("самовывоз", {"delivery_by_supplier": False, "has_stages": False}),
]

# ---------------------------------------------------------------------------
# Базовый контекст — все строковые значения заменены на уникальные сентинелы
# вида «КЛЮЧ». Это гарантирует, что проверка ловит реальные подстановки,
# а не совпадение с зашитой статикой в шаблоне.
# Сентинелы НЕ содержат < > & — иначе docxtpl портит XML.
# ---------------------------------------------------------------------------

_BASE_STRINGS = {
    "customer_inn": "customer_inn",
    "customer_kpp": "customer_kpp",
    "customer_ogrn": "customer_ogrn",
    "customer_full_name": "customer_full_name",
    "customer_short_name": "customer_short_name",
    "customer_address": "customer_address",
    "customer_bank_name": "customer_bank_name",
    "customer_settlement_account": "customer_settlement_account",
    "customer_correspondent_account": "customer_correspondent_account",
    "customer_bik": "customer_bik",
    "customer_personal_account": "customer_personal_account",
    "customer_signatory_position": "customer_signatory_position",
    "customer_signatory_initials": "customer_signatory_initials",
    "customer_signatory_name_initials": "customer_signatory_name_initials",
    "customer_signatory_name": "customer_signatory_name",
    "customer_signatory_basis": "customer_signatory_basis",
    "customer_email": "customer_email",
    "contract_number": "contract_number",
    "contract_date_day": "contract_date_day",
    "contract_date_month": "contract_date_month",
    "contract_date_year": "contract_date_year",
    "contract_city": "contract_city",
    "contractor_full_name": "contractor_full_name",
    "contractor_short_name": "contractor_short_name",
    "contractor_signatory_name": "contractor_signatory_name",
    "contractor_signatory_position": "contractor_signatory_position",
    "contractor_signatory_basis": "contractor_signatory_basis",
    "contractor_signatory_position_genitive": "contractor_signatory_position_genitive",
    "contractor_signatory_name_genitive": "contractor_signatory_name_genitive",
    "contractor_signatory_line": "contractor_signatory_line",
    "contractor_signatory_name_initials": "contractor_signatory_name_initials",
    "contractor_ogrnip": "contractor_ogrnip",
    "contractor_org_type": "contractor_org_type",
    "contractor_inn": "contractor_inn",
    "contractor_kpp": "contractor_kpp",
    "contractor_ogrn": "contractor_ogrn",
    "contractor_address": "contractor_address",
    "service_subject": "service_subject",
    "service_start_date": "service_start_date",
    "service_end_date": "service_end_date",
    "service_deadline_date": "service_deadline_date",
    "service_term": "service_term",
    "delivery_location": "delivery_location",
    "contract_price_num": "contract_price_num",
    "contract_price_words": "contract_price_words",
    "vat_exemption_article": "vat_exemption_article",
    "penalty_rate": "penalty_rate",
    "service_name": "service_name",
    "subsidy_agreement_text": "subsidy_agreement_text",
    "subsidy_agreement_number": "subsidy_agreement_number",
    "subsidy_agreement_date": "subsidy_agreement_date",
    "subsidy_grantor_name": "subsidy_grantor_name",
    "subsidy_ministry_name": "subsidy_ministry_name",
    "subsidy_extra_clause_1": "subsidy_extra_clause_1",
    "subsidy_extra_clause_2": "subsidy_extra_clause_2",
    "repair_request_number": "repair_request_number",
    "vehicle_body_number": "vehicle_body_number",
    "today": "today",
    "delivery_date": "delivery_date",
    "subject": "subject",
    "subject_kind": "subject_kind",
    "submission_deadline_datetime": "submission_deadline_datetime",
    "vat_amount_num": "vat_amount_num",
    "vat_amount_words": "vat_amount_words",
    "customer_postal_address": "customer_postal_address",
    "customer_phone": "customer_phone",
    "contractor_bank_name": "contractor_bank_name",
    "contractor_settlement_account": "contractor_settlement_account",
    "contractor_correspondent_account": "contractor_correspondent_account",
    "contractor_phone": "contractor_phone",
    "contractor_email": "contractor_email",
    "contractor_signatory_initials": "contractor_signatory_initials",
}

_ITEMS = [
    {
        "num": i,
        "name": "ПОЗИЦИЯ-СЕНТИНЕЛ-%d" % i,
        "quantity": i,
        "unit": "шт",
        "unit_price": "100,00",
        "total": "%d00,00" % i,
        "total_numeric": i * 100.0,
        "code": "КОД%d" % i,
        "norm_hours": "1,5",
    }
    for i in (1, 2, 3)
]


def _make_ctx(**override):
    """Собрать контекст со сентинелами, поверх применить override."""
    ctx = {}
    for k, v in _BASE_STRINGS.items():
        ctx[k] = "«%s»" % k.upper()[:26]
    # Нечисловые дефолты
    ctx.update(
        vat_applicable=True,
        third_party_involved=True,
        is_retroactive=False,
        vat_rate=20,
        warranty_period_days=15,
        acceptance_term_days=5,
        payment_term_days=10,
        service_term_days=30,
        items=_ITEMS,
        contract_items=_ITEMS,
        item={},
        contractor_org_type="ЮЛ",
        has_stages=True,
        delivery_by_supplier=True,
        commission_members=[],
    )
    ctx.update(override)
    return ctx


# ---------------------------------------------------------------------------
# Утилита рендера → plain text
# ---------------------------------------------------------------------------

def _render_text(template_path: str, ctx: dict) -> str:
    t = DocxTemplate(template_path)
    t.render(ctx)
    fd = tempfile.mktemp(suffix=".docx")
    t.save(fd)
    try:
        d = docx.Document(fd)
        parts = [p.text for p in d.paragraphs]
        for tb in d.tables:
            for row in tb.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    finally:
        try:
            os.remove(fd)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# Словари для проверок
# ---------------------------------------------------------------------------

# Зашитые реквизиты ВСКС — не должны встречаться в тексте при сентинел-контексте
STATIC_FORBIDDEN = {
    "ИНН ВСКС": "7731178803",
    "ОГРН ВСКС": "1037739198672",
    "КПП ВСКС": "772901001",
    "Козеев": "Козеев",
    "Росмолодёжь-статик": "Федеральным агентством по делам молодежи",
    "№ соглашения 091": "091-10-2025-032",
    "№ соглашения МТ 149": "149-10-2025-043",
    "ВСКС-полное": "Всероссийский студенческий корпус спасателей",
}

_SERVICES = frozenset(("contract_services", "contract_services_food"))
_GPH = frozenset(("contract_gph_individual", "contract_gph_individual_rid"))

_R5_STAGE_MARKERS = [
    "В случае наличия этапов оказания Услуг",
    "или этапа оказания Услуг",
    "в т.ч. этапа оказания Услуг",
]

# Маркеры «Методических рекомендаций» — методичка вынесена в отдельные
# документы (methodology_large/methodology_small) и НЕ должна физически
# присутствовать ни в одном из семи договорных шаблонов. Эмпирически
# проверенные, не выдумывать новые.
_METHODOLOGY_MARKERS = [
    "МЕТОДИЧЕСКИЕ РЕКОМЕНДАЦИИ",
    "к методическим рекомендациям по",
    "СОДЕРЖАТЕЛЬНЫЙ ОТЧЕТ",
    "Приложение №6",
    "Требования к фотографиям",
]

# Маркеры формы «питание» — присутствуют в contract_services_food БЕЗУСЛОВНО
# (никакого флага food_service — он отменён вместе со слиянием).
_FOOD_MARKERS = [
    "продукты питания",
    "медицинск",
]

# Маркеры передачи прав на РИД — присутствуют в contract_gph_individual_rid
# БЕЗУСЛОВНО и отсутствуют в contract_gph_individual (флаг rid_transfer
# отменён вместе со слиянием).
_RID_MARKERS = [
    "1229",
    "интеллектуальной собственности",
]

# Нерешённые альтернативы — всегда дефект
_UNRESOLVED_ALT = [
    "без привлечения третьих лиц / Услуги",
    "копеек) / в том числе НДС",
    "(включительно) / с момента заключения",
    "г. / в течение",
]


# ---------------------------------------------------------------------------
# Детекторы дефектов
# ---------------------------------------------------------------------------

def _text_defects(txt: str) -> list[str]:
    """Возвращает список описаний найденных дефектов текста."""
    errors = []

    def _snippets(pattern, label):
        found = re.findall(pattern, txt)
        if found:
            sample = found[0]
            errors.append(f"{label} — фрагмент: {' '.join(sample.split())[:120]!r}")

    _snippets(r".{40}\\g<\d+>.{20}", r"литерал \g<N> (regex не раскрыт)")
    _snippets(r".{50}«»", "пустая подстановка «»")
    _snippets(r".{0,40}№\s{1,}(?:от|г\.|\n|$).{0,20}", "«№» без значения")
    _snippets(r".{30}\b(?:в в|с с|по по|на на)\b.{30}", "двойной предлог")
    _snippets(r".{45}\(\s*\)", "пустые скобки ()")
    _snippets(r".{40}%%", "двойной %%")
    _snippets(r".{30}\{[\{%].{30}", "остаток Jinja {{ / {%")
    _snippets(r".{30}»«.{30}", "склейка маркеров »«")
    _snippets(r"[а-яё]{4,}\.[а-яё]", "слово склеено с точкой (порча текста)")
    _snippets(r"(?i)\b([А-Яа-яё]{3,})\s+\1\b", "слово продублировано (порча текста)")

    return errors


def _static_defects(txt: str) -> list[str]:
    errors = []
    for label, value in STATIC_FORBIDDEN.items():
        if value in txt:
            errors.append(f"зашитая статика {label!r}: {value!r} найдена в тексте")
    return errors


def _branch_defects(doc_name: str, ctx: dict, txt: str) -> list[str]:
    errors = []

    def must_have(s, rule):
        if s not in txt:
            errors.append(f"{rule}: ОТСУТСТВУЕТ обязательное «{s}»")

    def must_not(s, rule):
        if s in txt:
            errors.append(f"{rule}: ПРИСУТСТВУЕТ запрещённое «{s}»")

    # R2 — третьи лица (только договоры услуг)
    if doc_name in _SERVICES:
        if ctx.get("third_party_involved") is False:
            must_have("без привлечения третьих лиц", "R2-без3лиц")
            must_not("или с привлечением третьих лиц", "R2-без3лиц")
        else:
            must_have("или с привлечением третьих лиц", "R2-с3лиц")

    # R1 — НДС-ставка
    if ctx.get("vat_applicable") and ctx.get("vat_rate") == 20:
        must_not("НДС – 5%", "R1-НДС20")
        must_not("НДС 5%", "R1-НДС20")
    if ctx.get("vat_applicable") and ctx.get("vat_rate") == 5:
        must_not("НДС – 20%", "R1-НДС5")
        must_not("НДС 20%", "R1-НДС5")
    if not ctx.get("vat_applicable"):
        must_not("НДС – 20%", "R1-безНДС")
        must_not("НДС 20%", "R1-безНДС")
        must_not("НДС – 5%", "R1-безНДС")

    # R4 — доставка / самовывоз (только поставка товара)
    if doc_name == "contract_goods_single":
        if ctx.get("delivery_by_supplier") is False:
            must_have("выборки по месту нахождения Поставщика", "R4-самовывоз")
            must_not("Доставка товара осуществляется силами Поставщика", "R4-самовывоз")
        else:
            must_have("Доставка товара осуществляется силами Поставщика", "R4-доставка")
            must_not("выборки по месту нахождения Поставщика", "R4-доставка")

    # R6 — ретроактивность ст.425 (только услуги)
    if doc_name in _SERVICES:
        if ctx.get("is_retroactive") is False:
            must_not("425", "R6-не-ретро")
        else:
            must_have("425", "R6-ретро")

    # T4 — позиции сметы ремонта разворачиваются в таблицу
    if doc_name == "contract_repair_framework":
        must_have("ПОЗИЦИЯ-СЕНТИНЕЛ-1", "T4-смета-1")
        must_have("ПОЗИЦИЯ-СЕНТИНЕЛ-3", "T4-смета-3")

    # Устаревшая ставка НДС 18% — запрещена везде
    must_not("НДС, 18%", "устарел-НДС18")
    must_not("НДС 18%", "устарел-НДС18")
    must_not("НДС – 18%", "устарел-НДС18")

    # R5 — этапность ГПХ
    if doc_name in _GPH:
        if ctx.get("has_stages") is False:
            for marker in _R5_STAGE_MARKERS:
                must_not(marker, "R5-без-этапов")
        else:
            for marker in _R5_STAGE_MARKERS:
                must_have(marker, "R5-с-этапами")

    # Методичка вынесена в отдельные документы — ни в одном из семи
    # договорных шаблонов её маркеры не должны появляться ни при каком ctx.
    for marker in _METHODOLOGY_MARKERS:
        must_not(marker, "методичка-отсутствует")

    # Питание — абзацы в contract_services_food присутствуют БЕЗУСЛОВНО
    # (никакого флага food_service).
    if doc_name == "contract_services_food":
        for marker in _FOOD_MARKERS:
            must_have(marker, "питание-безусловно")

    # РИД — присутствует БЕЗУСЛОВНО в contract_gph_individual_rid,
    # отсутствует вовсе в contract_gph_individual (никакого флага
    # rid_transfer — формы разные файлы, не варианты одного).
    if doc_name == "contract_gph_individual_rid":
        for marker in _RID_MARKERS:
            must_have(marker, "РИД-безусловно")
    elif doc_name == "contract_gph_individual":
        for marker in _RID_MARKERS:
            must_not(marker, "РИД-отсутствует")

    # Нерешённые альтернативы «/»
    for alt in _UNRESOLVED_ALT:
        if alt in txt:
            errors.append(f"нерешённая альтернатива через «/»: «{alt}»")

    return errors


# ---------------------------------------------------------------------------
# Параметризация
# ---------------------------------------------------------------------------

_PARAMS = [
    pytest.param(doc_type, combo_label, combo_override, id=f"{doc_type}::{combo_label}")
    for doc_type in DOC_TYPES
    for combo_label, combo_override in COMBOS
]


@pytest.mark.parametrize("doc_type,combo_label,combo_override", _PARAMS)
def test_contract_template(doc_type: str, combo_label: str, combo_override: dict):
    """Рендер шаблона с сентинел-контекстом + набор проверок качества.

    При падении сообщение показывает шаблон, комбо и конкретный дефект.
    """
    template_path = os.path.join(_TEMPLATES_DIR, f"{doc_type}.docx")
    assert os.path.isfile(template_path), (
        f"Шаблон не найден: {template_path}"
    )

    ctx = _make_ctx(**combo_override)

    try:
        txt = _render_text(template_path, ctx)
    except Exception as exc:
        pytest.fail(
            f"[{doc_type} / {combo_label}] Рендер упал: {type(exc).__name__}: {exc}"
        )

    all_errors: list[str] = []
    all_errors.extend(_static_defects(txt))
    all_errors.extend(_text_defects(txt))
    all_errors.extend(_branch_defects(doc_type, ctx, txt))

    if all_errors:
        report = "\n".join(f"  • {e}" for e in all_errors)
        pytest.fail(
            f"[{doc_type} / {combo_label}] {len(all_errors)} дефект(а):\n{report}"
        )


# ---------------------------------------------------------------------------
# Откат Этапа 1 слияния форм договоров — отдельные целевые тесты.
# Проверяют «сырой» docx (без рендера docxtpl), чтобы доказать, что
# отменённые конструкции физически отсутствуют в шаблоне, а не просто
# «не сработали» при конкретном наборе значений контекста.
# ---------------------------------------------------------------------------

def _raw_docx_text(template_path: str) -> str:
    """Текст шаблона как есть, без рендера docxtpl (Jinja-теги остаются
    литеральным текстом) — абзацы + ячейки таблиц."""
    d = docx.Document(template_path)
    parts = [p.text for p in d.paragraphs]
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


_CANCELLED_FLAGS = ["large_reporting", "food_service", "rid_transfer"]


@pytest.mark.parametrize("doc_type", DOC_TYPES)
def test_no_cancelled_merge_flags_in_template(doc_type: str):
    """Флаги слияния Этапа 1 (large_reporting/food_service/rid_transfer)
    отменены владельцем — не должны физически присутствовать в тексте
    ни одного из семи шаблонов, в т.ч. внутри Jinja-тегов."""
    template_path = os.path.join(_TEMPLATES_DIR, f"{doc_type}.docx")
    assert os.path.isfile(template_path), f"Шаблон не найден: {template_path}"

    txt = _raw_docx_text(template_path)
    for flag in _CANCELLED_FLAGS:
        assert flag not in txt, (
            f"[{doc_type}] отменённый флаг «{flag}» найден в тексте шаблона"
        )


@pytest.mark.parametrize("doc_type", DOC_TYPES)
def test_no_methodology_markers_raw(doc_type: str):
    """Маркеры методических рекомендаций физически отсутствуют во всех семи
    договорных шаблонах (методичка вынесена в methodology_large/small) —
    проверка на «сыром» docx (абзацы + таблицы), без рендера."""
    template_path = os.path.join(_TEMPLATES_DIR, f"{doc_type}.docx")
    assert os.path.isfile(template_path), f"Шаблон не найден: {template_path}"

    txt = _raw_docx_text(template_path)
    for marker in _METHODOLOGY_MARKERS:
        assert marker not in txt, (
            f"[{doc_type}] маркер методички «{marker}» найден в шаблоне"
        )


def test_services_food_markers_unconditional_raw():
    """Абзацы про питание в contract_services_food НЕ обёрнуты в Jinja-условие
    — присутствуют в тексте шаблона безусловно (сырой docx, без рендера)."""
    template_path = os.path.join(_TEMPLATES_DIR, "contract_services_food.docx")
    assert os.path.isfile(template_path), f"Шаблон не найден: {template_path}"

    txt = _raw_docx_text(template_path)
    for marker in _FOOD_MARKERS:
        assert marker in txt, (
            f"«{marker}» отсутствует в contract_services_food (сырой docx)"
        )


def test_gph_rid_markers_unconditional_raw():
    """Маркеры РИД присутствуют безусловно в contract_gph_individual_rid и
    отсутствуют вовсе в contract_gph_individual (сырой docx, без рендера) —
    это два отдельных файла, не варианты одного под флагом rid_transfer."""
    rid_path = os.path.join(_TEMPLATES_DIR, "contract_gph_individual_rid.docx")
    no_rid_path = os.path.join(_TEMPLATES_DIR, "contract_gph_individual.docx")
    assert os.path.isfile(rid_path), f"Шаблон не найден: {rid_path}"
    assert os.path.isfile(no_rid_path), f"Шаблон не найден: {no_rid_path}"

    rid_txt = _raw_docx_text(rid_path)
    no_rid_txt = _raw_docx_text(no_rid_path)
    for marker in _RID_MARKERS:
        assert marker in rid_txt, (
            f"«{marker}» отсутствует в contract_gph_individual_rid (сырой docx)"
        )
        assert marker not in no_rid_txt, (
            f"«{marker}» неожиданно найден в contract_gph_individual (сырой docx)"
        )


def test_services_vat_rate_change_clause_gated_by_vat_applicable():
    """contract_services: пункт «изменится применяемая ставка НДС» вносится
    правилом R1 независимо от отката слияния питания (п.4 задания) — должен
    присутствовать при vat_applicable=True и отсутствовать при False."""
    template_path = os.path.join(_TEMPLATES_DIR, "contract_services.docx")
    assert os.path.isfile(template_path), f"Шаблон не найден: {template_path}"

    marker = "изменится применяемая ставка НДС"
    txt_true = _render_text(
        template_path, _make_ctx(vat_applicable=True, vat_rate=20)
    )
    txt_false = _render_text(
        template_path, _make_ctx(vat_applicable=False, vat_rate=0)
    )

    assert marker in txt_true, f"«{marker}» отсутствует при vat_applicable=True"
    assert marker not in txt_false, f"«{marker}» присутствует при vat_applicable=False"
