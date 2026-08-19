"""Purchase items-import router — extracted from purchases.py (Phase 16-03).

Handles:
  GET  /api/purchases/items/import/template       — download blank xlsx template
  POST /api/purchases/{pid}/items/import          — bulk import from Excel (legacy)
  POST /api/purchases/items/import-preview        — parse file, return headers/samples for mapping
  POST /api/purchases/{pid}/items/import-mapped   — import using user-specified column mapping
  POST /api/purchases/{pid}/items/import-smart    — smart AI-assisted import (markitdown + fallback)
  GET  /api/purchases/import/feo-format/template  — download FEO-format template
  POST /api/purchases/import/feo-format           — import purchases from FEO 57-column format

OCR helpers (_ocr_pdf_to_rows, _legacy_extract_tables, _legacy_detect_best_table) and
the product-catalog upsert helper (_upsert_product_to_catalog) live in this module.
"""
import logging
from urllib.parse import quote as _url_quote
from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession
from io import BytesIO
from typing import Optional
from decimal import Decimal
from datetime import datetime, date

from app.database import get_db
from app.models.purchase import Purchase
from app.models.purchase_item import PurchaseItem
from app.models.product import Product
from app.models.contractor import Contractor
from app.models.feo_category import FeoCategory
from app.models.subsidy import Subsidy
from app.auth.jwt import get_current_user, require_role, get_single_org_id, MANAGER_ROLES
from app.auth.permissions import require_tab, has_org_key
from app.routers.purchases import _has_purchase_write_access
from app.models.user import User
from app.services.product_matcher import score as _fuzzy_score, SCORE_AUTO as _SCORE_AUTO
from app.services.feo_plan import assert_tz_not_over_plan

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment
except ImportError:
    Workbook = None
    load_workbook = None

try:
    import pdfplumber as _pdfplumber
except ImportError:
    _pdfplumber = None

try:
    from docx import Document as _DocxDocument
except ImportError:
    _DocxDocument = None

try:
    from bs4 import BeautifulSoup as _BeautifulSoup
except ImportError:
    _BeautifulSoup = None

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/purchases", tags=["purchase-items-import"])


@router.get("/items/import-debug")
async def import_debug(current_user: User = Depends(get_current_user)):
    """Diagnostic: report which OCR/parsing libraries are loaded on the server.

    Use to debug 'PDF не распознаётся' issues. Hit GET /api/purchases/items/import-debug
    after autodeploy completes to verify tesseract/ocrmypdf binaries are installed.
    """
    out: dict = {}
    try:
        import pdfplumber  # noqa: F401
        out["pdfplumber"] = "ok"
    except Exception as e:
        out["pdfplumber"] = f"FAIL: {e}"
    try:
        import pytesseract
        out["pytesseract_python"] = "ok"
        try:
            v = pytesseract.get_tesseract_version()
            out["tesseract_binary"] = f"ok (v{v})"
            langs = pytesseract.get_languages()
            out["tesseract_langs"] = ",".join(sorted(langs))
            out["has_rus"] = "rus" in langs
        except Exception as e:
            out["tesseract_binary"] = f"FAIL: {e}"
    except Exception as e:
        out["pytesseract_python"] = f"FAIL: {e}"
    try:
        from pdf2image import convert_from_bytes  # noqa: F401
        out["pdf2image"] = "ok"
    except Exception as e:
        out["pdf2image"] = f"FAIL: {e}"
    try:
        import ocrmypdf  # noqa: F401
        out["ocrmypdf"] = "ok"
    except Exception as e:
        out["ocrmypdf"] = f"FAIL: {e}"
    try:
        from markitdown import MarkItDown  # noqa: F401
        out["markitdown"] = "ok"
    except Exception as e:
        out["markitdown"] = f"FAIL: {e}"
    try:
        from bs4 import BeautifulSoup  # noqa: F401
        out["beautifulsoup4"] = "ok"
    except Exception as e:
        out["beautifulsoup4"] = f"FAIL: {e}"
    try:
        from openpyxl import load_workbook  # noqa: F401
        out["openpyxl"] = "ok"
    except Exception as e:
        out["openpyxl"] = f"FAIL: {e}"
    # Check ghostscript (needed by ocrmypdf)
    import shutil
    out["ghostscript_bin"] = "ok" if shutil.which("gs") else "MISSING"
    out["pdftoppm_bin"] = "ok" if shutil.which("pdftoppm") else "MISSING"
    out["unpaper_bin"] = "ok" if shutil.which("unpaper") else "MISSING"
    try:
        import pdf_inspector
        out["pdf_inspector"] = f"ok (v{getattr(pdf_inspector, '__version__', '?')})"
    except Exception as e:
        out["pdf_inspector"] = f"FAIL: {e}"
    try:
        import cv2
        out["cv2"] = f"ok (v{cv2.__version__})"
    except Exception as e:
        out["cv2"] = f"FAIL: {e}"
    try:
        from app.utils.pdf_classify import inspect_pdf
        # Minimal single-page PDF with a couple of words, generated via reportlab
        # (already a dependency) rather than hand-built — just to exercise the
        # subprocess plumbing, not to assert anything about classification quality.
        from reportlab.pdfgen import canvas as _rl_canvas
        _buf = BytesIO()
        _c = _rl_canvas.Canvas(_buf)
        _c.drawString(72, 700, "VSKS import-debug test PDF")
        _c.save()
        _tiny_pdf = _buf.getvalue()
        _insp = inspect_pdf(_tiny_pdf, timeout=15)
        if _insp.ok:
            out["pdf_inspector_subprocess"] = "ok"
            # Real values from the library, not documentation guesses — for a
            # 1-page reportlab text PDF we'd *expect* pdf_type=="text",
            # ocr_pages==[], page_count==1, but show whatever actually comes back.
            out["pdf_inspector_result"] = {
                "pdf_type": _insp.pdf_type,
                "confidence": _insp.confidence,
                "page_count": _insp.page_count,
                "has_encoding_issues": _insp.has_encoding_issues,
                "ocr_pages": _insp.ocr_pages,
            }
        else:
            out["pdf_inspector_subprocess"] = f"FAIL: {_insp.error}"
    except Exception as e:
        out["pdf_inspector_subprocess"] = f"FAIL: {e}"
    return out


# ---------------------------------------------------------------------------
# OCR / legacy table-extraction helpers
# ---------------------------------------------------------------------------

def _ocr_pdf_to_rows(content: bytes) -> tuple[list, str | None]:
    """Thin delegating wrapper around app.utils.pdf_ocr.ocr_pdf_to_rows.

    The real implementation was moved there (Wave 1 of the OCR-classification
    work) so it isn't duplicated between this router and
    app/utils/document_to_markdown.py. Kept here — same name, same signature,
    unpaged — as the module's own OCR-fallback entry point (the two in-module
    call sites that need page-targeted OCR now call
    app.utils.pdf_ocr.ocr_pdf_to_rows_enhanced directly instead, since this
    wrapper's signature has no `pages` argument to pass one through).
    Import is local so importing this router doesn't pull in OCR-heavy deps.
    """
    from app.utils.pdf_ocr import ocr_pdf_to_rows as _impl
    return _impl(content)


def _ocrmypdf_then_extract_tables(content: bytes) -> list[list[list[str]]]:
    """Thin delegating wrapper around app.utils.pdf_ocr.ocrmypdf_then_extract_tables.

    See `_ocr_pdf_to_rows` above — same reasoning, same "keep name/signature,
    delegate the body" approach.
    """
    from app.utils.pdf_ocr import ocrmypdf_then_extract_tables as _impl
    return _impl(content)


def _extract_html_tables(content: bytes, filename: str = "doc.html") -> list[list[list[str]]]:
    """Single source of truth for HTML→tables extraction.

    Used by BOTH /items/import-preview AND /items/import-mapped so that table
    indices ("Таблица 1", "Таблица 2"...) refer to the same tables in both
    endpoints. Order: markitdown → parse_markdown_tables, fallback BeautifulSoup.
    """
    raw_tables: list[list[list[str]]] = []
    try:
        from app.utils.document_to_markdown import file_to_markdown, parse_markdown_tables
        md_text = file_to_markdown(content, filename)
        raw_tables = parse_markdown_tables(md_text)
    except Exception as e:
        logger.warning("markitdown HTML parse failed: %s", e)
        raw_tables = []
    if not raw_tables:
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            for tbl in soup.find_all('table'):
                rows = []
                for tr in tbl.find_all('tr'):
                    cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                    if cells:
                        rows.append(cells)
                if rows:
                    raw_tables.append(rows)
        except Exception as e:
            logger.warning("BeautifulSoup HTML fallback failed: %s", e)
    return raw_tables


def _legacy_extract_tables(content: bytes, filename: str, file_type: str) -> list[list[list[str]]]:
    """Fallback table extraction using original libraries (pdfplumber, python-docx, openpyxl)."""
    raw_tables: list[list[list[str]]] = []
    if file_type == "excel":
        if load_workbook:
            wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
            ws = wb.active
            rows = [[str(c) if c is not None else "" for c in row] for row in ws.iter_rows(values_only=True)]
            if rows:
                raw_tables.append(rows)
    elif file_type == "pdf":
        if _pdfplumber:
            with _pdfplumber.open(BytesIO(content)) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables() or []
                    for tbl in tables:
                        if tbl:
                            raw_tables.append([[str(c) if c is not None else "" for c in row] for row in tbl])
        # Classify once, only if pdfplumber didn't already find tables — the
        # classification subprocess isn't free, no point paying for it when
        # pdfplumber's native table extraction already succeeded.
        _ocr_pages: list[int] | None = None
        if not raw_tables:
            try:
                from app.utils.pdf_classify import inspect_pdf as _inspect_pdf_classify
                _insp = _inspect_pdf_classify(content)
                if _insp.ok:
                    _ocr_pages = _insp.ocr_pages or None
                    logger.info(
                        "_legacy_extract_tables: pdf_classify pdf_type=%s ocr_pages=%s (classifier path)",
                        _insp.pdf_type, _insp.ocr_pages,
                    )
                else:
                    logger.info("_legacy_extract_tables: pdf classification failed (%s), OCR-ing whole document", _insp.error)
            except Exception as e:
                logger.warning("_legacy_extract_tables: pdf classification unavailable, OCR-ing whole document: %s", e)
        if not raw_tables:
            # Try ocrmypdf — adds OCR layer, then pdfplumber extracts tables natively
            ocr_tables = _ocrmypdf_then_extract_tables(content)
            if ocr_tables:
                raw_tables.extend(ocr_tables)
        if not raw_tables:
            # Last resort: raw OCR via image_to_data, with preprocessing + PSM
            # auto-selection, restricted to pdf-inspector's flagged pages when
            # classification succeeded (whole document otherwise).
            from app.utils.pdf_ocr import ocr_pdf_to_rows_enhanced
            ocr_rows, _ = ocr_pdf_to_rows_enhanced(content, pages=_ocr_pages)
            if ocr_rows:
                raw_tables.append(ocr_rows)
    elif file_type == "docx":
        if _DocxDocument:
            doc = _DocxDocument(BytesIO(content))
            for table in doc.tables:
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if rows:
                    raw_tables.append(rows)
    elif file_type == "html":
        if _BeautifulSoup:
            soup = _BeautifulSoup(content, 'html.parser')
            for tbl in soup.find_all('table'):
                rows = []
                for tr in tbl.find_all('tr'):
                    cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                    if cells:
                        rows.append(cells)
                if rows:
                    raw_tables.append(rows)
        else:
            logger.warning("BeautifulSoup not installed, skipping HTML fallback parsing")
    return raw_tables


def _detect_upd_layout(row: list[str]) -> dict | None:
    """Detect УПД (Универсальный передаточный документ) numeric label row.

    УПД standard column layout (Постановление Правительства РФ № 1137):
      1   — № п/п
      1а  — код товара/работ/услуг (артикул)
      1б  — Наименование товара (описание выполненных работ, оказанных услуг)
      1в  — код вида товара
      2   — единица измерения, код
      2а  — единица измерения, условное обозначение
      3   — Количество (объем)
      4   — Цена (тариф) за единицу
      5   — Стоимость без налога
      6   — В т.ч. сумма акциза
      7   — Налоговая ставка
      8   — Сумма налога
      9   — Стоимость С НАЛОГОМ - всего
      10  — Страна происхождения - код
      10а — Страна происхождения - название
      11  — Регистрационный номер декларации

    Returns column mapping dict or None if row doesn't match УПД pattern.
    Match rule: ≥4 of {'1а','1б','2','2а','3','4','5','9'} present.
    """
    LABELS = {'1а', '1б', '2', '2а', '3', '4', '5', '9'}
    found = {}
    for i, h in enumerate(row):
        h_s = h.strip().lower().replace(' ', '')
        if h_s in LABELS or h_s in {'1', '6', '7', '8', '10', '10а', '11', '12', '12а', '13'}:
            found[h_s] = i
    if len(LABELS & set(found.keys())) < 4:
        return None
    col: dict = {}
    if '1б' in found:
        col['item_name'] = found['1б']
    if '3' in found:
        col['quantity'] = found['3']
    if '4' in found:
        col['unit_price'] = found['4']
    # Prefer total WITH tax (col 9), fallback to without-tax (col 5)
    if '9' in found:
        col['total_price'] = found['9']
    elif '5' in found:
        col['total_price'] = found['5']
    if '2а' in found:
        col['unit'] = found['2а']
    elif '2' in found:
        col['unit'] = found['2']
    return col if 'item_name' in col else None


def _legacy_detect_best_table(raw_tables: list[list[list[str]]]) -> tuple:
    """Fallback column detection on raw tables."""
    def _detect_columns_legacy(header_row: list[str]) -> dict:
        col: dict = {}
        for i, h in enumerate(header_row):
            h_s = h.strip().lower()
            # УПД precise headers (priority)
            if 'наименование товара' in h_s or 'описание выполненных' in h_s or 'описание оказанных' in h_s:
                col.setdefault("item_name", i)
            elif any(x in h_s for x in (
                "наименован", "назван", "name", "товар", "предмет", "описан",
                "услуг", "работ",
                # import-pdf-debug C2: билетный формат
                "маршрут", "направлен", "рейс", "билет",
            )):
                col.setdefault("item_name", i)
            elif any(x in h_s for x in ("тип", "type", "вид")):
                col.setdefault("item_type", i)
            elif any(x in h_s for x in ("количество (объем)", "кол-во", "количеств", "qty", "quantity")) or h_s.startswith("кол"):
                col.setdefault("quantity", i)
            elif "ед. изм" in h_s or "единиц" in h_s or "ед.изм" in h_s or h_s == "ед." or h_s == "unit":
                col.setdefault("unit", i)
            elif any(x in h_s for x in (
                "цена", "тариф", "price",
                # import-pdf-debug C2: стоимость единицы / тариф билета
                "стоимость", "сумма билет", "цена билет",
            )):
                col.setdefault("unit_price", i)
            # VAT columns (import-vat-cols C2)
            elif any(x in h_s for x in ("ставка ндс", "налоговая ставка", "% ндс", "ндс %", "vat rate")):
                col.setdefault("vat_rate", i)
            elif any(x in h_s for x in ("сумма ндс", "vat amount")):
                col.setdefault("vat_amount", i)
            # Total: prefer "с налогом", else "без налога"/"стоимость"/"сумма"
            elif "с налогом" in h_s and "всего" in h_s:
                col["total_price"] = i  # always overwrite — highest priority
            elif any(x in h_s for x in ("к оплате", "сумма с налогом", "итого с ндс")):
                col["total_price"] = i  # import-pdf-debug C2: билетные итоги
            elif "total_price" not in col and any(x in h_s for x in ("сумма", "итог", "total", "amount", "всего", "стоимость")):
                col["total_price"] = i
        return col

    best_table: list[list[str]] = []
    best_col: dict = {}
    best_header_row = 0
    # First pass: try УПД numeric-label detection (positional mapping)
    for table in raw_tables:
        for r_idx, row in enumerate(table[:8]):
            upd_col = _detect_upd_layout(row)
            if upd_col and len(upd_col) > len(best_col):
                best_col = upd_col
                best_table = table
                best_header_row = r_idx
    if best_col:
        return best_table, best_col, best_header_row
    # Second pass: keyword-based detection (legacy)
    for table in raw_tables:
        for r_idx, row in enumerate(table[:6]):
            col = _detect_columns_legacy(row)
            if "item_name" in col and len(col) > len(best_col):
                best_col = col
                best_table = table
                best_header_row = r_idx
    return best_table, best_col, best_header_row


# ---------------------------------------------------------------------------
# Product-catalog upsert helper
# ---------------------------------------------------------------------------

async def _upsert_product_to_catalog(
    db, item_name: str, item_type: str, unit_price, description: str = "",
    category: str | None = None, product_type: str | None = None,
    import_note: str | None = None, updated_by: str | None = None,
) -> int:
    """Find or create a product in the global catalog. Returns product.id.

    Правила конфликтов при импорте из файла:
    - цена: обновляется из файла (файл — источник актуальной цены);
    - категория/вид: БД главнее — из файла берём только если в БД пусто
      (для категории дефолт «Прочее» считается пустым);
    - import_note: кто/как/когда загрузил — перезаписывается свежим импортом.
    """
    from datetime import datetime as _dt
    norm = item_name.strip().lower()
    existing = (await db.execute(
        select(Product).where(func.lower(Product.name) == norm)
    )).scalar_one_or_none()
    if existing:
        new_price = Decimal(str(unit_price)) if unit_price else None
        if new_price and existing.price != new_price:
            existing.price = new_price
        if category and (not existing.category or existing.category == 'Прочее'):
            existing.category = category
        if product_type and not existing.product_type:
            existing.product_type = product_type
        if import_note:
            existing.import_note = import_note
            existing.updated_at = _dt.utcnow()
            if updated_by:
                existing.updated_by = updated_by
        return existing.id
    p = Product(
        name=item_name.strip(),
        description=description or "",
        category=category or 'Прочее',
        product_type=product_type or item_type or "товар",
        item_kind=item_type or "товар",
        price=Decimal(str(unit_price)) if unit_price else Decimal("0"),
        is_active=True,
        import_note=import_note,
        updated_at=_dt.utcnow() if import_note else None,
        updated_by=updated_by if import_note else None,
    )
    db.add(p)
    await db.flush()
    return p.id


# ---------------------------------------------------------------------------
# Purchase items import from Excel
# ---------------------------------------------------------------------------

@router.get("/items/import/template")
async def items_import_template(_=Depends(require_tab('purchases'))):
    """Download xlsx template for bulk purchase items import."""
    if not Workbook:
        raise HTTPException(500, "openpyxl не установлен")

    wb = Workbook()
    ws = wb.active
    ws.title = "Позиции"

    headers = ["Наименование", "Описание / ТЗ", "Тип (товар/услуга/работа)", "Количество", "Ед. изм.", "Цена за единицу"]
    required = {"Наименование"}

    header_fill = PatternFill("solid", fgColor="1E40AF")
    req_fill = PatternFill("solid", fgColor="1D4ED8")
    header_font = Font(bold=True, color="FFFFFF")

    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.font = header_font
        cell.fill = req_fill if h in required else header_fill
        cell.alignment = Alignment(horizontal="center")

    example = ["Ноутбук Lenovo ThinkPad", "Технические характеристики...", "товар", "5", "шт", "85000"]
    for ci, val in enumerate(example, 1):
        ws.cell(row=2, column=ci, value=val)

    col_widths = [45, 50, 25, 15, 15, 20]
    for ci, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=1, column=ci).column_letter].width = w

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{_url_quote('Шаблон_импорта_позиций_закупки.xlsx', safe='-_.~')}"},
    )


@router.post("/{pid}/items/import")
async def import_items_excel(
    pid: int,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Bulk import items into a purchase from Excel."""
    fname = (file.filename or '').lower()
    if not fname.endswith(('.xlsx', '.xls')):
        raise HTTPException(400, "Поддерживаются только файлы .xlsx / .xls")

    purchase = await db.get(Purchase, pid)
    if not purchase:
        raise HTTPException(404, "Закупка не найдена")

    content = await file.read()
    try:
        sheets = _read_excel_rows(content, fname)
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл: {e}")
    if not sheets or not sheets[0]:
        raise HTTPException(400, "Файл пустой")
    all_rows = sheets[0]  # первый лист
    header_row = all_rows[0] if all_rows else None
    data_iter = all_rows[1:] if len(all_rows) > 1 else []
    if not header_row:
        raise HTTPException(400, "Файл пустой")

    def _norm(v) -> str:
        return str(v).strip().lower() if v else ''

    # Fuzzy column mapping
    col: dict[str, int] = {}
    for i, h in enumerate(header_row):
        h_str = _norm(h)
        if any(x in h_str for x in ('наименован', 'назван', 'name', 'товар', 'предмет')):
            col.setdefault('item_name', i)
        elif any(x in h_str for x in ('описан', 'description', 'тз', 'техническ', 'specification', 'характерист')):
            col.setdefault('description', i)
        elif any(x in h_str for x in ('тип', 'type', 'вид')):
            col.setdefault('item_type', i)
        elif any(x in h_str for x in ('кол', 'количеств', 'qty', 'quantity')):
            col.setdefault('quantity', i)
        # Шаг 5 (владелец, 2026-08-07): проверка цены ПЕРЕД проверкой ед.изм. — заголовок
        # шаблона «Цена за единицу» (см. items_import_template выше) содержит подстроку
        # «единицу», которая раньше матчилась веткой ед.изм. ('единиц' in h_str) первой
        # (elif сверху вниз) — колонка цены никогда не находилась, unit_price/total_price
        # молча оставались NULL. Обнаружено при проверке шага 5 «ТЗ не выше плана»: гейт
        # не мог сработать, т.к. цена не читалась из файла вообще.
        elif any(x in h_str for x in ('цена', 'price', 'стоимость', 'за единиц')):
            col.setdefault('unit_price', i)
        elif any(x in h_str for x in ('ед.', 'единиц', 'unit', 'изм')):
            col.setdefault('unit', i)

    if 'item_name' not in col:
        raise HTTPException(400, "Не найдена колонка с наименованием.")

    def _cell(row, field):
        idx = col.get(field)
        if idx is None or idx >= len(row):
            return None
        v = row[idx]
        if v is None:
            return None
        s = str(v).strip()
        if not s or s.lower() in ('none', 'null', '-', '—'):
            return None
        return s

    def _to_dec(v):
        if v is None:
            return None
        try:
            return Decimal(str(v).replace(',', '.').replace(' ', ''))
        except Exception:
            return None

    # Load products for auto-matching by name
    org_id = get_single_org_id(current_user)
    prod_q = select(Product)
    if org_id:
        prod_q = prod_q.where((Product.org_id == org_id) | (Product.org_id.is_(None)))
    prod_result = await db.execute(prod_q)
    products = prod_result.scalars().all()
    # Build name lookup (lowercase → product)
    product_by_name: dict[str, Product] = {}
    for p in products:
        if p.name:
            product_by_name[p.name.lower().strip()] = p

    TYPE_MAP = {
        'товар': 'товар', 'товары': 'товар', 'product': 'товар', 'goods': 'товар',
        'услуга': 'услуга', 'услуги': 'услуга', 'service': 'услуга',
        'работа': 'работа', 'работы': 'работа', 'work': 'работа',
    }

    added = 0
    matched_catalog = 0
    new_in_catalog = 0
    errors_list = []

    for row_idx, row in enumerate(data_iter, start=2):  # +1 заголовок, +1 — 1-based для пользователя
        item_name = _cell(row, 'item_name')
        if not item_name:
            continue

        description = _cell(row, 'description')
        item_type_raw = (_cell(row, 'item_type') or 'товар').lower().strip()
        item_type = TYPE_MAP.get(item_type_raw, 'товар')
        quantity = _to_dec(_cell(row, 'quantity')) or Decimal('1')
        unit = _cell(row, 'unit') or 'шт'
        unit_price = _to_dec(_cell(row, 'unit_price'))
        total_price = (quantity * unit_price) if unit_price else None

        # Шаг 5 «цена ТЗ не выше плановой» (владелец, 2026-08-07): позиция импорта
        # наследует ФЭО-категорию закупки (feo_planned_item_id импорт не проставляет —
        # это делается отдельно, автоподбором/вручную после импорта). Ошибки
        # аггрегируем по строкам — импорт не должен падать на первой же проблемной
        # позиции, проблемные строки просто не добавляются, остальные — добавляются.
        try:
            await assert_tz_not_over_plan(
                db,
                feo_planned_item_id=None,
                feo_category_id=purchase.feo_category_id,
                quantity=quantity,
                unit_price=unit_price,
                total_price=total_price,
                item_name=item_name,
            )
        except HTTPException as _tz_exc:
            errors_list.append(f"Строка {row_idx}: {_tz_exc.detail}")
            continue

        # Auto-match or create in catalog
        # 1) exact match (fast path)
        matched_product = product_by_name.get(item_name.lower().strip())
        if not matched_product:
            # 2) fuzzy fallback — find best candidate above SCORE_AUTO threshold
            best_score = 0.0
            best_candidate = None
            for _key, _p in product_by_name.items():
                _s = _fuzzy_score(item_name, _p.name if hasattr(_p, 'name') else _key)
                if _s > best_score:
                    best_score = _s
                    best_candidate = _p
            if best_score >= _SCORE_AUTO and best_candidate is not None:
                matched_product = best_candidate
        if matched_product:
            product_id = matched_product.id
            matched_catalog += 1
            if not unit_price and matched_product.price:
                unit_price = matched_product.price
                total_price = quantity * unit_price
        else:
            _uname = getattr(current_user, 'full_name', None) or getattr(current_user, 'username', '') or ''
            product_id = await _upsert_product_to_catalog(
                db, item_name, item_type, unit_price, description or "",
                import_note=f"Импорт из файла «{file.filename}» (шаблон), {_uname}, {datetime.now().strftime('%d.%m.%Y %H:%M')}",
                updated_by=_uname,
            )
            product_by_name[item_name.lower().strip()] = type('_P', (), {'id': product_id, 'name': item_name, 'price': unit_price})()
            new_in_catalog += 1

        item = PurchaseItem(
            purchase_id=pid,
            product_id=product_id,
            item_name=item_name,
            item_type=item_type,
            quantity=quantity,
            unit=unit,
            unit_price=unit_price,
            total_price=total_price,
        )
        db.add(item)
        added += 1

    await db.commit()
    return {"added": added, "matched_catalog": matched_catalog, "new_in_catalog": new_in_catalog, "errors": errors_list}


# ---------------------------------------------------------------------------
# Excel import with column mapping (preview + mapped import)
# ---------------------------------------------------------------------------

@router.post("/items/import-preview")
async def import_items_preview(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    """Read Excel/PDF/DOCX file and return sheets, headers, and sample rows for column mapping."""
    fname = (file.filename or '').lower()
    if not fname.endswith(('.xlsx', '.xls', '.pdf', '.docx', '.doc', '.html', '.htm')):
        raise HTTPException(400, "Поддерживаются файлы .xlsx, .xls, .pdf, .docx, .html")

    content = await file.read()

    _NAME_HINTS = ('наименован', 'назван', 'товар', 'предмет', 'name', 'title', 'услуг', 'работ')
    _ALL_HINTS = _NAME_HINTS + ('цена', 'описан', 'кол', 'тип', 'price', 'стоимост', 'ед.', 'единиц', 'катег', 'сумм', 'количеств', 'ед. изм')

    def _detect_hdr(rows):
        best_score, best_idx = 0, 0
        for ri, row in enumerate(rows):
            norm = [str(h).strip().lower() if h is not None else "" for h in row]
            score = sum(1 for h in norm if h and any(x in h for x in _ALL_HINTS))
            if score > best_score:
                best_score = score; best_idx = ri
        return best_idx

    try:
        # ── PDF ──
        if fname.endswith('.pdf'):
            try:
                import pdfplumber
            except ImportError:
                raise HTTPException(500, "pdfplumber не установлен")
            pdf = pdfplumber.open(BytesIO(content))
            all_rows = []
            text_lines = []
            for page in pdf.pages:
                # Try tables
                for t in (page.extract_tables() or []):
                    if t:
                        all_rows.extend([[str(c).strip() if c else "" for c in row] for row in t])
                # Also collect text lines as fallback
                for line in (page.extract_text() or "").split('\n'):
                    line = line.strip()
                    if line:
                        text_lines.append(line)
            pdf.close()

            # Heuristic: detect "garbage text layer" from scanner OCR.
            # Sharp/Xerox scanners often embed a low-quality OCR layer that pdfplumber
            # reads as text but the result is unparseable. Skip text_lines and force
            # ocrmypdf if no line has ≥10 chars with a meaningful Russian/English word.
            # This is the fallback path — kept as-is for when pdf-inspector classification
            # is unavailable, but the classifier below takes priority when it succeeds.
            import re
            def _looks_like_real_text(lines: list[str]) -> bool:
                for ln in lines:
                    # Strip non-letter chars, count alphabetic runs of length ≥4
                    words = re.findall(r"[А-Яа-яA-Za-z]{4,}", ln)
                    if len(words) >= 2 and len(ln) >= 10:
                        return True
                return False

            # Classify once per request (not once per fallback attempt — the
            # subprocess isn't free) and let it decide "garbage text layer"
            # when it's available; the home-grown heuristic above stays as the
            # fallback for when classification itself is unavailable.
            _insp = None
            _ocr_pages: list[int] | None = None
            try:
                from app.utils.pdf_classify import inspect_pdf as _inspect_pdf_classify, needs_ocr as _needs_ocr_classify
                _insp = _inspect_pdf_classify(content)
            except Exception as e:
                logger.warning("PDF preview: pdf classification unavailable, using legacy heuristic: %s", e)
                _insp = None

            if _insp is not None and _insp.ok:
                text_layer_garbage = _needs_ocr_classify(_insp)
                text_layer_usable = bool(all_rows) or not text_layer_garbage
                _ocr_pages = _insp.ocr_pages or None
                logger.info(
                    "PDF preview: decision via classifier (pdf_type=%s, has_encoding_issues=%s, "
                    "ocr_pages=%s) → text_layer_usable=%s",
                    _insp.pdf_type, _insp.has_encoding_issues, _insp.ocr_pages, text_layer_usable,
                )
            else:
                text_layer_usable = bool(all_rows) or _looks_like_real_text(text_lines)
                logger.info(
                    "PDF preview: decision via legacy heuristic (classifier %s) → text_layer_usable=%s",
                    "unavailable" if _insp is None else f"failed: {_insp.error}", text_layer_usable,
                )

            if not all_rows and text_layer_usable and text_lines:
                # Real text PDF without explicit tables — split lines by gaps
                for line in text_lines:
                    parts = re.split(r'\t|  +', line)
                    all_rows.append([p.strip() for p in parts if p.strip()])

            # If no tables AND text layer is missing or garbage → ocrmypdf
            if not all_rows or not text_layer_usable:
                if not all_rows:
                    logger.info("PDF preview: pdfplumber found no tables, trying ocrmypdf")
                else:
                    logger.info("PDF preview: text layer looks like garbage, forcing ocrmypdf")
                    all_rows = []  # discard garbage rows
                ocr_tables = _ocrmypdf_then_extract_tables(content)
                if ocr_tables:
                    for tbl in ocr_tables:
                        all_rows.extend(tbl)

            if not all_rows:
                # Last resort: raw OCR via image_to_data, with preprocessing +
                # PSM auto-selection, restricted to pdf-inspector's flagged
                # pages when classification succeeded (whole document otherwise).
                logger.info("PDF preview: ocrmypdf returned nothing, trying raw pytesseract OCR (pages=%s)", _ocr_pages)
                from app.utils.pdf_ocr import ocr_pdf_to_rows_enhanced
                all_rows, ocr_error = ocr_pdf_to_rows_enhanced(content, pages=_ocr_pages)
                if not all_rows:
                    detail = ocr_error or "OCR не смог распознать таблицу."
                    raise HTTPException(
                        400,
                        f"Этот PDF — скан (изображение). {detail} "
                        "Попробуйте сохранить данные в Excel (.xlsx) или Word (.docx) "
                        "или конвертировать PDF→HTML через Adobe Acrobat."
                    )
            hdr_idx = _detect_hdr(all_rows)
            headers = [str(h).strip() if h else f"Столбец {j+1}" for j, h in enumerate(all_rows[hdr_idx])]
            data = all_rows[hdr_idx + 1:]
            sample = [[str(c) if c else "" for c in r] for r in data[:5]]
            return {"sheets": [{"name": "PDF", "headers": headers, "sample": sample, "total_rows": len(data), "header_row_offset": hdr_idx}]}

        # ── DOCX ──
        if fname.endswith(('.docx', '.doc')):
            try:
                from docx import Document as _DDoc
            except ImportError:
                raise HTTPException(500, "python-docx не установлен")
            doc = _DDoc(BytesIO(content))
            all_rows = []
            for table in doc.tables:
                for row in table.rows:
                    all_rows.append([cell.text.strip() for cell in row.cells])
            if not all_rows:
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_rows.append([text])
            if not all_rows:
                raise HTTPException(400, "Не удалось извлечь данные из документа")
            hdr_idx = _detect_hdr(all_rows)
            headers = [str(h).strip() if h else f"Столбец {j+1}" for j, h in enumerate(all_rows[hdr_idx])]
            data = all_rows[hdr_idx + 1:]
            sample = [[str(c) if c else "" for c in r] for r in data[:5]]
            return {"sheets": [{"name": "Document", "headers": headers, "sample": sample, "total_rows": len(data), "header_row_offset": hdr_idx}]}

        # ── HTML ──
        if fname.endswith(('.html', '.htm')):
            raw_tables = _extract_html_tables(content, file.filename or fname)
            if not raw_tables:
                raise HTTPException(400, "В HTML-файле не найдено таблиц с данными.")
            # Return ALL tables as separate "sheets" so user can pick which one to import.
            # IMPORTANT: numbering MUST match _extract_html_tables — same helper is used
            # in /items/import-mapped to look up the table by name "Таблица N".
            sheets_html = []
            for ti, tbl_rows in enumerate(raw_tables, start=1):
                if len(tbl_rows) < 2:
                    continue  # skip tables with no data
                hdr_idx = _detect_hdr(tbl_rows)
                headers = [str(h).strip() if h else f"Столбец {j+1}" for j, h in enumerate(tbl_rows[hdr_idx])]
                data = tbl_rows[hdr_idx + 1:]
                sample = [[str(c) if c else "" for c in r] for r in data[:5]]
                sheets_html.append({
                    "name": f"Таблица {ti}",
                    "headers": headers,
                    "sample": sample,
                    "total_rows": len(data),
                    "header_row_offset": hdr_idx,
                })
            if not sheets_html:
                raise HTTPException(400, "В HTML-файле не найдено таблиц с достаточным количеством строк.")
            return {"sheets": sheets_html}

        # ── Excel ──
        raw_sheets = _read_excel_rows(content, fname)
        sheets = []
        for si, sheet_rows in enumerate(raw_sheets):
            sheet_name = f"Лист{si+1}"
            all_rows = sheet_rows
            if not all_rows:
                continue
            hdr_idx = _detect_hdr(all_rows)
            hdr_rows = all_rows[hdr_idx:]
            if not hdr_rows:
                continue
            headers = [str(c).strip() if c else f"Столбец {j+1}" for j, c in enumerate(hdr_rows[0])]
            sample = [[str(c).strip() if c is not None else "" for c in row] for row in hdr_rows[1:min(6, len(hdr_rows))]]
            sheets.append({"name": sheet_name, "headers": headers, "sample": sample,
                           "total_rows": len(all_rows) - hdr_idx - 1, "header_row_offset": hdr_idx})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл ({file.filename}): {e}")

    if not sheets:
        raise HTTPException(400, "Файл не содержит листов с данными")

    return {"sheets": sheets}


@router.post("/items/import-mapped-nopid")
async def import_items_mapped_nopid(
    file: UploadFile = File(...),
    sheet_name: str = Query(""),
    col_item_name: int = Query(-1),
    col_description: int = Query(-1),
    col_quantity: int = Query(-1),
    col_unit_price: int = Query(-1),
    col_total_price: int = Query(-1),
    col_vat: int = Query(-1),
    col_unit: int = Query(-1),
    col_row_num: Optional[int] = Query(default=None),        # import-vat-cols: № строки (info only)
    col_vat_rate: Optional[int] = Query(default=None),       # import-vat-cols: ставка НДС
    col_vat_amount: Optional[int] = Query(default=None),     # import-vat-cols: сумма НДС
    col_total_with_vat: Optional[int] = Query(default=None), # import-vat-cols: стоимость с НДС
    col_category: Optional[int] = Query(default=None),       # Категория товара
    col_product_type: Optional[int] = Query(default=None),   # Вид товара
    header_row_offset: int = Query(0),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Like import-mapped but for new-purchase/wish context.

    Returns parsed items (не создаёт PurchaseItem), но товары upsert'ится в каталог:
    цена обновляется из файла, категория/вид берутся из файла только если в БД пусто
    (БД главнее), пишется import_note кто/как/когда загрузил."""
    if col_item_name < 0:
        raise HTTPException(400, "Не указан столбец Наименование")

    fname = (file.filename or '').lower()
    content = await file.read()

    try:
        if fname.endswith(('.docx', '.doc')):
            try:
                from docx import Document as _DDoc
            except ImportError:
                raise HTTPException(500, "python-docx не установлен")
            doc = _DDoc(BytesIO(content))
            all_rows_doc = []
            for table in doc.tables:
                for row in table.rows:
                    all_rows_doc.append(tuple(cell.text.strip() for cell in row.cells))
            if not all_rows_doc:
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_rows_doc.append((text,))
            skip = header_row_offset + 1
            data_iter = all_rows_doc[skip:] if len(all_rows_doc) > skip else []
        elif fname.endswith('.pdf'):
            try:
                import pdfplumber
            except ImportError:
                raise HTTPException(500, "pdfplumber не установлен")
            pdf = pdfplumber.open(BytesIO(content))
            all_rows_pdf = []
            for page in pdf.pages:
                for t in (page.extract_tables() or []):
                    if t:
                        all_rows_pdf.extend([tuple(str(c).strip() if c else "" for c in row) for row in t])
            pdf.close()
            skip = header_row_offset + 1
            data_iter = all_rows_pdf[skip:] if len(all_rows_pdf) > skip else []
        elif fname.endswith(('.html', '.htm')):
            raw_tables = _extract_html_tables(content, file.filename or fname)
            if not raw_tables:
                raise HTTPException(400, "В HTML не найдено таблиц")
            chosen_rows: list = []
            if sheet_name and sheet_name.lower().startswith('таблица'):
                try:
                    idx = int(sheet_name.split()[-1]) - 1
                    if 0 <= idx < len(raw_tables):
                        chosen_rows = raw_tables[idx]
                except (ValueError, IndexError):
                    pass
            if not chosen_rows:
                chosen_rows = max(raw_tables, key=len)
            all_rows_html = [tuple(row) for row in chosen_rows]
            skip = header_row_offset + 1
            data_iter = all_rows_html[skip:] if len(all_rows_html) > skip else []
        else:
            try:
                sheets = _read_excel_rows(content, fname)
            except Exception as e:
                raise HTTPException(400, f"Не удалось прочитать файл: {e}")
            if not sheets or not sheets[0]:
                raise HTTPException(400, "Файл пустой")
            sheet_idx = 0
            if sheet_name and sheet_name.lower().startswith('лист'):
                try:
                    sheet_idx = int(sheet_name[len('лист'):]) - 1
                except ValueError:
                    pass
            if sheet_idx < 0 or sheet_idx >= len(sheets):
                sheet_idx = 0
            all_rows = sheets[sheet_idx]
            skip = header_row_offset + 1
            data_iter = all_rows[skip:] if len(all_rows) > skip else []
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл ({file.filename}): {e}")

    def _cell(row, idx):
        if idx < 0 or idx >= len(row):
            return None
        v = row[idx]
        if v is None:
            return None
        s = str(v).strip()
        if not s or s.lower() in ('none', 'null', '-', '—', '0'):
            return None
        return s

    def _to_dec(v):
        if v is None:
            return None
        try:
            s = str(v).replace(',', '.').replace(' ', '').replace('\xa0', '')
            import re
            m = re.match(r'^([0-9]+\.?[0-9]*)', s)
            if not m:
                return None
            return Decimal(m.group(1))
        except Exception:
            return None

    _SKIP_KEYWORDS_NP = {
        'итого', 'всего', 'итог', 'total', 'подитог', 'subtotal',
        'поставщик', 'покупатель', 'заказчик', 'исполнитель',
        'генеральный директор', 'директор', 'бухгалтер', 'подпись',
        'м.п.', 'м.п', 'печать', 'ооо', 'оао', 'зао', 'ип ',
        'инн', 'кпп', 'огрн', 'р/с', 'к/с', 'бик',
        'адрес', 'телефон', 'email', 'банк',
        'примечание', 'основание', 'договор №', 'счёт №', 'счет №',
    }

    def _is_junk_row_np(name_val: str) -> bool:
        low = name_val.lower().strip()
        for kw in _SKIP_KEYWORDS_NP:
            if low.startswith(kw) or low == kw:
                return True
        if low.startswith('итого'):
            return True
        return False

    items_out = []
    skipped_empty = 0
    skipped_junk = 0

    _user_name = getattr(current_user, 'full_name', None) or getattr(current_user, 'username', '') or ''
    _import_note = (
        f"Импорт из файла «{file.filename}» (маппинг столбцов), "
        f"{_user_name}, {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )

    for row in data_iter:
        item_name = _cell(row, col_item_name)
        if not item_name:
            skipped_empty += 1
            continue
        if _is_junk_row_np(item_name):
            skipped_junk += 1
            continue
        description = _cell(row, col_description) if col_description >= 0 else None
        quantity = _to_dec(_cell(row, col_quantity)) if col_quantity >= 0 else None
        if not quantity:
            quantity = Decimal('1')
        unit_price = _to_dec(_cell(row, col_unit_price)) if col_unit_price >= 0 else None
        total_price = _to_dec(_cell(row, col_total_price)) if col_total_price >= 0 else None
        unit = (_cell(row, col_unit) if col_unit >= 0 else None) or 'шт'
        if not total_price and unit_price:
            total_price = quantity * unit_price
        elif not unit_price and total_price and quantity:
            unit_price = total_price / quantity
        vat_str = _cell(row, col_vat) if col_vat >= 0 else None
        if vat_str and description:
            description = f"{description} (НДС: {vat_str})"
        elif vat_str:
            description = f"НДС: {vat_str}"
        # import-vat-cols: новые НДС-поля
        vat_rate_str = _cell(row, col_vat_rate) if (col_vat_rate is not None and col_vat_rate >= 0) else None
        vat_amount_dec = _to_dec(_cell(row, col_vat_amount)) if (col_vat_amount is not None and col_vat_amount >= 0) else None
        total_with_vat_dec = _to_dec(_cell(row, col_total_with_vat)) if (col_total_with_vat is not None and col_total_with_vat >= 0) else None

        row_category = _cell(row, col_category) if (col_category is not None and col_category >= 0) else None
        row_product_type = _cell(row, col_product_type) if (col_product_type is not None and col_product_type >= 0) else None

        # Upsert в каталог: цена из файла, категория/вид — БД главнее, примечание об импорте
        product_id = None
        eff_category, eff_product_type = row_category, row_product_type
        try:
            product_id = await _upsert_product_to_catalog(
                db, item_name, 'товар', unit_price, description or "",
                category=row_category, product_type=row_product_type,
                import_note=_import_note, updated_by=_user_name,
            )
            prod = await db.get(Product, product_id)
            if prod:
                eff_category = prod.category
                eff_product_type = prod.product_type
        except Exception:
            logging.getLogger(__name__).warning("nopid import: upsert to catalog failed for %r", item_name, exc_info=True)

        items_out.append({
            'item_name': item_name[:500],
            'item_type': 'товар',
            'description': description,
            'quantity': float(quantity) if quantity else None,
            'unit': unit,
            'unit_price': float(unit_price) if unit_price else None,
            'total_price': float(total_price) if total_price else None,
            'vat_rate': vat_rate_str or (vat_str if not description else None),
            'vat_amount': float(vat_amount_dec) if vat_amount_dec else None,
            'total_with_vat': float(total_with_vat_dec) if total_with_vat_dec else None,
            'product_id': product_id,
            'category': eff_category,
            'product_type': eff_product_type,
        })

    try:
        await db.commit()
    except Exception:
        await db.rollback()

    return {
        "items": items_out,
        "added": len(items_out),
        "debug": {
            "total_rows_after_header": len(data_iter),
            "skipped_empty_name": skipped_empty,
            "skipped_junk_row": skipped_junk,
        },
    }


@router.post("/{pid}/items/import-mapped")
async def import_items_mapped(
    pid: int,
    file: UploadFile = File(...),
    sheet_name: str = Query(""),
    col_item_name: int = Query(-1, description="Индекс столбца Наименование (0-based)"),
    col_description: int = Query(-1, description="Индекс столбца Описание"),
    col_quantity: int = Query(-1, description="Индекс столбца Количество"),
    col_unit_price: int = Query(-1, description="Индекс столбца Цена"),
    col_total_price: int = Query(-1, description="Индекс столбца Сумма"),
    col_vat: int = Query(-1, description="Индекс столбца НДС"),
    col_unit: int = Query(-1, description="Индекс столбца Ед. изм."),
    col_row_num: Optional[int] = Query(default=None, description="Индекс столбца № строки (info only)"),
    col_vat_rate: Optional[int] = Query(default=None, description="Индекс столбца Ставка НДС"),
    col_vat_amount: Optional[int] = Query(default=None, description="Индекс столбца Сумма НДС"),
    col_total_with_vat: Optional[int] = Query(default=None, description="Индекс столбца Стоимость с НДС"),
    col_category: Optional[int] = Query(default=None, description="Индекс столбца Категория товара"),
    col_product_type: Optional[int] = Query(default=None, description="Индекс столбца Вид товара"),
    header_row_offset: int = Query(0, description="Сколько строк пропустить до заголовка (авто-определено при preview)"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Import items using user-specified column mapping."""
    if col_item_name < 0:
        raise HTTPException(400, "Не указан столбец Наименование")

    purchase = await db.get(Purchase, pid)
    if not purchase:
        raise HTTPException(404, "Закупка не найдена")

    fname = (file.filename or '').lower()
    content = await file.read()

    try:
        if fname.endswith(('.docx', '.doc')):
            # Word document — extract table rows
            try:
                from docx import Document as _DDoc
            except ImportError:
                raise HTTPException(500, "python-docx не установлен")
            doc = _DDoc(BytesIO(content))
            all_rows_doc = []
            for table in doc.tables:
                for row in table.rows:
                    all_rows_doc.append(tuple(cell.text.strip() for cell in row.cells))
            if not all_rows_doc:
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_rows_doc.append((text,))
            skip = header_row_offset + 1
            data_iter = all_rows_doc[skip:] if len(all_rows_doc) > skip else []
        elif fname.endswith('.pdf'):
            # PDF — extract table rows
            try:
                import pdfplumber
            except ImportError:
                raise HTTPException(500, "pdfplumber не установлен")
            pdf = pdfplumber.open(BytesIO(content))
            all_rows_pdf = []
            for page in pdf.pages:
                for t in (page.extract_tables() or []):
                    if t:
                        all_rows_pdf.extend([tuple(str(c).strip() if c else "" for c in row) for row in t])
            pdf.close()
            skip = header_row_offset + 1
            data_iter = all_rows_pdf[skip:] if len(all_rows_pdf) > skip else []
        elif fname.endswith(('.html', '.htm')):
            # Use SAME extraction as preview so "Таблица N" indices match.
            raw_tables = _extract_html_tables(content, file.filename or fname)
            if not raw_tables:
                raise HTTPException(400, "В HTML не найдено таблиц")
            chosen_rows: list = []
            if sheet_name and sheet_name.lower().startswith('таблица'):
                try:
                    idx = int(sheet_name.split()[-1]) - 1
                    if 0 <= idx < len(raw_tables):
                        chosen_rows = raw_tables[idx]
                except (ValueError, IndexError):
                    pass
            if not chosen_rows:
                # Fallback: pick the table with the most rows
                chosen_rows = max(raw_tables, key=len)
            # Convert to tuples for uniform downstream handling
            all_rows_html = [tuple(row) for row in chosen_rows]
            skip = header_row_offset + 1
            data_iter = all_rows_html[skip:] if len(all_rows_html) > skip else []
        else:
            try:
                sheets = _read_excel_rows(content, fname)
            except Exception as e:
                raise HTTPException(400, f"Не удалось прочитать файл: {e}")
            if not sheets or not sheets[0]:
                raise HTTPException(400, "Файл пустой")
            # sheet_name здесь — числовой индекс в виде "ЛистN" или первый лист
            sheet_idx = 0
            if sheet_name and sheet_name.lower().startswith('лист'):
                try:
                    sheet_idx = int(sheet_name[len('лист'):]) - 1
                except ValueError:
                    pass
            if sheet_idx < 0 or sheet_idx >= len(sheets):
                sheet_idx = 0
            all_rows = sheets[sheet_idx]
            skip = header_row_offset + 1
            data_iter = all_rows[skip:] if len(all_rows) > skip else []
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл ({file.filename}): {e}")

    def _cell(row, idx):
        if idx < 0 or idx >= len(row):
            return None
        v = row[idx]
        if v is None:
            return None
        s = str(v).strip()
        if not s or s.lower() in ('none', 'null', '-', '—', '0'):
            return None
        return s

    def _to_dec(v):
        if v is None:
            return None
        try:
            s = str(v).replace(',', '.').replace(' ', '').replace('\xa0', '')
            # Strip non-numeric suffix (e.g. "руб.", "шт.", "р.")
            import re
            m = re.match(r'^([0-9]+\.?[0-9]*)', s)
            if not m:
                return None
            return Decimal(m.group(1))
        except Exception:
            return None

    # Load products for auto-matching
    org_id = get_single_org_id(current_user)
    prod_q = select(Product)
    if org_id:
        prod_q = prod_q.where((Product.org_id == org_id) | (Product.org_id.is_(None)))
    prod_result = await db.execute(prod_q)
    products = prod_result.scalars().all()
    product_by_name: dict[str, Product] = {}
    for p in products:
        if p.name:
            product_by_name[p.name.lower().strip()] = p

    added = 0
    matched_catalog = 0
    new_in_catalog = 0
    errors_list = []
    skipped_empty = 0      # row[col_item_name] is None/empty
    skipped_junk = 0       # _is_junk_row matched
    total_data_rows = 0    # счётчик прошедших data_iter

    _user_name = getattr(current_user, 'full_name', None) or getattr(current_user, 'username', '') or ''
    _import_note = (
        f"Импорт из файла «{file.filename}» (маппинг столбцов в закупке), "
        f"{_user_name}, {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )

    # Keywords that indicate non-product rows (totals, footers, signatures)
    _SKIP_KEYWORDS = {
        'итого', 'всего', 'итог', 'total', 'подитог', 'subtotal',
        'поставщик', 'покупатель', 'заказчик', 'исполнитель',
        'генеральный директор', 'директор', 'бухгалтер', 'подпись',
        'м.п.', 'м.п', 'печать', 'ооо', 'оао', 'зао', 'ип ',
        'инн', 'кпп', 'огрн', 'р/с', 'к/с', 'бик',
        'адрес', 'телефон', 'email', 'банк',
        'примечание', 'основание', 'договор №', 'счёт №', 'счет №',
    }

    def _is_junk_row(name_val: str) -> bool:
        """Check if this looks like a footer/total/signature row, not a product."""
        low = name_val.lower().strip()
        # Direct match with skip keywords
        for kw in _SKIP_KEYWORDS:
            if low.startswith(kw) or low == kw:
                return True
        # Row starts with "итого" variants like "Итого с НДС:", "Итого:"
        if low.startswith('итого'):
            return True
        return False

    for row_idx, row in enumerate(data_iter):
        try:
            total_data_rows += 1
            item_name = _cell(row, col_item_name)
            if not item_name:
                skipped_empty += 1
                continue

            # Skip junk rows (totals, footers, signatures)
            if _is_junk_row(item_name):
                skipped_junk += 1
                continue

            description = _cell(row, col_description) if col_description >= 0 else None
            quantity = _to_dec(_cell(row, col_quantity)) if col_quantity >= 0 else Decimal('1')
            if not quantity:
                quantity = Decimal('1')
            unit_price = _to_dec(_cell(row, col_unit_price)) if col_unit_price >= 0 else None
            total_price = _to_dec(_cell(row, col_total_price)) if col_total_price >= 0 else None
            unit = (_cell(row, col_unit) if col_unit >= 0 else None) or 'шт'

            # Calculate missing values
            if not total_price and unit_price:
                total_price = quantity * unit_price
            elif not unit_price and total_price and quantity:
                unit_price = total_price / quantity

            # Шаг 5 «цена ТЗ не выше плановой» (владелец, 2026-08-07): позиция
            # импорта наследует ФЭО-категорию закупки (feo_planned_item_id импорт
            # не проставляет). Ловим ИМЕННО эту ошибку до общего `except Exception`
            # ниже — тот делает db.rollback(), который стёр бы уже добавленные
            # (ещё не закоммиченные) строки предыдущих итераций; здесь просто
            # пропускаем строку и продолжаем — агрегация ошибок по строкам.
            try:
                await assert_tz_not_over_plan(
                    db,
                    feo_planned_item_id=None,
                    feo_category_id=purchase.feo_category_id,
                    quantity=quantity,
                    unit_price=unit_price,
                    total_price=total_price,
                    item_name=item_name,
                )
            except HTTPException as _tz_exc:
                errors_list.append(f"Строка {row_idx + 1}: {_tz_exc.detail}")
                continue

            # VAT info → append to description
            vat_str = _cell(row, col_vat) if col_vat >= 0 else None
            if vat_str and description:
                description = f"{description} (НДС: {vat_str})"
            elif vat_str:
                description = f"НДС: {vat_str}"

            # import-vat-cols: новые НДС-поля
            vat_rate_val = _cell(row, col_vat_rate) if (col_vat_rate is not None and col_vat_rate >= 0) else None
            vat_amount_val = _to_dec(_cell(row, col_vat_amount)) if (col_vat_amount is not None and col_vat_amount >= 0) else None
            total_with_vat_val = _to_dec(_cell(row, col_total_with_vat)) if (col_total_with_vat is not None and col_total_with_vat >= 0) else None

            row_category = _cell(row, col_category) if (col_category is not None and col_category >= 0) else None
            row_product_type = _cell(row, col_product_type) if (col_product_type is not None and col_product_type >= 0) else None

            # Auto-match or create in catalog
            # 1) exact match (fast path)
            matched_product = product_by_name.get(item_name.lower().strip())
            if not matched_product:
                # 2) fuzzy fallback — find best candidate above SCORE_AUTO threshold
                best_score = 0.0
                best_candidate = None
                for _key, _p in product_by_name.items():
                    _s = _fuzzy_score(item_name, _p.name if hasattr(_p, 'name') else _key)
                    if _s > best_score:
                        best_score = _s
                        best_candidate = _p
                if best_score >= _SCORE_AUTO and best_candidate is not None:
                    matched_product = best_candidate
            if matched_product:
                product_id = matched_product.id
                matched_catalog += 1
                if not unit_price and matched_product.price:
                    unit_price = matched_product.price
                    total_price = quantity * unit_price
                elif unit_price and isinstance(matched_product, Product):
                    # Цена из файла актуальнее; категория/вид из БД не трогаем (БД главнее)
                    if matched_product.price != unit_price:
                        matched_product.price = unit_price
                    matched_product.import_note = _import_note
                    matched_product.updated_at = datetime.utcnow()
                    matched_product.updated_by = _user_name
                    if row_category and (not matched_product.category or matched_product.category == 'Прочее'):
                        matched_product.category = row_category
                    if row_product_type and not matched_product.product_type:
                        matched_product.product_type = row_product_type
            else:
                product_id = await _upsert_product_to_catalog(
                    db, item_name, 'товар', unit_price, description or "",
                    category=row_category, product_type=row_product_type,
                    import_note=_import_note, updated_by=_user_name,
                )
                product_by_name[item_name.lower().strip()] = type('_P', (), {'id': product_id, 'name': item_name, 'price': unit_price})()
                new_in_catalog += 1

            item = PurchaseItem(
                purchase_id=pid,
                product_id=product_id,
                item_name=item_name,
                item_type='товар',
                quantity=quantity,
                unit=unit,
                unit_price=unit_price,
                total_price=total_price,
                vat_rate=vat_rate_val or vat_str,
                vat_amount=vat_amount_val,
                total_with_vat=total_with_vat_val,
            )
            db.add(item)
            added += 1
        except Exception as e:
            errors_list.append(f"Строка {row_idx + 1}: {e}")
            await db.rollback()
            continue

    try:
        await db.commit()
    except Exception as e:
        raise HTTPException(500, f"Ошибка сохранения: {e}")
    return {
        "added": added,
        "matched_catalog": matched_catalog,
        "new_in_catalog": new_in_catalog,
        "errors": errors_list,
        "debug": {
            "total_rows_after_header": len(data_iter),
            "rows_processed": total_data_rows,
            "skipped_empty_name": skipped_empty,
            "skipped_junk_row": skipped_junk,
            "first_3_rows_sample": [list(r)[:8] for r in data_iter[:3]],
        },
    }


# ---------------------------------------------------------------------------
# Image import helpers (Phase 26-PP)
# ---------------------------------------------------------------------------

def _try_decode_qr(image_bytes: bytes) -> str | None:
    """Попытаться декодировать QR-код из изображения.

    Сначала pyzbar, потом cv2.QRCodeDetector. Возвращает строку QR или None.
    Все импорты в try/except — если библиотек нет, вернуть None без ошибок.
    """
    # pyzbar
    try:
        from pyzbar.pyzbar import decode as _pyzbar_decode
        from PIL import Image as _PILImage
        import io as _io
        img = _PILImage.open(_io.BytesIO(image_bytes))
        decoded = _pyzbar_decode(img)
        for d in decoded:
            data = d.data
            if isinstance(data, bytes):
                data = data.decode("utf-8", errors="replace")
            if data:
                return data
    except Exception:
        pass

    # cv2 fallback
    try:
        import cv2 as _cv2
        import numpy as _np
        arr = _np.frombuffer(image_bytes, dtype=_np.uint8)
        img_cv = _cv2.imdecode(arr, _cv2.IMREAD_COLOR)
        if img_cv is not None:
            detector = _cv2.QRCodeDetector()
            data, _, _ = detector.detectAndDecode(img_cv)
            if data:
                return data
    except Exception:
        pass

    return None


def _smart_import_image_ocr(content: bytes, filename: str) -> dict:
    """Stage 1 OCR для изображений: tesseract → regex-парсинг строк чека.

    Возвращает preview-словарь (без сохранения в БД).
    Бросает HTTPException если OCR пустой или ничего не распознано.
    """
    import re

    # Открыть изображение через Pillow
    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(400, f"Не удалось открыть изображение: {e}")

    # OCR
    try:
        import pytesseract
        ocr_text = pytesseract.image_to_string(img, lang="rus+eng")
    except Exception as e:
        raise HTTPException(
            400,
            f"Ошибка OCR (tesseract установлен?): {e}. "
            "Попробуйте QR-чек ФНС или загрузите Excel/PDF."
        )

    if not ocr_text or not ocr_text.strip():
        raise HTTPException(
            400,
            "На изображении не распознан текст. "
            "Попробуйте QR-чек ФНС или загрузите Excel/PDF с позициями."
        )

    # Эвристика: каждая строка — потенциальная позиция
    # Паттерны: «Название  N шт  X руб», «Название  X руб», «1  Название  N  X  Y»
    items = []
    lines = [ln.strip() for ln in ocr_text.splitlines() if ln.strip()]

    # Паттерн: число-разделитель-название-число(кол-во)-число(цена)-число(сумма)
    # Или проще: ищем строки с хотя бы одним числом-рублём
    PRICE_RE = re.compile(
        r"(?P<name>.+?)\s+"
        r"(?:(?P<qty>\d+(?:[.,]\d+)?)\s*(?:шт|кг|л|м|уп|уп\.|ед|ед\.)\s*)?"
        r"(?:x|х|×)?\s*"
        r"(?P<price>\d[\d\s]*(?:[.,]\d{1,2})?)\s*(?:руб|₽|р\.?)?",
        re.IGNORECASE | re.UNICODE,
    )
    SKIP_WORDS = re.compile(
        r"^(итого|total|сумма|кассир|кассa|чек|receipt|дата|дата:|inn|инн|"
        r"тел|телефон|спасибо|магазин|адрес|режим|номер|фн:|фд:|фпд:|"
        r"кку|ккт|ооо|ип |ооо )",
        re.IGNORECASE,
    )

    for line in lines:
        if SKIP_WORDS.match(line):
            continue
        m = PRICE_RE.search(line)
        if not m:
            continue
        name = m.group("name").strip(" -–•·")
        # Слишком короткое название — пропустить
        if len(name) < 3:
            continue
        qty_raw = m.group("qty")
        price_raw = m.group("price")
        try:
            price_clean = price_raw.replace(" ", "").replace(",", ".")
            price_val = float(price_clean)
        except Exception:
            price_val = None
        try:
            qty_val = float((qty_raw or "1").replace(",", "."))
        except Exception:
            qty_val = 1.0

        items.append({
            "item_name": name[:500],
            "item_type": "товар",
            "quantity": qty_val,
            "unit": "шт",
            "unit_price": price_val,
            "total_price": round(price_val * qty_val, 2) if price_val else None,
        })

    if not items:
        raise HTTPException(
            400,
            "На изображении текст распознан, но позиции товаров не обнаружены. "
            "Качество OCR для чеков невысокое — рекомендуем использовать QR-код ФНС "
            "(приложение «Проверка чека ФНС России») или загрузить Excel/PDF."
        )

    return {
        "preview": items,
        "total_rows": len(items),
        "file_type": "image_ocr",
        "columns_found": ["item_name", "quantity", "unit_price", "total_price"],
        "warning": (
            "Распознавание через OCR имеет ограниченную точность. "
            "Рекомендуется проверить позиции перед сохранением. "
            "Для лучшего результата используйте QR-код ФНС на чеке."
        ),
    }


def _looks_mojibake(wb) -> bool:
    """Эвристика: если в первом листе ≥50% строковых ячеек состоят целиком из latin-1 акцент-символов
    (U+00C0..U+00FF) либо содержат U+FFFD — считаем что cp1251 декодировалась как latin-1."""
    try:
        ws = wb.sheet_by_index(0)
    except Exception:
        return False
    total = 0
    bad = 0
    for r in range(min(ws.nrows, 20)):
        for v in ws.row_values(r):
            if not isinstance(v, str) or len(v.strip()) < 2:
                continue
            total += 1
            s = v.strip()
            if '\ufffd' in s:
                bad += 1
                continue
            # доля символов в диапазоне latin-1 supplement (0xC0..0xFF) — типичный признак cp1251→latin1
            latin_supp = sum(1 for ch in s if 0x00C0 <= ord(ch) <= 0x00FF)
            if latin_supp >= max(2, len(s) // 2):
                bad += 1
    return total > 0 and bad / total >= 0.5


def _read_excel_rows(content: bytes, fname: str) -> list[list[list]]:
    """Универсальное чтение Excel: возвращает list[sheets], каждый sheet = list[rows], row = list[cells].
    Поддерживает .xlsx (openpyxl) и .xls BIFF8 (xlrd с авто-cp1251-override при mojibake)."""
    is_xls = (fname or '').lower().endswith('.xls') or content[:4] == b'\xd0\xcf\x11\xe0'
    if is_xls:
        import xlrd as _xlrd
        wb = _xlrd.open_workbook(file_contents=content, formatting_info=False)
        if _looks_mojibake(wb):
            wb = _xlrd.open_workbook(file_contents=content, encoding_override='cp1251')
        sheets = []
        for si in range(wb.nsheets):
            ws = wb.sheet_by_index(si)
            sheets.append([list(ws.row_values(r)) for r in range(ws.nrows)])
        return sheets
    if load_workbook is None:
        raise RuntimeError("openpyxl не установлен")
    wb = load_workbook(BytesIO(content), read_only=False, data_only=True)
    return [[list(r) for r in ws.iter_rows(values_only=True)] for ws in wb.worksheets]


def _smart_import_xlsx_direct(content: bytes, fname: str = '') -> tuple[list[dict], list[str]]:
    """Direct XLSX/XLS parser without markitdown — устойчив к опечаткам в заголовке,
    разделам-подзаголовкам в середине, multi-line cells, merged headers.
    Поддерживает .xls BIFF8 с авто-cp1251-override при mojibake.

    Returns (preview_rows, columns_found).
    Каждый dict в preview_rows: item_name, item_type, quantity, unit, unit_price, total_price.
    """
    import re as _re

    # Header keywords (substring match, case-insensitive, tolerant к опечаткам через 'in')
    # Шаг 5 (владелец, 2026-08-07): 'unit_price' проверяется РАНЬШЕ 'unit' —
    # заголовок «Цена за единицу» содержит подстроку «единиц», которая раньше
    # матчилась ключом 'unit' первой (dict сохраняет порядок вставки, matching
    # идёт по порядку ключей) — колонка цены никогда не находилась, unit_price/
    # total_price молча оставались None. Тот же баг, что и в _detect_columns_legacy
    # (см. комментарий там) — независимая копия той же логики (проект уже
    # предупреждал о дублировании normalize/tokenize, см. план шаг 4).
    HEADER_PATTERNS = {
        'item_name': ['наимен', 'товар', 'позици', 'описан', 'материал', 'предмет'],
        'item_type': ['тип', 'вид'],
        'quantity': ['колич', 'кол-во', 'кол.', 'кол ', 'кол.во', 'qty'],  # ловит "Количечество"
        'unit_price': ['цена за ед', 'цена ед', 'цена/ед', 'цена', 'unit price', 'стоимость ед'],
        'unit': ['ед.из', 'ед. изм', 'едизм', 'единиц', 'unit'],
        'total_price': ['сумма', 'итого', 'стоимость', 'total'],
    }

    try:
        sheets = _read_excel_rows(content, fname)
    except Exception:
        return [], []
    all_rows: list[list] = [r for sh in sheets for r in sh]

    def _classify_header(row: list) -> dict:
        """Return dict {col_idx -> field_key} для cells матчащихся к HEADER_PATTERNS."""
        mapping: dict[int, str] = {}
        for idx, cell in enumerate(row):
            if cell is None:
                continue
            text = str(cell).lower().strip()
            if not text:
                continue
            for field, patterns in HEADER_PATTERNS.items():
                if any(p in text for p in patterns):
                    if field not in mapping.values():  # первый match выигрывает
                        mapping[idx] = field
                    break
        return mapping

    # Выбираем строку с наибольшим числом распознанных колонок (≥3 полей лучше 2)
    # — защита от ложного срабатывания на строки-итоги типа «Товар по листу…»
    best_idx = -1
    best_map: dict[int, str] = {}
    best_score = 0
    for i, row in enumerate(all_rows):
        mapping = _classify_header(row)
        if 'item_name' not in mapping.values() or len(mapping) < 2:
            continue
        score = len(mapping)
        if score > best_score:
            best_score = score
            best_idx = i
            best_map = mapping
    header_idx = best_idx
    col_map = best_map

    if header_idx == -1:
        return [], []

    # Inverse map: field -> col_idx
    field_to_idx = {v: k for k, v in col_map.items()}
    columns_found = list(field_to_idx.keys())

    def _to_dec(v):
        if v is None:
            return None
        if isinstance(v, (int, float)):
            return Decimal(str(v))
        try:
            s = str(v).strip().replace(',', '.').replace(' ', '').replace('\xa0', '')
            if not s or s in ('-', '—', '–'):
                return None
            return Decimal(s)
        except Exception:
            return None

    def _get_cell(row: list, field: str):
        idx = field_to_idx.get(field)
        if idx is None or idx >= len(row):
            return None
        v = row[idx]
        return v if v not in (None, '') else None

    TYPE_MAP = {
        'товар': 'товар', 'товары': 'товар',
        'услуга': 'услуга', 'услуги': 'услуга',
        'работа': 'работа', 'работы': 'работа',
    }

    preview: list[dict] = []
    for row in all_rows[header_idx + 1:]:
        name_val = _get_cell(row, 'item_name')
        if name_val is None:
            continue
        name = str(name_val).strip().replace('\n', ' ').replace('\r', ' ')
        name = _re.sub(r'\s+', ' ', name)
        if not name:
            continue
        # Пропускаем разделы-подзаголовки: нет ни qty, ни цены, ни суммы
        qty = _to_dec(_get_cell(row, 'quantity'))
        unit_price = _to_dec(_get_cell(row, 'unit_price'))
        total_price = _to_dec(_get_cell(row, 'total_price'))
        if qty is None and unit_price is None and total_price is None:
            continue
        # Вычисляем недостающее
        if unit_price is None and total_price is not None and qty:
            try:
                unit_price = total_price / qty
            except Exception:
                pass
        if total_price is None and unit_price is not None and qty:
            total_price = unit_price * qty
        unit_val = _get_cell(row, 'unit')
        unit = str(unit_val).strip() if unit_val else 'шт'
        type_val = _get_cell(row, 'item_type')
        item_type = TYPE_MAP.get(str(type_val).lower().strip() if type_val else '', 'товар')
        preview.append({
            'item_name': name,
            'item_type': item_type,
            'quantity': float(qty) if qty else None,
            'unit': unit,
            'unit_price': float(unit_price) if unit_price else None,
            'total_price': float(total_price) if total_price else None,
        })
    return preview, columns_found


async def _save_smart_preview_to_purchase(
    pid: int,
    preview: list[dict],
    purchase,
    db: AsyncSession,
    current_user,
    skip_catalog: bool = False,
) -> dict:
    """Сохраняет preview-строки в БД как PurchaseItem'ы.
    Переиспользуется как из xlsx-ветки, так и (потенциально) из markitdown-ветки.
    skip_catalog=True: не вызывать _upsert для несматченных → product_id=None.
    """
    org_id = get_single_org_id(current_user)
    prod_q = select(Product)
    if org_id:
        prod_q = prod_q.where((Product.org_id == org_id) | (Product.org_id.is_(None)))
    products = (await db.execute(prod_q)).scalars().all()
    product_by_name = {(p.name or "").lower().strip(): p for p in products}

    added = matched_catalog = new_in_catalog = 0
    errors_list: list[str] = []
    for row_idx, row_data in enumerate(preview, start=1):
        item_name = (row_data["item_name"] or "")[:500]
        qty = Decimal(str(row_data["quantity"])) if row_data["quantity"] else Decimal("1")
        unit_price = Decimal(str(row_data["unit_price"])) if row_data["unit_price"] else None
        total_price = Decimal(str(row_data["total_price"])) if row_data["total_price"] else None
        # Шаг 5 «цена ТЗ не выше плановой» (владелец, 2026-08-07): позиция смарт-
        # импорта наследует ФЭО-категорию закупки (feo_planned_item_id импорт не
        # проставляет). Аггрегация ошибок по строкам — строка пропускается, импорт
        # остальных продолжается.
        try:
            await assert_tz_not_over_plan(
                db,
                feo_planned_item_id=None,
                feo_category_id=getattr(purchase, "feo_category_id", None),
                quantity=qty,
                unit_price=unit_price,
                total_price=total_price if total_price is not None else (qty * unit_price if unit_price else None),
                item_name=item_name,
            )
        except HTTPException as _tz_exc:
            errors_list.append(f"Строка {row_idx}: {_tz_exc.detail}")
            continue
        if skip_catalog:
            # «Не добавлять в каталог» (напр. авансовые платежи): позиции должны быть
            # один-в-один как в чеке, без какой-либо привязки к каталогу. Не матчим вовсе.
            product_id = None
            matched = None
        else:
            # 1) exact match (fast path)
            matched = product_by_name.get(item_name.lower().strip())
            if not matched:
                # 2) fuzzy fallback
                best_score = 0.0
                best_candidate = None
                for _key, _p in product_by_name.items():
                    _s = _fuzzy_score(item_name, _p.name if hasattr(_p, 'name') else _key)
                    if _s > best_score:
                        best_score = _s
                        best_candidate = _p
                if best_score >= _SCORE_AUTO and best_candidate is not None:
                    matched = best_candidate
            if matched:
                product_id = matched.id
                matched_catalog += 1
                if not unit_price and matched.price:
                    unit_price = matched.price
                    total_price = qty * unit_price
            else:
                _uname = getattr(current_user, 'full_name', None) or getattr(current_user, 'username', '') or ''
                product_id = await _upsert_product_to_catalog(
                    db, item_name, row_data["item_type"], unit_price,
                    import_note=f"Смарт-импорт из файла, {_uname}, {datetime.now().strftime('%d.%m.%Y %H:%M')}",
                    updated_by=_uname,
                )
                new_in_catalog += 1
        if total_price is None and unit_price:
            total_price = qty * unit_price
        db.add(PurchaseItem(
            purchase_id=pid, product_id=product_id,
            item_name=item_name, item_type=row_data["item_type"],
            quantity=qty, unit=row_data["unit"],
            unit_price=unit_price, total_price=total_price,
        ))
        added += 1
    await db.commit()
    return {"ok": True, "added": added, "matched_catalog": matched_catalog, "new_in_catalog": new_in_catalog, "errors": errors_list}


@router.post("/items/import-smart-nopid")
async def import_items_smart_nopid(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    """27.4-26b: XLSX preview БЕЗ purchaseId — для wish / новой закупки.
    Возвращает все позиции через direct openpyxl парсер (не sample-only)."""
    fname = (file.filename or "").lower()
    if not fname.endswith((".xlsx", ".xls")):
        raise HTTPException(400, "Этот endpoint только для XLSX/XLS. Используйте /import-preview для других форматов.")
    content = await file.read()
    try:
        preview, columns = _smart_import_xlsx_direct(content, fname=fname)
    except Exception as e:
        logger.warning("import-smart-nopid failed: %s", e)
        ext = '.xls' if fname.endswith('.xls') else '.xlsx'
        raise HTTPException(400, f"Не удалось распознать файл ({ext}): {e}. Если файл .xls — попробуйте сохранить его как .xlsx (Excel: Файл → Сохранить как → Книга Excel) и повторите.")
    return {
        "preview": preview[:200],
        "total_rows": len(preview),
        "file_type": "excel",
        "columns_found": columns,
    }


@router.post("/{pid}/items/import-smart")
async def import_items_smart(
    pid: int,
    file: UploadFile = File(...),
    confirm: bool = Query(default=False),
    skip_catalog: bool = Query(default=False, description="Не добавлять несматченные позиции в каталог"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Smart import: extract items table from PDF / DOCX / XLSX."""
    purchase = await db.get(Purchase, pid)
    if not purchase:
        raise HTTPException(404, "Закупка не найдена")
    # Employees can import to any purchase they have access to
    if not await _has_purchase_write_access(current_user, db):
        raise HTTPException(403, "Нет прав импортировать позиции в эту закупку.")

    content = await file.read()
    filename = (file.filename or "").lower()

    IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif")
    allowed_ext = (".pdf", ".docx", ".doc", ".xlsx", ".xls", ".html", ".htm") + IMAGE_EXTS
    if not any(filename.endswith(ext) for ext in allowed_ext):
        raise HTTPException(400, "Поддерживаются файлы: PDF, DOCX, XLSX/XLS, HTML, JPG/PNG (фото чека)")

    # ── Image pipeline ──────────────────────────────────────────────────────────
    if any(filename.endswith(ext) for ext in IMAGE_EXTS):
        # HEIC: graceful error if pillow-heif not available
        if filename.endswith((".heic", ".heif")):
            try:
                import pillow_heif  # noqa: F401
                pillow_heif.register_heif_opener()
            except ImportError:
                raise HTTPException(
                    400,
                    "Формат HEIC временно не поддерживается. "
                    "Пожалуйста, конвертируйте изображение в JPG перед загрузкой."
                )

        # Stage 0: попробовать QR-декодирование (чек ФНС)
        qr_str = _try_decode_qr(content)
        if qr_str:
            # Делегируем на proverkacheka pipeline (те же helpers что from-qr-fetch)
            from app.routers.purchase_receipts import (
                _create_receipt_with_items as _r_create_receipt,
                _parse_fns_json_receipt,
                _parse_qr_string,
            )
            import os as _os
            import httpx as _httpx
            token = _os.getenv("PROVERKACHEKA_TOKEN", "").strip()
            if not token:
                raise HTTPException(
                    503,
                    "QR-код найден, но PROVERKACHEKA_TOKEN не настроен. "
                    "Обратитесь к администратору или импортируйте Excel/PDF."
                )
            try:
                async with _httpx.AsyncClient(timeout=30.0) as _client:
                    _resp = await _client.post(
                        "https://proverkacheka.com/api/v1/check/get",
                        data={"qrraw": qr_str, "token": token},
                    )
                    _resp.raise_for_status()
                    _payload = _resp.json()
            except Exception as exc:
                raise HTTPException(502, f"Ошибка запроса к proverkacheka.com: {exc}")
            if not isinstance(_payload, dict) or _payload.get("code") != 1:
                _msg = (_payload or {}).get("data", "") if isinstance(_payload, dict) else ""
                raise HTTPException(400, f"Чек не найден в ФНС: {_msg or 'неизвестная ошибка'}")
            _body = _payload.get("data") or {}
            _receipt_obj = _body.get("json") if isinstance(_body, dict) else None
            if not isinstance(_receipt_obj, dict):
                raise HTTPException(502, "Ответ proverkacheka.com без данных чека")
            _data = _parse_fns_json_receipt(_receipt_obj)
            if not _data.get("fiscal_drive_number"):
                _qr_data = _parse_qr_string(qr_str)
                for _k, _v in _qr_data.items():
                    if _v and not _data.get(_k):
                        _data[_k] = _v
            receipt = await _r_create_receipt(
                pid, _data, "qr_scan", {"qr": qr_str, "proverkacheka": _payload}, db
            )
            return {
                "ok": True,
                "source": "qr_fns",
                "receipt_id": receipt.id,
                "message": "Чек успешно импортирован по QR-коду ФНС. Позиции добавлены в закупку.",
            }

        # Stage 1: OCR через tesseract
        return _smart_import_image_ocr(content, filename)

    if filename.endswith(".pdf"):
        file_type = "pdf"
    elif filename.endswith((".xlsx", ".xls")):
        file_type = "excel"
    elif filename.endswith((".html", ".htm")):
        file_type = "html"
    else:
        file_type = "docx"

    # 27.4-26: для XLSX обходим markitdown — direct openpyxl парсинг устойчивее
    # к опечаткам в header'е (напр. "Количечество"), разделам-подзаголовкам и multi-line cells.
    if file_type == "excel":
        try:
            xlsx_preview, xlsx_columns = _smart_import_xlsx_direct(content, fname=filename)
        except Exception as _e:
            logger.warning("Direct XLSX parser failed: %s — fallback to markitdown", _e)
            xlsx_preview, xlsx_columns = [], []
        if xlsx_preview:
            if not confirm:
                return {
                    "preview": xlsx_preview[:200],
                    "total_rows": len(xlsx_preview),
                    "file_type": file_type,
                    "columns_found": xlsx_columns,
                }
            return await _save_smart_preview_to_purchase(pid, xlsx_preview, purchase, db, current_user, skip_catalog=skip_catalog)
        # Если direct-parser не нашёл строк — fallback на markitdown (ниже)

    # --- Stage 1: Convert to Markdown via markitdown ---
    from app.utils.document_to_markdown import (
        file_to_markdown, parse_markdown_tables, pick_best_table, detect_columns,
    )

    try:
        md_text = file_to_markdown(content, file.filename or filename)
    except Exception as e:
        logger.warning("markitdown conversion failed: %s", e)
        raise HTTPException(400, f"Не удалось обработать файл: {e}")

    # --- Stage 2: Extract tables from Markdown ---
    raw_tables = parse_markdown_tables(md_text)

    # Fallback: if markitdown found no tables, try legacy direct parsing
    if not raw_tables:
        raw_tables = _legacy_extract_tables(content, filename, file_type)

    if not raw_tables:
        raise HTTPException(
            400,
            "Таблицы в документе не найдены. "
            "Убедитесь что документ содержит таблицу с колонкой «Наименование»."
        )

    # --- Stage 3: Detect columns ---
    result = pick_best_table(raw_tables)
    if result is None:
        # Fallback: manual column detection on raw tables
        best_table, best_col, best_header_row = _legacy_detect_best_table(raw_tables)
    else:
        best_table, best_header_row = result
        best_col = detect_columns(best_table[best_header_row])

    if not best_table or "item_name" not in best_col:
        raise HTTPException(400, "Не удалось найти таблицу с позициями. Убедитесь что документ содержит колонку «Наименование».")

    TYPE_MAP = {
        "товар": "товар", "товары": "товар", "product": "товар",
        "услуга": "услуга", "услуги": "услуга", "service": "услуга",
        "работа": "работа", "работы": "работа",
    }

    def _to_dec(v: str):
        if not v:
            return None
        try:
            cleaned = v.replace(",", ".").replace(" ", "").replace("\xa0", "").replace("–", "").replace("—", "")
            return Decimal(cleaned)
        except Exception:
            return None

    def _parse_row(row: list[str]):
        def _get(field: str) -> str:
            idx = best_col.get(field)
            if idx is None or idx >= len(row):
                return ""
            v = row[idx].strip()
            return "" if v.lower() in ("none", "null", "-", "—", "") else v

        item_name = _get("item_name")
        if not item_name:
            return None
        item_type = TYPE_MAP.get(_get("item_type").lower(), "товар")
        quantity = _to_dec(_get("quantity"))
        unit = _get("unit") or "шт"
        unit_price = _to_dec(_get("unit_price"))
        total_price = _to_dec(_get("total_price"))
        if unit_price is None and total_price is not None and quantity:
            try:
                unit_price = total_price / quantity
            except Exception:
                pass
        if total_price is None and unit_price is not None and quantity:
            total_price = unit_price * (quantity or Decimal("1"))
        return {
            "item_name": item_name,
            "item_type": item_type,
            "quantity": float(quantity) if quantity else None,
            "unit": unit,
            "unit_price": float(unit_price) if unit_price else None,
            "total_price": float(total_price) if total_price else None,
        }

    data_rows = best_table[best_header_row + 1:]
    preview = [r for r in (_parse_row(row) for row in data_rows[:200]) if r]

    if not confirm:
        return {"preview": preview, "total_rows": len(preview), "file_type": file_type, "columns_found": list(best_col.keys())}

    # Save items to DB (markitdown path — xlsx теперь идёт через _save_smart_preview_to_purchase)
    return await _save_smart_preview_to_purchase(pid, preview, purchase, db, current_user, skip_catalog=skip_catalog)


# ---------------------------------------------------------------------------
# PDF Debug endpoint (import-pdf-debug)
# ---------------------------------------------------------------------------

@router.post("/items/import-pdf-debug")
async def import_pdf_debug(
    file: UploadFile = File(...),
    _: User = Depends(get_current_user),
):
    """Возвращает диагностику парсера для PDF/файла без импорта.

    Полезно для отладки когда Smart Import возвращает 0 строк.
    """
    import io as _io
    content = await file.read()
    result = {
        "filename": file.filename,
        "size_bytes": len(content),
        "pdfplumber_tables": 0,
        "pdfplumber_text_length": 0,
        "ocrmypdf_applied": False,
        "ocrmypdf_available": False,
        "raw_text_preview": "",
        "detected_headers": [],
        "rows_found": 0,
        "errors": [],
    }
    try:
        import pdfplumber
        with pdfplumber.open(_io.BytesIO(content)) as pdf:
            all_tables = []
            all_text = []
            for page in pdf.pages:
                tables = page.extract_tables() or []
                all_tables.extend(tables)
                txt = page.extract_text() or ""
                all_text.append(txt)
            result["pdfplumber_tables"] = len(all_tables)
            full_text = "\n".join(all_text)
            result["pdfplumber_text_length"] = len(full_text)
            result["raw_text_preview"] = full_text[:2000]
            if all_tables and all_tables[0]:
                result["detected_headers"] = [str(c) for c in (all_tables[0][0] or [])]
                result["rows_found"] = sum(max(0, len(t) - 1) for t in all_tables)
    except ImportError:
        result["errors"].append("pdfplumber не установлен")
    except Exception as e:
        result["errors"].append(f"pdfplumber: {type(e).__name__}: {str(e)[:200]}")

    # Если таблиц нет — проверить наличие ocrmypdf
    if result["pdfplumber_tables"] == 0 and result["pdfplumber_text_length"] < 100:
        try:
            import ocrmypdf  # noqa: F401
            result["ocrmypdf_available"] = True
        except ImportError:
            result["ocrmypdf_available"] = False
            result["errors"].append("ocrmypdf не установлен — сканированные PDF не распознаются")

    return result


# ---------------------------------------------------------------------------
# FEO-format import (57-column layout, headers in row 6)
# ---------------------------------------------------------------------------

FEO_HEADERS = [
    "№ п.п.", "Вид расходов", "Номер субсидии",
    "Номер строки плана закупок", "Категория расходов (ФЭО направление)",
    "Наименование товара", "Контрагент", "Количество поставлено",
    "Ед. изм.", "Стоимость единицы по смете", "Стоимость по плану всего",
    "Подтверждено", "Контрактная цена за единицу", "Стоимость контракта",
    "НМЦК всего", "НМЦК по субсидии", "КПП поставщика",
    "Номер договора", "Дата договора", "Номер ПП",
    "Описание платёжного поручения", "Дата платежа", "Сумма оплаты",
    "Всего исполнено", "Резерв 25", "Резерв 26", "Резерв 27",
    "Статус закупки", "Соответствие НДС (законодательство)", "НДС включён",
    "Соответствие требованиям НДС", "Страна происхождения", "Модель",
    "Номер категории ФЭО", "Статус исполнения", "Процент исполнения",
    "Остаток по обязательствам", "Срок исполнения", "Ссылка на торговую площадку",
    "НДС-расчёт", "Совместные", "Прямой контракт",
    "Прямой контракт (повторный)", "Малые контракты", "Конкурсные",
    "Дата последнего изменения", "Стоимость единицы по оплате", "Итого к оплате",
    "Субсидия", "Часть расходов (направление)", "Период",
    "Оплата", "Сумма НДС", "Счётчик",
    "Описание контракта", "Номер и дата платежа", "",
]


def _feo_find_header_row(rows: list) -> int:
    """Return 0-based index of the header row (first row with >=5 non-empty cells)."""
    for i, row in enumerate(rows[:10]):
        non_empty = sum(1 for c in row if c is not None and str(c).strip())
        if non_empty >= 5:
            return i
    return 0


@router.get("/import/feo-format/template")
async def download_feo_template(_=Depends(require_tab('purchases'))):
    """Скачать шаблон Excel в ФЭО-формате (57 колонок, заголовки в строке 6)."""
    if Workbook is None:
        raise HTTPException(500, "openpyxl не установлен")

    wb = Workbook()
    ws = wb.active
    ws.title = "ФЭО закупки"

    # Rows 1-5: instruction header (merged across all columns)
    ws.merge_cells("A1:BE5")
    instr_cell = ws["A1"]
    instr_cell.value = (
        "ШАБЛОН ДЛЯ ЗАГРУЗКИ ЗАКУПОК В ФЭО-ФОРМАТЕ\n"
        "Строка 6 — заголовки колонок (не изменять).\n"
        "Начиная со строки 7 — данные.\n"
        "Субсидия определяется автоматически по категории ФЭО (колонка 5).\n"
        "Обязательные колонки: 5 (ФЭО направление), 6 (Наименование товара)."
    )
    instr_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    instr_cell.font = Font(bold=True, size=11)
    ws.row_dimensions[1].height = 80

    # Row 6: headers
    ws.append(FEO_HEADERS)
    header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=10)
    for cell in ws[6]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[6].height = 35

    # Row 7: example
    ws.append([
        1, "Товары", "", "1",
        "Направление 1 - Техническое оснащение",
        "Компьютер офисный", "ООО Поставщик", 5, "шт",
        "50000", "250000", "да",
        "48000", "240000", "260000", "260000",
        "771234567890", "Д-001/2026", "15.03.2026",
        "70", "Оплата по договору Д-001/2026", "20.03.2026", "240000",
        "", "", "", "",
        "исполнено", "", "", "", "РФ", "", "",
        "оплачено", "1", "0",
        "30.06.2026", "", "", "", "", "", "", "",
        "", "", "", "", "", "", "", "2026",
        "240000", "0", "", "70 от 20.03.2026", "",
    ])
    ws.row_dimensions[7].height = 20

    ws.freeze_panes = "A7"

    # Column widths
    col_widths = [6, 12, 10, 12, 40, 35, 30, 10, 8, 14, 14, 12,
                  14, 14, 14, 14, 14, 18, 14, 10, 30, 14, 14, 14,
                  10, 10, 10, 14, 25, 14, 25, 18, 14, 14, 16, 12,
                  14, 14, 30, 12, 12, 16, 18, 14, 12, 16, 14, 14,
                  16, 20, 10, 14, 12, 20, 30, 20, 8]
    for i, w in enumerate(col_widths, 1):
        if i <= ws.max_column:
            ws.column_dimensions[ws.cell(6, i).column_letter].width = w

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{_url_quote('Шаблон_импорта_закупок_формат_ФЭО.xlsx', safe='-_.~')}"},
    )


@router.post("/import/feo-format")
async def import_feo_format(
    file: UploadFile = File(...),
    assigned_user_id: Optional[int] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_tab('purchases')),
):
    """Импорт закупок из ФЭО-формата (57 колонок, заголовки в строке 6).
    Субсидия определяется автоматически через feo_category.subsidy_id.
    assigned_user_id: если передан — присваивается всем созданным закупкам; иначе — текущий пользователь.
    """
    if load_workbook is None:
        raise HTTPException(500, "openpyxl не установлен")
    if not (file.filename or "").lower().endswith((".xlsx", ".xls")):
        raise HTTPException(400, "Поддерживаются только файлы .xlsx и .xls")

    # Resolve assigned user
    if assigned_user_id is not None:
        target_user = await db.get(User, assigned_user_id)
        if target_user is None:
            raise HTTPException(400, f"Пользователь {assigned_user_id} не найден")
        resolved_assigned_user_id = assigned_user_id
    else:
        resolved_assigned_user_id = current_user.id

    content = await file.read()
    wb = load_workbook(BytesIO(content), data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < 2:
        raise HTTPException(400, "Файл пустой")

    header_idx = _feo_find_header_row(rows)
    header_row = rows[header_idx]

    # Build column index by position keyword matching
    def _norm(v) -> str:
        return str(v).strip().lower() if v is not None else ""

    col: dict[str, int] = {}
    for i, h in enumerate(header_row):
        hn = _norm(h)
        if not hn:
            continue
        if "наименован" in hn or "товар" in hn or "услуг" in hn:
            col.setdefault("item_name", i)
        elif "вид расход" in hn or "вид" == hn:
            col.setdefault("item_type", i)
        elif "категори" in hn and ("фэо" in hn or "расход" in hn or "направлен" in hn):
            col.setdefault("feo_name", i)
        elif "контрагент" in hn:
            col.setdefault("contractor_name", i)
        elif "количеств" in hn or "кол-во" in hn:
            col.setdefault("quantity", i)
        elif "ед. изм" in hn or "единиц" in hn:
            col.setdefault("unit", i)
        elif "стоимость единицы по смете" in hn or ("стоимость" in hn and "единиц" in hn and "смет" in hn):
            col.setdefault("planned_unit_price", i)
        elif "стоимость по плану" in hn or ("стоимость" in hn and "план" in hn):
            col.setdefault("planned_total_price", i)
        elif "подтвержд" in hn:
            col.setdefault("confirmed", i)
        elif "стоимость контракта" in hn or ("стоимость" in hn and "контракт" in hn):
            col.setdefault("contract_price", i)
        elif "нмцк" in hn and "субсидии" not in hn:
            col.setdefault("nmck", i)
        elif "кпп" in hn:
            col.setdefault("kpp", i)
        elif "номер договора" in hn or ("номер" in hn and "договор" in hn):
            col.setdefault("contract_number", i)
        elif "дата договора" in hn or ("дата" in hn and "договор" in hn):
            col.setdefault("contract_date", i)
        elif "номер пп" in hn or ("номер" in hn and "платёжн" in hn) or ("номер" in hn and "поручен" in hn):
            col.setdefault("payment_doc_number", i)
        elif "дата платежа" in hn or ("дата" in hn and "платеж" in hn):
            col.setdefault("payment_doc_date", i)
        elif "сумма оплаты" in hn or ("сумма" in hn and "оплат" in hn):
            col.setdefault("payment_amount", i)
        elif "срок исполнения" in hn or ("срок" in hn and "исполнен" in hn):
            col.setdefault("execution_term", i)
        elif "сумма ндс" in hn:
            col.setdefault("vat_amount", i)

    # Fallback: use positional mapping (column indices 0-based from FEO_HEADERS order)
    # Col 5(idx=4)=feo, 6(5)=item_name, 7(6)=contractor, 8(7)=qty, 9(8)=unit,
    # 10(9)=unit_price, 11(10)=planned_total, 12(11)=confirmed, 14(13)=contract_price,
    # 15(14)=nmck, 17(16)=kpp, 18(17)=contract_num, 19(18)=contract_date,
    # 20(19)=pp_num, 22(21)=pp_date, 23(22)=payment_amount, 38(37)=execution_term, 53(52)=vat
    POSITIONAL = {
        "feo_name": 4, "item_name": 5, "item_type": 1,
        "contractor_name": 6, "quantity": 7, "unit": 8,
        "planned_unit_price": 9, "planned_total_price": 10, "confirmed": 11,
        "contract_price": 13, "nmck": 14, "kpp": 16,
        "contract_number": 17, "contract_date": 18,
        "payment_doc_number": 19, "payment_doc_date": 21, "payment_amount": 22,
        "execution_term": 37, "vat_amount": 52,
    }
    if "item_name" not in col and len(header_row) >= 57:
        # Headers didn't match keywords — use positional
        col = {k: v for k, v in POSITIONAL.items()}

    if "item_name" not in col and "feo_name" not in col:
        raise HTTPException(400, "Не удалось найти колонки наименования или ФЭО. Проверьте формат файла.")

    # Load lookup tables
    feo_rows = (await db.execute(select(FeoCategory))).scalars().all()
    feo_by_name: dict[str, FeoCategory] = {}
    for f in feo_rows:
        feo_by_name[f.name.lower().strip()] = f

    cont_rows = (await db.execute(select(Contractor))).scalars().all()
    cont_by_name: dict[str, int] = {c.name.lower().strip(): c.id for c in cont_rows}
    cont_by_kpp: dict[str, int] = {c.kpp.strip(): c.id for c in cont_rows if c.kpp}

    org_id = get_single_org_id(current_user)

    def _cell(row, field) -> Optional[str]:
        idx = col.get(field)
        if idx is None or idx >= len(row):
            return None
        v = row[idx]
        if v is None:
            return None
        s = str(v).strip()
        return s if s and s.lower() not in ("none", "null", "-", "—", "") else None

    def _to_dec(v) -> Optional[Decimal]:
        if v is None:
            return None
        try:
            return Decimal(str(v).replace(",", ".").replace(" ", "").replace("\xa0", ""))
        except Exception:
            return None

    def _to_date(v) -> Optional[date]:
        if v is None:
            return None
        if isinstance(v, (date, datetime)):
            return v.date() if isinstance(v, datetime) else v
        s = str(v).strip()
        for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y"):
            try:
                return datetime.strptime(s, fmt).date()
            except ValueError:
                pass
        return None

    created = 0
    skipped = 0
    errors_list = []

    data_start = header_idx + 1

    # ---- Permission pre-pass ----
    # Субсидия определяется только через feo_category.subsidy_id, поэтому
    # какие субсидии затронуты — известно только ПОСЛЕ разбора ФЭО-колонки.
    # Право «редактировать субсидию» (subsidy.edit) проверяется по КАЖДОЙ
    # затронутой субсидии ДО создания хоть одной закупки — нет права хотя бы
    # по одной → 403 с названием этой субсидии, ничего не создаётся
    # (требование владельца, 2026-08-19).
    def _resolve_feo_for_row(row):
        feo_name_raw = _cell(row, "feo_name")
        if not feo_name_raw:
            return None
        feo_obj = feo_by_name.get(feo_name_raw.lower().strip())
        if not feo_obj:
            for k, v in feo_by_name.items():
                if feo_name_raw.lower() in k or k in feo_name_raw.lower():
                    feo_obj = v
                    break
        return feo_obj

    touched_subsidy_ids: set[int] = set()
    for row in rows[data_start:]:
        if not _cell(row, "item_name"):
            continue
        feo_obj = _resolve_feo_for_row(row)
        if feo_obj is not None:
            touched_subsidy_ids.add(feo_obj.subsidy_id)

    if touched_subsidy_ids:
        touched_subsidies = (
            await db.execute(select(Subsidy).where(Subsidy.id.in_(touched_subsidy_ids)))
        ).scalars().all()
        subsidies_by_id = {s.id: s for s in touched_subsidies}
        for sub_id in touched_subsidy_ids:
            sub = subsidies_by_id.get(sub_id)
            sub_name = sub.name if sub else f"id {sub_id}"
            has_perm = (
                sub is not None
                and await has_org_key(current_user, db, sub.org_id, 'subsidy.edit', subsidy_id=sub_id)
            )
            if not has_perm:
                raise HTTPException(
                    403,
                    f"Импортировать закупки в субсидию «{sub_name}» может только тот, у кого "
                    "есть право её редактирования",
                )

    for row_num, row in enumerate(rows[data_start:], start=data_start + 1):
        item_name = _cell(row, "item_name")
        if not item_name:
            skipped += 1
            continue

        # FEO lookup → subsidy
        feo_name_raw = _cell(row, "feo_name")
        feo_obj = None
        if feo_name_raw:
            feo_obj = feo_by_name.get(feo_name_raw.lower().strip())
            if not feo_obj:
                # Partial match
                for k, v in feo_by_name.items():
                    if feo_name_raw.lower() in k or k in feo_name_raw.lower():
                        feo_obj = v
                        break

        if feo_obj is None:
            errors_list.append({"row": row_num, "name": item_name, "message": f"ФЭО категория не найдена: {feo_name_raw!r}"})
            skipped += 1
            continue

        feo_category_id = feo_obj.id
        subsidy_id = feo_obj.subsidy_id

        # Contractor lookup
        cont_name = _cell(row, "contractor_name")
        kpp_raw = _cell(row, "kpp")
        contractor_id = None
        if cont_name:
            contractor_id = cont_by_name.get(cont_name.lower().strip())
            if contractor_id is None and kpp_raw:
                contractor_id = cont_by_kpp.get(kpp_raw.strip())
            if contractor_id is None:
                # Create new contractor
                new_cont = Contractor(name=cont_name, kpp=kpp_raw, org_id=org_id)
                db.add(new_cont)
                await db.flush()
                contractor_id = new_cont.id
                cont_by_name[cont_name.lower().strip()] = contractor_id
                if kpp_raw:
                    cont_by_kpp[kpp_raw.strip()] = contractor_id

        # Numeric fields
        planned_qty = _to_dec(_cell(row, "quantity"))
        planned_unit = _to_dec(_cell(row, "planned_unit_price"))
        planned_total = _to_dec(_cell(row, "planned_total_price"))
        contract_price = _to_dec(_cell(row, "contract_price"))
        nmck = _to_dec(_cell(row, "nmck"))
        payment_amount = _to_dec(_cell(row, "payment_amount"))
        vat_amount = _to_dec(_cell(row, "vat_amount"))

        # Dates
        def _raw_date(field):
            idx = col.get(field)
            if idx is None or idx >= len(row):
                return None
            return _to_date(row[idx])

        contract_date = _raw_date("contract_date")
        payment_doc_date = _raw_date("payment_doc_date")
        execution_term = _raw_date("execution_term")

        # Status inference
        contract_number = _cell(row, "contract_number")
        payment_doc_number = _cell(row, "payment_doc_number")
        if payment_amount and payment_amount > 0:
            status = "paid"
        elif contract_number:
            status = "contracted"
        else:
            status = "work_in_progress"

        # Confirmed flag (legacy column — maps to boolean field Purchase.confirmed, not status)
        conf_raw = (_cell(row, "confirmed") or "").lower()
        confirmed = conf_raw in ("да", "yes", "1", "true", "+")

        # Item type
        item_type_raw = (_cell(row, "item_type") or "").lower()
        item_type = "услуга" if "услуг" in item_type_raw else "товар"

        # VAT
        vat_applicable = bool(vat_amount and vat_amount > 0)

        p = Purchase(
            item_name=item_name,
            item_type=item_type,
            feo_category_id=feo_category_id,
            subsidy_id=subsidy_id,
            contractor_id=contractor_id,
            planned_quantity=planned_qty,
            unit=_cell(row, "unit"),
            planned_unit_price=planned_unit,
            planned_total_price=planned_total,
            confirmed=confirmed,
            contract_price=contract_price,
            nmck=nmck,
            total_nmck=nmck,
            contract_number=contract_number,
            contract_date=contract_date,
            payment_doc_number=payment_doc_number,
            payment_doc_date=payment_doc_date,
            payment_amount=payment_amount,
            execution_term=execution_term,
            vat_applicable=vat_applicable,
            status=status,
            assigned_user_id=resolved_assigned_user_id,
        )
        db.add(p)
        created += 1

    await db.commit()
    return {"created": created, "skipped": skipped, "errors": errors_list}
