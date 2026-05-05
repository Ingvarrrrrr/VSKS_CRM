"""
Генератор цветных шаблонов Технических Заданий (ТЗ) — Phase 23.2.

Запуск:
    py backend/templates/make_tech_spec.py

Создаёт 2 файла в backend/templates/:
    tech_spec_request.docx   — ТЗ для запроса коммерческих предложений
    tech_spec_contract.docx  — ТЗ для договора (Приложение №1)

Цветовая схема (как в contract.docx):
    {{синие}}   — placeholder'ы (RGBColor #1E40AF, жирный)
    зелёный     — ветка «услуги» условного блока
    оранжевый   — ветка «товары» условного блока
    серый курсив — маркеры {% if %} / {% else %} / {% endif %}
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── ЦВЕТА ───────────────────────────────────────────────────────────────────
COLOR_PLACEHOLDER = RGBColor(0x1E, 0x40, 0xAF)   # синий — {{var}}
COLOR_BRANCH_TRUE = RGBColor(0x15, 0x80, 0x3D)   # зелёный — ветка услуги
COLOR_BRANCH_FALSE = RGBColor(0xC2, 0x41, 0x0C)  # оранжевый — ветка товары
COLOR_MARKER = RGBColor(0x6B, 0x72, 0x80)         # серый — {% %}

FONT_NAME = "Times New Roman"


# ─── HELPER'Ы ─────────────────────────────────────────────────────────────────

def _base_run(paragraph, text, size=12):
    r = paragraph.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    return r


def add_placeholder(para, text, size=12):
    """{{var}} — синий жирный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_PLACEHOLDER
    r.bold = True
    return r


def add_marker(para, text, size=12):
    """{% if %} / {% else %} / {% endif %} — серый курсив."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_MARKER
    r.italic = True
    return r


def add_branch_true(para, text, size=12):
    """Ветка услуги — зелёный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_BRANCH_TRUE
    return r


def add_branch_false(para, text, size=12):
    """Ветка товары — оранжевый."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_BRANCH_FALSE
    return r


def add_text(para, text, bold=False, size=12, italic=False):
    r = _base_run(para, text, size)
    r.bold = bold
    r.italic = italic
    return r


# ─── УТИЛИТЫ ДОКУМЕНТА ───────────────────────────────────────────────────────

def page_margins(doc):
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)


def new_para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
             first_line_indent=None, line_spacing=13.8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    return p


def add_text_run(para, text, bold=False, size=12, italic=False):
    r = _base_run(para, text, size)
    r.bold = bold
    r.italic = italic
    return r


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = Pt(13)
    return p


def clause(doc, num_text, space_before=4):
    p = new_para(doc, space_before=space_before, space_after=4,
                 first_line_indent=Cm(1))
    r = p.add_run(f"{num_text}. ")
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = Pt(12)
    return p


# ─── ЛЕГЕНДА ─────────────────────────────────────────────────────────────────

def add_legend(doc):
    p_leg_title = doc.add_paragraph()
    p_leg_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_leg_title.paragraph_format.space_after = Pt(4)
    p_leg_title.paragraph_format.space_before = Pt(0)
    r_lt = p_leg_title.add_run("🎨 ЛЕГЕНДА — как читать этот шаблон")
    r_lt.bold = True
    r_lt.font.name = FONT_NAME
    r_lt.font.size = Pt(11)
    r_lt.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    LEGEND_LINES = [
        ("placeholder", "{{синие переменные}}    — автоподстановка из БД"),
        ("true",        "зелёный текст           — вариант «услуги» в условном блоке"),
        ("false",       "оранжевый текст         — вариант «товары»"),
        ("marker",      "серые {% маркеры %}    — технические условия Jinja2 (не трогать)"),
    ]

    for kind, text in LEGEND_LINES:
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pl.paragraph_format.space_before = Pt(0)
        pl.paragraph_format.space_after = Pt(2)
        pl.paragraph_format.first_line_indent = Cm(0.5)
        r = pl.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        if kind == "placeholder":
            r.font.color.rgb = COLOR_PLACEHOLDER
            r.bold = True
        elif kind == "true":
            r.font.color.rgb = COLOR_BRANCH_TRUE
        elif kind == "false":
            r.font.color.rgb = COLOR_BRANCH_FALSE
        elif kind == "marker":
            r.font.color.rgb = COLOR_MARKER
            r.italic = True

    p_warn = doc.add_paragraph()
    p_warn.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_warn.paragraph_format.space_before = Pt(4)
    p_warn.paragraph_format.space_after = Pt(2)
    p_warn.paragraph_format.first_line_indent = Cm(0.5)
    r_warn = p_warn.add_run(
        "⚠️ После открытия в Word удалите эту секцию (от «ЛЕГЕНДА» до строки «—————»).\n"
        "Полный справочник переменных: SubsidiesView → Шаблоны → «Руководство по переменным»"
    )
    r_warn.font.name = FONT_NAME
    r_warn.font.size = Pt(10)
    r_warn.italic = True

    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sep.paragraph_format.space_before = Pt(2)
    p_sep.paragraph_format.space_after = Pt(10)
    r_sep = p_sep.add_run("—" * 40)
    r_sep.font.name = FONT_NAME
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = COLOR_MARKER


# ─── ТАБЛИЦА ПОЗИЦИЙ (цикл) ──────────────────────────────────────────────────

def add_items_table(doc, columns, col_headers):
    """
    columns — список имён полей item (без 'item.'), например ['num', 'name', 'unit', 'quantity']
    col_headers — заголовки колонок
    """
    p_for = new_para(doc, space_after=0)
    add_marker(p_for, "{% tr for item in items %}", size=10)

    tbl = doc.add_table(rows=2, cols=len(columns))
    tbl.style = "Table Grid"

    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(col_headers):
        p_h = hdr_cells[i].paragraphs[0]
        r = p_h.add_run(h)
        r.bold = True
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data_cells = tbl.rows[1].cells
    for i, col in enumerate(columns):
        p_c = data_cells[i].paragraphs[0]
        r = p_c.add_run(f"{{{{item.{col}}}}}")
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_PLACEHOLDER
        r.bold = True
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_endfor = new_para(doc, space_before=2, space_after=6)
    add_marker(p_endfor, "{% tr endfor %}", size=10)


# ─── ТЗ ДЛЯ ЗАПРОСА КП ───────────────────────────────────────────────────────

def make_ts_request(doc):
    section_heading(doc, "ТЕХНИЧЕСКОЕ ЗАДАНИЕ")
    p_sub = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_text_run(p_sub, "для запроса коммерческих предложений")

    # Заказчик
    p_cust = new_para(doc, space_after=6)
    add_text_run(p_cust, "Заказчик: ")
    add_placeholder(p_cust, "{{customer_full_name}}")
    add_text_run(p_cust, " (ИНН ")
    add_placeholder(p_cust, "{{customer_inn}}")
    add_text_run(p_cust, ")")

    # 1. Объект закупки
    p1 = clause(doc, "1")
    add_text_run(p1, "Объект закупки: ")
    add_placeholder(p1, "{{subject}}")

    # 2. Тип — условный
    p2 = clause(doc, "2")
    add_text_run(p2, "Тип: ")
    add_marker(p2, "{% if subject_kind == 'goods' %}")
    add_branch_false(p2, "поставка товаров")
    add_marker(p2, "{% else %}")
    add_branch_true(p2, "оказание услуг")
    add_marker(p2, "{% endif %}")

    # 3. Срок подачи КП
    p3 = clause(doc, "3")
    add_text_run(p3, "Срок предоставления КП до: ")
    add_placeholder(p3, "{{submission_deadline_datetime}}")

    # 4. Срок — условный
    p4 = clause(doc, "4")
    add_text_run(p4, "Срок ")
    add_marker(p4, "{% if subject_kind == 'goods' %}")
    add_branch_false(p4, "поставки")
    add_marker(p4, "{% else %}")
    add_branch_true(p4, "оказания")
    add_marker(p4, "{% endif %}")
    add_text_run(p4, ": ")
    add_placeholder(p4, "{{service_term}}")

    # 5. Место — условный
    p5 = clause(doc, "5")
    add_text_run(p5, "Место ")
    add_marker(p5, "{% if subject_kind == 'goods' %}")
    add_branch_false(p5, "поставки")
    add_marker(p5, "{% else %}")
    add_branch_true(p5, "оказания")
    add_marker(p5, "{% endif %}")
    add_text_run(p5, ": ")
    add_placeholder(p5, "{{delivery_location}}")

    # 6. Спецификация — таблица позиций
    p6_hdr = clause(doc, "6")
    add_text_run(p6_hdr, "Спецификация:")

    add_items_table(
        doc,
        columns=["num", "name", "description", "unit", "quantity"],
        col_headers=["№", "Наименование", "Описание / Требования", "Ед. изм.", "Кол-во"]
    )

    # 7. Требования к КП
    p7 = clause(doc, "7")
    add_text_run(p7, "Требования к КП: цена за единицу + НДС, общая сумма, срок действия предложения.")

    # 8. Контактное лицо
    p8 = clause(doc, "8")
    add_text_run(p8, "Контактное лицо Заказчика: ")
    add_placeholder(p8, "{{customer_signatory_name}}")
    add_text_run(p8, ", тел.: ")
    add_placeholder(p8, "{{customer_phone}}")
    add_text_run(p8, ", email: ")
    add_placeholder(p8, "{{customer_email}}")


# ─── ТЗ ДЛЯ ДОГОВОРА (Приложение №1) ────────────────────────────────────────

def make_ts_contract(doc):
    # Шапка «Приложение №1»
    p_app = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
    add_text_run(p_app, "Приложение №1 к Договору № ")
    add_placeholder(p_app, "{{contract_number}}")

    p_app2 = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)
    add_text_run(p_app2, "от «")
    add_placeholder(p_app2, "{{contract_date_day}}")
    add_text_run(p_app2, "» ")
    add_placeholder(p_app2, "{{contract_date_month}}")
    add_text_run(p_app2, " ")
    add_placeholder(p_app2, "{{contract_date_year}}")
    add_text_run(p_app2, " г.")

    section_heading(doc, "ТЕХНИЧЕСКОЕ ЗАДАНИЕ")

    # 1. Объект — условный
    p1 = clause(doc, "1")
    add_text_run(p1, "Объект ")
    add_marker(p1, "{% if subject_kind == 'goods' %}")
    add_branch_false(p1, "поставки")
    add_marker(p1, "{% else %}")
    add_branch_true(p1, "услуг")
    add_marker(p1, "{% endif %}")
    add_text_run(p1, ": ")
    add_placeholder(p1, "{{subject}}")

    # 2. Срок
    p2 = clause(doc, "2")
    add_text_run(p2, "Срок: ")
    add_placeholder(p2, "{{service_term}}")

    # 3. Место — условный
    p3 = clause(doc, "3")
    add_text_run(p3, "Место ")
    add_marker(p3, "{% if subject_kind == 'goods' %}")
    add_branch_false(p3, "поставки")
    add_marker(p3, "{% else %}")
    add_branch_true(p3, "оказания")
    add_marker(p3, "{% endif %}")
    add_text_run(p3, ": ")
    add_placeholder(p3, "{{delivery_location}}")

    # 4. Третьи лица — условный
    p4 = clause(doc, "4")
    add_marker(p4, "{% if third_party_involved %}")
    add_branch_true(p4, "Допускается привлечение третьих лиц")
    add_marker(p4, "{% else %}")
    add_branch_false(p4, "Без привлечения третьих лиц")
    add_marker(p4, "{% endif %}")
    add_text_run(p4, ".")

    # 5. Спецификация — таблица позиций с вложенными полями
    p5_hdr = clause(doc, "5")
    add_text_run(p5_hdr, "Спецификация:")

    # Расширенная таблица для договорного ТЗ
    p_for = new_para(doc, space_after=0)
    add_marker(p_for, "{% tr for item in items %}", size=10)

    tbl = doc.add_table(rows=2, cols=6)
    tbl.style = "Table Grid"

    col_headers = ["№", "Наименование", "Описание", "Кол-во / Ед.", "Цена, руб.", "Сумма, руб."]
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(col_headers):
        p_h = hdr_cells[i].paragraphs[0]
        r = p_h.add_run(h)
        r.bold = True
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cols_data = ["item.num", "item.name", "item.description",
                 "item.quantity ~ item.unit", "item.unit_price", "item.total_price"]
    placeholders = [
        "{{item.num}}", "{{item.name}}", "{{item.description}}",
        "{{item.quantity}} {{item.unit}}", "{{item.unit_price}}", "{{item.total_price}}"
    ]
    data_cells = tbl.rows[1].cells
    for i, ph in enumerate(placeholders):
        p_c = data_cells[i].paragraphs[0]
        r = p_c.add_run(ph)
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_PLACEHOLDER
        r.bold = True
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_endfor = new_para(doc, space_before=2, space_after=6)
    add_marker(p_endfor, "{% tr endfor %}", size=10)

    # Итого
    p_total = new_para(doc, space_after=4)
    add_text_run(p_total, "ИТОГО: ")
    add_placeholder(p_total, "{{contract_price_num}}")
    add_text_run(p_total, " (")
    add_placeholder(p_total, "{{contract_price_words}}")
    add_text_run(p_total, ")")

    # НДС — условный
    p_vat = new_para(doc, space_after=6)
    add_marker(p_vat, "{% if vat_applicable %}")
    add_branch_true(p_vat, "в т.ч. НДС ")
    add_placeholder(p_vat, "{{vat_rate}}")
    add_branch_true(p_vat, "% — ")
    add_placeholder(p_vat, "{{vat_amount_num}}")
    add_marker(p_vat, "{% else %}")
    add_branch_false(p_vat, "НДС не облагается на основании ")
    add_placeholder(p_vat, "{{vat_exemption_article}}")
    add_marker(p_vat, "{% endif %}")

    # 6. Требования к качеству
    p6 = clause(doc, "6")
    add_text_run(p6, "Требования к качеству: согласно ГОСТ, ТУ и условиям Договора.")

    # 7. Отчётность
    p7 = clause(doc, "7")
    add_text_run(p7, "Отчётность: акт сдачи-приёмки в 2 экз.")

    # Подписи
    doc.add_paragraph()
    tbl_sign = doc.add_table(rows=1, cols=2)
    tbl_sign.style = "Table Grid"

    left_c = tbl_sign.rows[0].cells[0]
    right_c = tbl_sign.rows[0].cells[1]

    def sign_cell(cell, label, ph_initials):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(f"{label}: ___________________ / ")
        r.font.name = FONT_NAME
        r.font.size = Pt(11)
        r2 = p.add_run(ph_initials)
        r2.font.name = FONT_NAME
        r2.font.size = Pt(11)
        r2.font.color.rgb = COLOR_PLACEHOLDER
        r2.bold = True
        p.paragraph_format.space_after = Pt(4)

    sign_cell(left_c, "Заказчик", "{{customer_signatory_initials}}")
    sign_cell(right_c, "Исполнитель", "{{contractor_signatory_initials}}")


# ─── ГЛАВНАЯ ФУНКЦИЯ ─────────────────────────────────────────────────────────

def make_tech_spec(kind, output_path):
    """
    Создать шаблон технического задания.
    kind ∈ {'request', 'contract'}
    """
    doc = Document()
    page_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(12)

    add_legend(doc)

    if kind == 'request':
        make_ts_request(doc)
    elif kind == 'contract':
        make_ts_contract(doc)
    else:
        raise ValueError(f"Unknown kind: {kind!r}")

    doc.save(output_path)
    print(f"OK: {output_path}")


# ─── MAIN ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    KINDS = {
        'request':  "tech_spec_request.docx",
        'contract': "tech_spec_contract.docx",
    }
    for kind, filename in KINDS.items():
        out = os.path.join(TEMPLATES_DIR, filename)
        make_tech_spec(kind, out)
