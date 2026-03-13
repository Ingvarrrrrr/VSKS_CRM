import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import lxml.etree as etree

INPUT  = r'C:/Users/1/VSKS_CRM/типовые документы/ФАДМ/Договор на оказание услуг (Типовой) для 2025 (ФАДМ) - большая отчетность.docx'
OUTPUT = r'C:/Users/1/VSKS_CRM/backend/templates/contract_fadm.docx'

doc = Document(INPUT)

def set_para(para, text):
    """Clear runs, put all text in first run."""
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    first_run = runs[0]
    first_run.text = text
    for run in runs[1:]:
        r_elem = run._r
        r_elem.getparent().remove(r_elem)

def find_para(doc, startswith):
    for p in doc.paragraphs:
        if p.text.startswith(startswith):
            return p
    return None

# ── P0: Номер договора ───────────────────────────────────────────────────────
set_para(doc.paragraphs[0], 'Договор № {{ contract_number }}')

# ── P2: Дата ─────────────────────────────────────────────────────────────────
p2 = doc.paragraphs[2]
for run in p2.runs:
    if '«___» _______ 20__ г.' in run.text:
        run.text = run.text.replace(
            '«___» _______ 20__ г.',
            '«{{ contract_date_day }}» {{ contract_date_month }} {{ contract_date_year }} г.')
        break

# ── P5: Преамбула для ЮЛ ─────────────────────────────────────────────────────
p5 = doc.paragraphs[5]
set_para(p5,
    '{%p if contractor_org_type == "Юр.лицо" %}'
    '{{ contractor_name }} ({{ contractor_short_name }}), именуемое/ая в дальнейшем «Исполнитель», '
    'в лице {{ contractor_signatory_position }} {{ contractor_signatory }}, '
    'действующего на основании {{ contractor_signatory_basis }},')

# ── P7: Преамбула для ИП ─────────────────────────────────────────────────────
p7 = doc.paragraphs[7]
set_para(p7,
    '{%p if contractor_org_type == "ИП" %}'
    'Индивидуальный предприниматель {{ contractor_name }}, именуемый в дальнейшем «Исполнитель», '
    'в своем лице и действующий в своих собственных интересах, зарегистрированный в качестве '
    'индивидуального предпринимателя в соответствии с законодательством Российской Федерации '
    '(ОГРНИП {{ contractor_ogrn }}),')

# ── Вставка параграфа Самозанятый после P7 ───────────────────────────────────
p7_elem = doc.paragraphs[7]._p
new_p_elem = deepcopy(p7_elem)
# Удаляем все runs в копии
for r in new_p_elem.findall(qn('w:r')):
    new_p_elem.remove(r)
# Добавляем один run с нужным текстом
r_new = etree.SubElement(new_p_elem, qn('w:r'))
t_new = etree.SubElement(r_new, qn('w:t'))
t_new.text = (
    '{%p if contractor_org_type == "Самозанятый" %}'
    '{{ contractor_name }}, применяющий специальный налоговый режим '
    '«Налог на профессиональный доход» (ИНН {{ contractor_inn }}), '
    'именуемый в дальнейшем «Исполнитель», в своем лице,'
)
t_new.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
p7_elem.addnext(new_p_elem)

# ── 1.1 Предмет договора ─────────────────────────────────────────────────────
p_11 = find_para(doc, '\t1.1.')
if p_11:
    for run in p_11.runs:
        if '____________________________' in run.text and 'Услуги' not in run.text:
            run.text = '{{ service_name }}'
            break

# ── 1.2 Срок (период) ────────────────────────────────────────────────────────
p_period = find_para(doc, '\t1.2. Срок')
if p_period:
    set_para(p_period,
        '{%p if period_type == "period" %}'
        '\t1.2. Срок оказания Услуг: '
        '{% if service_start_date %}с «{{ service_start_date }}» {% else %}с момента заключения Договора {% endif %}'
        'по «{{ service_end_date }}» (включительно).')

# ── 1.2 Дата (разовая) ───────────────────────────────────────────────────────
p_date = find_para(doc, '\t1.2. Дата')
if p_date:
    set_para(p_date,
        '{%p if period_type == "date" %}'
        '\t1.2. Дата оказания услуг: {{ service_date }}.')

# ── 1.4 Третьи лица ──────────────────────────────────────────────────────────
p_third = find_para(doc, '1.4.')
if p_third:
    set_para(p_third,
        '1.4. Услуги по настоящему Договору оказываются Исполнителем своими силами и средствами '
        '{% if third_party_involved %}или с привлечением третьих лиц'
        '{% else %}без привлечения третьих лиц{% endif %}.')

# ── 4.1 Цена договора ────────────────────────────────────────────────────────
p_41 = find_para(doc, '4.1. Цена')
if p_41:
    set_para(p_41,
        '4.1. Цена Договора включает в себя все расходы Исполнителя, связанные с оказанием Услуг '
        'по настоящему Договору (в т.ч. выполнение всех сопутствующих работ и всех сопутствующих услуг, '
        'оказываемых в соответствии с Техническим заданием (Приложение №1)), налоги, сборы и другие '
        'обязательные платежи и составляет {{ contract_price_num }} ({{ contract_price_words }}) рублей 00 копеек, '
        '{% if vat_applicable %}'
        'в том числе НДС {{ vat_rate }}% — {{ vat_amount_num }} ({{ vat_amount_words }}) рублей 00 копеек'
        '{% else %}'
        'НДС не облагается на основании {{ vat_exemption_article }}'
        '{% endif %}.')

# ── 4.1.1 НДС оговорка (только без НДС) ─────────────────────────────────────
p_411 = find_para(doc, '4.1.1.')
if p_411:
    set_para(p_411,
        '{%p if not vat_applicable %}'
        '4.1.1. В случае, если в течение срока действия Договора возникнет НДС, Цена Договора, '
        'указанная в п. 4.1. будет включать в себя НДС (НДС считается включенным) по ставке, '
        'применяемой Исполнителем и/или применимой для Исполнителя в соответствующий период/срок.')

# ── Таблица реквизитов (T0, C1 — Исполнитель) ───────────────────────────────
tbl = doc.tables[0]
cell = tbl.rows[0].cells[1]
paras = cell.paragraphs

if len(paras) > 1:
    set_para(paras[1], 'Наименование: {{ contractor_name }}')
if len(paras) > 2:
    set_para(paras[2], 'Адрес местонахождения: {{ contractor_address }}')
if len(paras) > 3:
    set_para(paras[3], 'Почтовый адрес: {{ contractor_postal_address }}')
# ИНН/КПП — только для ЮЛ
if len(paras) > 4:
    set_para(paras[4],
        '{%p if contractor_org_type == "Юр.лицо" %}'
        'ИНН/КПП: {{ contractor_inn }}/{{ contractor_kpp }}')
# ОГРН — только для ЮЛ
if len(paras) > 5:
    set_para(paras[5],
        '{%p if contractor_org_type == "Юр.лицо" %}'
        'ОГРН: {{ contractor_ogrn }}')
# ИНН — для ИП и Самозанятого
if len(paras) > 7:
    set_para(paras[7],
        '{%p if contractor_org_type != "Юр.лицо" %}'
        'ИНН: {{ contractor_inn }}')
# ОГРНИП — только для ИП
if len(paras) > 8:
    set_para(paras[8],
        '{%p if contractor_org_type == "ИП" %}'
        'ОГРНИП: {{ contractor_ogrn }}')
if len(paras) > 11:
    set_para(paras[11], 'р/с: {{ contractor_settlement_account }}')
if len(paras) > 12:
    set_para(paras[12], 'в {{ contractor_bank_name }}')
if len(paras) > 13:
    set_para(paras[13], 'БИК: {{ contractor_bik }}')
if len(paras) > 14:
    set_para(paras[14], 'к/с: {{ contractor_correspondent_account }}')
if len(paras) > 16:
    set_para(paras[16], 'Тел.: {{ contractor_phone }}')
if len(paras) > 17:
    set_para(paras[17], 'Эл.почта: {{ contractor_email }}')
if len(paras) > 19:
    set_para(paras[19],
        '{% if contractor_org_type == "ИП" %}Индивидуальный предприниматель'
        '{% elif contractor_org_type == "Самозанятый" %}Самозанятый'
        '{% else %}{{ contractor_signatory_position }}{% endif %}')
if len(paras) > 20:
    set_para(paras[20], '_________________ / {{ contractor_signatory }}')

# ── Таблица подписей (T1, C1 — Исполнитель) ─────────────────────────────────
tbl1 = doc.tables[1]
cell1 = tbl1.rows[0].cells[1]
paras1 = cell1.paragraphs
if len(paras1) > 1:
    set_para(paras1[1], '{{ contractor_name }}')
if len(paras1) > 3:
    set_para(paras1[3],
        '{% if contractor_org_type == "ИП" %}Индивидуальный предприниматель'
        '{% elif contractor_org_type == "Самозанятый" %}Самозанятый'
        '{% else %}{{ contractor_signatory_position }}{% endif %}')
if len(paras1) > 5:
    set_para(paras1[5], '_________________ / {{ contractor_signatory }}')

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
print('Paragraphs check:')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t and ('{{' in t or '{%' in t):
        print(f'  P{i}: {t[:120]}')
