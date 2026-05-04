"""
Генератор универсального шаблона «Договор» (contract.docx) — Phase 23.1.

Запуск (один раз):
    py backend/templates/make_contract.py

Создаёт backend/templates/contract.docx — шаблон docxtpl с Jinja2
placeholder'ами и условными блоками {% if subject_kind == 'goods' %}{% else %}{% endif %}.

subject_kind определяется автоматически в documents.py:
  - 'services' — если ВСЕ позиции закупки имеют item_kind='услуга'
  - 'goods'    — если есть хотя бы одна позиция 'товар' (или нет позиций)

Цветовая схема:
    {{синие}}   — placeholder'ы (RGBColor #1E40AF, жирный)
    зелёный     — ветка services ({% if subject_kind != 'goods' %})
    оранжевый   — ветка goods  ({% else %} / {% if subject_kind == 'goods' %})
    серый курсив — маркеры {% if %} / {% else %} / {% endif %}
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(TEMPLATES_DIR, "contract.docx")

# ─── ЦВЕТА ───────────────────────────────────────────────────────────────────
COLOR_PLACEHOLDER = RGBColor(0x1E, 0x40, 0xAF)   # синий — {{var}}
COLOR_BRANCH_TRUE = RGBColor(0x15, 0x80, 0x3D)   # зелёный — ветка services
COLOR_BRANCH_FALSE = RGBColor(0xC2, 0x41, 0x0C)  # оранжевый — ветка goods
COLOR_MARKER = RGBColor(0x6B, 0x72, 0x80)         # серый — {% %}

FONT_NAME = "Times New Roman"


# ─── HELPER'Ы ─────────────────────────────────────────────────────────────────

def _base_run(paragraph, text, size=12):
    r = paragraph.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    return r


def add_placeholder(para, text, size=12):
    """{{var}} — синий жирный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_PLACEHOLDER
    r.bold = True
    return r


def add_marker(para, text, size=12):
    """{% if %} / {% else %} / {% endif %} — серый курсив."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_MARKER
    r.italic = True
    return r


def add_branch_true(para, text, size=12):
    """Ветка services — зелёный обычный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_BRANCH_TRUE
    return r


def add_branch_false(para, text, size=12):
    """Ветка goods — оранжевый обычный."""
    r = _base_run(para, text, size)
    r.font.color.rgb = COLOR_BRANCH_FALSE
    return r


def add_text(para, text, bold=False, size=12, italic=False):
    r = _base_run(para, text, size)
    r.bold = bold
    r.italic = italic
    return r


# ─── УТИЛИТЫ ДОКУМЕНТА ───────────────────────────────────────────────────────

def page_margins(doc):
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def new_para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
             first_line_indent=None, line_spacing=13.8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    return p


def para_simple(doc, text="", bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0, space_after=6, italic=False):
    p = new_para(doc, align=align, space_before=space_before, space_after=space_after)
    if text:
        add_text(p, text, bold=bold, size=size, italic=italic)
    return p


def section_title(doc, text, num=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    full = f"{num}. {text}" if num else text
    r = p.add_run(full)
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = Pt(12)
    return p


def clause_para(doc, num, space_before=2):
    p = new_para(doc, space_before=space_before, space_after=3,
                 first_line_indent=Cm(1))
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = Pt(12)
    return p


def clause_plain(doc, num, text, space_before=2):
    p = clause_para(doc, num, space_before)
    add_text(p, text)
    return p


# ─── СБОРКА ДОКУМЕНТА ────────────────────────────────────────────────────────

doc = Document()
page_margins(doc)

style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = Pt(12)

# ══════════════════════════════════════════════════════════════════════════════
# ЛЕГЕНДА
# ══════════════════════════════════════════════════════════════════════════════

p_leg_title = doc.add_paragraph()
p_leg_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_leg_title.paragraph_format.space_after = Pt(4)
p_leg_title.paragraph_format.space_before = Pt(0)
r_lt = p_leg_title.add_run("🎨 ЛЕГЕНДА — как читать этот шаблон")
r_lt.bold = True
r_lt.font.name = FONT_NAME
r_lt.font.size = Pt(11)
r_lt.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

LEGEND_LINES = [
    ("placeholder", "{{синие переменные}}  — автоподстановка из БД (Заказчик / Исполнитель / Цена / Сроки)"),
    ("true",        "зелёный текст         — вариант «услуги» (Заказчик/Исполнитель, Акт сдачи-приёмки, ТЗ)"),
    ("false",       "оранжевый текст       — вариант «товары» (Покупатель/Поставщик, накладная, Спецификация)"),
    ("marker",      "серые {% маркеры %}   — технические условия Jinja2 (оставить как есть)"),
]

for kind, text in LEGEND_LINES:
    pl = doc.add_paragraph()
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pl.paragraph_format.space_before = Pt(0)
    pl.paragraph_format.space_after = Pt(2)
    pl.paragraph_format.first_line_indent = Cm(0.5)
    r = pl.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = Pt(10)
    if kind == "placeholder":
        r.font.color.rgb = COLOR_PLACEHOLDER
        r.bold = True
    elif kind == "true":
        r.font.color.rgb = COLOR_BRANCH_TRUE
    elif kind == "false":
        r.font.color.rgb = COLOR_BRANCH_FALSE
    elif kind == "marker":
        r.font.color.rgb = COLOR_MARKER
        r.italic = True

# Пояснение про автоопределение
p_auto = doc.add_paragraph()
p_auto.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_auto.paragraph_format.space_before = Pt(4)
p_auto.paragraph_format.space_after = Pt(2)
p_auto.paragraph_format.first_line_indent = Cm(0.5)
r_auto = p_auto.add_run(
    "⚙️ ТИП ДОГОВОРА выбирается автоматически по позициям закупки:\n"
    "   • все позиции 'услуга' → договор оказания услуг (зелёные ветки)\n"
    "   • есть позиция 'товар' или нет позиций → договор поставки (оранжевые ветки)\n"
    "В сгенерированном документе останется только одна ветка."
)
r_auto.font.name = FONT_NAME
r_auto.font.size = Pt(10)
r_auto.italic = True

p_warn = doc.add_paragraph()
p_warn.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_warn.paragraph_format.space_before = Pt(4)
p_warn.paragraph_format.space_after = Pt(2)
p_warn.paragraph_format.first_line_indent = Cm(0.5)
r_warn = p_warn.add_run(
    "⚠️ После открытия в Word удалите ЭТУ секцию (от «ЛЕГЕНДА» до строки «—————») "
    "если хотите чистый шаблон. Сама разметка должна остаться нетронутой."
)
r_warn.font.name = FONT_NAME
r_warn.font.size = Pt(10)
r_warn.italic = True

p_sep = doc.add_paragraph()
p_sep.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_sep.paragraph_format.space_before = Pt(2)
p_sep.paragraph_format.space_after = Pt(10)
r_sep = p_sep.add_run("—" * 40)
r_sep.font.name = FONT_NAME
r_sep.font.size = Pt(10)
r_sep.font.color.rgb = COLOR_MARKER

# ══════════════════════════════════════════════════════════════════════════════
# ЗАГОЛОВОК ДОГОВОРА — условный
# ══════════════════════════════════════════════════════════════════════════════

# {% if subject_kind == 'goods' %} — ДОГОВОР ПОСТАВКИ
p_if_title = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_marker(p_if_title, "{% if subject_kind == 'goods' %}", size=10)

p_title_goods = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_branch_false(p_title_goods, "ДОГОВОР ПОСТАВКИ № ", size=14)
add_placeholder(p_title_goods, "{{contract_number}}", size=14)

# {% else %} — ДОГОВОР ОКАЗАНИЯ УСЛУГ
p_else_title = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_marker(p_else_title, "{% else %}", size=10)

p_title_services = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_branch_true(p_title_services, "ДОГОВОР № ", size=14)
add_placeholder(p_title_services, "{{contract_number}}", size=14)

p_sub_services = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_branch_true(p_sub_services, "на оказание услуг", size=12)

# {% endif %}
p_endif_title = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_marker(p_endif_title, "{% endif %}", size=10)

# Город и дата
p_city = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
add_text(p_city, "г. ")
add_placeholder(p_city, "{{contract_city}}")
add_text(p_city, "                                    «")
add_placeholder(p_city, "{{contract_date_day}}")
add_text(p_city, "» ")
add_placeholder(p_city, "{{contract_date_month}}")
add_text(p_city, " ")
add_placeholder(p_city, "{{contract_date_year}}")
add_text(p_city, " г.")

# ══════════════════════════════════════════════════════════════════════════════
# ПРЕАМБУЛА
# ══════════════════════════════════════════════════════════════════════════════

# Заказчик (одинаково для обоих вариантов — меняется только роль)
p_pre1 = new_para(doc, space_after=6)
add_placeholder(p_pre1, "{{customer_full_name}}")
add_text(p_pre1, " (")
add_placeholder(p_pre1, "{{customer_short_name}}")
add_text(p_pre1, "), именуемое в дальнейшем ")
add_marker(p_pre1, "{% if subject_kind == 'goods' %}")
add_branch_false(p_pre1, "«Покупатель»")
add_marker(p_pre1, "{% else %}")
add_branch_true(p_pre1, "«Заказчик»")
add_marker(p_pre1, "{% endif %}")
add_text(p_pre1, ", в лице ")
add_placeholder(p_pre1, "{{customer_signatory_position}}")
add_text(p_pre1, " ")
add_placeholder(p_pre1, "{{customer_signatory_name_genitive}}")
add_text(p_pre1, ", действующего на основании ")
add_placeholder(p_pre1, "{{customer_signatory_basis}}")
add_text(p_pre1, ", с одной стороны, и")

# Исполнитель/Поставщик — условный блок ИП / Юр.лицо
p_cond = new_para(doc, space_after=6)
add_marker(p_cond, "{%- if contractor_org_type == 'ИП' %}")
add_branch_true(p_cond, "Индивидуальный предприниматель ")
add_placeholder(p_cond, "{{contractor_signatory_name}}")
add_branch_true(p_cond, " (ИП ")
add_placeholder(p_cond, "{{contractor_short_name}}")
add_branch_true(p_cond, "), ИНН ")
add_placeholder(p_cond, "{{contractor_inn}}")
add_branch_true(p_cond, ", ОГРНИП ")
add_placeholder(p_cond, "{{contractor_ogrnip}}")
add_marker(p_cond, "{%- else %}")
add_placeholder(p_cond, "{{contractor_full_name}}")
add_branch_false(p_cond, " (")
add_placeholder(p_cond, "{{contractor_short_name}}")
add_branch_false(p_cond, "), в лице ")
add_placeholder(p_cond, "{{contractor_signatory_position}}")
add_branch_false(p_cond, " ")
add_placeholder(p_cond, "{{contractor_signatory_name_genitive}}")
add_branch_false(p_cond, ", действующего на основании ")
add_placeholder(p_cond, "{{contractor_signatory_basis}}")
add_marker(p_cond, "{%- endif %}")
add_text(p_cond, ", именуемое в дальнейшем ")
add_marker(p_cond, "{% if subject_kind == 'goods' %}")
add_branch_false(p_cond, "«Поставщик»")
add_marker(p_cond, "{% else %}")
add_branch_true(p_cond, "«Исполнитель»")
add_marker(p_cond, "{% endif %}")
add_text(p_cond, ", с другой стороны, совместно именуемые «Стороны», заключили настоящий ")
add_marker(p_cond, "{% if subject_kind == 'goods' %}")
add_branch_false(p_cond, "Договор поставки")
add_marker(p_cond, "{% else %}")
add_branch_true(p_cond, "Договор оказания услуг")
add_marker(p_cond, "{% endif %}")
add_text(p_cond, " о нижеследующем:")

# ══════════════════════════════════════════════════════════════════════════════
# 1. ПРЕДМЕТ ДОГОВОРА
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "ПРЕДМЕТ ДОГОВОРА", "1")

# 1.1 — условный блок
p_11 = clause_para(doc, "1.1")
add_marker(p_11, "{% if subject_kind == 'goods' %}")
add_branch_false(p_11, "Поставщик обязуется передать в собственность Покупателю товар (далее — Товар), "
                       "указанный в Спецификации (Приложение №1 к Договору), а Покупатель обязуется "
                       "принять Товар и оплатить его в порядке, предусмотренном Договором.")
add_marker(p_11, "{% else %}")
add_branch_true(p_11, "Исполнитель обязуется оказать Заказчику услуги по ")
add_placeholder(p_11, "{{service_subject}}")
add_branch_true(p_11, " (далее — Услуги), а Заказчик обязуется принять и оплатить оказанные Услуги. "
                      "Содержание, объём и требования к Услугам определяются Техническим заданием "
                      "(Приложение №1 к Договору), являющимся неотъемлемой частью настоящего Договора.")
add_marker(p_11, "{% endif %}")

# 1.2 — условный блок
p_12 = clause_para(doc, "1.2")
add_marker(p_12, "{% if subject_kind == 'goods' %}")
add_branch_false(p_12, "Перечень, наименование, ассортимент, количество, цена и технические характеристики "
                       "Товара установлены Спецификацией (Приложение №1).")
add_marker(p_12, "{% else %}")
add_branch_true(p_12, "Содержание, объём и порядок оказания Услуг установлены Техническим заданием (Приложение №1).")
add_marker(p_12, "{% endif %}")

# 1.3 — условный блок
p_13 = clause_para(doc, "1.3")
add_marker(p_13, "{% if subject_kind == 'goods' %}")
add_branch_false(p_13, "Поставщик гарантирует, что Товар является новым, не бывшим в эксплуатации, "
                       "свободным от прав третьих лиц.")
add_marker(p_13, "{% else %}")
add_branch_true(p_13, "Услуги оказываются Исполнителем ")
add_marker(p_13, "{% if third_party_involved %}")
add_branch_true(p_13, "с привлечением третьих лиц")
add_marker(p_13, "{% else %}")
add_branch_true(p_13, "своими силами, без привлечения третьих лиц")
add_marker(p_13, "{% endif %}")
add_branch_true(p_13, ".")
add_marker(p_13, "{% endif %}")

# 1.4 — общий для обоих
p_14 = clause_para(doc, "1.4")
add_text(p_14, "Договор заключён в рамках реализации ")
add_placeholder(p_14, "{{subsidy_agreement_text}}")
add_text(p_14, ".")

# ══════════════════════════════════════════════════════════════════════════════
# 2. СРОК И МЕСТО ИСПОЛНЕНИЯ
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "СРОК И МЕСТО ИСПОЛНЕНИЯ", "2")

# 2.1
p_21 = clause_para(doc, "2.1")
add_marker(p_21, "{% if subject_kind == 'goods' %}")
add_branch_false(p_21, "Срок поставки Товара: ")
add_marker(p_21, "{% else %}")
add_branch_true(p_21, "Срок оказания Услуг: ")
add_marker(p_21, "{% endif %}")
add_placeholder(p_21, "{{service_term}}")
add_text(p_21, ".")

# 2.2
p_22 = clause_para(doc, "2.2")
add_marker(p_22, "{% if subject_kind == 'goods' %}")
add_branch_false(p_22, "Место поставки: ")
add_marker(p_22, "{% else %}")
add_branch_true(p_22, "Место оказания Услуг: ")
add_marker(p_22, "{% endif %}")
add_placeholder(p_22, "{{delivery_location}}")
add_text(p_22, ".")

# 2.3
p_23 = clause_para(doc, "2.3")
add_marker(p_23, "{% if subject_kind == 'goods' %}")
add_branch_false(p_23, "Передача Товара оформляется товарной накладной / универсальным передаточным документом (УПД).")
add_marker(p_23, "{% else %}")
add_branch_true(p_23, "Приёмка Услуг оформляется Актом сдачи-приёмки оказанных услуг.")
add_marker(p_23, "{% endif %}")

# ══════════════════════════════════════════════════════════════════════════════
# 3. ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЁТОВ
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЁТОВ", "3")

p_31 = clause_para(doc, "3.1")
add_text(p_31, "Цена Договора составляет ")
add_placeholder(p_31, "{{contract_price_num}}")
add_text(p_31, " (")
add_placeholder(p_31, "{{contract_price_words}}")
add_text(p_31, ") рублей. ")
add_marker(p_31, "{% if vat_applicable %}")
add_branch_true(p_31, "В том числе НДС ")
add_placeholder(p_31, "{{vat_rate}}")
add_branch_true(p_31, "%: ")
add_placeholder(p_31, "{{vat_amount_num}}")
add_branch_true(p_31, " (")
add_placeholder(p_31, "{{vat_amount_words}}")
add_branch_true(p_31, ") рублей.")
add_marker(p_31, "{% else %}")
add_branch_false(p_31, "НДС не облагается на основании ")
add_placeholder(p_31, "{{vat_exemption_article}}")
add_branch_false(p_31, ".")
add_marker(p_31, "{% endif %}")

p_32 = clause_para(doc, "3.2")
add_text(p_32, "Оплата производится в безналичной форме путём перечисления денежных средств "
               "на расчётный счёт ")
add_marker(p_32, "{% if subject_kind == 'goods' %}")
add_branch_false(p_32, "Поставщика")
add_marker(p_32, "{% else %}")
add_branch_true(p_32, "Исполнителя")
add_marker(p_32, "{% endif %}")
add_text(p_32, ", указанный в разделе 7 настоящего Договора, "
               "в течение 30 (тридцати) календарных дней с даты подписания "
               "соответствующего закрывающего документа.")

clause_plain(doc, "3.3",
             "Цена Договора является фиксированной и изменению не подлежит, за исключением "
             "случаев, предусмотренных действующим законодательством Российской Федерации.")

# ══════════════════════════════════════════════════════════════════════════════
# 4. КАЧЕСТВО / ОТВЕТСТВЕННОСТЬ (условный раздел)
# ══════════════════════════════════════════════════════════════════════════════

# Заголовок раздела — условный
p_s4_if = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=0)
add_marker(p_s4_if, "{% if subject_kind == 'goods' %}", size=10)

p_s4_goods = doc.add_paragraph()
p_s4_goods.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_s4_goods.paragraph_format.space_before = Pt(0)
p_s4_goods.paragraph_format.space_after = Pt(4)
r_s4g = p_s4_goods.add_run("4. КАЧЕСТВО И ГАРАНТИИ")
r_s4g.bold = True
r_s4g.font.name = FONT_NAME
r_s4g.font.size = Pt(12)
r_s4g.font.color.rgb = COLOR_BRANCH_FALSE

p_s4_else = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
add_marker(p_s4_else, "{% else %}", size=10)

p_s4_svc = doc.add_paragraph()
p_s4_svc.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_s4_svc.paragraph_format.space_before = Pt(0)
p_s4_svc.paragraph_format.space_after = Pt(4)
r_s4s = p_s4_svc.add_run("4. ПРАВА И ОБЯЗАННОСТИ СТОРОН")
r_s4s.bold = True
r_s4s.font.name = FONT_NAME
r_s4s.font.size = Pt(12)
r_s4s.font.color.rgb = COLOR_BRANCH_TRUE

p_s4_endif = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
add_marker(p_s4_endif, "{% endif %}", size=10)

# 4.1
p_41 = clause_para(doc, "4.1")
add_marker(p_41, "{% if subject_kind == 'goods' %}")
add_branch_false(p_41, "Качество Товара должно соответствовать ГОСТ, ТУ и иным обязательным требованиям "
                       "законодательства Российской Федерации.")
add_marker(p_41, "{% else %}")
add_branch_true(p_41, "Исполнитель обязан оказать Услуги надлежащего качества в установленные сроки, "
                      "соблюдать требования действующего законодательства РФ, а также предоставлять Заказчику "
                      "по его требованию информацию о ходе оказания Услуг.")
add_marker(p_41, "{% endif %}")

# 4.2
p_42 = clause_para(doc, "4.2")
add_marker(p_42, "{% if subject_kind == 'goods' %}")
add_branch_false(p_42, "Гарантийный срок устанавливается изготовителем либо иным образом, указанным в Спецификации.")
add_marker(p_42, "{% else %}")
add_branch_true(p_42, "Заказчик обязан своевременно передавать Исполнителю материалы и информацию, необходимые "
                      "для оказания Услуг; принять результаты оказанных Услуг и оплатить их в соответствии "
                      "с разделом 3 настоящего Договора.")
add_marker(p_42, "{% endif %}")

# ══════════════════════════════════════════════════════════════════════════════
# 5. ОТВЕТСТВЕННОСТЬ СТОРОН
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "ОТВЕТСТВЕННОСТЬ СТОРОН", "5")

p_51 = clause_para(doc, "5.1")
add_marker(p_51, "{% if subject_kind == 'goods' %}")
add_branch_false(p_51, "За нарушение сроков поставки Поставщик уплачивает Покупателю неустойку "
                       "в размере 0,1% от цены Договора за каждый день просрочки.")
add_marker(p_51, "{% else %}")
add_branch_true(p_51, "За нарушение сроков оказания Услуг Исполнитель уплачивает Заказчику неустойку "
                      "в размере 0,1% от цены Договора за каждый день просрочки.")
add_marker(p_51, "{% endif %}")

p_52 = clause_para(doc, "5.2")
add_marker(p_52, "{% if subject_kind == 'goods' %}")
add_branch_false(p_52, "За нарушение сроков оплаты Покупатель уплачивает Поставщику пени "
                       "в размере 1/300 ставки рефинансирования Банка России за каждый день просрочки.")
add_marker(p_52, "{% else %}")
add_branch_true(p_52, "За нарушение сроков оплаты Заказчик уплачивает Исполнителю пени "
                      "в размере 1/300 ставки рефинансирования Банка России за каждый день просрочки.")
add_marker(p_52, "{% endif %}")

clause_plain(doc, "5.3",
             "Стороны освобождаются от ответственности за частичное или полное неисполнение "
             "обязательств, если оно явилось следствием обстоятельств непреодолимой силы (форс-мажор), "
             "возникших после заключения настоящего Договора.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. СРОК ДЕЙСТВИЯ И ПОРЯДОК РАСТОРЖЕНИЯ
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "СРОК ДЕЙСТВИЯ И ПОРЯДОК РАСТОРЖЕНИЯ ДОГОВОРА", "6")

clause_plain(doc, "6.1",
             "Договор вступает в силу с момента подписания Сторонами и действует до полного "
             "исполнения Сторонами принятых на себя обязательств.")

clause_plain(doc, "6.2",
             "Договор может быть расторгнут по соглашению Сторон, а также в одностороннем "
             "порядке по основаниям, предусмотренным действующим законодательством РФ.")

# ══════════════════════════════════════════════════════════════════════════════
# 7. АДРЕСА И РЕКВИЗИТЫ СТОРОН
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, "АДРЕСА И РЕКВИЗИТЫ СТОРОН", "7")

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"


def set_col_width(cell, cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


LEFT = tbl.rows[0].cells[0]
RIGHT = tbl.rows[0].cells[1]
set_col_width(LEFT, 8.5)
set_col_width(RIGHT, 8.5)


def req_block_colored(cell, items):
    """
    items — список: строка или tuple (kind, payload).
    kind: 'ph', 'marker', 'branch_true', 'branch_false'
    branch_true/branch_false payload: список [(seg_kind, text), ...]
    """
    para_idx = 0
    for item in items:
        if para_idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

        if isinstance(item, str):
            r = p.add_run(item)
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            if para_idx == 0:
                r.bold = True
        elif isinstance(item, tuple):
            kind = item[0]
            if kind == "marker":
                r = p.add_run(item[1])
                r.font.name = FONT_NAME
                r.font.size = Pt(9)
                r.font.color.rgb = COLOR_MARKER
                r.italic = True
            elif kind == "branch_true":
                for seg in item[1]:
                    r = p.add_run(seg[1])
                    r.font.name = FONT_NAME
                    r.font.size = Pt(10)
                    if seg[0] == "p":
                        r.font.color.rgb = COLOR_PLACEHOLDER
                        r.bold = True
                    else:
                        r.font.color.rgb = COLOR_BRANCH_TRUE
            elif kind == "branch_false":
                for seg in item[1]:
                    r = p.add_run(seg[1])
                    r.font.name = FONT_NAME
                    r.font.size = Pt(10)
                    if seg[0] == "p":
                        r.font.color.rgb = COLOR_PLACEHOLDER
                        r.bold = True
                    else:
                        r.font.color.rgb = COLOR_BRANCH_FALSE
            elif kind == "ph":
                r = p.add_run(item[1])
                r.font.name = FONT_NAME
                r.font.size = Pt(10)
                r.font.color.rgb = COLOR_PLACEHOLDER
                r.bold = True
        para_idx += 1


# Левая колонка — ЗАКАЗЧИК / ПОКУПАТЕЛЬ
LEFT_ITEMS = [
    # Заголовок — условный
    ("marker", "{% if subject_kind == 'goods' %}"),
    ("branch_false", [("t", "ПОКУПАТЕЛЬ:")]),
    ("marker", "{% else %}"),
    ("branch_true", [("t", "ЗАКАЗЧИК:")]),
    ("marker", "{% endif %}"),
    ("ph", "{{customer_full_name}}"),
    ("ph", "Адрес: {{customer_address}}"),
    ("ph", "Почтовый адрес: {{customer_postal_address}}"),
    ("ph", "ИНН: {{customer_inn}} / КПП: {{customer_kpp}}"),
    ("ph", "ОГРН: {{customer_ogrn}}"),
    ("ph", "Банк: {{customer_bank_name}}"),
    ("ph", "БИК: {{customer_bik}}"),
    ("ph", "р/с: {{customer_settlement_account}}"),
    ("ph", "к/с: {{customer_correspondent_account}}"),
    ("ph", "Тел.: {{customer_phone}}"),
    ("ph", "E-mail: {{customer_email}}"),
    "",
    ("ph", "{{customer_signatory_position}}"),
    ("ph", "________________ / {{customer_signatory_initials}}"),
    "М.П.",
]

req_block_colored(LEFT, LEFT_ITEMS)

# Правая колонка — ИСПОЛНИТЕЛЬ / ПОСТАВЩИК с условным блоком ИП / ООО
RIGHT_ITEMS = [
    # Заголовок — условный
    ("marker", "{% if subject_kind == 'goods' %}"),
    ("branch_false", [("t", "ПОСТАВЩИК:")]),
    ("marker", "{% else %}"),
    ("branch_true", [("t", "ИСПОЛНИТЕЛЬ:")]),
    ("marker", "{% endif %}"),
    ("marker", "{%- if contractor_org_type == 'ИП' %}"),
    ("branch_true", [("t", "ИП "), ("p", "{{contractor_short_name}}")]),
    ("branch_true", [("t", "ИНН: "), ("p", "{{contractor_inn}}")]),
    ("branch_true", [("t", "ОГРНИП: "), ("p", "{{contractor_ogrnip}}")]),
    ("marker", "{%- else %}"),
    ("branch_false", [("p", "{{contractor_full_name}}")]),
    ("branch_false", [("t", "ИНН: "), ("p", "{{contractor_inn}}"), ("t", " / КПП: "), ("p", "{{contractor_kpp}}")]),
    ("branch_false", [("t", "ОГРН: "), ("p", "{{contractor_ogrn}}")]),
    ("marker", "{%- endif %}"),
    ("ph", "Адрес: {{contractor_address}}"),
    ("ph", "Банк: {{contractor_bank_name}}"),
    ("ph", "БИК: {{contractor_bik}}"),
    ("ph", "р/с: {{contractor_settlement_account}}"),
    ("ph", "к/с: {{contractor_correspondent_account}}"),
    ("ph", "Тел.: {{contractor_phone}}"),
    ("ph", "E-mail: {{contractor_email}}"),
    "",
    ("ph", "{{contractor_signatory_position}}"),
    ("ph", "________________ / {{contractor_signatory_initials}}"),
    "М.П.",
]

req_block_colored(RIGHT, RIGHT_ITEMS)

# Приложение
doc.add_paragraph()
p_app = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=6)
add_marker(p_app, "{% if subject_kind == 'goods' %}")
add_branch_false(p_app, "Приложение №1: Спецификация товаров (на __ л.).")
add_marker(p_app, "{% else %}")
add_branch_true(p_app, "Приложение №1: Техническое задание (на __ л.).")
add_marker(p_app, "{% endif %}")

# ── СОХРАНЕНИЕ ───────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"OK: {OUTPUT}")
