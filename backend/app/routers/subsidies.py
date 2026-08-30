import io
import os
import re
import shutil
import logging
from urllib.parse import quote
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query, Request
from fastapi.responses import FileResponse, StreamingResponse


def _content_disposition(filename: str) -> str:
    """RFC 5987 — кириллица в имени файла недопустима в latin-1 заголовке."""
    ascii_fallback = filename.encode('ascii', 'ignore').decode('ascii').strip() or 'export'
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{quote(filename)}"

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
except ImportError:
    openpyxl = None

try:
    from docxtpl import DocxTemplate
except ImportError:
    DocxTemplate = None

logger = logging.getLogger(__name__)
from sqlalchemy import select, delete, func, text, or_
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.exc import IntegrityError, ProgrammingError
from app.database import get_db, engine
from app.models.subsidy import Subsidy
from app.models.feo_category import FeoCategory
from app.models.feo_planned_item import FeoPlannedItem
from app.models.purchase import Purchase
from app.models.subsidy_contractor_override import SubsidyContractorOverride
from app.models.contractor import Contractor
from app.schemas.schemas import (
    SubsidyCreate, SubsidyOut,
    SubsidyContractorOverrideCreate, SubsidyContractorOverrideOut,
)
from app.auth.jwt import get_current_user, require_role, get_org_filter, get_single_org_id, MANAGER_ROLES, ADMIN_ROLES, ALL_ROLES
from app.auth.permissions import require_tab, require_action
from app.models.user import User
from app.services.fio import compose_fio
from app.services.feo_plan import calculate_ceiling_forecasts_bulk, calculate_ceiling_forecast
from typing import List, Optional
from datetime import date
from pydantic import BaseModel

router = APIRouter(prefix="/api/subsidies", tags=["subsidies"])


@router.get("/diag/columns")
async def diag_columns(db: AsyncSession = Depends(get_db)):
    """Diagnostic endpoint — возвращает реальные колонки таблицы subsidies из information_schema."""
    result = await db.execute(text(
        "SELECT column_name, data_type, is_nullable "
        "FROM information_schema.columns "
        "WHERE table_name = 'subsidies' "
        "ORDER BY ordinal_position"
    ))
    rows = result.fetchall()
    columns = [{"column_name": r[0], "data_type": r[1], "is_nullable": r[2]} for r in rows]
    logger.info("diag_columns: subsidies has %d columns: %s", len(columns), [c["column_name"] for c in columns])
    return {"table": "subsidies", "columns": columns}


def _budget_from_tree(all_categories) -> float:
    """
    Рекурсивный подсчёт бюджета из ФЭО-дерева одной субсидии:
    - Листовой узел (нет детей): берём его budget
    - Родительский узел: если budget задан вручную (not null) → берём его,
                         иначе сумма детей (рекурсивно)
    Итог = сумма по всем корневым (level-1) узлам.
    """
    if not all_categories:
        return 0.0

    by_id = {c.id: c for c in all_categories}
    children_map: dict = {}
    for c in all_categories:
        children_map.setdefault(c.id, [])
        if c.parent_id and c.parent_id in by_id:
            children_map.setdefault(c.parent_id, []).append(c)

    def _calc_node(cat) -> float:
        kids = children_map.get(cat.id, [])
        if not kids:
            return float(cat.budget) if cat.budget is not None else 0.0
        if cat.budget is not None:
            return float(cat.budget)
        return sum(_calc_node(k) for k in kids)

    roots = [c for c in all_categories if c.level == 1]
    return sum(_calc_node(r) for r in roots)


async def calculate_budget_from_categories(db: AsyncSession, subsidy_id: int) -> float:
    result = await db.execute(
        select(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id)
    )
    return _budget_from_tree(result.scalars().all())


async def calculate_budgets_bulk(db: AsyncSession, subsidy_ids: list[int]) -> dict[int, float]:
    """Бюджеты для набора субсидий одним запросом (вместо N запросов в списках)."""
    if not subsidy_ids:
        return {}
    result = await db.execute(
        select(FeoCategory).where(FeoCategory.subsidy_id.in_(subsidy_ids))
    )
    by_subsidy: dict[int, list] = {}
    for c in result.scalars().all():
        by_subsidy.setdefault(c.subsidy_id, []).append(c)
    return {sid: _budget_from_tree(by_subsidy.get(sid, [])) for sid in subsidy_ids}


# Статусы, исключаемые из подсчёта потраченного (отменённые закупки не занимают бюджет)
_EXCLUDED_FROM_SPENT: tuple = ("cancelled",)


async def _calculate_spent(
    db: AsyncSession,
    subsidy_id: int,
    exclude_purchase_id: Optional[int] = None,
) -> float:
    """Σ planned_total_price по активным закупкам субсидии (кроме cancelled)."""
    q = select(func.coalesce(func.sum(Purchase.planned_total_price), 0)).where(
        Purchase.subsidy_id == subsidy_id,
        Purchase.status.notin_(_EXCLUDED_FROM_SPENT),
    )
    if exclude_purchase_id:
        q = q.where(Purchase.id != exclude_purchase_id)
    result = await db.execute(q)
    return float(result.scalar() or 0)


async def _calculate_planned_amount(db: AsyncSession, subsidy_id: int) -> float:
    """Σ FeoPlannedItem.amount для субсидии (ФЭО плановая сумма)."""
    q = select(func.coalesce(func.sum(FeoPlannedItem.amount), 0)).join(
        FeoCategory, FeoPlannedItem.feo_category_id == FeoCategory.id
    ).where(
        FeoCategory.subsidy_id == subsidy_id,
        FeoPlannedItem.is_active == True,
    )
    result = await db.execute(q)
    return float(result.scalar() or 0)


async def _calculate_planned_amounts_bulk(
    db: AsyncSession, subsidy_ids: list[int]
) -> dict[int, float]:
    """Batch Σ FeoPlannedItem.amount для набора субсидий."""
    if not subsidy_ids:
        return {}
    q = (
        select(FeoCategory.subsidy_id, func.coalesce(func.sum(FeoPlannedItem.amount), 0).label("total"))
        .join(FeoPlannedItem, FeoPlannedItem.feo_category_id == FeoCategory.id)
        .where(
            FeoCategory.subsidy_id.in_(subsidy_ids),
            FeoPlannedItem.is_active == True,
        )
        .group_by(FeoCategory.subsidy_id)
    )
    rows = (await db.execute(q)).all()
    result = {sid: 0.0 for sid in subsidy_ids}
    for r in rows:
        result[r.subsidy_id] = float(r.total)
    return result


async def _calculate_spent_bulk(
    db: AsyncSession, subsidy_ids: list[int]
) -> dict[int, float]:
    """Batch Σ planned_total_price для набора субсидий (кроме cancelled)."""
    if not subsidy_ids:
        return {}
    q = (
        select(Purchase.subsidy_id, func.coalesce(func.sum(Purchase.planned_total_price), 0).label("spent"))
        .where(
            Purchase.subsidy_id.in_(subsidy_ids),
            Purchase.status.notin_(_EXCLUDED_FROM_SPENT),
        )
        .group_by(Purchase.subsidy_id)
    )
    rows = (await db.execute(q)).all()
    result = {sid: 0.0 for sid in subsidy_ids}
    for r in rows:
        result[r.subsidy_id] = float(r.spent)
    return result


async def _calculate_feo_planned_tree_bulk(
    db: AsyncSession, subsidy_ids: list[int]
) -> dict[int, float]:
    """Плановая сумма дерева ФЭО для набора субсидий (единый источник для «Запланировано»).

    Тонкая обёртка над app.services.feo_plan.feo_plan_subsidy_totals — Σ display
    корневых узлов дерева ФЭО, где для каждого узла (листа и группы):
      display = MAX(план, выбрано) + сверх_плана
    План/выбрано/сверх план считаются рекурсивно по поддереву (см. подробный
    docstring compute_feo_plan_tree в app/services/feo_plan.py). Формула вынесена
    в общий сервис, чтобы:
      1) устранить задвоение, когда позиция закупки лежит ОДНОВРЕМЕННО на группе
         и на её дочернем листе (раньше их суммы складывались напрямую — группа
         «Внедорожник повышенной проходимости» показывала 16 380 000 = план листа
         8 380 000 + позиция заявки на самой группе 8 000 000, вместо MAX = 8 380 000);
      2) не расходиться с _create_plan_graph_version (purchases.py) — снапшот
         истории версий плана закупок использует ту же compute_feo_plan_tree.

    Совпадает с тем, что показывает панель ФЭО вкладки «Субсидии»
    (feoPlannedDisplayFor(root) в SubsidiesView.vue, режим 'all' — реализует ту же
    формулу MAX(план, выбрано)+сверх_план на фронте, см. docstring там).
    """
    if not subsidy_ids:
        return {}
    from app.services.feo_plan import feo_plan_subsidy_totals
    return await feo_plan_subsidy_totals(db, subsidy_ids)


@router.get("/", response_model=List[SubsidyOut])
async def list_subsidies(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
    scope: Optional[str] = Query(None, description="strict|wishes — строгий скоуп пикера: прямое членство + управляемые орги + гранты, БЕЗ контура-детей"),
):
    # GET list is intentionally open to all authenticated users (incl. employee) so
    # that WishesView can populate the subsidy selector. Org-filter below already
    # limits visible subsidies to the user's own organisations. Write operations
    # (POST/PUT/DELETE/templates) remain gated by require_tab('subsidies').
    q = select(Subsidy).order_by(Subsidy.year.desc(), Subsidy.name)
    # Shared picker endpoint (orders/wishes/etc): member-union scope, NOT role-scoped.
    # Сужение до орг-админ-орг ломало бы выбор субсидии у сотрудника в /orders.
    org_ids = get_org_filter(current_user)
    # scope="wishes"/"strict" — строгий скоуп пикера заявки: пользователь видит
    # субсидии ТОЛЬКО тех орг, где реально состоит (членство/primary/UOA) или которыми
    # управляет, + личные гранты. Контур аккаунта (root+дети) НЕ применяется: менеджер
    # ВСКС не должен видеть субсидии дочернего Центрпоиска только из-за дерева орг —
    # наследование по дереву признано неверным (видимость даёт членство/управление).
    if scope in ("wishes", "strict") and current_user.role not in ('superadmin', 'account_owner'):
        from app.models.user_organization import UserOrganization
        from app.models.user_org_access import UserOrgAccess
        from app.models.manager_organization import ManagerOrganization
        strict: set[int] = set()
        if current_user.org_id:
            strict.add(int(current_user.org_id))
        strict |= {int(r[0]) for r in (await db.execute(
            select(UserOrganization.org_id).where(UserOrganization.user_id == current_user.id)
        )).all() if r[0]}
        strict |= {int(r[0]) for r in (await db.execute(
            select(UserOrgAccess.org_id).where(UserOrgAccess.user_id == current_user.id)
        )).all() if r[0]}
        strict |= {int(r[0]) for r in (await db.execute(
            select(ManagerOrganization.org_id).where(ManagerOrganization.manager_user_id == current_user.id)
        )).all() if r[0]}
        org_ids = list(strict)
    # 27.4-06: hard fallback для не-SaaS-ролей. Если JWT не содержит org_id/org_ids
    # И users.org_id = NULL — get_org_filter вернёт None → раньше employee видел ВСЕ
    # субсидии. Подтягиваем memberships из user_organizations.
    if org_ids is None and current_user.role not in ('superadmin', 'account_owner'):
        from app.models.user_organization import UserOrganization
        uo_rows = (await db.execute(
            select(UserOrganization.org_id).where(UserOrganization.user_id == current_user.id)
        )).all()
        org_ids = list({r[0] for r in uo_rows if r[0]})
        if not org_ids and current_user.org_id:
            org_ids = [current_user.org_id]
    # Per-user explicit subsidy grants (expand visibility regardless of org)
    from app.models.user_subsidy_access import UserSubsidyAccess
    granted_ids = set((await db.execute(
        select(UserSubsidyAccess.subsidy_id).where(UserSubsidyAccess.user_id == current_user.id)
    )).scalars().all())
    # Пикер (orders/wishes): пермиссивный member-union scope + кросс-орг гранты.
    # Управление ВИДИМОСТЬЮ субсидий (страница «Субсидии») — отдельно, в dashboard
    # charts через get_visible_subsidy_ids. Здесь НЕ сужаем до грантов, иначе ломается
    # выбор субсидии у рядового сотрудника.
    if org_ids is not None:
        if not org_ids and not granted_ids:
            return []
        q = q.where(or_(Subsidy.org_id.in_(org_ids), Subsidy.id.in_(granted_ids)))
    result = await db.execute(q)
    subsidies = result.scalars().all()

    # Batch-загрузка вместо N+1: бюджеты/контрагенты/орги одним IN-запросом каждый
    from app.models.organization import Organization
    sid_list = [s.id for s in subsidies]
    budgets = await calculate_budgets_bulk(db, sid_list)
    spent_map = await _calculate_spent_bulk(db, sid_list)
    planned_amounts = await _calculate_planned_amounts_bulk(db, sid_list)
    # Владелец (2026-08-30): предупреждение «сумма заказанного приближается к
    # потолку субсидии» — считаем батчем на список, не по одной (см. docstring
    # calculate_ceiling_forecasts_bulk в app/services/feo_plan.py).
    ceiling_forecasts = await calculate_ceiling_forecasts_bulk(db, sid_list)
    contractor_ids = {s.contractor_id for s in subsidies if s.contractor_id}
    contractors = {}
    if contractor_ids:
        rows = (await db.execute(
            select(Contractor).where(Contractor.id.in_(contractor_ids))
        )).scalars().all()
        contractors = {c.id: c for c in rows}
    subsidy_org_ids = {s.org_id for s in subsidies if s.org_id}
    orgs = {}
    if subsidy_org_ids:
        rows = (await db.execute(
            select(Organization).where(Organization.id.in_(subsidy_org_ids))
        )).scalars().all()
        orgs = {o.id: o for o in rows}

    out = []
    for s in subsidies:
        calc = budgets.get(s.id, 0.0)
        effective_budget = calc if calc > 0 else float(s.budget or 0)
        s.calculated_budget = effective_budget
        # Решение 14.07: budget — ручное значение, деревом ФЭО НЕ перезаписывается
        spent = spent_map.get(s.id, 0.0)
        planned_amt = planned_amounts.get(s.id, 0.0)
        remaining = effective_budget - spent
        discrepancy = (effective_budget - planned_amt) if abs(effective_budget - planned_amt) > 0.01 else None

        d = {c.name: getattr(s, c.name) for c in s.__table__.columns}
        d["calculated_budget"] = effective_budget
        d["feo_filled"] = calc > 0
        d["feo_budget_total"] = effective_budget
        d["remaining"] = remaining
        d["planned_amount"] = planned_amt
        d["budget_discrepancy"] = discrepancy
        contractor = contractors.get(s.contractor_id) if s.contractor_id else None
        d["contractor_name"] = contractor.name if contractor else None
        d["contractor_inn"] = contractor.inn if contractor else None
        org = orgs.get(s.org_id) if s.org_id else None
        d["org_inn"] = org.inn if org else None
        d.update(ceiling_forecasts.get(s.id, {}))
        out.append(d)

    await db.commit()
    return out

@router.get("/{subsidy_id}", response_model=SubsidyOut)
async def get_subsidy(
    subsidy_id: int,
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))
    subsidy = result.scalar_one_or_none()
    if not subsidy:
        raise HTTPException(status_code=404, detail="Subsidy not found")

    calc = await calculate_budget_from_categories(db, subsidy.id)
    effective_budget = calc if calc > 0 else float(subsidy.budget or 0)
    subsidy.calculated_budget = effective_budget
    # Решение 14.07: budget — ручное значение, деревом ФЭО НЕ перезаписывается
    await db.commit()

    spent = await _calculate_spent(db, subsidy.id)
    planned_amt = await _calculate_planned_amount(db, subsidy.id)
    remaining = effective_budget - spent
    discrepancy = (effective_budget - planned_amt) if abs(effective_budget - planned_amt) > 0.01 else None

    d = {c.name: getattr(subsidy, c.name) for c in subsidy.__table__.columns}
    d["calculated_budget"] = effective_budget
    d["feo_filled"] = calc > 0
    d["feo_budget_total"] = effective_budget
    d["remaining"] = remaining
    d["planned_amount"] = planned_amt
    d["budget_discrepancy"] = discrepancy
    if subsidy.contractor_id:
        contractor = await db.get(Contractor, subsidy.contractor_id)
        d["contractor_name"] = contractor.name if contractor else None
        d["contractor_inn"] = contractor.inn if contractor else None
    else:
        d["contractor_name"] = None
        d["contractor_inn"] = None
    d.update(await calculate_ceiling_forecast(db, subsidy.id))
    return d


@router.get("/{subsidy_id}/budget-check")
async def budget_check(
    subsidy_id: int,
    exclude_purchase_id: Optional[int] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Canonical budget-check endpoint (D-17).

    Returns limit (ФЭО-tree budget), spent (Σ active purchases), remaining, discrepancy.
    T-31-05-01: read-access validated — only subsidies visible to user.
    """
    from app.auth.visibility import get_visible_subsidy_ids, build_visibility_clause
    # Security: ensure user has read-access to this subsidy
    visible = await get_visible_subsidy_ids(current_user, db)
    if visible is not None and subsidy_id not in visible:
        # Пользователь может не иметь вкладки «Субсидии» в орге субсидии, но видеть
        # закупку по ней (назначен/участник/в чате/задача) — тогда бюджет субсидии
        # ему доступен для контекста закупки. Иначе открытие чужой (по орг) закупки
        # падало 404 «Subsidy not found» у исполнителя.
        clause = await build_visibility_clause(current_user, db, "purchase")
        allowed = clause is None
        if clause is not None:
            cnt = (await db.execute(
                select(func.count()).select_from(Purchase).where(
                    Purchase.subsidy_id == subsidy_id, clause
                )
            )).scalar() or 0
            allowed = cnt > 0
        if not allowed:
            raise HTTPException(status_code=404, detail="Subsidy not found")

    subsidy = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not subsidy:
        raise HTTPException(status_code=404, detail="Subsidy not found")

    limit = await calculate_budget_from_categories(db, subsidy_id)
    spent = await _calculate_spent(db, subsidy_id, exclude_purchase_id=exclude_purchase_id)
    planned_amt = await _calculate_planned_amount(db, subsidy_id)
    remaining = limit - spent
    discrepancy = (limit - planned_amt) if abs(limit - planned_amt) > 0.01 else 0.0

    return {
        "limit": limit,
        "spent": spent,
        "remaining": remaining,
        "planned_amount": planned_amt,
        "discrepancy": discrepancy,
    }


async def _materialize_org_from_contractor(db: AsyncSession, contractor):
    """Find-or-create Organization, зеркалящую контрагента (получатель субсидии),
    чтобы он попадал в контур (Персонал/Иерархия). Матч по contractor_id, затем по ИНН.
    Идемпотентно. Возвращает Organization (или None) — вызывающие используют её как
    org_id субсидии: субсидия принадлежит организации-грантополучателю, а не активной
    орг из свитчера (кейс ДНР_2026: владелец был ЦЕНТРПОИСК вместо Донецкого)."""
    if contractor is None:
        return None
    from app.models.organization import Organization
    org = (await db.execute(
        select(Organization).where(Organization.contractor_id == contractor.id)
    )).scalars().first()
    if org is None and contractor.inn:
        org = (await db.execute(
            select(Organization).where(Organization.inn == contractor.inn)
        )).scalars().first()
    if org is None:
        org = Organization(
            name=(contractor.name or "")[:255] or f"Контрагент #{contractor.id}",
            full_name=(contractor.full_name or None),
            inn=contractor.inn,
            kpp=contractor.kpp,
            ogrn=contractor.ogrn,
            address=((contractor.address or "")[:500] or None),
            signatory=((contractor.signatory or "")[:500] or None),
            contractor_id=contractor.id,
            is_active=True,
        )
        db.add(org)
        await db.flush()
    elif org.contractor_id is None:
        org.contractor_id = contractor.id
    await db.commit()
    return org


async def _merge_duplicate_orgs_by_inn(db: AsyncSession) -> None:
    """Идемпотентный мерж дублей organizations по ИНН.

    Алгоритм:
    - survivor = MIN(id) в группе (стабильный якорь).
    - Победа «свежих данных»: для каждого поля берём непустое значение из строки с MAX(id).
    - Все FK перепривязываются через information_schema (generic), SAVEPOINT защищает
      от конфликтов уникальных индексов (junction-таблицы).
    - losers удаляются, commit выполняется после всех групп.
    - Функция НЕ падает фатально — ошибка группы логируется и пропускается.
    """
    _log = logging.getLogger(__name__)

    # 1. Найти группы дублей по ИНН
    dup_result = await db.execute(text(
        "SELECT inn, array_agg(id ORDER BY id) AS ids "
        "FROM organizations "
        "WHERE inn IS NOT NULL AND btrim(inn) <> '' "
        "GROUP BY inn HAVING count(*) > 1"
    ))
    groups = dup_result.fetchall()

    if not groups:
        _log.info("org-merge по ИНН: дублей не найдено")
        return

    # 2. Собрать FK-ссылки на organizations.id через information_schema (один раз)
    fk_result = await db.execute(text(
        "SELECT tc.table_name, kcu.column_name "
        "FROM information_schema.table_constraints tc "
        "JOIN information_schema.key_column_usage kcu "
        "  ON tc.constraint_name = kcu.constraint_name "
        "  AND tc.table_schema = kcu.table_schema "
        "JOIN information_schema.constraint_column_usage ccu "
        "  ON tc.constraint_name = ccu.constraint_name "
        "  AND tc.table_schema = ccu.table_schema "
        "WHERE tc.constraint_type = 'FOREIGN KEY' "
        "  AND ccu.table_name = 'organizations' "
        "  AND ccu.column_name = 'id'"
    ))
    fk_refs = fk_result.fetchall()  # список (table_name, column_name)

    # Поля для «свежих данных»
    _merge_fields = [
        "name", "full_name", "kpp", "ogrn", "address",
        "signatory", "color", "contractor_id", "is_active",
    ]

    total_deleted = 0
    n_groups = len(groups)

    for row in groups:
        inn_val = row[0]
        ids = row[1]  # уже упорядочены по возрастанию
        survivor_id = ids[0]
        loser_ids = ids[1:]

        try:
            # 3. Загрузить все строки группы и выбрать «свежие» непустые поля
            all_rows_result = await db.execute(
                text("SELECT id, name, full_name, kpp, ogrn, address, signatory, "
                     "color, contractor_id, is_active "
                     "FROM organizations WHERE id = ANY(:ids) ORDER BY id DESC"),
                {"ids": ids}
            )
            all_rows = all_rows_result.fetchall()

            # Для каждого поля: берём первое непустое значение при обходе от MAX к MIN id
            updates = {}
            for field_idx, field in enumerate(_merge_fields):
                for r in all_rows:
                    val = r[field_idx + 1]  # +1 т.к. id в col 0
                    if val is None:
                        continue
                    if isinstance(val, str) and val.strip() == "":
                        continue
                    updates[field] = val
                    break

            # Получить текущие значения survivor для сравнения
            surv_row_result = await db.execute(
                text("SELECT name, full_name, kpp, ogrn, address, signatory, "
                     "color, contractor_id, is_active "
                     "FROM organizations WHERE id = :sid"),
                {"sid": survivor_id}
            )
            surv_row = surv_row_result.fetchone()
            surv_dict = dict(zip(_merge_fields, surv_row)) if surv_row else {}

            # Применить только изменившиеся поля
            changed = {k: v for k, v in updates.items() if surv_dict.get(k) != v}
            if changed:
                set_clause = ", ".join(f"{k} = :{k}" for k in changed)
                changed["_sid"] = survivor_id
                await db.execute(
                    text(f"UPDATE organizations SET {set_clause} WHERE id = :_sid"),
                    changed
                )

            # 4. Перепривязать все FK (generic)
            for fk_table, fk_col in fk_refs:
                # Пропускаем self-reference organizations.id на organizations.id
                # (т.е. если таблица=organizations и колонка=id — это PK, не FK)
                # Но root_org_id — это FK и должен перепривязаться, он пройдёт как обычно
                try:
                    async with db.begin_nested():
                        await db.execute(
                            text(f'UPDATE "{fk_table}" SET "{fk_col}" = :surv '
                                 f'WHERE "{fk_col}" = ANY(:losers)'),
                            {"surv": survivor_id, "losers": loser_ids}
                        )
                except Exception as upd_err:
                    # Конфликт уникального индекса — fallback: удалить loser-строки
                    _log.warning(
                        f"org-merge: UPDATE {fk_table}.{fk_col} конфликт ({upd_err!r}), "
                        f"fallback DELETE loser-строк"
                    )
                    try:
                        async with db.begin_nested():
                            await db.execute(
                                text(f'DELETE FROM "{fk_table}" '
                                     f'WHERE "{fk_col}" = ANY(:losers)'),
                                {"losers": loser_ids}
                            )
                    except Exception as del_err:
                        _log.warning(
                            f"org-merge: DELETE fallback {fk_table}.{fk_col} тоже провалился: {del_err!r}"
                        )

            # 5. Удалить losers
            await db.execute(
                text("DELETE FROM organizations WHERE id = ANY(:losers)"),
                {"losers": loser_ids}
            )
            total_deleted += len(loser_ids)

        except Exception as group_err:
            _log.warning(
                f"org-merge: группа ИНН={inn_val!r} пропущена из-за ошибки: {group_err!r}"
            )
            continue

    await db.commit()
    _log.info(
        f"org-merge по ИНН: групп {n_groups}, удалено дублей {total_deleted}"
    )


@router.post("/", response_model=SubsidyOut)
async def create_subsidy(
    subsidy: SubsidyCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_tab('subsidies')),
):
    data = subsidy.dict()
    # org_id субсидии = организация-грантополучатель (из контрагента), НЕ активная орг
    # из свитчера: иначе субсидия «уезжает» под чужую орг (кейс ДНР_2026 → ЦЕНТРПОИСК).
    _contractor = None
    _grantee_org = None
    if data.get('contractor_id'):
        _contractor = await db.get(Contractor, data['contractor_id'])
        _grantee_org = await _materialize_org_from_contractor(db, _contractor)
    if _grantee_org is not None:
        data['org_id'] = _grantee_org.id
    else:
        data['org_id'] = get_single_org_id(current_user) or current_user.org_id
    _new_name = (data.get('name') or '').strip()
    if _new_name:
        _dup = (await db.execute(
            select(Subsidy.id).where(func.lower(func.trim(Subsidy.name)) == _new_name.lower())
        )).first()
        if _dup:
            raise HTTPException(status_code=409, detail=f"Субсидия с названием «{_new_name}» уже существует")
    db_subsidy = Subsidy(**data)
    db.add(db_subsidy)
    await db.commit()
    await db.refresh(db_subsidy)

    d = {c.name: getattr(db_subsidy, c.name) for c in db_subsidy.__table__.columns}
    d["calculated_budget"] = 0.0
    d["feo_filled"] = False
    d["feo_budget_total"] = 0.0
    if _contractor is not None:
        d["contractor_name"] = _contractor.name
        d["contractor_inn"] = _contractor.inn
    else:
        d["contractor_name"] = None
        d["contractor_inn"] = None
    d.update(await calculate_ceiling_forecast(db, db_subsidy.id))
    return d

@router.put("/{subsidy_id}", response_model=SubsidyOut)
async def update_subsidy(
    subsidy_id: int,
    subsidy: SubsidyCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_action('subsidy.edit')),
):
    result = await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))
    db_subsidy = result.scalar_one_or_none()
    if not db_subsidy:
        raise HTTPException(status_code=404, detail="Subsidy not found")

    # Wave 3: орг-осознанная проверка — право subsidy.edit должно действовать
    # именно в орге ЭТОЙ субсидии (орг-роль не даёт власти в чужих оргах).
    from app.auth.permissions import has_org_key
    if not await has_org_key(current_user, db, db_subsidy.org_id, 'subsidy.edit', subsidy_id=subsidy_id):
        raise HTTPException(
            status_code=403,
            detail="Нет права редактировать субсидии этой организации: право «Редактирование субсидий» не выдано для организации-грантополучателя",
        )

    old_budget = db_subsidy.budget  # capture BEFORE setattr loop

    # Step 1: log incoming payload
    payload = subsidy.model_dump() if hasattr(subsidy, 'model_dump') else subsidy.dict()
    logger.info("update_subsidy id=%s incoming payload: %s", subsidy_id, payload)

    calc = await calculate_budget_from_categories(db, subsidy_id)

    # Step 2: try normal setattr → commit
    _upd_name = (payload.get('name') or '').strip()
    if _upd_name:
        _dup = (await db.execute(
            select(Subsidy.id).where(
                func.lower(func.trim(Subsidy.name)) == _upd_name.lower(),
                Subsidy.id != subsidy_id,
            )
        )).first()
        if _dup:
            raise HTTPException(status_code=409, detail=f"Субсидия с названием «{_upd_name}» уже существует")
    for key, value in payload.items():
        setattr(db_subsidy, key, value)
    db_subsidy.calculated_budget = calc

    # Budget history write hook — track subsidy limit changes only (NOT calculated_budget)
    if old_budget != db_subsidy.budget:
        from app.models.budget_history import BudgetHistory as _BH
        db.add(_BH(
            subsidy_id=subsidy_id,
            purchase_id=None,
            entity_type="subsidy",
            old_value=float(old_budget) if old_budget is not None else None,
            new_value=float(db_subsidy.budget) if db_subsidy.budget is not None else None,
            changed_by_id=current_user.id,
            changed_by_name=getattr(current_user, 'full_name', None) or current_user.username,
            reason=None,
        ))

    logger.info(
        "update_subsidy id=%s db_subsidy state before commit: name=%r year=%r budget=%r "
        "basis_doc_number=%r basis_doc_date=%r",
        subsidy_id,
        db_subsidy.name, db_subsidy.year, db_subsidy.budget,
        getattr(db_subsidy, 'basis_doc_number', '<attr_missing>'),
        getattr(db_subsidy, 'basis_doc_date', '<attr_missing>'),
    )

    try:
        await db.commit()
        await db.refresh(db_subsidy)
    except (ProgrammingError, IntegrityError) as exc:
        logger.warning(
            "update_subsidy id=%s commit failed (%s: %s), attempting ALTER fallback",
            subsidy_id, type(exc).__name__, exc,
        )
        await db.rollback()

        # ALTER fallback — отдельная транзакция вне текущей сессии
        from app.database import ensure_phase22_columns as _ensure_p22
        await _ensure_p22()

        # Перезагружаем объект и повторяем setattr
        result2 = await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))
        db_subsidy = result2.scalar_one_or_none()
        if not db_subsidy:
            raise HTTPException(status_code=404, detail="Subsidy not found after ALTER fallback")
        db.expire_all()

        for key, value in payload.items():
            setattr(db_subsidy, key, value)
        db_subsidy.calculated_budget = calc

        if old_budget != db_subsidy.budget:
            from app.models.budget_history import BudgetHistory as _BH2
            db.add(_BH2(
                subsidy_id=subsidy_id,
                purchase_id=None,
                entity_type="subsidy",
                old_value=float(old_budget) if old_budget is not None else None,
                new_value=float(db_subsidy.budget) if db_subsidy.budget is not None else None,
                changed_by_id=current_user.id,
                changed_by_name=getattr(current_user, 'full_name', None) or current_user.username,
                reason=None,
            ))

        await db.commit()
        await db.refresh(db_subsidy)
        logger.info("update_subsidy id=%s ALTER fallback commit succeeded", subsidy_id)

    # Step 4: raw SELECT to verify final DB state
    raw = await db.execute(
        text("SELECT id, name, year, budget, basis_doc_number, basis_doc_date FROM subsidies WHERE id = :id"),
        {"id": subsidy_id},
    )
    raw_row = raw.fetchone()
    logger.info("update_subsidy id=%s post-commit raw SELECT: %s", subsidy_id, raw_row)

    d = {c.name: getattr(db_subsidy, c.name) for c in db_subsidy.__table__.columns}
    d["calculated_budget"] = calc
    d["feo_filled"] = calc > 0
    d["feo_budget_total"] = calc
    if db_subsidy.contractor_id:
        contractor = await db.get(Contractor, db_subsidy.contractor_id)
        d["contractor_name"] = contractor.name if contractor else None
        d["contractor_inn"] = contractor.inn if contractor else None
        grantee_org = await _materialize_org_from_contractor(db, contractor)
        # Субсидия принадлежит организации-грантополучателю — перепривязываем
        # владельца при смене контрагента (expire_on_commit=False, доступ безопасен).
        if grantee_org is not None and db_subsidy.org_id != grantee_org.id:
            db_subsidy.org_id = grantee_org.id
            await db.commit()
            d["org_id"] = grantee_org.id
    else:
        d["contractor_name"] = None
        d["contractor_inn"] = None
    d.update(await calculate_ceiling_forecast(db, db_subsidy.id))
    return d

@router.get("/{subsidy_id}/delete-impact")
async def subsidy_delete_impact(
    subsidy_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_tab('subsidies')),
):
    """Counts of dependent rows so the UI can warn before deleting a subsidy."""
    from app.models.purchase import Purchase
    from app.models.contract import Contract
    from app.models.feo_planned_item import FeoPlannedItem
    feo_count = await db.scalar(
        select(func.count()).select_from(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id)
    )
    planned_count = await db.scalar(
        select(func.count()).select_from(FeoPlannedItem)
        .join(FeoCategory, FeoPlannedItem.feo_category_id == FeoCategory.id)
        .where(FeoCategory.subsidy_id == subsidy_id)
    )
    p_count = await db.scalar(
        select(func.count()).select_from(Purchase).where(Purchase.subsidy_id == subsidy_id)
    )
    c_count = await db.scalar(
        select(func.count()).select_from(Contract).where(Contract.subsidy_id == subsidy_id)
    )
    return {
        "feo_categories": feo_count or 0,
        "planned_items": planned_count or 0,
        "purchases": p_count or 0,
        "contracts": c_count or 0,
    }

@router.delete("/{subsidy_id}")
async def delete_subsidy(
    subsidy_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_tab('subsidies')),
):
    result = await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))
    db_subsidy = result.scalar_one_or_none()
    if not db_subsidy:
        raise HTTPException(status_code=404, detail="Subsidy not found")

    # Wave 3: удалять субсидию можно только при праве на вкладку «Субсидии»
    # в орге ЭТОЙ субсидии (орг-роль не даёт власти в чужих оргах).
    from app.auth.permissions import has_org_key
    if not await has_org_key(current_user, db, db_subsidy.org_id, 'subsidies', subsidy_id=subsidy_id):
        raise HTTPException(
            status_code=403,
            detail="Нет права удалить субсидию: доступ к субсидиям не выдан для организации-грантополучателя",
        )

    # Pre-check FK references to avoid 500 ForeignKeyViolationError
    # Superadmin bypasses this check — DB will SET NULL automatically (b7e1 migration).
    if current_user.role != 'superadmin':
        from app.models.purchase import Purchase
        from app.models.contract import Contract

        blocking_purchases = (await db.execute(
            select(Purchase.id, Purchase.status).where(Purchase.subsidy_id == subsidy_id)
        )).all()
        blocking_contracts = (await db.execute(
            select(Contract.id).where(Contract.subsidy_id == subsidy_id)
        )).all()
    else:
        blocking_purchases = []
        blocking_contracts = []
    if blocking_purchases or blocking_contracts:
        parts = []
        if blocking_purchases:
            ids = ", ".join(f"#{p.id}" for p in blocking_purchases[:20])
            more = f" и ещё {len(blocking_purchases) - 20}" if len(blocking_purchases) > 20 else ""
            parts.append(f"{len(blocking_purchases)} закупок (ID: {ids}{more})")
        if blocking_contracts:
            ids = ", ".join(f"#{c.id}" for c in blocking_contracts[:20])
            more = f" и ещё {len(blocking_contracts) - 20}" if len(blocking_contracts) > 20 else ""
            parts.append(f"{len(blocking_contracts)} договоров (ID: {ids}{more})")
        raise HTTPException(
            status_code=409,
            detail=(
                f"Нельзя удалить субсидию «{db_subsidy.name}»: связано "
                f"{' и '.join(parts)}. Часть закупок может быть скрыта фильтрами в списке "
                f"(например, разделённые закупки со статусом «split»). Сначала удалите или перепривяжите их."
            ),
        )

    # Bulk-delete dependents via SQL (not ORM cascade): feo_planned_items.feo_category_id
    # is NOT NULL, and ORM-level cascade would try to NULL it on parent delete -> IntegrityError.
    # Direct DELETE relies on the DB and avoids the autoflush nullify path entirely.
    from app.models.feo_planned_item import FeoPlannedItem
    try:
        cat_ids_subq = select(FeoCategory.id).where(FeoCategory.subsidy_id == subsidy_id)
        await db.execute(
            delete(FeoPlannedItem).where(FeoPlannedItem.feo_category_id.in_(cat_ids_subq))
        )
        await db.execute(
            delete(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id)
        )
        await db.delete(db_subsidy)
        await db.commit()
    except IntegrityError as e:
        await db.rollback()
        logger.warning("delete_subsidy %s blocked by FK: %s", subsidy_id, e)
        raise HTTPException(
            status_code=409,
            detail=(
                "Не удалось удалить субсидию: на неё ссылаются другие записи "
                "(закупки, договоры или плановые позиции). Сначала удалите или "
                "перепривяжите связанные данные, затем повторите."
            ),
        )
    return {"message": "Субсидия удалена"}


# ── Per-subsidy contractor override endpoints ──

@router.get("/{subsidy_id}/contractor-override", response_model=SubsidyContractorOverrideOut)
async def get_contractor_override(
    subsidy_id: int,
    db: AsyncSession = Depends(get_db),
):
    """Get contractor detail overrides for a subsidy."""
    subsidy = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not subsidy or not subsidy.contractor_id:
        raise HTTPException(404, "Субсидия или контрагент не найден")

    override = (await db.execute(
        select(SubsidyContractorOverride).where(
            SubsidyContractorOverride.subsidy_id == subsidy_id,
            SubsidyContractorOverride.contractor_id == subsidy.contractor_id,
        )
    )).scalar_one_or_none()

    if not override:
        # Return base contractor data as initial override
        contractor = await db.get(Contractor, subsidy.contractor_id)
        if not contractor:
            raise HTTPException(404, "Контрагент не найден")
        return {
            "id": 0,
            "subsidy_id": subsidy_id,
            "contractor_id": subsidy.contractor_id,
            "org_type": contractor.org_type,
            "inn": contractor.inn,
            "kpp": contractor.kpp,
            "ogrn": contractor.ogrn,
            "signatory": contractor.signatory,
            "signatory_basis": contractor.signatory_basis,
            "address": contractor.address,
            "postal_address": contractor.postal_address,
            "bank_details": contractor.bank_details,
            "settlement_account": contractor.settlement_account,
            "bank_name": contractor.bank_name,
            "bik": contractor.bik,
            "correspondent_account": contractor.correspondent_account,
            "contact_person": contractor.contact_person,
            "phone": contractor.phone,
            "email": contractor.email,
        }
    return override


@router.put("/{subsidy_id}/contractor-override", response_model=SubsidyContractorOverrideOut)
async def upsert_contractor_override(
    subsidy_id: int,
    data: SubsidyContractorOverrideCreate,
    db: AsyncSession = Depends(get_db),
):
    """Create or update contractor detail overrides for a subsidy."""
    subsidy = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not subsidy or not subsidy.contractor_id:
        raise HTTPException(404, "Субсидия или контрагент не найден")

    override = (await db.execute(
        select(SubsidyContractorOverride).where(
            SubsidyContractorOverride.subsidy_id == subsidy_id,
            SubsidyContractorOverride.contractor_id == subsidy.contractor_id,
        )
    )).scalar_one_or_none()

    d = data.dict(exclude_unset=True) if override else data.dict()
    # Пересобираем signatory из структурированных частей (только ФИО)
    _fio_parts = (
        d.get('signatory_last_name') or (override.signatory_last_name if override else None),
        d.get('signatory_first_name') or (override.signatory_first_name if override else None),
        d.get('signatory_middle_name') or (override.signatory_middle_name if override else None),
    )
    if any(_fio_parts):
        d['signatory'] = compose_fio(*_fio_parts) or d.get('signatory')

    if override:
        for key, value in d.items():
            setattr(override, key, value)
    else:
        override = SubsidyContractorOverride(
            subsidy_id=subsidy_id,
            contractor_id=subsidy.contractor_id,
            **d
        )
        db.add(override)

    await db.commit()
    await db.refresh(override)
    return override


# ── Per-subsidy document templates ──────────────────────────────────────────

TEMPLATES_BASE = "/app/templates"
SUBSIDY_TEMPLATES_BASE = "/app/uploads/templates"
SUPPORTED_DOC_TYPES = {
    "contract":                 "Договор",
    "contract_tz":              "ТЗ (общий шаблон)",
    # tech_spec slot removed from SubsidiesView UI 2026-04-21 — the endpoint
    # still resolves to contract_tz.docx (see documents.py DOC_TYPES fallback)
    # for any client that requests /documents/tech_spec directly.
    "service_note_delivery":    "СЗ на выдачу",
    "service_note_payment":     "СЗ на оплату",
    # Phase 19.05: split ТЗ and dedicated SZ на закупку
    "service_note_procurement": "СЗ на закупку",
    # Phase 19.07: СЗ на аванс
    "service_note_advance":     "СЗ на аванс",
    "tech_spec_request":        "ТЗ для запроса цен",
    "tech_spec_contract":       "ТЗ для договора",
    "approval_sheet":           "Лист согласования",
    "order_purchase":           "Приказ на закупку",
    # Phase 28: typed contract forms per-subsidy
    # Форма «услуги» объединена в один шаблон; большая/малая отчётность теперь
    # выбирается отдельной методичкой (methodology_large/small), приклеиваемой
    # к любому из семи договоров, а не отдельным шаблоном текста договора.
    "contract_services":             "Договор услуг",
    "contract_services_food":       "Договор услуг (питание)",
    "contract_goods_single":        "Договор поставки (разовый)",
    "contract_gph_individual":      "Договор ГПХ с физ.лицом (без РИД)",
    "contract_gph_individual_rid":  "Договор ГПХ с физ.лицом (+РИД)",
    "contract_repair_vehicle":      "Договор на ремонт ТС",
    "contract_repair_framework":    "Рамочный договор на ремонт ТС",
    # Методические рекомендации — приклеиваются к договору по Purchase.methodology
    "methodology_large":            "Методические рекомендации — большие",
    "methodology_small":            "Методические рекомендации — малые",
    # Fabrikant ЭТП package
    "fabrikant_instruction":        "Инструкция по заполнению заявки (Фабрикант)",
    "fabrikant_application_form":   "Форма заявки (Фабрикант)",
    "fabrikant_documentation":      "Документация к закупке (Фабрикант)",
    "fabrikant_contract_project":   "Проект договора (Фабрикант)",
}


def _check_template_render_ok(path: str) -> bool:
    """Try a trial render with empty context; return True if template is valid."""
    if not os.path.exists(path):
        return False
    try:
        if DocxTemplate is None:
            return True  # docxtpl not installed — assume ok
        tpl = DocxTemplate(path)
        tpl.render({})
        return True
    except Exception:
        return False


@router.get("/{subsidy_id}/templates")
async def list_subsidy_templates(
    subsidy_id: int,
    current_user=Depends(require_tab('subsidies')),
):
    """List which doc types have a subsidy-specific template override."""
    result = []
    subsidy_dir = os.path.join(SUBSIDY_TEMPLATES_BASE, "subsidies", str(subsidy_id))
    for doc_type, label in SUPPORTED_DOC_TYPES.items():
        path = os.path.join(subsidy_dir, f"{doc_type}.docx")
        global_path = os.path.join(TEMPLATES_BASE, f"{doc_type}.docx")
        has_custom = os.path.exists(path)
        render_ok: bool | None = None
        if has_custom:
            render_ok = _check_template_render_ok(path)
        result.append({
            "doc_type": doc_type,
            "label": label,
            "has_custom": has_custom,
            "has_global": os.path.exists(global_path),
            "render_ok": render_ok,
        })
    return result


def _normalize_docx_template(path: str) -> dict:
    """Strip Word-internal markers that split jinja-placeholders across runs.

    Word inserts <w:proofErr/>, bookmarks, comments, lastRenderedPageBreak
    between runs inside `{{ var }}` / `{% ... %}` blocks during user editing.
    For text placeholders docxtpl can recover, but for InlineImage the tag
    MUST be inside a single contiguous run — otherwise the drawing element
    is silently dropped.

    Touches ONLY xml files where jinja syntax may live (document/header/
    footer/footnotes/endnotes). Avoiding [Content_Types].xml and *.rels —
    rewriting those breaks the package and was the root cause of the
    Phase 26-VV revert.

    Writes the cleaned bytes through a tempfile + atomic replace so the
    on-disk file stays valid even if a write is interrupted.
    """
    import zipfile as _zipfile
    import tempfile as _tempfile

    JINJA_BEARING = ('word/document.xml',)
    JINJA_BEARING_PREFIXES = ('word/header', 'word/footer')
    JINJA_BEARING_SUFFIXES = ('word/footnotes.xml', 'word/endnotes.xml')

    def _is_jinja_bearing(name: str) -> bool:
        if name in JINJA_BEARING or name in JINJA_BEARING_SUFFIXES:
            return True
        return any(name.startswith(p) and name.endswith('.xml') for p in JINJA_BEARING_PREFIXES)

    stats = {"proofErr": 0, "bookmark": 0, "comment": 0, "pageBreak": 0}
    tmp_fd, tmp_path = _tempfile.mkstemp(suffix=".docx", dir=os.path.dirname(path) or None)
    os.close(tmp_fd)
    try:
        with _zipfile.ZipFile(path, 'r') as zin:
            with _zipfile.ZipFile(tmp_path, 'w', _zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if _is_jinja_bearing(item.filename):
                        xml = data.decode('utf-8', errors='replace')
                        before_proof = xml.count('<w:proofErr')
                        xml = re.sub(r'<w:proofErr\s[^/]*?/>', '', xml)
                        stats["proofErr"] += before_proof - xml.count('<w:proofErr')

                        before_bm = xml.count('<w:bookmark')
                        xml = re.sub(r'<w:bookmarkStart\s[^/]*?/>', '', xml)
                        xml = re.sub(r'<w:bookmarkEnd\s[^/]*?/>', '', xml)
                        stats["bookmark"] += before_bm - xml.count('<w:bookmark')

                        before_cm = xml.count('<w:comment')
                        xml = re.sub(r'<w:commentRangeStart\s[^/]*?/>', '', xml)
                        xml = re.sub(r'<w:commentRangeEnd\s[^/]*?/>', '', xml)
                        xml = re.sub(r'<w:commentReference\s[^/]*?/>', '', xml)
                        stats["comment"] += before_cm - xml.count('<w:comment')

                        before_pb = xml.count('<w:lastRenderedPageBreak')
                        xml = re.sub(r'<w:lastRenderedPageBreak\s*/>', '', xml)
                        stats["pageBreak"] += before_pb - xml.count('<w:lastRenderedPageBreak')

                        data = xml.encode('utf-8')
                    zout.writestr(item, data)
        os.replace(tmp_path, path)
        logger.info("docx normalize %s: %s", path, stats)
    except Exception as e:
        logger.warning("docx normalize failed for %s: %s", path, e)
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass
    return stats


def _repair_docx_template(path: str) -> list[str]:
    """Fix common Jinja2 errors in uploaded .docx templates.

    Word often splits {{ variable }} across multiple XML runs, and users
    accidentally put Russian text like 'г.' inside Jinja tags.
    Returns list of fixes applied.
    """
    fixes = []
    try:
        from docx import Document
        doc = Document(path)

        def _fix_paragraph(para):
            if not para.runs:
                return False
            full = "".join(r.text for r in para.runs)
            if "{{" not in full and "{%" not in full:
                return False

            original = full
            # Fix: {{ var something_russian }} -> {{ var }} something_russian
            # Pattern: inside {{ }}, after a valid variable name, remove non-ASCII chars
            def _clean_tag(m):
                inner = m.group(1)
                # Split into variable name and trailing junk
                parts = inner.strip().split()
                if len(parts) <= 1:
                    return m.group(0)  # just {{ var }} — ok
                var_name = parts[0]
                # Check if first part looks like a variable (ascii, underscores, dots)
                if not re.match(r'^[a-zA-Z_][a-zA-Z0-9_.]*$', var_name):
                    return m.group(0)
                # Remaining parts — check if they contain non-ASCII (Russian text)
                rest = " ".join(parts[1:])
                if any(ord(c) > 127 for c in rest):
                    return "{{ " + var_name + " }} " + rest
                return m.group(0)

            fixed = re.sub(r'\{\{([^}]+)\}\}', _clean_tag, full)

            # Fix: {{полностью_русский текст}} — снимаем маркеры, оставляем текст
            def _strip_if_all_non_ascii(m):
                inner = m.group(1).strip()
                if not inner:
                    return m.group(0)
                # Если ни одного ASCII-идентификаторного символа — это ремарка, не переменная
                if not re.search(r'[a-zA-Z_]', inner):
                    return inner
                return m.group(0)
            fixed = re.sub(r'\{\{([^}]+)\}\}', _strip_if_all_non_ascii, fixed)

            # Fix: {% русский текст %} — снимаем маркеры если не Jinja keyword
            _JINJA_KEYWORDS = {
                'if', 'else', 'elif', 'endif', 'for', 'endfor', 'set',
                'with', 'endwith', 'tr', 'block', 'endblock', 'macro',
                'endmacro', 'include', 'extends', 'import', 'from',
                'do', 'raw', 'endraw', 'call', 'endcall', 'filter',
                'endfilter', 'autoescape', 'endautoescape',
            }
            def _strip_pct_if_no_keyword(m):
                inner = m.group(1).strip().lstrip('-').rstrip('-').strip()
                if not inner:
                    return m.group(0)
                first = inner.split()[0].lower() if inner.split() else ''
                # Если первое слово — не Jinja keyword И не похоже на ASCII identifier — ремарка
                if first not in _JINJA_KEYWORDS and not re.match(r'^[a-zA-Z_]', first):
                    return inner
                return m.group(0)
            fixed = re.sub(r'\{%([^%]+)%\}', _strip_pct_if_no_keyword, fixed)

            # Fix doubled text like "г. г." that can result from prior fixes
            fixed = re.sub(r'(\b\w+\.)\s+\1', r'\1', fixed)

            if fixed == original:
                return False
            # Write back: all text into first run, clear rest
            para.runs[0].text = fixed
            for r in para.runs[1:]:
                r.text = ""
            return True

        count = 0
        for para in doc.paragraphs:
            if _fix_paragraph(para):
                count += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if _fix_paragraph(para):
                            count += 1

        if count:
            doc.save(path)
            fixes.append(f"Исправлено {count} Jinja-тегов")

        # Validate the template renders
        from docxtpl import DocxTemplate
        tpl = DocxTemplate(path)
        tpl.get_undeclared_template_variables()

    except Exception as e:
        logger.warning("Template repair/validation warning for %s: %s", path, e)
        fixes.append(f"Предупреждение: {e}")

    return fixes


@router.put("/{subsidy_id}/templates/{doc_type}")
async def upload_subsidy_template(
    subsidy_id: int,
    doc_type: str,
    file: UploadFile = File(...),
    current_user=Depends(require_tab('subsidies')),
):
    """Upload a .docx template override for a specific subsidy and doc type."""
    if doc_type not in SUPPORTED_DOC_TYPES:
        raise HTTPException(400, f"Неизвестный тип документа: {doc_type}")
    if not file.filename.endswith(".docx"):
        raise HTTPException(400, "Допускаются только .docx файлы")

    subsidy_dir = os.path.join(SUBSIDY_TEMPLATES_BASE, "subsidies", str(subsidy_id))
    os.makedirs(subsidy_dir, exist_ok=True)
    dest = os.path.join(subsidy_dir, f"{doc_type}.docx")

    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    norm_stats = _normalize_docx_template(dest)
    repairs = _repair_docx_template(dest)

    # Trial render: validate docxtpl syntax before confirming upload
    if DocxTemplate is not None:
        try:
            tpl = DocxTemplate(dest)
            tpl.render({})
        except Exception as e:
            try:
                os.remove(dest)
            except OSError:
                pass
            raise HTTPException(
                status_code=400,
                detail=f"Шаблон некорректен: {type(e).__name__}: {str(e)[:200]}"
            )

    return {
        "ok": True,
        "doc_type": doc_type,
        "label": SUPPORTED_DOC_TYPES[doc_type],
        "repairs": repairs,
        "normalized": norm_stats,
    }


@router.get("/{subsidy_id}/templates/{doc_type}/download")
async def download_subsidy_template(
    subsidy_id: int,
    doc_type: str,
    current_user=Depends(require_tab('subsidies')),
):
    """Download the current subsidy-specific template (or global if no override)."""
    if doc_type not in SUPPORTED_DOC_TYPES:
        raise HTTPException(400, f"Неизвестный тип документа: {doc_type}")

    subsidy_path = os.path.join(SUBSIDY_TEMPLATES_BASE, "subsidies", str(subsidy_id), f"{doc_type}.docx")
    global_path = os.path.join(TEMPLATES_BASE, f"{doc_type}.docx")

    if os.path.exists(subsidy_path):
        path = subsidy_path
    elif os.path.exists(global_path):
        path = global_path
    else:
        raise HTTPException(404, "Шаблон не найден")

    response = FileResponse(
        path,
        filename=f"{doc_type}_{subsidy_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    return response


@router.delete("/{subsidy_id}/templates/{doc_type}")
async def delete_subsidy_template(
    subsidy_id: int,
    doc_type: str,
    current_user=Depends(require_tab('subsidies')),
):
    """Delete subsidy-specific template override (falls back to global)."""
    if doc_type not in SUPPORTED_DOC_TYPES:
        raise HTTPException(400, f"Неизвестный тип документа: {doc_type}")

    path = os.path.join(SUBSIDY_TEMPLATES_BASE, "subsidies", str(subsidy_id), f"{doc_type}.docx")
    if not os.path.exists(path):
        raise HTTPException(404, "Индивидуальный шаблон не найден")

    os.remove(path)
    return {"ok": True}


# ── Global template upload (admin) ──────────────────────────────────────────

@router.put("/global-templates/{doc_type}")
async def upload_global_template(
    doc_type: str,
    file: UploadFile = File(...),
    current_user=Depends(require_action('subsidy.edit')),
):
    """Upload a global .docx template (superadmin/account_owner only)."""
    if doc_type not in SUPPORTED_DOC_TYPES:
        raise HTTPException(400, f"Неизвестный тип документа: {doc_type}")
    if not file.filename.endswith(".docx"):
        raise HTTPException(400, "Допускаются только .docx файлы")

    dest = os.path.join(TEMPLATES_BASE, f"{doc_type}.docx")
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    norm_stats = _normalize_docx_template(dest)
    repairs = _repair_docx_template(dest)

    # Trial render: validate docxtpl syntax before confirming upload
    if DocxTemplate is not None:
        try:
            tpl = DocxTemplate(dest)
            tpl.render({})
        except Exception as e:
            try:
                os.remove(dest)
            except OSError:
                pass
            raise HTTPException(
                status_code=400,
                detail=f"Шаблон некорректен: {type(e).__name__}: {str(e)[:200]}"
            )

    return {"ok": True, "doc_type": doc_type, "repairs": repairs, "normalized": norm_stats}


@router.post("/templates/normalize-all")
async def normalize_all_existing_templates(
    current_user=Depends(require_action('subsidy.edit')),
):
    """One-shot migration: normalize every existing .docx template in place.

    Walks SUBSIDY_TEMPLATES_BASE and TEMPLATES_BASE, strips Word-internal
    markers from each file. Idempotent — safe to call multiple times.
    """
    processed = []
    errors = []

    def _walk_and_normalize(root: str):
        if not os.path.isdir(root):
            return
        for dirpath, _dirs, files in os.walk(root):
            for fn in files:
                if not fn.endswith(".docx"):
                    continue
                fpath = os.path.join(dirpath, fn)
                try:
                    stats = _normalize_docx_template(fpath)
                    processed.append({"path": fpath, "stripped": stats})
                    try:
                        repairs = _repair_docx_template(fpath)
                        if repairs:
                            processed[-1]["repairs"] = repairs
                    except Exception as re_err:
                        errors.append({"path": fpath, "error": f"repair: {re_err}"})
                except Exception as e:
                    errors.append({"path": fpath, "error": str(e)})

    _walk_and_normalize(SUBSIDY_TEMPLATES_BASE)
    _walk_and_normalize(TEMPLATES_BASE)

    return {"ok": True, "processed": len(processed), "errors": len(errors), "details": processed, "error_details": errors}


@router.get("/{subsidy_id}/history")
async def get_budget_history(
    subsidy_id: int,
    offset: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_tab('subsidies')),
):
    from app.models.budget_history import BudgetHistory as BudgetHistoryModel
    from sqlalchemy import func as safunc

    base_q = (
        select(BudgetHistoryModel)
        .where(BudgetHistoryModel.subsidy_id == subsidy_id)
        .order_by(BudgetHistoryModel.changed_at.desc())
    )

    total = (
        await db.execute(select(safunc.count()).select_from(base_q.subquery()))
    ).scalar() or 0

    rows = (await db.execute(base_q.offset(offset).limit(limit))).scalars().all()

    return {
        "total": total,
        "items": [
            {
                "id": r.id,
                "entity_type": r.entity_type,
                "purchase_id": r.purchase_id,
                "old_value": float(r.old_value) if r.old_value is not None else None,
                "new_value": float(r.new_value) if r.new_value is not None else None,
                "changed_by_name": r.changed_by_name,
                "reason": r.reason,
                "changed_at": r.changed_at.isoformat() if r.changed_at else None,
            }
            for r in rows
        ],
    }


# ── Plan-Graph Version endpoints (Phase 12-03) ────────────────────────────────

def _snapshot_effective_total(snapshot: dict) -> float:
    """Compute total_effective from snapshot: fact replaces plan where fact > 0."""
    if snapshot.get("total_effective") is not None:
        return float(snapshot["total_effective"])
    items = snapshot.get("items") or []
    if items:
        return sum(
            (float(it.get("used_actual_amount") or 0) if float(it.get("used_actual_amount") or 0) > 0
             else float(it.get("planned_amount") or 0))
            for it in items
        )
    def _leaf_effective(nodes):
        total = 0.0
        for n in nodes:
            children = n.get("children") or []
            if children:
                total += _leaf_effective(children)
            else:
                ua = float(n.get("used_actual_amount") or 0)
                total += ua if ua > 0 else float(n.get("budget") or 0)
        return total
    return _leaf_effective(snapshot.get("tree") or [])


@router.get("/{subsidy_id}/plan-graph/versions")
async def list_plan_graph_versions(
    subsidy_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Return list of plan-graph versions for a subsidy (summary, no full snapshot)."""
    from app.models.plan_graph_version import PlanGraphVersion as _PGV

    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")

    org_ids = get_org_filter(current_user)
    if org_ids is not None and sub.org_id not in org_ids:
        raise HTTPException(403, "Нет доступа к субсидии")

    vers = (await db.execute(
        select(_PGV)
        .where(_PGV.subsidy_id == subsidy_id)
        .order_by(_PGV.version_number.desc())
        .limit(100)
    )).scalars().all()

    out = []
    for v in vers:
        snap = v.snapshot or {}
        out.append({
            "id": v.id,
            "version_number": v.version_number,
            "created_at": v.created_at.isoformat() if v.created_at else None,
            "created_by_name": v.created_by_name,
            "note": v.note,
            "effective_date": v.effective_date.isoformat() if v.effective_date else None,
            "total_planned": snap.get("total_plan_combined", snap.get("total_planned", 0)),
            "total_used": snap.get("total_used", 0),
            "total_effective": _snapshot_effective_total(snap),
            "item_count": len(snap.get("items", [])),
        })
    return out


@router.get("/{subsidy_id}/plan-graph/versions/{version_id:int}")
async def get_plan_graph_version(
    subsidy_id: int,
    version_id: int,
    with_reconciliation: bool = Query(False, description="Добавить факт по текущим закупкам"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Return full snapshot for a specific plan-graph version.

    Optional ?with_reconciliation=true adds current actual PurchaseItem totals
    matched to snapshot tree nodes via composite-key matcher.
    """
    from app.models.plan_graph_version import PlanGraphVersion as _PGV

    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")

    org_ids = get_org_filter(current_user)
    if org_ids is not None and sub.org_id not in org_ids:
        raise HTTPException(403, "Нет доступа к субсидии")

    ver = (await db.execute(
        select(_PGV).where(
            _PGV.id == version_id,
            _PGV.subsidy_id == subsidy_id,
        )
    )).scalar_one_or_none()

    if not ver:
        raise HTTPException(404, "Версия плана закупок не найдена")

    resp = {
        "id": ver.id,
        "subsidy_id": ver.subsidy_id,
        "version_number": ver.version_number,
        "created_at": ver.created_at.isoformat() if ver.created_at else None,
        "created_by_id": ver.created_by_id,
        "created_by_name": ver.created_by_name,
        "note": ver.note,
        "effective_date": ver.effective_date.isoformat() if ver.effective_date else None,
        "snapshot": ver.snapshot,
    }

    if with_reconciliation:
        snap = ver.snapshot or {}
        snap_tree = snap.get("tree")
        reconciliation: dict = {}

        if snap_tree:
            from app.models.feo_planned_item import FeoPlannedItem as _FPI
            from app.models.purchase_item import PurchaseItem as _PI

            # Load current live FeoCategory tree for this subsidy
            live_cats = (await db.execute(
                select(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id).order_by(FeoCategory.id)
            )).scalars().all()

            def _build_live_tree_simple(cats, parent_id=None):
                nodes = []
                for c in cats:
                    if c.parent_id == parent_id:
                        nodes.append({
                            "id": c.id,
                            "name": c.name,
                            "level": c.level,
                            "code": c.code,
                            "children": _build_live_tree_simple(cats, parent_id=c.id),
                        })
                return nodes

            live_tree = _build_live_tree_simple(list(live_cats))
            live_cats_by_id = {c.id: c for c in live_cats}

            # Match snapshot tree → live tree
            match_result = _match_feo_nodes(snap_tree, live_tree)
            # match_result["matches"] = [(snap_node, live_node, match_type), ...]

            # For each live category, get PurchaseItem totals via FeoPlannedItem OR feo_category_id
            live_cat_ids = [c.id for c in live_cats]
            # FCAT-B3: SUM via feo_planned_item_id (legacy) + feo_category_id (new)
            cat_actual_map: dict[int, float] = {}
            if live_cat_ids:
                # Source 1: via FeoPlannedItem join (legacy ФАДМ_2026)
                actual_rows_via_fpi = (await db.execute(
                    select(
                        _FPI.feo_category_id,
                        func.coalesce(func.sum(_PI.total_price), 0).label("actual"),
                    )
                    .join(_PI, _PI.feo_planned_item_id == _FPI.id)
                    .where(_FPI.feo_category_id.in_(live_cat_ids))
                    .group_by(_FPI.feo_category_id)
                )).all()
                for r in actual_rows_via_fpi:
                    cat_actual_map[r.feo_category_id] = float(r.actual)

                # Source 2: via PurchaseItem.feo_category_id (FCAT-B1 new column)
                actual_rows_via_cat = (await db.execute(
                    select(
                        _PI.feo_category_id,
                        func.coalesce(func.sum(_PI.total_price), 0).label("actual"),
                    )
                    .where(_PI.feo_category_id.in_(live_cat_ids))
                    .group_by(_PI.feo_category_id)
                )).all()
                for r in actual_rows_via_cat:
                    cat_id = r.feo_category_id
                    cat_actual_map[cat_id] = cat_actual_map.get(cat_id, 0.0) + float(r.actual)

            def _subtree_actual(cat_id: int) -> float:
                """Recursively sum actual from cat and all its descendants."""
                total = cat_actual_map.get(cat_id, 0.0)
                cat = live_cats_by_id.get(cat_id)
                if cat:
                    for child in live_cats:
                        if child.parent_id == cat_id:
                            total += _subtree_actual(child.id)
                return total

            # Build reconciliation map keyed by snapshot node id
            for snap_node, live_node, match_type in match_result["matches"]:
                snap_id = snap_node.get("id")
                if snap_id is None:
                    continue
                live_id = live_node.get("id")
                budget_snap = float(snap_node.get("budget") or 0)
                actual = _subtree_actual(live_id) if live_id else 0.0
                reconciliation[str(snap_id)] = {
                    "matched_current_id": live_id,
                    "actual_used": round(actual, 2),
                    "actual_residual": round(budget_snap - actual, 2),
                    "match_type": match_type,
                }

            # Unmatched snapshot nodes → null
            for snap_node in match_result["only_a"]:
                snap_id = snap_node.get("id")
                if snap_id is not None:
                    reconciliation[str(snap_id)] = {
                        "matched_current_id": None,
                        "actual_used": 0.0,
                        "actual_residual": float(snap_node.get("budget") or 0),
                        "match_type": None,
                    }

        resp["reconciliation"] = reconciliation

    return resp


class PlanGraphVersionCreate(BaseModel):
    note: Optional[str] = None
    effective_date: Optional[date] = None


@router.post("/{subsidy_id}/plan-graph/versions")
async def create_plan_graph_version_manual(
    subsidy_id: int,
    body: PlanGraphVersionCreate = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(*ADMIN_ROLES)),
):
    """Manually publish a plan-graph version (admin only)."""
    from app.routers.purchases import _create_plan_graph_version
    body = body or PlanGraphVersionCreate()
    await _create_plan_graph_version(
        subsidy_id,
        db,
        current_user,
        note=body.note or "Ручная публикация",
        effective_date=body.effective_date,
    )
    await db.commit()
    return {"ok": True, "message": "Версия плана закупок сохранена"}


# ── Plan-Graph Export (Phase 12-04) ───────────────────────────────────────────

TEMPLATE_DIR = "media/plan_graph_templates"


def _render_plan_graph_workbook(tree: list, items: list, meta: dict):
    """
    Common renderer for plan-graph Excel workbook.

    Args:
        tree: recursive list of FeoCategory nodes
              [{id, name, level, code, appendix, budget, planned_amount,
                planned_quantity, unit, children:[...]}]
              If empty — fallback to flat `items` (v1 backward-compat).
        items: flat list of snapshot items (v1 backward-compat)
               [{name, planned_amount, used_amount, residual}]
        meta: dict with keys:
              subsidy_name, subsidy_year (optional), effective_date (str|None),
              version_number (int|None), note (str|None), generated_at (str|None)

    Returns:
        openpyxl.Workbook
    """
    if openpyxl is None:
        raise RuntimeError("openpyxl не установлен")

    HEADER_FILL  = PatternFill("solid", fgColor="1E3A5F")
    HEADER_FONT  = Font(color="FFFFFF", bold=True, size=9)
    L1_FILL      = PatternFill("solid", fgColor="DBEAFE")
    L1_FONT      = Font(bold=True, size=9)
    L2_FILL      = PatternFill("solid", fgColor="F0F9FF")
    L2_FONT      = Font(bold=True, size=9, color="0C4A6E")
    L3_FILL      = PatternFill("solid", fgColor="F0FDF4")
    L3_FONT      = Font(size=9, color="166534")
    ITEM_FONT    = Font(size=9)
    RED_FONT     = Font(size=9, color="EF4444", bold=True)
    META_FONT    = Font(size=9, italic=True, color="374151")
    CENTER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
    LEFT_ALIGN   = Alignment(horizontal="left", vertical="center", wrap_text=True)
    RIGHT_ALIGN  = Alignment(horizontal="right", vertical="center")
    THIN_BORDER  = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    HEADERS = [
        "№", "Направление расходов", "Тип расходов", "Наименование",
        "Ед.", "Кол-во план", "Плановая сумма, ₽",
        "Фактическая сумма, ₽", "Остаток, ₽",
        "% исполнения", "Исполнитель", "Статус",
    ]
    COL_WIDTHS = [5, 30, 25, 40, 8, 10, 18, 18, 18, 12, 30, 15]
    n_cols = len(HEADERS)
    last_col_letter = chr(64 + n_cols)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "План закупок"

    # ── Title row ─────────────────────────────────────────────────────────────
    subsidy_name = meta.get("subsidy_name", "")
    subsidy_year = meta.get("subsidy_year", "")
    title_text = f"ПЛАН-ГРАФИК — {subsidy_name}"
    if subsidy_year:
        title_text += f" ({subsidy_year})"
    ws.append([title_text] + [""] * (n_cols - 1))
    title_cell = ws.cell(row=1, column=1)
    title_cell.font = Font(bold=True, size=12, color="1E3A5F")
    ws.merge_cells(f"A1:{last_col_letter}1")
    title_cell.alignment = CENTER_ALIGN
    ws.row_dimensions[1].height = 28

    # ── Meta info rows (version, effective_date, note, generated_at) ──────────
    meta_start_row = 2
    meta_rows = []
    if meta.get("version_number") is not None:
        meta_rows.append(f"Версия: {meta['version_number']}")
    if meta.get("effective_date"):
        meta_rows.append(f"Дата редакции: {meta['effective_date']}")
    if meta.get("note"):
        meta_rows.append(f"Примечание: {meta['note']}")
    if meta.get("generated_at"):
        meta_rows.append(f"Сформировано: {meta['generated_at']}")

    for i, mtext in enumerate(meta_rows):
        r = meta_start_row + i
        ws.append([mtext] + [""] * (n_cols - 1))
        cell = ws.cell(row=r, column=1)
        cell.font = META_FONT
        cell.alignment = LEFT_ALIGN
        ws.merge_cells(f"A{r}:{last_col_letter}{r}")
        ws.row_dimensions[r].height = 16

    # ── Column headers ────────────────────────────────────────────────────────
    header_row = meta_start_row + len(meta_rows)
    ws.append(HEADERS)
    for col_idx, (header, width) in enumerate(zip(HEADERS, COL_WIDTHS), 1):
        cell = ws.cell(row=header_row, column=col_idx)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER
        ws.column_dimensions[cell.column_letter].width = width
    ws.row_dimensions[header_row].height = 32
    ws.freeze_panes = f"A{header_row + 1}"

    row_num = header_row + 1
    seq = 0

    def _write_row(values, fill, font, height=20):
        nonlocal row_num
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_num, column=col_idx, value=val)
            cell.fill = fill
            cell.font = font
            cell.border = THIN_BORDER
            cell.alignment = RIGHT_ALIGN if col_idx >= 6 else LEFT_ALIGN
        ws.row_dimensions[row_num].height = height
        row_num += 1

    if tree:
        # ── Tree mode (schema_version=2) ──────────────────────────────────────
        def _traverse_node(node, direction_name="", type_name=""):
            nonlocal seq
            level = node.get("level", 1)
            name = node.get("name", "")
            code = node.get("code") or ""
            budget = node.get("budget")

            if level == 1:
                direction_name = name
                _write_row(
                    [code, name, "", "", "", "", budget or "", "", "", "", "", ""],
                    L1_FILL, L1_FONT, height=22,
                )
            elif level == 2:
                type_name = name
                _write_row(
                    ["", direction_name, name, "", "", "", budget or "", "", "", "", "", ""],
                    L2_FILL, L2_FONT,
                )
            elif level == 3:
                _write_row(
                    ["", direction_name, type_name, name, "", "", budget or "", "", "", "", "", ""],
                    L3_FILL, L3_FONT,
                )

            for child in node.get("children", []):
                _traverse_node(child, direction_name, type_name)

        for root in tree:
            _traverse_node(root)
    else:
        # ── Flat mode (v1 backward-compat) ─────────────────────────────────────
        for item in items:
            seq += 1
            planned = float(item.get("planned_amount") or 0)
            used = float(item.get("used_amount") or 0)
            residual = float(item.get("residual") or (planned - used))
            pct = round(used / planned * 100) if planned > 0 else 0
            status = "Выполнено" if pct >= 100 else ("В работе" if used > 0 else "Не начато")
            font = RED_FONT if used > planned else ITEM_FONT
            _write_row(
                [
                    seq, "", "", item.get("name", ""),
                    "", "",
                    round(planned, 2), round(used, 2), round(residual, 2),
                    f"{pct}%", "", status,
                ],
                PatternFill(), font,
            )

    return wb


@router.get("/{subsidy_id}/plan-graph/export")
async def export_plan_graph_excel(
    subsidy_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Export plan-graph as Excel file with full FEO hierarchy (live data)."""
    if openpyxl is None:
        raise HTTPException(500, "openpyxl не установлен")

    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")

    org_ids = get_org_filter(current_user)
    if org_ids is not None and sub.org_id not in org_ids:
        raise HTTPException(403, "Нет доступа")

    from app.models.feo_planned_item import FeoPlannedItem as _FPI
    from app.models.purchase_item import PurchaseItem as _PI
    from app.models.purchase import Purchase as _P
    from app.models.contractor import Contractor as _C

    cats = (await db.execute(
        select(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id)
        .order_by(FeoCategory.level, FeoCategory.id)
    )).scalars().all()

    feo_items = (await db.execute(
        select(_FPI)
        .join(FeoCategory, _FPI.feo_category_id == FeoCategory.id)
        .where(FeoCategory.subsidy_id == subsidy_id)
        .where(_FPI.is_active == True)
        .order_by(_FPI.id)
    )).scalars().all()

    item_ids = [i.id for i in feo_items]
    used_map: dict[int, float] = {}
    if item_ids:
        used_rows = (await db.execute(
            select(
                _PI.feo_planned_item_id,
                func.coalesce(func.sum(_PI.total_price), 0).label("used"),
            )
            .where(_PI.feo_planned_item_id.in_(item_ids))
            .group_by(_PI.feo_planned_item_id)
        )).all()
        used_map = {r.feo_planned_item_id: float(r.used) for r in used_rows}

    # Status breakdown: sum(PurchaseItem.total_price) per (feo_planned_item_id, Purchase.status)
    # статусы: plan_schedule, work_in_progress → «Запланировано»;
    #          contracted → «Договор»; ordered → «Заказано»;
    #          delivered → «Поставлено»; paid → «Оплачено»
    status_sums_map: dict[int, dict[str, float]] = {}
    if item_ids:
        st_rows = (await db.execute(
            select(
                _PI.feo_planned_item_id,
                _P.status,
                func.coalesce(func.sum(_PI.total_price), 0).label("s"),
            )
            .join(_P, _PI.purchase_id == _P.id)
            .where(_PI.feo_planned_item_id.in_(item_ids))
            .group_by(_PI.feo_planned_item_id, _P.status)
        )).all()
        for r in st_rows:
            status_sums_map.setdefault(r.feo_planned_item_id, {})[r.status] = float(r.s)

    def _st(item_id: int, *statuses: str) -> float:
        d = status_sums_map.get(item_id, {})
        return sum(d.get(s, 0.0) for s in statuses)

    contractor_map: dict[int, str] = {}
    if item_ids:
        c_rows = (await db.execute(
            select(_PI.feo_planned_item_id, _C.name.label("cname"))
            .join(_P, _PI.purchase_id == _P.id)
            .join(_C, _P.contractor_id == _C.id)
            .where(_PI.feo_planned_item_id.in_(item_ids))
            .distinct()
        )).all()
        for r in c_rows:
            if r.feo_planned_item_id not in contractor_map:
                contractor_map[r.feo_planned_item_id] = r.cname

    # Фактически закупленные позиции (PurchaseItem) по каждой плановой статье —
    # чтобы в выгрузке было видно ЧТО куплено и ПО КАКОЙ ЦЕНЕ, а не только сумма.
    _STATUS_HUMAN = {
        "wishes": "Запланировано", "plan_schedule": "Запланировано",
        "work_in_progress": "Запланировано",
        "contracted": "Договор", "ordered": "Заказано",
        "delivered": "Поставлено", "paid": "Оплачено",
    }
    def _act_number(acc_number, acc_docs) -> str:
        if acc_number:
            return str(acc_number)
        if isinstance(acc_docs, list) and acc_docs:
            first = acc_docs[0]
            if isinstance(first, dict) and first.get("number"):
                return str(first["number"])
        return ""

    purchased_by_item: dict[int, list] = {}
    if item_ids:
        pi_rows = (await db.execute(
            select(
                _PI.feo_planned_item_id,
                _PI.item_name,
                _PI.unit,
                _PI.quantity,
                _PI.unit_price,
                _PI.total_price,
                _PI.contractor_name,
                _P.id.label("purchase_id"),
                _P.status,
                _P.acceptance_doc_number,
                _P.acceptance_docs,
            )
            .join(_P, _PI.purchase_id == _P.id)
            .where(_PI.feo_planned_item_id.in_(item_ids))
            .order_by(_PI.feo_planned_item_id, _PI.id)
        )).all()
        for r in pi_rows:
            purchased_by_item.setdefault(r.feo_planned_item_id, []).append({
                "name": r.item_name or "",
                "unit": r.unit or "",
                "qty": float(r.quantity or 0),
                "unit_price": float(r.unit_price or 0),
                "total": float(r.total_price or 0),
                "contractor": r.contractor_name or "",
                "raw_status": r.status or "",
                "status": _STATUS_HUMAN.get(r.status, r.status or ""),
                "purchase_id": r.purchase_id,
                "act_number": _act_number(r.acceptance_doc_number, r.acceptance_docs),
            })

    # FCAT: закупки часто привязаны к КАТЕГОРИИ (feo_category_id), а не к плановой
    # позиции. Агрегируем факт по категории — иначе колонки факта пустые.
    cat_ids = [c.id for c in cats]
    cat_status_map: dict[int, dict[str, float]] = {}
    cat_contractor_map: dict[int, str] = {}
    cat_monthly_map: dict[int, float] = {}  # факт по закупкам-регулярным (ежемес.) платежам
    purchased_by_cat: dict[int, list] = {}
    # Привязка к категории живёт на ТРЁХ уровнях: PurchaseItem.feo_category_id,
    # Purchase.feo_category_id (так работает вкладка «план vs факт») и через
    # плановую статью (feo_planned_item_id → её категория). Берём coalesce,
    # иначе часть факта выпадает из роллапа.
    _ecat = func.coalesce(_PI.feo_category_id, _P.feo_category_id, _FPI.feo_category_id)
    if cat_ids:
        cst_rows = (await db.execute(
            select(
                _ecat.label("ecat"),
                _P.status,
                func.coalesce(func.sum(_PI.total_price), 0).label("s"),
            )
            .join(_P, _PI.purchase_id == _P.id)
            .outerjoin(_FPI, _PI.feo_planned_item_id == _FPI.id)
            .where(_ecat.in_(cat_ids))
            .group_by(_ecat, _P.status)
        )).all()
        for r in cst_rows:
            cat_status_map.setdefault(r.ecat, {})[r.status] = float(r.s)

        mth_rows = (await db.execute(
            select(
                _ecat.label("ecat"),
                func.coalesce(func.sum(_PI.total_price), 0).label("s"),
            )
            .join(_P, _PI.purchase_id == _P.id)
            .outerjoin(_FPI, _PI.feo_planned_item_id == _FPI.id)
            .where(_ecat.in_(cat_ids))
            .where(_P.is_monthly_payment.is_(True))
            .group_by(_ecat)
        )).all()
        for r in mth_rows:
            cat_monthly_map[r.ecat] = float(r.s)

        cpi_rows = (await db.execute(
            select(
                _ecat.label("ecat"),
                _PI.feo_planned_item_id,
                _PI.item_name,
                _PI.unit,
                _PI.quantity,
                _PI.unit_price,
                _PI.total_price,
                func.coalesce(_PI.contractor_name, _C.name).label("contractor_name"),
                _P.id.label("purchase_id"),
                _P.status,
                _P.is_monthly_payment,
                _P.monthly_payment_count,
                _P.monthly_payment_amount,
                _P.acceptance_doc_number,
                _P.acceptance_docs,
            )
            .join(_P, _PI.purchase_id == _P.id)
            .outerjoin(_C, _P.contractor_id == _C.id)
            .outerjoin(_FPI, _PI.feo_planned_item_id == _FPI.id)
            .where(_ecat.in_(cat_ids))
            .order_by(_ecat, _PI.id)
        )).all()
        for r in cpi_rows:
            if r.contractor_name and r.ecat not in cat_contractor_map:
                cat_contractor_map[r.ecat] = r.contractor_name
            # Детализацию по категории показываем только для позиций БЕЗ привязки к
            # плановой статье — иначе задвоится со строками под плановой позицией.
            if r.feo_planned_item_id is None:
                purchased_by_cat.setdefault(r.ecat, []).append({
                    "name": r.item_name or "",
                    "unit": r.unit or "",
                    "qty": float(r.quantity or 0),
                    "unit_price": float(r.unit_price or 0),
                    "total": float(r.total_price or 0),
                    "contractor": r.contractor_name or "",
                    "raw_status": r.status or "",
                    "status": _STATUS_HUMAN.get(r.status, r.status or ""),
                    "purchase_id": r.purchase_id,
                    "act_number": _act_number(r.acceptance_doc_number, r.acceptance_docs),
                    "is_monthly": bool(r.is_monthly_payment),
                    "monthly_count": r.monthly_payment_count,
                    "monthly_amount": float(r.monthly_payment_amount or 0),
                })

    # ── Закупки БЕЗ позиций: факт живёт на самой закупке (item_name/суммы) ──
    # Иначе такие закупки полностью выпадают из выгрузки.
    itemless_extra: list[dict] = []      # для «Сводной»
    unlinked_purchases: list[dict] = []  # секция «без категории ФЭО»
    pl_rows = (await db.execute(
        select(
            _P.id.label("purchase_id"), _P.feo_category_id, _P.item_name,
            _P.subject, _P.status,
            _P.planned_quantity, _P.unit, _P.planned_unit_price,
            _P.planned_total_price, _P.final_total_amount,
            _P.acceptance_doc_number, _P.acceptance_docs,
            _P.is_monthly_payment, _P.monthly_payment_count, _P.monthly_payment_amount,
            _P.item_type, _P.is_likely_needed,
            _C.name.label("contractor_name"),
        )
        .outerjoin(_C, _P.contractor_id == _C.id)
        .where(or_(
            _P.feo_category_id.in_(cat_ids or [-1]),
            _P.subsidy_id == subsidy_id,
        ))
        .where(~select(_PI.id).where(_PI.purchase_id == _P.id).exists())
        .order_by(_P.id)
    )).all()
    _cat_id_set = set(cat_ids)
    for r in pl_rows:
        amount = float(r.final_total_amount or r.planned_total_price or 0)
        qty = float(r.planned_quantity or 0)
        d = {
            "name": (r.item_name or r.subject or f"Закупка №{r.purchase_id}"),
            "unit": r.unit or "",
            "qty": qty,
            "unit_price": float(r.planned_unit_price or 0) or (amount / qty if qty else amount),
            "total": amount,
            "contractor": r.contractor_name or "",
            "raw_status": r.status or "",
            "status": _STATUS_HUMAN.get(r.status, r.status or ""),
            "purchase_id": r.purchase_id,
            "act_number": _act_number(r.acceptance_doc_number, r.acceptance_docs),
            "has_act": bool(r.acceptance_doc_number) or bool(
                isinstance(r.acceptance_docs, list) and r.acceptance_docs),
            "is_monthly": bool(r.is_monthly_payment),
            "monthly_count": r.monthly_payment_count,
            "monthly_amount": float(r.monthly_payment_amount or 0),
            "item_type": (r.item_type or "").strip().lower(),
            "is_likely_needed": bool(r.is_likely_needed),
        }
        itemless_extra.append(d)
        if r.feo_category_id in _cat_id_set:
            purchased_by_cat.setdefault(r.feo_category_id, []).append(d)
            if amount:
                m = cat_status_map.setdefault(r.feo_category_id, {})
                st_key = r.status or ""
                m[st_key] = m.get(st_key, 0.0) + amount
                if r.is_monthly_payment:
                    cat_monthly_map[r.feo_category_id] = \
                        cat_monthly_map.get(r.feo_category_id, 0.0) + amount
        else:
            unlinked_purchases.append(d)

    # ── Позиции закупок, привязанных к субсидии ТОЛЬКО через subsidy_id ──────
    # (ни у закупки, ни у позиций нет категории/плановой статьи)
    _ecat_u = func.coalesce(_PI.feo_category_id, _P.feo_category_id, _FPI.feo_category_id)
    upi_rows = (await db.execute(
        select(
            _PI.item_name, _PI.unit, _PI.quantity, _PI.unit_price, _PI.total_price,
            func.coalesce(_PI.contractor_name, _C.name).label("contractor_name"),
            func.coalesce(_PI.item_type, _P.item_type).label("item_type"),
            _P.id.label("purchase_id"), _P.status,
            _P.is_monthly_payment, _P.monthly_payment_count, _P.monthly_payment_amount,
            _P.acceptance_doc_number, _P.acceptance_docs, _P.is_likely_needed,
        )
        .join(_P, _PI.purchase_id == _P.id)
        .outerjoin(_C, _P.contractor_id == _C.id)
        .outerjoin(_FPI, _PI.feo_planned_item_id == _FPI.id)
        .where(_P.subsidy_id == subsidy_id)
        .where(_ecat_u.is_(None))
        .order_by(_P.id, _PI.id)
    )).all()
    for r in upi_rows:
        unlinked_purchases.append({
            "name": r.item_name or "",
            "unit": r.unit or "",
            "qty": float(r.quantity or 0),
            "unit_price": float(r.unit_price or 0),
            "total": float(r.total_price or 0),
            "contractor": r.contractor_name or "",
            "raw_status": r.status or "",
            "status": _STATUS_HUMAN.get(r.status, r.status or ""),
            "purchase_id": r.purchase_id,
            "act_number": _act_number(r.acceptance_doc_number, r.acceptance_docs),
            "is_monthly": bool(r.is_monthly_payment),
            "monthly_count": r.monthly_payment_count,
            "monthly_amount": float(r.monthly_payment_amount or 0),
        })

    def _cst(cat_id: int, *statuses: str) -> float:
        d = cat_status_map.get(cat_id, {})
        return sum(d.get(s, 0.0) for s in statuses)

    items_by_cat: dict[int, list] = {}
    for item in feo_items:
        items_by_cat.setdefault(item.feo_category_id, []).append(item)

    # Build live tree for _render_plan_graph_workbook
    def _build_live_tree(all_cats, parent_id=None):
        nodes = []
        for c in all_cats:
            if c.parent_id == parent_id:
                cat_items = items_by_cat.get(c.id, [])
                children = _build_live_tree(all_cats, parent_id=c.id)
                node = {
                    "id": c.id,
                    "name": c.name,
                    "level": c.level,
                    "code": c.code,
                    "appendix": c.appendix,
                    "budget": float(c.budget) if c.budget is not None else None,
                    "planned_amount": float(c.planned_amount) if c.planned_amount is not None else None,
                    "planned_quantity": float(c.planned_quantity) if c.planned_quantity is not None else None,
                    "unit": c.unit,
                    "children": children,
                    # live items embedded for leaf rendering
                    "_live_items": cat_items,
                    "_used_map": used_map,
                    "_contractor_map": contractor_map,
                }
                nodes.append(node)
        return nodes

    live_tree = _build_live_tree(list(cats))

    # Use dedicated live-render path (keeps contractor/used columns populated)
    from datetime import datetime as _dt
    meta = {
        "subsidy_name": sub.name,
        "subsidy_year": str(sub.year) if sub.year else "",
        "generated_at": _dt.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
    }

    # Build workbook manually using existing logic (live tree with items)
    if openpyxl is None:
        raise HTTPException(500, "openpyxl не установлен")

    HEADER_FILL  = PatternFill("solid", fgColor="1E3A5F")
    HEADER_FONT  = Font(color="FFFFFF", bold=True, size=9)
    L1_FILL      = PatternFill("solid", fgColor="DBEAFE")
    L1_FONT      = Font(bold=True, size=9)
    L2_FILL      = PatternFill("solid", fgColor="F0F9FF")
    L2_FONT      = Font(bold=True, size=9, color="0C4A6E")
    L3_FILL      = PatternFill("solid", fgColor="F0FDF4")
    L3_FONT      = Font(size=9, color="166534")
    ITEM_FONT    = Font(size=9)
    RED_FONT     = Font(size=9, color="EF4444", bold=True)
    SUB_FONT     = Font(size=8, italic=True, color="6B7280")
    SUB_FILL     = PatternFill("solid", fgColor="FAFAFA")
    CENTER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
    LEFT_ALIGN   = Alignment(horizontal="left", vertical="center", wrap_text=True)
    RIGHT_ALIGN  = Alignment(horizontal="right", vertical="center")
    THIN_BORDER  = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    HEADERS = [
        "№", "Направление расходов", "Тип расходов", "Наименование",
        "Ед.", "Кол-во план", "Бюджет ФЭО, ₽", "Запланировано, ₽",
        "Договор, ₽", "Заказано, ₽", "Поставлено, ₽", "Оплачено, ₽",
        "Фактически (итого), ₽", "Остаток, ₽",
        "% исполнения", "Исполнитель", "Статус",
        "Ежемес. (регуляр.) платежи, ₽",
        "№ акта", "Документы",
    ]
    COL_WIDTHS = [5, 30, 25, 40, 8, 10, 18, 18, 18, 18, 18, 18, 20, 18, 12, 30, 15, 22, 16, 14]

    cats_by_parent: dict = {}
    for c in cats:
        cats_by_parent.setdefault(c.parent_id, []).append(c)

    # Роллап факта/бюджета вверх по дереву: направление (ур.1) и группа (ур.2)
    # показывают суммарные итоги по всем листовым категориям-потомкам.
    # Факт листа берём по feo_category_id (cat_status_map) — это включает все
    # закупки категории, поэтому суммирование по потомкам не задваивается.
    _cat_by_id = {c.id: c for c in cats}
    subtree_map: dict[int, dict] = {}

    def _compute_subtree(cat) -> dict:
        if cat.id in subtree_map:
            return subtree_map[cat.id]
        children = cats_by_parent.get(cat.id, [])
        status_agg: dict[str, float] = {}
        if not children:
            for k, v in cat_status_map.get(cat.id, {}).items():
                status_agg[k] = status_agg.get(k, 0.0) + v
            monthly = cat_monthly_map.get(cat.id, 0.0)
            if cat.budget is not None:
                budget = float(cat.budget)
            else:
                budget = sum(float(it.amount or 0) for it in items_by_cat.get(cat.id, []))
        else:
            budget = 0.0
            monthly = 0.0
            for ch in children:
                r = _compute_subtree(ch)
                for k, v in r["status"].items():
                    status_agg[k] = status_agg.get(k, 0.0) + v
                budget += r["budget"]
                monthly += r["monthly"]
        res = {"status": status_agg, "budget": budget, "monthly": monthly}
        subtree_map[cat.id] = res
        return res

    for c in cats:
        _compute_subtree(c)

    def _sub(cat_id: int, *statuses: str) -> float:
        d = subtree_map.get(cat_id, {}).get("status", {})
        return sum(d.get(s, 0.0) for s in statuses)

    base_url = str(request.base_url).rstrip("/")
    LINK_FONT = Font(size=8, italic=True, color="2563EB", underline="single")

    def _cascade(raw_status: str, total: float):
        """Каскад суммы поставки по стадиям: сумма падает во все столбцы до текущей
        стадии включительно; недостигнутые = 0. Столбцы: Заказано/Поставлено/Оплачено."""
        t = round(total or 0, 2)
        ordered = t if raw_status in ("ordered", "delivered", "paid") else 0
        delivered = t if raw_status in ("delivered", "paid") else 0
        paid = t if raw_status == "paid" else 0
        return ordered, delivered, paid

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "План закупок"

    # Заголовок-титул в строке 1 (не через insert_rows — иначе гиперлинки не
    # сдвигаются вместе с ячейками и «съезжают» на строку выше).
    title_cell = ws.cell(row=1, column=1, value=f"ПЛАН-ГРАФИК — {sub.name} ({sub.year})")
    title_cell.font = Font(bold=True, size=12, color="1E3A5F")
    ws.merge_cells(f"A1:{chr(64 + len(HEADERS))}1")
    title_cell.alignment = CENTER_ALIGN
    ws.row_dimensions[1].height = 28

    for col_idx, (header, width) in enumerate(zip(HEADERS, COL_WIDTHS), 1):
        cell = ws.cell(row=2, column=col_idx, value=header)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER
        ws.column_dimensions[cell.column_letter].width = width
    ws.row_dimensions[2].height = 32
    ws.freeze_panes = "A3"

    row_num = 3
    seq = 0

    def _write_row(values, fill, font, height=20):
        nonlocal row_num
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_num, column=col_idx, value=val)
            cell.fill = fill
            cell.font = font
            cell.border = THIN_BORDER
            cell.alignment = RIGHT_ALIGN if col_idx >= 6 else LEFT_ALIGN
        ws.row_dimensions[row_num].height = height
        row_num += 1

    def _set_doc_link(purchase_id):
        """Проставить гиперлинк «Открыть» в колонку 20 последней записанной строки —
        ведёт на страницу заказа, где под авторизацией доступны все документы."""
        if not purchase_id:
            return
        cell = ws.cell(row=row_num - 1, column=20)
        cell.value = "Открыть"
        cell.hyperlink = f"{base_url}/orders/{purchase_id}"
        cell.font = LINK_FONT

    E = ""  # пустая ячейка

    def _money_cols(cat_id: int, contractor: str) -> list:
        """Денежные колонки 7..18 из роллапа поддерева по статусам."""
        st = subtree_map.get(cat_id, {})
        budget = st.get("budget", 0.0)
        monthly = st.get("monthly", 0.0)
        plan = _sub(cat_id, "plan_schedule", "work_in_progress", "wishes")
        contract = _sub(cat_id, "contracted")
        ordered = _sub(cat_id, "ordered")
        delivered = _sub(cat_id, "delivered")
        paid = _sub(cat_id, "paid")
        used = plan + contract + ordered + delivered + paid
        resid = budget - used

        def _nz(v):
            return round(v, 2) if abs(v) > 0.005 else E

        return [
            round(budget, 2) if budget else E,
            _nz(plan), _nz(contract), _nz(ordered), _nz(delivered), _nz(paid),
            _nz(used),
            round(resid, 2) if (budget or abs(used) > 0.005) else E,
            (f"{round(used / budget * 100)}%" if budget > 0 else E),
            contractor, E,
            _nz(monthly),
        ]

    def _traverse(cat, direction_name="", type_name=""):
        nonlocal seq
        if cat.level == 1:
            direction_name = cat.name
            _write_row(
                [cat.code or "", cat.name, E, E, E, E] + _money_cols(cat.id, ""),
                L1_FILL, L1_FONT, height=22
            )
        elif cat.level == 2:
            type_name = cat.name
            _write_row(
                [E, direction_name, cat.name, E, E, E] + _money_cols(cat.id, ""),
                L2_FILL, L2_FONT
            )
        elif cat.level == 3:
            _write_row(
                [E, direction_name, type_name, cat.name, E, E]
                + _money_cols(cat.id, cat_contractor_map.get(cat.id, "")),
                L3_FILL, L3_FONT
            )
            # Под-строки: закупки, привязанные напрямую к категории (без плановой статьи).
            # По каждой фактической позиции — цена, кол-во, сумма, контрагент, № акта,
            # ссылка на документы + каскад Заказано/Поставлено/Оплачено.
            cat_purchases = purchased_by_cat.get(cat.id, [])
            cat_ord = cat_del = cat_paid = 0.0
            for pi in cat_purchases:
                price_note = f"    └ {pi['name']} — {round(pi['unit_price'], 2):,.2f} ₽/ед."
                status_txt = pi["status"]
                monthly_val = ""
                if pi.get("is_monthly"):
                    cnt = pi.get("monthly_count")
                    amt = pi.get("monthly_amount") or 0
                    status_txt = f"{status_txt} · ежемес." + (f" ×{cnt}" if cnt else "")
                    monthly_val = round(pi["total"], 2)
                    if amt:
                        price_note += f" ({round(amt, 2):,.2f} ₽/мес)"
                ord_a, del_a, paid_a = _cascade(pi["raw_status"], pi["total"])
                cat_ord += ord_a; cat_del += del_a; cat_paid += paid_a
                _write_row(
                    [
                        "", "", "", price_note,
                        pi["unit"], round(pi["qty"], 3),
                        "", "", "",
                        ord_a, del_a, paid_a,
                        round(pi["total"], 2), "",
                        "", pi["contractor"], status_txt,
                        monthly_val, pi["act_number"],
                    ],
                    SUB_FILL, SUB_FONT, height=16
                )
                _set_doc_link(pi["purchase_id"])
            if cat_purchases:
                _write_row(
                    [
                        "", "", "", "    Итого по факту:",
                        "", "",
                        "", "", "",
                        round(cat_ord, 2), round(cat_del, 2), round(cat_paid, 2),
                        "", "", "", "", "", "",
                    ],
                    SUB_FILL, Font(size=8, bold=True, color="374151"), height=16
                )
            for item in items_by_cat.get(cat.id, []):
                seq += 1
                feo_budget = float(item.amount or 0)
                plan_sum = _st(item.id, "plan_schedule", "work_in_progress", "wishes")
                contract_sum = _st(item.id, "contracted")
                ordered_sum = _st(item.id, "ordered")
                delivered_sum = _st(item.id, "delivered")
                paid_sum = _st(item.id, "paid")
                used = used_map.get(item.id, 0.0)
                residual = feo_budget - used
                pct = round(used / feo_budget * 100) if feo_budget > 0 else 0
                contractor = contractor_map.get(item.id, "")
                status = "Выполнено" if pct >= 100 else ("В работе" if used > 0 else "Не начато")
                font = RED_FONT if used > feo_budget else ITEM_FONT
                _write_row(
                    [
                        seq, direction_name, type_name, item.name,
                        item.unit or "", float(item.quantity or 0),
                        round(feo_budget, 2), round(plan_sum, 2),
                        round(contract_sum, 2), round(ordered_sum, 2),
                        round(delivered_sum, 2), round(paid_sum, 2),
                        round(used, 2), round(residual, 2),
                        f"{pct}%", contractor, status,
                    ],
                    PatternFill(), font
                )
                # Под-строки: реально закупленные позиции (что, по какой цене, № акта,
                # каскад статусов Заказано/Поставлено/Оплачено + ссылка на документы)
                for pi in purchased_by_item.get(item.id, []):
                    price_note = f"    └ {pi['name']} — {round(pi['unit_price'], 2):,.2f} ₽/ед."
                    ord_a, del_a, paid_a = _cascade(pi["raw_status"], pi["total"])
                    _write_row(
                        [
                            "", "", "", price_note,
                            pi["unit"], round(pi["qty"], 3),
                            "", "", "",
                            ord_a, del_a, paid_a,
                            round(pi["total"], 2), "",
                            "", pi["contractor"], pi["status"],
                            "", pi["act_number"],
                        ],
                        SUB_FILL, SUB_FONT, height=16
                    )
                    _set_doc_link(pi["purchase_id"])

        for child in cats_by_parent.get(cat.id, []):
            _traverse(child, direction_name, type_name)

    for root in cats_by_parent.get(None, []):
        _traverse(root)

    # ── Секция: закупки, привязанные к субсидии, но без категории ФЭО ────────
    if unlinked_purchases:
        _write_row(
            ["", "Закупки по субсидии без привязки к категории ФЭО", "", "", "", ""]
            + [""] * 12,
            L1_FILL, L1_FONT, height=22
        )
        u_ord = u_del = u_paid = 0.0
        for pi in unlinked_purchases:
            price_note = f"    └ {pi['name']} — {round(pi['unit_price'], 2):,.2f} ₽/ед."
            status_txt = pi["status"]
            monthly_val = ""
            if pi.get("is_monthly"):
                cnt = pi.get("monthly_count")
                amt = pi.get("monthly_amount") or 0
                status_txt = f"{status_txt} · ежемес." + (f" ×{cnt}" if cnt else "")
                monthly_val = round(pi["total"], 2)
                if amt:
                    price_note += f" ({round(amt, 2):,.2f} ₽/мес)"
            ord_a, del_a, paid_a = _cascade(pi["raw_status"], pi["total"])
            u_ord += ord_a; u_del += del_a; u_paid += paid_a
            _write_row(
                [
                    "", "", "", price_note,
                    pi["unit"], round(pi["qty"], 3),
                    "", "", "",
                    ord_a, del_a, paid_a,
                    round(pi["total"], 2), "",
                    "", pi["contractor"], status_txt,
                    monthly_val, pi["act_number"],
                ],
                SUB_FILL, SUB_FONT, height=16
            )
            _set_doc_link(pi["purchase_id"])
        _write_row(
            [
                "", "", "", "    Итого по факту:",
                "", "",
                "", "", "",
                round(u_ord, 2), round(u_del, 2), round(u_paid, 2),
                "", "", "", "", "", "",
            ],
            SUB_FILL, Font(size=8, bold=True, color="374151"), height=16
        )

    # ── Лист «Сводная»: деньги по Товары/Услуги × корзины обязательств ────────
    def _empty_buckets():
        return {"paid": 0.0, "delivered": 0.0, "accepted_unpaid": 0.0,
                "planned": 0.0, "monthly": 0.0, "likely": 0.0, "total": 0.0}

    summary = {"услуга": _empty_buckets(), "товар": _empty_buckets()}
    _planned_statuses = ("wishes", "plan_schedule", "work_in_progress", "contracted", "ordered")
    sum_rows = (await db.execute(
        select(
            func.coalesce(_PI.item_type, _P.item_type).label("item_type"),
            _PI.total_price,
            _P.status,
            _P.acceptance_doc_number,
            _P.acceptance_docs,
            _P.is_monthly_payment,
            _P.is_likely_needed,
        )
        .join(_P, _PI.purchase_id == _P.id)
        .where(or_(
            func.coalesce(_PI.feo_category_id, _P.feo_category_id).in_(cat_ids or [-1]),
            _PI.feo_planned_item_id.in_(item_ids or [-1]),
            _P.subsidy_id == subsidy_id,
        ))
    )).all()
    for r in sum_rows:
        key = "услуга" if (r.item_type or "").strip().lower() == "услуга" else "товар"
        b = summary[key]
        amt = float(r.total_price or 0)
        st = r.status or ""
        b["total"] += amt
        if st == "paid":
            b["paid"] += amt
        if st == "delivered":
            b["delivered"] += amt
        has_act = bool(r.acceptance_doc_number) or bool(
            isinstance(r.acceptance_docs, list) and r.acceptance_docs)
        if has_act and st != "paid":
            b["accepted_unpaid"] += amt
        if st in _planned_statuses:
            b["planned"] += amt
        if r.is_monthly_payment:
            b["monthly"] += amt
        if r.is_likely_needed:
            b["likely"] += amt

    # Закупки без позиций — их суммы живут на самой закупке
    for d in itemless_extra:
        key = "услуга" if d["item_type"] == "услуга" else "товар"
        b = summary[key]
        amt = d["total"]
        st = d["raw_status"]
        b["total"] += amt
        if st == "paid":
            b["paid"] += amt
        if st == "delivered":
            b["delivered"] += amt
        if d["has_act"] and st != "paid":
            b["accepted_unpaid"] += amt
        if st in _planned_statuses:
            b["planned"] += amt
        if d["is_monthly"]:
            b["monthly"] += amt
        if d["is_likely_needed"]:
            b["likely"] += amt

    sm = wb.create_sheet("Сводная", 0)
    SUM_HEADERS = [
        "Категория",
        "Оплаченные+поставленные (треб. оплата), ₽",
        "Оплаченные, ₽",
        "Принятые, не оплаченные, ₽",
        "Планируемые, ₽",
        "Ежемесячные оплаты, ₽",
        "Скорее всего понадобятся, ₽",
        "Всего, ₽",
    ]
    SUM_WIDTHS = [16, 28, 18, 24, 20, 22, 26, 20]
    sm_title = sm.cell(row=1, column=1, value=f"СВОДНАЯ — {sub.name} ({sub.year})")
    sm_title.font = Font(bold=True, size=12, color="1E3A5F")
    sm.merge_cells(f"A1:{chr(64 + len(SUM_HEADERS))}1")
    sm_title.alignment = CENTER_ALIGN
    sm.row_dimensions[1].height = 28
    for ci, (h, w) in enumerate(zip(SUM_HEADERS, SUM_WIDTHS), 1):
        cell = sm.cell(row=2, column=ci, value=h)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER
        sm.column_dimensions[cell.column_letter].width = w
    sm.row_dimensions[2].height = 34

    _order = ["delivered", "paid", "accepted_unpaid", "planned", "monthly", "likely", "total"]

    def _sum_row(sm_row, label, b, fill, font):
        sm.cell(row=sm_row, column=1, value=label)
        for j, k in enumerate(_order, 2):
            c = sm.cell(row=sm_row, column=j, value=round(b[k], 2) if b[k] else "")
            c.number_format = "#,##0.00"
            c.alignment = RIGHT_ALIGN
            c.border = THIN_BORDER
        lc = sm.cell(row=sm_row, column=1)
        lc.fill = fill; lc.font = font; lc.border = THIN_BORDER
        for j in range(2, 9):
            sm.cell(row=sm_row, column=j).fill = fill

    _sum_row(3, "Услуги", summary["услуга"], L2_FILL, L2_FONT)
    _sum_row(4, "Товары", summary["товар"], L3_FILL, L3_FONT)
    total_b = {k: summary["услуга"][k] + summary["товар"][k] for k in _empty_buckets()}
    _sum_row(5, "ИТОГО", total_b, L1_FILL, Font(bold=True, size=10))
    sm.freeze_panes = "A3"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    safe_name = sub.name.replace(" ", "_").replace("/", "_")[:40]
    filename = f"План_график_{safe_name}_{sub.year}.xlsx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": _content_disposition(filename)},
    )


@router.get("/{subsidy_id}/plan-graph/versions/{version_id:int}/export")
async def export_plan_graph_version_excel(
    subsidy_id: int,
    version_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Export a specific plan-graph version snapshot as Excel."""
    if openpyxl is None:
        raise HTTPException(500, "openpyxl не установлен")

    from app.models.plan_graph_version import PlanGraphVersion as _PGV
    from datetime import datetime as _dt

    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")

    org_ids = get_org_filter(current_user)
    if org_ids is not None and sub.org_id not in org_ids:
        raise HTTPException(403, "Нет доступа")

    ver = (await db.execute(
        select(_PGV).where(
            _PGV.id == version_id,
            _PGV.subsidy_id == subsidy_id,
        )
    )).scalar_one_or_none()
    if not ver:
        raise HTTPException(404, "Версия плана закупок не найдена")

    snap = ver.snapshot or {}
    tree = snap.get("tree", [])
    flat_items = snap.get("items", [])

    eff_date = ver.effective_date.isoformat() if ver.effective_date else (
        snap.get("effective_date") or None
    )
    created_at_str = ver.created_at.strftime("%Y-%m-%d %H:%M UTC") if ver.created_at else None

    meta = {
        "subsidy_name": sub.name,
        "subsidy_year": str(sub.year) if sub.year else "",
        "version_number": ver.version_number,
        "effective_date": eff_date,
        "note": ver.note,
        "generated_at": created_at_str,
    }

    wb = _render_plan_graph_workbook(tree, flat_items, meta)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    date_part = (eff_date or (ver.created_at.strftime("%Y-%m-%d") if ver.created_at else "nodate"))
    safe_name = sub.name.replace(" ", "_").replace("/", "_")[:30]
    filename = f"Субсидия_{subsidy_id}_план_v{ver.version_number}_{date_part}.xlsx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": _content_disposition(filename)},
    )


# ── Plan-Graph Compare (Phase 12-05) ─────────────────────────────────────────


def _normalize_name(name: str) -> str:
    """Normalize FeoCategory name for matching: lowercase, strip, collapse spaces."""
    return re.sub(r"\s+", " ", (name or "").lower().strip())


def _match_feo_nodes(tree_a: list, tree_b: list) -> dict:
    """
    Match nodes from two FeoCategory trees using composite key priority:
      1. code (if both non-empty)
      2. (level, parent_path, normalized_name) — parent_path = tuple of normalized ancestor names
      3. (level, normalized_name) — fallback if parent renamed → marked as 'moved'

    Returns:
      {
        "matches": [(node_a, node_b, match_type)],  # match_type: 'code'|'path'|'fallback'
        "only_a": [node_a, ...],
        "only_b": [node_b, ...],
      }
    """
    def _flatten(tree, parent_path=()):
        nodes = []
        for node in tree:
            path = parent_path + (_normalize_name(node.get("name", "")),)
            nodes.append((node, path))
            nodes.extend(_flatten(node.get("children", []), path))
        return nodes

    flat_a = _flatten(tree_a)
    flat_b = _flatten(tree_b)

    matched_b = set()
    matched_a = set()
    matches = []

    # Pass 1: match by code
    code_map_b: dict[str, tuple] = {}
    for nb, pb in flat_b:
        code = (nb.get("code") or "").strip()
        if code:
            code_map_b[code] = (nb, pb)

    for na, pa in flat_a:
        code = (na.get("code") or "").strip()
        if code and code in code_map_b:
            nb, pb = code_map_b[code]
            id_b = id(nb)
            id_a = id(na)
            if id_b not in matched_b and id_a not in matched_a:
                matches.append((na, nb, "code"))
                matched_a.add(id_a)
                matched_b.add(id_b)

    # Pass 2: match by (level, parent_path, normalized_name)
    path_map_b: dict[tuple, tuple] = {}
    for nb, pb in flat_b:
        if id(nb) not in matched_b:
            key = (nb.get("level", 0), pb[:-1], _normalize_name(nb.get("name", "")))
            if key not in path_map_b:
                path_map_b[key] = (nb, pb)

    for na, pa in flat_a:
        if id(na) in matched_a:
            continue
        key = (na.get("level", 0), pa[:-1], _normalize_name(na.get("name", "")))
        if key in path_map_b:
            nb, pb = path_map_b[key]
            if id(nb) not in matched_b:
                matches.append((na, nb, "path"))
                matched_a.add(id(na))
                matched_b.add(id(nb))

    # Pass 3: fallback (level, normalized_name) — marks as 'fallback' → UI shows 'moved'
    name_level_map_b: dict[tuple, tuple] = {}
    for nb, pb in flat_b:
        if id(nb) not in matched_b:
            key = (nb.get("level", 0), _normalize_name(nb.get("name", "")))
            if key not in name_level_map_b:
                name_level_map_b[key] = (nb, pb)

    for na, pa in flat_a:
        if id(na) in matched_a:
            continue
        key = (na.get("level", 0), _normalize_name(na.get("name", "")))
        if key in name_level_map_b:
            nb, pb = name_level_map_b[key]
            if id(nb) not in matched_b:
                matches.append((na, nb, "fallback"))
                matched_a.add(id(na))
                matched_b.add(id(nb))

    only_a = [na for na, pa in flat_a if id(na) not in matched_a]
    only_b = [nb for nb, pb in flat_b if id(nb) not in matched_b]

    return {"matches": matches, "only_a": only_a, "only_b": only_b}


def _build_compare_rows(tree_a: list, tree_b: list) -> list:
    """
    Build flat list of comparison rows for JSON/Excel output.
    Each row: {path, level, name_v1, name_v2, code, budget_v1, budget_v2,
               delta, delta_pct, status}
    status: unchanged | changed | new | removed | moved
    """
    match_result = _match_feo_nodes(tree_a, tree_b)
    match_by_a_id = {id(na): (nb, mtype) for na, nb, mtype in match_result["matches"]}
    matched_b_ids = {id(nb) for _, nb, _ in match_result["matches"]}

    rows = []

    def _walk_a(nodes, parent_path=()):
        for node in nodes:
            path = parent_path + (node.get("name", ""),)
            v1 = float(node.get("budget") or 0)
            mid = id(node)
            if mid in match_by_a_id:
                nb, mtype = match_by_a_id[mid]
                v2 = float(nb.get("budget") or 0)
                delta = v2 - v1
                delta_pct = round(delta / v1 * 100, 2) if v1 != 0 else None
                if mtype == "fallback":
                    status = "moved"
                elif abs(delta) < 0.01:
                    status = "unchanged"
                else:
                    status = "changed"
                rows.append({
                    "path": list(path),
                    "level": node.get("level", 0),
                    "name_v1": node.get("name"),
                    "name_v2": nb.get("name"),
                    "code": node.get("code") or nb.get("code"),
                    "budget_v1": round(v1, 2),
                    "budget_v2": round(v2, 2),
                    "delta": round(delta, 2),
                    "delta_pct": delta_pct,
                    "status": status,
                })
            else:
                rows.append({
                    "path": list(path),
                    "level": node.get("level", 0),
                    "name_v1": node.get("name"),
                    "name_v2": None,
                    "code": node.get("code"),
                    "budget_v1": round(v1, 2),
                    "budget_v2": 0.0,
                    "delta": round(-v1, 2),
                    "delta_pct": -100.0 if v1 != 0 else None,
                    "status": "removed",
                })
            _walk_a(node.get("children", []), path)

    def _walk_only_b(nodes):
        for node in nodes:
            if id(node) not in matched_b_ids:
                v2 = float(node.get("budget") or 0)
                rows.append({
                    "path": [node.get("name", "")],
                    "level": node.get("level", 0),
                    "name_v1": None,
                    "name_v2": node.get("name"),
                    "code": node.get("code"),
                    "budget_v1": 0.0,
                    "budget_v2": round(v2, 2),
                    "delta": round(v2, 2),
                    "delta_pct": None,
                    "status": "new",
                })

    _walk_a(tree_a)

    def _collect_b_nodes(nodes):
        result = []
        for n in nodes:
            result.append(n)
            result.extend(_collect_b_nodes(n.get("children", [])))
        return result

    all_b = _collect_b_nodes(tree_b)
    for nb in all_b:
        if id(nb) not in matched_b_ids:
            v2 = float(nb.get("budget") or 0)
            rows.append({
                "path": [nb.get("name", "")],
                "level": nb.get("level", 0),
                "name_v1": None,
                "name_v2": nb.get("name"),
                "code": nb.get("code"),
                "budget_v1": 0.0,
                "budget_v2": round(v2, 2),
                "delta": round(v2, 2),
                "delta_pct": None,
                "status": "new",
            })

    return rows


@router.get("/{subsidy_id}/plan-graph/versions/compare")
async def compare_plan_graph_versions(
    subsidy_id: int,
    v1: int = Query(..., description="ID первой версии"),
    v2: int = Query(..., description="ID второй версии"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Compare two plan-graph versions and return JSON diff."""
    from app.models.plan_graph_version import PlanGraphVersion as _PGV

    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")

    org_ids = get_org_filter(current_user)
    if org_ids is not None and sub.org_id not in org_ids:
        raise HTTPException(403, "Нет доступа к субсидии")

    ver1 = (await db.execute(
        select(_PGV).where(_PGV.id == v1, _PGV.subsidy_id == subsidy_id)
    )).scalar_one_or_none()
    ver2 = (await db.execute(
        select(_PGV).where(_PGV.id == v2, _PGV.subsidy_id == subsidy_id)
    )).scalar_one_or_none()

    if not ver1:
        raise HTTPException(404, f"Версия {v1} не найдена")
    if not ver2:
        raise HTTPException(404, f"Версия {v2} не найдена")

    snap1 = ver1.snapshot or {}
    snap2 = ver2.snapshot or {}
    tree1 = snap1.get("tree")
    tree2 = snap2.get("tree")

    if not tree1:
        raise HTTPException(
            422,
            f"Старая версия v{ver1.version_number} не содержит дерева ФЭО "
            f"(создана до Phase 12-05). Сравнение недоступно.",
        )
    if not tree2:
        raise HTTPException(
            422,
            f"Старая версия v{ver2.version_number} не содержит дерева ФЭО "
            f"(создана до Phase 12-05). Сравнение недоступно.",
        )

    rows = _build_compare_rows(tree1, tree2)

    def _ver_meta(ver, snap):
        return {
            "id": ver.id,
            "version_number": ver.version_number,
            "effective_date": ver.effective_date.isoformat() if ver.effective_date else snap.get("effective_date"),
            "note": ver.note,
            "created_at": ver.created_at.isoformat() if ver.created_at else None,
            "total_planned": snap.get("total_planned", 0),
        }

    return {
        "v1_meta": _ver_meta(ver1, snap1),
        "v2_meta": _ver_meta(ver2, snap2),
        "rows": rows,
    }


@router.get("/{subsidy_id}/plan-graph/versions/compare.xlsx")
async def compare_plan_graph_versions_excel(
    subsidy_id: int,
    v1: int = Query(..., description="ID первой версии"),
    v2: int = Query(..., description="ID второй версии"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Compare two plan-graph versions and return Excel diff."""
    if openpyxl is None:
        raise HTTPException(500, "openpyxl не установлен")

    from app.models.plan_graph_version import PlanGraphVersion as _PGV

    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")

    org_ids = get_org_filter(current_user)
    if org_ids is not None and sub.org_id not in org_ids:
        raise HTTPException(403, "Нет доступа к субсидии")

    ver1 = (await db.execute(
        select(_PGV).where(_PGV.id == v1, _PGV.subsidy_id == subsidy_id)
    )).scalar_one_or_none()
    ver2 = (await db.execute(
        select(_PGV).where(_PGV.id == v2, _PGV.subsidy_id == subsidy_id)
    )).scalar_one_or_none()

    if not ver1:
        raise HTTPException(404, f"Версия {v1} не найдена")
    if not ver2:
        raise HTTPException(404, f"Версия {v2} не найдена")

    snap1 = ver1.snapshot or {}
    snap2 = ver2.snapshot or {}
    tree1 = snap1.get("tree")
    tree2 = snap2.get("tree")

    if not tree1:
        raise HTTPException(
            422,
            f"Старая версия v{ver1.version_number} не содержит дерева ФЭО "
            f"(создана до Phase 12-05). Сравнение недоступно.",
        )
    if not tree2:
        raise HTTPException(
            422,
            f"Старая версия v{ver2.version_number} не содержит дерева ФЭО "
            f"(создана до Phase 12-05). Сравнение недоступно.",
        )

    rows = _build_compare_rows(tree1, tree2)

    # ── Build Excel ──────────────────────────────────────────────────────────
    HEADER_FILL   = PatternFill("solid", fgColor="1E3A5F")
    HEADER_FONT   = Font(color="FFFFFF", bold=True, size=9)
    META_FONT     = Font(size=9, italic=True, color="374151")
    ITEM_FONT     = Font(size=9)
    CENTER_ALIGN  = Alignment(horizontal="center", vertical="center", wrap_text=True)
    LEFT_ALIGN    = Alignment(horizontal="left", vertical="center", wrap_text=True)
    RIGHT_ALIGN   = Alignment(horizontal="right", vertical="center")
    THIN_BORDER   = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    STATUS_FILLS = {
        "changed":   PatternFill("solid", fgColor="FFF3CD"),
        "new":       PatternFill("solid", fgColor="D4EDDA"),
        "removed":   PatternFill("solid", fgColor="F8D7DA"),
        "moved":     PatternFill("solid", fgColor="D1ECF1"),
        "unchanged": PatternFill(),
    }
    STATUS_LABELS = {
        "changed": "Изменено",
        "new": "Новое",
        "removed": "Удалено",
        "moved": "Перемещено",
        "unchanged": "Без изменений",
    }

    HEADERS = [
        "№", "Уровень", "Наименование", "Код",
        f"План v{ver1.version_number} (₽)", f"План v{ver2.version_number} (₽)",
        "Дельта (₽)", "Дельта (%)", "Статус",
    ]
    COL_WIDTHS = [5, 8, 50, 15, 18, 18, 18, 12, 15]
    n_cols = len(HEADERS)
    last_col = chr(64 + n_cols)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Сравнение версий"

    # Title
    title_v1_date = ver1.effective_date.isoformat() if ver1.effective_date else (snap1.get("effective_date") or ver1.created_at.strftime("%Y-%m-%d") if ver1.created_at else "—")
    title_v2_date = ver2.effective_date.isoformat() if ver2.effective_date else (snap2.get("effective_date") or ver2.created_at.strftime("%Y-%m-%d") if ver2.created_at else "—")
    ws.append([f"СРАВНЕНИЕ ВЕРСИЙ — {sub.name}"] + [""] * (n_cols - 1))
    title_cell = ws.cell(row=1, column=1)
    title_cell.font = Font(bold=True, size=12, color="1E3A5F")
    ws.merge_cells(f"A1:{last_col}1")
    title_cell.alignment = CENTER_ALIGN
    ws.row_dimensions[1].height = 28

    # Meta rows
    meta_texts = [
        f"v{ver1.version_number}: {title_v1_date}" + (f" — {ver1.note}" if ver1.note else ""),
        f"v{ver2.version_number}: {title_v2_date}" + (f" — {ver2.note}" if ver2.note else ""),
        f"Итого v{ver1.version_number}: {snap1.get('total_planned', 0):,.2f} ₽  |  "
        f"Итого v{ver2.version_number}: {snap2.get('total_planned', 0):,.2f} ₽  |  "
        f"Дельта: {(snap2.get('total_planned', 0) - snap1.get('total_planned', 0)):+,.2f} ₽",
    ]
    for i, mtext in enumerate(meta_texts):
        r = 2 + i
        ws.append([mtext] + [""] * (n_cols - 1))
        cell = ws.cell(row=r, column=1)
        cell.font = META_FONT
        cell.alignment = LEFT_ALIGN
        ws.merge_cells(f"A{r}:{last_col}{r}")
        ws.row_dimensions[r].height = 16

    # Column headers
    header_row = 2 + len(meta_texts)
    ws.append(HEADERS)
    for col_idx, (h, w) in enumerate(zip(HEADERS, COL_WIDTHS), 1):
        cell = ws.cell(row=header_row, column=col_idx)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER
        ws.column_dimensions[cell.column_letter].width = w
    ws.row_dimensions[header_row].height = 32
    ws.freeze_panes = f"A{header_row + 1}"

    # Data rows
    for seq_i, row in enumerate(rows, 1):
        level = row["level"]
        indent = "    " * (level - 1) if level > 0 else ""
        name = row["name_v2"] or row["name_v1"] or ""
        status = row["status"]
        fill = STATUS_FILLS.get(status, PatternFill())
        delta_pct_str = f"{row['delta_pct']:+.1f}%" if row["delta_pct"] is not None else "—"

        data_row_num = header_row + seq_i
        values = [
            seq_i, level, indent + name, row.get("code") or "",
            row["budget_v1"], row["budget_v2"],
            row["delta"], delta_pct_str,
            STATUS_LABELS.get(status, status),
        ]
        ws.append(values)
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=data_row_num, column=col_idx)
            cell.fill = fill
            cell.font = ITEM_FONT
            cell.border = THIN_BORDER
            if col_idx == 3:
                cell.alignment = LEFT_ALIGN
            elif col_idx >= 5:
                cell.alignment = RIGHT_ALIGN
            else:
                cell.alignment = CENTER_ALIGN
        ws.row_dimensions[data_row_num].height = 18

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    filename = f"Субсидия_{subsidy_id}_сравнение_v{ver1.version_number}_и_v{ver2.version_number}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": _content_disposition(filename)},
    )


# ── Multi-edition ФЭО compare (Phase 12-06) ──────────────────────────────────


async def _build_current_feo_tree(subsidy_id: int, db: AsyncSession) -> list:
    """Строит живое дерево FeoCategory для субсидии (аналог snapshot["tree"])."""
    cats = (await db.execute(
        select(FeoCategory).where(FeoCategory.subsidy_id == subsidy_id).order_by(FeoCategory.id)
    )).scalars().all()

    def _recurse(parent_id):
        nodes = []
        for c in cats:
            if c.parent_id == parent_id:
                nodes.append({
                    "id": c.id,
                    "name": c.name,
                    "level": c.level,
                    "code": c.code,
                    "budget": float(c.budget) if c.budget is not None else None,
                    "children": _recurse(c.id),
                })
        return nodes

    return _recurse(None)


def _build_multi_compare_rows(versions: list) -> list:
    """
    versions = [{"tree": [...]}, ...]  — уже в хронологическом порядке.
    Возвращает список строк: {"level", "name", "code", "budgets": [v0_val, v1_val, ...]}.
    Canonical key:
      - ("code", code)          если у узла есть code
      - ("path", level, path)   иначе (path = tuple нормализованных предков + имя узла)
    Ключи упорядочены по первому появлению (первая редакция задаёт порядок).
    """
    order: list = []            # список ключей в порядке появления
    meta: dict = {}             # key -> {"level", "name", "code"}
    budgets: dict = {}          # key -> {version_index: float}
    facts: dict = {}            # key -> {version_index: float} — фактически освоено (used_amount)

    def _dfs(node, path, ver_idx):
        norm_name = _normalize_name(node.get("name", ""))
        full_path = path + (norm_name,)

        code = (node.get("code") or "").strip()
        if code:
            key = ("code", code)
        else:
            key = ("path", node.get("level", 0), full_path)

        if key not in meta:
            order.append(key)
            meta[key] = {
                "level": node.get("level", 0),
                "name": node.get("name", ""),
                "code": code,
            }
        else:
            # Обновляем имя/код из последней редакции, если они непустые
            if node.get("name", ""):
                meta[key]["name"] = node["name"]
            if code:
                meta[key]["code"] = code

        if key not in budgets:
            budgets[key] = {}
        budgets[key][ver_idx] = float(node.get("budget") or 0)

        if key not in facts:
            facts[key] = {}
        facts[key][ver_idx] = float(node.get("used_amount") or 0)

        for child in node.get("children", []):
            _dfs(child, full_path, ver_idx)

    for ver_idx, ver in enumerate(versions):
        for root_node in ver.get("tree", []):
            _dfs(root_node, (), ver_idx)

    n_vers = len(versions)
    rows = []
    for key in order:
        m = meta[key]
        rows.append({
            "level": m["level"],
            "name": m["name"],
            "code": m["code"],
            "budgets": [budgets[key].get(i, 0.0) for i in range(n_vers)],
            "facts": [facts[key].get(i, 0.0) for i in range(n_vers)],
        })
    return rows


@router.get("/{subsidy_id}/plan-graph/versions/export-multi.xlsx")
async def export_plan_graph_versions_multi_excel(
    subsidy_id: int,
    ids: str = Query("", description="ID редакций через запятую"),
    include_current: bool = Query(False, description="Добавить колонку текущей живой ФЭО"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Экспорт N редакций ФЭО бок-о-бок в один Excel-файл с дельтами."""
    if openpyxl is None:
        raise HTTPException(500, "openpyxl не установлен")

    from app.models.plan_graph_version import PlanGraphVersion as _PGV
    from datetime import datetime as _dt

    # ── Авторизация ──────────────────────────────────────────────────────────
    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")

    org_ids = get_org_filter(current_user)
    if org_ids is not None and sub.org_id not in org_ids:
        raise HTTPException(403, "Нет доступа к субсидии")

    # ── Парсим id редакций ───────────────────────────────────────────────────
    try:
        id_list: list[int] = [int(x) for x in ids.split(",") if x.strip()]
    except ValueError:
        raise HTTPException(422, "Некорректный список id редакций")
    if not id_list and not include_current:
        raise HTTPException(422, "Выберите хотя бы одну редакцию")

    # ── Загружаем версии из БД ───────────────────────────────────────────────
    skipped_legacy: list[str] = []
    versions: list[dict] = []

    if id_list:
        ver_rows = (await db.execute(
            select(_PGV).where(_PGV.id.in_(id_list), _PGV.subsidy_id == subsidy_id)
        )).scalars().all()

        def _ver_date(v) -> date:
            if v.effective_date:
                return v.effective_date
            if v.created_at:
                return v.created_at.date()
            return date.today()

        # Хронологическая сортировка: сначала по дате, потом по номеру версии
        ver_rows_sorted = sorted(ver_rows, key=lambda v: (_ver_date(v), v.version_number))

        for ver in ver_rows_sorted:
            snap = ver.snapshot or {}
            tree = snap.get("tree")
            if not tree:
                skipped_legacy.append(f"v{ver.version_number}")
                continue
            d = _ver_date(ver)
            versions.append({
                "ver": ver,
                "label": f"v{ver.version_number}",
                "date_str": d.isoformat(),
                "date_display": d.strftime("%d.%m.%Y"),
                "note": ver.note,
                "tree": tree,
                "total_planned": snap.get("total_planned", 0),
            })

    # ── Текущая живая ФЭО ────────────────────────────────────────────────────
    if include_current:
        live_tree = await _build_current_feo_tree(subsidy_id, db)

        def _sum_tree(nodes):
            total = 0.0
            for n in nodes:
                total += float(n.get("budget") or 0)
                total += _sum_tree(n.get("children", []))
            return total

        today = date.today()
        versions.append({
            "ver": None,
            "label": "Текущая",
            "date_str": today.isoformat(),
            "date_display": today.strftime("%d.%m.%Y"),
            "note": None,
            "tree": live_tree,
            "total_planned": _sum_tree(live_tree),
        })

    if not versions:
        raise HTTPException(422, "Нет редакций с деревом ФЭО для выгрузки")

    rows = _build_multi_compare_rows(versions)

    # ── Стили (копируем из compare.xlsx) ────────────────────────────────────
    HEADER_FILL  = PatternFill("solid", fgColor="1E3A5F")
    HEADER_FONT  = Font(color="FFFFFF", bold=True, size=9)
    META_FONT    = Font(size=9, italic=True, color="374151")
    ITEM_FONT    = Font(size=9)
    CENTER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
    LEFT_ALIGN   = Alignment(horizontal="left", vertical="center", wrap_text=True)
    RIGHT_ALIGN  = Alignment(horizontal="right", vertical="center")
    THIN_BORDER  = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    n_vers = len(versions)
    has_deltas = n_vers >= 2

    # ── Заголовки колонок ────────────────────────────────────────────────────
    fixed_headers = ["№", "Уровень", "Наименование", "Код"]
    plan_headers = [
        f"План ₽\n{v['date_display']}" for v in versions
    ]
    fact_headers = [
        f"Факт ₽\n{v['date_display']}" for v in versions
    ]
    delta_adj_headers = []
    if has_deltas:
        for i in range(1, n_vers):
            delta_adj_headers.append(
                f"Δ {versions[i]['date_display']}−{versions[i - 1]['date_display']}"
            )
    total_delta_headers = []
    if has_deltas:
        total_delta_headers = [
            f"Δ итог\n{versions[-1]['date_display']}−{versions[0]['date_display']}"
        ]

    all_headers = fixed_headers + plan_headers + fact_headers + delta_adj_headers + total_delta_headers
    n_cols = len(all_headers)
    last_col = chr(64 + n_cols) if n_cols <= 26 else (
        chr(64 + (n_cols - 1) // 26) + chr(64 + (n_cols - 1) % 26 + 1)
    )

    # Ширины колонок: фиксированные 5,8,50,15; план 18; факт 18; дельта 16
    col_widths = [5, 8, 50, 15] + [18] * n_vers + [18] * n_vers + [16] * (len(delta_adj_headers) + len(total_delta_headers))

    # ── Workbook ─────────────────────────────────────────────────────────────
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Редакции ФЭО"

    # Строка-заголовок (merged)
    ws.append([f"РЕДАКЦИИ ФЭО — {sub.name}"] + [""] * (n_cols - 1))
    title_cell = ws.cell(row=1, column=1)
    title_cell.font = Font(bold=True, size=12, color="1E3A5F")
    ws.merge_cells(f"A1:{last_col}1")
    title_cell.alignment = CENTER_ALIGN
    ws.row_dimensions[1].height = 28

    # Мета-строки: по одной на каждую редакцию
    meta_lines = []
    for v in versions:
        line = f"{v['label']} — {v['date_display']}"
        if v["note"]:
            line += f" — {v['note']}"
        meta_lines.append(line)

    # Строка с итогами по редакциям
    totals_parts = [
        f"{v['label']}: {float(v['total_planned']):,.2f} ₽" for v in versions
    ]
    meta_lines.append("Итого: " + "  |  ".join(totals_parts))

    # Предупреждение о пропущенных legacy-версиях
    if skipped_legacy:
        meta_lines.append(
            f"⚠ Пропущены (нет дерева ФЭО): {', '.join(skipped_legacy)}"
        )

    for i, mtext in enumerate(meta_lines):
        r = 2 + i
        ws.append([mtext] + [""] * (n_cols - 1))
        cell = ws.cell(row=r, column=1)
        cell.font = META_FONT
        cell.alignment = LEFT_ALIGN
        ws.merge_cells(f"A{r}:{last_col}{r}")
        ws.row_dimensions[r].height = 16

    # Строка заголовков колонок
    header_row = 2 + len(meta_lines)
    ws.append(all_headers)
    for col_idx, (h, w) in enumerate(zip(all_headers, col_widths), 1):
        cell = ws.cell(row=header_row, column=col_idx)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER
        ws.column_dimensions[cell.column_letter].width = w
    ws.row_dimensions[header_row].height = 32
    ws.freeze_panes = f"A{header_row + 1}"

    # ── Строки данных ────────────────────────────────────────────────────────
    GREEN_FONT = Font(size=9, color="1B7F3B")
    RED_FONT   = Font(size=9, color="B00020")

    for seq_i, row in enumerate(rows, 1):
        level = row["level"]
        indent = ("    " * (level - 1)) if level > 0 else ""
        name_str = indent + (row["name"] or "")
        code_str = row["code"] or ""
        bgets = row["budgets"]  # list[float] длиной n_vers
        fcts = row.get("facts") or [0.0] * n_vers  # фактически освоено по каждой редакции

        # Вычисляем дельта-значения
        adj_deltas = [bgets[i] - bgets[i - 1] for i in range(1, n_vers)] if has_deltas else []
        total_delta = [bgets[-1] - bgets[0]] if has_deltas else []

        row_values = [seq_i, level, name_str, code_str] + bgets + fcts + adj_deltas + total_delta
        data_row_num = header_row + seq_i
        ws.append(row_values)

        for col_idx, val in enumerate(row_values, 1):
            cell = ws.cell(row=data_row_num, column=col_idx)
            cell.border = THIN_BORDER

            if col_idx == 3:
                # Наименование — левое выравнивание
                cell.font = ITEM_FONT
                cell.alignment = LEFT_ALIGN
            elif col_idx <= 4:
                # №, Уровень, Код
                cell.font = ITEM_FONT
                cell.alignment = CENTER_ALIGN
            elif col_idx <= 4 + n_vers:
                # Плановые колонки
                cell.font = ITEM_FONT
                cell.alignment = RIGHT_ALIGN
            elif col_idx <= 4 + 2 * n_vers:
                # Фактические колонки (used_amount)
                cell.font = GREEN_FONT if isinstance(val, float) and val > 0.005 else ITEM_FONT
                cell.alignment = RIGHT_ALIGN
            else:
                # Дельта-колонки: цветной шрифт
                if isinstance(val, float) and val > 0.005:
                    cell.font = GREEN_FONT
                elif isinstance(val, float) and val < -0.005:
                    cell.font = RED_FONT
                else:
                    cell.font = ITEM_FONT
                cell.alignment = RIGHT_ALIGN

        ws.row_dimensions[data_row_num].height = 18

    # ── Сохраняем и отдаём ──────────────────────────────────────────────────
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    filename = f"Субсидия_{subsidy_id}_ФЭО_редакции.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": _content_disposition(filename)},
    )


@router.post("/{subsidy_id}/plan-graph/template")
async def upload_plan_graph_template(
    subsidy_id: int,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(*ADMIN_ROLES)),
):
    """Upload a .docx Word template for this subsidy's plan-graph export."""
    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Только .docx файлы поддерживаются")

    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    dest = os.path.join(TEMPLATE_DIR, f"subsidy_{subsidy_id}.docx")
    content = await file.read()
    with open(dest, "wb") as f:
        f.write(content)

    return {"ok": True, "template_path": dest, "message": "Шаблон загружен"}


@router.get("/{subsidy_id}/plan-graph/export-docx")
async def export_plan_graph_docx(
    subsidy_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Fill the uploaded .docx template via docxtpl and return the filled document."""
    if DocxTemplate is None:
        raise HTTPException(500, "docxtpl не установлен")

    sub = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Субсидия не найдена")

    template_path = os.path.join(TEMPLATE_DIR, f"subsidy_{subsidy_id}.docx")
    if not os.path.exists(template_path):
        raise HTTPException(404, "Шаблон не загружен. Загрузите через POST /plan-graph/template")

    org_ids = get_org_filter(current_user)
    if org_ids is not None and sub.org_id not in org_ids:
        raise HTTPException(403, "Нет доступа")

    from app.models.plan_graph_version import PlanGraphVersion as _PGV
    from datetime import datetime as _dt

    latest_ver = (await db.execute(
        select(_PGV)
        .where(_PGV.subsidy_id == subsidy_id)
        .order_by(_PGV.version_number.desc())
        .limit(1)
    )).scalar_one_or_none()

    if latest_ver and latest_ver.snapshot:
        snap = latest_ver.snapshot
        items_ctx = snap.get("items", [])
        total_planned = snap.get("total_planned", 0)
        total_used = snap.get("total_used", 0)
    else:
        items_ctx = []
        total_planned = 0.0
        total_used = 0.0

    context = {
        "subsidy_name": sub.name,
        "subsidy_year": sub.year,
        "items": items_ctx,
        "total_planned": f"{total_planned:,.2f}",
        "total_used": f"{total_used:,.2f}",
        "total_residual": f"{total_planned - total_used:,.2f}",
        "export_date": _dt.now().strftime("%d.%m.%Y"),
    }

    doc = DocxTemplate(template_path)
    doc.render(context)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = sub.name.replace(" ", "_").replace("/", "_")[:40]
    filename = f"План_график_{safe_name}_{sub.year}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _content_disposition(filename)},
    )


# ---------------------------------------------------------------------------
# GET /api/subsidies/{subsidy_id}/payment-summary
# Сверка платежей закупок vs банковской выписки по субсидии.
# ---------------------------------------------------------------------------

@router.get("/{subsidy_id}/payment-summary")
async def payment_summary(
    subsidy_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Сводка платежей субсидии: закупки vs реестр (BankPayment).

    Скоупинг BankPayment к субсидии:
    - Используем BankPayment.matched_subsidy_id == subsidy_id (прямое поле match-состояния).
    - Дополнительно берём BankPayment с payment_number, совпадающим с payment_doc_number
      закупок этой субсидии — на случай, если matched_subsidy_id ещё не проставлен.
    Объединяем оба набора (дедуп по id) для максимального покрытия.
    """
    from decimal import Decimal
    from app.models.bank_statement import BankPayment
    from app.services.payment_reconciliation import build_reconciliation, _normalize_num
    from app.auth.visibility import get_visible_subsidy_ids, build_visibility_clause

    # --- security: visibility check (mirror of budget-check) ---
    visible = await get_visible_subsidy_ids(current_user, db)
    if visible is not None and subsidy_id not in visible:
        clause = await build_visibility_clause(current_user, db, "purchase")
        allowed = clause is None
        if clause is not None:
            cnt = (await db.execute(
                select(func.count()).select_from(Purchase).where(
                    Purchase.subsidy_id == subsidy_id, clause
                )
            )).scalar() or 0
            allowed = cnt > 0
        if not allowed:
            raise HTTPException(status_code=404, detail="Subsidy not found")

    subsidy = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not subsidy:
        raise HTTPException(status_code=404, detail="Subsidy not found")

    # --- 1. Purchase payments for this subsidy ---
    purchase_rows = (await db.execute(
        select(Purchase).where(
            Purchase.subsidy_id == subsidy_id,
            Purchase.payment_doc_number.isnot(None),
        )
    )).scalars().all()

    purchases_total = float(sum(
        Decimal(str(p.payment_amount or 0)) for p in purchase_rows
    ))

    # Normalised ПП numbers from purchases (for bank lookup)
    purchase_pp_set: set[str] = set()
    for p in purchase_rows:
        norm = _normalize_num(p.payment_doc_number)
        if norm:
            purchase_pp_set.add(norm)

    # --- 2. Bank payments scoped to this subsidy ---
    # Primary: matched_subsidy_id directly links BankPayment to subsidy
    bp_by_matched_q = select(BankPayment).where(
        BankPayment.matched_subsidy_id == subsidy_id,
        BankPayment.payment_number.isnot(None),
    )
    bp_matched = (await db.execute(bp_by_matched_q)).scalars().all()
    bp_ids_seen: set[int] = {bp.id for bp in bp_matched}

    # Secondary: BankPayment whose payment_number normalises to one of the purchase ПП numbers
    bp_extra: list = []
    if purchase_pp_set:
        bp_all_q = select(BankPayment).where(
            BankPayment.payment_number.isnot(None),
            BankPayment.matched_subsidy_id.is_(None),  # skip already counted
        )
        for bp in (await db.execute(bp_all_q)).scalars().all():
            if _normalize_num(bp.payment_number) in purchase_pp_set and bp.id not in bp_ids_seen:
                bp_extra.append(bp)
                bp_ids_seen.add(bp.id)

    all_bank_payments = list(bp_matched) + bp_extra
    bank_total = float(sum(Decimal(str(bp.amount or 0)) for bp in all_bank_payments))

    difference = round(purchases_total - bank_total, 2)

    # --- 3. Reconciliation via existing service ---
    # build_reconciliation works over ALL bank payments and payments. We call it for
    # the whole dataset and then filter to rows relevant to this subsidy's ПП numbers.
    # The service is NOT subsidy-aware, so we post-filter by normalised ПП set.
    all_recon = await build_reconciliation(db)

    # ПП numbers appearing in this subsidy's purchases (already built above)
    relevant_rows = [
        row for row in all_recon
        if _normalize_num(row.get("payment_number", "")) in purchase_pp_set
    ] if purchase_pp_set else []

    # Split into buckets
    matched_count = sum(1 for r in relevant_rows if r["status"] == "match")
    discrepancy_rows = [r for r in relevant_rows if r["status"] != "match"]

    discrepancies = []
    for row in discrepancy_rows:
        item: dict = {
            "payment_number": row["payment_number"],
            "status": row["status"],
            "purchase_amount": row["purchases_amount"],
            "bank_amount": row["registry_amount"],
            "amount_diff": row["amount_diff"],
        }
        if row["status"] == "amount_mismatch":
            item["reason"] = "Суммы в закупках и реестре расходятся"
        elif row["status"] == "registry_only":
            item["reason"] = "Есть в банковском реестре, отсутствует в закупках"
        elif row["status"] == "purchases_only":
            item["reason"] = "Привязан к закупке, отсутствует в банковском реестре"
        else:
            item["reason"] = row["status"]
        discrepancies.append(item)

    return {
        "subsidy_id": subsidy_id,
        "purchases_total": purchases_total,
        "bank_total": bank_total,
        "difference": difference,
        "matched_count": matched_count,
        "discrepancies": discrepancies,
    }


# ---------------------------------------------------------------------------
# Финансовый план субсидии — помесячный cash-flow
# ---------------------------------------------------------------------------

@router.get("/{subsidy_id}/financial-plan")
async def financial_plan(
    subsidy_id: int,
    year: Optional[int] = Query(None, description="Год (по умолчанию текущий)"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Помесячный финансовый план субсидии.

    Слои данных:
      plan        — плановые выплаты из ФЭО (expand_planned_item)
      obligations — принятые обязательства (contract_price) по месяцу obligation_date
      paid        — оплачено (payment_amount) по месяцу payment_doc_date, статус paid
      debt_cumul  — накопленный долг = Σ(obligations) − Σ(paid) нарастающим итогом
      var_plan_obl, var_plan_paid — отклонения план−обяз, план−факт за месяц

    Также возвращает разбивку по directions (level-1 ФЭО).
    """
    from datetime import date as _date
    from decimal import Decimal as _Dec
    from app.auth.visibility import get_visible_subsidy_ids
    from app.services.plan_cashflow import (
        cashflow_for_subsidy,
        cashflow_with_directions,
    )
    from app.routers.dashboard import obligation_date as _obligation_date

    target_year = year or _date.today().year

    # ----- авторизация: видимость субсидии -----
    visible = await get_visible_subsidy_ids(current_user, db)
    if visible is not None and subsidy_id not in visible:
        raise HTTPException(status_code=404, detail="Subsidy not found")

    subsidy = (await db.execute(select(Subsidy).where(Subsidy.id == subsidy_id))).scalar_one_or_none()
    if not subsidy:
        raise HTTPException(status_code=404, detail="Subsidy not found")

    # ----- Плановые выплаты (plan) -----
    plan_monthly = await cashflow_for_subsidy(db, subsidy_id)
    # Разбивка по directions
    dir_monthly = await cashflow_with_directions(db, subsidy_id)

    # Загрузить names level-1 категорий
    dir_cat_ids = [cid for cid in dir_monthly if cid != 0]
    dir_names: dict[int, str] = {}
    if dir_cat_ids:
        dir_cats = (await db.execute(
            select(FeoCategory.id, FeoCategory.name).where(FeoCategory.id.in_(dir_cat_ids))
        )).all()
        dir_names = {r.id: r.name for r in dir_cats}

    # ----- Закупки (obligations + paid) -----
    OBL_STATUSES = ("contracted", "delivered", "paid")
    purchases_q = select(Purchase).where(
        Purchase.subsidy_id == subsidy_id,
        Purchase.status.in_(OBL_STATUSES),
    )
    purchases = (await db.execute(purchases_q)).scalars().all()

    # Разбиваем по месяцам за target_year
    obligations_monthly: dict[int, _Dec] = {}   # month -> Decimal
    paid_monthly: dict[int, _Dec] = {}

    for p in purchases:
        # obligations — по obligation_date, только contracted/delivered/paid
        obl_d = _obligation_date(p)
        if obl_d and obl_d.year == target_year:
            amt = _Dec(str(p.contract_price or 0))
            obligations_monthly[obl_d.month] = (
                obligations_monthly.get(obl_d.month, _Dec(0)) + amt
            )
        # paid — по payment_doc_date, только status=paid
        if p.status == "paid" and p.payment_doc_date:
            pd = p.payment_doc_date
            if hasattr(pd, 'date'):
                pd = pd.date()
            if pd.year == target_year:
                amt_p = _Dec(str(p.payment_amount or 0))
                paid_monthly[pd.month] = paid_monthly.get(pd.month, _Dec(0)) + amt_p

    # ----- Сборка ответа: 12 месяцев -----
    months_out = []
    cum_obl = _Dec(0)
    cum_paid = _Dec(0)
    annual_plan = _Dec(0)
    annual_obl = _Dec(0)
    annual_paid = _Dec(0)

    quarter_plan: dict[int, _Dec] = {}
    quarter_obl: dict[int, _Dec]  = {}
    quarter_paid: dict[int, _Dec] = {}

    for m in range(1, 13):
        q = (m - 1) // 3 + 1
        plan_amt  = plan_monthly.get((target_year, m), _Dec(0))
        obl_amt   = obligations_monthly.get(m, _Dec(0))
        paid_amt  = paid_monthly.get(m, _Dec(0))

        cum_obl  += obl_amt
        cum_paid += paid_amt
        debt_cum  = cum_obl - cum_paid

        annual_plan += plan_amt
        annual_obl  += obl_amt
        annual_paid += paid_amt

        quarter_plan[q]  = quarter_plan.get(q, _Dec(0)) + plan_amt
        quarter_obl[q]   = quarter_obl.get(q, _Dec(0)) + obl_amt
        quarter_paid[q]  = quarter_paid.get(q, _Dec(0)) + paid_amt

        # Разбивка по directions для этого месяца
        dir_breakdown = {}
        for dir_id, dir_data in dir_monthly.items():
            dir_val = dir_data.get((target_year, m), _Dec(0))
            if dir_val:
                dir_breakdown[str(dir_id)] = {
                    "name": dir_names.get(dir_id, ""),
                    "plan": float(dir_val),
                }

        months_out.append({
            "month": m,
            "plan": float(plan_amt),
            "obligations": float(obl_amt),
            "paid": float(paid_amt),
            "debt_cumulative": float(debt_cum),
            "variance_plan_obligation": float(plan_amt - obl_amt),
            "variance_plan_paid": float(plan_amt - paid_amt),
            "directions": dir_breakdown,
        })

    # Квартальные подытоги
    quarters_out = []
    for q in range(1, 5):
        q_plan = quarter_plan.get(q, _Dec(0))
        q_obl  = quarter_obl.get(q, _Dec(0))
        q_paid = quarter_paid.get(q, _Dec(0))
        quarters_out.append({
            "quarter": q,
            "plan": float(q_plan),
            "obligations": float(q_obl),
            "paid": float(q_paid),
            "variance_plan_obligation": float(q_plan - q_obl),
            "variance_plan_paid": float(q_plan - q_paid),
        })

    return {
        "subsidy_id": subsidy_id,
        "year": target_year,
        "months": months_out,
        "quarters": quarters_out,
        "annual": {
            "plan": float(annual_plan),
            "obligations": float(annual_obl),
            "paid": float(annual_paid),
            "variance_plan_obligation": float(annual_plan - annual_obl),
            "variance_plan_paid": float(annual_plan - annual_paid),
        },
    }
